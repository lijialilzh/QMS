#!/usr/bin/env python
# encoding: utf-8

# 代码审查记录服务层（开发文件）。整份文档以 content(JSON) 存储；导出复用 docx_util.fonted_txt 生成 Word。
# 自动获取：检查日期/签字日期←时间线(代码审查)，被审核人←TPM，审核人←研发负责人，审核签字←研发负责人签名章。

import base64
import copy
import logging
import re
from io import BytesIO
from typing import List

from sqlalchemy import delete, func, select
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE

from ..model.product import Product
from ..model.crr_doc import CrrDoc
from ..model.project_member import ProjectMember
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_crr_doc import CrrDocForm
from ..obj.vobj_crr_doc import CrrDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from .serv_utils import new_version, sync_file_no_version
from .serv_utils import docx_util
from . import serv_review_util

logger = logging.getLogger(__name__)

# 检查表分类行（整行只有首格有内容，导出/编辑时作为分组标题横向合并）。
_CATEGORIES = ("结构", "文档", "变量", "算法操作", "循环和分支")

# 代码地址默认模板（取自《代码审查记录》，可手动编辑；用空格分隔，避免制表符导致导出内容不左对齐）。
_DEFAULT_CODE_URL = (
    "DP http://gitlab.quality.info:8081/infercare-recist/repacs/tree/1.0.0.0\n"
    "RePACS http://gitlab.quality.info:8081/infercare-recist/repacs/tree/1.0.0.0\n"
    "DLServer http://gitlab.quality.info:8081/infercare-recist/dlserver/tree/1.0.0.0\n"
    "NeoViewer http://gitlab.quality.info:8081/infercare-recist/neoviewer/tree/1.0.0.0"
)

# 标准模板默认内容（取自《代码审查记录》模板），新增文档时预填、可改。
DEFAULT_CRR_CONTENT = {
    "code_url": _DEFAULT_CODE_URL,
    "check_date": "",
    "auditee": "",
    "auditor": "",
    "basis": "《代码管理制度》",
    "checklist": [
        ["编号", "问题", "是", "否", "不适用", "备注"],
        ["结构", "", "", "", "", ""],
        ["1", "代码是否符合相关的编码标准?", "", "", "", ""],
        ["2", "代码结构是否适当，风格和格式是否保持一致?", "", "", "", ""],
        ["3", "代码中是否有没有被调用的或无用的程序，或没有被执行的代码?", "", "", "", ""],
        ["4", "是否有过于复杂的模块需要重新构造或拆分成多个程序?", "", "", "", ""],
        ["文档", "", "", "", "", ""],
        ["1", "代码是否已被用易于维护的注释方式清晰充分的文档化?", "", "", "", ""],
        ["2", "注释是否与代码协调一致?", "", "", "", ""],
        ["变量", "", "", "", "", ""],
        ["1", "所有变量的命名是否清晰，一致并且有意义?", "", "", "", ""],
        ["2", "是否有冗余或无用的变量?", "", "", "", ""],
        ["算法操作", "", "", "", "", ""],
        ["1", "被除数是否做了零值测试?", "", "", "", ""],
        ["循环和分支", "", "", "", "", ""],
        ["1", "所有的循环，分支和逻辑构造是否完整，正确并且嵌套适当?", "", "", "", ""],
        ["2", "每种状况是否都有缺省值?", "", "", "", ""],
    ],
    "conclusion": "",
    "sign_img": "",
    "sign_date": "",
}


class Server(object):

    def __normalize_content(self, content):
        base = copy.deepcopy(DEFAULT_CRR_CONTENT)
        if not isinstance(content, dict):
            return base
        for key in base.keys():
            if key == "checklist":
                val = content.get("checklist")
                if isinstance(val, list) and val:
                    base["checklist"] = [list(r) if isinstance(r, list) else [r] for r in val]
            elif content.get(key) is not None:
                base[key] = content.get(key)
        # 代码地址：制表符转空格，保证左对齐不缩进
        base["code_url"] = str(base.get("code_url") or "").replace("\t", "  ")
        return base

    def __autofill(self, content, prod_id):
        """按产品参与人员/时间线自动获取（仅填空，不覆盖已填）：
        检查日期/签字日期←时间线(代码审查)；被审核人←TPM；审核人←研发负责人；审核签字←研发负责人签名章。"""
        if not prod_id or not isinstance(content, dict):
            return content
        members = db.session.execute(
            select(ProjectMember).where(ProjectMember.prod_id == prod_id)
        ).scalars().all()

        def role_name(*keys):
            for m in members:
                if any(k in str(m.role or "") for k in keys):
                    return (m.name or "").strip()
            return ""

        rev_date = serv_review_util.cover_date(prod_id, "crr") or ""
        # 检查日期：始终从时间线获取（时间线有值即覆盖，保证与时间线同步）
        if rev_date:
            content["check_date"] = rev_date
        if not str(content.get("auditee") or "").strip():
            content["auditee"] = role_name("TPM")
        if not str(content.get("auditor") or "").strip():
            content["auditor"] = role_name("研发负责人")
        auditor = str(content.get("auditor") or "").strip()
        # 审核签字：研发负责人签名章（仅在当前非签名图时填充）
        if not str(content.get("sign_img") or "").startswith("data:image"):
            content["sign_img"] = (serv_review_util._sign_by_name(auditor) if auditor else "") or ""
        # 审核日期与检查日期保持一致
        content["sign_date"] = str(content.get("check_date") or "")
        return content

    def __to_obj(self, row: CrrDoc, product: Product = None):
        obj = CrrDocObj(**row.dict())
        obj.content = self.__autofill(self.__normalize_content(obj.content), row.product_id)
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
        return obj

    async def add_crr_doc(self, form: CrrDocForm):
        try:
            sql = select(func.count(CrrDoc.id)).where(
                CrrDoc.product_id == form.product_id,
                CrrDoc.version == form.version,
            )
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = CrrDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=CrrDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_crr_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: CrrDoc = db.session.execute(select(CrrDoc).where(CrrDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(CrrDoc.version).where(CrrDoc.product_id == target_pid)).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = CrrDoc(
                product_id=target_pid,
                version=version,
                file_no=sync_file_no_version(fromdoc.file_no, version),
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            return Resp.resp_ok(data=CrrDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_crr_doc(self, form: CrrDocForm):
        try:
            row: CrrDoc = db.session.execute(select(CrrDoc).where(CrrDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_crr_doc(self, id: int):
        db.session.execute(delete(CrrDoc).where(CrrDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_crr_doc(self, id: int):
        sql = select(CrrDoc, Product).join(Product, CrrDoc.product_id == Product.id).where(CrrDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_crr_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(CrrDoc.product_id == product_id)
        if version:
            wheres.append(CrrDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(CrrDoc.id)).join(Product, CrrDoc.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (
            select(CrrDoc, Product)
            .join(Product, CrrDoc.product_id == Product.id)
            .where(*wheres)
            .order_by(CrrDoc.id.desc())
            .offset(page_index * page_size)
            .limit(page_size)
        )
        rows: List[CrrDocObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    # ---------------- 导出 Word ----------------
    async def export_crr_doc(self, output, id: int):
        resp = await self.get_crr_doc(id)
        obj: CrrDocObj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        c = self.__normalize_content(obj.content)
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")
        docx_util.add_page_number_footer(section, obj.file_no or "", skip_first=False)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            s = str(text or "")
            if s.startswith("data:image"):
                try:
                    b64 = s.split(",", 1)[1] if "," in s else ""
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = align
                    para.add_run().add_picture(BytesIO(base64.b64decode(b64)), height=Pt(30))
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    return
                except Exception:
                    pass
            cell.text = ""
            for i, line in enumerate(s.split("\n")):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        def set_check(cell, checked):
            # 复选框：选中☑(框内对号)、未选☐(空框)；强制宋体避免渲染成彩色 emoji。
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # ☑ 加文本呈现选择符(U+FE0E)强制单色白底，避免渲染成灰色 emoji
            run = para.add_run("\u2611\ufe0e" if checked else "\u2610")
            run.font.size = Pt(12)
            run.font.name = "宋体"
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            for _attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                rfonts.set(qn(_attr), "宋体")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # 标题
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        docx_util.fonted_txt(title, "代码审查记录", font_size=18.0, bold=True)

        # 抬头表：代码地址 / 检查日期 / 被审核人 / 审核人 / 审核依据
        head = document.add_table(rows=0, cols=4)
        head.style = "Table Grid"
        head.alignment = WD_TABLE_ALIGNMENT.CENTER

        def head_row(label1, val1, label2="", val2="", merge_val=False):
            cells = head.add_row().cells
            set_cell(cells[0], label1, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            if merge_val:
                merged = cells[1].merge(cells[2]).merge(cells[3])
                set_cell(merged, val1)
            else:
                set_cell(cells[1], val1)
                set_cell(cells[2], label2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell(cells[3], val2)

        # 代码地址：去掉制表符（避免 Word 里被顶到中间），保证左对齐
        head_row("代码地址", str(c.get("code_url", "") or "").replace("\t", " "), "检查日期", c.get("check_date", ""))
        head_row("被审核人", c.get("auditee", ""), "审核人", c.get("auditor", ""))
        head_row("审核依据", c.get("basis", ""), merge_val=True)
        document.add_paragraph()

        # 检查表
        checklist = [r for r in (c.get("checklist") or []) if isinstance(r, list)]
        cols = 6
        tbl = document.add_table(rows=0, cols=cols)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        # 列宽：问题列加宽（编号/问题/是/否/不适用/备注），单位 dxa(1/20 pt)
        col_dxa = [720, 4680, 720, 720, 1008, 2160]
        cat_flags = []
        for r_idx, row in enumerate(checklist):
            cells = tbl.add_row().cells
            first = str(row[0] if row else "").strip()
            is_cat = first in _CATEGORIES and all(not str(row[i] if i < len(row) else "").strip() for i in range(1, cols))
            cat_flags.append(is_cat)
            if is_cat:
                merged = cells[0]
                for i in range(1, cols):
                    merged = merged.merge(cells[i])
                set_cell(merged, first, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
            else:
                for c_idx in range(cols):
                    val = row[c_idx] if c_idx < len(row) else ""
                    # 是/否/不适用 三列以复选框展示：选中框内对号、未选空框（数据行，非表头）
                    if r_idx != 0 and c_idx in (2, 3, 4):
                        set_check(cells[c_idx], bool(str(val).strip()))
                        continue
                    align = WD_ALIGN_PARAGRAPH.CENTER if c_idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
                    set_cell(cells[c_idx], val, bold=(r_idx == 0), align=align)
        # 固定列宽：设置 tblLayout=fixed + tblGrid，并逐格写入 tcW
        tbl.autofit = False
        _tblPr = tbl._tbl.tblPr
        _layout = _tblPr.find(qn("w:tblLayout"))
        if _layout is None:
            _layout = OxmlElement("w:tblLayout")
            _tblPr.append(_layout)
        _layout.set(qn("w:type"), "fixed")
        _grid = tbl._tbl.find(qn("w:tblGrid"))
        if _grid is not None:
            for _gc in list(_grid):
                _grid.remove(_gc)
            for _w in col_dxa:
                _gc = OxmlElement("w:gridCol")
                _gc.set(qn("w:w"), str(_w))
                _grid.append(_gc)
        for _ri, _r in enumerate(tbl.rows):
            # 跳过整行合并的分类行
            if _ri < len(cat_flags) and cat_flags[_ri]:
                continue
            _cells = _r.cells
            for _i, _w in enumerate(col_dxa):
                if _i < len(_cells):
                    _tcpr = _cells[_i]._tc.get_or_add_tcPr()
                    _tcw = _tcpr.find(qn("w:tcW"))
                    if _tcw is None:
                        _tcw = OxmlElement("w:tcW")
                        _tcpr.append(_tcw)
                    _tcw.set(qn("w:w"), str(_w))
                    _tcw.set(qn("w:type"), "dxa")
        document.add_paragraph()

        # 结论
        concl_tbl = document.add_table(rows=1, cols=1)
        concl_tbl.style = "Table Grid"
        concl_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_cell(concl_tbl.rows[0].cells[0], f"结论：{c.get('conclusion', '') or ''}")
        document.add_paragraph()

        # 审核人（签字）/日期
        sign_tbl = document.add_table(rows=1, cols=3)
        sign_tbl.style = "Table Grid"
        sign_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        srow = sign_tbl.rows[0].cells
        # 加大行高，避免签名图超出单元格
        sign_tbl.rows[0].height = Pt(52)
        sign_tbl.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        set_cell(srow[0], "审核人（签字）/日期", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(srow[1], c.get("sign_img", "") or "", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(srow[2], c.get("sign_date", "") or "", align=WD_ALIGN_PARAGRAPH.CENTER)

        document.save(output)
        output.seek(0)
