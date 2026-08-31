#!/usr/bin/env python
# encoding: utf-8

# 模型文件服务层，详见 docs/function_docs/99_模型文件管理.md。
# 单表 model_doc + doc_type；导出结构与产品立项报告一致：封面→分页→修订记录→分页→目录→分页→正文。

import base64
import copy
import logging
import os
import re
from datetime import date, timedelta
from io import BytesIO
from typing import List
from sqlalchemy import delete, func, select
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE

from ..model.product import Product
from ..model.model_doc import ModelDoc
from ..model.project_timeline import ProjectTimelineRow, ProjectTimelineCell
from ..model.project_member import ProjectMember
from ..model.srs_doc import SrsDoc
from ..model.srs_req import SrsReq
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_model_doc import ModelDocForm
from ..obj.vobj_model_doc import ModelDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from . import serv_review_util
from .serv_utils import new_version, sync_file_no_version
from .serv_utils import docx_util
from .model_doc_templates import DOC_META as WORD_META, DEFAULT_CONTENTS as WORD_CONTENTS, REVIEW_TABLES as WORD_REVIEW_TABLES
from .model_doc_xlsx_templates import XLSX_META, XLSX_CONTENTS
from .model_doc_build_templates import BUILD_DOC_TYPES, BUILD_DEFAULTS
from .model_doc_train_templates import TRAIN_DOC_TYPES, TRAIN_DEFAULTS

DOC_META = {**WORD_META, **XLSX_META}
DEFAULT_CONTENTS = {**WORD_CONTENTS, **XLSX_CONTENTS}
REVIEW_TABLES = WORD_REVIEW_TABLES
EQ_DOC_TYPES = ("md_deq", "md_teq", "md_eq")
CRR_DOC_TYPES = ("md_008_01", "md_008_02")
_CRR_CATEGORIES = ("结构", "文档", "变量", "算法操作", "循环和分支")
_CRR_CONCLUSIONS = ("通过", "有条件通过", "不通过")
_MD008_URL = {
    "md_008_01": "http://172.16.6.3:8081/model/pe/pe-segmetation",
    "md_008_02": "http://172.16.6.3:8081/model/pe/lobe_segmentation",
}
_MD008_CHECKLIST = [
    ["编号", "问题", "是", "否", "不适用", "备注"],
    ["结构", "", "", "", "", ""],
    ["1", "代码是否符合相关的编码标准?", "√", "", "", ""],
    ["2", "代码结构是否适当，风格和格式是否保持一致?", "√", "", "", ""],
    ["3", "代码中是否有没有被调用的或无用的程序，或没有被执行的代码?", "", "√", "", ""],
    ["4", "是否有过于复杂的模块需要重新构造或拆分成多个程序?", "", "√", "", ""],
    ["文档", "", "", "", "", ""],
    ["1", "代码是否已被用易于维护的注释方式清晰充分的文档化?", "√", "", "", ""],
    ["2", "注释是否与代码协调一致?", "√", "", "", ""],
    ["变量", "", "", "", "", ""],
    ["1", "所有变量的命名是否清晰，一致并且有意义?", "√", "", "", ""],
    ["2", "是否有冗余或无用的变量?", "", "√", "", ""],
    ["算法操作", "", "", "", "", ""],
    ["1", "被除数是否做了零值测试?", "√", "", "", ""],
    ["循环和分支", "", "", "", "", ""],
    ["1", "所有的循环，分支和逻辑构造是否完整，正确并且嵌套适当?", "√", "", "", ""],
    ["2", "每种状况是否都有缺省值?", "√", "", "", ""],
]

logger = logging.getLogger(__name__)

COVER_DEPT = "模型部"
SKIP_ANNEX_NUM = {"md_001", "md_004"}
DELETED_SRS_VERSION_PREFIX = "__deleted_srs__"
MD022_ID_COLS = ("算法设计ID", "训练集构建", "调优集构建ID", "算法训练ID", "测试集构建ID", "算法测试ID")
MD022_MODULE_DOC_TYPES = {
    "肺栓塞分诊": {
        "算法设计ID": "md_004",
        "训练集构建": "md_009_01",
        "调优集构建ID": "md_010_01",
        "算法训练ID": "md_012_01",
        "测试集构建ID": "md_011_01",
        "算法测试ID": "md_013_01",
    },
    "肺叶分割": {
        "算法设计ID": "md_004",
        "训练集构建": "md_009_02",
        "调优集构建ID": "md_010_02",
        "算法训练ID": "md_012_02",
        "测试集构建ID": "md_011_02",
        "算法测试ID": "md_013_02",
    },
}
MD022_FILE_TYPES = sorted({dt for mapping in MD022_MODULE_DOC_TYPES.values() for dt in mapping.values()})
MD022_MODULES = ("肺栓塞分诊", "肺叶分割", "气管分割", "肺血管分割")
ENV_CHECK_GROUPS = {
    ("md_019", "server"): [
        ("日期", []),
        ("硬件环境", ["CPU", "GPU", "内存", "网卡"]),
        ("软件环境", ["操作系统\n运行是否正常", "数据库\n运行是否正常", "应用服务\n运行是否正常"]),
        ("开发环境\n是否更新升级", []),
        ("服务器\n是否杀毒", []),
        ("网络环境\n是否正常", []),
        ("开发工具", ["是否正常运行", "是否更新升级"]),
        ("服务器\n是否备份", []),
        ("服务器\n日志是否错误", []),
        ("出现的问题及处理方式", []),
        ("检查人", []),
    ],
    ("md_019", "dev"): [
        ("日期", []),
        ("硬件环境", ["CPU", "GPU", "内存", "网卡"]),
        ("软件环境", ["操作系统\n运行是否正常", "浏览器\n运行是否正常"]),
        ("开发环境\n是否更新升级", []),
        ("开发机\n是否杀毒", []),
        ("网络环境\n是否正常", []),
        ("开发工具", ["是否正常运行", "是否更新升级"]),
        ("出现的问题及处理方式", []),
        ("检查人", []),
    ],
    ("md_020", "server"): [
        ("日期", []),
        ("硬件环境", ["CPU", "GPU", "内存", "网卡"]),
        ("软件环境", ["操作系统\n运行是否正常", "数据库\n运行是否正常", "应用服务\n运行是否正常"]),
        ("测试环境\n是否更新升级", []),
        ("服务器\n是否杀毒", []),
        ("网络环境\n是否正常", []),
        ("测试工具", ["是否正常运行", "是否更新升级"]),
        ("服务器\n是否备份", []),
        ("服务器\n日志是否错误", []),
        ("出现的问题及处理方式", []),
        ("检查人", []),
    ],
    ("md_020", "dev"): [
        ("日期", []),
        ("硬件环境", ["CPU", "GPU", "内存", "网卡"]),
        ("软件环境", ["操作系统\n运行是否正常", "浏览器\n运行是否正常"]),
        ("测试环境\n是否更新升级", []),
        ("测试机\n是否杀毒", []),
        ("网络环境\n是否正常", []),
        ("测试工具", ["是否正常运行", "是否更新升级"]),
        ("出现的问题及处理方式", []),
        ("检查人", []),
    ],
}
_MD007_IMG_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "src-res", "model_doc", "md_007"
)
_MD007_IMG_FILES = {
    "fig1": "fig1_overview.png",
    "pe_flow": "pe_flow.png",
    "fig2": "fig2_patch.png",
    "fig3": "fig3_unet.png",
    "fig4": "fig4_receptive.png",
    "lobe_flow": "lobe_flow.png",
    "recon_flow": "recon_flow.png",
    "cube": "cube.png",
}
_MD007_IMG_CACHE = {}


def doc_title(doc_type):
    return (DOC_META.get(doc_type) or {}).get("title") or "模型文件"


def doc_keywords(doc_type):
    return list((DOC_META.get(doc_type) or {}).get("keywords") or [])


def doc_format(doc_type):
    return (DOC_META.get(doc_type) or {}).get("format") or "docx"


def _empty_template(doc_type):
    title = doc_title(doc_type)
    return {
        "sections": [
            {
                "title": title, "ref_type": "cover", "body": "", "children": [],
                "tables": [[
                    ["编制部门", COVER_DEPT, "文件版本", "A0"],
                    ["编制人", "", "日期", ""],
                    ["审核人", "", "日期", ""],
                    ["批准人", "", "日期", ""],
                    ["生效日期", "", "", ""],
                ]],
            },
            {
                "title": "文件修订记录", "ref_type": "revision", "body": "", "children": [],
                "tables": [[
                    ["修改日期", "版本号", "修订说明", "修订人", "批准人"],
                    ["", "", "首次发布", "", ""],
                    ["", "", "", "", ""],
                    ["", "", "", "", ""],
                    ["", "", "", "", ""],
                    ["", "", "", "", ""],
                ]],
            },
        ]
    }


class Server(object):

    def __default_content(self, doc_type):
        if doc_type in EQ_DOC_TYPES:
            return {"rows": self.__eq_default_rows(doc_type)}
        if doc_type in CRR_DOC_TYPES:
            return self.__crr_default_content(doc_type)
        if doc_type in BUILD_DOC_TYPES:
            return self.__build_default_content(doc_type)
        if doc_type in TRAIN_DOC_TYPES:
            return self.__train_default_content(doc_type)
        raw = DEFAULT_CONTENTS.get(doc_type)
        content = copy.deepcopy(raw) if raw else _empty_template(doc_type)
        self.__drop_product_info(content)
        self.__ensure_review_annex(content, doc_type)
        if doc_type == "md_007":
            self.__fill_md007_algo_info(content)
        if doc_type in ("md_019", "md_020"):
            self.__complete_env_maint_chapter(content, doc_type)
        return content

    def __to_obj(self, row: ModelDoc, product: Product = None):
        obj = ModelDocObj(**row.dict())
        obj.content = self.__normalize_content(obj.content, row.doc_type)
        key = row.doc_type or ""
        if key not in EQ_DOC_TYPES and key not in CRR_DOC_TYPES and key not in BUILD_DOC_TYPES and key not in TRAIN_DOC_TYPES:
            self.__fill_cover_meta(obj.content, obj.version)
            serv_review_util.fill_cover_dates(
                obj.content, serv_review_util.cover_date(row.product_id, key) if row.product_id else ""
            )
            serv_review_util.fill_cover_signers(
                obj.content, serv_review_util.cover_signers(row.product_id, key) if row.product_id else {}
            )
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
            if not (obj.file_no or "").strip():
                resolved = serv_review_util.resolve_doc_file_no(product.id, obj.file_no, obj.version, key)
                if resolved:
                    obj.file_no = resolved
        return obj

    @staticmethod
    def __migrate_cover_table(rows):
        """把旧版 2 列或「使用部门/版本号」封面迁移为 4 列：编制部门 / 模型部。"""
        rows = [r for r in (rows or []) if isinstance(r, list)]
        if rows and len(rows[0]) >= 4 and str(rows[0][0]).strip() == "编制部门":
            if not str(rows[0][1] if len(rows[0]) > 1 else "").strip():
                rows[0][1] = COVER_DEPT
            return rows
        items = [(str(r[0]).strip(), str(r[1]).strip() if len(r) > 1 else "") for r in rows if r]
        def val(*labels):
            for want in labels:
                for l, v in items:
                    if l == want:
                        return v
            return ""
        dates = [v for l, v in items if l == "日期"]
        d = lambda i: dates[i] if i < len(dates) else ""
        dept = val("编制部门", "使用部门", "编写部门") or COVER_DEPT
        ver = val("文件版本", "版本号") or "A0"
        return [
            ["编制部门", dept, "文件版本", ver],
            ["编制人", val("编制人"), "日期", d(0)],
            ["审核人", val("审核人"), "日期", d(1)],
            ["批准人", val("批准人"), "日期", d(2)],
            ["生效日期", val("生效日期"), "", ""],
        ]

    def __normalize_node(self, node):
        if not isinstance(node, dict):
            return {"title": str(node or ""), "body": "", "tables": [], "children": []}
        result = dict(node)
        result["title"] = str(result.get("title") or "")
        result["body"] = str(result.get("body") or "")
        tables = result.get("tables")
        if not isinstance(tables, list):
            tables = []
        norm_tables = []
        for table in tables:
            if isinstance(table, list):
                norm_tables.append([[str(c) if c is not None else "" for c in (row or [])] for row in table if isinstance(row, list)])
        if result.get("ref_type") == "cover" and norm_tables:
            norm_tables = [self.__migrate_cover_table(norm_tables[0])] + norm_tables[1:]
        result["tables"] = norm_tables
        children = result.get("children")
        result["children"] = [self.__normalize_node(c) for c in children] if isinstance(children, list) else []
        return result

    def __eq_default_rows(self, doc_type):
        found = self.__extract_eq_rows(DEFAULT_CONTENTS.get(doc_type) or {})
        if found:
            return found
        return [["序号", "名称", "规格型号", "品牌", "资产编码/SN码", "类别", "用途", "地点", "使用人", "状态"]]

    def __extract_eq_rows(self, content):
        if not isinstance(content, dict):
            return None
        rows = content.get("rows")
        if isinstance(rows, list) and rows and isinstance(rows[0], list):
            hdr = [str(c or "") for c in rows[0]]
            if any("资产编码" in h for h in hdr) or (hdr and hdr[0].strip() == "序号"):
                return [[str(c or "") for c in r] for r in rows if isinstance(r, list)]

        def walk(ns):
            for n in ns or []:
                if not isinstance(n, dict):
                    continue
                for tb in n.get("tables") or []:
                    if not isinstance(tb, list):
                        continue
                    for i, row in enumerate(tb):
                        if not isinstance(row, list):
                            continue
                        cells = [str(c or "").strip() for c in row]
                        if cells and cells[0] == "序号" and any("资产编码" in c for c in cells):
                            return [[str(c or "") for c in r] for r in tb[i:] if isinstance(r, list)]
                hit = walk(n.get("children") or [])
                if hit:
                    return hit
            return None

        return walk(content.get("sections") or [])

    def __normalize_eq_content(self, content, doc_type):
        rows = self.__extract_eq_rows(content)
        return {"rows": rows or self.__eq_default_rows(doc_type)}

    def __crr_default_content(self, doc_type):
        return {
            "code_url": _MD008_URL.get(doc_type) or "",
            "check_date": "",
            "auditee": "",
            "auditor": "",
            "basis": "《代码管理制度》",
            "method": "代码审查",
            "checklist": copy.deepcopy(_MD008_CHECKLIST),
            "conclusion": "",
            "sign_img": "",
            "sign_date": "",
        }

    def __extract_crr_from_grid(self, tb):
        if not isinstance(tb, list):
            return None
        out = self.__crr_default_content("md_008_01")
        out["checklist"] = [["编号", "问题", "是", "否", "不适用", "备注"]]
        found = False
        for row in tb:
            if not isinstance(row, list):
                continue
            cells = [str(c or "").strip() for c in row]
            a = cells[0] if cells else ""
            if a == "代码地址":
                found = True
                out["code_url"] = cells[2] if len(cells) > 2 and cells[2] else (cells[1] if len(cells) > 1 else "")
                if len(cells) > 5 and cells[5]:
                    out["check_date"] = cells[5]
                elif len(cells) > 3 and cells[2] == "检查日期":
                    out["check_date"] = cells[3] if len(cells) > 3 else ""
            elif a == "被审核人":
                found = True
                out["auditee"] = cells[2] if len(cells) > 2 else (cells[1] if len(cells) > 1 else "")
                if len(cells) > 5:
                    out["auditor"] = cells[5]
            elif a == "审核依据":
                out["basis"] = cells[2] if len(cells) > 2 and cells[2] else (cells[1] if len(cells) > 1 else out["basis"])
            elif a == "审核方式":
                out["method"] = cells[2] if len(cells) > 2 and cells[2] else (cells[1] if len(cells) > 1 else out["method"])
            elif a in _CRR_CATEGORIES:
                out["checklist"].append([a, "", "", "", "", ""])
            elif a.startswith("结论"):
                joined = "".join(cells)
                for name in _CRR_CONCLUSIONS:
                    if name in joined:
                        out["conclusion"] = name
                        break
            elif "签字" in a:
                sign = cells[2] if len(cells) > 2 else ""
                if sign:
                    out["sign_img"] = sign
            elif a == "编号" or (a and "问题" in "".join(cells)):
                continue
            elif a.isdigit():
                if len(cells) >= 7:
                    out["checklist"].append([cells[0], cells[1], cells[3], cells[4], cells[5], cells[6]])
                elif len(cells) >= 6:
                    out["checklist"].append(cells[:6])
        if found:
            out["checklist"] = self.__complete_crr_checklist(out.get("checklist"))
        return out if found else None

    def __crr_pad_row(self, row):
        cells = [str(c or "") for c in (row or [])]
        while len(cells) < 6:
            cells.append("")
        return cells[:6]

    def __crr_parse_groups(self, rows):
        groups = {cat: [] for cat in _CRR_CATEGORIES}
        current = None
        pending = []
        for row in rows or []:
            if not isinstance(row, list):
                continue
            cells = self.__crr_pad_row(row)
            a = cells[0].strip()
            if a == "编号" or ("问题" in "".join(cells[:2]) and not a.isdigit()):
                continue
            if a in _CRR_CATEGORIES:
                current = a
                continue
            if a.isdigit():
                if current:
                    groups[current].append(cells)
                else:
                    pending.append(cells)
        if pending:
            groups["结构"] = pending + groups["结构"]
        return groups

    def __complete_crr_checklist(self, rows):
        """丢掉文件编号/标题空行；缺的分类用默认检查表补，已有分类（含勾选）保留。"""
        groups = self.__crr_parse_groups(rows)
        default_groups = self.__crr_parse_groups(_MD008_CHECKLIST)
        out = [["编号", "问题", "是", "否", "不适用", "备注"]]
        for cat in _CRR_CATEGORIES:
            out.append([cat, "", "", "", "", ""])
            out.extend(groups[cat] or default_groups[cat])
        return out

    def __normalize_crr_content(self, content, doc_type):
        base = self.__crr_default_content(doc_type)
        if not isinstance(content, dict):
            return base
        if isinstance(content.get("checklist"), list) and content.get("checklist"):
            for key in base:
                if content.get(key) is not None:
                    base[key] = content.get(key)
            base["code_url"] = str(base.get("code_url") or "").replace("\t", "  ")
            base["checklist"] = self.__complete_crr_checklist(base.get("checklist"))
            return base
        extracted = None

        def walk(ns):
            nonlocal extracted
            for n in ns or []:
                if extracted or not isinstance(n, dict):
                    continue
                for tb in n.get("tables") or []:
                    extracted = self.__extract_crr_from_grid(tb)
                    if extracted:
                        return
                walk(n.get("children") or [])

        walk(content.get("sections") or [])
        if extracted:
            extracted["code_url"] = extracted.get("code_url") or base["code_url"]
            extracted["basis"] = extracted.get("basis") or base["basis"]
            extracted["method"] = extracted.get("method") or base["method"]
            extracted["checklist"] = self.__complete_crr_checklist(extracted.get("checklist"))
            return extracted
        return base

    def __build_default_content(self, doc_type):
        raw = BUILD_DEFAULTS.get(doc_type) or BUILD_DEFAULTS.get("md_009_01") or {}
        return copy.deepcopy(raw)

    def __build_fmt_pct(self, v):
        s = str(v if v is not None else "").strip()
        if not s:
            return ""
        if s.endswith("%"):
            return s
        try:
            n = float(s)
        except Exception:
            return s
        if abs(n - 1) < 1e-9:
            return "100.00%"
        if abs(n) <= 1.0001:
            return f"{n * 100:.2f}%"
        return f"{n:.2f}%"

    def __build_pad_row(self, row):
        cells = [str(c if c is not None else "") for c in (row or [])]
        while len(cells) < 4:
            cells.append("")
        return cells[:4]

    def __build_parse_qty(self, v):
        s = str(v if v is not None else "").replace(",", "").strip()
        if not s:
            return 0.0
        try:
            return float(s)
        except Exception:
            return 0.0

    def __build_parse_pct(self, v):
        s = str(v if v is not None else "").strip()
        if not s:
            return 0.0
        if s.endswith("%"):
            try:
                return float(s[:-1].strip() or 0)
            except Exception:
                return 0.0
        try:
            n = float(s)
        except Exception:
            return 0.0
        return n * 100.0 if abs(n) <= 1.0001 else n

    def __build_fill_total(self, dist):
        """总计数量/占比按紧挨其上的最后一组因素求和，避免把性别+年龄+设备重复相加。"""
        rows = [self.__build_pad_row(r) for r in (dist or []) if isinstance(r, list)]
        if not rows:
            return [["因素", "类别", "数量", "占比"], ["总计", "", "0", "0.00%"]]
        if str(rows[0][0]).strip() != "因素":
            rows = [["因素", "类别", "数量", "占比"]] + rows
        total_at = next((i for i, r in enumerate(rows) if str(r[0]).strip() == "总计"), -1)
        end = total_at if total_at >= 0 else len(rows)
        start = 1
        for i in range(end - 1, 0, -1):
            a = str(rows[i][0]).strip()
            if a and a != "总计":
                start = i
                break
        qty = 0.0
        pct = 0.0
        for i in range(start, end):
            qty += self.__build_parse_qty(rows[i][2])
            pct += self.__build_parse_pct(rows[i][3])
        qty_s = str(int(round(qty))) if abs(qty - round(qty)) < 1e-9 else str(qty)
        if pct == 0 and qty > 0:
            pct_s = "100.00%"
        elif abs(pct - 100) < 0.05:
            pct_s = "100.00%"
        else:
            pct_s = f"{pct:.2f}%"
        total_row = ["总计", "", qty_s, pct_s]
        if total_at >= 0:
            rows[total_at] = total_row
        else:
            rows.append(total_row)
        return rows

    def __extract_build_from_grid(self, tb):
        if not isinstance(tb, list):
            return None
        out = {
            "author": "", "write_date": "", "data_use": "", "data_type": "",
            "method": "", "case_count": "", "annotator": "",
            "dist_rows": [["因素", "类别", "数量", "占比"]],
            "author_sign": "", "auditor_sign": "",
        }
        found = False
        in_dist = False
        for row in tb:
            if not isinstance(row, list):
                continue
            cells = self.__build_pad_row(row)
            a, b, c, d = [x.strip() for x in cells]
            if a == "编写人":
                found = True
                out["author"] = b
                if c == "编写时间":
                    out["write_date"] = d
            elif a == "数据用途":
                found = True
                out["data_use"] = b
                if c == "数据类型":
                    out["data_type"] = d
            elif a == "构建方法":
                out["method"] = b or c
            elif a == "病例数量":
                out["case_count"] = b
                if "标记" in c:
                    out["annotator"] = d
            elif a == "数据分布":
                in_dist = True
            elif a == "因素" and "类别" in (b + c):
                in_dist = True
            elif "签字" in a:
                out["author_sign"] = b
                if "审核" in c:
                    out["auditor_sign"] = d
                in_dist = False
            elif in_dist and (a or b or c or d):
                if a == "总计":
                    out["dist_rows"].append(["总计", "", c or b, self.__build_fmt_pct(d or "1")])
                else:
                    out["dist_rows"].append([a, b, c, self.__build_fmt_pct(d)])
        return out if found else None

    def __normalize_build_content(self, content, doc_type):
        base = self.__build_default_content(doc_type)
        if not isinstance(content, dict):
            base["dist_rows"] = self.__build_fill_total(base.get("dist_rows") or [])
            return base
        if isinstance(content.get("dist_rows"), list) and content.get("dist_rows"):
            for key in base:
                if content.get(key) is not None:
                    base[key] = content.get(key)
            dist = []
            for r in base.get("dist_rows") or []:
                if not isinstance(r, list):
                    continue
                cells = self.__build_pad_row(r)
                dist.append(cells)
            if not dist or str(dist[0][0]).strip() != "因素":
                dist = [["因素", "类别", "数量", "占比"]] + dist
            has_total = any(str(r[0]).strip() == "总计" for r in dist)
            if not has_total:
                base["dist_rows"] = self.__build_default_content(doc_type)["dist_rows"]
            else:
                for r in dist:
                    if str(r[0]).strip() not in ("因素", "总计"):
                        r[3] = self.__build_fmt_pct(r[3])
                base["dist_rows"] = dist
            base["dist_rows"] = self.__build_fill_total(base.get("dist_rows") or [])
            return base
        extracted = None

        def walk(ns):
            nonlocal extracted
            for n in ns or []:
                if extracted or not isinstance(n, dict):
                    continue
                for tb in n.get("tables") or []:
                    extracted = self.__extract_build_from_grid(tb)
                    if extracted:
                        return
                walk(n.get("children") or [])

        walk(content.get("sections") or [])
        if extracted:
            for k in ("author", "write_date", "data_use", "data_type", "method", "case_count", "annotator"):
                extracted[k] = extracted.get(k) or base[k]
            has_total = any(r and str(r[0]).strip() == "总计" for r in extracted.get("dist_rows") or [])
            if not has_total:
                extracted["dist_rows"] = base["dist_rows"]
            extracted["dist_rows"] = self.__build_fill_total(extracted.get("dist_rows") or [])
            return extracted
        base["dist_rows"] = self.__build_fill_total(base.get("dist_rows") or [])
        return base

    def __train_default_content(self, doc_type):
        raw = TRAIN_DEFAULTS.get(doc_type) or TRAIN_DEFAULTS.get("md_012_01") or {}
        return copy.deepcopy(raw)

    def __train_pad_row(self, row, n=3):
        cells = [str(c if c is not None else "") for c in (row or [])]
        while len(cells) < n:
            cells.append("")
        return cells[:n]

    def __train_fill_count(self, content):
        pts = content.get("eval_points") or []
        last = ""
        for r in pts:
            if not isinstance(r, list) or not r:
                continue
            a = str(r[0]).strip()
            if not a or a in ("数据量", "step"):
                continue
            last = a
        if last:
            content["case_count"] = last
        return content

    def __extract_train_from_grid(self, tb, doc_type):
        if not isinstance(tb, list):
            return None
        out = self.__train_default_content(doc_type)
        found = False
        for row in tb:
            if not isinstance(row, list):
                continue
            cells = [str(c or "").strip() for c in row]
            while len(cells) < 7:
                cells.append("")
            a = cells[0]
            if a == "编写人":
                found = True
                out["author"] = cells[1]
                if "编写日期" in cells:
                    i = cells.index("编写日期")
                    out["write_date"] = cells[i + 1] if i + 1 < len(cells) else ""
                if "审核人" in cells:
                    i = cells.index("审核人")
                    out["auditor"] = cells[i + 1] if i + 1 < len(cells) else ""
            elif a == "模型名称":
                found = True
                out["model_name"] = cells[1]
            elif a == "模型功能":
                out["model_func"] = cells[1]
            elif a == "训练集":
                out["train_set"] = cells[1]
                if "数量" in cells:
                    i = cells.index("数量")
                    out["case_count"] = cells[i + 1] if i + 1 < len(cells) else out["case_count"]
            elif a == "训练时间":
                out["train_time"] = cells[1]
            elif a == "硬件环境":
                out["hw_env"] = cells[1]
            elif a == "软件环境":
                out["sw_env"] = cells[1]
            elif a == "结论":
                out["conclusion"] = cells[1]
            elif "编写人" in a and "签字" in a:
                out["author_sign"] = cells[1]
                if any("审核人" in x for x in cells):
                    for i, x in enumerate(cells):
                        if "审核人" in x and i + 1 < len(cells):
                            out["auditor_sign"] = cells[i + 1]
                            break
        return out if found else None

    def __normalize_train_content(self, content, doc_type):
        base = self.__train_default_content(doc_type)
        if not isinstance(content, dict):
            return self.__train_fill_count(base)
        if isinstance(content.get("eval_points"), list) and content.get("eval_points"):
            for key in base:
                if content.get(key) is not None:
                    base[key] = content.get(key)
            if not isinstance(base.get("eval_points"), list) or not base["eval_points"]:
                base["eval_points"] = self.__train_default_content(doc_type)["eval_points"]
            if not isinstance(base.get("process_points"), list) or not base["process_points"]:
                base["process_points"] = self.__train_default_content(doc_type)["process_points"]
            return self.__train_fill_count(base)
        extracted = None

        def walk(ns):
            nonlocal extracted
            for n in ns or []:
                if extracted or not isinstance(n, dict):
                    continue
                for tb in n.get("tables") or []:
                    extracted = self.__extract_train_from_grid(tb, doc_type)
                    if extracted:
                        return
                walk(n.get("children") or [])

        walk(content.get("sections") or [])
        if extracted:
            return self.__train_fill_count(extracted)
        return self.__train_fill_count(base)

    def __normalize_content(self, content, doc_type=None):
        if doc_type in EQ_DOC_TYPES:
            return self.__normalize_eq_content(content, doc_type)
        if doc_type in CRR_DOC_TYPES:
            return self.__normalize_crr_content(content, doc_type)
        if doc_type in BUILD_DOC_TYPES:
            return self.__normalize_build_content(content, doc_type)
        if doc_type in TRAIN_DOC_TYPES:
            return self.__normalize_train_content(content, doc_type)
        if not isinstance(content, dict) or not isinstance(content.get("sections"), list):
            return self.__default_content(doc_type)
        out = {"sections": [self.__normalize_node(s) for s in content["sections"]]}
        if doc_type == "md_001":
            self.__relocate_md001_tables(out)
        if doc_type == "md_006":
            self.__relocate_md006_tables(out)
        if doc_type == "md_007":
            self.__fill_md007_algo_info(out)
        if doc_type in WORD_CONTENTS:
            self.__fill_empty_template_tables(out, doc_type)
        if doc_type in ("md_008_01", "md_008_02"):
            self.__complete_md008_checklist(out, doc_type)
        if doc_type in ("md_019", "md_020"):
            self.__complete_env_maint_chapter(out, doc_type)
        self.__drop_product_info(out)
        self.__ensure_review_annex(out, doc_type)
        return out

    @classmethod
    def __drop_product_info(cls, content):
        """去掉原 Word 没有的「产品信息」章（已存文档打开/导出时也去掉）。"""
        def drop(ns):
            out = []
            for n in ns or []:
                t = cls.__strip_num(n.get("title"))
                if n.get("ref_type") == "basic_info" or t == "产品信息":
                    continue
                n["children"] = drop(n.get("children") or [])
                out.append(n)
            return out
        content["sections"] = drop((content or {}).get("sections") or [])

    @classmethod
    def __is_annex_title(cls, title):
        t = cls.__strip_num(title).replace(" ", "")
        return t.startswith("附件") and "评审记录" in t

    @classmethod
    def __ensure_review_annex(cls, content, doc_type=None):
        """原 Word 有评审表的，附件为空则补上；误挂在其它章节的评审表挪回附件。"""
        sections = (content or {}).get("sections")
        if not isinstance(sections, list):
            return

        def strip_annex_line(n):
            body = str(n.get("body") or "")
            lines = [ln for ln in body.split("\n") if ln.strip() not in ("附件 1 评审记录", "附件1 评审记录")]
            n["body"] = "\n".join(lines).rstrip()

        annex_nodes = []

        def find(ns):
            for n in ns or []:
                if not isinstance(n, dict):
                    continue
                strip_annex_line(n)
                if cls.__is_annex_title(n.get("title")):
                    annex_nodes.append(n)
                find(n.get("children") or [])

        find(sections)
        annex = annex_nodes[0] if annex_nodes else None
        annex_ids = {id(n) for n in annex_nodes}
        pulled = []

        def pull(ns):
            for n in ns or []:
                if not isinstance(n, dict):
                    continue
                if id(n) in annex_ids:
                    pull(n.get("children") or [])
                    continue
                keep = []
                for tb in (n.get("tables") or []):
                    if cls.__is_review_grid(tb):
                        pulled.append(tb)
                    else:
                        keep.append(tb)
                n["tables"] = keep
                pull(n.get("children") or [])

        pull(sections)
        src = copy.deepcopy(REVIEW_TABLES.get(doc_type) or []) if doc_type else []
        table = None
        if annex:
            existing = [tb for tb in (annex.get("tables") or []) if cls.__is_review_grid(tb)]
            table = existing[0] if existing else (pulled[0] if pulled else (src or None))
            if table:
                annex["tables"] = [table]
                annex["title"] = "附件 1 评审记录"
        elif pulled or src:
            table = pulled[0] if pulled else src
            sections.append({"title": "附件 1 评审记录", "body": "", "tables": [table], "children": []})

    @staticmethod
    def __is_review_grid(tb):
        return isinstance(tb, list) and tb and isinstance(tb[0], list) and "评审记录" in str(tb[0][0] or "")

    @staticmethod
    def __table_hdr(tb):
        if not (isinstance(tb, list) and tb and isinstance(tb[0], list) and tb[0]):
            return "", 0
        return str(tb[0][0] or "").strip(), len(tb[0])

    @staticmethod
    def __table_head_cells(tb):
        if not (isinstance(tb, list) and tb and isinstance(tb[0], list)):
            return []
        return [str(c or "").strip() for c in tb[0]]

    def __relocate_md006_tables(self, content):
        """按原 Word：人员表→人员资源；设备表→开发平台；里程碑表→项目开发计划及里程碑。"""
        sections = (content or {}).get("sections") or []
        nodes = {}

        def visit(ns):
            for n in ns or []:
                nodes[self.__strip_num(n.get("title"))] = n
                visit(n.get("children") or [])

        visit(sections)

        def is_person(tb):
            return self.__table_head_cells(tb)[:4] == ["编号", "姓名", "所属部门", "角色"]

        def is_equip(tb):
            return self.__table_head_cells(tb)[:3] == ["编号", "设备", "设备名称"]

        def is_mile(tb):
            return self.__table_head_cells(tb)[:5] == ["阶段", "任务划分", "负责人", "计划完成时间", "阶段性交付物"]

        def collect_and_strip(pred, keep_node):
            moved = []

            def walk(ns):
                for n in ns or []:
                    stay, take = [], []
                    for tb in (n.get("tables") or []):
                        if pred(tb) and n is not keep_node:
                            take.append(tb)
                        else:
                            stay.append(tb)
                    n["tables"] = stay
                    moved.extend(take)
                    walk(n.get("children") or [])

            walk(sections)
            return moved

        for title, pred in (
            ("人员资源", is_person),
            ("开发平台", is_equip),
            ("项目开发计划及里程碑", is_mile),
        ):
            target = nodes.get(title)
            if target is None:
                continue
            misplaced = collect_and_strip(pred, target)
            if misplaced and not any(pred(tb) for tb in (target.get("tables") or [])):
                target["tables"] = misplaced + (target.get("tables") or [])

    def __relocate_md001_tables(self, content):
        """按原 Word 把误挂章节的表挪回：评审表→附件；SCI 表→标识配置；工具表→版本更新原则。"""
        sections = (content or {}).get("sections") or []
        nodes = {}

        def visit(ns):
            for n in ns or []:
                nodes[self.__strip_num(n.get("title"))] = n
                visit(n.get("children") or [])

        visit(sections)

        def pull(title, pred):
            node = nodes.get(title)
            if not node:
                return []
            tbs = node.get("tables") or []
            moved = [tb for tb in tbs if pred(tb)]
            if moved:
                node["tables"] = [tb for tb in tbs if not pred(tb)]
            return moved

        tool, annex = nodes.get("配置管理工具"), nodes.get("附件 1 评审记录")
        if tool and annex:
            moved = pull("配置管理工具", self.__is_review_grid)
            if moved and not any(self.__is_review_grid(tb) for tb in (annex.get("tables") or [])):
                annex["tables"] = moved + (annex.get("tables") or [])

        ident = nodes.get("标识配置")
        if ident is not None and not (ident.get("tables") or []):
            def is_sci5(tb):
                h, n = self.__table_hdr(tb)
                return h == "SCI名称" and n >= 5
            def is_sci4(tb):
                h, n = self.__table_hdr(tb)
                return h == "SCI名称" and n == 4
            moved = pull("目的", is_sci5) + pull("范围", is_sci4) + pull("目的", is_sci4) + pull("范围", is_sci5)
            if moved:
                ident["tables"] = moved

        verp = nodes.get("版本更新原则")
        if verp is not None and not (verp.get("tables") or []):
            def is_tool(tb):
                return self.__table_hdr(tb)[0] == "工具类型"
            moved = pull("缩写", is_tool)
            if moved:
                verp["tables"] = moved

        prod_name = ""
        def take_name(ns):
            nonlocal prod_name
            for n in ns or []:
                t = self.__strip_num(n.get("title"))
                if n.get("ref_type") == "basic_info" or t == "产品信息":
                    for tb in n.get("tables") or []:
                        for row in tb or []:
                            if isinstance(row, list) and str(row[0] if row else "").strip() == "产品名称":
                                v = str(row[1] if len(row) > 1 else "").strip()
                                if v:
                                    prod_name = v
                                    return
                take_name(n.get("children") or [])
        take_name(sections)

        def drop_info(ns):
            out = []
            for n in ns or []:
                t = self.__strip_num(n.get("title"))
                if n.get("ref_type") == "basic_info" or t == "产品信息":
                    continue
                n["children"] = drop_info(n.get("children") or [])
                out.append(n)
            return out
        content["sections"] = drop_info(sections)
        nodes.clear()
        visit(content["sections"])

        scope = nodes.get("范围")
        if scope is not None:
            from_tbl = ""
            keep = []
            for tb in (scope.get("tables") or []):
                if self.__is_prod_name_table(tb):
                    for row in tb:
                        if isinstance(row, list) and str(row[0] if row else "").strip() == "产品名称":
                            v = str(row[1] if len(row) > 1 else "").strip()
                            if v and not from_tbl:
                                from_tbl = v
                    continue
                keep.append(tb)
            scope["tables"] = keep
            scope["body"] = self.__fill_scope_body(scope.get("body") or "", from_tbl or prod_name)

    @classmethod
    def __md007_data_url(cls, key):
        if key in _MD007_IMG_CACHE:
            return _MD007_IMG_CACHE[key]
        name = _MD007_IMG_FILES.get(key) or ""
        path = os.path.join(_MD007_IMG_DIR, name) if name else ""
        data = ""
        if path and os.path.isfile(path):
            with open(path, "rb") as f:
                data = "data:image/png;base64," + base64.b64encode(f.read()).decode("ascii")
        _MD007_IMG_CACHE[key] = data
        return data

    @staticmethod
    def __tables_have_text(tables):
        for tb in tables or []:
            for row in tb or []:
                for c in row or []:
                    s = str(c or "").strip()
                    if s and not s.startswith("data:image"):
                        return True
        return False

    @staticmethod
    def __has_figure_table(tables):
        for tb in tables or []:
            if not isinstance(tb, list) or not tb:
                continue
            cols = max((len(row) for row in tb if isinstance(row, list)), default=0)
            if cols != 1:
                continue
            for row in tb:
                if isinstance(row, list) and row and str(row[0] or "").startswith("data:image"):
                    return True
        return False

    @staticmethod
    def __fill_flow_cell(tb, url):
        if not url or not isinstance(tb, list):
            return
        for row in tb:
            if not isinstance(row, list) or not row:
                continue
            if "算法流程图" not in str(row[0] or ""):
                continue
            while len(row) < 2:
                row.append("")
            if not str(row[1] or "").strip():
                row[1] = url

    @classmethod
    def __collect_md007_src(cls):
        src = {}

        def visit(ns, parents):
            for n in ns or []:
                if not isinstance(n, dict):
                    continue
                title = cls.__strip_num(n.get("title"))
                path = parents + [title]
                if title == "算法基本信息":
                    if "肺栓塞分割模块" in path:
                        src["pe"] = n
                    elif "肺叶分割模块" in path:
                        src["lobe"] = n
                    elif "三维重建模块" in path:
                        src["recon"] = n
                    else:
                        src["h1"] = n
                elif title == "模块设计描述":
                    if "肺栓塞分割模块" in path:
                        src["pe_desc"] = n
                    elif "肺叶分割模块" in path:
                        src["lobe_desc"] = n
                    elif "三维重建模块" in path:
                        src["recon_desc"] = n
                visit(n.get("children") or [], path)

        visit((WORD_CONTENTS.get("md_007") or {}).get("sections") or [], [])
        return src

    def __fill_md007_algo_info(self, content):
        """原 Word：三个「算法基本信息」为两列表；图 1～6 与立方体示意图按章节补上。空才填。"""
        sections = (content or {}).get("sections")
        if not isinstance(sections, list):
            return
        src = self.__collect_md007_src()
        imgs = {k: self.__md007_data_url(k) for k in _MD007_IMG_FILES}

        def ensure_tables(node, src_node, flow_key=None, fig_keys=None):
            if not isinstance(node, dict):
                return
            tables = node.get("tables") if isinstance(node.get("tables"), list) else []
            if not self.__tables_have_text(tables) and not self.__has_figure_table(tables):
                src_tables = copy.deepcopy((src_node or {}).get("tables") or []) if src_node else []
                if src_tables:
                    tables = src_tables
            if flow_key:
                for tb in tables:
                    self.__fill_flow_cell(tb, imgs.get(flow_key) or "")
            if fig_keys and not self.__has_figure_table(tables):
                extra = [[[imgs[k]]] for k in fig_keys if imgs.get(k)]
                tables = list(tables) + extra
            node["tables"] = tables

        def visit(ns, parents):
            for n in ns or []:
                if not isinstance(n, dict):
                    continue
                title = self.__strip_num(n.get("title"))
                path = parents + [title]
                if title == "算法基本信息":
                    if "肺栓塞分割模块" in path:
                        ensure_tables(n, src.get("pe"), flow_key="pe_flow")
                    elif "肺叶分割模块" in path:
                        ensure_tables(n, src.get("lobe"), flow_key="lobe_flow")
                    elif "三维重建模块" in path:
                        ensure_tables(n, src.get("recon"), flow_key="recon_flow")
                    else:
                        ensure_tables(n, src.get("h1"), fig_keys=["fig1"])
                elif title == "模块设计描述":
                    if "肺栓塞分割模块" in path:
                        ensure_tables(n, None, fig_keys=["fig2", "fig3", "fig4"])
                    elif "肺叶分割模块" in path:
                        ensure_tables(n, None, fig_keys=["lobe_flow"])
                    elif "三维重建模块" in path:
                        ensure_tables(n, None, fig_keys=["recon_flow", "cube"])
                visit(n.get("children") or [], path)

        visit(sections, [])

    def __fill_empty_template_tables(self, content, doc_type):
        """原 Word 章节表：按标题路径从模板补到空章节。封面/修订/产品信息不走此逻辑。空才填。"""
        src = (WORD_CONTENTS.get(doc_type) or {}).get("sections")
        sections = (content or {}).get("sections")
        if not isinstance(src, list) or not isinstance(sections, list):
            return
        skip = {"cover", "revision", "basic_info"}
        src_map = {}

        def visit_src(ns, parents):
            for n in ns or []:
                if not isinstance(n, dict):
                    continue
                title = self.__strip_num(n.get("title"))
                src_map[tuple(parents + [title])] = n
                visit_src(n.get("children") or [], parents + [title])

        visit_src(src, [])

        def visit(ns, parents):
            for n in ns or []:
                if not isinstance(n, dict):
                    continue
                title = self.__strip_num(n.get("title"))
                path = tuple(parents + [title])
                src_n = src_map.get(path)
                if (
                    src_n
                    and src_n.get("ref_type") not in skip
                    and n.get("ref_type") not in skip
                ):
                    src_tables = copy.deepcopy(src_n.get("tables") or [])
                    if src_tables and not self.__tables_have_text(n.get("tables")) and not self.__has_figure_table(n.get("tables")):
                        n["tables"] = src_tables
                visit(n.get("children") or [], parents + [title])

        visit(sections, [])

    @staticmethod
    def __md008_row0(row):
        return str(row[0] if row else "").strip()

    @classmethod
    def __is_md008_checklist(cls, tb):
        if not isinstance(tb, list) or len(tb) < 3:
            return False
        has_header = False
        has_addr = False
        for row in tb:
            if not isinstance(row, list):
                continue
            if cls.__md008_row0(row) == "代码地址":
                has_addr = True
            cells = [str(c or "").replace(" ", "") for c in row]
            if "编号" in cells and "是" in cells and "否" in cells and any("不适用" in x for x in cells):
                has_header = True
        return has_header and has_addr

    def __complete_md008_checklist(self, content, doc_type):
        src_tb = None
        for n in (DEFAULT_CONTENTS.get(doc_type) or {}).get("sections") or []:
            for tb in n.get("tables") or []:
                if self.__is_md008_checklist(tb):
                    src_tb = tb
                    break
            if src_tb:
                break
        if not src_tb:
            return

        def start_key(tb):
            keys = [self.__md008_row0(r) for r in tb if isinstance(r, list)]
            if "文档" not in keys:
                return "文档"
            if not any(k.startswith("结论") for k in keys):
                return "结论"
            if not any("签字" in k for k in keys):
                return "审核人"
            return ""

        def visit(ns):
            for n in ns or []:
                if not isinstance(n, dict):
                    continue
                tables = n.get("tables") or []
                for i, tb in enumerate(tables):
                    if not self.__is_md008_checklist(tb):
                        continue
                    key = start_key(tb)
                    if not key:
                        continue
                    while tb and isinstance(tb[-1], list) and not any(str(c or "").strip() for c in tb[-1]):
                        tb.pop()
                    src_i = -1
                    for j, r in enumerate(src_tb):
                        a = self.__md008_row0(r)
                        if key == "文档" and a == "文档":
                            src_i = j
                            break
                        if key == "结论" and a.startswith("结论"):
                            src_i = j
                            break
                        if key == "审核人" and "签字" in a:
                            src_i = j
                            break
                    if src_i < 0:
                        continue
                    tables[i] = tb + copy.deepcopy(src_tb[src_i:])
                n["tables"] = tables
                visit(n.get("children") or [])

        visit((content or {}).get("sections") or [])

    @staticmethod
    def __is_env_check_grid(tb):
        return isinstance(tb, list) and tb and isinstance(tb[0], list) and str(tb[0][0] or "").strip() == "env_check"

    @staticmethod
    def __is_asset_grid(tb):
        if not (isinstance(tb, list) and tb and isinstance(tb[0], list) and tb[0]):
            return False
        hdr = [str(c or "") for c in tb[0]]
        return str(hdr[0] or "").strip() == "资产编码" and any("设备信息" in h for h in hdr)

    @staticmethod
    def __env_check_leaves(doc_type, kind):
        cols = []
        for gl, leaves in ENV_CHECK_GROUPS.get((doc_type, kind), ENV_CHECK_GROUPS[("md_019", "dev")]):
            if leaves:
                for lf in leaves:
                    cols.append({"label": lf, "type": "check"})
            else:
                t = "date" if gl == "日期" else "problem" if gl.startswith("出现的问题") else "checker" if gl == "检查人" else "check"
                cols.append({"label": gl, "type": t})
        return cols

    @classmethod
    def __env_check_defaults(cls, doc_type, kind):
        out = []
        for c in cls.__env_check_leaves(doc_type, kind):
            if c["type"] != "check":
                continue
            lb = c["label"]
            out.append("否" if ("更新升级" in lb or "日志是否错误" in lb) else "是")
        return out

    def __complete_env_maint_chapter(self, content, doc_type):
        if doc_type not in ("md_019", "md_020"):
            return
        want = "开发环境维护记录" if doc_type == "md_019" else "测试环境维护记录"
        after = "开发环境定期检查" if doc_type == "md_019" else "测试环境定期检查"
        sections = (content or {}).get("sections")
        if not isinstance(sections, list):
            return

        def find_title(ns, name):
            for n in ns or []:
                if not isinstance(n, dict):
                    continue
                if self.__strip_num(n.get("title")) == name:
                    return n
                hit = find_title(n.get("children") or [], name)
                if hit:
                    return hit
            return None

        if find_title(sections, want):
            return
        new_node = {"title": want, "body": "", "tables": [], "children": []}
        idx = next((i for i, n in enumerate(sections) if isinstance(n, dict) and self.__strip_num(n.get("title")) == after), -1)
        if idx >= 0:
            sections.insert(idx + 1, new_node)
        else:
            sections.append(new_node)

    @staticmethod
    def __cell_eq_nonempty(a, b):
        sa, sb = str(a or ""), str(b or "")
        return sa == sb and sa.strip() != ""

    def __grid_span_origins(self, grid):
        rows = [list(r) for r in grid]
        r_n = len(rows)
        c_n = max((len(r) for r in rows), default=0)
        for r in rows:
            while len(r) < c_n:
                r.append("")
        skip = [[False] * c_n for _ in range(r_n)]
        colspan = [[1] * c_n for _ in range(r_n)]
        rowspan = [[1] * c_n for _ in range(r_n)]
        for r in range(r_n):
            c = 0
            while c < c_n:
                if skip[r][c]:
                    c += 1
                    continue
                c2 = c
                while c2 + 1 < c_n and self.__cell_eq_nonempty(rows[r][c], rows[r][c2 + 1]):
                    c2 += 1
                colspan[r][c] = c2 - c + 1
                for k in range(c + 1, c2 + 1):
                    skip[r][k] = True
                c = c2 + 1
        for r in range(r_n):
            for c in range(c_n):
                if skip[r][c]:
                    continue
                if c != 0:
                    continue
                if not str(rows[r][c] or "").strip():
                    continue
                cs = colspan[r][c]
                r2 = r
                while r2 + 1 < r_n:
                    if skip[r2 + 1][c] or colspan[r2 + 1][c] != cs:
                        break
                    if not self.__cell_eq_nonempty(rows[r][c], rows[r2 + 1][c]):
                        break
                    r2 += 1
                rs = r2 - r + 1
                if rs > 1:
                    rowspan[r][c] = rs
                    for k in range(r + 1, r2 + 1):
                        skip[k][c] = True
        origins = []
        for r in range(r_n):
            for c in range(c_n):
                if skip[r][c]:
                    continue
                origins.append((r, c, rowspan[r][c], colspan[r][c]))
        return rows, origins

    def __md008_span_origins(self, grid):
        categories = {"结构", "文档", "变量", "算法操作", "循环和分支"}
        meta4 = {"代码地址", "被审核人"}
        meta2 = {"审核依据", "审核方式"}
        rows = [list(r) if isinstance(r, list) else [] for r in (grid or [])]
        r_n = len(rows)
        c_n = max([7] + [len(r) for r in rows])
        for r in rows:
            while len(r) < c_n:
                r.append("")
        skip = [[False] * c_n for _ in range(r_n)]
        colspan = [[1] * c_n for _ in range(r_n)]
        rowspan = [[1] * c_n for _ in range(r_n)]

        def empty_row(r):
            return not any(str(rows[r][c] or "").strip() for c in range(c_n))

        def only_first(r):
            if not str(rows[r][0] or "").strip():
                return False
            return not any(str(rows[r][c] or "").strip() for c in range(1, c_n))

        def merge(r, c, rs, cs):
            if r >= r_n or c >= c_n:
                return
            rr = min(rs, r_n - r)
            cc = min(cs, c_n - c)
            rowspan[r][c] = rr
            colspan[r][c] = cc
            skip[r][c] = False
            for i in range(rr):
                for j in range(cc):
                    if i == 0 and j == 0:
                        continue
                    skip[r + i][c + j] = True

        for r in range(r_n):
            if skip[r][0]:
                continue
            a = str(rows[r][0] or "").strip()
            if not a:
                continue
            if a.startswith("结论"):
                merge(r, 0, 2 if r + 1 < r_n and empty_row(r + 1) else 1, c_n)
                continue
            if "审核人" in a and "签字" in a:
                rs = 2 if r + 1 < r_n and empty_row(r + 1) else 1
                merge(r, 0, rs, 2)
                merge(r, 2, rs, c_n - 2)
                continue
            if a in meta4:
                merge(r, 0, 1, 2)
                merge(r, 3, 1, 2)
                merge(r, 5, 1, 2)
                continue
            if a in meta2:
                merge(r, 0, 1, 2)
                merge(r, 2, 1, c_n - 2)
                continue
            if a == "编号" or a.isdigit():
                merge(r, 1, 1, 2)
                continue
            if a in categories or only_first(r):
                merge(r, 0, 1, c_n)
        origins = []
        for r in range(r_n):
            for c in range(c_n):
                if skip[r][c]:
                    continue
                origins.append((r, c, rowspan[r][c], colspan[r][c]))
        return rows, origins

    def __env_maint_span_origins(self, grid):
        rows = [list(r) if isinstance(r, list) else [] for r in (grid or [])]
        r_n = len(rows)
        c_n = max([17] + [len(r) for r in rows])
        for r in rows:
            while len(r) < c_n:
                r.append("")
        skip = [[False] * c_n for _ in range(r_n)]
        colspan = [[1] * c_n for _ in range(r_n)]
        rowspan = [[1] * c_n for _ in range(r_n)]

        def only_first(r):
            if not str(rows[r][0] or "").strip():
                return False
            return not any(str(rows[r][c] or "").strip() for c in range(1, c_n))

        def merge(r, c, rs, cs):
            if r >= r_n or c >= c_n:
                return
            rr = min(rs, r_n - r)
            cc = min(cs, c_n - c)
            rowspan[r][c] = rr
            colspan[r][c] = cc
            skip[r][c] = False
            for i in range(rr):
                for j in range(cc):
                    if i == 0 and j == 0:
                        continue
                    skip[r + i][c + j] = True

        h = -1
        for r in range(r_n):
            if str(rows[r][0] or "").strip() == "日期" and "检查内容" in str(rows[r][1] or ""):
                h = r
                break
        for r in range(r_n):
            if h >= 0 and h <= r <= h + 2:
                continue
            if only_first(r):
                merge(r, 0, 1, c_n)
        if h >= 0:
            merge(h, 0, 3, 1)
            merge(h, 1, 1, 14)
            merge(h, 15, 3, 1)
            merge(h, 16, 3, 1)
            if h + 1 < r_n:
                merge(h + 1, 1, 1, 4)
                merge(h + 1, 5, 1, 3)
                merge(h + 1, 8, 2, 1)
                merge(h + 1, 9, 1, 2)
                merge(h + 1, 11, 2, 1)
                merge(h + 1, 12, 2, 1)
                merge(h + 1, 13, 2, 1)
                merge(h + 1, 14, 2, 1)
        origins = []
        for r in range(r_n):
            for c in range(c_n):
                if skip[r][c]:
                    continue
                origins.append((r, c, rowspan[r][c], colspan[r][c]))
        return rows, origins

    @staticmethod
    def __fill_cover_meta(content, version):
        """封面编制部门 / 文件版本：仅填空。"""
        for section in (content or {}).get("sections") or []:
            if not isinstance(section, dict):
                continue
            for table in (section.get("tables") or []):
                if not isinstance(table, list):
                    continue
                for row in table:
                    if not isinstance(row, list) or not row:
                        continue
                    label = str(row[0] or "").strip()
                    if label in ("编制部门", "使用部门", "编写部门") and len(row) >= 2:
                        if not str(row[1] or "").strip():
                            row[1] = COVER_DEPT
                    if label in ("文件版本", "版本号") and len(row) >= 4:
                        if version and not str(row[3] or "").strip():
                            row[3] = version

    def __autofill_for_export(self, content, obj: ModelDocObj):
        if obj.doc_type in EQ_DOC_TYPES:
            return content
        if obj.doc_type in CRR_DOC_TYPES:
            return self.__fill_crr_fields(content, obj)
        if obj.doc_type in BUILD_DOC_TYPES:
            return self.__normalize_build_content(content, obj.doc_type)
        if obj.doc_type in TRAIN_DOC_TYPES:
            return self.__normalize_train_content(content, obj.doc_type)
        sections = (content or {}).get("sections") or []
        prod_id = obj.product_id
        if not prod_id:
            return content
        product = db.session.execute(select(Product).where(Product.id == prod_id)).scalars().first()
        info = self.__collect_autofill(prod_id, product, obj.version, obj.doc_type)
        for node in sections:
            self.__fill_node(node, info)
        if obj.doc_type in ("md_019", "md_020"):
            self.__rebuild_env_checks(content, info)
        self.__fill_cover_meta(content, obj.version)
        key = obj.doc_type or ""
        serv_review_util.fill_cover_dates(content, serv_review_util.cover_date(prod_id, key))
        serv_review_util.fill_cover_signers(content, serv_review_util.cover_signers(prod_id, key))
        return content

    def __collect_autofill(self, prod_id, product, doc_version, doc_type):
        prod_name = (getattr(product, "name", "") or "").strip()
        full_version = (getattr(product, "full_version", "") or "").strip()
        product_code = (getattr(product, "product_code", "") or "").strip()
        scope = (getattr(product, "scope", "") or "").strip()

        tl_rows = db.session.execute(
            select(ProjectTimelineRow).where(ProjectTimelineRow.prod_id == prod_id)
        ).scalars().all()
        cell_map = {}
        model_row_ids = set()

        def is_model_dev_out(text):
            s = str(text or "")
            if re.search(r"模型开发(?!计划)", s):
                return True
            if "模型训练" in s:
                return True
            if re.search(r"模型测试(?!方案)", s):
                return True
            if "模型封装" in s or "模型服务提交" in s:
                return True
            return False

        if tl_rows:
            for c in db.session.execute(
                select(ProjectTimelineCell).where(ProjectTimelineCell.row_id.in_([r.id for r in tl_rows]))
            ).scalars().all():
                cell_map.setdefault(c.row_id, []).append(c.output_result or "")
                if (c.dept or "") == "模型部" and is_model_dev_out(c.output_result):
                    model_row_ids.add(c.row_id)

        def to_int(v):
            digits = re.sub(r"[^\d]", "", str(v or ""))
            return int(digits) if digits else None

        date_rows = [r for r in tl_rows if (r.row_type or "date") == "date" and to_int(r.year) and to_int(r.month)]

        def date_key(r):
            return to_int(r.year) * 10000 + to_int(r.month) * 100 + (to_int(r.day) or 0)

        kws = doc_keywords(doc_type) or [doc_title(doc_type)]
        file_rows = [
            r for r in date_rows
            if any(any(k in str(v or "") for k in kws) for v in cell_map.get(r.id, []))
        ]
        file_date = ""
        if file_rows:
            fr = min(file_rows, key=date_key)
            file_date = f"{to_int(fr.year)}年{to_int(fr.month)}月{to_int(fr.day)}日"

        cycle_text = ""
        model_dates = []
        if doc_type in ("md_006", "md_019", "md_020"):
            for r in date_rows:
                if r.id not in model_row_ids:
                    continue
                y, m, d = to_int(r.year), to_int(r.month), to_int(r.day) or 1
                try:
                    model_dates.append(date(y, m, d))
                except Exception:
                    continue
            if doc_type == "md_006" and model_dates:
                days = (max(model_dates) - min(model_dates)).days + 1
                if days > 0:
                    cycle_text = "共用时约%d天。" % days

        members = db.session.execute(select(ProjectMember).where(ProjectMember.prod_id == prod_id)).scalars().all()
        def find_member(pred):
            for m in members:
                if pred(str(m.role or "")):
                    return (m.name or "").strip()
            return ""
        modeler = find_member(lambda r: r in ("模型部负责人", "模型负责人")) or find_member(lambda r: "模型" in r)
        algo = find_member(lambda r: "算法" in r)
        approver = find_member(lambda r: "研发负责人" in r) or find_member(lambda r: "负责人" in r)

        md022_file_nos = {}
        md022_srs = {}
        if doc_type == "md_022":
            md022_file_nos = {t: self.__md022_file_no(prod_id, doc_version, t) for t in MD022_FILE_TYPES}
            md022_srs = self.__md022_srs_by_module(prod_id)

        md008_date = md008_auditee = md008_auditor = md008_sign = ""
        if doc_type in ("md_008_01", "md_008_02"):
            auditees = self.__member_names(members, lambda r: r == "算法工程师")
            auditors = self.__member_names(members, lambda r: r == "高级算法工程师")
            md008_date = self.__to_dotted_date(file_date)
            md008_auditee = " ".join(auditees)
            md008_auditor = " ".join(auditors)
            md008_sign = serv_review_util._sign_by_name(auditors[0] if auditors else "") or md008_auditor

        env_weeks = []
        env_assets = None
        env_checker = env_title = ""
        if doc_type in ("md_019", "md_020"):
            env_title = "开发环境维护记录" if doc_type == "md_019" else "测试环境维护记录"
            dev_ds, test_ds = [], []
            for r in date_rows:
                y, m, d = to_int(r.year), to_int(r.month), to_int(r.day) or 1
                try:
                    dt = date(y, m, d)
                except Exception:
                    continue
                vals = cell_map.get(r.id, [])
                if any(("产品开发" in str(v)) and ("计划" not in str(v)) for v in vals):
                    dev_ds.append(dt)
                if any("测试" in str(v) for v in vals):
                    test_ds.append(dt)
            if dev_ds and test_ds:
                env_weeks = self.__week_ranges_from_dates([min(dev_ds), max(test_ds)])
            eq_type = "md_deq" if doc_type == "md_019" else "md_teq"
            eq_doc = db.session.execute(
                select(ModelDoc).where(ModelDoc.product_id == prod_id, ModelDoc.doc_type == eq_type).order_by(ModelDoc.id.desc())
            ).scalars().first()
            if eq_doc:
                env_assets = self.__parse_eq_codes(eq_doc.content if isinstance(eq_doc.content, dict) else {})
            checkers = self.__member_names(members, lambda r: r == "模型部负责人") or self.__member_names(members, lambda r: r == "模型负责人")
            env_checker = serv_review_util._sign_by_name(checkers[0] if checkers else "") or (checkers[0] if checkers else "")

        return {
            "prod_name": prod_name, "full_version": full_version, "product_code": product_code,
            "scope": scope, "file_date": file_date, "version": doc_version,
            "reviser": modeler or algo, "approver": approver, "cycle_text": cycle_text,
            "members": members, "doc_type": doc_type,
            "md022_file_nos": md022_file_nos, "md022_srs": md022_srs,
            "md008_date": md008_date, "md008_auditee": md008_auditee,
            "md008_auditor": md008_auditor, "md008_sign": md008_sign,
            "env_weeks": env_weeks, "env_assets": env_assets, "env_checker": env_checker,
            "env_title": env_title,
        }

    @staticmethod
    def __fill_scope_body(body, prod_name):
        name = str(prod_name or "").strip()
        s = str(body or "")
        if re.search(r"产品名称[：:]\s*\S", s):
            return s
        if re.search(r"产品名称[：:]", s):
            return re.sub(r"产品名称[：:]\s*", ("产品名称：" + name) if name else "产品名称：", s, count=1)
        line = f"产品名称：{name}" if name else "产品名称："
        return f"{line}\n{s}" if s else line

    @staticmethod
    def __is_prod_name_table(tb):
        return (
            isinstance(tb, list) and tb and isinstance(tb[0], list)
            and str(tb[0][0] if tb[0] else "").strip() == "产品名称"
            and len(tb[0]) <= 2
        )

    @staticmethod
    def __strip_num(title):
        return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()

    def __fill_node(self, node, info):
        ref = node.get("ref_type")
        title = self.__strip_num(node.get("title"))
        if ref == "revision" or title == "文件修订记录":
            tables = node.get("tables") or []
            if tables and isinstance(tables[0], list):
                t = tables[0]
                cols = len(t[0]) if t and t[0] else 5
                while len(t) < 6:
                    t.append([""] * cols)
                row = t[1]
                def set_if(i, val):
                    if val and not str(row[i] if i < len(row) else "").strip():
                        row[i] = val
                set_if(0, info["file_date"])
                set_if(1, info["version"])
                if not str(row[2] if len(row) > 2 else "").strip():
                    row[2] = "首次发布"
                set_if(3, info["reviser"])
                set_if(4, info["approver"])
        if ref == "basic_info" or title == "产品信息":
            label_map = {
                "产品名称": info["prod_name"],
                "软件版本": info["full_version"],
                "完整版本": info["full_version"],
                "产品标识": info["product_code"],
                "产品代码": info["product_code"],
                "适用范围": info["scope"],
                "预期用途": info["scope"],
                "项目名称": info["prod_name"],
            }
            for table in (node.get("tables") or []):
                for row in table:
                    if not isinstance(row, list) or len(row) < 2:
                        continue
                    key = str(row[0]).strip()
                    if key in label_map and label_map[key] and not str(row[1] or "").strip():
                        row[1] = label_map[key]
        if title == "范围":
            node["body"] = self.__fill_scope_body(node.get("body") or "", info.get("prod_name") or "")
        is_cycle = ref == "prod_cycle" or (title == "项目开发时间" and not (node.get("children") or []))
        if is_cycle and info.get("cycle_text"):
            node["body"] = info["cycle_text"]
        if info.get("doc_type") == "md_006":
            self.__fill_md006_people_node(node, info.get("members") or [])
        if info.get("doc_type") == "md_017":
            self.__fill_md017_people_node(node, info.get("members") or [])
        if info.get("doc_type") == "md_022":
            self.__fill_md022_trace_node(node, info)
        if info.get("doc_type") in ("md_008_01", "md_008_02"):
            self.__fill_md008_meta_node(node, info)
        if info.get("doc_type") in ("md_019", "md_020"):
            self.__fill_env_maint_node(node, info)
        for child in (node.get("children") or []):
            self.__fill_node(child, info)

    @staticmethod
    def __member_names(members, pred):
        out = []
        for m in members or []:
            role = str(getattr(m, "role", "") or "").strip()
            name = str(getattr(m, "name", "") or "").strip()
            if name and pred(role):
                out.append(name)
        return out

    def __fill_md006_people_node(self, node, members):
        title = self.__strip_num(node.get("title"))
        staff_defs = [
            (lambda r: r in ("模型部负责人", "模型负责人"), "模型部负责人"),
            (lambda r: r == "高级算法工程师", "高级算法工程师"),
            (lambda r: r == "算法工程师", "算法工程师"),
            (lambda r: r == "项目专员", "项目专员"),
        ]
        staff_rows = []
        for pred, label in staff_defs:
            for name in self.__member_names(members, pred):
                staff_rows.append([str(len(staff_rows) + 1), name, "模型部", label])
        pm = (self.__member_names(members, lambda r: "产品经理" in r) or [""])[0]
        testers = self.__member_names(members, lambda r: r == "项目专员")
        algos = self.__member_names(members, lambda r: r == "算法工程师")
        tpm = (self.__member_names(members, lambda r: "TPM" in r.upper()) or [""])[0]
        if not tpm:
            tpm = " ".join(self.__member_names(members, lambda r: "开发人员" in r))
        data_names = " ".join(self.__member_names(members, lambda r: "数据" in r))
        model_dept = " ".join(r[1] for r in staff_rows)

        if title == "项目简介" and pm:
            body = str(node.get("body") or "")
            if re.search(r"产品经理[：:]", body):
                node["body"] = re.sub(r"产品经理[：:][^\n]*", "产品经理： " + pm, body, count=1)
            else:
                node["body"] = (body.rstrip() + ("\n" if body.strip() else "") + "产品经理： " + pm)

        tables = node.get("tables") or []
        if title == "人员资源" and tables and isinstance(tables[0], list) and tables[0]:
            hdr = tables[0][0]
            if isinstance(hdr, list) and "编号" in str(hdr[0] or ""):
                tables[0] = [hdr] + staff_rows

        if "里程碑" in title and tables and isinstance(tables[0], list) and tables[0]:
            t = tables[0]
            header = t[0] if isinstance(t[0], list) else []
            hi = next((i for i, h in enumerate(header) if "负责人" in str(h or "")), -1)
            si = next((i for i, h in enumerate(header) if "阶段" in str(h or "")), -1)
            if hi >= 0:
                for row in t[1:]:
                    if not isinstance(row, list):
                        continue
                    stage = str(row[si] or "") if si >= 0 else " ".join(str(c or "") for c in row)
                    names = testers if "测试" in stage else algos
                    if names:
                        while len(row) <= hi:
                            row.append("")
                        row[hi] = "\n".join(names)

        for i, tb in enumerate(tables):
            if not self.__is_review_grid(tb):
                continue
            new_tb = []
            for row in tb:
                if not isinstance(row, list) or str(row[0] or "").strip() != "参评人员":
                    new_tb.append(row)
                    continue
                next_row = list(row)

                def put(idx):
                    if idx >= len(next_row):
                        return
                    dept = str(next_row[idx] or "").strip()
                    names = ""
                    if dept == "模型部":
                        names = model_dept
                    elif dept == "产品部":
                        names = pm
                    elif "产品开发" in dept:
                        names = tpm
                    elif dept == "数据部":
                        names = data_names
                    if names and idx + 1 < len(next_row):
                        next_row[idx + 1] = names

                put(1)
                put(3)
                new_tb.append(next_row)
            tables[i] = new_tb
        node["tables"] = tables

    def __fill_md017_people_node(self, node, members):
        if self.__strip_num(node.get("title")) != "测试人员":
            return
        tables = node.get("tables") or []
        for ti, tb in enumerate(tables):
            if not isinstance(tb, list) or not tb or not isinstance(tb[0], list):
                continue
            hdr = [str(h or "") for h in tb[0]]
            pi = next((i for i, h in enumerate(hdr) if "资源数量" in h or "具体人员" in h), -1)
            ri = next((i for i, h in enumerate(hdr) if "角色" in h), 0)
            if pi < 0:
                continue
            new_tb = [tb[0]]
            for row in tb[1:]:
                if not isinstance(row, list):
                    new_tb.append(row)
                    continue
                next_row = list(row)
                role = str(next_row[ri] if ri < len(next_row) else "").strip()
                names = self.__member_names(members, lambda r, role=role: r == role) if role else []
                while len(next_row) <= pi:
                    next_row.append("")
                next_row[pi] = ("%d人/%s" % (len(names), " ".join(names))) if names else ""
                new_tb.append(next_row)
            tables[ti] = new_tb
        node["tables"] = tables

    def __md022_file_no(self, prod_id, version, doc_type):
        row = db.session.execute(
            select(ModelDoc).where(ModelDoc.product_id == prod_id, ModelDoc.doc_type == doc_type).order_by(ModelDoc.id.desc())
        ).scalars().first()
        stored = (row.file_no or "").strip() if row else ""
        ver = version or ((row.version or "") if row else "")
        return serv_review_util.resolve_doc_file_no(prod_id, stored, ver, doc_type) or ""

    def __md022_srs_by_module(self, prod_id):
        out = {m: "" for m in MD022_MODULES}
        if not prod_id:
            return out
        doc = db.session.execute(
            select(SrsDoc).where(
                SrsDoc.product_id == prod_id,
                ~SrsDoc.version.like(f"{DELETED_SRS_VERSION_PREFIX}%"),
            ).order_by(SrsDoc.id.desc())
        ).scalars().first()
        if not doc:
            return out
        reqs = db.session.execute(
            select(SrsReq).where(SrsReq.doc_id == doc.id, SrsReq.type_code != "reqd")
        ).scalars().all()

        def blob(row):
            return " ".join(str(getattr(row, f, "") or "") for f in ("module", "function", "sub_function"))

        fallback = ""
        for module in MD022_MODULES:
            hits = [
                str(r.code or "").strip()
                for r in reqs
                if module in blob(r) and str(r.code or "").strip().upper().startswith("SRS-")
            ]
            if hits:
                hits.sort()
                out[module] = hits[0]
                if not fallback:
                    fallback = hits[0]
        if fallback:
            for module in MD022_MODULES:
                if not out[module]:
                    out[module] = fallback
        return out

    def __fill_md022_trace_node(self, node, info):
        if self.__strip_num(node.get("title")) != "模型可追溯性分析表":
            return
        file_nos = info.get("md022_file_nos") or {}
        srs_map = info.get("md022_srs") or {}
        tables = node.get("tables") or []
        for ti, tb in enumerate(tables):
            if not isinstance(tb, list) or not tb or not isinstance(tb[0], list):
                continue
            hdr = [str(h or "").strip() for h in tb[0]]
            if "算法需求" not in hdr or "模块" not in hdr:
                continue
            req_i = hdr.index("算法需求")
            mod_i = hdr.index("模块")
            col_i = {name: hdr.index(name) for name in MD022_ID_COLS if name in hdr}
            new_tb = [tb[0]]
            for row in tb[1:]:
                if not isinstance(row, list):
                    new_tb.append(row)
                    continue
                next_row = list(row)
                while len(next_row) < len(hdr):
                    next_row.append("")
                module = str(next_row[mod_i] if mod_i < len(next_row) else "").strip()
                mapping = MD022_MODULE_DOC_TYPES.get(module) or {}
                next_row[req_i] = srs_map.get(module) or ""
                for name, idx in col_i.items():
                    dt = mapping.get(name)
                    next_row[idx] = (file_nos.get(dt) or "") if dt else ""
                new_tb.append(next_row)
            tables[ti] = new_tb
        node["tables"] = tables

    @staticmethod
    def __to_dotted_date(s):
        m = re.search(r"(\d+)\s*年\s*(\d+)\s*月\s*(\d+)\s*日", str(s or ""))
        if m:
            return "%s.%d.%d" % (m.group(1), int(m.group(2)), int(m.group(3)))
        return str(s or "").strip()

    def __fill_md008_meta_node(self, node, info):
        tables = node.get("tables") or []
        date = info.get("md008_date") or ""
        auditee = info.get("md008_auditee") or ""
        auditor = info.get("md008_auditor") or ""
        sign = info.get("md008_sign") or auditor
        for ti, tb in enumerate(tables):
            if not self.__is_md008_checklist(tb):
                continue
            new_tb = []
            for row in tb:
                if not isinstance(row, list):
                    new_tb.append(row)
                    continue
                next_row = list(row)
                while len(next_row) < 7:
                    next_row.append("")
                a = str(next_row[0] or "").strip()
                if a == "代码地址":
                    next_row[5] = date
                if a == "被审核人":
                    next_row[2] = auditee
                    next_row[5] = auditor
                if "审核人" in a and "签字" in a:
                    next_row[2] = sign
                new_tb.append(next_row)
            tables[ti] = new_tb
        node["tables"] = tables

    @staticmethod
    def __week_ranges_from_dates(dates):
        if not dates:
            return []
        start_d, end_d = min(dates), max(dates)
        if start_d > end_d:
            return []

        def fmt(d):
            return f"{d.year}.{d.month:02d}.{d.day:02d}"

        ranges = []
        cur = start_d - timedelta(days=start_d.weekday())
        while cur <= end_d:
            monday = cur
            friday = monday + timedelta(days=4)
            ws = max(monday, start_d)
            we = min(friday, end_d)
            if ws.weekday() >= 5:
                cur = monday + timedelta(days=7)
                continue
            if we.weekday() >= 5:
                we = friday
            if ws <= we:
                ranges.append(f"{fmt(ws)}- {fmt(we)}")
            cur = monday + timedelta(days=7)
        return ranges

    def __parse_eq_table(self, tb):
        out, seen = [], set()
        if not isinstance(tb, list):
            return out
        hi = brand_i = code_i = name_i = usage_i = -1
        for i, row in enumerate(tb):
            if not isinstance(row, list):
                continue
            cells = [str(c or "").strip() for c in row]
            if any(c == "品牌" for c in cells) and any("资产编码" in c for c in cells):
                hi = i
                brand_i = next(j for j, c in enumerate(cells) if c == "品牌")
                code_i = next(j for j, c in enumerate(cells) if "资产编码" in c)
                name_i = next((j for j, c in enumerate(cells) if c == "名称"), -1)
                usage_i = next((j for j, c in enumerate(cells) if c == "用途"), -1)
                break
        if hi < 0:
            return out
        for row in tb[hi + 1:]:
            if not isinstance(row, list):
                continue
            brand = str(row[brand_i] if brand_i < len(row) else "").strip()
            code = str(row[code_i] if code_i < len(row) else "").strip()
            name = str(row[name_i] if 0 <= name_i < len(row) else "").strip()
            usage = str(row[usage_i] if 0 <= usage_i < len(row) else "").strip()
            if name == "显示器":
                continue
            if brand in ("组装机", "Apple") and code and code not in seen:
                seen.add(code)
                out.append((code, usage))
        return out

    def __parse_eq_codes(self, content):
        if isinstance(content, dict) and isinstance(content.get("rows"), list):
            return self.__parse_eq_table(content.get("rows"))
        out, seen = [], set()

        def walk(ns):
            for n in ns or []:
                if not isinstance(n, dict):
                    continue
                for tb in n.get("tables") or []:
                    if not isinstance(tb, list):
                        continue
                    hi = brand_i = code_i = name_i = -1
                    for i, row in enumerate(tb):
                        if not isinstance(row, list):
                            continue
                        cells = [str(c or "").strip() for c in row]
                        if any(c == "品牌" for c in cells) and any("资产编码" in c for c in cells):
                            hi = i
                            brand_i = next(j for j, c in enumerate(cells) if c == "品牌")
                            code_i = next(j for j, c in enumerate(cells) if "资产编码" in c)
                            name_i = next((j for j, c in enumerate(cells) if c == "名称"), -1)
                            break
                    if hi < 0:
                        continue
                    for row in tb[hi + 1:]:
                        if not isinstance(row, list):
                            continue
                        brand = str(row[brand_i] if brand_i < len(row) else "").strip()
                        code = str(row[code_i] if code_i < len(row) else "").strip()
                        name = str(row[name_i] if 0 <= name_i < len(row) else "").strip()
                        if name == "显示器":
                            continue
                        usage = ""
                        if any("用途" in str(c or "") for c in (tb[hi] if hi >= 0 else [])):
                            usage_i = next((j for j, c in enumerate(tb[hi]) if "用途" in str(c or "")), -1)
                            usage = str(row[usage_i] if 0 <= usage_i < len(row) else "").strip()
                        if brand in ("组装机", "Apple") and code and code not in seen:
                            seen.add(code)
                            out.append((code, usage))
                walk(n.get("children") or [])

        walk((content or {}).get("sections") or [])
        return out

    def __fill_env_maint_node(self, node, info):
        eq_assets = info.get("env_assets")
        prod_name = info.get("prod_name") or ""
        full_version = info.get("full_version") or ""
        tables = node.get("tables") or []
        for ti, tb in enumerate(tables):
            if not isinstance(tb, list) or not self.__is_asset_grid(tb):
                continue
            old = {}
            for row in tb[1:]:
                if isinstance(row, list) and row:
                    code = str(row[0] or "").strip()
                    if code:
                        old[code] = str(row[1] if len(row) > 1 else "")
            hdr = [str(c or "") for c in tb[0]]
            hdr = hdr[:4] if len(hdr) >= 4 else ["资产编码", "设备信息", "产品名称", "完整版本"]
            if eq_assets is not None:
                body = [[code, old.get(code, ""), prod_name, full_version] for code, _u in eq_assets]
                tables[ti] = [hdr] + body
            else:
                new_tb = [hdr]
                for row in tb[1:]:
                    if not isinstance(row, list):
                        continue
                    next_row = list(row)
                    while len(next_row) < 4:
                        next_row.append("")
                    next_row[2] = prod_name
                    next_row[3] = full_version
                    new_tb.append(next_row)
                tables[ti] = new_tb
        node["tables"] = tables

    def __collect_asset_codes(self, content):
        out, seen = [], set()

        def walk(ns):
            for n in ns or []:
                if not isinstance(n, dict):
                    continue
                for tb in n.get("tables") or []:
                    if not self.__is_asset_grid(tb):
                        continue
                    for row in tb[1:]:
                        if not isinstance(row, list) or not row:
                            continue
                        code = str(row[0] or "").strip()
                        if code and code not in seen:
                            seen.add(code)
                            out.append((code, ""))
                walk(n.get("children") or [])

        walk((content or {}).get("sections") or [])
        return out

    def __rebuild_env_checks(self, content, info):
        want = info.get("env_title") or ""
        doc_type = info.get("doc_type") or "md_019"
        weeks = info.get("env_weeks") or []
        checker = info.get("env_checker") or ""
        assets = info.get("env_assets")
        if assets is None:
            assets = self.__collect_asset_codes(content)

        def find_title(ns, name):
            for n in ns or []:
                if not isinstance(n, dict):
                    continue
                if self.__strip_num(n.get("title")) == name:
                    return n
                hit = find_title(n.get("children") or [], name)
                if hit:
                    return hit
            return None

        node = find_title((content or {}).get("sections") or [], want)
        if not node:
            return
        old = {}
        for tb in node.get("tables") or []:
            if not self.__is_env_check_grid(tb):
                continue
            code = str(tb[0][2] if len(tb[0]) > 2 else "")
            by_date = {}
            for row in tb[1:]:
                if isinstance(row, list) and row:
                    by_date[str(row[0] or "")] = [str(c or "") for c in row]
            old[code] = by_date
        tables = []
        for code, usage in assets or []:
            kind = "server" if "共用" in str(usage or "") else "dev"
            defaults = self.__env_check_defaults(doc_type, kind)
            prev = old.get(code) or {}
            rows = [["env_check", kind, code]]
            for w in weeks:
                p = prev.get(w) or []
                marks = []
                for i, d in enumerate(defaults):
                    v = str(p[i + 1] if i + 1 < len(p) else "").strip()
                    marks.append(v if v in ("是", "否") else d)
                problem = str(p[len(defaults) + 1] if len(p) > len(defaults) + 1 else "").strip() or "无"
                rows.append([w] + marks + [problem, checker])
            tables.append(rows)
        node["tables"] = tables

    def __apply_env_eq_assets(self, obj: ModelDocObj, product: Product = None):
        """打开详情时按最新设备清单覆盖资产编码并重建周检（与 DEM get 一致，不落库）。"""
        if obj.doc_type not in ("md_019", "md_020") or not obj.product_id:
            return
        info = self.__collect_autofill(obj.product_id, product, obj.version, obj.doc_type)

        def walk(ns):
            for n in ns or []:
                if isinstance(n, dict):
                    self.__fill_env_maint_node(n, info)
                    walk(n.get("children") or [])

        walk((obj.content or {}).get("sections") or [])
        self.__rebuild_env_checks(obj.content, info)

    def __fill_crr_fields(self, content, obj: ModelDocObj):
        content = self.__normalize_crr_content(content, obj.doc_type)
        if not obj.product_id:
            return content
        product = db.session.execute(select(Product).where(Product.id == obj.product_id)).scalars().first()
        info = self.__collect_autofill(obj.product_id, product, obj.version, obj.doc_type)
        content["check_date"] = info.get("md008_date") or ""
        content["sign_date"] = content["check_date"]
        content["auditee"] = info.get("md008_auditee") or ""
        content["auditor"] = info.get("md008_auditor") or ""
        content["sign_img"] = info.get("md008_sign") or content.get("auditor") or ""
        return content

    def __apply_crr_autofill(self, obj: ModelDocObj, product: Product = None):
        if obj.doc_type not in CRR_DOC_TYPES:
            return
        obj.content = self.__fill_crr_fields(obj.content, obj)

    def __exists(self, product_id, doc_type, version, exclude_id=None):
        sql = select(func.count(ModelDoc.id)).where(
            ModelDoc.product_id == product_id,
            ModelDoc.doc_type == doc_type,
            ModelDoc.version == version,
        )
        if exclude_id:
            sql = sql.where(ModelDoc.id != exclude_id)
        return (db.session.execute(sql).scalar() or 0) > 0

    async def add_model_doc(self, form: ModelDocForm):
        try:
            doc_type = (form.doc_type or "").strip()
            if doc_type not in DOC_META:
                return Resp.resp_err(msg=ts("msg_err_param"))
            if self.__exists(form.product_id, doc_type, form.version):
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            payload = form.dict(exclude_none=True)
            payload["doc_type"] = doc_type
            row = ModelDoc(**payload)
            row.id = None
            row.file_no = serv_review_util.resolve_doc_file_no(form.product_id, form.file_no, form.version, doc_type) or None
            row.content = self.__normalize_content(row.content, doc_type)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=ModelDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_model_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: ModelDoc = db.session.execute(select(ModelDoc).where(ModelDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(
                select(ModelDoc.version).where(ModelDoc.product_id == target_pid, ModelDoc.doc_type == fromdoc.doc_type)
            ).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = ModelDoc(
                product_id=target_pid,
                doc_type=fromdoc.doc_type,
                version=version,
                file_no=sync_file_no_version(
                    (fromdoc.file_no or "").strip()
                    or serv_review_util.resolve_doc_file_no(target_pid, "", version, fromdoc.doc_type)
                    or "",
                    version,
                ) or None,
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content, fromdoc.doc_type)),
            )
            db.session.add(newdoc)
            db.session.commit()
            return Resp.resp_ok(data=ModelDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_model_doc(self, form: ModelDocForm):
        try:
            row: ModelDoc = db.session.execute(select(ModelDoc).where(ModelDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            payload = form.dict(exclude_none=True)
            payload.pop("doc_type", None)
            next_pid = payload.get("product_id", row.product_id)
            next_ver = payload.get("version", row.version)
            if next_pid != row.product_id or next_ver != row.version:
                if self.__exists(next_pid, row.doc_type, next_ver, exclude_id=row.id):
                    return Resp.resp_err(msg=ts("msg_obj_exist"))
            for key, value in payload.items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value, row.doc_type)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_model_doc(self, id: int):
        db.session.execute(delete(ModelDoc).where(ModelDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_model_doc(self, id: int):
        sql = select(ModelDoc, Product).join(Product, ModelDoc.product_id == Product.id).where(ModelDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        obj = self.__to_obj(doc, product)
        self.__apply_env_eq_assets(obj, product)
        self.__apply_crr_autofill(obj, product)
        return Resp.resp_ok(data=obj)

    async def list_model_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None,
                             doc_type: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if doc_type:
            wheres.append(ModelDoc.doc_type == doc_type)
        if product_id:
            wheres.append(ModelDoc.product_id == product_id)
        if version:
            wheres.append(ModelDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(ModelDoc.id)).join(Product, ModelDoc.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (
            select(ModelDoc, Product)
            .join(Product, ModelDoc.product_id == Product.id)
            .where(*wheres)
            .order_by(ModelDoc.id.desc())
            .offset(page_index * page_size)
            .limit(page_size)
        )
        rows: List[ModelDocObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    def __export_xlsx(self, output, obj: ModelDocObj, content):
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        default = wb.active
        if obj.doc_type in EQ_DOC_TYPES or (isinstance(content, dict) and isinstance(content.get("rows"), list)):
            thin_eq = Border(
                left=Side(style="thin"), right=Side(style="thin"),
                top=Side(style="thin"), bottom=Side(style="thin"),
            )
            default.title = (doc_title(obj.doc_type) or "设备清单")[:31]
            for r_i, row in enumerate(content.get("rows") or [], 1):
                if not isinstance(row, list):
                    continue
                for c_i, val in enumerate(row, 1):
                    cell = default.cell(r_i, c_i, str(val or ""))
                    cell.alignment = Alignment(wrap_text=True, vertical="center")
                    cell.border = thin_eq
                    cell.font = Font(bold=(r_i == 1), name="宋体")
            wb.save(output)
            output.seek(0)
            return
        used_names = set()
        thin = Border(
            left=Side(style="thin"), right=Side(style="thin"),
            top=Side(style="thin"), bottom=Side(style="thin"),
        )

        def sheet_name(title):
            raw = self.__strip_num(title) or "Sheet"
            name = re.sub(r'[:\\/\?\*\[\]]', "_", raw)[:31] or "Sheet"
            base = name
            idx = 2
            while name.lower() in used_names:
                suffix = str(idx)
                name = (base[: 31 - len(suffix)] + suffix)
                idx += 1
            used_names.add(name.lower())
            return name

        def write_table(ws, grid, start_row=1):
            if self.__is_md008_checklist(grid):
                rows, origins = self.__md008_span_origins(grid)
                r_n = len(rows)
                c_n = max((len(r) for r in rows), default=7)
                for r_i, row in enumerate(rows):
                    for c_i in range(c_n):
                        raw = row[c_i] if c_i < len(row) else ""
                        s = str(raw or "")
                        if s.startswith("data:image"):
                            s = "[签名]"
                        cell = ws.cell(start_row + r_i, c_i + 1, s)
                        cell.alignment = Alignment(wrap_text=True, vertical="center")
                        cell.border = thin
                        cell.font = Font(name="宋体")
                for r, c, rs, cs in origins:
                    if rs > 1 or cs > 1:
                        ws.merge_cells(
                            start_row=start_row + r, start_column=c + 1,
                            end_row=start_row + r + rs - 1, end_column=c + cs,
                        )
                widths = (10.5, 11.4, 41.7, 6.9, 6.5, 6.9, 13.6)
                for i, w in enumerate(widths, 1):
                    ws.column_dimensions[get_column_letter(i)].width = w
                return start_row + r_n
            r_idx = start_row
            for r_i, row in enumerate(grid or []):
                if not isinstance(row, list):
                    continue
                for c_i, val in enumerate(row, 1):
                    s = str(val or "")
                    if s.startswith("data:image"):
                        s = "[签名]"
                    cell = ws.cell(r_idx, c_i, s)
                    cell.alignment = Alignment(wrap_text=True, vertical="center")
                    cell.border = thin
                    if r_i == 0:
                        cell.font = Font(bold=True, name="宋体")
                    else:
                        cell.font = Font(name="宋体")
                r_idx += 1
            return r_idx

        first = True
        for node in (content or {}).get("sections") or []:
            ws = default if first else wb.create_sheet()
            first = False
            ws.title = sheet_name(node.get("title"))
            title = self.__strip_num(node.get("title"))
            tables = node.get("tables") or []
            crr_sheet = any(self.__is_md008_checklist(t) for t in tables)
            if crr_sheet:
                row = 1
                for table in tables:
                    row = write_table(ws, table, row) + 1
            else:
                ws.cell(1, 1, title).font = Font(bold=True, size=14, name="宋体")
                row = 3
                if (node.get("body") or "").strip():
                    ws.cell(row, 1, node.get("body"))
                    row += 2
                for table in tables:
                    row = write_table(ws, table, row) + 2
            for child in (node.get("children") or []):
                ws.cell(row, 1, self.__strip_num(child.get("title"))).font = Font(bold=True, name="宋体")
                row += 1
                if (child.get("body") or "").strip():
                    ws.cell(row, 1, child.get("body"))
                    row += 1
                for table in (child.get("tables") or []):
                    row = write_table(ws, table, row) + 2
        if first:
            default.title = "模型文件"
        wb.save(output)
        output.seek(0)

    def __export_crr_docx(self, output, obj: ModelDocObj, content):
        c = content if isinstance(content, dict) else {}
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")
        docx_util.add_page_number_footer(section, obj.file_no or "", skip_first=False)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            s = str(text or "")
            if s.startswith("data:image"):
                try:
                    b64 = s.split(",", 1)[1] if "," in s else ""
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = align
                    para.add_run().add_picture(BytesIO(base64.b64decode(b64)), height=Pt(30))
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    return
                except Exception:
                    pass
            cell.text = ""
            for i, line in enumerate(s.split("\n")):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        def set_check(cell, checked):
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run("\u2611\ufe0e" if checked else "\u2610")
            run.font.size = Pt(12)
            run.font.name = "宋体"
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            for _attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                rfonts.set(qn(_attr), "宋体")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        docx_util.fonted_txt(title, doc_title(obj.doc_type), font_size=18.0, bold=True)

        head = document.add_table(rows=0, cols=4)
        head.style = "Table Grid"
        head.alignment = WD_TABLE_ALIGNMENT.CENTER

        def head_row(label1, val1, label2="", val2="", merge_val=False):
            cells = head.add_row().cells
            set_cell(cells[0], label1, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            if merge_val:
                merged = cells[1].merge(cells[2]).merge(cells[3])
                set_cell(merged, val1)
            else:
                set_cell(cells[1], val1)
                set_cell(cells[2], label2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell(cells[3], val2)

        head_row("代码地址", str(c.get("code_url", "") or "").replace("\t", " "), "检查日期", c.get("check_date", ""))
        head_row("被审核人", c.get("auditee", ""), "审核人", c.get("auditor", ""))
        head_row("审核依据", c.get("basis", "") or "", merge_val=True)
        head_row("审核方式", c.get("method", "") or "代码审查", merge_val=True)
        document.add_paragraph()

        checklist = [r for r in (c.get("checklist") or []) if isinstance(r, list)]
        cols = 6
        tbl = document.add_table(rows=0, cols=cols)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        col_dxa = [720, 4680, 720, 720, 1008, 2160]
        cat_flags = []
        for r_idx, row in enumerate(checklist):
            cells = tbl.add_row().cells
            first = str(row[0] if row else "").strip()
            is_cat = first in _CRR_CATEGORIES and all(not str(row[i] if i < len(row) else "").strip() for i in range(1, cols))
            cat_flags.append(is_cat)
            if is_cat:
                merged = cells[0]
                for i in range(1, cols):
                    merged = merged.merge(cells[i])
                set_cell(merged, first, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
                tc_pr = merged._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "FFFF99")
                tc_pr.append(shd)
            else:
                for c_idx in range(cols):
                    val = row[c_idx] if c_idx < len(row) else ""
                    if r_idx != 0 and c_idx in (2, 3, 4):
                        set_check(cells[c_idx], bool(str(val).strip()))
                        continue
                    align = WD_ALIGN_PARAGRAPH.CENTER if c_idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
                    set_cell(cells[c_idx], val, bold=(r_idx == 0), align=align)
        tbl.autofit = False
        _tblPr = tbl._tbl.tblPr
        _layout = _tblPr.find(qn("w:tblLayout"))
        if _layout is None:
            _layout = OxmlElement("w:tblLayout")
            _tblPr.append(_layout)
        _layout.set(qn("w:type"), "fixed")
        _grid = tbl._tbl.find(qn("w:tblGrid"))
        if _grid is not None:
            for _gc in list(_grid):
                _grid.remove(_gc)
            for _w in col_dxa:
                _gc = OxmlElement("w:gridCol")
                _gc.set(qn("w:w"), str(_w))
                _grid.append(_gc)
        for _ri, _r in enumerate(tbl.rows):
            if _ri < len(cat_flags) and cat_flags[_ri]:
                continue
            _cells = _r.cells
            for _i, _w in enumerate(col_dxa):
                if _i < len(_cells):
                    _tcpr = _cells[_i]._tc.get_or_add_tcPr()
                    _tcw = _tcpr.find(qn("w:tcW"))
                    if _tcw is None:
                        _tcw = OxmlElement("w:tcW")
                        _tcpr.append(_tcw)
                    _tcw.set(qn("w:w"), str(_w))
                    _tcw.set(qn("w:type"), "dxa")
        document.add_paragraph()

        concl = str(c.get("conclusion") or "").strip()
        concl_tbl = document.add_table(rows=1, cols=1)
        concl_tbl.style = "Table Grid"
        concl_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        marks = " ".join([("%s%s" % ("\u2611\ufe0e" if concl == name else "\u2610", name)) for name in _CRR_CONCLUSIONS])
        set_cell(concl_tbl.rows[0].cells[0], "结论： " + marks)
        document.add_paragraph()

        sign_tbl = document.add_table(rows=1, cols=3)
        sign_tbl.style = "Table Grid"
        sign_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        srow = sign_tbl.rows[0].cells
        sign_tbl.rows[0].height = Pt(52)
        sign_tbl.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        set_cell(srow[0], "审核人（签字）/日期", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(srow[1], c.get("sign_img", "") or "", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(srow[2], c.get("sign_date", "") or c.get("check_date", "") or "", align=WD_ALIGN_PARAGRAPH.CENTER)

        document.save(output)
        output.seek(0)

    def __export_build_docx(self, output, obj: ModelDocObj, content):
        c = content if isinstance(content, dict) else {}
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")
        docx_util.add_page_number_footer(section, obj.file_no or "", skip_first=False)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            s = str(text or "")
            if s.startswith("data:image"):
                try:
                    b64 = s.split(",", 1)[1] if "," in s else ""
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = align
                    para.add_run().add_picture(BytesIO(base64.b64decode(b64)), height=Pt(30))
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    return
                except Exception:
                    pass
            cell.text = ""
            for i, line in enumerate(s.split("\n")):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        docx_util.fonted_txt(title, doc_title(obj.doc_type), font_size=18.0, bold=True)

        tbl = document.add_table(rows=0, cols=4)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        def add_row():
            return tbl.add_row().cells

        def pair(l1, v1, l2, v2):
            cells = add_row()
            set_cell(cells[0], l1, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell(cells[1], v1, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell(cells[2], l2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell(cells[3], v2, align=WD_ALIGN_PARAGRAPH.CENTER)

        pair("编写人", c.get("author") or "", "编写时间", c.get("write_date") or "")
        pair("数据用途", c.get("data_use") or "", "数据类型", c.get("data_type") or "")
        cells = add_row()
        set_cell(cells[0], "构建方法", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        merged = cells[1].merge(cells[2]).merge(cells[3])
        set_cell(merged, c.get("method") or "")
        pair("病例数量", c.get("case_count") or "", "标记人员及方式", c.get("annotator") or "")
        cells = add_row()
        bar = cells[0].merge(cells[1]).merge(cells[2]).merge(cells[3])
        set_cell(bar, "数据分布", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        dist = [self.__build_pad_row(r) for r in (c.get("dist_rows") or []) if isinstance(r, list)]
        if not dist:
            dist = [["因素", "类别", "数量", "占比"]]
        dist_start = len(tbl.rows)
        for r_idx, row in enumerate(dist):
            cells = add_row()
            is_head = r_idx == 0
            is_total = str(row[0]).strip() == "总计"
            for ci in range(4):
                set_cell(
                    cells[ci], row[ci],
                    bold=(is_head or is_total),
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )
        i = 1
        while i < len(dist):
            a = str(dist[i][0] or "").strip()
            if not a or a == "总计":
                i += 1
                continue
            j = i + 1
            while j < len(dist) and not str(dist[j][0] or "").strip():
                j += 1
            if j > i + 1:
                start_cell = tbl.cell(dist_start + i, 0)
                end_cell = tbl.cell(dist_start + j - 1, 0)
                start_cell.merge(end_cell)
                set_cell(start_cell, a, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            i = j

        cells = add_row()
        tbl.rows[-1].height = Pt(44)
        tbl.rows[-1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        set_cell(cells[0], "编写人签字（日期）", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(cells[1], c.get("author_sign") or "", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(cells[2], "审核人签字（日期）", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(cells[3], c.get("auditor_sign") or "", align=WD_ALIGN_PARAGRAPH.CENTER)

        col_dxa = [1600, 2800, 1600, 2800]
        tbl.autofit = False
        _tblPr = tbl._tbl.tblPr
        _layout = _tblPr.find(qn("w:tblLayout"))
        if _layout is None:
            _layout = OxmlElement("w:tblLayout")
            _tblPr.append(_layout)
        _layout.set(qn("w:type"), "fixed")
        _grid = tbl._tbl.find(qn("w:tblGrid"))
        if _grid is not None:
            for _gc in list(_grid):
                _grid.remove(_gc)
            for _w in col_dxa:
                _gc = OxmlElement("w:gridCol")
                _gc.set(qn("w:w"), str(_w))
                _grid.append(_gc)

        document.save(output)
        output.seek(0)

    def __export_train_docx(self, output, obj: ModelDocObj, content):
        c = content if isinstance(content, dict) else {}
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")
        docx_util.add_page_number_footer(section, obj.file_no or "", skip_first=False)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, img_w=None):
            s = str(text or "")
            if s.startswith("data:image"):
                try:
                    b64 = s.split(",", 1)[1] if "," in s else ""
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.add_run().add_picture(BytesIO(base64.b64decode(b64)), width=Inches(img_w or 6.2))
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    return
                except Exception:
                    pass
            cell.text = ""
            for i, line in enumerate(s.split("\n")):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        docx_util.fonted_txt(title, doc_title(obj.doc_type), font_size=18.0, bold=True)

        tbl = document.add_table(rows=0, cols=6)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        def add_row():
            return tbl.add_row().cells

        def label_cell(cell, text):
            set_cell(cell, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        cells = add_row()
        label_cell(cells[0], "编写人")
        set_cell(cells[1], c.get("author") or "", align=WD_ALIGN_PARAGRAPH.CENTER)
        label_cell(cells[2], "编写日期")
        set_cell(cells[3], c.get("write_date") or "", align=WD_ALIGN_PARAGRAPH.CENTER)
        label_cell(cells[4], "审核人")
        set_cell(cells[5], c.get("auditor") or "", align=WD_ALIGN_PARAGRAPH.CENTER)

        def span_val(label, value):
            row = add_row()
            label_cell(row[0], label)
            merged = row[1].merge(row[2]).merge(row[3]).merge(row[4]).merge(row[5])
            set_cell(merged, value)

        span_val("模型名称", c.get("model_name") or "")
        span_val("模型功能", c.get("model_func") or "")
        cells = add_row()
        label_cell(cells[0], "训练集")
        set_cell(cells[1].merge(cells[2]), c.get("train_set") or "", align=WD_ALIGN_PARAGRAPH.CENTER)
        label_cell(cells[3], "数量")
        set_cell(cells[4].merge(cells[5]), c.get("case_count") or "", align=WD_ALIGN_PARAGRAPH.CENTER)
        span_val("训练时间", c.get("train_time") or "")

        def bar(text):
            row = add_row()
            merged = row[0].merge(row[1]).merge(row[2]).merge(row[3]).merge(row[4]).merge(row[5])
            set_cell(merged, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        bar("训练环境")
        span_val("硬件环境", c.get("hw_env") or "")
        span_val("软件环境", c.get("sw_env") or "")
        bar("训练数据量评估曲线")
        cells = add_row()
        tbl.rows[-1].height = Pt(220)
        tbl.rows[-1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        pic = cells[0].merge(cells[1]).merge(cells[2]).merge(cells[3]).merge(cells[4]).merge(cells[5])
        set_cell(pic, c.get("eval_img") or "", img_w=6.2)
        bar("训练过程曲线")
        cells = add_row()
        tbl.rows[-1].height = Pt(180)
        tbl.rows[-1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        pic = cells[0].merge(cells[1]).merge(cells[2]).merge(cells[3]).merge(cells[4]).merge(cells[5])
        set_cell(pic, c.get("process_img") or "", img_w=6.2)
        span_val("结论", c.get("conclusion") or "")
        cells = add_row()
        tbl.rows[-1].height = Pt(44)
        tbl.rows[-1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        label_cell(cells[0], "编写人（签字）/日期")
        set_cell(cells[1].merge(cells[2]), c.get("author_sign") or "", align=WD_ALIGN_PARAGRAPH.CENTER)
        label_cell(cells[3], "审核人（签字）/日期")
        set_cell(cells[4].merge(cells[5]), c.get("auditor_sign") or "", align=WD_ALIGN_PARAGRAPH.CENTER)

        document.save(output)
        output.seek(0)

    async def export_model_doc(self, output, id: int):
        resp = await self.get_model_doc(id)
        obj: ModelDocObj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return "模型文件", "docx"
        c = self.__autofill_for_export(self.__normalize_content(obj.content, obj.doc_type), obj)
        title = doc_title(obj.doc_type)
        if obj.doc_type in CRR_DOC_TYPES:
            self.__export_crr_docx(output, obj, c)
            return title, "docx"
        if obj.doc_type in BUILD_DOC_TYPES:
            self.__export_build_docx(output, obj, c)
            return title, "docx"
        if obj.doc_type in TRAIN_DOC_TYPES:
            self.__export_train_docx(output, obj, c)
            return title, "docx"
        if doc_format(obj.doc_type) == "xlsx":
            self.__export_xlsx(output, obj, c)
            return title, "xlsx"
        sections = c.get("sections") or []
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        document.settings.element.append(update_fields)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")
        docx_util.add_page_number_footer(section, obj.file_no or "")

        def write_center_title(text, size=22.0, bold=False):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.right_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, text, font_size=size, bold=bold)

        def add_blank_lines(count):
            for _ in range(max(0, int(count or 0))):
                document.add_paragraph("")

        def add_text(text):
            docx_util.save_txt2docx(str(text or ""), document)

        def png_wh(raw):
            if raw[:8] == b"\x89PNG\r\n\x1a\n" and len(raw) >= 24:
                return int.from_bytes(raw[16:20], "big"), int.from_bytes(raw[20:24], "big")
            return 0, 0

        def add_picture_fit(run, raw, max_w, max_h):
            w, h = png_wh(raw)
            if w > 0 and h > 0 and (h * max_w) > (w * max_h):
                run.add_picture(BytesIO(raw), height=Inches(max_h))
            else:
                run.add_picture(BytesIO(raw), width=Inches(max_w))

        def is_figure_grid(grid):
            cols = max((len(row) for row in grid if isinstance(row, list)), default=0)
            if cols != 1:
                return False
            return any(
                isinstance(row, list) and row and str(row[0] or "").startswith("data:image")
                for row in grid
            )

        def add_figure_grid(grid):
            for row in grid:
                val = str(row[0] if row else "")
                if val.startswith("data:image"):
                    try:
                        raw = base64.b64decode(val.split(",", 1)[1] if "," in val else "")
                        p = document.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        add_picture_fit(p.add_run(), raw, 5.5, 3.6)
                    except Exception:
                        logger.exception("md007_export_figure_failed")
                elif val.strip():
                    p = document.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    docx_util.fonted_txt(p, val, font_size=10.5)
            document.add_paragraph()

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, figure=False):
            s = str(text or "")
            if s.startswith("data:image"):
                try:
                    raw = base64.b64decode(s.split(",", 1)[1] if "," in s else "")
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = align
                    if figure:
                        add_picture_fit(para.add_run(), raw, 4.2, 2.6)
                    else:
                        para.add_run().add_picture(BytesIO(raw), height=Pt(33))
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    return
                except Exception:
                    pass
            cell.text = ""
            lines = s.split("\n")
            for i, line in enumerate(lines):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER if align == WD_ALIGN_PARAGRAPH.CENTER else WD_CELL_VERTICAL_ALIGNMENT.TOP

        def set_yesno(cell, mark):
            cell.text = ""
            yes = str(mark or "").strip() == "是"
            no = str(mark or "").strip() == "否"
            p1 = cell.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p1.add_run(("\u2611\ufe0e" if yes else "\u2610") + " 是")
            r1.font.size = Pt(9)
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(("\u2610" if not no else "\u2611\ufe0e") + " 否")
            r2.font.size = Pt(9)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        def add_env_check_grid(grid):
            kind = str(grid[0][1] if len(grid[0]) > 1 else "dev") or "dev"
            code = str(grid[0][2] if len(grid[0]) > 2 else "")
            doc_t = obj.doc_type or "md_019"
            leaves = self.__env_check_leaves(doc_t, kind)
            groups = ENV_CHECK_GROUPS.get((doc_t, kind), ENV_CHECK_GROUPS[("md_019", "dev")])
            ncols = len(leaves)
            if doc_t == "md_020":
                title_txt = "测试共用-%s检查表（%s）" % ("服务器" if kind == "server" else "测试机", code)
            else:
                title_txt = "开发共用-%s检查表（%s）" % ("服务器" if kind == "server" else "开发机", code)
            tb = document.add_table(rows=0, cols=ncols)
            tb.style = "Table Grid"
            tb.alignment = WD_TABLE_ALIGNMENT.CENTER
            trow = tb.add_row().cells
            tmerge = trow[0]
            for i in range(1, ncols):
                tmerge = tmerge.merge(trow[i])
            set_cell(tmerge, title_txt, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            grow = tb.add_row().cells
            lrow = tb.add_row().cells
            ci = 0
            for gl, gleaves in groups:
                if gleaves:
                    gm = grow[ci]
                    for k in range(1, len(gleaves)):
                        gm = gm.merge(grow[ci + k])
                    set_cell(gm, gl, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                    for k, lf in enumerate(gleaves):
                        set_cell(lrow[ci + k], lf, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                    ci += len(gleaves)
                else:
                    vm = grow[ci].merge(lrow[ci])
                    set_cell(vm, gl, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                    ci += 1
            for row in grid[1:]:
                cells = tb.add_row().cells
                j = 0
                for idx, col in enumerate(leaves):
                    t = col["type"]
                    if t == "date":
                        set_cell(cells[idx], str(row[0] if row else "").replace("- ", "-\n"), align=WD_ALIGN_PARAGRAPH.CENTER)
                    elif t == "check":
                        set_yesno(cells[idx], row[j + 1] if j + 1 < len(row) else "")
                        j += 1
                    elif t == "problem":
                        set_cell(cells[idx], row[j + 1] if j + 1 < len(row) else "", align=WD_ALIGN_PARAGRAPH.CENTER)
                    elif t == "checker":
                        set_cell(cells[idx], row[j + 2] if j + 2 < len(row) else "", align=WD_ALIGN_PARAGRAPH.CENTER)
            document.add_paragraph()

        def add_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            if self.__is_env_check_grid(grid):
                add_env_check_grid(grid)
                return
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            if is_figure_grid(grid):
                add_figure_grid(grid)
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            if self.__is_review_grid(grid):
                padded, origins = self.__grid_span_origins(grid)
                for _ in padded:
                    table.add_row()
                for r, c, rs, cs in origins:
                    cell = table.cell(r, c)
                    if rs > 1 or cs > 1:
                        cell = cell.merge(table.cell(r + rs - 1, c + cs - 1))
                    set_cell(cell, padded[r][c], bold=(r == 0))
            else:
                for r_idx, row in enumerate(grid):
                    cells = table.add_row().cells
                    left = str(row[0] if row else "")
                    for c_idx in range(cols):
                        set_cell(
                            cells[c_idx],
                            row[c_idx] if c_idx < len(row) else "",
                            bold=(r_idx == 0),
                            figure=("算法流程图" in left and c_idx > 0),
                        )
            document.add_paragraph()

        def add_cover_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for row in grid:
                cells = table.add_row().cells
                for c_idx in range(cols):
                    text = row[c_idx] if c_idx < len(row) else ""
                    set_cell(cells[c_idx], text, bold=(c_idx % 2 == 0), align=WD_ALIGN_PARAGRAPH.CENTER)
                if (str(row[0]).strip() if row else "") == "生效日期" and cols > 2:
                    merged = cells[1]
                    for c_idx in range(2, cols):
                        merged = merged.merge(cells[c_idx])
                    set_cell(merged, row[1] if len(row) > 1 else "", align=WD_ALIGN_PARAGRAPH.CENTER)
            document.add_paragraph()

        def add_body_heading(title, level):
            size = {1: 16.0, 2: 14.0, 3: 12.0}.get(level, 11.0)
            p = document.add_heading("", level=max(1, min(level, 9)))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, title, font_size=size, bold=True)

        def render_body_section(node, level, number=""):
            name = self.__strip_num(node.get("title"))
            heading = f"{number} {name}".strip() if number else name
            add_body_heading(heading, level=max(1, min(level, 9)))
            if (node.get("body") or "").strip():
                add_text(node.get("body"))
            for table in (node.get("tables") or []):
                add_grid(table)
            idx = 0
            for child in (node.get("children") or []):
                idx += 1
                child_num = f"{number}.{idx}" if number else f"{idx}"
                render_body_section(child, level + 1, child_num)

        cover = next((s for s in sections if s.get("ref_type") == "cover"), None)
        revision = next((s for s in sections if s.get("ref_type") == "revision"), None)
        body = [s for s in sections if s.get("ref_type") not in ("cover", "revision")]
        title = (self.__strip_num(cover.get("title")) if cover else "") or doc_title(obj.doc_type)

        add_blank_lines(6)
        write_center_title(title, size=22.0, bold=True)
        add_blank_lines(4)
        if cover:
            for table in (cover.get("tables") or []):
                add_cover_grid(table)

        document.add_page_break()
        write_center_title("文件修订记录", size=14.0, bold=True)
        add_blank_lines(2)
        if revision:
            for table in (revision.get("tables") or []):
                add_grid(table)

        document.add_page_break()
        write_center_title("目录", size=16.0, bold=True)
        docx_util.insert_toc_field(document)

        document.add_page_break()
        if obj.doc_type in SKIP_ANNEX_NUM:
            idx = 0
            for node in body:
                t = self.__strip_num(node.get("title"))
                if node.get("ref_type") == "basic_info" or t == "产品信息":
                    render_body_section(node, 1, "")
                    continue
                if t.startswith("附件"):
                    render_body_section(node, 1, "")
                    continue
                idx += 1
                render_body_section(node, 1, str(idx))
        else:
            for i, node in enumerate(body):
                render_body_section(node, 1, str(i + 1))

        docx_util.fill_toc_cache(document)
        document.save(output)
        output.seek(0)
        return title, "docx"
