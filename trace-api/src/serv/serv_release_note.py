#!/usr/bin/env python
# encoding: utf-8

# 产品发布说明服务层。整份文档以 content(JSON) 存储；导出复用 docx_util.fonted_txt 生成 Word。
# 自动获取：全文产品名称、产品概述(总体描述)、文件编号(产品DHF)、发布时间(时间逻辑线)。

import copy
import logging
import re
from typing import List
from sqlalchemy import delete, func, select
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.prod_dhf import ProdDhf
from ..model.release_note import ReleaseNote
from ..model.project_timeline import ProjectTimelineRow, ProjectTimelineCell
from ..model.project_member import ProjectMember
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_release_note import ReleaseNoteForm
from ..obj.vobj_release_note import ReleaseNoteObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from .serv_utils import new_version
from .serv_utils import docx_util

logger = logging.getLogger(__name__)

# 本文档名（用于从产品 DHF 匹配文件编号 / 从时间逻辑线匹配发布时间）
DOC_NAME = "产品发布说明"

PURPOSE_TEXT = (
    "确保所有设计开发活动和任务连同所有相关文件是完整的，并且可被交付给相应的使用人员。"
    "按照《产品发布管理制度》将验收通过的软件产品、配置项和文档归档。"
)
ACCEPT_TEXT = (
    "为保证产品符合需求，应对产品进行验收，详见《产品验收记录》。\n"
    "研发人员完成从研发到生产母盘（全新储存媒介）拷贝产品的活动，"
    "生产人员完成将母盘中产品拷贝至生产电脑指定位置的活动。"
)
TRANSFER_NOTE = "注：其他DMR文档已作为通用技术文件受控和发放。"


def _archive_text(name):
    name = str(name or "").strip()
    return f"通过评审，{name}的全套设计开发历史文档（DHF）和全套器械主记录（DMR）已完成归档。"


def _overview_text(name, release_version, full_version, func_desc):
    """组装产品概述：产品名称/发布版本/完整版本(自动) + 交付形式/存储媒介(固定) + 产品功能概述(总体描述)。"""
    func_desc = str(func_desc or "").strip()
    lines = [
        f"产品名称：{str(name or '').strip()}",
        f"发布版本：{str(release_version or '').strip()}",
        f"完整版本：{str(full_version or '').strip()}",
        "交付形式：物理交付",
        "存储媒介：U盘",
        "产品功能概述：",
    ]
    if func_desc:
        lines.append(func_desc)
    return "\n".join(lines)


DEFAULT_RELEASE_NOTE_CONTENT = {
    "sections": [
        {
            "title": DOC_NAME, "ref_type": "cover", "body": "", "children": [],
            "tables": [[
                ["编制部门", "产品部", "文件版本", "A0"],
                ["编制人", "", "日期", ""],
                ["审核人", "", "日期", ""],
                ["批准人", "", "日期", ""],
                ["生效日期", "", "", ""],
            ]],
        },
        {
            "title": "文件修订记录", "ref_type": "revision", "body": "", "children": [],
            "tables": [[
                ["修改日期", "版本号", "修订说明", "修订人", "批准人"],
                ["", "", "首次发布", "", ""],
                ["", "", "", "", ""],
                ["", "", "", "", ""],
                ["", "", "", "", ""],
                ["", "", "", "", ""],
            ]],
        },
        {"title": "目的", "body": PURPOSE_TEXT, "tables": [], "children": []},
        {
            "title": "产品概述", "ref_type": "rn_overview", "tables": [], "children": [],
            "body": _overview_text(
                "肿瘤CT图像随访与评估软件", "2", "2.0.0.0",
                "肿瘤CT图像随访与评估软件主要包括登录、数据上传、系统管理、时间线、工作站和报告六个模块。",
            ),
        },
        {
            "title": "发布活动", "body": "", "tables": [], "children": [
                {"title": "发布时间", "ref_type": "rn_release_time", "body": "", "tables": [], "children": []},
                {"title": "产品验收和交付", "body": ACCEPT_TEXT, "tables": [], "children": []},
                {"title": "文档归档", "ref_type": "rn_archive", "body": _archive_text("肿瘤CT图像随访与评估软件"), "tables": [], "children": []},
                {
                    "title": "产品和文档移交记录", "body": "", "tables": [], "children": [
                        {
                            "title": "产品移交记录", "ref_type": "rn_transfer_product", "body": "", "children": [],
                            "tables": [[
                                ["产品交付", "安装包名称", "InferCare_RECIST-2.0.0.0.zip"],
                                ["产品交付", "说明书名称", "TX-TF-RCN3V2000-PD-006-A0 用户说明书"],
                                ["产品交付", "交付方式", "U盘交付"],
                                ["产品交付", "交付时间", "2025年10月17日"],
                                ["产品交付", "交付过程说明", "研发人员将装有上述安装包的母盘交付给生产人员。\n生产人员将母盘中的安装包拷贝至生产计算机中“生产专用”文件夹。\n由生产人员进行生产。\n由授权人员将产品交付给用户并进行安装（U盘安装或云部署）。"],
                                ["母盘交付记录", "研发人员/日期：", ""],
                                ["母盘交付记录", "生产人员/日期：", ""],
                            ]],
                        },
                        {
                            "title": "文件移交记录", "ref_type": "rn_transfer_files", "body": TRANSFER_NOTE, "children": [],
                            "tables": [[
                                ["文件编号", "文件名称", "交付人/日期", "接收人/日期"],
                                ["TX-TF-RCN3V2000-RD-008-A0", "产品标签样稿", "", ""],
                                ["TX-TF-RCN3V2000-PD-009-A0", "成品检验规程", "", ""],
                                ["TX-TF-RCN3V2000-PD-009-QR-01-A0", "成品检验记录", "", ""],
                                ["TX-TF-RCN3V2000-PD-006-A0", "用户说明书", "", ""],
                                ["TX-TF-RCN3V2000-VV-005-A0", "安装维护手册", "", ""],
                                ["TX-TF-RCN3V2000-VV-006-QR-01-A0", "现场测试记录", "", ""],
                                ["TX-TF-RCN3V2000-VV-006-A0", "现场测试规程", "", ""],
                            ]],
                        },
                    ],
                },
            ],
        },
    ],
}


class Server(object):

    def __normalize_node(self, node):
        if not isinstance(node, dict):
            return {"title": str(node or ""), "body": "", "tables": [], "children": []}
        result = dict(node)
        result["title"] = str(result.get("title") or "")
        result["body"] = str(result.get("body") or "")
        tables = result.get("tables")
        if not isinstance(tables, list):
            tables = []
        norm_tables = []
        for table in tables:
            if isinstance(table, list):
                norm_tables.append([[str(c) if c is not None else "" for c in (row or [])] for row in table if isinstance(row, list)])
        result["tables"] = norm_tables
        children = result.get("children")
        result["children"] = [self.__normalize_node(c) for c in children] if isinstance(children, list) else []
        return result

    def __normalize_content(self, content):
        if not isinstance(content, dict) or not isinstance(content.get("sections"), list):
            return copy.deepcopy(DEFAULT_RELEASE_NOTE_CONTENT)
        return {"sections": [self.__normalize_node(s) for s in content["sections"]]}

    @staticmethod
    def __strip_num(title):
        return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()

    # ---------------- 自动获取 ----------------
    def __collect_autofill(self, prod_id, product, doc_version):
        name = (getattr(product, "name", "") or "").strip()
        overall_desc = (getattr(product, "overall_desc", "") or "").strip()
        release_version = (getattr(product, "release_version", "") or "").strip()
        full_version = (getattr(product, "full_version", "") or "").strip()

        def to_int(v):
            digits = re.sub(r"[^\d]", "", str(v or ""))
            return int(digits) if digits else None

        tl_rows = db.session.execute(
            select(ProjectTimelineRow).where(ProjectTimelineRow.prod_id == prod_id)
        ).scalars().all()
        cell_map = {}
        if tl_rows:
            for c in db.session.execute(
                select(ProjectTimelineCell).where(ProjectTimelineCell.row_id.in_([r.id for r in tl_rows]))
            ).scalars().all():
                cell_map.setdefault(c.row_id, []).append(c.output_result or "")
        date_rows = [r for r in tl_rows if (r.row_type or "date") == "date" and to_int(r.year) and to_int(r.month)]

        def date_key(r):
            return to_int(r.year) * 10000 + to_int(r.month) * 100 + (to_int(r.day) or 0)

        def latest_date(keyword):
            rows = [r for r in date_rows if any(keyword in str(v or "") for v in cell_map.get(r.id, []))]
            if not rows:
                return ""
            r = max(rows, key=date_key)
            return f"{to_int(r.year)}年{to_int(r.month)}月{to_int(r.day)}日"

        release_date = latest_date("发布说明")
        acceptance_date = latest_date("产品验收记录")

        members = db.session.execute(select(ProjectMember).where(ProjectMember.prod_id == prod_id)).scalars().all()

        def find_member(pred):
            for mem in members:
                if pred(str(mem.role or "")):
                    return (mem.name or "").strip()
            return ""

        pm = find_member(lambda r: "产品经理" in r)
        approver = find_member(lambda r: "负责人" in r and "产品" in r)

        dhf_map = {}
        for d in db.session.execute(select(ProdDhf).where(ProdDhf.prod_id == prod_id)).scalars().all():
            nm = (d.name or "").strip()
            if nm and d.code:
                dhf_map[nm] = str(d.code).strip()
        return {
            "name": name,
            "release_date": release_date,
            "overview": _overview_text(name, release_version, full_version, overall_desc),
            "archive": _archive_text(name),
            "version": doc_version,
            "pm": pm,
            "approver": approver,
            "dhf_map": dhf_map,
            "acceptance_date": acceptance_date,
        }

    @staticmethod
    def __dhf_code_of(dhf_map, name):
        k = str(name or "").strip()
        if not k:
            return ""
        if k in dhf_map:
            return dhf_map[k]
        for nm, code in dhf_map.items():
            if nm and (nm in k or k in nm):
                return code
        return ""

    @staticmethod
    def __resolve_doc_cell(dhf_map, val):
        # 「编号 文档名」或纯「文档名」单元格 → 按文档名(含匹配)取 DHF 最新编号，保留原文档名文案
        s = str(val or "").strip()
        if not s:
            return val
        m = re.match(r"^([A-Z][A-Z0-9]*(?:-[A-Z0-9]+)+)\s+(\S.*)$", s)
        if m:
            name = m.group(2).strip()
            code = Server.__dhf_code_of(dhf_map, name)
            return f"{code} {name}" if code else val
        if s in dhf_map:
            return f"{dhf_map[s]} {s}"
        return val

    def __fill_node(self, node, info):
        ref = node.get("ref_type")
        title = self.__strip_num(node.get("title"))
        if ref == "rn_overview" or title == "产品概述":
            node["body"] = info["overview"]
        elif ref == "rn_release_time" or title == "发布时间":
            node["body"] = info["release_date"] or ""
        elif ref == "rn_archive" or title == "文档归档":
            if info["name"]:
                node["body"] = info["archive"]
        if ref == "rn_transfer_files" or title == "文件移交记录":
            for tbl in (node.get("tables") or []):
                if not tbl or not isinstance(tbl[0], list):
                    continue
                header = tbl[0]
                is_file_tbl = any("文件编号" in str(h) for h in header) and any("文件名称" in str(h) for h in header)
                if not is_file_tbl:
                    continue
                for row in tbl[1:]:
                    if not isinstance(row, list) or len(row) < 2:
                        continue
                    code = self.__dhf_code_of(info["dhf_map"], row[1])
                    if code:
                        row[0] = code
        if ref == "rn_transfer_product" or title == "产品移交记录":
            for tbl in (node.get("tables") or []):
                if not isinstance(tbl, list):
                    continue
                for row in tbl:
                    if not isinstance(row, list):
                        continue
                    for ci in range(len(row)):
                        row[ci] = self.__resolve_doc_cell(info["dhf_map"], row[ci])
                    # 交付时间：取时间线「产品验收记录」日期，取不到则清空
                    if len(row) > 2 and "交付时间" in str(row[1] or ""):
                        row[2] = info.get("acceptance_date") or ""
        if ref == "revision" or title == "文件修订记录":
            tables = node.get("tables") or []
            if tables and isinstance(tables[0], list):
                t = tables[0]
                cols = len(t[0]) if t and t[0] else 5
                while len(t) < 6:
                    t.append([""] * cols)
                row = t[1]
                while len(row) < 5:
                    row.append("")

                row[0] = info["release_date"] or ""
                if info["version"]:
                    row[1] = info["version"]
                if not str(row[2] or "").strip():
                    row[2] = "首次发布"
                row[3] = info["pm"] or ""
                if info["approver"]:
                    row[4] = info["approver"]
        for child in (node.get("children") or []):
            self.__fill_node(child, info)

    def __autofill_for_export(self, content, obj: ReleaseNoteObj):
        sections = (content or {}).get("sections") or []
        if not obj.product_id:
            return content
        product = db.session.execute(select(Product).where(Product.id == obj.product_id)).scalars().first()
        info = self.__collect_autofill(obj.product_id, product, obj.version)
        for node in sections:
            self.__fill_node(node, info)
        return content

    def __dhf_file_no(self, prod_id):
        row = db.session.execute(
            select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name.like(f"%{DOC_NAME}%"))
        ).scalars().first()
        return (row.code or "").strip() if row and row.code else ""

    def __to_obj(self, row: ReleaseNote, product: Product = None):
        obj = ReleaseNoteObj(**row.dict())
        obj.content = self.__normalize_content(obj.content)
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
            # 文件编号优先用文档已填值，未填时从产品 DHF 中按文档名匹配获取
            if not (obj.file_no or "").strip():
                dhf_no = self.__dhf_file_no(product.id)
                if dhf_no:
                    obj.file_no = dhf_no
        return obj

    # ---------------- CRUD ----------------
    async def add_release_note(self, form: ReleaseNoteForm):
        try:
            sql = select(func.count(ReleaseNote.id)).where(
                ReleaseNote.product_id == form.product_id,
                ReleaseNote.version == form.version,
            )
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = ReleaseNote(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=ReleaseNoteForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_release_note(self, id: int, product_id: int = None):
        try:
            fromdoc: ReleaseNote = db.session.execute(select(ReleaseNote).where(ReleaseNote.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(ReleaseNote.version).where(ReleaseNote.product_id == target_pid)).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = ReleaseNote(
                product_id=target_pid,
                version=version,
                file_no=fromdoc.file_no,
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            return Resp.resp_ok(data=ReleaseNoteForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_release_note(self, form: ReleaseNoteForm):
        try:
            row: ReleaseNote = db.session.execute(select(ReleaseNote).where(ReleaseNote.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_release_note(self, id: int):
        db.session.execute(delete(ReleaseNote).where(ReleaseNote.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_release_note(self, id: int):
        sql = select(ReleaseNote, Product).join(Product, ReleaseNote.product_id == Product.id).where(ReleaseNote.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_release_note(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(ReleaseNote.product_id == product_id)
        if version:
            wheres.append(ReleaseNote.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(ReleaseNote.id)).join(Product, ReleaseNote.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (
            select(ReleaseNote, Product)
            .join(Product, ReleaseNote.product_id == Product.id)
            .where(*wheres)
            .order_by(ReleaseNote.id.desc())
            .offset(page_index * page_size)
            .limit(page_size)
        )
        rows: List[ReleaseNoteObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    # ---------------- 导出 Word ----------------
    async def export_release_note(self, output, id: int):
        resp = await self.get_release_note(id)
        obj: ReleaseNoteObj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        c = self.__autofill_for_export(self.__normalize_content(obj.content), obj)
        sections = c.get("sections") or []
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")

        def write_center_title(text, size=22.0, bold=False):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.right_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, text, font_size=size, bold=bold)

        def add_blank_lines(count):
            for _ in range(max(0, int(count or 0))):
                document.add_paragraph("")

        def add_text(text):
            docx_util.save_txt2docx(str(text or ""), document)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            cell.text = ""
            lines = str(text or "").split("\n")
            for i, line in enumerate(lines):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER if align == WD_ALIGN_PARAGRAPH.CENTER else WD_CELL_VERTICAL_ALIGNMENT.TOP

        def merge_col_runs(table, col):
            # 纵向合并某列中相邻且相同(非空)的单元格
            rows = table.rows
            n = len(rows)
            r = 0
            while r < n:
                text = rows[r].cells[col].text
                r2 = r
                while r2 + 1 < n and rows[r2 + 1].cells[col].text == text:
                    r2 += 1
                if r2 > r and text.strip():
                    merged = rows[r].cells[col].merge(rows[r2].cells[col])
                    set_cell(merged, text, align=WD_ALIGN_PARAGRAPH.CENTER)
                    merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                r = r2 + 1

        def add_grid(grid, merge_col0=False, has_header=True):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for r_idx, row in enumerate(grid):
                cells = table.add_row().cells
                for c_idx in range(cols):
                    set_cell(cells[c_idx], row[c_idx] if c_idx < len(row) else "", bold=(has_header and r_idx == 0))
            if merge_col0:
                merge_col_runs(table, 0)
            document.add_paragraph()

        def add_cover_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for row in grid:
                cells = table.add_row().cells
                for c_idx in range(cols):
                    text = row[c_idx] if c_idx < len(row) else ""
                    set_cell(cells[c_idx], text, bold=(c_idx % 2 == 0), align=WD_ALIGN_PARAGRAPH.CENTER)
                if (str(row[0]).strip() if row else "") == "生效日期" and cols > 2:
                    merged = cells[1]
                    for c_idx in range(2, cols):
                        merged = merged.merge(cells[c_idx])
                    set_cell(merged, row[1] if len(row) > 1 else "", align=WD_ALIGN_PARAGRAPH.CENTER)
            document.add_paragraph()

        def add_body_heading(title, level):
            size = {1: 16.0, 2: 14.0, 3: 12.0}.get(level, 11.0)
            p = document.add_heading("", level=max(1, min(level, 9)))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, title, font_size=size, bold=True)

        def render_body_section(node, level, number=""):
            name = self.__strip_num(node.get("title"))
            heading = f"{number} {name}".strip() if number else name
            add_body_heading(heading, level=max(1, min(level, 9)))
            if (node.get("body") or "").strip():
                add_text(node.get("body"))
            is_prod_rec = node.get("ref_type") == "rn_transfer_product" or self.__strip_num(node.get("title")) == "产品移交记录"
            for table in (node.get("tables") or []):
                add_grid(table, merge_col0=is_prod_rec, has_header=not is_prod_rec)
            idx = 0
            for child in (node.get("children") or []):
                idx += 1
                child_num = f"{number}.{idx}" if number else f"{idx}"
                render_body_section(child, level + 1, child_num)

        cover = next((s for s in sections if s.get("ref_type") == "cover"), None)
        revision = next((s for s in sections if s.get("ref_type") == "revision"), None)
        body = [s for s in sections if s.get("ref_type") not in ("cover", "revision")]

        # 封面：上方留白 + 居中大标题 + 留白 + 封面表
        add_blank_lines(6)
        write_center_title((self.__strip_num(cover.get("title")) if cover else "") or DOC_NAME, size=22.0, bold=True)
        add_blank_lines(4)
        if cover:
            for table in (cover.get("tables") or []):
                add_cover_grid(table)

        # 文件修订记录：独立页、居中标题
        document.add_page_break()
        write_center_title("文件修订记录", size=14.0, bold=True)
        add_blank_lines(2)
        if revision:
            for table in (revision.get("tables") or []):
                add_grid(table)

        # 正文：另起一页（源文档无目录，故不插入目录）
        document.add_page_break()
        for i, node in enumerate(body):
            render_body_section(node, 1, str(i + 1))

        document.save(output)
        output.seek(0)
