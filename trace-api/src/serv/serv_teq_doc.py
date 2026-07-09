#!/usr/bin/env python
# encoding: utf-8

# 测试设备清单服务层（测试文件）。整份文档为单张设备清单表，以 content(JSON) 存储。

import copy
import logging
import re
from typing import List

from sqlalchemy import delete, func, select
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.teq_doc import TeqDoc
from ..model.prod_dhf import ProdDhf
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_teq_doc import TeqDocForm
from ..obj.vobj_teq_doc import TeqDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from .serv_utils import new_version, sync_file_no_version
from .serv_utils import docx_util

logger = logging.getLogger(__name__)

_HEADER = ["序号", "名称", "规格型号", "品牌", "资产编码", "类别", "用途", "地点", "使用人"]
# 测试设备清单默认模板（取自《测试设备清单》，可编辑）。
_ROWS = [
    ["1", "测试机", "组装机", "组装机", "SER01405", "办公电脑", "测试共用", "北京", "宁随军"],
    ["2", "测试机", "组装机", "组装机", "SER01358", "办公电脑", "测试自用", "北京", "徐秋实"],
    ["3", "测试机", "27英寸", "DELL", "ILCD22060028", "办公电脑", "测试自用", "北京", "徐秋实"],
    ["4", "测试机", "CU34G2X", "DELL", "ILCD21120001", "办公电脑", "测试自用", "北京", "徐秋实"],
    ["5", "测试机", "四卡", "组装机", "SER01268", "办公电脑", "测试自用", "北京", "李鹏"],
    ["6", "测试机", "DELL", "DELL", "LCD01371", "办公电脑", "测试自用", "北京", "李鹏"],
    ["7", "测试机", "DELL", "DELL", "LCD01375", "办公电脑", "测试自用", "北京", "李鹏"],
    ["8", "测试机", "两卡", "组装机", "SER01100", "办公电脑", "测试自用", "北京", "杨学峰"],
    ["9", "测试机", "P2722H", "DELL", "ILCD21110098", "办公电脑", "测试自用", "北京", "杨学峰"],
    ["10", "测试机", "P2722H", "DELL", "ILCD21110100", "办公电脑", "测试自用", "北京", "杨学峰"],
]


def _default_content():
    return {"rows": [list(_HEADER)] + [list(r) for r in _ROWS]}


class Server(object):

    def __normalize_content(self, content):
        base = _default_content()
        if isinstance(content, dict) and isinstance(content.get("rows"), list) and content["rows"]:
            base["rows"] = [list(r) if isinstance(r, list) else [r] for r in content["rows"]]
        return base

    def __dhf_file_no(self, prod_id):
        if not prod_id:
            return ""
        row = db.session.execute(
            select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name == "测试设备清单")
            .order_by(ProdDhf.id.asc())
        ).scalars().first()
        if not row:
            row = db.session.execute(
                select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name.like("%测试设备清单%"))
                .order_by(ProdDhf.id.asc())
            ).scalars().first()
        return (row.code or "").strip() if row and row.code else ""

    def __to_obj(self, row: TeqDoc, product: Product = None):
        obj = TeqDocObj(**row.dict())
        obj.content = self.__normalize_content(obj.content)
        if not (obj.file_no or "").strip():
            obj.file_no = self.__dhf_file_no(row.product_id)
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
        return obj

    async def add_teq_doc(self, form: TeqDocForm):
        try:
            sql = select(func.count(TeqDoc.id)).where(
                TeqDoc.product_id == form.product_id,
                TeqDoc.version == form.version,
            )
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = TeqDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=TeqDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_teq_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: TeqDoc = db.session.execute(select(TeqDoc).where(TeqDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(TeqDoc.version).where(TeqDoc.product_id == target_pid)).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = TeqDoc(
                product_id=target_pid,
                version=version,
                file_no=sync_file_no_version(fromdoc.file_no, version),
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            return Resp.resp_ok(data=TeqDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_teq_doc(self, form: TeqDocForm):
        try:
            row: TeqDoc = db.session.execute(select(TeqDoc).where(TeqDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_teq_doc(self, id: int):
        db.session.execute(delete(TeqDoc).where(TeqDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_teq_doc(self, id: int):
        sql = select(TeqDoc, Product).join(Product, TeqDoc.product_id == Product.id).where(TeqDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_teq_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(TeqDoc.product_id == product_id)
        if version:
            wheres.append(TeqDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(TeqDoc.id)).join(Product, TeqDoc.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (
            select(TeqDoc, Product)
            .join(Product, TeqDoc.product_id == Product.id)
            .where(*wheres)
            .order_by(TeqDoc.id.desc())
            .offset(page_index * page_size)
            .limit(page_size)
        )
        rows: List[TeqDocObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    # ---------------- 导出 Word ----------------
    async def export_teq_doc(self, output, id: int):
        resp = await self.get_teq_doc(id)
        obj: TeqDocObj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        c = self.__normalize_content(obj.content)
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")
        docx_util.add_page_number_footer(section, obj.file_no or "", skip_first=False)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = align
            para.paragraph_format.line_spacing = 1.2
            docx_util.fonted_txt(para, str(text or ""), font_size=10.0, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        docx_util.fonted_txt(title, "测试设备清单", font_size=18.0, bold=True)
        document.add_paragraph()

        rows = [r for r in (c.get("rows") or []) if isinstance(r, list)]
        cols = max((len(r) for r in rows), default=len(_HEADER))
        tb = document.add_table(rows=0, cols=cols)
        tb.style = "Table Grid"
        tb.alignment = WD_TABLE_ALIGNMENT.CENTER
        for r_idx, row in enumerate(rows):
            cells = tb.add_row().cells
            for c_idx in range(cols):
                set_cell(cells[c_idx], row[c_idx] if c_idx < len(row) else "", bold=(r_idx == 0))
        document.save(output)
        output.seek(0)