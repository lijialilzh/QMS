#!/usr/bin/env python
# encoding: utf-8

# 风险管理计划服务层（参考网络安全维护计划 serv_nsmp_doc + 网络安全研究报告的正文图片渲染）。
# 整份文档以 content(JSON)「章节树」存储；产品信息（名称/版本/项目时间/人员/日期）自动获取注入，其余模板化。
# 导出：封面「风险管理计划」→ 封面信息表 → 文件修订记录 → 目录(TOC) → 正文章节（带章节号，含正文表格与正文图片）。

import base64
import copy
import io
import json
import logging
import os
import re
from typing import List

from sqlalchemy import delete, func, select
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from ..model.product import Product
from ..model.rmp_doc import RmpDoc
from ..model.prod_dhf import ProdDhf
from ..model.company_info import CompanyInfo
from ..model.project_member import ProjectMember
from ..model.project_timeline import ProjectTimelineRow, ProjectTimelineCell
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_rmp_doc import RmpDocForm
from ..obj.vobj_rmp_doc import RmpDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from . import serv_review_util
from .serv_utils import new_version, sync_file_no_version, docx_util

logger = logging.getLogger(__name__)

DOC_TITLE = "风险管理计划"
# 修订日期从时间线「输出结果」匹配的文档名关键字（命中项取最早日期，参考产品开发计划）
DATE_KEYWORDS = ["风险管理计划"]

# 表1「风险管理小组」姓名列自动获取：按「项目角色」精确匹配项目人员职能关键字（按优先级取首个命中）
TEAM_ROLE_KEYWORDS = {
    "产品经理": ["产品经理", "项目经理"],
    "模型负责人": ["模型负责人", "算法负责人"],
    "产品开发负责人": ["产品开发负责人", "研发负责人", "开发负责人"],
    "RA负责人": ["RA负责人", "RA", "法规负责人"],
    "QA负责人": ["QA负责人", "QA", "质量负责人"],
    "验证和确认负责人": ["验证和确认负责人", "验证负责人", "确认负责人"],
    "执行负责人": ["执行负责人"],
    "临床专家": ["临床专家"],
}

DEFAULT_RMP_CONTENT = {"sections": []}
_DEFAULT_CONTENT_FILE = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "src-res", "rmp_default_content.json"
)
try:
    with open(_DEFAULT_CONTENT_FILE, encoding="utf-8") as _f:
        _loaded = json.load(_f)
    if isinstance(_loaded, dict) and isinstance(_loaded.get("sections"), list) and _loaded["sections"]:
        DEFAULT_RMP_CONTENT = _loaded
except Exception:
    logger.exception("加载风险管理计划默认内容资源失败")


class Server(object):

    # ---------------- 归一化 ----------------
    def __normalize_node(self, node):
        if not isinstance(node, dict):
            return {"title": str(node or ""), "body": "", "tables": [], "images": [], "children": []}
        result = dict(node)
        result["title"] = str(result.get("title") or "")
        result["body"] = str(result.get("body") or "")
        if "body_after" in result:
            result["body_after"] = str(result.get("body_after") or "")
        tables = result.get("tables")
        norm_tables = []
        if isinstance(tables, list):
            for table in tables:
                if isinstance(table, list):
                    norm_tables.append([[str(c) if c is not None else "" for c in (row or [])] for row in table if isinstance(row, list)])
        result["tables"] = norm_tables
        imgs = result.get("images")
        result["images"] = [str(x) for x in imgs if isinstance(x, str) and x] if isinstance(imgs, list) else []
        children = result.get("children")
        result["children"] = [self.__normalize_node(c) for c in children] if isinstance(children, list) else []
        return result

    def __normalize_content(self, content):
        if not isinstance(content, dict) or not isinstance(content.get("sections"), list):
            return copy.deepcopy(DEFAULT_RMP_CONTENT)
        return {"sections": [self.__normalize_node(s) for s in content["sections"]]}

    @staticmethod
    def __strip_num(title):
        return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()

    # ---------------- 自动获取 ----------------
    def __company_name(self, product):
        prod_registrant = (getattr(product, "registrant", "") or "").strip() if product else ""
        company = None
        if prod_registrant:
            company = db.session.execute(
                select(CompanyInfo).where(CompanyInfo.registrant == prod_registrant)
            ).scalars().first()
        if not company:
            company = db.session.execute(select(CompanyInfo).order_by(CompanyInfo.id.asc())).scalars().first()
        return (getattr(company, "registrant", "") or "").strip() if company else prod_registrant

    def __member_name(self, prod_id, keywords):
        for kw in keywords:
            row = db.session.execute(
                select(ProjectMember).where(ProjectMember.prod_id == prod_id, ProjectMember.role.like(f"%{kw}%"))
                .order_by(ProjectMember.sort_order.asc(), ProjectMember.id.asc())
            ).scalars().first()
            if row and (row.name or "").strip():
                return row.name.strip()
        return ""

    @staticmethod
    def __to_int(v):
        digits = re.sub(r"[^\d]", "", str(v or ""))
        return int(digits) if digits else None

    def __date_rows(self, prod_id):
        tl_rows = db.session.execute(
            select(ProjectTimelineRow).where(ProjectTimelineRow.prod_id == prod_id)
        ).scalars().all()
        return [r for r in tl_rows if (r.row_type or "date") == "date" and self.__to_int(r.year) and self.__to_int(r.month)]

    def __release_date(self, prod_id):
        # 修订日期：从项目时间线取「输出结果含文档名」的日期行，取最早一条；取不到留空
        date_rows = self.__date_rows(prod_id)
        if not date_rows:
            return ""
        cell_map = {}
        for c in db.session.execute(
            select(ProjectTimelineCell).where(ProjectTimelineCell.row_id.in_([r.id for r in date_rows]))
        ).scalars().all():
            cell_map.setdefault(c.row_id, []).append(c.output_result or "")

        def date_key(r):
            return self.__to_int(r.year) * 10000 + self.__to_int(r.month) * 100 + (self.__to_int(r.day) or 0)

        file_rows = [r for r in date_rows if any(k in str(v or "") for k in DATE_KEYWORDS for v in cell_map.get(r.id, []))]
        if not file_rows:
            return ""
        fr = min(file_rows, key=date_key)
        return f"{self.__to_int(fr.year)}年{self.__to_int(fr.month)}月{self.__to_int(fr.day)}日"

    def __time_range(self, prod_id):
        # 项目时间范围：时间线最早～最晚（年月），供「风险管理活动计划」章节使用；取不到留空
        date_rows = self.__date_rows(prod_id)
        if not date_rows:
            return ""

        def ym_key(r):
            return self.__to_int(r.year) * 100 + self.__to_int(r.month)

        start = min(date_rows, key=ym_key)
        end = max(date_rows, key=ym_key)
        s = f"{self.__to_int(start.year)}年{self.__to_int(start.month)}月"
        e = f"{self.__to_int(end.year)}年{self.__to_int(end.month)}月"
        return f"{s}至{e}"

    def __collect_autofill(self, product_id):
        product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
        if not product:
            return {}
        return {
            "prod_id": product_id,
            "prod_name": (product.name or "").strip(),
            "full_version": (product.full_version or "").strip(),
            "company": self.__company_name(product),
            "reviser": self.__member_name(product_id, ("产品经理", "项目经理")),
            "auditor": self.__member_name(product_id, ("QA负责人", "质量负责人", "QA")),
            "approver": self.__member_name(product_id, ("产品负责人", "研发负责人", "管理者代表")),
            "rev_date": self.__release_date(product_id),
            "time_range": self.__time_range(product_id),
        }

    def __fill_team_table(self, node, prod_id):
        # 表1「风险管理小组」：按表头(项目角色/姓名/职责)识别，姓名列按项目角色从项目人员自动获取（命中即覆盖）
        if not prod_id:
            return
        for table in (node.get("tables") or []):
            if not table or not isinstance(table[0], list):
                continue
            header = [str(c).strip() for c in table[0][:3]]
            if header != ["项目角色", "姓名", "职责"]:
                continue
            for row in table[1:]:
                if not isinstance(row, list) or len(row) < 2:
                    continue
                keywords = TEAM_ROLE_KEYWORDS.get(str(row[0]).strip())
                if not keywords:
                    continue
                name = self.__member_name(prod_id, keywords)
                if name:
                    row[1] = name

    def __fill_node(self, node, info, version):
        ref = node.get("ref_type")
        title = self.__strip_num(node.get("title"))
        self.__fill_team_table(node, info.get("prod_id"))
        if ref == "scope" or title == "范围":
            node["body"] = (
                "本文件适用的产品为的软件产品：\n"
                f"产品名称：{info.get('prod_name', '')}\n"
                f"完整版本：{info.get('full_version', '')}\n"
                "寿命周期内产品需要开展的风险管理活动的范围包括：\n"
                "从实施过程的角度包括：设计和研发；采购；制造；包装；检查和测试等。\n"
                "交付过程：贮存和运输。\n"
                "交付后：操作说明，软件升级维护，使用过程中可能面临的问题。\n"
                "产品失效后的处理。"
            )
        elif ref == "activity_plan":
            name = info.get("prod_name", "") or "产品"
            time_range = info.get("time_range", "")
            node["body"] = (
                f"风险管理的计划的阶段和里程碑与{name}的项目管理计划保持一致，"
                f"项目时间范围为{time_range}，软件产品发布后进入到确认阶段。\n表 2 风险管理活动计划"
            )
        elif ref == "cover":
            for table in (node.get("tables") or []):
                for row in table:
                    if not isinstance(row, list) or not row:
                        continue
                    label = str(row[0]).strip()
                    # 文件版本填文档版本
                    if len(row) >= 4 and str(row[2]).strip() == "文件版本" and version and not str(row[3] or "").strip():
                        row[3] = version

                    def set_name(val):
                        if val and len(row) >= 2 and not str(row[1] or "").strip():
                            row[1] = val

                    def set_date(val):
                        if val and len(row) >= 4 and not str(row[3] or "").strip():
                            row[3] = val
                    # 编制/审核/批准人「姓名」列由签名规则统一填充（见 __apply_autofill），此处仅填日期
                    if label in ("编制人", "审核人", "批准人"):
                        set_date(info.get("rev_date", ""))
                    elif label == "生效日期":
                        set_name(info.get("rev_date", ""))
        elif ref == "revision" or title == "文件修订记录":
            for table in (node.get("tables") or []):
                if len(table) >= 2 and isinstance(table[1], list) and len(table[1]) >= 5:
                    row = table[1]

                    def set_if(i, val):
                        if val and not str(row[i] if i < len(row) else "").strip():
                            row[i] = val
                    set_if(0, info.get("rev_date", ""))
                    set_if(1, version)
                    if not str(row[2] or "").strip():
                        row[2] = "首次发布"
                    set_if(3, info.get("reviser", ""))
                    set_if(4, info.get("approver", ""))
        for child in (node.get("children") or []):
            self.__fill_node(child, info, version)

    def __apply_autofill(self, content, info, version):
        for node in (content.get("sections") or []):
            self.__fill_node(node, info, version)
        # 封面「编制/审核/批准人」按部门签名规则填充签名图（仅填空）
        serv_review_util.fill_cover_signers(content, serv_review_util.cover_signers(info.get("prod_id"), "rmp"))
        # 评审记录（内置于默认内容的「附：评审记录」）：自动填参评人签字、批准人签字/日期、其他参评人员合并/
        rev = serv_review_util.review_date(info.get("prod_id"), ["风险管理计划"]) if info.get("prod_id") else ""
        for section in (content.get("sections") or []):
            for tbl in (section.get("tables") or []):
                serv_review_util.autofill_review_person_table(tbl, "rmp", rev, info.get("prod_id"))
        return content

    # ---------------- 文件编号（未手填时从产品 DHF 匹配） ----------------
    def __dhf_file_no(self, prod_id):
        row = db.session.execute(
            select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name.like("%风险管理计划%"))
        ).scalars().first()
        return (row.code or "").strip() if row and row.code else ""

    def __to_obj(self, row: RmpDoc, product: Product = None, with_autofill=True):
        obj = RmpDocObj(**row.dict())
        content = self.__normalize_content(obj.content)
        if with_autofill and row.product_id:
            info = self.__collect_autofill(row.product_id)
            content = self.__apply_autofill(content, info, obj.version)
        obj.content = content
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
            if not (obj.file_no or "").strip():
                dhf_no = self.__dhf_file_no(product.id)
                if dhf_no:
                    obj.file_no = dhf_no
        return obj

    # ---------------- CRUD ----------------
    async def add_rmp_doc(self, form: RmpDocForm):
        try:
            sql = select(func.count(RmpDoc.id)).where(RmpDoc.product_id == form.product_id, RmpDoc.version == form.version)
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = RmpDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=RmpDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_rmp_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: RmpDoc = db.session.execute(select(RmpDoc).where(RmpDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(RmpDoc.version).where(RmpDoc.product_id == target_pid)).scalars().all()
            existing = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing:
                version = new_version(version)
            base_file_no = (fromdoc.file_no or "").strip() or self.__dhf_file_no(target_pid)
            newdoc = RmpDoc(
                product_id=target_pid, version=version, file_no=sync_file_no_version(base_file_no, version) or None,
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            return Resp.resp_ok(data=RmpDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_rmp_doc(self, form: RmpDocForm):
        try:
            row: RmpDoc = db.session.execute(select(RmpDoc).where(RmpDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            if form.product_id is not None:
                row.product_id = form.product_id
            if form.version is not None:
                row.version = form.version
            if form.file_no is not None:
                row.file_no = (form.file_no or "").strip() or None
            if form.change_log is not None:
                row.change_log = form.change_log
            if form.content is not None:
                row.content = self.__normalize_content(form.content)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_rmp_doc(self, id: int):
        db.session.execute(delete(RmpDoc).where(RmpDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_rmp_doc(self, id: int):
        sql = select(RmpDoc, Product).join(Product, RmpDoc.product_id == Product.id).where(RmpDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_rmp_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(RmpDoc.product_id == product_id)
        if version:
            wheres.append(RmpDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        total = db.session.execute(
            select(func.count(RmpDoc.id)).join(Product, RmpDoc.product_id == Product.id).where(*wheres)
        ).scalar() or 0
        sql = (
            select(RmpDoc, Product).join(Product, RmpDoc.product_id == Product.id)
            .where(*wheres).order_by(RmpDoc.id.desc()).offset(page_index * page_size).limit(page_size)
        )
        rows: List[RmpDocObj] = [self.__to_obj(doc, product, with_autofill=False) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    async def rmp_autofill(self, product_id: int, version: str = ""):
        content = self.__normalize_content(None)
        info = self.__collect_autofill(product_id)
        content = self.__apply_autofill(content, info, version or "")
        return Resp.resp_ok(data=content)

    # ---------------- 导出 Word ----------------
    async def export_rmp_doc(self, output, id: int):
        resp = await self.get_rmp_doc(id)
        obj: RmpDocObj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        content = obj.content if isinstance(obj.content, dict) else self.__normalize_content(obj.content)
        sections = content.get("sections") or []

        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        document.settings.element.append(update_fields)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")
        docx_util.add_page_number_footer(section, obj.file_no or "")

        def add_text(text):
            if str(text or "").strip():
                docx_util.save_txt2docx(str(text or ""), document)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            s = str(text or "")
            # 签名图（编制/审核/批准人）：等比嵌入图片
            if s.startswith("data:image"):
                try:
                    b64 = s.split(",", 1)[1] if "," in s else ""
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = align
                    para.add_run().add_picture(io.BytesIO(base64.b64decode(b64)), height=Pt(33))
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    return
                except Exception:
                    pass
            cell.text = ""
            for i, line in enumerate(s.split("\n")):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        def set_grid_widths(table, grid, cols):
            def cell_units(text):
                m = 0
                for line in str(text or "").split("\n"):
                    w = sum(2 if ord(ch) > 127 else 1 for ch in line)
                    m = max(m, w)
                return max(m, 2)
            col_units = []
            for c in range(cols):
                u = 2
                for row in grid:
                    if c < len(row):
                        u = max(u, cell_units(row[c]))
                col_units.append(u)
            widths = [u * 120 + 440 for u in col_units]
            sect = document.sections[0]
            usable = int((sect.page_width - sect.left_margin - sect.right_margin) / 635)
            total = sum(widths)
            if total > usable and total > 0:
                scale = usable / total
                widths = [max(int(w * scale), 600) for w in widths]
                total = sum(widths)
            tbl_pr = table._tbl.tblPr
            layout = tbl_pr.find(qn("w:tblLayout"))
            if layout is None:
                layout = OxmlElement("w:tblLayout")
                tbl_pr.append(layout)
            layout.set(qn("w:type"), "fixed")
            tbl_w = tbl_pr.find(qn("w:tblW"))
            if tbl_w is None:
                tbl_w = OxmlElement("w:tblW")
                tbl_pr.append(tbl_w)
            tbl_w.set(qn("w:w"), str(total))
            tbl_w.set(qn("w:type"), "dxa")
            grid_el = table._tbl.find(qn("w:tblGrid"))
            if grid_el is not None:
                for i, gc in enumerate(grid_el.findall(qn("w:gridCol"))):
                    if i < len(widths):
                        gc.set(qn("w:w"), str(widths[i]))
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    if i >= len(widths):
                        continue
                    tc_pr = cell._tc.get_or_add_tcPr()
                    tc_w = tc_pr.find(qn("w:tcW"))
                    if tc_w is None:
                        tc_w = OxmlElement("w:tcW")
                        tc_pr.append(tc_w)
                    tc_w.set(qn("w:w"), str(widths[i]))
                    tc_w.set(qn("w:type"), "dxa")

        def add_grid(grid, merge_full_rows=False):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return

            def rest_all_empty(row):
                for c in range(1, cols):
                    if str((row[c] if c < len(row) else "") or "").strip():
                        return False
                return True

            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            h_merged = []
            for r_idx, row in enumerate(grid):
                cells = table.add_row().cells
                first = row[0] if row else ""
                first_txt = str(first or "").strip()
                is_other = first_txt.startswith("其他参会人员") or first_txt.startswith("其他参评人员")
                # 仅评审记录：整行只有第一格有内容时，跨整行合并
                if merge_full_rows and cols > 1 and first_txt and rest_all_empty(row):
                    merged = cells[0].merge(cells[cols - 1])
                    center = first_txt.startswith("参评人员签字") or first_txt.startswith("评审时间")
                    align = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
                    set_cell(merged, first, bold=(r_idx == 0), align=align)
                    h_merged.append(True)
                elif merge_full_rows and cols > 2 and is_other:
                    # 「其他参会人员/其他参评人员」行：标签后单元格合并为一格并居中显示「/」
                    set_cell(cells[0], first, align=WD_ALIGN_PARAGRAPH.CENTER)
                    merged = cells[1].merge(cells[cols - 1])
                    set_cell(merged, "/", align=WD_ALIGN_PARAGRAPH.CENTER)
                    h_merged.append(False)
                else:
                    for c_idx in range(cols):
                        set_cell(cells[c_idx], row[c_idx] if c_idx < len(row) else "", bold=(r_idx == 0))
                    h_merged.append(False)
            # 仅评审记录：第一列连续相同的非空项纵向合并（跳过整行合并的行与表头）
            if merge_full_rows and cols > 0:
                n = len(grid)
                r = 1
                while r < n:
                    val = str((grid[r][0] if grid[r] else "") or "").strip()
                    if h_merged[r] or not val:
                        r += 1
                        continue
                    j = r + 1
                    while j < n and not h_merged[j] and str((grid[j][0] if grid[j] else "") or "").strip() == val:
                        j += 1
                    if j - 1 > r:
                        m = table.cell(r, 0).merge(table.cell(j - 1, 0))
                        set_cell(m, val)
                    r = j
            set_grid_widths(table, grid, cols)
            document.add_paragraph()

        def add_image(image_url):
            raw_url = str(image_url or "").strip()
            if not raw_url or not raw_url.startswith("data:image/"):
                return
            try:
                image_data = raw_url.split(",", 1)[1]
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic = paragraph.add_run().add_picture(io.BytesIO(base64.b64decode(image_data)))
                max_w = Inches(5.5)
                max_h = Inches(7.0)
                if pic.width and pic.height:
                    ratio = min(max_w / pic.width, max_h / pic.height, 1)
                    pic.width = int(pic.width * ratio)
                    pic.height = int(pic.height * ratio)
                document.add_paragraph()
            except Exception:
                logger.exception("导出风险管理计划图片失败")

        def write_center_title(text, size=22.0, bold=True):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, text, font_size=size, bold=bold)

        def add_blank_lines(count):
            for _ in range(max(0, int(count or 0))):
                document.add_paragraph("")

        def insert_toc_field():
            p = document.add_paragraph()
            run_begin = p.add_run()
            fld_begin = OxmlElement("w:fldChar")
            fld_begin.set(qn("w:fldCharType"), "begin")
            fld_begin.set(qn("w:dirty"), "true")
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = ' TOC \\o "1-3" \\h \\z \\u '
            fld_separate = OxmlElement("w:fldChar")
            fld_separate.set(qn("w:fldCharType"), "separate")
            run_end = p.add_run()
            fld_end = OxmlElement("w:fldChar")
            fld_end.set(qn("w:fldCharType"), "end")
            run_begin._r.append(fld_begin)
            run_begin._r.append(instr)
            run_begin._r.append(fld_separate)
            p.add_run("目录将在打开文档后自动更新")
            run_end._r.append(fld_end)

        def add_cover_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for row in grid:
                cells = table.add_row().cells
                for c_idx in range(cols):
                    text = row[c_idx] if c_idx < len(row) else ""
                    set_cell(cells[c_idx], text, bold=(c_idx % 2 == 0))
                    cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if (str(row[0]).strip() if row else "") == "生效日期" and cols > 2:
                    merged = cells[1]
                    for c_idx in range(2, cols):
                        merged = merged.merge(cells[c_idx])
                    set_cell(merged, row[1] if len(row) > 1 else "", align=WD_ALIGN_PARAGRAPH.CENTER)
            document.add_paragraph()

        def add_heading(text, level):
            size = {1: 16.0, 2: 14.0, 3: 12.0}.get(level, 11.0)
            p = document.add_heading("", level=max(1, min(level, 9)))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, text, font_size=size, bold=True)

        def render_section(node, level, number=""):
            name = self.__strip_num(node.get("title"))
            heading = f"{number} {name}".strip() if number else name
            add_heading(heading, level)
            add_text(node.get("body"))
            merge_full_rows = node.get("ref_type") == "appendix"
            for table in (node.get("tables") or []):
                add_grid(table, merge_full_rows=merge_full_rows)
            for image_url in (node.get("images") or []):
                add_image(image_url)
            add_text(node.get("body_after"))
            idx = 0
            for child in (node.get("children") or []):
                idx += 1
                child_num = f"{number}.{idx}" if number else f"{idx}"
                render_section(child, level + 1, child_num)

        cover = next((s for s in sections if s.get("ref_type") == "cover"), None)
        revision = next((s for s in sections if s.get("ref_type") == "revision"), None)
        body = [s for s in sections if s.get("ref_type") not in ("cover", "revision", "appendix")]
        appendices = [s for s in sections if s.get("ref_type") == "appendix"]

        # 封面：上方留白 + 居中大标题 + 留白 + 封面表
        add_blank_lines(6)
        write_center_title((self.__strip_num(cover.get("title")) if cover else "") or DOC_TITLE, size=22.0, bold=True)
        add_blank_lines(4)
        if cover:
            for table in (cover.get("tables") or []):
                add_cover_grid(table)

        # 文件修订记录
        document.add_page_break()
        write_center_title("文件修订记录", size=14.0, bold=True)
        add_blank_lines(2)
        if revision:
            for table in (revision.get("tables") or []):
                add_grid(table)

        # 目录
        document.add_page_break()
        write_center_title("目录", size=16.0, bold=True)
        insert_toc_field()

        # 正文：另起一页，按 1/2/3 编号
        document.add_page_break()
        for i, node in enumerate(body):
            render_section(node, 1, str(i + 1))

        # 附录：不参与章节编号，排在正文之后
        for node in appendices:
            document.add_page_break()
            render_section(node, 1, "")

        docx_util.fill_toc_cache(document)
        document.save(output)
        output.seek(0)
