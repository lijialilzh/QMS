#!/usr/bin/env python
# encoding: utf-8

# 软件开发计划服务层，模板参考产品开发计划(serv_pdp_doc)。
# 整份文档以 content(JSON) 存储；导出复用 docx_util.fonted_txt 生成 Word。

import base64
import copy
import logging
import re
from datetime import date, timedelta
from io import BytesIO
from typing import List
from sqlalchemy import delete, func, select
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.sd_doc import SdDoc
from ..model.project_timeline import ProjectTimelineRow, ProjectTimelineCell
from ..model.project_member import ProjectMember
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_sd_doc import SdDocForm
from ..obj.vobj_sd_doc import SdDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from .serv_utils import new_version, sync_file_no_version
from .serv_utils import docx_util
from . import serv_review_util

logger = logging.getLogger(__name__)


# 标准模板默认内容（取自《软件开发计划》模板），新增文档时预填、可改。
# 结构与产品开发计划一致：content.sections 为可递归章节树，节点 {title, body, tables, children, ref_type}。
# ref_type 仅影响导出/编号：cover=封面（居中大标题、不编号），revision=修订记录（不编号），其余为正文章节。
DEFAULT_SD_CONTENT = {
    "sections": [
        {
            "title": "软件开发计划", "ref_type": "cover", "body": "", "children": [],
            "tables": [[
                ["编写部门", "产品开发部", "文件版本", "A0"],
                ["编制人", "", "日期", ""],
                ["审核人", "", "日期", ""],
                ["批准人", "", "日期", ""],
                ["生效日期", "", "", ""],
            ]],
        },
        {
            "title": "文件修订记录", "ref_type": "revision", "body": "", "children": [],
            "tables": [[
                ["修改日期", "版本号", "修订说明", "修订人", "批准人"],
                ["", "", "首次发布", "", ""],
                ["", "", "", "", ""],
                ["", "", "", "", ""],
                ["", "", "", "", ""],
                ["", "", "", "", ""],
            ]],
        },
        {
            "title": "内容简介", "body": "", "tables": [], "children": [
                {"title": "文档目的", "tables": [], "children": [], "body": (
                    "本软件开发计划用于从总体上指导该项目的顺利进行并最终通过评审。"
                    "本软件开发计划面向产品开发部全部开发工程师和测试工程师。"
                )},
                {"title": "文档范围", "tables": [], "children": [], "body": (
                    "本文档描述了在该产品项目中产品开发部需要的人员资源、设备资源等，"
                    "软件开发计划和里程碑的阶段工作任务、时间、项目交付物。"
                )},
            ],
        },
        {
            "title": "项目概况", "body": "", "tables": [], "children": [
                {"title": "项目简介", "ref_type": "prod_name", "body": "产品名称：", "tables": [], "children": []},
                {"title": "项目需求", "body": "详细项目需求请参考《需求规格说明》。", "tables": [], "children": []},
                {"title": "项目开发时间", "ref_type": "prod_cycle", "body": "", "tables": [], "children": []},
            ],
        },
        {
            "title": "开发标准", "tables": [], "children": [],
            "body": (
                "开发过程，参照以下规章制度进行：\n"
                "《软件开发管理制度》\n《开发环境管理制度》\n"
                "《第三方工具操作规范》\n《代码管理制度》"
            ),
        },
        {
            "title": "开发方法", "tables": [], "children": [],
            "body": (
                "在开发过程中，采用经典的瀑布式开发方法，将开发的过程分为需求分析、设计、编码、测试等阶段。\n"
                "在设计和编码过程中，主要使用面向对象的分析方法，进行模块划分、功能拆解等。\n"
                "应当保证软件开发的人员及环境与软件开发要求相适宜。"
            ),
        },
        {
            "title": "人员资源", "ref_type": "personnel", "body": "人员资源如下表所示。", "children": [],
            "tables": [[
                ["编号", "姓名", "所属部门", "角色"],
                ["1", "沈宏", "产品开发部", "经理"],
                ["2", "宁随军", "产品开发部", "开发工程师"],
                ["3", "胡晓光", "产品开发部", "开发工程师"],
                ["4", "徐帅", "产品开发部", "开发工程师"],
            ]],
        },
        {
            "title": "软件资源", "body": "开发过程中用到的软件资源如下表所示。", "children": [],
            "tables": [[
                ["编号", "资源", "资源名称"],
                ["1", "操作系统", "Ubuntu24.04"],
                ["2", "开发语言", "JavaScript，less，Html"],
            ]],
        },
        {
            "title": "软件开发计划及里程碑", "ref_type": "milestone", "body": "软件开发计划里程碑如下表所示。", "children": [],
            "tables": [[
                ["阶段", "任务划分", "负责人", "计划完成时间", "交付物", "使用语言"],
                ["评审开发计划", "编写评审开发计划", "宁随军", "2025年04月07日", "评审记录", ""],
                ["基础代码开发", "DP 架构开发", "李鹏", "2025年05月09日", "/", "python"],
                ["基础代码开发", "Repacs 数据库设计与开发", "成少阳", "2025年05月09日", "/", "python"],
                ["基础代码开发", "DLServer 架构开发", "杨学峰", "2025年05月09日", "/", "python"],
                ["基础代码开发", "NeoViewer 架构开发", "王亮、徐帅", "2025年05月09日", "/", "JavaScript、less、Html"],
                ["模块内开发", "DP 功能实现", "李鹏", "2025年05月26日", "/", "python"],
                ["模块内开发", "Repacs 接口实现", "成少阳", "2025年05月26日", "/", "python"],
                ["模块内开发", "DLServer 功能实现", "杨学峰", "2025年05月26日", "/", "python"],
                ["模块内开发", "NeoViewer 功能实现", "王亮、徐帅", "2025年05月26日", "/", "JavaScript、less、Html"],
                ["整体联调", "各模块进行联合调试", "全体参与", "2025年05月29日", "代码开发结束", "/"],
                ["整体联调", "软件整体联调", "全体参与", "2025年05月29日", "生成安装包", "/"],
            ]],
        },
    ],
}


class Server(object):

    @staticmethod
    def __migrate_cover_table(rows):
        """把旧版 2 列封面表（项目/内容）迁移为新版 4 列结构，保留已填值。"""
        rows = [r for r in (rows or []) if isinstance(r, list)]
        if rows and len(rows[0]) >= 4 and str(rows[0][0]).strip() == "编写部门":
            return rows
        items = [(str(r[0]).strip(), str(r[1]).strip() if len(r) > 1 else "") for r in rows if r]
        def val(label):
            for l, v in items:
                if l == label:
                    return v
            return ""
        dates = [v for l, v in items if l == "日期"]
        d = lambda i: dates[i] if i < len(dates) else ""
        return [
            ["编写部门", val("编写部门") or "产品开发部", "文件版本", val("文件版本") or "A0"],
            ["编制人", val("编制人"), "日期", d(0)],
            ["审核人", val("审核人"), "日期", d(1)],
            ["批准人", val("批准人"), "日期", d(2)],
            ["生效日期", val("生效日期"), "", ""],
        ]

    def __normalize_node(self, node):
        if not isinstance(node, dict):
            return {"title": str(node or ""), "body": "", "tables": [], "children": []}
        result = dict(node)
        result["title"] = str(result.get("title") or "")
        result["body"] = str(result.get("body") or "")
        tables = result.get("tables")
        if not isinstance(tables, list):
            tables = []
        norm_tables = []
        for table in tables:
            if isinstance(table, list):
                norm_tables.append([[str(c) if c is not None else "" for c in (row or [])] for row in table if isinstance(row, list)])
        if result.get("ref_type") == "cover" and norm_tables:
            norm_tables = [self.__migrate_cover_table(norm_tables[0])] + norm_tables[1:]
        result["tables"] = norm_tables
        children = result.get("children")
        result["children"] = [self.__normalize_node(c) for c in children] if isinstance(children, list) else []
        return result

    def __normalize_content(self, content):
        if not isinstance(content, dict) or not isinstance(content.get("sections"), list):
            return copy.deepcopy(DEFAULT_SD_CONTENT)
        return {"sections": [self.__normalize_node(s) for s in content["sections"]]}

    def __autofill_for_export(self, content, obj: SdDocObj, force=False):
        """导出时按产品/时间线/参与人员实时填充。
        force=False（默认）：仅填空，不覆盖已填。
        force=True（切换产品）：强制覆盖时间/参与人员/里程碑/修订/封面，无数据则置空。"""
        sections = (content or {}).get("sections") or []
        prod_id = obj.product_id
        if not prod_id:
            return content
        product = db.session.execute(select(Product).where(Product.id == prod_id)).scalars().first()
        prod_name = (getattr(product, "name", "") or "").strip()
        overall_desc = (getattr(product, "overall_desc", "") or "").strip()
        tl_rows = db.session.execute(
            select(ProjectTimelineRow).where(ProjectTimelineRow.prod_id == prod_id)
        ).scalars().all()
        cell_map = {}
        dev_row_ids = set()
        if tl_rows:
            for c in db.session.execute(
                select(ProjectTimelineCell).where(ProjectTimelineCell.row_id.in_([r.id for r in tl_rows]))
            ).scalars().all():
                cell_map.setdefault(c.row_id, []).append(c.output_result or "")
                # 产品开发活动：输出含「产品开发」但排除「产品开发计划」文档
                if re.search(r"产品开发(?!计划)", str(c.output_result or "")):
                    dev_row_ids.add(c.row_id)

        def to_int(v):
            digits = re.sub(r"[^\d]", "", str(v or ""))
            return int(digits) if digits else None

        date_rows = [r for r in tl_rows if (r.row_type or "date") == "date" and to_int(r.year) and to_int(r.month)]

        def date_key(r):
            return to_int(r.year) * 10000 + to_int(r.month) * 100 + (to_int(r.day) or 0)

        # 产品开发周期：取「产品开发」活动的最早日期=开始、最晚日期=结束
        cycle = ""
        dev_rows = [r for r in date_rows if r.id in dev_row_ids]
        if dev_rows:
            s = min(dev_rows, key=date_key)
            e = max(dev_rows, key=date_key)
            sy, sm, ey, em = to_int(s.year), to_int(s.month), to_int(e.year), to_int(e.month)
            cycle = f"{sy}年{sm}月~{em}月" if sy == ey else f"{sy}年{sm}月~{ey}年{em}月"

        file_rows = [r for r in date_rows if any("软件开发计划" in str(v or "") for v in cell_map.get(r.id, []))]
        file_date = ""
        if file_rows:
            fr = min(file_rows, key=date_key)
            file_date = f"{to_int(fr.year)}年{to_int(fr.month)}月{to_int(fr.day)}日"

        members = db.session.execute(select(ProjectMember).where(ProjectMember.prod_id == prod_id)).scalars().all()
        def find_member(pred):
            for m in members:
                if pred(str(m.role or "")):
                    return (m.name or "").strip()
            return ""
        pm = find_member(lambda r: "产品经理" in r or "经理" in r)
        approver = find_member(lambda r: "负责人" in r)
        # 备注模块映射：备注 前端-NeoViewer / 后端-DP 等 → 取"-"后模块名(小写) → 参与人员姓名
        module_map = {}
        for m in members:
            note = (m.note or "").strip()
            nm = (m.name or "").strip()
            idx = note.rfind("-")
            if not nm or idx < 0:
                continue
            mod = note[idx + 1:].strip().lower()
            if mod:
                module_map.setdefault(mod, []).append(nm)

        # 里程碑表自动填充所需：开发结束日、开发计划评审日、开发工程师
        def fmt_date(r):
            return f"{to_int(r.year)}年{to_int(r.month)}月{to_int(r.day) or 1}日"

        dev_end = fmt_date(max(dev_rows, key=date_key)) if dev_rows else ""

        # 里程碑阶段日期：以产品开发周期(开始~结束)按比例拆分
        # 基础代码开发=开始+1/4、模块开发=开始+3/4、整体联调=结束日、封装安装包=结束日+1天
        ph_base = ph_module = ph_integ = ph_pkg = ""
        if dev_rows:
            s_row = min(dev_rows, key=date_key)
            e_row = max(dev_rows, key=date_key)
            ds = date(to_int(s_row.year), to_int(s_row.month), to_int(s_row.day) or 1)
            de = date(to_int(e_row.year), to_int(e_row.month), to_int(e_row.day) or 1)
            total = (de - ds).days

            def fmt_d(d):
                return f"{d.year}年{d.month}月{d.day}日"

            ph_base = fmt_d(ds + timedelta(days=int(total * 0.25 + 0.5)))
            ph_module = fmt_d(ds + timedelta(days=int(total * 0.75 + 0.5)))
            ph_integ = fmt_d(de)
            ph_pkg = fmt_d(de + timedelta(days=1))

        def latest_date_for(keys):
            matched = [r for r in date_rows if any(any(k in str(v or "") for k in keys) for v in cell_map.get(r.id, []))]
            return fmt_date(max(matched, key=date_key)) if matched else ""

        plan_date = latest_date_for(["软件开发计划"])
        tpm_name = find_member(lambda r: "TPM" in r.upper())

        def strip_num(title):
            return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()

        def fill_milestone(node):
            tables = node.get("tables") or []
            if not tables or not isinstance(tables[0], list) or not tables[0]:
                return
            table = tables[0]
            header = table[0]
            def col_of(name):
                for i, h in enumerate(header):
                    if name in str(h or ""):
                        return i
                return -1
            c_time, c_owner = col_of("计划完成时间"), col_of("负责人")
            c_stage = col_of("阶段")
            for row in table[1:]:
                if not isinstance(row, list):
                    continue
                row_text = " ".join(str(c or "") for c in row)
                is_review = "评审" in row_text and "开发计划" in row_text
                # 计划完成时间（可判定则覆盖）：评审→评审日；基础/模块/联调按周期比例；封装安装包→结束日+1天
                if c_time >= 0 and c_time < len(row):
                    stage = str(row[c_stage]) if 0 <= c_stage < len(row) else row_text
                    dt = ""
                    if is_review:
                        dt = plan_date or dev_end
                    elif re.search(r"安装包|封装", row_text):
                        dt = ph_pkg
                    elif "基础" in stage:
                        dt = ph_base
                    elif "模块" in stage:
                        dt = ph_module
                    elif "联调" in stage or "整体" in stage:
                        dt = ph_integ
                    else:
                        dt = dev_end
                    if force or dt:
                        row[c_time] = dt if not force else (dt or "")
                # 负责人（可判定则覆盖）：评审行→TPM；任务含模块名→该模块参与人员；联调/整体行→全体参与
                if c_owner >= 0 and c_owner < len(row):
                    owner = ""
                    if is_review:
                        owner = tpm_name
                    else:
                        lower = row_text.lower()
                        names = []
                        for k, nms in module_map.items():
                            if k and k in lower:
                                for nm in nms:
                                    if nm and nm not in names:
                                        names.append(nm)
                        if names:
                            owner = "、".join(names)
                        elif re.search(r"联调|整体|全体", row_text):
                            owner = "全体参与"
                    if force or owner:
                        row[c_owner] = owner if not force else (owner or "")

        def fill(node):
            ref = node.get("ref_type")
            title = strip_num(node.get("title"))
            children = node.get("children") or []
            is_overview = ref == "prod_overview" or (title == "产品概况" and not children)
            is_cycle = ref == "prod_cycle" or (title == "项目开发时间" and not children)
            if (ref == "prod_name" or title == "项目简介") and (prod_name or force):
                node["body"] = f"产品名称：{prod_name}" if prod_name else ""
            elif is_overview and (overall_desc or force):
                node["body"] = overall_desc
            elif is_cycle and (cycle or force):
                node["body"] = cycle
            if ref == "revision" or title == "文件修订记录":
                tables = node.get("tables") or []
                if tables and isinstance(tables[0], list):
                    t = tables[0]
                    cols = len(t[0]) if t and t[0] else 5
                    while len(t) < 6:
                        t.append([""] * cols)
                    row = t[1]
                    def set_if(i, val):
                        if force:
                            row[i if i < len(row) else (len(row) - 1)] = val or ""
                        elif val and not str(row[i] if i < len(row) else "").strip():
                            row[i] = val
                    set_if(0, file_date)
                    set_if(1, obj.version)
                    if force or not str(row[2] if len(row) > 2 else "").strip():
                        if len(row) > 2: row[2] = "首次发布"
                    set_if(3, pm)
                    set_if(4, approver)
            if ref == "milestone" or "里程碑" in title:
                fill_milestone(node)
            for child in children:
                fill(child)

        for node in sections:
            fill(node)
        serv_review_util.ensure_review(
            content, "sd", serv_review_util.review_date(prod_id, serv_review_util.REVIEW_DEFS["sd"]["name_keywords"]), prod_id
        )
        serv_review_util.fill_cover_dates(content, serv_review_util.cover_date(prod_id, "sd"), force=force)
        serv_review_util.fill_cover_signers(content, serv_review_util.cover_signers(prod_id, "sd"), force=force)
        return content

    def __to_obj(self, row: SdDoc, product: Product = None):
        obj = SdDocObj(**row.dict())
        obj.content = self.__normalize_content(obj.content)
        serv_review_util.ensure_review(
            obj.content, "sd",
            serv_review_util.review_date(row.product_id, serv_review_util.REVIEW_DEFS["sd"]["name_keywords"]) if row.product_id else "",
            row.product_id,
        )
        serv_review_util.fill_cover_dates(
            obj.content, serv_review_util.cover_date(row.product_id, "sd") if row.product_id else ""
        )
        serv_review_util.fill_cover_signers(
            obj.content, serv_review_util.cover_signers(row.product_id, "sd") if row.product_id else {}
        )
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
        return obj

    async def add_sd_doc(self, form: SdDocForm):
        try:
            sql = select(func.count(SdDoc.id)).where(
                SdDoc.product_id == form.product_id,
                SdDoc.version == form.version,
            )
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = SdDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=SdDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_sd_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: SdDoc = db.session.execute(select(SdDoc).where(SdDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(SdDoc.version).where(SdDoc.product_id == target_pid)).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = SdDoc(
                product_id=target_pid,
                version=version,
                file_no=sync_file_no_version(fromdoc.file_no, version),
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            return Resp.resp_ok(data=SdDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_sd_doc(self, form: SdDocForm):
        try:
            row: SdDoc = db.session.execute(select(SdDoc).where(SdDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def rebind_product(self, id: int, product_id: int):
        """切换产品：更新 product_id 并强制用新产品信息重新获取时间/参与人员/里程碑/封面/修订后保存，返回新 obj。"""
        try:
            row: SdDoc = db.session.execute(select(SdDoc).where(SdDoc.id == id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            product: Product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
            if not product:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            row.product_id = product_id
            content = self.__normalize_content(row.content)
            # 构造临时 obj 供 __autofill_for_export 使用
            tmp_obj = SdDocObj(**row.dict())
            tmp_obj.product_id = product_id
            content = self.__autofill_for_export(content, tmp_obj, force=True)
            row.content = content
            db.session.commit()
            return Resp.resp_ok(data=self.__to_obj(row, product))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_sd_doc(self, id: int):
        db.session.execute(delete(SdDoc).where(SdDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_sd_doc(self, id: int):
        sql = select(SdDoc, Product).join(Product, SdDoc.product_id == Product.id).where(SdDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_sd_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(SdDoc.product_id == product_id)
        if version:
            wheres.append(SdDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(SdDoc.id)).join(Product, SdDoc.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (
            select(SdDoc, Product)
            .join(Product, SdDoc.product_id == Product.id)
            .where(*wheres)
            .order_by(SdDoc.id.desc())
            .offset(page_index * page_size)
            .limit(page_size)
        )
        rows: List[SdDocObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    # ---------------- 导出 Word ----------------
    async def export_sd_doc(self, output, id: int):
        resp = await self.get_sd_doc(id)
        obj: SdDocObj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        c = self.__autofill_for_export(self.__normalize_content(obj.content), obj)
        sections = c.get("sections") or []
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        document.settings.element.append(update_fields)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")
        docx_util.add_page_number_footer(section, obj.file_no or "")

        def write_center_title(text, size=22.0, bold=False):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.right_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, text, font_size=size, bold=bold)

        def add_blank_lines(count):
            for _ in range(max(0, int(count or 0))):
                document.add_paragraph("")

        def add_text(text):
            docx_util.save_txt2docx(str(text or ""), document)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            s = str(text or "")
            if s.startswith("data:image"):
                try:
                    b64 = s.split(",", 1)[1] if "," in s else ""
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = align
                    para.add_run().add_picture(BytesIO(base64.b64decode(b64)), height=Pt(33))
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    return
                except Exception:
                    pass
            cell.text = ""
            lines = s.split("\n")
            for i, line in enumerate(lines):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER if align == WD_ALIGN_PARAGRAPH.CENTER else WD_CELL_VERTICAL_ALIGNMENT.TOP

        def add_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for r_idx, row in enumerate(grid):
                cells = table.add_row().cells
                for c_idx in range(cols):
                    set_cell(cells[c_idx], row[c_idx] if c_idx < len(row) else "", bold=(r_idx == 0))
            document.add_paragraph()

        def add_cover_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for row in grid:
                cells = table.add_row().cells
                for c_idx in range(cols):
                    text = row[c_idx] if c_idx < len(row) else ""
                    set_cell(cells[c_idx], text, bold=(c_idx % 2 == 0), align=WD_ALIGN_PARAGRAPH.CENTER)
                if (str(row[0]).strip() if row else "") == "生效日期" and cols > 2:
                    merged = cells[1]
                    for c_idx in range(2, cols):
                        merged = merged.merge(cells[c_idx])
                    set_cell(merged, row[1] if len(row) > 1 else "", align=WD_ALIGN_PARAGRAPH.CENTER)
            document.add_paragraph()

        def strip_num(title):
            return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()

        def add_body_heading(title, level):
            size = {1: 16.0, 2: 14.0, 3: 12.0}.get(level, 11.0)
            p = document.add_heading("", level=max(1, min(level, 9)))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, title, font_size=size, bold=True)

        def render_body_section(node, level, number=""):
            name = strip_num(node.get("title"))
            heading = f"{number} {name}".strip() if number else name
            add_body_heading(heading, level=max(1, min(level, 9)))
            if (node.get("body") or "").strip():
                add_text(node.get("body"))
            if node.get("ref_type") == "review":
                for t_idx, table in enumerate(node.get("tables") or []):
                    serv_review_util.render_review_grid(document, table, set_cell, merge_col0=(t_idx == 0), merge_full=True)
            else:
                for table in (node.get("tables") or []):
                    add_grid(table)
            idx = 0
            for child in (node.get("children") or []):
                idx += 1
                child_num = f"{number}.{idx}" if number else f"{idx}"
                render_body_section(child, level + 1, child_num)

        cover = next((s for s in sections if s.get("ref_type") == "cover"), None)
        revision = next((s for s in sections if s.get("ref_type") == "revision"), None)
        body = [s for s in sections if s.get("ref_type") not in ("cover", "revision")]

        add_blank_lines(6)
        write_center_title((strip_num(cover.get("title")) if cover else "") or "软件开发计划", size=22.0, bold=True)
        add_blank_lines(4)
        if cover:
            for table in (cover.get("tables") or []):
                add_cover_grid(table)

        document.add_page_break()
        write_center_title("文件修订记录", size=14.0, bold=True)
        add_blank_lines(2)
        if revision:
            for table in (revision.get("tables") or []):
                add_grid(table)

        document.add_page_break()
        write_center_title("目录", size=16.0, bold=True)
        docx_util.insert_toc_field(document)

        document.add_page_break()
        seq = 0
        for node in body:
            if node.get("ref_type") == "review":
                render_body_section(node, 1, "")
            else:
                seq += 1
                render_body_section(node, 1, str(seq))

        document.save(output)
        output.seek(0)
