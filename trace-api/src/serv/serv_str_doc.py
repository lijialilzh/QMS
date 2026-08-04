#!/usr/bin/env python
# encoding: utf-8

# 软件测试报告服务层（开发文件，PDP 风格章节树）。默认内容取自 src-res/str_default_content.json。

import base64
import copy
import json
import logging
import os
import re
from io import BytesIO
from typing import List

from sqlalchemy import delete, func, select
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.str_doc import StrDoc
from ..model.doc_file import DocFile
from ..model.prod_dhf import ProdDhf
from ..model.project_member import ProjectMember
from ..model.test_set import TestSet
from ..model.test_case import TestCase
from ..model.srs_doc import SrsDoc
from ..model.srs_req import SrsReq
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_str_doc import StrDocForm
from ..obj.vobj_str_doc import StrDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from . import serv_review_util
from .serv_utils import new_version, sync_file_no_version
from .serv_utils import docx_util

logger = logging.getLogger(__name__)

_DEFAULT_FILE = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "src-res", "str_default_content.json")
try:
    with open(_DEFAULT_FILE, encoding="utf-8") as _f:
        DEFAULT_STR_CONTENT = json.load(_f)
except Exception:
    DEFAULT_STR_CONTENT = {"sections": []}

# 模板基准产品名（用于全文替换为当前产品名）。
BASE_NAME = "InferOperate Suite"
DOC_NAME = "软件测试报告"
DOC_KEY = "str"


class Server(object):

    def __normalize_node(self, node):
        if not isinstance(node, dict):
            return {"title": str(node or ""), "body": "", "tables": [], "children": []}
        result = dict(node)
        result["title"] = str(result.get("title") or "")
        result["body"] = str(result.get("body") or "")
        tables = result.get("tables")
        if not isinstance(tables, list):
            tables = []
        norm_tables = []
        for table in tables:
            if isinstance(table, list):
                norm_tables.append([[str(c) if c is not None else "" for c in (row or [])] for row in table if isinstance(row, list)])
        result["tables"] = norm_tables
        children = result.get("children")
        result["children"] = [self.__normalize_node(c) for c in children] if isinstance(children, list) else []
        return result

    def __normalize_content(self, content):
        if not isinstance(content, dict) or not isinstance(content.get("sections"), list):
            return copy.deepcopy(DEFAULT_STR_CONTENT)
        return {"sections": [self.__normalize_node(s) for s in content["sections"]]}

    def __replace_name(self, node, base, name, skip_titles=None):
        if not name or base == name:
            return
        if skip_titles and str(node.get("title") or "").strip() in skip_titles:
            return
        if node.get("body"):
            node["body"] = node["body"].replace(base, name)
        for tbl in (node.get("tables") or []):
            for row in tbl:
                for i in range(len(row)):
                    if isinstance(row[i], str) and base in row[i]:
                        row[i] = row[i].replace(base, name)
        for c in (node.get("children") or []):
            self.__replace_name(c, base, name, skip_titles=skip_titles)

    def __dhf_file_no(self, prod_id):
        if not prod_id:
            return ""
        row = db.session.execute(
            select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name == DOC_NAME).order_by(ProdDhf.id.asc())
        ).scalars().first()
        if not row:
            row = db.session.execute(
                select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name.like(f"%{DOC_NAME}%")).order_by(ProdDhf.id.asc())
            ).scalars().first()
        return (row.code or "").strip() if row and row.code else ""

    def __fill_revision(self, content, prod_id, version, force=False):
        """文件修订记录首行：修改日期(评审/封面日期)、版本号、首次发布、修订人(TPM)、批准人(研发负责人)。
        force=True（切换产品）时强制覆盖，无数据则置空。"""
        rev_date = serv_review_util.cover_date(prod_id, DOC_KEY) if prod_id else ""
        reviser = approver = ""
        if prod_id:
            members = db.session.execute(select(ProjectMember).where(ProjectMember.prod_id == prod_id)).scalars().all()
            reviser = next((m.name for m in members if "TPM" in str(m.role or "")), "")
            approver = next((m.name for m in members if "研发负责人" in str(m.role or "")), "")
        for s in (content.get("sections") or []):
            if s.get("ref_type") != "revision":
                continue
            tables = s.get("tables") or []
            if tables and isinstance(tables[0], list) and tables[0]:
                t = tables[0]
                cols = len(t[0]) if isinstance(t[0], list) and t[0] else 5
                while len(t) < 6:
                    t.append([""] * cols)
                if len(t) >= 2 and isinstance(t[1], list) and len(t[1]) >= 3:
                    row = t[1]
                    if force:
                        row[0] = rev_date or ""
                        if len(row) >= 2: row[1] = str(version or "")
                        if len(row) >= 3: row[2] = "首次发布"
                        if len(row) >= 4: row[3] = reviser
                        if len(row) >= 5: row[4] = approver
                    else:
                        if not str(row[0] or "").strip(): row[0] = rev_date
                        if version and len(row) >= 2 and not str(row[1] or "").strip(): row[1] = str(version)
                        if len(row) >= 3 and not str(row[2] or "").strip(): row[2] = "首次发布"
                        if len(row) >= 4 and reviser and not str(row[3] or "").strip(): row[3] = reviser
                        if len(row) >= 5 and approver and not str(row[4] or "").strip(): row[4] = approver
            break
        return content

    TESTITEM_HEADER = ["需求编号", "Case编号", "功能点概述", "用例个数", "是否通过", "备注"]

    def __stage_rows(self, prod_id, stage, prev):
        """按阶段(测试集)取用例，按功能点分组：需求编号/Case编号取段/功能点/用例个数/是否通过/备注。"""
        set_ids = db.session.execute(
            select(TestSet.id).where(TestSet.product_id == prod_id, TestSet.stage == stage)
        ).scalars().all()
        if not set_ids:
            return []
        cases = db.session.execute(
            select(TestCase).where(TestCase.set_id.in_(set_ids)).order_by(TestCase.id)
        ).scalars().all()
        order, gmap = [], {}
        for c in cases:
            func = (c.function or "").strip()
            key = func or (c.code or "").strip()
            if key not in gmap:
                gmap[key] = {"func": func, "codes": [], "srs": (c.srs_code or "").strip()}
                order.append(key)
            code = (c.code or "").strip()
            if code:
                gmap[key]["codes"].append(code)
            if not gmap[key]["srs"]:
                gmap[key]["srs"] = (c.srs_code or "").strip()
        rows = []
        for key in order:
            g = gmap[key]
            cs = sorted(g["codes"])
            case_no = "" if not cs else (cs[0] if len(cs) == 1 else f"{cs[0]} 至 {cs[-1]}")
            p = prev.get(g["func"], {})
            rows.append([g["srs"], case_no, g["func"], str(len(g["codes"])), p.get("pass", "是"), p.get("note", "")])
        return rows

    def __fill_test_items(self, content, prod_id):
        if not prod_id or not isinstance(content, dict):
            return content

        def walk(node):
            if node.get("ref_type") == "test_items":
                prev = {}
                for ch in (node.get("children") or []):
                    for tbl in (ch.get("tables") or []):
                        for r in (tbl or [])[1:]:
                            if isinstance(r, list) and len(r) >= 6 and str(r[2] or "").strip():
                                prev[str(r[2]).strip()] = {"pass": str(r[4] or "").strip() or "是", "note": str(r[5] or "").strip()}
                children = []
                for stage in (node.get("stages") or []):
                    grid = [list(self.TESTITEM_HEADER)] + self.__stage_rows(prod_id, stage, prev)
                    children.append({"title": stage, "body": "", "tables": [grid], "children": []})
                node["children"] = children
                node["tables"] = []
            for ch in (node.get("children") or []):
                walk(ch)

        for s in (content.get("sections") or []):
            walk(s)
        return content

    def __fill_workload(self, content, prod_id):
        """「测试工作量」表：按任务(测试计划/测试用例设计/测试执行/测试总结)从时间线取日期，填开始/结束时间。"""
        if not prod_id or not isinstance(content, dict):
            return content
        doc_date = serv_review_util.cover_date(prod_id, DOC_KEY) or ""
        is_user = DOC_KEY == "utr"
        if is_user:
            range_kw = {"计划": ["用户测试计划"], "用例": ["用户测试用例"],
                        "执行": ["用户测试记录"], "报告": ["用户测试报告"]}
        else:
            range_kw = {"计划": ["软件测试计划", "软开测试计划"],
                        "用例": ["单元测试用例", "集成测试用例", "系统测试用例"],
                        "执行": ["单元测试记录", "集成测试记录", "系统测试记录"],
                        "报告": ["软件测试报告"]}

        def kind_of(task):
            a = str(task or "")
            if "总结" in a or "报告" in a:
                return "报告"
            if "用例" in a or "设计" in a:
                return "用例"
            if "执行" in a:
                return "执行"
            if "计划" in a:
                return "计划"
            return ""

        def _range_days(a, b):
            m1 = re.match(r"(\d{4})\D+(\d{1,2})\D+(\d{1,2})", str(a or ""))
            m2 = re.match(r"(\d{4})\D+(\d{1,2})\D+(\d{1,2})", str(b or ""))
            if not (m1 and m2):
                return 0
            from datetime import date as _d
            try:
                return (_d(int(m2.group(1)), int(m2.group(2)), int(m2.group(3)))
                        - _d(int(m1.group(1)), int(m1.group(2)), int(m1.group(3)))).days + 1
            except Exception:
                return 0

        def walk(node):
            for tbl in (node.get("tables") or []):
                if not tbl or not isinstance(tbl[0], list):
                    continue
                header = [str(c or "") for c in tbl[0]]
                hj = " ".join(header)
                if "开始时间" in hj and "结束时间" in hj:
                    ci_s = next((i for i, h in enumerate(header) if "开始时间" in h), 2)
                    ci_e = next((i for i, h in enumerate(header) if "结束时间" in h), 3)
                    ci_days = next((i for i, h in enumerate(header) if "天数" in h), -1)
                    for r in tbl[1:]:
                        if not isinstance(r, list):
                            continue
                        task = ""
                        for i in range(min(ci_s, len(r))):
                            v = str(r[i] or "").strip()
                            if v and v != "实际":
                                task = v
                        kind = kind_of(task)
                        start, end = ("", "")
                        if kind and kind in range_kw:
                            start, end = serv_review_util.date_range(prod_id, range_kw[kind])
                        start = start or doc_date
                        end = end or start
                        if ci_s < len(r):
                            r[ci_s] = start
                        if ci_e < len(r):
                            r[ci_e] = end
                        if 0 <= ci_days < len(r) and kind and kind in range_kw:
                            _days = serv_review_util.date_count(prod_id, range_kw[kind])
                            if _days:
                                r[ci_days] = "%d天" % _days
            for ch in (node.get("children") or []):
                walk(ch)

        for s in (content.get("sections") or []):
            walk(s)
        return content

    def __fill_compat(self, content, prod_id, product=None):
        """兼容性测试报告附件：产品名称/版本、测试执行人(宋月/孙家旭)、测试时间自动获取；正文产品名替换。"""
        if not isinstance(content, dict):
            return content
        name = (getattr(product, "name", "") or "").strip()
        ver = (getattr(product, "full_version", "") or "").strip()
        rev = ""
        if prod_id:
            rev = serv_review_util.review_date(prod_id, ["兼容性测试", "兼容性"]) or serv_review_util.cover_date(prod_id, DOC_KEY)
        engineer = "宋月" if serv_review_util._before_202509(rev) else "孙家旭"
        pv = (name + "/" + ver) if (name and ver) else name
        srs_leaves = self.__srs_leaves(prod_id)

        def repl(s):
            s = str(s or "")
            return s.replace("InferCare RECIST", name) if name else s

        def relabel(tbl):
            """测试项列表：测试模块列按 SRS 标准需求叶子名(顺序)重贴（跳过表头行）。"""
            if not srs_leaves or not tbl:
                return
            hdr = 0
            for r in tbl:
                if isinstance(r, list) and r and str(r[0] or "").strip() == "测试模块":
                    hdr += 1
                else:
                    break
            data = [r for r in tbl[hdr:] if isinstance(r, list) and r]
            order = []
            for r in data:
                m = str(r[0] or "").strip()
                if m and m not in order:
                    order.append(m)
            mapping = {m: srs_leaves[i] for i, m in enumerate(order) if i < len(srs_leaves)}
            for r in data:
                m = str(r[0] or "").strip()
                if m in mapping:
                    r[0] = mapping[m]

        def proc(node, relabel_on):
            node["body"] = repl(node.get("body"))
            if node.get("body_tail"):
                node["body_tail"] = repl(node.get("body_tail"))
            title = str(node.get("title") or "")
            do_relabel = relabel_on or ("测试项列表" in title) or ("用户端" in title)
            for tbl in (node.get("tables") or []):
                for row in tbl:
                    if not isinstance(row, list):
                        continue
                    for i in range(len(row)):
                        row[i] = repl(row[i])
                    lab = str(row[0] or "").strip()
                    if lab in ("产品名称/版本", "产品名称", "产品名称/型号") and len(row) >= 2 and pv:
                        row[1] = pv
                    elif lab == "测试执行人" and len(row) >= 2:
                        row[1] = engineer
                    elif lab == "测试时间" and len(row) >= 2 and rev:
                        row[1] = rev
                if do_relabel:
                    relabel(tbl)
            for ch in (node.get("children") or []):
                proc(ch, do_relabel)

        for s in (content.get("sections") or []):
            if s.get("ref_type") == "attachment":
                proc(s, False)
        return content

    def __srs_leaves(self, prod_id):
        """该产品 SRS 标准需求(type_code=1)的叶子名(子功能优先，其次功能、模块)，按需求顺序。"""
        if not prod_id:
            return []
        rows = db.session.execute(
            select(SrsReq).join(SrsDoc, SrsReq.doc_id == SrsDoc.id)
            .where(SrsDoc.product_id == prod_id, SrsReq.type_code == "1").order_by(SrsReq.id)
        ).scalars().all()
        placeholders = {"", "/", "\\", "-", "--", "—", "N/A", "n/a", "无", "暂无"}
        leaves = []
        for r in rows:
            leaf = ""
            for val in (getattr(r, "sub_function", ""), getattr(r, "function", ""), getattr(r, "module", "")):
                t = str(val or "").strip()
                if t and t not in placeholders:
                    leaf = t
                    break
            if leaf:
                leaves.append(leaf)
        return leaves

    def __total_cases(self, content):
        total = 0

        def collect(node):
            nonlocal total
            if node.get("ref_type") == "test_items":
                for ch in (node.get("children") or []):
                    for tbl in (ch.get("tables") or []):
                        if not tbl or not isinstance(tbl[0], list):
                            continue
                        header = [str(c or "") for c in tbl[0]]
                        ci_n = next((i for i, h in enumerate(header) if "用例个数" in h), -1)
                        for r in tbl[1:]:
                            if isinstance(r, list) and 0 <= ci_n < len(r):
                                digits = re.sub(r"[^\d]", "", str(r[ci_n]))
                                total += int(digits) if digits else 0
            for ch in (node.get("children") or []):
                collect(ch)

        for s in (content.get("sections") or []):
            collect(s)
        return total

    def __fill_defects(self, content, prod_id):
        """缺陷统计分析表：从「Bug管理及回归测试」文档的缺陷统计自动回填（Bug总数/已解决/遗留/各级别 + 用例总数）。"""
        if not prod_id or not isinstance(content, dict):
            return content
        from .serv_bug_doc import Server as BugServer
        stats = BugServer().stats_for_product(prod_id)
        total_cases = self.__total_cases(content)
        levels = ["一级", "二级", "三级", "四级", "五级"]
        chars = "一二三四五"

        def walk(node):
            for tbl in (node.get("tables") or []):
                if not tbl or not isinstance(tbl[0], list):
                    continue
                header = [str(c or "") for c in tbl[0]]
                labels = [str((r[0] if isinstance(r, list) and r else "") or "") for r in tbl]
                is_defect = ("Bug级别" in " ".join(header)) or any(lb.startswith("Bug总数") for lb in labels)
                if not is_defect:
                    continue
                lvl_col = {}
                for ci in range(2, len(header)):
                    for lv, ch in zip(levels, chars):
                        if ch in header[ci]:
                            lvl_col[lv] = ci
                for r in tbl[1:]:
                    if not isinstance(r, list) or not r:
                        continue
                    lab = str(r[0] or "").strip()

                    def fill_row(by_level, tot):
                        if len(r) > 1:
                            r[1] = str(tot)
                        for lv, ci in lvl_col.items():
                            if ci < len(r):
                                r[ci] = str(by_level.get(lv, 0))

                    if not stats and not lab.startswith("用例总数"):
                        # 无 bug 管理数据：清空 Bug 总数/已解决/遗留行的数值
                        if len(r) > 1:
                            r[1] = ""
                        for lv, ci in lvl_col.items():
                            if ci < len(r):
                                r[ci] = ""
                        continue
                    if lab.startswith("Bug总数"):
                        fill_row(stats.get("by_level", {}), stats.get("total", 0))
                    elif lab.startswith("已解决"):
                        fill_row(stats.get("resolved_by_level", {}), stats.get("resolved", 0))
                    elif lab.startswith("遗留"):
                        fill_row(stats.get("remaining_by_level", {}), stats.get("remaining", 0))
                    elif lab.startswith("用例总数"):
                        if len(r) > 1 and total_cases:
                            r[1] = str(total_cases)
            for ch in (node.get("children") or []):
                walk(ch)

        for s in (content.get("sections") or []):
            walk(s)
        return content

    def __fill_conclusion(self, content, prod_id):
        """测试结论正文：填「在总共N项用例中，通过M项，占P%」，数据来自测试项列表(用例个数/是否通过)。"""
        if not isinstance(content, dict):
            return content
        total = passed = 0

        def collect(node):
            nonlocal total, passed
            if node.get("ref_type") == "test_items":
                for ch in (node.get("children") or []):
                    for tbl in (ch.get("tables") or []):
                        if not tbl or not isinstance(tbl[0], list):
                            continue
                        header = [str(c or "") for c in tbl[0]]
                        ci_n = next((i for i, h in enumerate(header) if "用例个数" in h), -1)
                        ci_p = next((i for i, h in enumerate(header) if "是否通过" in h), -1)
                        for r in tbl[1:]:
                            if not isinstance(r, list):
                                continue
                            digits = re.sub(r"[^\d]", "", str(r[ci_n])) if 0 <= ci_n < len(r) else ""
                            n = int(digits) if digits else 0
                            total += n
                            pv = str(r[ci_p]) if 0 <= ci_p < len(r) else ""
                            if ("通过" in pv) or ("是" == pv.strip()) or (pv.strip() == ""):
                                passed += n
            for ch in (node.get("children") or []):
                collect(ch)

        for s in (content.get("sections") or []):
            collect(s)
        if total <= 0:
            return content
        pct = round(passed * 100 / total)

        def fill(node):
            if "结论" in str(node.get("title", "")) and (node.get("body") or ""):
                node["body"] = re.sub(
                    r"在总共\d*项用例中，通过\d*项，占\d*%",
                    "在总共%d项用例中，通过%d项，占%d%%" % (total, passed, pct),
                    node["body"],
                )
            for ch in (node.get("children") or []):
                fill(ch)

        for s in (content.get("sections") or []):
            fill(s)
        return content

    def __fill_hr(self, content, prod_id, force=False):
        """人力资源表：按「角色」列自动匹配当前产品参与人员，填「资源数量/具体人员」列(N人/姓名)。
        force=False：匹配不到保留原值；force=True：强制覆盖，无人员则置空。"""
        if not isinstance(content, dict):
            return content
        doc_date = serv_review_util.cover_date(prod_id, DOC_KEY) if prod_id else ""
        members = db.session.execute(
            select(ProjectMember).where(ProjectMember.prod_id == prod_id)
        ).scalars().all() if prod_id else []
        primary = "宋月" if serv_review_util._before_202509(doc_date) else "孙家旭"

        def by_kw(pred):
            out = []
            for m in members:
                if pred(str(m.role or "")):
                    nm = (m.name or "").strip()
                    if nm and nm not in out:
                        out.append(nm)
            return out

        def names_for(label, duty):
            lab, d = str(label or ""), str(duty or "")
            if "用户" in lab and "测试" in lab:
                return by_kw(lambda r: "用户测试" in r)
            if "测试" in lab and "用户" not in lab:
                others = [n for n in by_kw(lambda r: ("测试" in r) and ("用户" not in r)) if n not in ("宋月", "孙家旭")]
                return ([primary] + others) if any(k in d for k in ("执行", "记录", "回归", "首轮")) else [primary]
            if "产品经理" in lab:
                return by_kw(lambda r: "产品经理" in r)
            if "开发" in lab and "负责" not in lab:
                return by_kw(lambda r: "开发人员" in r or "开发工程师" in r)
            if "QA" in lab:
                return by_kw(lambda r: "QA" in r or "质量" in r)
            if "RA" in lab:
                return by_kw(lambda r: "RA" in r or "法规" in r)
            if "临床" in lab:
                return by_kw(lambda r: "临床" in r)
            return by_kw(lambda r: bool(lab) and (lab in r or r in lab))

        def walk(node):
            for tbl in (node.get("tables") or []):
                if not tbl or not isinstance(tbl[0], list):
                    continue
                header = [str(c or "") for c in tbl[0]]
                if not (header and "角色" in header[0] and any(("具体人员" in h or "资源数量" in h) for h in header)):
                    continue
                ci_p = next((i for i, h in enumerate(header) if ("具体人员" in h or "资源数量" in h)), 1)
                ci_d = next((i for i, h in enumerate(header) if "职责" in h), 2)
                for r in tbl[1:]:
                    if not isinstance(r, list) or not r:
                        continue
                    nm = names_for(r[0], r[ci_d] if ci_d < len(r) else "")
                    if force:
                        if ci_p < len(r):
                            r[ci_p] = ("%d人/%s" % (len(nm), "、".join(nm))) if nm else ""
                    elif nm and ci_p < len(r):
                        r[ci_p] = "%d人/%s" % (len(nm), "、".join(nm))
            for ch in (node.get("children") or []):
                walk(ch)

        for s in (content.get("sections") or []):
            walk(s)
        return content

    def __fill_refs(self, content, prod_id):
        """引用文档/参考文件表：「作者或来源」列若为已知模板姓名，按角色替换为当前产品对应成员。"""
        if not prod_id or not isinstance(content, dict):
            return content
        doc_date = serv_review_util.cover_date(prod_id, DOC_KEY) or ""
        members = db.session.execute(select(ProjectMember).where(ProjectMember.prod_id == prod_id)).scalars().all()
        primary = "宋月" if serv_review_util._before_202509(doc_date) else "孙家旭"
        name2role = {"吴福乐": "pm", "余航": "pm", "杨静": "pm", "沈宏": "devlead", "宁随军": "tpm",
                     "林金贵": "qa", "夏晨": "pdir", "张淑芳": "ra", "宋月": "test", "王小敏": "test",
                     "孙家旭": "test", "徐冕": "utest", "罗鹏宇": "utest", "刘锦龙": "utest",
                     "徐学强": "utest", "刘娜": "utest"}

        def first(pred):
            for m in members:
                if pred(str(m.role or "")):
                    nm = (m.name or "").strip()
                    if nm:
                        return nm
            return ""

        def resolve(role):
            return {
                "test": primary,
                "utest": first(lambda r: "用户测试" in r),
                "pm": first(lambda r: "产品经理" in r),
                "devlead": first(lambda r: "研发负责人" in r),
                "tpm": first(lambda r: "TPM" in r),
                "qa": first(lambda r: "QA" in r or "质量" in r),
                "pdir": first(lambda r: "产品负责人" in r or "产品总监" in r),
                "ra": first(lambda r: "RA" in r or "法规" in r),
            }.get(role, "")

        def walk(node):
            for tbl in (node.get("tables") or []):
                if not tbl or not isinstance(tbl[0], list):
                    continue
                header = [str(c or "") for c in tbl[0]]
                if not (any("文档名称" in h for h in header) and any(("作者" in h or "来源" in h) for h in header)):
                    continue
                ci_a = next((i for i, h in enumerate(header) if ("作者" in h or "来源" in h)), 1)
                for r in tbl[1:]:
                    if not isinstance(r, list) or ci_a >= len(r):
                        continue
                    cur = str(r[ci_a] or "").strip()
                    role = name2role.get(cur)
                    if role:
                        nm = resolve(role)
                        if nm:
                            r[ci_a] = nm
            for ch in (node.get("children") or []):
                walk(ch)

        for s in (content.get("sections") or []):
            walk(s)
        return content

    def __autofill(self, content, prod_id, product=None, version="", force=False):
        if not isinstance(content, dict):
            return content
        if product and product.name:
            for s in (content.get("sections") or []):
                self.__replace_name(s, BASE_NAME, product.name)
                # 兼容性测试报告相关章节用 InferCare RECIST 作为产品名占位，也替换为当前产品名
                self.__replace_name(s, "InferCare RECIST", product.name)
        self.__fill_test_items(content, prod_id)
        self.__fill_hr(content, prod_id, force=force)
        self.__fill_refs(content, prod_id)
        self.__fill_conclusion(content, prod_id)
        self.__fill_defects(content, prod_id)
        self.__fill_workload(content, prod_id)
        self.__fill_compat(content, prod_id, product)
        self.__fill_revision(content, prod_id, version, force=force)
        serv_review_util.ensure_review(
            content, DOC_KEY,
            serv_review_util.review_date(prod_id, serv_review_util.REVIEW_DEFS[DOC_KEY]["name_keywords"]) if prod_id else "",
            prod_id,
        )
        serv_review_util.fill_cover_dates(content, serv_review_util.cover_date(prod_id, DOC_KEY) if prod_id else "", force=force)
        serv_review_util.fill_cover_signers(content, serv_review_util.cover_signers(prod_id, DOC_KEY) if prod_id else {}, force=force)
        # 附件(兼容性测试报告)排到评审记录之后
        secs = content.get("sections") or []
        atts = [s for s in secs if s.get("ref_type") == "attachment"]
        if atts:
            content["sections"] = [s for s in secs if s.get("ref_type") != "attachment"] + atts
        return content

    def __to_obj(self, row: StrDoc, product: Product = None):
        obj = StrDocObj(**row.dict())
        obj.content = self.__autofill(self.__normalize_content(obj.content), row.product_id, product, row.version)
        if not (obj.file_no or "").strip():
            obj.file_no = self.__dhf_file_no(row.product_id)
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
        return obj

    async def add_str_doc(self, form: StrDocForm):
        try:
            sql = select(func.count(StrDoc.id)).where(StrDoc.product_id == form.product_id, StrDoc.version == form.version)
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = StrDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=StrDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_str_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: StrDoc = db.session.execute(select(StrDoc).where(StrDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(StrDoc.version).where(StrDoc.product_id == target_pid)).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = StrDoc(
                product_id=target_pid, version=version,
                file_no=sync_file_no_version((fromdoc.file_no or "").strip() or self.__dhf_file_no(target_pid), version) or None,
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            if target_pid != fromdoc.product_id:
                await self.rebind_product(newdoc.id, target_pid)
            return Resp.resp_ok(data=StrDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_str_doc(self, form: StrDocForm):
        try:
            row: StrDoc = db.session.execute(select(StrDoc).where(StrDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def rebind_product(self, id: int, product_id: int):
        """切换产品：更新 product_id 并强制用新产品信息重新获取封面/修订/产品名后保存，返回新 obj。"""
        try:
            row: StrDoc = db.session.execute(select(StrDoc).where(StrDoc.id == id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            old_product: Product = db.session.execute(select(Product).where(Product.id == row.product_id)).scalars().first() if row.product_id else None
            product: Product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
            if not product:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            db.session.execute(delete(StrDoc).where(StrDoc.product_id == product_id, StrDoc.version == row.version, StrDoc.id != id))
            content = self.__normalize_content(row.content)
            # 重置含产品名/完整版本的固定章节为模板原值（恢复基准名+原版本号，避免被旧产品污染）
            tpl = copy.deepcopy(DEFAULT_STR_CONTENT) if isinstance(DEFAULT_STR_CONTENT, dict) else {"sections": []}
            tpl_map = {}
            fixed_titles = {"用户端测试"}
            def collect_tpl(node):
                key = str(node.get("title") or "").strip()
                body = str(node.get("body") or "")
                caps = node.get("table_captions") or []
                tables = node.get("tables") or []
                if BASE_NAME in body or "完整版本" in body or "InferCare RECIST" in body or key in fixed_titles or any(BASE_NAME in str(r) for tbl in tables for r in tbl) or any("1.1.0.1" in str(r) for tbl in tables for r in tbl) or any("1.1.0.1" in str(c) for c in caps):
                    tpl_map[key] = {"body": node.get("body", ""), "tables": copy.deepcopy(tables), "table_captions": copy.deepcopy(caps)}
                for c in (node.get("children") or []):
                    collect_tpl(c)
            for s in (tpl.get("sections") or []):
                collect_tpl(s)
            def reset_fixed(node):
                key = str(node.get("title") or "").strip()
                if key in tpl_map:
                    node["body"] = tpl_map[key]["body"]
                    node["tables"] = copy.deepcopy(tpl_map[key]["tables"])
                    if tpl_map[key].get("table_captions"):
                        node["table_captions"] = copy.deepcopy(tpl_map[key]["table_captions"])
                for c in (node.get("children") or []):
                    reset_fixed(c)
            for s in (content.get("sections") or []):
                reset_fixed(s)
            # 完整版本号更新为新产品版本（body + tables + table_captions）
            # 跳过：版本信息表的"最终安装包路径"行（保留模板包名）、用户端测试章节表头（浏览器版本号保留）
            new_full_version = (product.full_version or "").strip()
            def update_version(node):
                node_title = str(node.get("title") or "").strip()
                if node.get("body") and "完整版本" in str(node.get("body")):
                    node["body"] = re.sub(r"完整版本：[^\n]*", f"完整版本：{new_full_version}", str(node["body"]))
                # 用户端测试章节的表格不替换版本号（表头是浏览器版本号，保留模板）
                if "用户端测试" not in node_title:
                    for tbl in (node.get("tables") or []):
                        for r in tbl:
                            # 版本信息表跳过"最终安装包路径"行（保留模板包名）
                            if "版本信息" in node_title and any("安装包路径" in str(c) for c in r):
                                continue
                            for i in range(len(r)):
                                if isinstance(r[i], str) and re.search(r"\d+\.\d+\.\d+\.\d+", r[i]):
                                    r[i] = re.sub(r"\d+\.\d+\.\d+\.\d+", new_full_version, r[i])
                caps = node.get("table_captions") or []
                if caps:
                    node["table_captions"] = [re.sub(r"\d+\.\d+\.\d+\.\d+", new_full_version, str(c)) if re.search(r"\d+\.\d+\.\d+\.\d+", str(c)) else c for c in caps]
                for c in (node.get("children") or []):
                    update_version(c)
            for s in (content.get("sections") or []):
                update_version(s)
            row.product_id = product_id
            # 测试分布图：未获取时从该产品 doc_file（img_flow）获取，已获取的保留
            # 缺陷统计分析分布图：有 bug 管理数据才获取，没 bug 清空
            from .serv_bug_doc import Server as BugServer
            has_bug = bool(BugServer().stats_for_product(product_id))
            def fill_test_chart(nodes):
                for n in nodes:
                    title = str(n.get("title") or "").strip()
                    if title == "测试分布图":
                        imgs = n.get("images") or []
                        if not imgs:
                            row_df = db.session.execute(
                                select(DocFile).where(DocFile.product_id == product_id, DocFile.category == "img_flow").order_by(DocFile.id)
                            ).scalars().first()
                            if row_df and row_df.file_url:
                                url = str(row_df.file_url).strip()
                                if url.startswith("data.trace/"):
                                    url = "/" + url
                                n["images"] = [url]
                    elif title == "缺陷统计分析":
                        # 有 bug 管理数据才生成分布图，没 bug 提示"还未上传Bug管理及回归测试"
                        if has_bug:
                            row_df = db.session.execute(
                                select(DocFile).where(DocFile.product_id == product_id, DocFile.category == "img_flow").order_by(DocFile.id)
                            ).scalars().first()
                            if row_df and row_df.file_url:
                                url = str(row_df.file_url).strip()
                                if url.startswith("data.trace/"):
                                    url = "/" + url
                                n["images"] = [url]
                            n["body"] = ""
                        else:
                            n["images"] = []
                            n["body"] = "还未上传Bug管理及回归测试"
                    for c in (n.get("children") or []):
                        fill_test_chart([c])
            fill_test_chart(content.get("sections") or [])
            content = self.__autofill(content, product_id, product, row.version, force=True)
            row.content = content
            db.session.commit()
            return Resp.resp_ok(data=self.__to_obj(row, product))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_str_doc(self, id: int):
        db.session.execute(delete(StrDoc).where(StrDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_str_doc(self, id: int):
        sql = select(StrDoc, Product).join(Product, StrDoc.product_id == Product.id).where(StrDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_str_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(StrDoc.product_id == product_id)
        if version:
            wheres.append(StrDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(StrDoc.id)).join(Product, StrDoc.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (select(StrDoc, Product).join(Product, StrDoc.product_id == Product.id).where(*wheres)
               .order_by(StrDoc.id.desc()).offset(page_index * page_size).limit(page_size))
        rows: List[StrDocObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    # ---------------- 导出 Word（PDP 风格章节树） ----------------
    async def export_str_doc(self, output, id: int):
        resp = await self.get_str_doc(id)
        obj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        c = self.__normalize_content(obj.content)
        sections = c.get("sections") or []
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        document.settings.element.append(update_fields)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")
        docx_util.add_page_number_footer(section, obj.file_no or "")

        def add_blank_lines(n):
            for _ in range(max(0, int(n or 0))):
                document.add_paragraph("")

        def write_center_title(text, size=22.0, bold=False):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            docx_util.fonted_txt(p, str(text or ""), font_size=size, bold=bold)

        def add_text(text):
            docx_util.save_txt2docx(str(text or ""), document)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            s = str(text or "")
            if s.startswith("data:image"):
                try:
                    b64 = s.split(",", 1)[1] if "," in s else ""
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = align
                    para.add_run().add_picture(BytesIO(base64.b64decode(b64)), height=Pt(33))
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    return
                except Exception:
                    pass
            cell.text = ""
            for i, line in enumerate(s.split("\n")):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER if align == WD_ALIGN_PARAGRAPH.CENTER else WD_CELL_VERTICAL_ALIGNMENT.TOP

        def _set_fixed_widths(table, widths_dxa):
            table.autofit = False
            table.allow_autofit = False
            tbl_pr = table._tbl.tblPr
            layout = tbl_pr.find(qn("w:tblLayout"))
            if layout is None:
                layout = OxmlElement("w:tblLayout")
                tbl_pr.append(layout)
            layout.set(qn("w:type"), "fixed")
            grid_el = table._tbl.find(qn("w:tblGrid"))
            if grid_el is not None:
                for gc in list(grid_el):
                    grid_el.remove(gc)
                for w in widths_dxa:
                    gc = OxmlElement("w:gridCol")
                    gc.set(qn("w:w"), str(w))
                    grid_el.append(gc)
            # 逐格写 tcW，确保固定列宽生效
            for row in table.rows:
                for i, w in enumerate(widths_dxa):
                    if i < len(row.cells):
                        tcpr = row.cells[i]._tc.get_or_add_tcPr()
                        tcw = tcpr.find(qn("w:tcW"))
                        if tcw is None:
                            tcw = OxmlElement("w:tcW")
                            tcpr.append(tcw)
                        tcw.set(qn("w:w"), str(w))
                        tcw.set(qn("w:type"), "dxa")

        def add_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for r_idx, row in enumerate(grid):
                cells = table.add_row().cells
                for c_idx in range(cols):
                    set_cell(cells[c_idx], row[c_idx] if c_idx < len(row) else "", bold=(r_idx == 0))
            # SCI 清单表：固定列宽，SCI名字/类型不换行、存储地点较宽
            first = str(grid[0][0]).strip() if grid and grid[0] else ""
            if first == "SCI名字":
                if cols == 5:
                    _set_fixed_widths(table, [1700, 1100, 2000, 1500, 3000])
                elif cols == 4:
                    _set_fixed_widths(table, [1700, 2200, 1500, 3600])
            elif first == "类别" and cols >= 3:
                # 配置项存储地址表：类别列收窄、存储路径较宽；首列连续相同「类别」纵向合并；文字左对齐
                if cols == 3:
                    _set_fixed_widths(table, [1300, 1300, 5000])
                rows = table.rows
                n = len(rows)
                r = 1
                while r < n:
                    val = str((grid[r][0] if r < len(grid) and grid[r] else "") or "").strip()
                    if not val:
                        r += 1
                        continue
                    r2 = r
                    while r2 + 1 < n and str((grid[r2 + 1][0] if grid[r2 + 1] else "") or "").strip() == val:
                        r2 += 1
                    if r2 > r:
                        merged = rows[r].cells[0].merge(rows[r2].cells[0])
                        set_cell(merged, val, align=WD_ALIGN_PARAGRAPH.LEFT)
                        merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    r = r2 + 1
            elif first == "测试模块":
                # 兼容性测试项：表头合并 + 首列「测试模块」连续相同纵向合并
                rows = table.rows
                n = len(rows)
                hdr = 0
                for gr in grid:
                    if gr and str(gr[0] or "").strip() == "测试模块":
                        hdr += 1
                    else:
                        break
                # 表头：测试模块/测试点跨表头行；版本组横向合并
                if hdr >= 2:
                    for col in (0, 1):
                        try:
                            mh = rows[0].cells[col].merge(rows[hdr - 1].cells[col])
                            set_cell(mh, str(grid[0][col] if col < len(grid[0]) else ""), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                            mh.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        except Exception:
                            pass
                if grid and len(grid[0]) > 2:
                    ncol = len(grid[0])
                    c = 2
                    while c < ncol:
                        v = str(grid[0][c] or "")
                        c2 = c
                        while c2 + 1 < ncol and str(grid[0][c2 + 1] or "") == v:
                            c2 += 1
                        if c2 > c:
                            try:
                                mh = rows[0].cells[c].merge(rows[0].cells[c2])
                                set_cell(mh, v, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                                mh.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            except Exception:
                                pass
                        c = c2 + 1
                r = hdr
                while r < n:
                    val = str((grid[r][0] if r < len(grid) and grid[r] else "") or "").strip()
                    if not val:
                        r += 1
                        continue
                    r2 = r
                    while r2 + 1 < n and str((grid[r2 + 1][0] if grid[r2 + 1] else "") or "").strip() == val:
                        r2 += 1
                    if r2 > r:
                        merged = rows[r].cells[0].merge(rows[r2].cells[0])
                        set_cell(merged, val, align=WD_ALIGN_PARAGRAPH.CENTER)
                        merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    r = r2 + 1
            elif first in ("测试目标", "测试方法和技术") and cols == 2:
                _set_fixed_widths(table, [2200, 7100])
            elif first == "配置":
                if cols == 2:
                    _set_fixed_widths(table, [1900, 7400])
                elif cols == 3:
                    _set_fixed_widths(table, [1900, 3700, 3700])
            document.add_paragraph()

        def add_cover_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for row in grid:
                cells = table.add_row().cells
                for c_idx in range(cols):
                    text = row[c_idx] if c_idx < len(row) else ""
                    set_cell(cells[c_idx], text, bold=(c_idx % 2 == 0), align=WD_ALIGN_PARAGRAPH.CENTER)
                if (str(row[0]).strip() if row else "") == "生效日期" and cols > 2:
                    merged = cells[1]
                    for c_idx in range(2, cols):
                        merged = merged.merge(cells[c_idx])
                    set_cell(merged, row[1] if len(row) > 1 else "", align=WD_ALIGN_PARAGRAPH.CENTER)
            document.add_paragraph()

        def strip_num(title):
            return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()

        def add_body_heading(title, level):
            size = {1: 16.0, 2: 14.0, 3: 12.0}.get(level, 11.0)
            p = document.add_heading("", level=max(1, min(level, 9)))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, title, font_size=size, bold=True)

        def render_body_section(node, level, number=""):
            name = strip_num(node.get("title"))
            heading = f"{number} {name}".strip() if number else name
            add_body_heading(heading, level=max(1, min(level, 9)))
            if (node.get("body") or "").strip():
                add_text(node.get("body"))
            if node.get("ref_type") == "review":
                for t_idx, table in enumerate(node.get("tables") or []):
                    serv_review_util.render_review_grid(document, table, set_cell, merge_col0=(t_idx == 0), merge_full=True)
            else:
                caps = node.get("table_captions") or []
                for t_idx, table in enumerate(node.get("tables") or []):
                    cap = caps[t_idx] if t_idx < len(caps) else ""
                    if str(cap or "").strip():
                        add_text(cap)
                    add_grid(table)
                # 分布图等节点图片（data URL）内嵌，排在表格之后、表格后正文之前
                for _img in (node.get("charts") or []):
                    if isinstance(_img, str) and _img.startswith("data:image"):
                        try:
                            _p = document.add_paragraph()
                            _p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            _p.add_run().add_picture(BytesIO(base64.b64decode(_img.split(",", 1)[1])), width=Inches(6.2))
                        except Exception:
                            pass
                if str(node.get("body_tail") or "").strip():
                    add_text(node.get("body_tail"))
            idx = 0
            for child in (node.get("children") or []):
                idx += 1
                child_num = f"{number}.{idx}" if number else f"{idx}"
                render_body_section(child, level + 1, child_num)

        cover = next((s for s in sections if s.get("ref_type") == "cover"), None)
        revision = next((s for s in sections if s.get("ref_type") == "revision"), None)
        body = [s for s in sections if s.get("ref_type") not in ("cover", "revision")]

        add_blank_lines(6)
        write_center_title((strip_num(cover.get("title")) if cover else "") or DOC_NAME, size=22.0, bold=True)
        add_blank_lines(4)
        if cover:
            for table in (cover.get("tables") or []):
                add_cover_grid(table)

        document.add_page_break()
        write_center_title("文件修订记录", size=14.0, bold=True)
        add_blank_lines(2)
        if revision:
            for table in (revision.get("tables") or []):
                add_grid(table)

        document.add_page_break()
        write_center_title("目录", size=16.0, bold=True)
        docx_util.insert_toc_field(document)

        document.add_page_break()
        seq = 0
        for node in body:
            if node.get("ref_type") in ("review", "attachment"):
                render_body_section(node, 1, "")
            else:
                seq += 1
                render_body_section(node, 1, str(seq))

        docx_util.fill_toc_cache(document)
        document.save(output)
        output.seek(0)
