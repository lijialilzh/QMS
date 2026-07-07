#!/usr/bin/env python
# encoding: utf-8

# 产品开发计划服务层，详见 docs/function_docs/52_产品开发计划.md。
# 整份文档以 content(JSON) 存储；导出复用 docx_util.fonted_txt 生成 Word。

import base64
import copy
import logging
import re
from io import BytesIO
from typing import List
from sqlalchemy import delete, func, select
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.pdp_doc import PdpDoc
from ..model.project_timeline import ProjectTimelineRow, ProjectTimelineCell
from ..model.project_member import ProjectMember
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_pdp_doc import PdpDocForm
from ..obj.vobj_pdp_doc import PdpDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from .serv_utils import new_version, sync_file_no_version
from .serv_utils import docx_util
from . import serv_review_util

logger = logging.getLogger(__name__)


# 标准模板默认内容（取自《产品开发计划》模板），新增文档时预填、可改。
# 采用「目录树」结构：content.sections 为可递归章节树，节点 {title, body, tables, children, ref_type}。
# ref_type 仅影响导出/编号：cover=封面（居中大标题、不编号），revision=修订记录（不编号），其余为正文章节。
# 表格(tables)为二维数组列表，每个表第一行为表头；用户可在前端任意增删章节/行/列。

DEFAULT_PDP_CONTENT = {
    "sections": [
        {
            "title": "产品开发计划", "ref_type": "cover", "body": "", "children": [],
            "tables": [[
                ["编写部门", "产品部", "文件版本", "A0"],
                ["编制人", "", "日期", ""],
                ["审核人", "", "日期", ""],
                ["批准人", "", "日期", ""],
                ["生效日期", "", "", ""],
            ]],
        },
        {
            "title": "文件修订记录", "ref_type": "revision", "body": "", "children": [],
            "tables": [[
                ["修改日期", "版本号", "修订说明", "修订人", "批准人"],
                ["", "", "首次发布", "", ""],
                ["", "", "", "", ""],
                ["", "", "", "", ""],
                ["", "", "", "", ""],
                ["", "", "", "", ""],
            ]],
        },
        {
            "title": "内容简介", "body": "", "tables": [], "children": [
                {"title": "文档目的", "tables": [], "children": [], "body": (
                    "为了保证团队按时保质地完成目标，便于团队成员更好地了解情况，使工作开展的各个过程合理有序，"
                    "有必要以文件化的形式，把产品开发生命周期内的工作任务范围、各项工作的任务分解、团队组织结构、"
                    "各团队成员的工作责任等内容以书面的方式描述出来，作为团队成员以及干系人之间的共识与约定，"
                    "产品开发生命周期内的所有活动的行动基础，团队开展和检查产品开发工作的依据。"
                )},
                {"title": "文档范围", "tables": [], "children": [], "body": (
                    "本文档描述了在该产品项目中需要的人员资源、设备资源等，产品开发计划和里程碑的阶段工作任务、"
                    "时间、交付物、相关人员的角色和职责等。应当保证软件开发和测试的人员及环境与软件开发要求相适宜。"
                )},
            ],
        },
        {
            "title": "产品概况", "body": "", "tables": [], "children": [
                {"title": "产品简介", "ref_type": "prod_name", "body": "产品名称：", "tables": [], "children": []},
                {"title": "产品概况", "ref_type": "prod_overview", "body": "", "tables": [], "children": []},
                {"title": "产品开发周期", "ref_type": "prod_cycle", "body": "", "tables": [], "children": []},
            ],
        },
        {
            "title": "人员资源", "ref_type": "personnel", "body": "", "children": [],
            "tables": [[
                ["人数", "所属部门", "人员编制", "角色/岗位", "职责"],
                ["1", "产品部", "余航", "产品经理", "负责需求的收集、整理和分析\n对需求基线后的需求变更进行控制\n安排产品资源；\n组织协调工作；\n组织产品开发的实施；\n监督总体进度；\n提供后勤支持。"],
                ["8", "产品开发部", "宁随军", "软件开发工程师", "按照需求文档进行系统规划和设计；\n设计编码开发。"],
                ["2", "产品开发部", "宋月", "软件测试工程师", "单元测试、集成测试、系统测试、用户测试"],
                ["1", "模型部", "王瑜", "算法研究员及算法工程师", "按照需求文档进行算法模块的规划和设计，包括算法概要设计和详细设计；\n设计编码开发；\n数据处理、数据标注规则定义、数据质控等"],
                ["1", "质量部", "林金贵", "质量工程师", "负责管控风险管理的过程符合风险管理控制程序的规定"],
                ["1", "生产部", "陈福临", "生产人员", "负责转换阶段的生产工作"],
                ["1", "/", "杨冰", "临床专家（或者有临床背景的客户经理）", "负责从临床角度和用户的角度评审危害源识别的正确性和风险控制措施的合理性，确保剩余风险可以接受"],
                ["1", "/", "李佳励", "安全工程师（或者有风险管理背景软件开发工程师）", "负责风险管理文件的撰写"],
                ["3", "/", "徐学强、刘锦龙、杨冰", "用户测试人员", "负责产品的用户测试相关工作"],
            ]],
        },
        {
            "title": "设备资源", "body": "", "children": [],
            "tables": [[
                ["设备及用途", "设备名称", "数量", "设备说明"],
                ["操作系统", "Ubuntu24.04、Windows 10", "3", "现有满足"],
                ["开发语言", "Python、JavaScript、less、html", "4", "现有满足"],
                ["数据库", "PostgreSQL", "2", "现有满足"],
                ["开发工具", "VS Code", "1", "现有满足"],
                ["测试工具", "Chrome、JMeter、Nmap", "3", "现有满足"],
                ["配置管理工具", "Jira、GitLab、NextCloud、Nas", "4", "现有满足"],
                ["开发设备", "计算机", "7", "现有满足"],
                ["测试设备", "计算机", "2", "现有满足"],
                ["生产设备", "计算机", "1", "现有满足"],
                ["检验设备", "计算机", "1", "现有满足"],
            ]],
        },
        {
            "title": "产品开发计划及里程碑", "body": "", "children": [],
            "tables": [[
                ["阶段", "活动", "评审部门", "负责部门", "计划完成时间", "主要活动和任务", "阶段性交付物"],
                ["可行性研究与定义阶段", "市场调研", "研发部、质量部", "产品部", "", "通过调研市场需求明确产品功能定位、预期用途、经济及社会意义，判断产品开发是否可行。", "《产品立项报告》"],
                ["可行性研究与定义阶段", "策划", "研发部、质量部", "产品部", "", "确定产品开发计划", "《产品开发计划》"],
                ["可行性研究与定义阶段", "需求调研", "研发部、质量部", "产品部", "", "调研产品需求并进行需求分析", "《需求规格说明》《初步危害分析清单》《软件配置管理计划》"],
                ["可行性研究与定义阶段", "输入评审", "研发部、质量部", "产品部", "", "评审产品的技术指标、安全要求、法规要求、风险信息、网络安全信息。", "《设计开发输入清单》评审记录（设计开发输入）"],
                ["产品实现阶段", "产品开发", "/", "产品开发部", "", "实现除模型以外的软件功能", "产品"],
                ["产品实现阶段", "产品测试", "产品开发部、产品部", "产品开发部", "", "进行软件的单元、集成、系统测试，以保证软件的正常运行。", "《软件测试计划》测试用例、测试记录《软件测试报告》"],
                ["产品实现阶段", "用户测试", "产品开发部、产品部", "产品开发部", "", "进行产品的用户测试", "《用户测试计划》《用户测试用例》《用户测试记录》《用户测试报告》"],
                ["产品实现阶段", "输出评审", "研发部、质量部", "研发部", "", "输出保证产品质量的生产、服务相关信息，如：工艺流程图、采购物资明细、图纸、作业指导书、检验规范、技术要求、用户说明书、包装标签的样稿、软件的开发和测试。", "《可追溯性分析报告》《用户说明书》《安装维护手册》《软件配置状态报告》《软件维护计划》《生产工艺流程图》《产品标签样稿》《现场测试规程》《采购要求及产品BOM》《成品检验规程》《产品技术要求》"],
                ["产品转换阶段", "发布", "研发部、质量部", "产品部", "", "发布测试通过的产品。", "《产品发布说明》"],
                ["产品转换阶段", "转换", "质量部、研发部、执行部", "质量部", "", "通过确立文件规定、人员培训、设备验证等手段，实现研发阶段到规模化生产阶段。", "《生产作业指导书》《成品检验规程》"],
            ]],
        },
        {
            "title": "相关计划", "tables": [], "children": [],
            "body": (
                "配置管理、风险管理、维护、产品开发等计划详见以下计划：\n"
                "《软件配置管理计划》\n《软件维护计划》\n《软件开发计划》\n"
                "《软件测试计划》\n《用户测试计划》\n《风险管理计划》"
            ),
        },
    ],
}


class Server(object):

    @staticmethod
    def __migrate_cover_table(rows):
        """把旧版 2 列封面表（项目/内容）迁移为新版 4 列结构，保留已填值。"""
        rows = [r for r in (rows or []) if isinstance(r, list)]
        # 已是新结构：首行 4 列且首格为「编写部门」
        if rows and len(rows[0]) >= 4 and str(rows[0][0]).strip() == "编写部门":
            return rows
        items = [(str(r[0]).strip(), str(r[1]).strip() if len(r) > 1 else "") for r in rows if r]
        def val(label):
            for l, v in items:
                if l == label:
                    return v
            return ""
        dates = [v for l, v in items if l == "日期"]
        d = lambda i: dates[i] if i < len(dates) else ""
        return [
            ["编写部门", val("编写部门") or "产品部", "文件版本", val("文件版本") or "A0"],
            ["编制人", val("编制人"), "日期", d(0)],
            ["审核人", val("审核人"), "日期", d(1)],
            ["批准人", val("批准人"), "日期", d(2)],
            ["生效日期", val("生效日期"), "", ""],
        ]

    def __normalize_node(self, node):
        if not isinstance(node, dict):
            return {"title": str(node or ""), "body": "", "tables": [], "children": []}
        result = dict(node)
        result["title"] = str(result.get("title") or "")
        result["body"] = str(result.get("body") or "")
        tables = result.get("tables")
        if not isinstance(tables, list):
            tables = []
        norm_tables = []
        for table in tables:
            if isinstance(table, list):
                norm_tables.append([[str(c) if c is not None else "" for c in (row or [])] for row in table if isinstance(row, list)])
        # 封面表统一迁移为新版 4 列结构
        if result.get("ref_type") == "cover" and norm_tables:
            norm_tables = [self.__migrate_cover_table(norm_tables[0])] + norm_tables[1:]
        result["tables"] = norm_tables
        children = result.get("children")
        result["children"] = [self.__normalize_node(c) for c in children] if isinstance(children, list) else []
        return result

    def __normalize_content(self, content):
        if not isinstance(content, dict) or not isinstance(content.get("sections"), list):
            return copy.deepcopy(DEFAULT_PDP_CONTENT)
        return {"sections": [self.__normalize_node(s) for s in content["sections"]]}

    def __autofill_for_export(self, content, obj: PdpDocObj):
        """导出时按产品/时间线/参与人员实时填充，保证导出与「新增/编辑」一致（仅填空、不覆盖已填）。"""
        sections = (content or {}).get("sections") or []
        prod_id = obj.product_id
        if not prod_id:
            return content
        # 产品信息
        product = db.session.execute(select(Product).where(Product.id == prod_id)).scalars().first()
        prod_name = (getattr(product, "name", "") or "").strip()
        overall_desc = (getattr(product, "overall_desc", "") or "").strip()
        # 时间线
        tl_rows = db.session.execute(
            select(ProjectTimelineRow).where(ProjectTimelineRow.prod_id == prod_id)
        ).scalars().all()
        cell_map = {}
        if tl_rows:
            for c in db.session.execute(
                select(ProjectTimelineCell).where(ProjectTimelineCell.row_id.in_([r.id for r in tl_rows]))
            ).scalars().all():
                cell_map.setdefault(c.row_id, []).append(c.output_result or "")

        def to_int(v):
            digits = re.sub(r"[^\d]", "", str(v or ""))
            return int(digits) if digits else None

        date_rows = [r for r in tl_rows if (r.row_type or "date") == "date" and to_int(r.year) and to_int(r.month)]

        def date_key(r):
            return to_int(r.year) * 10000 + to_int(r.month) * 100 + (to_int(r.day) or 0)

        cycle = ""
        if date_rows:
            start = min(date_rows, key=date_key)
            out_rows = [r for r in date_rows if any(str(v or "").strip() for v in cell_map.get(r.id, []))]
            pool = out_rows or date_rows
            end = max(pool, key=date_key)
            cycle = f"{to_int(start.year)} 年 {to_int(start.month)} 月~{to_int(end.year)} 年 {to_int(end.month)} 月"

        file_rows = [r for r in date_rows if any("产品开发计划" in str(v or "") for v in cell_map.get(r.id, []))]
        file_date = ""
        if file_rows:
            fr = min(file_rows, key=date_key)
            file_date = f"{to_int(fr.year)}年{to_int(fr.month)}月{to_int(fr.day)}日"

        # 参与人员
        members = db.session.execute(select(ProjectMember).where(ProjectMember.prod_id == prod_id)).scalars().all()
        def find_member(pred):
            for m in members:
                if pred(str(m.role or "")):
                    return (m.name or "").strip()
            return ""
        pm = find_member(lambda r: "产品经理" in r)
        approver = find_member(lambda r: "负责人" in r and "产品" in r)

        def strip_num(title):
            return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()

        def fill(node):
            ref = node.get("ref_type")
            title = strip_num(node.get("title"))
            children = node.get("children") or []
            is_overview = ref == "prod_overview" or (title == "产品概况" and not children)
            is_cycle = ref == "prod_cycle" or (title == "产品开发周期" and not children)
            if (ref == "prod_name" or title == "产品简介") and prod_name:
                node["body"] = f"产品名称：{prod_name}"
            elif is_overview and overall_desc:
                node["body"] = overall_desc
            elif is_cycle and cycle:
                node["body"] = cycle
            if ref == "revision" or title == "文件修订记录":
                tables = node.get("tables") or []
                if tables and isinstance(tables[0], list):
                    t = tables[0]
                    cols = len(t[0]) if t and t[0] else 5
                    while len(t) < 6:
                        t.append([""] * cols)
                    row = t[1]
                    def set_if(i, val):
                        if val and not str(row[i] if i < len(row) else "").strip():
                            row[i] = val
                    set_if(0, file_date)
                    set_if(1, obj.version)
                    if not str(row[2] if len(row) > 2 else "").strip():
                        row[2] = "首次发布"
                    set_if(3, pm)
                    set_if(4, approver)
            for child in children:
                fill(child)

        for node in sections:
            fill(node)
        serv_review_util.ensure_review(
            content, "pdp", serv_review_util.review_date(prod_id, serv_review_util.REVIEW_DEFS["pdp"]["name_keywords"]), prod_id
        )
        serv_review_util.fill_cover_dates(content, serv_review_util.cover_date(prod_id, "pdp"))
        serv_review_util.fill_cover_signers(content, serv_review_util.cover_signers(prod_id, "pdp"))
        return content

    def __to_obj(self, row: PdpDoc, product: Product = None):
        obj = PdpDocObj(**row.dict())
        obj.content = self.__normalize_content(obj.content)
        serv_review_util.ensure_review(
            obj.content, "pdp",
            serv_review_util.review_date(row.product_id, serv_review_util.REVIEW_DEFS["pdp"]["name_keywords"]) if row.product_id else "",
            row.product_id,
        )
        serv_review_util.fill_cover_dates(
            obj.content, serv_review_util.cover_date(row.product_id, "pdp") if row.product_id else ""
        )
        serv_review_util.fill_cover_signers(
            obj.content, serv_review_util.cover_signers(row.product_id, "pdp") if row.product_id else {}
        )
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
        return obj

    async def add_pdp_doc(self, form: PdpDocForm):
        try:
            sql = select(func.count(PdpDoc.id)).where(
                PdpDoc.product_id == form.product_id,
                PdpDoc.version == form.version,
            )
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = PdpDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=PdpDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_pdp_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: PdpDoc = db.session.execute(select(PdpDoc).where(PdpDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(PdpDoc.version).where(PdpDoc.product_id == target_pid)).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = PdpDoc(
                product_id=target_pid,
                version=version,
                file_no=sync_file_no_version(fromdoc.file_no, version),
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            return Resp.resp_ok(data=PdpDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_pdp_doc(self, form: PdpDocForm):
        try:
            row: PdpDoc = db.session.execute(select(PdpDoc).where(PdpDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_pdp_doc(self, id: int):
        db.session.execute(delete(PdpDoc).where(PdpDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_pdp_doc(self, id: int):
        sql = select(PdpDoc, Product).join(Product, PdpDoc.product_id == Product.id).where(PdpDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_pdp_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(PdpDoc.product_id == product_id)
        if version:
            wheres.append(PdpDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(PdpDoc.id)).join(Product, PdpDoc.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (
            select(PdpDoc, Product)
            .join(Product, PdpDoc.product_id == Product.id)
            .where(*wheres)
            .order_by(PdpDoc.id.desc())
            .offset(page_index * page_size)
            .limit(page_size)
        )
        rows: List[PdpDocObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    # ---------------- 导出 Word ----------------
    async def export_pdp_doc(self, output, id: int):
        resp = await self.get_pdp_doc(id)
        obj: PdpDocObj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        c = self.__autofill_for_export(self.__normalize_content(obj.content), obj)
        sections = c.get("sections") or []
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        document.settings.element.append(update_fields)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")
        docx_util.add_page_number_footer(section, obj.file_no or "")

        def write_center_title(text, size=22.0, bold=False):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.right_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, text, font_size=size, bold=bold)

        def add_blank_lines(count):
            for _ in range(max(0, int(count or 0))):
                document.add_paragraph("")

        def add_text(text):
            docx_util.save_txt2docx(str(text or ""), document)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            cell.text = ""
            lines = str(text or "").split("\n")
            for i, line in enumerate(lines):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER if align == WD_ALIGN_PARAGRAPH.CENTER else WD_CELL_VERTICAL_ALIGNMENT.TOP

        def add_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for r_idx, row in enumerate(grid):
                cells = table.add_row().cells
                for c_idx in range(cols):
                    set_cell(cells[c_idx], row[c_idx] if c_idx < len(row) else "", bold=(r_idx == 0))
            document.add_paragraph()

        def set_cover_value_cell(cell, text, bold, align):
            """封面填写格：若为签名图(data URL) 则渲染图片，否则按文本渲染。"""
            s = str(text or "")
            if s.startswith("data:image"):
                try:
                    b64 = s.split(",", 1)[1] if "," in s else ""
                    data = base64.b64decode(b64)
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = align
                    para.add_run().add_picture(BytesIO(data), height=Pt(33))
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    return
                except Exception:
                    pass
            set_cell(cell, text, bold=bold, align=align)

        def add_cover_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for row in grid:
                cells = table.add_row().cells
                for c_idx in range(cols):
                    text = row[c_idx] if c_idx < len(row) else ""
                    # 偶数列为标签（编写部门/文件版本/日期等）加粗，奇数列为填写值
                    set_cover_value_cell(cells[c_idx], text, bold=(c_idx % 2 == 0), align=WD_ALIGN_PARAGRAPH.CENTER)
                # 「生效日期」行：合并后面的填写格，跨满整行
                if (str(row[0]).strip() if row else "") == "生效日期" and cols > 2:
                    merged = cells[1]
                    for c_idx in range(2, cols):
                        merged = merged.merge(cells[c_idx])
                    set_cell(merged, row[1] if len(row) > 1 else "", align=WD_ALIGN_PARAGRAPH.CENTER)
            document.add_paragraph()

        def strip_num(title):
            return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()

        # 标题：用 Heading 样式（便于 Word 目录域识别），各级均加粗
        def add_body_heading(title, level):
            size = {1: 16.0, 2: 14.0, 3: 12.0}.get(level, 11.0)
            p = document.add_heading("", level=max(1, min(level, 9)))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, title, font_size=size, bold=True)

        # 正文章节：按树自动编号 1 / 1.1 / 1.1.1
        def render_body_section(node, level, number=""):
            name = strip_num(node.get("title"))
            heading = f"{number} {name}".strip() if number else name
            add_body_heading(heading, level=max(1, min(level, 9)))
            if (node.get("body") or "").strip():
                add_text(node.get("body"))
            if node.get("ref_type") == "review":
                for t_idx, table in enumerate(node.get("tables") or []):
                    serv_review_util.render_review_grid(document, table, set_cell, merge_col0=(t_idx == 0), merge_full=True)
            else:
                for table in (node.get("tables") or []):
                    add_grid(table)
            idx = 0
            for child in (node.get("children") or []):
                idx += 1
                child_num = f"{number}.{idx}" if number else f"{idx}"
                render_body_section(child, level + 1, child_num)

        cover = next((s for s in sections if s.get("ref_type") == "cover"), None)
        revision = next((s for s in sections if s.get("ref_type") == "revision"), None)
        body = [s for s in sections if s.get("ref_type") not in ("cover", "revision")]

        # 封面：上方留白 + 居中大标题 + 留白 + 封面表，使标题落在页面居中位置
        add_blank_lines(6)
        write_center_title((strip_num(cover.get("title")) if cover else "") or "产品开发计划", size=22.0, bold=True)
        add_blank_lines(4)
        if cover:
            for table in (cover.get("tables") or []):
                add_cover_grid(table)

        # 文件修订记录：独立页、居中标题
        document.add_page_break()
        write_center_title("文件修订记录", size=14.0, bold=True)
        add_blank_lines(2)
        if revision:
            for table in (revision.get("tables") or []):
                add_grid(table)

        # 目录：独立页、居中标题 + 目录域
        document.add_page_break()
        write_center_title("目录", size=16.0, bold=True)
        docx_util.insert_toc_field(document)

        # 正文：另起一页（评审记录章节不编号）
        document.add_page_break()
        seq = 0
        for node in body:
            if node.get("ref_type") == "review":
                render_body_section(node, 1, "")
            else:
                seq += 1
                render_body_section(node, 1, str(seq))

        document.save(output)
        output.seek(0)
