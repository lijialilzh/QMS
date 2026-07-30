#!/usr/bin/env python
# encoding: utf-8

# 安装维护手册服务层（测试文件 VV-005，PDP 风格章节树 + MD5 附件/评审）。
# 默认内容取自 src-res/imm_default_content.json。

import base64
import copy
import json
import logging
import os
import re
from io import BytesIO
from typing import List

from sqlalchemy import delete, func, select
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.imm_doc import ImmDoc
from ..model.prod_dhf import ProdDhf
from ..model.prod_runtime_env import ProdRuntimeEnv
from ..model.project_member import ProjectMember
from ..model.doc_file import DocFile
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_imm_doc import ImmDocForm
from ..obj.vobj_imm_doc import ImmDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from . import serv_review_util
from .serv_prod_runtime_env import DEFAULT_RUNTIME_ENV
from .serv_utils import new_version, sync_file_no_version
from .serv_utils import docx_util

logger = logging.getLogger(__name__)

_DEFAULT_FILE = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "src-res", "imm_default_content.json"
)
try:
    with open(_DEFAULT_FILE, encoding="utf-8") as _f:
        DEFAULT_IMM_CONTENT = json.load(_f)
except Exception:
    DEFAULT_IMM_CONTENT = {"md5_value": "", "package_name": "", "sections": []}

# 模板基准产品名（用于全文替换为当前产品名）。
BASE_NAME = "肿瘤CT图像随访与评估软件"
DOC_NAME = "安装维护手册"
DOC_KEY = "imm"
DATE_KEYWORDS = ["安装维护手册", "安装维护", "MD5"]


class Server(object):

    def __normalize_node(self, node):
        if not isinstance(node, dict):
            return {"title": str(node or ""), "body": "", "tables": [], "images": [], "children": []}
        result = dict(node)
        result["title"] = str(result.get("title") or "")
        result["body"] = str(result.get("body") or "")
        if "body_tail" in result:
            result["body_tail"] = str(result.get("body_tail") or "")
        tables = result.get("tables")
        if not isinstance(tables, list):
            tables = []
        norm_tables = []
        for table in tables:
            if isinstance(table, list):
                norm_tables.append([[str(c) if c is not None else "" for c in (row or [])] for row in table if isinstance(row, list)])
        result["tables"] = norm_tables
        imgs = result.get("images")
        result["images"] = [str(x) for x in imgs if isinstance(x, str) and x] if isinstance(imgs, list) else []
        children = result.get("children")
        result["children"] = [self.__normalize_node(c) for c in children] if isinstance(children, list) else []
        return result

    def __normalize_content(self, content):
        if not isinstance(content, dict) or not isinstance(content.get("sections"), list):
            return copy.deepcopy(DEFAULT_IMM_CONTENT)
        default = DEFAULT_IMM_CONTENT if isinstance(DEFAULT_IMM_CONTENT, dict) else {}
        return {
            "md5_value": str(content.get("md5_value") if content.get("md5_value") is not None else default.get("md5_value") or ""),
            "package_name": str(content.get("package_name") if content.get("package_name") is not None else default.get("package_name") or ""),
            "sections": [self.__normalize_node(s) for s in content["sections"]],
        }

    def __replace_in_node(self, node, replacer, skip_titles=None):
        if skip_titles and str(node.get("title") or "").strip() in skip_titles:
            return
        if node.get("body"):
            node["body"] = replacer(node["body"])
        if node.get("body_tail"):
            node["body_tail"] = replacer(node.get("body_tail"))
        for tbl in (node.get("tables") or []):
            for row in tbl:
                for i in range(len(row)):
                    if isinstance(row[i], str):
                        row[i] = replacer(row[i])
        for c in (node.get("children") or []):
            self.__replace_in_node(c, replacer, skip_titles=skip_titles)

    def __replace_name(self, node, base, name):
        if not name or base == name:
            return
        self.__replace_in_node(node, lambda s: s.replace(base, name) if isinstance(s, str) else s)

    def __replace_type_code(self, node, type_code):
        if not type_code:
            return
        # 2.2 MD5校验安装包 / 2.4.1 检查安装包 正文中的安装包名称保留模板原样；
        # 4.确认安装包信息 表格中的安装包名称也保留模板原样；
        # 软件总体描述 / 适用范围 body 取自产品 overall_desc/scope 原值，不应被型号替换污染。
        skip_titles = {"2.2 MD5校验安装包", "2.4.1 检查安装包", "4.确认安装包信息", "1.4 软件总体描述", "1.4软件总体描述", "1.5 适用范围", "1.5适用范围"}
        space_pat = "InferCare RECIST"
        under_pat = "InferCare_RECIST"
        space_new = f"InferCare {type_code}"
        under_new = f"InferCare_{type_code}"

        def repl(s):
            if not isinstance(s, str):
                return s
            return s.replace(space_pat, space_new).replace(under_pat, under_new)

        self.__replace_in_node(node, repl, skip_titles=skip_titles)

    def __member_name(self, prod_id, keywords):
        if not prod_id:
            return ""
        for kw in keywords:
            row = db.session.execute(
                select(ProjectMember).where(ProjectMember.prod_id == prod_id, ProjectMember.role.like(f"%{kw}%"))
                .order_by(ProjectMember.sort_order.asc(), ProjectMember.id.asc())
            ).scalars().first()
            if row and (row.name or "").strip():
                return row.name.strip()
        return ""

    def __dhf_file_no(self, prod_id):
        if not prod_id:
            return ""
        row = db.session.execute(
            select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name == DOC_NAME).order_by(ProdDhf.id.asc())
        ).scalars().first()
        if not row:
            row = db.session.execute(
                select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name.like(f"%{DOC_NAME}%")).order_by(ProdDhf.id.asc())
            ).scalars().first()
        return (row.code or "").strip() if row and row.code else ""

    def __runtime_env(self, prod_id):
        """从产品运行环境配置读取（无记录时用运行环境模板默认值）。"""
        env = dict(DEFAULT_RUNTIME_ENV)
        if not prod_id:
            return env
        row = db.session.execute(
            select(ProdRuntimeEnv).where(ProdRuntimeEnv.prod_id == prod_id)
        ).scalars().first()
        if row:
            for key in DEFAULT_RUNTIME_ENV.keys():
                val = getattr(row, key, None)
                if val is not None and str(val).strip():
                    env[key] = val
        return env

    @staticmethod
    def __strip_section_title(title):
        return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()

    @staticmethod
    def __overwrite_col1(table, label_map):
        for row in table:
            if not isinstance(row, list) or len(row) < 2:
                continue
            key = str(row[0]).strip()
            if key in label_map and str(label_map[key] or "").strip():
                row[1] = label_map[key]

    def __fill_runtime_node(self, node, rt):
        if not isinstance(node, dict) or not rt:
            return
        title = str(node.get("title") or "")
        plain = self.__strip_section_title(title)
        if "表1" in title or plain.startswith("服务器硬件"):
            for tbl in (node.get("tables") or []):
                self.__overwrite_col1(tbl, {
                    "CPU": rt.get("srv_cpu"),
                    "内存": rt.get("srv_memory"),
                    "GPU": rt.get("srv_gpu"),
                    "硬盘": rt.get("srv_disk"),
                    "网卡": rt.get("srv_nic"),
                })
        elif "表2" in title or "服务器软件" in plain:
            for tbl in (node.get("tables") or []):
                if len(tbl) >= 2 and isinstance(tbl[1], list) and len(tbl[1]) >= 3:
                    if str(rt.get("srv_os") or "").strip():
                        tbl[1][1] = rt["srv_os"]
                    if str(rt.get("srv_cuda") or "").strip():
                        tbl[1][2] = rt["srv_cuda"]
        elif "表3" in title or plain.startswith("用户端"):
            for tbl in (node.get("tables") or []):
                self.__overwrite_col1(tbl, {
                    "CPU": rt.get("cli_cpu"),
                    "内存": rt.get("cli_memory"),
                    "显示器分辨率": rt.get("cli_resolution"),
                    "操作系统": rt.get("cli_os"),
                    "浏览器": rt.get("cli_browser"),
                })
        elif "表4" in title or "网络" in plain:
            for tbl in (node.get("tables") or []):
                for row in tbl:
                    if not isinstance(row, list) or str(row[0]).strip() != "带宽" or len(row) < 3:
                        continue
                    if str(rt.get("net_lan") or "").strip():
                        row[1] = rt["net_lan"]
                    if str(rt.get("net_wan") or "").strip():
                        row[2] = rt["net_wan"]
        for child in (node.get("children") or []):
            self.__fill_runtime_node(child, rt)

    def __fill_runtime_env(self, content, prod_id):
        if not prod_id or not isinstance(content, dict):
            return content
        rt = self.__runtime_env(prod_id)
        for section in (content.get("sections") or []):
            plain = self.__strip_section_title(section.get("title"))
            if plain == "概述" or "运行环境" in plain:
                self.__fill_runtime_node(section, rt)
            else:
                for child in (section.get("children") or []):
                    if "运行环境" in str(child.get("title") or ""):
                        self.__fill_runtime_node(child, rt)
        return content

    def __fill_revision(self, content, prod_id, version, force=False):
        """文件修订记录首行：修改日期(评审/封面日期)、版本号、首次发布、修订人(TPM)、批准人(研发负责人)。
        force=True（切换产品）时强制用新产品数据覆盖，无时间线/人员的字段直接置空。"""
        rev_date = serv_review_util.cover_date(prod_id, DOC_KEY) if prod_id else ""
        if not rev_date and prod_id:
            rev_date = serv_review_util.review_date(prod_id, DATE_KEYWORDS)
        reviser = approver = ""
        if prod_id:
            members = db.session.execute(select(ProjectMember).where(ProjectMember.prod_id == prod_id)).scalars().all()
            reviser = next((m.name for m in members if "TPM" in str(m.role or "")), "")
            approver = next((m.name for m in members if "研发负责人" in str(m.role or "")), "")
        for s in (content.get("sections") or []):
            if s.get("ref_type") != "revision":
                continue
            tables = s.get("tables") or []
            if tables and isinstance(tables[0], list) and tables[0]:
                t = tables[0]
                cols = len(t[0]) if isinstance(t[0], list) and t[0] else 5
                while len(t) < 6:
                    t.append([""] * cols)
                if len(t) >= 2 and isinstance(t[1], list) and len(t[1]) >= 3:
                    row = t[1]
                    if force:
                        row[0] = rev_date
                        if len(row) >= 2:
                            row[1] = str(version or "")
                        if len(row) >= 3:
                            row[2] = "首次发布"
                        if len(row) >= 4:
                            row[3] = reviser
                        if len(row) >= 5:
                            row[4] = approver
                    else:
                        if not str(row[0] or "").strip():
                            row[0] = rev_date
                        if version and len(row) >= 2 and not str(row[1] or "").strip():
                            row[1] = str(version)
                        if len(row) >= 3 and not str(row[2] or "").strip():
                            row[2] = "首次发布"
                        if len(row) >= 4 and reviser and not str(row[3] or "").strip():
                            row[3] = reviser
                        if len(row) >= 5 and approver and not str(row[4] or "").strip():
                            row[4] = approver
            break
        return content

    def __fill_md5_attachment(self, content, package_name, md5_value):
        for s in (content.get("sections") or []):
            if s.get("ref_type") != "md5_attachment":
                continue
            for tbl in (s.get("tables") or []):
                if not tbl or not isinstance(tbl[0], list):
                    continue
                header = [str(c).strip() for c in tbl[0][:2]]
                if header != ["安装包名称", "MD5值"] or len(tbl) < 2:
                    continue
                row = tbl[1]
                if not isinstance(row, list):
                    continue
                if package_name and not str(row[0] or "").strip():
                    row[0] = package_name
                if md5_value and not str(row[1] or "").strip():
                    row[1] = md5_value
            break
        return content

    def __fill_md5_review(self, content, prod_id, product, package_name, md5_value, rev_date):
        prod_name = (getattr(product, "name", "") or "").strip() if product else ""
        # 评审结论里的安装包名称保留模板原样（从 4.确认安装包信息 表格读取，不自动获取/不型号替换）
        pkg = self.__read_attachment_pkg(content)
        md5 = (md5_value or "").strip()
        test_lead = self.__member_name(prod_id, ("测试负责人", "主测试工程师")) if prod_id else ""
        pm = self.__member_name(prod_id, ("产品经理",)) if prod_id else ""
        for s in (content.get("sections") or []):
            if s.get("ref_type") != "md5_review":
                continue
            # 标题显示"附件一 评审结论"，body 清空（正文不显示）
            s["title"] = "附件一 评审结论"
            s["body"] = ""
            tables = s.get("tables") or []
            if tables and isinstance(tables[0], list):
                # 评审内容表：第三列"评审结论"下的"通过"前加复选框勾选（☑通过 □存在问题）
                check_mark = serv_review_util.CHECK
                for row in tables[0]:
                    if not isinstance(row, list) or not row:
                        continue
                    if str(row[0] or "").startswith("评审结论"):
                        row[0] = (
                            f"评审结论：\n通过，{prod_name}安装包（{pkg}）的MD5值为{md5}，"
                            f"安装包信息及MD5值信息确认准确。"
                        )
                        for i in range(1, len(row)):
                            row[i] = ""
                    elif len(row) >= 3 and str(row[2] or "").strip() == "通过  □存在问题":
                        row[2] = check_mark
            if len(tables) > 1 and isinstance(tables[1], list):
                ptbl = tables[1]
                for row in ptbl:
                    if not isinstance(row, list) or len(row) < 2:
                        continue
                    role = str(row[0] or "").strip()
                    if role == "测试负责人" and test_lead and not str(row[1] or "").strip():
                        row[1] = test_lead
                    if len(row) >= 5 and str(row[3] or "").strip() == "产品经理" and pm and not str(row[4] or "").strip():
                        row[4] = pm
                serv_review_util.autofill_review_person_table(ptbl, DOC_KEY, rev_date, prod_id)
            break
        return content

    def __read_md5_attachment(self, content):
        """从 md5_attachment 章节的表格读取 (package_name, md5_value)。
        表头为 ["安装包名称", "MD5值"]，取第二行。"""
        for s in (content.get("sections") or []):
            if s.get("ref_type") != "md5_attachment":
                continue
            for tbl in (s.get("tables") or []):
                if not tbl or not isinstance(tbl[0], list):
                    continue
                header = [str(c).strip() for c in tbl[0][:2]]
                if header != ["安装包名称", "MD5值"] or len(tbl) < 2:
                    continue
                row = tbl[1]
                if not isinstance(row, list):
                    continue
                pkg = str(row[0] or "").strip() if len(row) > 0 else ""
                md5 = str(row[1] or "").strip() if len(row) > 1 else ""
                return pkg, md5
            break
        return "", ""

    def __read_attachment_pkg(self, content):
        """从 md5_attachment 下「4.确认安装包信息」子节点表格读取安装包名称（保留模板原值）。"""
        for s in (content.get("sections") or []):
            if s.get("ref_type") != "md5_attachment":
                continue
            for c in (s.get("children") or []):
                if self.__strip_section_title(str(c.get("title") or "")) != "确认安装包信息":
                    continue
                for tbl in (c.get("tables") or []):
                    if not tbl or not isinstance(tbl[0], list):
                        continue
                    header = [str(x).strip() for x in tbl[0][:2]]
                    if header != ["安装包名称", "MD5值"] or len(tbl) < 2:
                        continue
                    row = tbl[1]
                    if isinstance(row, list) and len(row) > 0:
                        return str(row[0] or "").strip()
                break
            break
        return ""

    def __autofill(self, content, prod_id, product=None, version="", force=False):
        if not isinstance(content, dict):
            return content
        type_code = (getattr(product, "type_code", "") or "").strip() if product else ""
        full_version = (getattr(product, "full_version", "") or "").strip() if product else ""
        # 安装包名称 / MD5 值以 md5_attachment 附件章节的表格为准（用户在附件里维护）。
        # 附件表格为空时，兜底按产品型号+版本自动生成安装包名称。
        package_name, md5_value = self.__read_md5_attachment(content)
        if not package_name and type_code and full_version:
            package_name = f"InferCare_{type_code}-{full_version}.zip"
        rev_date = ""
        if prod_id:
            rev_date = serv_review_util.cover_date(prod_id, DOC_KEY) or serv_review_util.review_date(prod_id, DATE_KEYWORDS)
        if product and product.name:
            for s in (content.get("sections") or []):
                self.__replace_name(s, BASE_NAME, product.name)
        if type_code:
            for s in (content.get("sections") or []):
                self.__replace_type_code(s, type_code)
        self.__fill_runtime_env(content, prod_id)
        self.__fill_revision(content, prod_id, version, force=force)
        self.__fill_md5_attachment(content, package_name, md5_value)
        self.__fill_md5_review(content, prod_id, product, package_name, md5_value, rev_date)
        serv_review_util.fill_cover_dates(content, rev_date, force=force)
        serv_review_util.fill_cover_signers(content, serv_review_util.cover_signers(prod_id, DOC_KEY) if prod_id else {}, force=force)
        return content

    def __to_obj(self, row: ImmDoc, product: Product = None):
        obj = ImmDocObj(**row.dict())
        obj.content = self.__autofill(self.__normalize_content(obj.content), row.product_id, product, row.version)
        if not (obj.file_no or "").strip():
            obj.file_no = self.__dhf_file_no(row.product_id)
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
        return obj

    async def add_imm_doc(self, form: ImmDocForm):
        try:
            sql = select(func.count(ImmDoc.id)).where(ImmDoc.product_id == form.product_id, ImmDoc.version == form.version)
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = ImmDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=ImmDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_imm_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: ImmDoc = db.session.execute(select(ImmDoc).where(ImmDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(ImmDoc.version).where(ImmDoc.product_id == target_pid)).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = ImmDoc(
                product_id=target_pid, version=version,
                file_no=sync_file_no_version(fromdoc.file_no, version),
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            # 跨产品复制：用新产品信息强制填充
            if target_pid != fromdoc.product_id:
                await self.rebind_product(newdoc.id, target_pid)
            return Resp.resp_ok(data=ImmDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_imm_doc(self, form: ImmDocForm):
        try:
            row: ImmDoc = db.session.execute(select(ImmDoc).where(ImmDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    def __fill_product_info(self, content, product):
        """切换产品时按新产品信息重建自动获取的章节正文（1.1软件说明/1.4总体描述/1.5适用范围/2.确认产品信息/概述）。
        参考 rmp 模式：按标题定位，用新产品字段强制覆盖，不依赖旧值匹配。"""
        if not product or not isinstance(content, dict):
            return content
        prod_name = (product.name or "").strip()
        type_code = (product.type_code or "").strip()
        full_version = (product.full_version or "").strip()
        release_version = (product.release_version or "").strip()
        # 制造商取产品 registrant（注册人），为空则显示空，不兜底固定公司名
        registrant = (product.registrant or "").strip()
        overall_desc = (product.overall_desc or "").strip()
        scope = (product.scope or "").strip()
        # 型号直接用 type_code 原值（字段本身可能已含/不含 InferCare 前缀，不再额外拼接）
        type_display = type_code
        # 物理拓扑图：从该产品的 doc_file（img_topo 类别）获取，URL 前缀 /data.trace 供前端静态访问
        topo_url = ""
        topo_row = db.session.execute(
            select(DocFile).where(DocFile.product_id == product.id, DocFile.category == "img_topo").order_by(DocFile.id)
        ).scalars().first()
        if topo_row and topo_row.file_url:
            # file_url 形如 data.trace/img_topo/51.png，转为 /data.trace/img_topo/51.png
            url = str(topo_row.file_url).strip()
            if url.startswith("data.trace/"):
                url = "/" + url
            topo_url = url

        def rebuild(node):
            title = str(node.get("title") or "").strip()
            plain = self.__strip_section_title(title)
            if plain == "软件说明":
                node["body"] = (
                    "软件标识：\n"
                    f"产品名称：{prod_name}\n"
                    f"产品型号：{type_display}\n"
                    f"发布版本：{release_version}\n"
                    f"完整版本：{full_version}\n"
                    f"制造商：{registrant}"
                )
            elif plain == "软件总体描述":
                # 自动获取：为空则显示空
                node["body"] = overall_desc
            elif plain == "适用范围":
                # 自动获取：为空则显示空
                node["body"] = scope
            elif plain == "确认产品信息":
                node["body"] = (
                    f"产品名称：{prod_name}\n"
                    f"产品型号：{type_display}\n"
                    f"发布版本：{release_version}\n"
                    f"完整版本：{full_version}"
                )
            elif plain == "概述":
                node["body"] = (
                    f"本文档适用于{registrant}正式发布的产品。本软件需要由本公司授权的安装人员进行安装。"
                    "软件的日常维护也需要由本公司的授权人员进行。\n"
                    f"本文档第二章介绍了正确安装“{prod_name}”产品的所有步骤，第三章介绍了软件安装完毕后的一些必要的软件配置项，"
                    f"第四章软件运行前的基本测试流程，第五章介绍了如何检查软件的日志以及软件所在服务器所需的日常维护。"
                )
            elif plain == "硬件拓扑":
                # 物理拓扑图：从 doc_file 获取，无则置空
                node["body"] = "物理拓扑图"
                node["images"] = [topo_url] if topo_url else []
            for c in (node.get("children") or []):
                rebuild(c)

        for s in (content.get("sections") or []):
            rebuild(s)
        return content

    async def rebind_product(self, id: int, product_id: int):
        """切换产品：重置为默认模板 + 用新产品信息重建自动获取字段 + 强制填充封面/修订/运行环境后保存，返回新 obj。
        参考 rmp 模式：重置模板避免旧值污染，按新产品字段强制覆盖。"""
        try:
            row: ImmDoc = db.session.execute(select(ImmDoc).where(ImmDoc.id == id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            product: Product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
            if not product:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            # 目标产品已有同版本记录时删除旧记录，避免唯一约束冲突
            db.session.execute(delete(ImmDoc).where(ImmDoc.product_id == product_id, ImmDoc.version == row.version, ImmDoc.id != id))
            # 重置为默认模板（丢弃旧产品污染的内容）
            content = self.__normalize_content(None)
            row.product_id = product_id
            # 先做型号替换/封面/修订/运行环境等自动填充
            content = self.__autofill(content, product_id, product, row.version, force=True)
            # 再用新产品字段重建自动获取章节（总体描述/适用范围等用产品原值整体覆盖，避免被型号替换污染）
            content = self.__fill_product_info(content, product)
            row.content = content
            db.session.commit()
            return Resp.resp_ok(data=self.__to_obj(row, product))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_imm_doc(self, id: int):
        db.session.execute(delete(ImmDoc).where(ImmDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_imm_doc(self, id: int):
        sql = select(ImmDoc, Product).join(Product, ImmDoc.product_id == Product.id).where(ImmDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_imm_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(ImmDoc.product_id == product_id)
        if version:
            wheres.append(ImmDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(ImmDoc.id)).join(Product, ImmDoc.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (select(ImmDoc, Product).join(Product, ImmDoc.product_id == Product.id).where(*wheres)
               .order_by(ImmDoc.id.desc()).offset(page_index * page_size).limit(page_size))
        rows: List[ImmDocObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    # ---------------- 导出 Word（PDP 风格章节树） ----------------
    def __export_docx(self, output, obj: ImmDocObj, mode: str = "main"):
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        c = obj.content if isinstance(obj.content, dict) else self.__normalize_content(obj.content)
        sections = c.get("sections") or []

        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

        if mode in ("main", "all"):
            update_fields = OxmlElement("w:updateFields")
            update_fields.set(qn("w:val"), "true")
            document.settings.element.append(update_fields)
            header_para = section.header.add_paragraph()
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            docx_util.fonted_txt(header_para, obj.file_no or "")
            docx_util.add_page_number_footer(section, obj.file_no or "")

        def add_blank_lines(n):
            for _ in range(max(0, int(n or 0))):
                document.add_paragraph("")

        def write_center_title(text, size=22.0, bold=False):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            docx_util.fonted_txt(p, str(text or ""), font_size=size, bold=bold)

        def add_text(text):
            docx_util.save_txt2docx(str(text or ""), document)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            s = str(text or "")
            if s.startswith("data:image"):
                try:
                    b64 = s.split(",", 1)[1] if "," in s else ""
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = align
                    para.add_run().add_picture(BytesIO(base64.b64decode(b64)), height=Pt(33))
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    return
                except Exception:
                    pass
            cell.text = ""
            for i, line in enumerate(s.split("\n")):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER if align == WD_ALIGN_PARAGRAPH.CENTER else WD_CELL_VERTICAL_ALIGNMENT.TOP

        def _set_fixed_widths(table, widths_dxa):
            table.autofit = False
            table.allow_autofit = False
            tbl_pr = table._tbl.tblPr
            layout = tbl_pr.find(qn("w:tblLayout"))
            if layout is None:
                layout = OxmlElement("w:tblLayout")
                tbl_pr.append(layout)
            layout.set(qn("w:type"), "fixed")
            grid_el = table._tbl.find(qn("w:tblGrid"))
            if grid_el is not None:
                for gc in list(grid_el):
                    grid_el.remove(gc)
                for w in widths_dxa:
                    gc = OxmlElement("w:gridCol")
                    gc.set(qn("w:w"), str(w))
                    grid_el.append(gc)
            for row in table.rows:
                for i, w in enumerate(widths_dxa):
                    if i < len(row.cells):
                        tcpr = row.cells[i]._tc.get_or_add_tcPr()
                        tcw = tcpr.find(qn("w:tcW"))
                        if tcw is None:
                            tcw = OxmlElement("w:tcW")
                            tcpr.append(tcw)
                        tcw.set(qn("w:w"), str(w))
                        tcw.set(qn("w:type"), "dxa")

        def add_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for r_idx, row in enumerate(grid):
                cells = table.add_row().cells
                for c_idx in range(cols):
                    set_cell(cells[c_idx], row[c_idx] if c_idx < len(row) else "", bold=(r_idx == 0))
            first = str(grid[0][0]).strip() if grid and grid[0] else ""
            if first == "SCI名字":
                if cols == 5:
                    _set_fixed_widths(table, [1700, 1100, 2000, 1500, 3000])
                elif cols == 4:
                    _set_fixed_widths(table, [1700, 2200, 1500, 3600])
            elif first == "类别" and cols >= 3:
                if cols == 3:
                    _set_fixed_widths(table, [1300, 1300, 5000])
                rows = table.rows
                n = len(rows)
                r = 1
                while r < n:
                    val = str((grid[r][0] if r < len(grid) and grid[r] else "") or "").strip()
                    if not val:
                        r += 1
                        continue
                    r2 = r
                    while r2 + 1 < n and str((grid[r2 + 1][0] if grid[r2 + 1] else "") or "").strip() == val:
                        r2 += 1
                    if r2 > r:
                        merged = rows[r].cells[0].merge(rows[r2].cells[0])
                        set_cell(merged, val, align=WD_ALIGN_PARAGRAPH.LEFT)
                        merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    r = r2 + 1
            elif first in ("测试目标", "测试方法和技术") and cols == 2:
                _set_fixed_widths(table, [2200, 7100])
            elif first == "配置":
                if cols == 2:
                    _set_fixed_widths(table, [1900, 7400])
                elif cols == 3:
                    _set_fixed_widths(table, [1900, 3700, 3700])
            document.add_paragraph()

        def add_cover_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for row in grid:
                cells = table.add_row().cells
                for c_idx in range(cols):
                    text = row[c_idx] if c_idx < len(row) else ""
                    set_cell(cells[c_idx], text, bold=(c_idx % 2 == 0), align=WD_ALIGN_PARAGRAPH.CENTER)
                if (str(row[0]).strip() if row else "") == "生效日期" and cols > 2:
                    merged = cells[1]
                    for c_idx in range(2, cols):
                        merged = merged.merge(cells[c_idx])
                    set_cell(merged, row[1] if len(row) > 1 else "", align=WD_ALIGN_PARAGRAPH.CENTER)
            document.add_paragraph()

        def add_image(image_url):
            raw_url = str(image_url or "").strip()
            if not raw_url or not raw_url.startswith("data:image/"):
                return
            try:
                image_data = raw_url.split(",", 1)[1]
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run().add_picture(BytesIO(base64.b64decode(image_data)), width=Inches(6.2))
                document.add_paragraph()
            except Exception:
                logger.exception("导出安装维护手册图片失败")

        def strip_num(title):
            return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()

        def add_body_heading(title, level):
            size = {1: 16.0, 2: 14.0, 3: 12.0}.get(level, 11.0)
            p = document.add_heading("", level=max(1, min(level, 9)))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, title, font_size=size, bold=True)

        def render_body_section(node, level, number="", numbered=True):
            name = strip_num(node.get("title"))
            heading = f"{number} {name}".strip() if number else name
            if heading:
                add_body_heading(heading, level=max(1, min(level, 9)))
            if (node.get("body") or "").strip():
                add_text(node.get("body"))
            caps = node.get("table_captions") or []
            for t_idx, table in enumerate(node.get("tables") or []):
                cap = caps[t_idx] if t_idx < len(caps) else ""
                if str(cap or "").strip():
                    add_text(cap)
                add_grid(table)
            for image_url in (node.get("images") or []):
                add_image(image_url)
            if str(node.get("body_tail") or "").strip():
                add_text(node.get("body_tail"))
            idx = 0
            for child in (node.get("children") or []):
                # "表N xxx" 是表格标题块（被上一章节引用），不作为独立章节编号，
                # 避免占用章节序号导致后续真实章节号错位。
                if re.match(r"^\s*表\d+([、.\s　]|$)", str(child.get("title") or "")):
                    render_body_section(child, level + 1, "", numbered=numbered)
                    continue
                idx += 1
                if numbered and number:
                    child_num = f"{number}.{idx}"
                elif numbered:
                    child_num = str(idx)
                else:
                    child_num = ""
                render_body_section(child, level + 1, child_num, numbered=numbered)

        if mode in ("main", "all"):
            cover = next((s for s in sections if s.get("ref_type") == "cover"), None)
            revision = next((s for s in sections if s.get("ref_type") == "revision"), None)
            body = [s for s in sections if s.get("ref_type") not in ("cover", "revision", "md5_attachment", "md5_review")]

            add_blank_lines(6)
            write_center_title((strip_num(cover.get("title")) if cover else "") or DOC_NAME, size=22.0, bold=True)
            add_blank_lines(4)
            if cover:
                for table in (cover.get("tables") or []):
                    add_cover_grid(table)

            document.add_page_break()
            write_center_title("文件修订记录", size=14.0, bold=True)
            add_blank_lines(2)
            if revision:
                for table in (revision.get("tables") or []):
                    add_grid(table)

            document.add_page_break()
            write_center_title("目录", size=16.0, bold=True)
            docx_util.insert_toc_field(document)

            document.add_page_break()
            seq = 0
            for node in body:
                seq += 1
                render_body_section(node, 1, str(seq), numbered=True)

            if mode == "all":
                document.add_page_break()

        if mode in ("md5_attachment", "all"):
            if mode == "md5_attachment":
                # 单独导出 MD5 附件时设页眉；all 模式下页眉已在 main 段设置
                header_para = section.header.add_paragraph()
                header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                docx_util.fonted_txt(header_para, obj.file_no or "")
                docx_util.add_page_number_footer(section, obj.file_no or "")
            node = next((s for s in sections if s.get("ref_type") == "md5_attachment"), None)
            if node:
                title = strip_num(node.get("title")) or "安装维护手册附件：MD5值"
                write_center_title(title, size=16.0, bold=True)
                add_blank_lines(2)
                body = str(node.get("body") or "").strip()
                if body and body != title:
                    add_text(body)
                caps = node.get("table_captions") or []
                for t_idx, table in enumerate(node.get("tables") or []):
                    cap = caps[t_idx] if t_idx < len(caps) else ""
                    if str(cap or "").strip():
                        add_text(cap)
                    add_grid(table)
                for image_url in (node.get("images") or []):
                    add_image(image_url)
                if str(node.get("body_tail") or "").strip():
                    add_text(node.get("body_tail"))
                for child in (node.get("children") or []):
                    # MD5 附件子节点标题自带编号（1./2./3./4.），保留原标题不 strip_num
                    child_title = str(child.get("title") or "").strip()
                    if child_title:
                        add_body_heading(child_title, level=1)
                    if str(child.get("body") or "").strip():
                        add_text(child.get("body"))
                    caps2 = child.get("table_captions") or []
                    for t_idx2, table2 in enumerate(child.get("tables") or []):
                        cap2 = caps2[t_idx2] if t_idx2 < len(caps2) else ""
                        if str(cap2 or "").strip():
                            add_text(cap2)
                        add_grid(table2)
                    for image_url2 in (child.get("images") or []):
                        add_image(image_url2)
                    if str(child.get("body_tail") or "").strip():
                        add_text(child.get("body_tail"))
                if mode == "all":
                    document.add_page_break()

        if mode in ("md5_review", "all"):
            if mode == "md5_review":
                # 单独导出评审记录时设页眉；all 模式下页眉已在 main 段设置
                header_para = section.header.add_paragraph()
                header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                docx_util.fonted_txt(header_para, obj.file_no or "")
                docx_util.add_page_number_footer(section, obj.file_no or "")
            node = next((s for s in sections if s.get("ref_type") == "md5_review"), None)
            if node:
                # 标题"附件一 评审结论"作为章节标题居中显示，不输出正文
                title_txt = self.__strip_section_title(node.get("title")) or str(node.get("body") or "").strip() or "附件一 评审结论"
                write_center_title(title_txt, size=16.0, bold=True)
                add_blank_lines(2)
                for table in (node.get("tables") or []):
                    serv_review_util.render_review_grid(document, table, set_cell)

        docx_util.fill_toc_cache(document)
        document.save(output)
        output.seek(0)

    async def export_imm_doc(self, output, id: int):
        resp = await self.get_imm_doc(id)
        # 列表导出完整版：主文档 + MD5值附件 + 评审记录（分页分隔）
        self.__export_docx(output, resp.data, mode="all")

    async def export_imm_md5_attachment(self, output, id: int):
        resp = await self.get_imm_doc(id)
        self.__export_docx(output, resp.data, mode="md5_attachment")

    async def export_imm_md5_review(self, output, id: int):
        resp = await self.get_imm_doc(id)
        self.__export_docx(output, resp.data, mode="md5_review")
