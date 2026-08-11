from enum import Enum
import logging
import json
import re
import io
import base64
import os
import builtins
import functools
import shutil
import subprocess
import tempfile
import time
from datetime import datetime
from types import SimpleNamespace
from typing import Any, Dict, List, Optional, Tuple, Union
from sqlalchemy import select, delete, func
from sqlalchemy.dialects.postgresql import insert as pg_insert
from sqlalchemy.sql import desc
try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph
    from docx.shared import Pt
    from docx import enum as dox_enum
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
except Exception:
    Document = None
    OxmlElement = None
    DocxTable = None
    Paragraph = None
    Pt = None
    dox_enum = None
    qn = None
    RGBColor = None
from ..obj.vobj_user import UserObj
from ..obj.vobj_sds_trace import SdsTraceObj
from ..model.srs_type import SrsType
from ..model.srs_reqd import SrsReqd
from ..obj.vobj_sds_reqd import SdsReqdObj
from ..model.srs_req import SrsReq
from ..model.sds_reqd import Logic, SdsReqd
from ..model.doc_file import DocFile
from ..model.sds_trace import SdsTrace
from ..model.srs_doc import SrsDoc
from ..obj.tobj_srs_doc import Table, TabHeader
from ..model.product import Product, UserProd
from ..obj.vobj_sds_doc import CompareObj, SdsDocObj
from ..model.sds_doc import SdsDoc, SdsNode
from ..obj.tobj_sds_doc import SdsDocForm, SdsNodeForm, SdsTable, SdsExtraTable
from ..utils.sql_ctx import db
from ..utils.i18n import ts
from ..utils import get_uuid
from .serv_utils.tree_util import find_parent
from .serv_utils import new_version
from .serv_sds_trace import Server as ServSdsTrace, NAME_DICT, fixed_rcn300_sds_codes
from .serv_sds_reqd import Server as ServSdsReqd
from .serv_srs_doc import Server as ServSrsDoc
from .serv_sds_srs_trace_sync import SdsSrsTraceSyncMixin

from ..obj import Page, Resp
from . import msg_err_db, save_file, serv_review_util

logger = logging.getLogger(__name__)


def _normalize_sds_img_url(url: Optional[str]) -> Optional[str]:
    """前端会为图片追加 ?_v= 缓存参数，多次保存会叠加超长；内部静态路径去掉 query 即可。"""
    if not url:
        return url
    s = str(url).strip()
    if "?" not in s:
        return s
    path, _ = s.split("?", 1)
    if path.startswith("/data.trace/"):
        return path
    return s


srsdoc_serv = ServSrsDoc()
sdstrace_serv = ServSdsTrace()
sdstreqd_serv = ServSdsReqd()
DELETED_SRS_VERSION_PREFIX = "__deleted_srs__"


class RefTypes(Enum):
    img_struct = "img_struct"
    img_flow = "img_flow"
    img_topo = "img_topo"
    sds_traces = "sds_traces"
    sds_reqds = "sds_reqds"

class Server(SdsSrsTraceSyncMixin, object):

    @staticmethod
    def _serialize_sds_table(table):
        if not table:
            return None
        if isinstance(table, SdsTable):
            return table.dict()
        if isinstance(table, dict):
            return table
        if isinstance(table, str):
            try:
                return json.loads(table)
            except Exception:
                return None
        if hasattr(table, "dict"):
            return table.dict()
        return table

    @staticmethod
    def _parse_sds_table(table):
        if not table:
            return None
        if isinstance(table, SdsTable):
            return table
        if isinstance(table, str):
            return SdsTable.parse_raw(table)
        if isinstance(table, dict):
            return SdsTable.parse_obj(table)
        return None

    @staticmethod
    def _strip_heading_number(title: str) -> str:
        return SdsSrsTraceSyncMixin._strip_sds_heading_text(title)

    def _normalize_word_imported_chapter_numbers(self, roots: List[SdsNodeForm]) -> List[SdsNodeForm]:
        """封面/修订记录不参与章节编号，正文从 1 章开始。"""
        if not roots:
            return roots

        def normalize_title(title: str) -> str:
            return re.sub(r"\s+", "", self._strip_heading_number(title or ""))

        def is_front_matter(node: SdsNodeForm) -> bool:
            title = normalize_title(getattr(node, "title", "") or "")
            return title in {"目录", "需求规格说明", "软件详细设计", "软件详细设计说明书", "文件修订记录"}

        def parse_heading(title: str) -> Optional[str]:
            matched = re.match(
                r"^\s*(\d+(?:\.\d+)*)(?:[\s、.．]+|(?=[\u4e00-\u9fffA-Za-z]))",
                str(title or "").strip(),
            )
            return matched.group(1) if matched else None

        first_body_major = 0
        for node in roots or []:
            if is_front_matter(node):
                continue
            heading = parse_heading(getattr(node, "title", "") or "")
            if heading:
                first_body_major = int(heading.split(".")[0])
                break
        offset = first_body_major - 1 if first_body_major > 1 else 0

        def shift_title(title: str, front: bool) -> str:
            raw = str(title or "")
            heading = parse_heading(raw)
            if front:
                return self._strip_heading_number(raw) if heading else raw
            if not heading or offset <= 0:
                return raw
            parts = [int(part) for part in heading.split(".")]
            if not parts or parts[0] <= offset:
                return raw
            parts[0] -= offset
            return raw.replace(heading, ".".join(str(part) for part in parts), 1)

        def walk(nodes: List[SdsNodeForm], front_parent: bool = False):
            for node in nodes or []:
                front = front_parent or is_front_matter(node)
                node.title = shift_title(getattr(node, "title", "") or "", front)
                walk(getattr(node, "children", None) or [], front)

        walk(roots)
        return roots

    @staticmethod
    def _is_word_imported_doc(roots: List[SdsNodeForm]) -> bool:
        def walk(nodes: List[SdsNodeForm]):
            for node in nodes or []:
                raw_title = (getattr(node, "title", "") or "").strip()
                title = re.sub(r"\s+", "", raw_title)
                if title == "目录" or title.startswith("目录"):
                    return True
                if Server.__is_imported_table_title(raw_title) or re.match(r"^图\s*\d+", raw_title):
                    return True
                if getattr(node, "table", None) is not None:
                    return True
                if walk(getattr(node, "children", None) or []):
                    return True
            return False
        return walk(roots or [])

    @staticmethod
    def __is_imported_table_title(value: str):
        return re.match(r"^导入表格\d*$", (value or "").strip()) is not None

    @staticmethod
    def __is_table_caption_line(line: str):
        txt = (line or "").strip()
        if not txt:
            return False
        # JSON 键值行不是表题（如 "code":0, / "filename":"a.zip"）
        if re.match(r'^[\'"]\s*[^\'"]+\s*[\'"]\s*:\s*.+$', txt):
            return False
        if re.match(r"^(表|table)\s*\d+", txt, re.I):
            return True
        if re.match(r"^图\s*\d+", txt):
            return False
        # 兼容“alembic_version 数据库迁移表：”这类末尾冒号标题
        if "表" in txt and re.match(r"^.+表\s*[:：]?$", txt):
            return True
        # 仅将“字段名: 值”这类无空格英文标识识别为表题，避免误判整句正文
        if re.match(r"^[A-Za-z][A-Za-z0-9_]{1,64}[:：]\s*.+$", txt):
            return True
        if re.search(r"[:：]", txt) and len(txt) <= 80 and re.search(r"[。！？]$", txt) is None:
            parts = [seg.strip() for seg in re.split(r"[:：]", txt)]
            left = parts[0] if parts else ""
            right = "".join(parts[1:]).strip() if len(parts) > 1 else ""
            left_is_identifier = re.match(r"^[A-Za-z][A-Za-z0-9_]{1,64}$", left or "") is not None
            if left and right and (left_is_identifier or "表" in left):
                return True
            # 冒号后为空时，仅“含表”才视作表名，避免“库2数据库：”误命中
            if left and not right and "表" in left:
                return True
        return False

    def __bind_imported_table_titles(self, nodes: List[SdsNodeForm]):
        def walk(node_list: List[SdsNodeForm]):
            for node in node_list or []:
                children = list(getattr(node, "children", None) or [])
                table_children = [
                    child for child in children
                    if getattr(child, "table", None) and getattr(getattr(child, "table", None), "headers", None)
                ]
                if table_children:
                    lines = str(getattr(node, "text", "") or "").replace("\r", "").split("\n")
                    caption_entries = [
                        (idx, (line or "").strip())
                        for idx, line in enumerate(lines)
                        if self.__is_table_caption_line(line)
                    ]
                    if caption_entries:
                        used_line_idx = set()
                        for idx, child in enumerate(table_children):
                            if idx >= len(caption_entries):
                                break
                            line_idx, caption = caption_entries[idx]
                            if not caption:
                                continue
                            child_title = str(getattr(child, "title", "") or "").strip()
                            child_label = str(getattr(child, "label", "") or "").strip()
                            if not child_title or self.__is_imported_table_title(child_title):
                                child.title = caption
                            elif not child_label:
                                child.label = caption
                            used_line_idx.add(line_idx)
                        if used_line_idx:
                            remained = [
                                (line or "").strip()
                                for idx, line in enumerate(lines)
                                if idx not in used_line_idx and str(line or "").strip()
                            ]
                            node.text = "\n".join(remained)
                if children:
                    walk(children)
        walk(nodes or [])

    def __bind_imported_image_titles(self, nodes: List[SdsNodeForm]):
        def is_imported_image_title(value: str):
            return re.match(r"^导入图片\d*$", (value or "").strip()) is not None

        def is_image_caption_line(value: str):
            return re.match(r"^\s*图\s*\d+\s*", (value or "").strip()) is not None

        def walk(node_list: List[SdsNodeForm]):
            for node in node_list or []:
                children = list(getattr(node, "children", None) or [])
                image_children = [
                    child for child in children
                    if str(getattr(child, "img_url", "") or "").strip()
                ]
                if image_children:
                    lines = str(getattr(node, "text", "") or "").replace("\r", "").split("\n")
                    caption_entries = [
                        (idx, (line or "").strip())
                        for idx, line in enumerate(lines)
                        if is_image_caption_line(line)
                    ]
                    if caption_entries:
                        used_line_idx = set()
                        for idx, child in enumerate(image_children):
                            if idx >= len(caption_entries):
                                break
                            line_idx, caption = caption_entries[idx]
                            if not caption:
                                continue
                            child_title = str(getattr(child, "title", "") or "").strip()
                            child_label = str(getattr(child, "label", "") or "").strip()
                            if not child_title or is_imported_image_title(child_title):
                                child.title = caption
                            elif not child_label:
                                child.label = caption
                            used_line_idx.add(line_idx)
                        if used_line_idx:
                            remained = [
                                (line or "").strip()
                                for idx, line in enumerate(lines)
                                if idx not in used_line_idx and str(line or "").strip()
                            ]
                            node.text = "\n".join(remained)
                if children:
                    walk(children)
        walk(nodes or [])

    @staticmethod
    def __has_table_payload(node: SdsNodeForm):
        table = getattr(node, "table", None)
        headers = getattr(table, "headers", None) if table else None
        return bool(headers)

    def __extract_data_structure_db_table_plan(self, docx: Document) -> List[Dict[str, Any]]:
        plans: List[Dict[str, Any]] = []
        if docx is None or Paragraph is None or DocxTable is None:
            return plans

        def strip_heading_no(value: str):
            return re.sub(r"^(\d+(?:\.\d+)*)(?:[\s、.．]+|(?=[\u4e00-\u9fffA-Za-z]))", "", value or "").strip()

        def is_db_heading(value: str):
            txt = strip_heading_no((value or "").strip())
            if not txt:
                return False
            return re.search(r"数据库\s*[:：]?$", txt) is not None and not self.__is_table_caption_line(txt)

        active_plan: Union[Dict[str, Any], None] = None
        current_db: Union[Dict[str, Any], None] = None
        pending_caption = ""

        for child in docx.element.body.iterchildren():
            tag = str(child.tag).lower()
            if tag.endswith("}p"):
                para = Paragraph(child, docx._body)
                txt = (para.text or "").replace("\xa0", " ").strip()
                if not txt:
                    continue
                heading_no = self.__extract_heading_no(txt)
                heading_level = (heading_no.count(".") + 1) if heading_no else None
                if active_plan and heading_level is not None and heading_level <= int(active_plan.get("level") or 1):
                    if heading_no != active_plan.get("chapter"):
                        active_plan = None
                        current_db = None
                        pending_caption = ""
                if "数据结构" in txt and heading_no:
                    active_plan = {
                        "chapter": heading_no,
                        "title": txt,
                        "level": heading_level or 1,
                        "dbs": [],
                    }
                    plans.append(active_plan)
                    current_db = None
                    pending_caption = ""
                    continue
                if not active_plan:
                    continue
                if is_db_heading(txt):
                    current_db = {
                        "title": txt,
                        "captions": [],
                    }
                    active_plan["dbs"].append(current_db)
                    pending_caption = ""
                    continue
                if current_db is not None and self.__is_table_caption_line(txt):
                    pending_caption = txt
                continue

            if tag.endswith("}tbl") and active_plan and current_db is not None:
                # Word 原始顺序里遇到一张表，就归到当前库标题下；表题用最近一行表名。
                current_db["captions"].append(pending_caption)
                pending_caption = ""

        for plan in plans:
            logger.info("[DB_PLAN] 数据结构=%r dbs=%s",
                plan.get("title"),
                [
                    {
                        "title": db.get("title"),
                        "count": len(db.get("captions") or []),
                        "captions": db.get("captions") or [],
                    }
                    for db in (plan.get("dbs") or [])
                ],
            )
        return plans

    @staticmethod
    def __extract_heading_no(title: str):
        matched = re.match(r"^(\d+(?:\.\d+)*)(?:[\s、.．]+|(?=[\u4e00-\u9fffA-Za-z]))", (title or "").strip())
        return (matched.group(1) if matched else "")

    def __split_data_structure_db_tables(self, nodes: List[SdsNodeForm], db_table_plans: List[Dict[str, Any]] = None):
        db_heading_re = re.compile(r"((?:[A-Za-z]+\s*)?库\s*[0-9一二三四五六七八九十]+\s*数据库\s*[:：])", re.I)
        caption_re = re.compile(r"((?:表\s*\d+(?:[.\-_]\d+)*|[A-Za-z][A-Za-z0-9_]{1,64})\s*[:：]\s*[^\n]{0,80})", re.I)
        plan_used_indexes = set()

        def is_db_heading_title(value: str):
            txt = (value or "").strip()
            if not txt:
                return False
            txt = re.sub(r"^(\d+(?:\.\d+)*)(?:[\s、.．]+|(?=[\u4e00-\u9fffA-Za-z]))", "", txt).strip()
            return re.search(r"数据库\s*[:：]?$", txt) is not None

        def is_real_table_title(value: str):
            txt = (value or "").strip()
            if not txt:
                return False
            if is_db_heading_title(txt):
                return False
            return self.__is_table_caption_line(txt)

        def is_placeholder_title(value: str):
            txt = (value or "").strip()
            if not txt:
                return True
            return self.__is_imported_table_title(txt)

        def extract_caption_matches(raw_text: str):
            txt = (raw_text or "")
            matches = []
            for m in caption_re.finditer(txt):
                cap = (m.group(1) or "").strip()
                if not cap:
                    continue
                if is_db_heading_title(cap):
                    continue
                if not is_real_table_title(cap):
                    continue
                matches.append({"text": cap, "pos": m.start()})
            return matches

        def extract_db_index(value: str):
            txt = (value or "").strip()
            hit = re.search(r"库\s*([0-9一二三四五六七八九十]+)\s*数据库", txt, re.I)
            if not hit:
                return 0
            raw = (hit.group(1) or "").strip()
            if raw.isdigit():
                return int(raw)
            zh_map = {
                "一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
                "六": 6, "七": 7, "八": 8, "九": 9, "十": 10
            }
            return zh_map.get(raw, 0)

        def normalize_compare(value: str):
            return re.sub(r"[\s\u3000:：]+", "", re.sub(r"^(\d+(?:\.\d+)*)(?:[\s、.．]+|(?=[\u4e00-\u9fffA-Za-z]))", "", value or "")).lower()

        def pick_docx_plan_for_node(node_title: str):
            plans = db_table_plans or []
            if not plans:
                return None
            node_heading = self.__extract_heading_no(node_title)
            node_plain = normalize_compare(node_title)
            for idx, plan in enumerate(plans):
                if idx in plan_used_indexes:
                    continue
                if node_heading and node_heading == plan.get("chapter"):
                    plan_used_indexes.add(idx)
                    return plan
            for idx, plan in enumerate(plans):
                if idx in plan_used_indexes:
                    continue
                if node_plain and node_plain == normalize_compare(str(plan.get("title") or "")):
                    plan_used_indexes.add(idx)
                    return plan
            available = [idx for idx in range(len(plans)) if idx not in plan_used_indexes]
            if len(available) == 1:
                idx = available[0]
                plan_used_indexes.add(idx)
                return plans[idx]
            return None

        def collect_table_nodes(node_list: List[SdsNodeForm]):
            result: List[SdsNodeForm] = []
            def _walk(items: List[SdsNodeForm]):
                for item in items or []:
                    if self.__has_table_payload(item):
                        result.append(item)
                        continue
                    _walk(list(getattr(item, "children", None) or []))
            _walk(node_list or [])
            return result

        def apply_docx_plan(node: SdsNodeForm, plan: Dict[str, Any]):
            db_defs = [db for db in (plan.get("dbs") or []) if db.get("title")]
            if len(db_defs) < 2:
                return False
            all_tables = collect_table_nodes(list(getattr(node, "children", None) or []))
            if not all_tables:
                return False
            base_heading = self.__extract_heading_no(str(getattr(node, "title", "") or ""))
            existing_anchors = [
                child for child in list(getattr(node, "children", None) or [])
                if is_db_heading_title(str(getattr(child, "title", "") or ""))
            ]
            rebuilt_children: List[SdsNodeForm] = []
            cursor = 0
            for idx, db_def in enumerate(db_defs):
                raw_title = str(db_def.get("title") or "").strip()
                db_title = raw_title
                if base_heading and not self.__extract_heading_no(db_title):
                    db_title = f"{base_heading}.{idx + 1} {db_title}".strip()
                captions = [str(c or "").strip() for c in (db_def.get("captions") or [])]
                take = len(captions)
                if idx == len(db_defs) - 1:
                    take = max(take, len(all_tables) - cursor)
                take = max(0, min(take, len(all_tables) - cursor))
                assigned = all_tables[cursor: cursor + take]
                cursor += take
                for tab_idx, tab in enumerate(assigned):
                    if tab_idx < len(captions) and captions[tab_idx]:
                        tab.title = captions[tab_idx]
                db_idx = extract_db_index(db_title)
                anchor = next(
                    (
                        item for item in existing_anchors
                        if db_idx > 0 and extract_db_index(str(getattr(item, "title", "") or "")) == db_idx
                    ),
                    None,
                ) or (existing_anchors[idx] if idx < len(existing_anchors) else None)
                if anchor is None:
                    anchor = SdsNodeForm(title=db_title, children=[])
                anchor.title = db_title
                anchor.children = assigned
                rebuilt_children.append(anchor)
            if cursor < len(all_tables) and rebuilt_children:
                rebuilt_children[-1].children = (rebuilt_children[-1].children or []) + all_tables[cursor:]
            node.children = rebuilt_children
            logger.info("[DB_SPLIT_APPLY] title=%r -> %s",
                getattr(node, "title", ""),
                [
                    {
                        "db": getattr(child, "title", ""),
                        "count": len(getattr(child, "children", None) or []),
                        "tables": [getattr(t, "title", "") for t in (getattr(child, "children", None) or [])],
                    }
                    for child in rebuilt_children
                ],
            )
            return True

        def fix_postgresql_prefix_split(node: SdsNodeForm):
            db_nodes = [
                child for child in list(getattr(node, "children", None) or [])
                if is_db_heading_title(str(getattr(child, "title", "") or ""))
            ]
            if len(db_nodes) != 2:
                return False
            first_db, second_db = db_nodes[0], db_nodes[1]
            first_tables = [c for c in list(getattr(first_db, "children", None) or []) if self.__has_table_payload(c)]
            second_tables = [c for c in list(getattr(second_db, "children", None) or []) if self.__has_table_payload(c)]
            first_title = str(getattr(first_db, "title", "") or "")
            if "Postgresql" not in first_title and "postgresql" not in first_title.lower():
                return False
            if len(first_tables) != 1 or len(second_tables) < 2:
                return False
            split_idx = -1
            for idx, tab in enumerate(second_tables):
                tab_title = str(getattr(tab, "title", "") or "")
                if "weekly_statistic" in tab_title or "周度综合统计表" in tab_title:
                    split_idx = idx
            if split_idx < 0:
                return False
            moved_to_first = second_tables[:split_idx + 1]
            remained_second = second_tables[split_idx + 1:]
            if not moved_to_first or not remained_second:
                return False
            first_non_tables = [c for c in list(getattr(first_db, "children", None) or []) if not self.__has_table_payload(c)]
            second_non_tables = [c for c in list(getattr(second_db, "children", None) or []) if not self.__has_table_payload(c)]
            first_db.children = [*first_non_tables, *first_tables, *moved_to_first]
            second_db.children = [*second_non_tables, *remained_second]
            logger.info("[DB_SPLIT_FIX] title=%r 库1=%d 库2=%d 库1表=%s 库2表=%s",
                getattr(node, "title", ""),
                len([c for c in (first_db.children or []) if self.__has_table_payload(c)]),
                len([c for c in (second_db.children or []) if self.__has_table_payload(c)]),
                [str(getattr(c, "title", "") or "") for c in (first_db.children or []) if self.__has_table_payload(c)],
                [str(getattr(c, "title", "") or "") for c in (second_db.children or []) if self.__has_table_payload(c)],
            )
            return True

        def walk(node_list: List[SdsNodeForm]):
            for node in node_list or []:
                title = str(getattr(node, "title", "") or "").strip()
                text = str(getattr(node, "text", "") or "")
                children = list(getattr(node, "children", None) or [])

                merged_hint = f"{title} {text}"
                is_data_structure_node = "数据结构" in merged_hint
                if is_data_structure_node and children:
                    docx_plan = pick_docx_plan_for_node(title)
                    if docx_plan and apply_docx_plan(node, docx_plan):
                        if getattr(node, "children", None):
                            walk(node.children or [])
                        continue
                    # 优先使用“已解析出的库标题节点”作为锚点重挂后续表节点：
                    # 结构应为 5.6 -> 5.6.1库1 / 5.6.2库2 -> 各自表节点（与编辑页层级一致）。
                    db_anchor_indexes = [
                        idx for idx, child in enumerate(children)
                        if is_db_heading_title(str(getattr(child, "title", "") or ""))
                    ]
                    if db_anchor_indexes:
                        rebuilt_children: List[SdsNodeForm] = []
                        current_anchor: Union[SdsNodeForm, None] = None
                        for child in children:
                            child_title = str(getattr(child, "title", "") or "").strip()
                            if is_db_heading_title(child_title):
                                current_anchor = child
                                current_anchor.children = list(getattr(current_anchor, "children", None) or [])
                                rebuilt_children.append(current_anchor)
                                continue
                            if current_anchor is not None and self.__has_table_payload(child):
                                current_anchor.children.append(child)
                                continue
                            rebuilt_children.append(child)
                        node.children = rebuilt_children
                        children = rebuilt_children

                    # 若已存在“库1/库2”锚点，按正文中的库标题区间重新分配各库下表数量（允许某库为0张）
                    # 解决“库2起始表被挂到库1”的问题。
                    db_anchor_nodes = [
                        child for child in children
                        if is_db_heading_title(str(getattr(child, "title", "") or ""))
                    ]
                    if len(db_anchor_nodes) >= 2:
                        heading_matches = [
                            {"text": (m.group(1) or "").strip(), "pos": m.start()}
                            for m in db_heading_re.finditer(text or "")
                        ]
                        heading_matches = [item for item in heading_matches if item["text"]]
                        caption_matches = extract_caption_matches(text or "")
                        logger.info("[DB_SPLIT_DETAIL] node=%r anchor节点=%s | text中找到库标题=%s | text中找到表标题(%d)=%s",
                            title,
                            [str(getattr(a, "title", "") or "") for a in db_anchor_nodes],
                            [h["text"] for h in heading_matches],
                            len(caption_matches),
                            [c["text"] for c in caption_matches],
                        )
                        if len(heading_matches) >= 2 and len(caption_matches) > 0:
                            # 与下方逻辑保持一致：显式库序号（库1/库2）必须保留
                            filtered_headings = [heading_matches[0]]
                            for item in heading_matches[1:]:
                                prev = filtered_headings[-1]
                                prev_idx = extract_db_index(str(prev.get("text") or ""))
                                curr_idx = extract_db_index(str(item.get("text") or ""))
                                if prev_idx > 0 and curr_idx == prev_idx + 1:
                                    filtered_headings.append(item)
                                    continue
                                between_caps = [
                                    c for c in caption_matches
                                    if int(prev["pos"]) < int(c["pos"]) < int(item["pos"])
                                ]
                                if len(between_caps) == 0:
                                    continue
                                filtered_headings.append(item)
                            heading_matches = filtered_headings

                            use_count = min(len(db_anchor_nodes), len(heading_matches))
                            if use_count >= 2:
                                flat_tables: List[SdsNodeForm] = []
                                anchor_non_tables: Dict[int, List[SdsNodeForm]] = {}
                                for idx, db_node in enumerate(db_anchor_nodes):
                                    original_children = list(getattr(db_node, "children", None) or [])
                                    non_tables = [c for c in original_children if not self.__has_table_payload(c)]
                                    tables = [c for c in original_children if self.__has_table_payload(c)]
                                    anchor_non_tables[idx] = non_tables
                                    flat_tables.extend(tables)

                                if flat_tables:
                                    counts: List[int] = []
                                    for i in range(use_count - 1):
                                        start = int(heading_matches[i]["pos"])
                                        end = int(heading_matches[i + 1]["pos"])
                                        cnt = len([c for c in caption_matches if start < int(c["pos"]) < end])
                                        counts.append(max(0, cnt))
                                    logger.info("[DB_SPLIT_DETAIL] flat_tables=%d counts(前n-1)=%s heading_matches_used=%s",
                                        len(flat_tables), counts, [h["text"] for h in heading_matches[:use_count]])
                                    used = sum(counts)
                                    counts.append(max(0, len(flat_tables) - used))

                                    cursor = 0
                                    for i in range(use_count):
                                        db_node = db_anchor_nodes[i]
                                        take = counts[i] if i < len(counts) else 0
                                        if i == use_count - 1:
                                            take = len(flat_tables) - cursor
                                        take = max(0, min(take, len(flat_tables) - cursor))
                                        assigned = flat_tables[cursor: cursor + take]
                                        cursor += take
                                        db_node.children = [*(anchor_non_tables.get(i, []) or []), *assigned]

                                        # 表名按该库区间的标题顺序回填
                                        h_start = int(heading_matches[i]["pos"])
                                        h_end = int(heading_matches[i + 1]["pos"]) if i + 1 < use_count else 10**9
                                        db_caps = [str(c["text"]) for c in caption_matches if h_start < int(c["pos"]) < h_end]
                                        cap_cursor = 0
                                        for tab in assigned:
                                            old_title = str(getattr(tab, "title", "") or "").strip()
                                            if is_placeholder_title(old_title) and cap_cursor < len(db_caps):
                                                tab.title = db_caps[cap_cursor]
                                                cap_cursor += 1
                        fix_postgresql_prefix_split(node)

                    table_children = [child for child in children if self.__has_table_payload(child)]
                    plain_children = [child for child in children if not self.__has_table_payload(child)]
                    already_grouped = any(
                        is_db_heading_title(str(getattr(child, "title", "") or ""))
                        and any(self.__has_table_payload(gc) for gc in (getattr(child, "children", None) or []))
                        for child in children
                    )
                    if table_children and not already_grouped:
                        heading_matches = [
                            {"text": (m.group(1) or "").strip(), "pos": m.start()}
                            for m in db_heading_re.finditer(text or "")
                        ]
                        heading_matches = [item for item in heading_matches if item["text"]]
                        caption_matches = extract_caption_matches(text or "")
                        if len(heading_matches) > 1:
                            # 清理“伪库标题”：若两个库标题之间没有任何表标题，后一个通常是正文误命中
                            filtered_headings = [heading_matches[0]]
                            for item in heading_matches[1:]:
                                prev = filtered_headings[-1]
                                prev_idx = extract_db_index(str(prev.get("text") or ""))
                                curr_idx = extract_db_index(str(item.get("text") or ""))
                                # 显式“库1/库2/库3 ...”标题必须保留，不能被误过滤
                                if prev_idx > 0 and curr_idx == prev_idx + 1:
                                    filtered_headings.append(item)
                                    continue
                                between_caps = [
                                    c for c in caption_matches
                                    if int(prev["pos"]) < int(c["pos"]) < int(item["pos"])
                                ]
                                if len(between_caps) == 0:
                                    continue
                                filtered_headings.append(item)
                            heading_matches = filtered_headings

                        if heading_matches:
                            use_count = min(len(heading_matches), len(table_children))
                            if use_count > 0:
                                guessed_counts: List[int] = []
                                if len(caption_matches) > 0 and use_count > 1:
                                    for i in range(use_count - 1):
                                        start = int(heading_matches[i]["pos"])
                                        end = int(heading_matches[i + 1]["pos"])
                                        guessed_counts.append(
                                            len([c for c in caption_matches if start < int(c["pos"]) < end])
                                        )

                                assign_counts: List[int] = []
                                remain_tables = len(table_children)
                                for i in range(use_count):
                                    if i == use_count - 1:
                                        cnt = remain_tables
                                    else:
                                        guessed = guessed_counts[i] if i < len(guessed_counts) else 0
                                        min_for_rest = use_count - i - 1
                                        cnt = max(1, guessed)
                                        cnt = min(cnt, max(1, remain_tables - min_for_rest))
                                    assign_counts.append(cnt)
                                    remain_tables -= cnt

                                base_heading = self.__extract_heading_no(title)
                                grouped_children: List[SdsNodeForm] = []
                                table_cursor = 0
                                for i in range(use_count):
                                    cnt = assign_counts[i]
                                    db_title_raw = str(heading_matches[i]["text"] or "").strip()
                                    db_title = db_title_raw
                                    if base_heading and not self.__extract_heading_no(db_title_raw):
                                        db_title = f"{base_heading}.{i + 1} {db_title_raw}".strip()
                                    db_tables = table_children[table_cursor: table_cursor + cnt]
                                    # 为每个库下的表回填标题，避免前端展示时出现“只有表结构没表名”
                                    h_start = int(heading_matches[i]["pos"])
                                    h_end = int(heading_matches[i + 1]["pos"]) if i + 1 < use_count else 10**9
                                    db_caps = [str(c["text"]) for c in caption_matches if h_start < int(c["pos"]) < h_end]
                                    cap_cursor = 0
                                    for tab in db_tables:
                                        old_title = str(getattr(tab, "title", "") or "").strip()
                                        if is_placeholder_title(old_title) and cap_cursor < len(db_caps):
                                            tab.title = db_caps[cap_cursor]
                                            cap_cursor += 1
                                    table_cursor += cnt
                                    grouped_children.append(SdsNodeForm(title=db_title, children=db_tables))
                                if table_cursor < len(table_children) and grouped_children:
                                    grouped_children[-1].children = (grouped_children[-1].children or []) + table_children[table_cursor:]

                                node.children = [*plain_children, *grouped_children]

                    # 最后兜底执行一次，确保“库1只有首表、库2吞了库1后续表”的情况在入库前被修正。
                    fix_postgresql_prefix_split(node)

                if getattr(node, "children", None):
                    walk(node.children or [])

        walk(nodes or [])

        def cleanup_duplicate_db_heading_text(node_list: List[SdsNodeForm]):
            # 库标题（如「Postgresql库1数据库：」）已被提为子节点(5.6.1/5.6.2)后，
            # 从数据结构父节点正文里删除与子节点标题重复的同名行，避免正文与子节点重复展示。
            for node in node_list or []:
                title = str(getattr(node, "title", "") or "").strip()
                text = str(getattr(node, "text", "") or "")
                children = list(getattr(node, "children", None) or [])
                if text and "数据结构" in f"{title} {text}":
                    db_titles = {
                        normalize_compare(str(getattr(c, "title", "") or ""))
                        for c in children
                        if is_db_heading_title(str(getattr(c, "title", "") or ""))
                    }
                    db_titles = {t for t in db_titles if t}
                    if db_titles:
                        kept_lines = []
                        for line in text.replace("\r", "").split("\n"):
                            norm = normalize_compare(line)
                            if norm and norm in db_titles:
                                continue
                            kept_lines.append(line)
                        node.text = "\n".join(kept_lines).strip()
                cleanup_duplicate_db_heading_text(children)

        cleanup_duplicate_db_heading_text(nodes or [])

    def __persist_data_url_images(self, nodes: List[SdsNodeForm]):
        ext_map = {
            "image/png": "png",
            "image/jpeg": "jpg",
            "image/jpg": "jpg",
            "image/gif": "gif",
            "image/bmp": "bmp",
            "image/webp": "webp",
        }

        def walk(node_list: List[SdsNodeForm]):
            for node in node_list or []:
                img_url = (getattr(node, "img_url", None) or "").strip()
                if img_url.startswith("data:"):
                    matched = re.match(r"^data:([^;]+);base64,(.+)$", img_url, re.S)
                    if matched:
                        mime = (matched.group(1) or "").lower()
                        b64 = matched.group(2) or ""
                        ext = ext_map.get(mime, "png")
                        try:
                            bys = base64.b64decode(b64)
                            path = os.path.join("data.trace", "sds_node_img", "import_sds", f"{get_uuid()}.{ext}")
                            os.makedirs(os.path.dirname(path), exist_ok=True)
                            with open(path, "wb") as fs:
                                fs.write(bys)
                            node.img_url = path
                        except Exception:
                            node.img_url = None
                    else:
                        node.img_url = None
                if getattr(node, "children", None):
                    walk(node.children or [])

        walk(nodes or [])

    @staticmethod
    def _normalize_code(code: str):
        txt = (code or "").strip().upper()
        txt = re.sub(r"\s+", "", txt)
        txt = re.sub(r"[，。；;、,.]+$", "", txt)
        return txt

    @staticmethod
    def __to_srs_code(code: str):
        txt = Server._normalize_code(code)
        if txt.startswith("SDS-"):
            return "SRS-" + txt[4:]
        return txt

    @staticmethod
    def __normalize_section_name(value: str):
        txt = (value or "").strip()
        txt = re.sub(r"^[（(]?[一二三四五六七八九十0-9]+[)）.\s、]*", "", txt)
        txt = re.sub(r"[\s:：\-_，。；;、]+", "", txt)
        return txt

    def __detect_sds_reqd_field(self, node: SdsNodeForm):
        merged = self.__normalize_section_name(f"{getattr(node, 'label', '')}{getattr(node, 'title', '')}")
        if not merged:
            return None
        if any(k in merged for k in ["总体描述", "需求概述", "概述"]):
            return "overview"
        if "程序逻辑" in merged or "逻辑" in merged:
            return "logic_txt"
        if "输入项" in merged or merged == "输入":
            return "intput"
        if "输出项" in merged or merged == "输出":
            return "output"
        if "接口" in merged:
            return "interface"
        # “功能”放在逻辑之后，避免“子功能”误判
        if "功能" in merged:
            return "func_detail"
        return None

    def __extract_sds_reqd_fields_from_text(self, text: str):
        content = (text or "").strip()
        if not content:
            return {}
        header_re = re.compile(
            r"^\s*[（(]?\s*(?:\d+|[一二三四五六七八九十]+)\s*[）)]?\s*(总体描述|需求概述|概述|功能|程序逻辑|输入项|输入|输出项|输出|接口)\s*$"
        )
        result: Dict[str, str] = {}
        current_field = ""
        bucket: List[str] = []

        def flush():
            nonlocal bucket, current_field
            if not current_field:
                bucket = []
                return
            txt = "\n".join([line for line in bucket if line.strip()]).strip()
            if txt:
                old = result.get(current_field, "")
                if not old or len(txt) > len(old):
                    result[current_field] = txt
            bucket = []

        for raw in content.splitlines():
            line = (raw or "").strip()
            if not line:
                continue
            matched = header_re.match(line)
            if matched:
                flush()
                sec = self.__normalize_section_name(matched.group(1))
                if any(k in sec for k in ["总体描述", "需求概述", "概述"]):
                    current_field = "overview"
                elif "程序逻辑" in sec or "逻辑" in sec:
                    current_field = "logic_txt"
                elif "输入项" in sec or sec == "输入":
                    current_field = "intput"
                elif "输出项" in sec or sec == "输出":
                    current_field = "output"
                elif "接口" in sec:
                    current_field = "interface"
                elif "功能" in sec:
                    current_field = "func_detail"
                else:
                    current_field = ""
                continue
            if current_field:
                bucket.append(line)
        flush()
        return result

    def __extract_sds_reqd_payload(self, nodes: List[SdsNodeForm]):
        payload: Dict[str, Dict[str, str]] = {}

        def save_value(code: str, field: str, text: str):
            if not code or not field or not text:
                return
            data = payload.setdefault(code, {})
            old = data.get(field, "")
            # 保留信息量更大的文本，避免被短标题覆盖
            if not old or len(text) > len(old):
                data[field] = text

        def walk(node_list: List[SdsNodeForm], current_code: str = ""):
            for node in node_list or []:
                node_code = self._normalize_code(getattr(node, "sds_code", "") or "")
                active_code = node_code or current_code
                field = self.__detect_sds_reqd_field(node)
                text = (getattr(node, "text", "") or "").strip()
                if active_code and field and text:
                    save_value(active_code, field, text)
                if active_code and text:
                    for f_key, f_val in self.__extract_sds_reqd_fields_from_text(text).items():
                        save_value(active_code, f_key, f_val)
                if getattr(node, "children", None):
                    walk(node.children or [], active_code)

        walk(nodes or [])
        return payload

    def __sync_imported_sds_reqd_fields(self, sds_doc_id: int, srs_doc_id: int, nodes: List[SdsNodeForm]):
        reqd_payload = self.__extract_sds_reqd_payload(nodes)
        if not reqd_payload:
            return
        srs_codes = [self.__to_srs_code(code) for code in reqd_payload.keys() if code]
        srs_codes = [code for code in srs_codes if code]
        if not srs_codes:
            return

        req_rows = db.session.execute(
            select(SrsReq).where(SrsReq.doc_id == srs_doc_id, SrsReq.code.in_(srs_codes))
        ).scalars().all()
        if not req_rows:
            return

        req_id_map = {row.id: row.code for row in req_rows}
        sds_reqd_rows = db.session.execute(
            select(SdsReqd).where(SdsReqd.doc_id == sds_doc_id, SdsReqd.req_id.in_(list(req_id_map.keys())))
        ).scalars().all()
        if not sds_reqd_rows:
            return

        for row in sds_reqd_rows:
            srs_code = req_id_map.get(row.req_id, "")
            sds_code = self._normalize_code(srs_code.replace("SRS-", "SDS-")) if srs_code else ""
            values = reqd_payload.get(sds_code) or reqd_payload.get(self._normalize_code(srs_code))
            if not values:
                continue
            for field in ["overview", "func_detail", "logic_txt", "intput", "output", "interface"]:
                val = (values.get(field) or "").strip()
                if val:
                    setattr(row, field, val)
        db.session.commit()

    async def import_sds_doc_word(self, product_id: int, srsdoc_id: int, version: str, change_log: str, file):
        if Document is None or DocxTable is None or Paragraph is None:
            return Resp.resp_err(msg="当前环境缺少 python-docx 依赖，暂不可用 Word 导入。")
        try:
            srs_row = None
            import_version = (version or "").strip()
            if srsdoc_id:
                srs_row = db.session.execute(
                    select(SrsDoc).where(
                        SrsDoc.id == srsdoc_id,
                        SrsDoc.product_id == product_id,
                        ~SrsDoc.version.like(f"{DELETED_SRS_VERSION_PREFIX}%"),
                    )
                ).scalars().first()
                if not srs_row:
                    return Resp.resp_err(msg="导入失败：未找到匹配的需求文档版本，请重新选择。")
            # 优先绑定“同版本”需求文档，避免导入 A0 时误关联到最新 A1
            if not srs_row and import_version:
                srs_row = db.session.execute(
                    select(SrsDoc)
                    .where(
                        SrsDoc.product_id == product_id,
                        SrsDoc.version == import_version,
                        ~SrsDoc.version.like(f"{DELETED_SRS_VERSION_PREFIX}%"),
                    )
                    .order_by(desc(SrsDoc.create_time), desc(SrsDoc.id))
                ).scalars().first()
            if not srs_row:
                srs_row = db.session.execute(
                    select(SrsDoc)
                    .where(SrsDoc.product_id == product_id, ~SrsDoc.version.like(f"{DELETED_SRS_VERSION_PREFIX}%"))
                    .order_by(desc(SrsDoc.create_time), desc(SrsDoc.id))
                ).scalars().first()
            if not srs_row:
                return Resp.resp_err(msg="导入失败：当前产品下未找到需求规格说明，请先导入需求规格说明。")

            bys = await file.read()
            docx = Document(io.BytesIO(bys))
            db_table_plans = self.__extract_data_structure_db_table_plan(docx)
            content, _ = srsdoc_serv._Server__parse_docx_content(docx)  # 复用 SRS 导入解析
            file_name = file.filename or ""
            _, file_no = srsdoc_serv._Server__extract_file_info(file_name)

            def extract_imported_catalog_text(source_docx: Document):
                lines = []
                in_catalog = False
                for para in source_docx.paragraphs:
                    txt = re.sub(r"\s+", " ", (para.text or "").strip())
                    if not txt:
                        continue
                    if re.sub(r"\s+", "", txt) == "目录":
                        in_catalog = True
                        continue
                    if not in_catalog:
                        continue
                    if re.match(r"^1(?:[\s、.．]+|(?=[\u4e00-\u9fffA-Za-z]))", txt):
                        break
                    if re.match(r"^\d+(?:\.\d+)*\s+\S.*(?:[.·…]{3,}|\s+)\d+\s*$", txt):
                        lines.append(txt)
                return "\n".join(lines).strip()

            def to_sds_node(node):
                data = {}
                for key in ["title", "label", "img_url", "text", "ref_type", "table", "sds_code"]:
                    val = getattr(node, key, None)
                    if val is not None:
                        data[key] = val
                if not data.get("sds_code"):
                    srs_code = getattr(node, "srs_code", None)
                    if srs_code:
                        data["sds_code"] = srs_code.replace("SRS-", "SDS-")
                data["children"] = [to_sds_node(child) for child in (getattr(node, "children", None) or [])]
                return SdsNodeForm(**data)

            sds_content = [to_sds_node(node) for node in (content or [])]
            imported_catalog_text = extract_imported_catalog_text(docx)
            if imported_catalog_text and not any(re.sub(r"\s+", "", str(getattr(node, "title", "") or "")) == "目录" for node in sds_content):
                sds_content.insert(0, SdsNodeForm(title="目录", text=imported_catalog_text, children=[]))
            # 导入入库前，把“图X 标题”绑定到对应图片节点标题，避免编辑页只看到“导入图片X”
            self.__bind_imported_image_titles(sds_content)
            # 导入入库前，先把“正文里的表名”绑定到对应表节点，避免后续查看/编辑再做文本猜测
            self.__bind_imported_table_titles(sds_content)
            # 入库前固定“数据结构 -> 库 -> 表”层级，后续展示直接读取树关系，不再依赖前端二次猜测
            self.__split_data_structure_db_tables(sds_content, db_table_plans)
            self._normalize_word_imported_chapter_numbers(sds_content)
            # === 导入调试日志：数据结构分组结果 ===
            def _log_db_split(nodes, path="root"):
                for n in (nodes or []):
                    t = str(getattr(n, "title", "") or "").strip()
                    children = list(getattr(n, "children", None) or [])
                    if "数据结构" in t or re.search(r"数据库\s*[:：]?$", re.sub(r"^\d[\d.]*\s*", "", t)):
                        child_titles = [str(getattr(c, "title", "") or "").strip() for c in children]
                        logger.info("[DB_SPLIT] path=%s title=%r children(%d)=%s", path, t, len(child_titles), child_titles)
                    for c in children:
                        _log_db_split([c], f"{path}/{t}")
            _log_db_split(sds_content)
            # === end ===
            self.__persist_data_url_images(sds_content)
            form = SdsDocForm(
                srsdoc_id=srs_row.id,
                product_id=product_id,
                version=version,
                file_no=file_no or None,
                change_log=change_log,
                content=sds_content,
            )
            resp = await self.add_sds_doc(form, preserve_word_structure=True)
            if resp.code == 200 and resp.data and resp.data.id:
                self.__sync_imported_sds_reqd_fields(resp.data.id, srs_row.id, sds_content)
                self.__sync_imported_sds_doc_images_to_doc_file(resp.data.id, sds_content)
                # 首次进入编辑页应展示 Word 自带追溯表；点击“获取SRS追溯”后才重算追溯表。
                sdstrace_serv.__ensure_sds_traces(doc_id=resp.data.id)
                db.session.commit()
            return resp
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    def _backfill_sds_code_from_trace_table(self, roots: List[SdsNodeForm]):
        """导入时用 word 自带的“需求-设计追溯表”按功能名称把设计编号回填到第6章下的功能节点。

        追溯表“需求/代码”列形如“用户登录(章节 6.6.1)”，取其功能名“用户登录”，与第6章
        （功能详细设计）下功能节点标题去掉章节号后的名称匹配，把“设计编号”列回填到节点
        sds_code（仅空值）。只处理第6章下功能（标准需求/变更需求对应的功能），其它章节的
        模块节点不回填；仅导入环节使用，不改标题/正文/结构，不触碰 word 解析逻辑。
        """
        if not roots:
            return
        name_to_code: Dict[str, str] = {}

        def norm_name(value: str) -> str:
            return re.sub(r"\s+", "", self._strip_sds_heading_text(str(value or "")).strip())

        def collect(nodes: List[SdsNodeForm]):
            for node in nodes or []:
                tbl = getattr(node, "table", None)
                headers = getattr(tbl, "headers", None) if tbl is not None else None
                rows = getattr(tbl, "rows", None) if tbl is not None else None
                if headers and rows:
                    code_col = None
                    ref_col = None
                    for h in headers:
                        hname = re.sub(r"\s+", "", str(getattr(h, "name", "") or ""))
                        hcode = getattr(h, "code", "") or ""
                        if not code_col and "设计编号" in hname:
                            code_col = hcode
                        elif not ref_col and ("需求/代码" in hname or "代码" in hname):
                            ref_col = hcode
                    if code_col and ref_col:
                        for row in rows:
                            if not isinstance(row, dict):
                                continue
                            raw_code = str(row.get(code_col, "") or "")
                            sds_code = raw_code.replace("\r", "").split("\n")[0].strip()
                            ref = str(row.get(ref_col, "") or "")
                            # “用户登录(章节 6.6.1)” → 仅取功能名“用户登录”
                            feat = re.sub(r"[（(]\s*章节.*$", "", ref).strip()
                            key = re.sub(r"\s+", "", feat)
                            if sds_code and key:
                                name_to_code.setdefault(key, sds_code)
                collect(getattr(node, "children", None) or [])

        collect(roots)
        if not name_to_code:
            return

        def backfill(nodes: List[SdsNodeForm]):
            for node in nodes or []:
                heading = self._parse_sds_node_heading(getattr(node, "title", "") or "")
                in_chapter6 = bool(heading) and heading.split(".")[0] == "6"
                if in_chapter6 and not (getattr(node, "sds_code", "") or "").strip():
                    key = norm_name(getattr(node, "title", "") or "")
                    if key and key in name_to_code:
                        node.sds_code = name_to_code[key]
                backfill(getattr(node, "children", None) or [])

        backfill(roots)

    def _update_nodes(self, doc: SdsDoc, p_id, nodes: List[SdsNodeForm]):
        for idx, node in enumerate(nodes):
            sql = select(SdsNode).where(SdsNode.doc_id == doc.id, SdsNode.n_id == node.n_id) if node.n_id else None
            row = db.session.execute(sql).scalars().first() if sql is not None else None
            if not row:
                doc.n_id += 1
                table = self._serialize_sds_table(node.table)
                row = SdsNode(doc_id=doc.id, n_id=doc.n_id, p_id=p_id, priority=idx, title=node.title, label=node.label, img_url=_normalize_sds_img_url(node.img_url), text=node.text, ref_type=node.ref_type,
                            table=table, sds_code=node.sds_code)
                db.session.add(row)
                node.n_id = doc.n_id
                logger.info("add_node: %s, %s, %s", p_id, doc.n_id, node.title)
            else:
                row.p_id = p_id
                for key, value in node.dict().items():
                    if key == "doc_id" or key == "n_id" or key == "p_id" or value is None:
                        continue
                    if key == "table":
                        value = self._serialize_sds_table(value)
                    if key == "img_url":
                        value = _normalize_sds_img_url(value)
                    setattr(row, key, value)
                row.priority = idx
                node.n_id = row.n_id
                logger.info("alt_node: %s, %s, %s", p_id, doc.n_id, node.title)
            if node.children:
                self._update_nodes(doc, row.n_id, node.children)

    @staticmethod
    def _reset_tree_node_ids(nodes: List[SdsNodeForm]):
        for node in nodes or []:
            node.n_id = 0
            Server._reset_tree_node_ids(getattr(node, "children", None) or [])

    async def duplicate_sds_doc(self, id: int, product_id: int = None):
        fromdoc:SdsDocObj = (await self.get_sds_doc(id, with_tree=True)).data
        if not fromdoc:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        srcrow = db.session.execute(select(SdsDoc).where(SdsDoc.id == id)).scalars().first()
        if not srcrow:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        src_srsdoc_id = srcrow.srsdoc_id or 0
        src_product_id = (srcrow.product_id or getattr(fromdoc, "product_id", 0) or 0)
        # 复制目标产品：默认沿用原产品，跨产品复制时使用指定产品
        target_pid = product_id or src_product_id
        cross = bool(product_id) and product_id != src_product_id
        if cross:
            # 跨产品：自动绑定目标产品下最新有效 SRS 文档，无则不绑定
            real_product_id = target_pid
            srs_latest = db.session.execute(
                select(SrsDoc)
                .where(SrsDoc.product_id == target_pid, ~SrsDoc.version.like(f"{DELETED_SRS_VERSION_PREFIX}%"))
                .order_by(desc(SrsDoc.create_time), desc(SrsDoc.id))
            ).scalars().first()
            real_srsdoc_id = srs_latest.id if srs_latest else 0
            # 版本：目标产品现有最大版本递增，目标产品无文档时沿用原版本
            def _version_seq(v):
                m = re.search(r"(\d+)(?!.*\d)", v or "")
                return int(m.group(1)) if m else -1
            tgt_versions = db.session.execute(select(SdsDoc.version).where(SdsDoc.product_id == target_pid)).scalars().all()
            valid = [v for v in tgt_versions if v]
            version = new_version(max(valid, key=_version_seq)) if valid else fromdoc.version
        else:
            real_product_id = src_product_id
            real_srsdoc_id = src_srsdoc_id
            version = new_version(fromdoc.version)
        # 兜底：确保目标产品/绑定SRS下版本唯一，撞号则继续递增
        def _sds_version_exists(ver):
            c1 = db.session.execute(select(func.count(SdsDoc.id)).where(SdsDoc.srsdoc_id == real_srsdoc_id, SdsDoc.version == ver)).scalar() or 0
            c2 = 0
            if real_product_id:
                c2 = db.session.execute(select(func.count(SdsDoc.id)).where(SdsDoc.product_id == real_product_id, SdsDoc.version == ver)).scalar() or 0
            return c1 > 0 or c2 > 0
        while _sds_version_exists(version):
            version = new_version(version)
        next_file_no = serv_review_util.resolve_doc_file_no(
            real_product_id,
            getattr(fromdoc, "file_no", None),
            version,
            "sds",
        )
        newdoc = SdsDoc(srsdoc_id=real_srsdoc_id, product_id=real_product_id, version=version, change_log=fromdoc.change_log, n_id=0, file_no=next_file_no)
        try:
            db.session.add(newdoc)
            db.session.flush()
            self._update_nodes(newdoc, 0, fromdoc.content)

            sdsreqds = db.session.execute(select(SdsReqd).where(SdsReqd.doc_id == fromdoc.id)).scalars().all()
            for sdsreqd in sdsreqds:
                newreqd = SdsReqd(**sdsreqd.dict())
                newreqd.id = None
                newreqd.doc_id = newdoc.id
                db.session.add(newreqd)
            sdstraces = db.session.execute(select(SdsTrace).where(SdsTrace.doc_id == fromdoc.id)).scalars().all()
            for sdstrace in sdstraces:
                newtrace = SdsTrace(**sdstrace.dict())
                newtrace.id = None
                newtrace.doc_id = newdoc.id
                db.session.add(newtrace)
            return Resp.resp_ok(data=SdsDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def add_sds_doc(self, form: SdsDocForm, preserve_word_structure: bool = False):
        def __chapter(req: SrsReq):
            return sdstrace_serv.compose_srs_req_chapter(req) or req.sub_function or req.function or req.module
        try:
            sql = select(func.count(SdsDoc.id)).where(SdsDoc.srsdoc_id == form.srsdoc_id, SdsDoc.version == form.version)
            count = db.session.execute(sql).scalar()
            if count > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))

            product_id = form.product_id or 0
            if not product_id and form.srsdoc_id:
                srs = db.session.execute(select(SrsDoc).where(SrsDoc.id == form.srsdoc_id)).scalars().first()
                product_id = (srs.product_id if srs else 0) or 0
            row = SdsDoc(srsdoc_id=form.srsdoc_id, product_id=product_id, version=form.version, change_log=form.change_log, n_id=0, file_no=serv_review_util.resolve_doc_file_no(product_id, form.file_no, form.version, "sds"))
            db.session.add(row)
            db.session.flush()
            word_imported = preserve_word_structure or self._is_word_imported_doc(form.content or [])
            if form.content:
                if not word_imported:
                    form.content = await self._ensure_trace_nodes_from_saved_locations(row.id, form.content)
                    form.content = await self._refresh_trace_table_nodes(row.id, form.content)
                    form.content = self.__normalize_heading_hierarchy(form.content)
                    form.content = self._dedupe_requirement_nodes(form.content)
                form.content = self._clear_node_ids(form.content)
                self._backfill_sds_code_from_trace_table(form.content)
                self._update_nodes(row, 0, form.content)
            srs_reqs: List[SrsReq] = db.session.execute(select(SrsReq).where(SrsReq.doc_id == form.srsdoc_id)).scalars().all()
    
            req_values = [dict(doc_id=row.id, req_id=req.id) for req in srs_reqs if req.type_code != "2"]
            if req_values:
                db.session.execute(pg_insert(SdsReqd).values(req_values).on_conflict_do_nothing())

            req_values = [dict(doc_id=row.id, req_id=req.id, sds_code=req.code.replace("SRS", "SDS"), chapter=__chapter(req)) for req in srs_reqs if req.type_code != "reqd"]
            if req_values:
                db.session.execute(pg_insert(SdsTrace).values(req_values).on_conflict_do_nothing())
                
            db.session.commit()
            return Resp.resp_ok(data=SdsDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))
    
    def __upsert_product_doc_image_from_sds_upload(
        self,
        doc_id: int,
        ref_type: str,
        file_name: str,
        file_size: int,
        file_url: str,
    ):
        if ref_type not in {"img_flow", "img_topo", "img_struct"} or not file_url:
            return
        sds_doc = db.session.execute(select(SdsDoc).where(SdsDoc.id == doc_id)).scalars().first()
        if not sds_doc:
            return
        srs_doc = db.session.execute(select(SrsDoc).where(SrsDoc.id == sds_doc.srsdoc_id)).scalars().first()
        product_id = getattr(srs_doc, "product_id", None) if srs_doc else None
        if not product_id:
            return
        row = db.session.execute(
            select(DocFile)
            .where(DocFile.product_id == product_id)
            .where(DocFile.category == ref_type)
            .order_by(desc(DocFile.id))
        ).scalars().first()
        if row is None:
            row = DocFile(product_id=product_id, category=ref_type)
            db.session.add(row)
        ext = os.path.splitext(str(file_name or ""))[1] or os.path.splitext(str(file_url or ""))[1] or ".png"
        row.file_name = f"{sds_doc.version or ref_type}_{ref_type}{ext}"
        row.file_size = file_size
        row.file_url = file_url
        row.update_time = datetime.now()
        db.session.flush()

    def __sync_imported_sds_doc_images_to_doc_file(self, doc_id: int, nodes: List[SdsNodeForm]):
        """Word 重新导入时，以 Word 内嵌图覆盖产品图表文件，避免沿用旧编辑页上传图。"""
        if not doc_id or not nodes:
            return

        def file_size_of(url: str) -> int:
            path = str(url or "").strip()
            if path.startswith("/"):
                path = path[1:]
            try:
                return os.path.getsize(path) if path and os.path.exists(path) else 0
            except Exception:
                return 0

        best_flow: Optional[SdsNodeForm] = None
        best_score = -1

        def score_flow_node(node: SdsNodeForm) -> int:
            title = f"{getattr(node, 'title', '') or ''} {getattr(node, 'label', '') or ''} {getattr(node, 'text', '') or ''}"
            norm = re.sub(r"\s+", "", title)
            if "网络安全流程图" in norm:
                return 100
            if "安全流程图" in norm:
                return 80
            if "流程图" in norm:
                return 50
            return 10

        def walk(items: List[SdsNodeForm]):
            nonlocal best_flow, best_score
            for node in items or []:
                ref_type = str(getattr(node, "ref_type", "") or "").strip()
                img_url = _normalize_sds_img_url(getattr(node, "img_url", "") or "")
                if ref_type == RefTypes.img_flow.value and img_url:
                    score = score_flow_node(node)
                    if score > best_score:
                        best_flow = node
                        best_score = score
                walk(getattr(node, "children", None) or [])

        walk(nodes or [])
        if best_flow is not None:
            img_url = _normalize_sds_img_url(getattr(best_flow, "img_url", "") or "")
            if img_url.startswith("/"):
                img_url = img_url[1:]
            self.__upsert_product_doc_image_from_sds_upload(
                doc_id,
                RefTypes.img_flow.value,
                getattr(best_flow, "title", "") or "网络安全流程图",
                file_size_of(img_url),
                img_url,
            )

    async def add_doc_file(self, doc_id: int, file, ref_type: str = None):
        size, path = await save_file("sds_node_img", doc_id, file)
        self.__upsert_product_doc_image_from_sds_upload(
            doc_id,
            str(ref_type or "").strip(),
            getattr(file, "filename", "") or "",
            size,
            path,
        )
        db.session.commit()
        return Resp.resp_ok(data=path)   
   
    async def delete_sds_doc(self, id):
        db.session.execute(delete(SdsReqd).where(SdsReqd.doc_id == id))
        db.session.execute(delete(SdsTrace).where(SdsTrace.doc_id == id))
        db.session.execute(delete(SdsNode).where(SdsNode.doc_id == id))
        db.session.execute(delete(SdsDoc).where(SdsDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def add_sds_node(self, node: SdsNodeForm):
        sql = select(SdsNode, SdsDoc).join(SdsDoc, SdsNode.doc_id == SdsDoc.id)
        sql = sql.where(SdsNode.doc_id == node.doc_id, SdsNode.n_id == node.p_id)
        result = db.session.execute(sql).first()
        if not result:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        _, doc = result
        doc.n_id += 1
        table = self._serialize_sds_table(node.table)
        row = SdsNode(doc_id=doc.id, n_id=doc.n_id, p_id=node.p_id, priority=doc.n_id, 
                            title=node.title, img_url=_normalize_sds_img_url(node.img_url), text=node.text, table=table)
        db.session.add(row)
        db.session.commit()
        return Resp.resp_ok(data=SdsNodeForm(doc_id=row.doc_id, n_id=row.n_id, p_id=row.p_id, priority=row.priority,
                            title=row.title, img_url=row.img_url, text=row.text, table=node.table))

    async def delete_sds_node(self, doc_id, n_id):
        db.session.execute(delete(SdsNode).where(SdsNode.doc_id == doc_id, SdsNode.n_id == n_id))
        db.session.commit()
        return Resp.resp_ok()
   
    async def update_sds_doc(self, form: SdsDocForm):
        try:
            sql = select(func.count(SdsDoc.id)).where(SdsDoc.srsdoc_id == form.srsdoc_id, SdsDoc.version == form.version, SdsDoc.id != form.id)
            count = db.session.execute(sql).scalar()
            if count > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            sql = select(SdsDoc).where(SdsDoc.id == form.id)
            row:SdsDoc = db.session.execute(sql).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict().items():
                if key == "id" or key == "n_id" or value is None:
                    continue
                setattr(row, key, value)
            row.file_no = srsdoc_serv._Server__sync_file_no_version(row.file_no, row.version)
            if form.content:
                word_imported = self._is_word_imported_doc(form.content or [])
                if not word_imported:
                    form.content = await self._ensure_trace_nodes_from_saved_locations(row.id, form.content)
                    form.content = await self._refresh_trace_table_nodes(row.id, form.content)
                    form.content = self.__normalize_heading_hierarchy(form.content)
                    form.content = self._dedupe_requirement_nodes(form.content)
                form.content = self._clear_node_ids(form.content)
                row.n_id = 0
                db.session.execute(delete(SdsNode).where(SdsNode.doc_id == row.id))
                self._update_nodes(row, 0, form.content)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_sds_doc_file_no(self, id: int, file_no: str):
        try:
            sql = select(SdsDoc).where(SdsDoc.id == id)
            row: SdsDoc = db.session.execute(sql).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            row.file_no = (file_no or "").strip() or None
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))
   
    def __query_imgs(self, product_id: int):
        subquery = select(DocFile.category, func.max(DocFile.id).label("max_id"))
        subquery = subquery.where(DocFile.product_id == product_id).group_by(DocFile.category).subquery()
        sql = select(DocFile).join(subquery, DocFile.id == subquery.c.max_id)
        rows: List[DocFile] = db.session.execute(sql).scalars().all()
        return {row.category: row.file_url for row in rows}


    def _load_sds_tree(self, doc_id: int) -> List[SdsNodeForm]:
        nodes: list[SdsNode] = db.session.execute(
            select(SdsNode).where(SdsNode.doc_id == doc_id).order_by(SdsNode.priority)
        ).scalars().all()
        objs_dict: Dict[int, SdsNodeForm] = {}
        roots: List[SdsNodeForm] = []
        for node in nodes:
            table = self._parse_sds_table(node.table)
            obj = SdsNodeForm(
                children=[], doc_id=node.doc_id, n_id=node.n_id, p_id=node.p_id,
                title=node.title, label=node.label, img_url=node.img_url, text=node.text,
                ref_type=node.ref_type, table=table, sds_code=node.sds_code,
            )
            objs_dict[obj.n_id] = obj
        for obj in objs_dict.values():
            if obj.p_id == 0:
                roots.append(obj)
            else:
                parent = objs_dict.get(obj.p_id)
                if parent:
                    parent.children.append(obj)
        if self._is_word_imported_doc(roots):
            self._normalize_word_imported_chapter_numbers(roots)
        return roots


    def __tree_sync_fingerprint(self, roots: List[SdsNodeForm]) -> tuple:
        parts: List[str] = []

        def walk(nodes: List[SdsNodeForm]):
            for node in nodes or []:
                parts.append(str(getattr(node, "title", "") or ""))
                parts.append(str(getattr(node, "sds_code", "") or ""))
                parts.append(str(getattr(node, "text", "") or "")[:80])
                walk(getattr(node, "children", None) or [])

        walk(roots or [])
        return tuple(parts)

    def _persist_sds_tree(self, doc: SdsDoc, roots: List[SdsNodeForm]):
        self._reset_tree_node_ids(roots)
        doc.n_id = 0
        db.session.execute(delete(SdsNode).where(SdsNode.doc_id == doc.id))
        db.session.flush()
        self._update_nodes(doc, 0, roots)
        db.session.commit()















    @staticmethod



    @staticmethod





    @staticmethod




    def __normalize_heading_hierarchy(self, roots: List[SdsNodeForm]):
        def parse_heading(value: str):
            matched = re.match(r"^\s*(\d+(?:\.\d+)*)", str(value or "").strip())
            return matched.group(1) if matched else ""

        def heading_sort_key(value: str):
            heading = parse_heading(value)
            if not heading:
                return (9999,)
            return tuple(int(part) for part in heading.split("."))

        parent_lists = {}
        by_heading = {}

        def collect(nodes: List[SdsNodeForm], parent_list: List[SdsNodeForm]):
            for node in nodes or []:
                heading = parse_heading(getattr(node, "title", "") or "")
                if heading:
                    by_heading[heading] = node
                    parent_lists[id(node)] = parent_list
                collect(getattr(node, "children", None) or [], getattr(node, "children", None) or [])

        collect(roots, roots)
        headings = sorted(by_heading.keys(), key=lambda value: len(value.split(".")))
        for heading in headings:
            if "." not in heading:
                continue
            node = by_heading.get(heading)
            parent_heading = ".".join(heading.split(".")[:-1])
            parent_node = by_heading.get(parent_heading)
            if not node or not parent_node:
                continue
            current_parent_list = parent_lists.get(id(node))
            if current_parent_list is getattr(parent_node, "children", None):
                continue
            if current_parent_list is not None and node in current_parent_list:
                current_parent_list.remove(node)
            parent_node.children = list(getattr(parent_node, "children", None) or [])
            if node not in parent_node.children:
                parent_node.children.append(node)
            parent_node.children.sort(key=lambda item: heading_sort_key(getattr(item, "title", "") or ""))
            parent_lists[id(node)] = parent_node.children
        return roots

    async def get_sds_doc(self, id:str, with_tree: bool = False):
        sql = select(SdsDoc, SrsDoc, Product).outerjoin(SrsDoc, SdsDoc.srsdoc_id == SrsDoc.id).outerjoin(Product, SdsDoc.product_id == Product.id).where(SdsDoc.id == id)
        row, row_srs, row_prd = db.session.execute(sql).first() or (None, None, None)
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        is_srs_deleted = bool(row_srs and (row_srs.version or "").startswith(DELETED_SRS_VERSION_PREFIX))
        
        tree = []
        if with_tree:
            tree = self._load_sds_tree(row.id)
            prod_imgs = self.__query_imgs(row.product_id) if row.product_id else (self.__query_imgs(row_srs.product_id) if row_srs else {})
            def hydrate_imgs(nodes: List[SdsNodeForm]):
                for obj in nodes or []:
                    if obj.ref_type in prod_imgs:
                        obj.img_url = prod_imgs[obj.ref_type]
                    hydrate_imgs(getattr(obj, "children", None) or [])
            hydrate_imgs(tree)
            tree = await self._refresh_trace_table_for_display(row.id, tree, persist=True)
        data = row.dict()
        data["srsdoc_id"] = 0 if is_srs_deleted else row.srsdoc_id
        data["product_id"] = row.product_id or (row_prd.id if row_prd else (row_srs.product_id if row_srs else 0))
        data["product_name"] = row_prd.name if row_prd else ""
        data["product_version"] = row_prd.full_version if row_prd else ""
        data["srs_version"] = "" if is_srs_deleted else (row_srs.version if row_srs else "")
        if not (data.get("file_no") or "").strip():
            data["file_no"] = serv_review_util.resolve_doc_file_no(data["product_id"], row.file_no, row.version, "sds")
        data["content"] = tree
        return Resp.resp_ok(data=SdsDocObj(**data))

    async def list_sds_doc(self, op_user: UserObj, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        page_index = page_index if page_index >= 0 else 0
        page_size = page_size if page_size > 0 else 10 
    
        sql = select(SdsDoc, SrsDoc, Product).outerjoin(SrsDoc, SdsDoc.srsdoc_id == SrsDoc.id).outerjoin(Product, SdsDoc.product_id == Product.id)
        if product_id:
            sql = sql.where(SdsDoc.product_id == product_id)
        if version:
            sql = sql.where(SdsDoc.version.like(f"%{version}%"))
        if not product_id and op_user and op_user.id != 1:
            subquery = select(UserProd.product_id).where(UserProd.user_id == op_user.id).scalar_subquery()
            sql = sql.where(Product.id.in_(subquery))
        
        sql_count = select(func.count()).select_from(sql)
        total = db.session.execute(sql_count).scalars().first()
        sql = sql.offset(page_size * page_index).limit(page_size).order_by(desc(SdsDoc.create_time))
        rows: list[SdsDoc] = db.session.execute(sql).all()

        objs = []
        for row, row_srs, row_prd in rows:
            obj = SdsDocObj(**row.dict())
            if row_prd:
                obj.product_id = row_prd.id
                obj.product_name = row_prd.name
                obj.product_version = row_prd.full_version
            if row_srs:
                obj.srs_version = "" if (row_srs.version or "").startswith(DELETED_SRS_VERSION_PREFIX) else row_srs.version
            if not (obj.file_no or "").strip():
                obj.file_no = serv_review_util.resolve_doc_file_no(obj.product_id or row.product_id, row.file_no, row.version, "sds")
            objs.append(obj)
        return Resp.resp_ok(data=Page(total=total, page_size=page_size, rows=objs, page_index=page_index))

    async def export_sds_doc(self, output, id: int = 0, *args, **kwargs):
        if Document is None or Pt is None or dox_enum is None:
            return
        from .serv_utils import docx_util
        def __normalize_req_code(value: str):
            return re.sub(r"\s+", "", str(value or "").strip().upper())

        def __compare_req_code(a: str, b: str):
            ax = [int(x) for x in re.findall(r"\d+", __normalize_req_code(a))]
            bx = [int(x) for x in re.findall(r"\d+", __normalize_req_code(b))]
            for idx in range(max(len(ax), len(bx))):
                av = ax[idx] if idx < len(ax) else 0
                bv = bx[idx] if idx < len(bx) else 0
                if av != bv:
                    return av - bv
            left = __normalize_req_code(a)
            right = __normalize_req_code(b)
            return (left > right) - (left < right)

        def __to_sds_code(value: str):
            code = __normalize_req_code(value)
            if not code:
                return ""
            return re.sub(r"^SRS-", "SDS-", code)

        def __norm_title(value: str):
            txt = (value or "").strip()
            txt = re.sub(r"\s+", " ", txt)
            # 仅清理异常前导符号，不改章节号数值
            txt = re.sub(r"^[\s\u3000•·▪■◆●○□◇\-–—\.．]+(?=[0-9０-９A-Za-z\u4e00-\u9fff])", "", txt)
            return txt

        def __is_cover_section_title(title: str):
            txt = __biz_title(title)
            return txt in ["软件详细设计", "软件详细设计说明书", "文件修订记录"]

        def __is_pure_punct_line(value: str):
            txt = (value or "").strip()
            return re.match(r"^[\s\u3000•·▪■◆●○□◇\-–—\.．:：,，;；_]+$", txt) is not None

        def __biz_title(value: str):
            txt = __norm_title(value)
            txt = re.sub(r"^([0-9０-９]+(?:[\.．][0-9０-９]+)*)(?:[\s、.．]+|(?=[\u4e00-\u9fffA-Za-z]))", "", txt)
            return re.sub(r"\s+", "", txt)

        def __is_revision_label(value: str):
            return __biz_title(value) == "文件修订记录"

        def __is_catalog(value: str):
            return __biz_title(value) == "目录"

        def __is_design_cover(value: str):
            return __biz_title(value) in ["软件详细设计", "软件详细设计说明书"]

        def __is_rev_title(value: str):
            return __biz_title(value) == "文件修订记录"

        def __parse_heading(value: str):
            txt = __norm_title(value)
            matched = re.match(r"^([0-9]+(?:\.[0-9]+)*)(?:[\s、.．]+|(?=[\u4e00-\u9fffA-Za-z]))(.*)$", txt)
            if not matched:
                return None, txt
            nums = [int(p) for p in (matched.group(1) or "").split(".") if p != ""]
            if not nums:
                return None, txt
            return nums, (matched.group(2) or "").strip()

        def __strip_heading_text(value: str):
            _nums, rest = __parse_heading(value)
            return (rest or __norm_title(value)).strip()

        def __normalize_req_title(value: str):
            raw = str(value or "").strip()
            stripped = __strip_heading_text(raw)
            return re.sub(r"\s+", "", stripped or raw).lower()

        def __compose_req_description(row: SdsReqdObj):
            return "\n".join([
                f"(1) 总体描述\n{(getattr(row, 'overview', None) or '').strip() or '无'}",
                f"(2) 功能\n{(getattr(row, 'func_detail', None) or '').strip() or '无'}",
                f"(3) 程序逻辑\n{(getattr(row, 'logic_txt', None) or '').strip() or '无'}",
                f"(4) 输入项\n{(getattr(row, 'intput', None) or '').strip() or '无'}",
                f"(5) 输出项\n{(getattr(row, 'output', None) or '').strip() or '无'}",
                f"(6) 接口\n{(getattr(row, 'interface', None) or '').strip() or '无'}",
            ])

        def __req_hierarchy_titles(row: SdsReqdObj):
            titles = []
            for value in [getattr(row, "module", None), getattr(row, "function", None), getattr(row, "sub_function", None)]:
                txt = str(value or "").strip()
                if not txt:
                    continue
                norm = __normalize_req_title(txt)
                if norm and norm not in [item[0] for item in titles]:
                    titles.append((norm, txt))
            fallback = str(getattr(row, "name", None) or getattr(row, "srs_code", None) or "").strip()
            return [item[1] for item in titles] or ([fallback] if fallback else [])

        def __find_function_area_node(nodes: List[SdsNodeForm]):
            for node in nodes or []:
                nums, rest = __parse_heading(getattr(node, "title", "") or "")
                title_txt = __normalize_req_title(rest or getattr(node, "title", "") or "")
                if nums == [6] or "功能设计" in title_txt:
                    return node
                found = __find_function_area_node(getattr(node, "children", None) or [])
                if found:
                    return found
            return None

        def __is_function_stopper(value: str):
            text = __normalize_req_title(value)
            return "限制条件" in text or "尚未解决的问题" in text

        def __with_heading(title: str, heading: str):
            return f"{heading} {__strip_heading_text(title)}".strip()

        def __assign_child_headings(node: SdsNodeForm):
            nums, _rest = __parse_heading(getattr(node, "title", "") or "")
            if not nums:
                return node
            for idx, child in enumerate(getattr(node, "children", None) or []):
                child_heading = ".".join([*(str(num) for num in nums), str(idx + 1)])
                child.title = __with_heading(getattr(child, "title", "") or "", child_heading)
                __assign_child_headings(child)
            return node

        def __renumber_direct_children(parent: SdsNodeForm):
            parent_nums, _ = __parse_heading(getattr(parent, "title", "") or "")
            if not parent_nums:
                return
            for idx, child in enumerate(getattr(parent, "children", None) or []):
                nums, _rest = __parse_heading(getattr(child, "title", "") or "")
                if len(nums or []) != len(parent_nums) + 1:
                    continue
                heading = ".".join([*(str(num) for num in parent_nums), str(idx + 1)])
                child.title = __with_heading(getattr(child, "title", "") or "", heading)

        def __append_hierarchy_row(roots: List[SdsNodeForm], row: SdsReqdObj, code: str):
            level_nodes = roots
            titles = __req_hierarchy_titles(row)
            for idx, title in enumerate(titles):
                is_leaf = idx == len(titles) - 1
                norm = __normalize_req_title(title)
                target = next((node for node in level_nodes if __normalize_req_title(getattr(node, "title", "") or "") == norm), None)
                if target is None:
                    target = SdsNodeForm(
                        title=title,
                        sds_code=code if is_leaf else "",
                        text=__compose_req_description(row) if is_leaf else "",
                        children=[],
                    )
                    level_nodes.append(target)
                elif is_leaf and not getattr(target, "sds_code", None):
                    target.sds_code = code
                    target.text = target.text or __compose_req_description(row)
                if target.children is None:
                    target.children = []
                level_nodes = target.children

        async def __sync_missing_reqd_nodes_for_export(roots: List[SdsNodeForm]):
            try:
                resp = await sdstreqd_serv.list_sds_reqd(None, doc_id=id, page_index=0, page_size=10000)
                req_rows: List[SdsReqdObj] = resp.data.rows if resp and resp.data else []
            except Exception:
                logger.exception("sync export sds reqd nodes failed")
                return roots
            if not req_rows:
                return roots
            existing_codes = set()
            def collect(nodes: List[SdsNodeForm]):
                for node in nodes or []:
                    code = __normalize_req_code(getattr(node, "sds_code", "") or "")
                    if code:
                        existing_codes.add(code)
                    collect(getattr(node, "children", None) or [])
            collect(roots)

            missing_rows = []
            for row in req_rows:
                type_code = str(getattr(row, "type_code", "") or "").strip()
                if type_code in ["1", "2"]:
                    continue
                code = __to_sds_code(getattr(row, "srs_code", "") or "")
                if code and code not in existing_codes:
                    missing_rows.append((code, row))
            if not missing_rows:
                return roots
            function_node = __find_function_area_node(roots)
            if not function_node:
                return roots
            if function_node.children is None:
                function_node.children = []
            virtual_roots: List[SdsNodeForm] = []
            for code, row in sorted(missing_rows, key=lambda item: item[0]):
                __append_hierarchy_row(virtual_roots, row, code)
            if not virtual_roots:
                return roots

            insert_index = len(function_node.children)
            for idx, child in enumerate(function_node.children):
                if __is_function_stopper(getattr(child, "title", "") or ""):
                    insert_index = idx
                    break
            parent_nums, _ = __parse_heading(getattr(function_node, "title", "") or "")
            if not parent_nums:
                return roots
            start_index = insert_index + 1
            for idx, node in enumerate(virtual_roots):
                node.title = __with_heading(getattr(node, "title", "") or "", ".".join([*(str(num) for num in parent_nums), str(start_index + idx)]))
                __assign_child_headings(node)
            function_node.children[insert_index:insert_index] = virtual_roots
            __renumber_direct_children(function_node)
            logger.info("export synced missing sds reqd nodes: doc=%s count=%s", id, len(virtual_roots))
            return roots

        def __split_trace_lines(value: str):
            lines = [line.strip() for line in str(value or "").replace("\r", "").split("\n")]
            while len(lines) > 1 and not lines[-1]:
                lines.pop()
            return lines or [""]

        def __build_sds_location_map_from_export_tree(nodes: List[SdsNodeForm]):
            result = {}
            def walk(items: List[SdsNodeForm]):
                for item in items or []:
                    nums, _rest = __parse_heading(getattr(item, "title", "") or "")
                    heading = ".".join([str(num) for num in nums]) if nums else ""
                    code = __normalize_req_code(getattr(item, "sds_code", "") or "")
                    if code and heading and code not in result:
                        result[code] = heading
                    walk(getattr(item, "children", None) or [])
            walk(nodes or [])
            return result

        def __is_trace_table_node(node: SdsNodeForm):
            title = __biz_title(getattr(node, "title", "") or "")
            ref_type = str(getattr(node, "ref_type", "") or "")
            return ref_type == RefTypes.sds_traces.value or "设计与需求追溯表" in title or "设计与需求追溯列表" in title

        def __is_change_trace_row(row: SdsTraceObj):
            type_code = str(getattr(row, "type_code", "") or "").strip()
            return bool(type_code) and type_code not in ["1", "2"]

        def __make_trace_change_table_title(product_full_version: str):
            version = str(product_full_version or "").strip()
            return f"{version or '产品'}变更需求"

        def __build_trace_table_from_rows(rows: List[SdsTraceObj], location_by_sds_code: Dict[str, str]):
            headers = [
                TabHeader(code="srs_code", name="需求编号"),
                TabHeader(code="sds_code", name="设计编号"),
                TabHeader(code="chapter", name="需求/代码"),
            ]

            def build_chapter_cell(row: SdsTraceObj):
                sds_codes = __split_trace_lines(getattr(row, "sds_code", "") or "")
                chapters = sdstrace_serv.trace_chapter_lines(
                    getattr(row, "chapter", "") or "",
                    srs_code=getattr(row, "srs_code", None),
                    sub_function=getattr(row, "sub_function", None),
                    function=getattr(row, "function", None),
                    module=getattr(row, "module", None),
                )
                locations = __split_trace_lines(getattr(row, "location", "") or "")
                count = max(1, len(sds_codes), len(chapters), len(locations))
                values = []
                for idx in range(count):
                    chapter = chapters[idx].strip() if idx < len(chapters) else (chapters[0].strip() if chapters else "")
                    sds_code = __normalize_req_code(sds_codes[idx] if idx < len(sds_codes) else "")
                    location = locations[idx].strip() if idx < len(locations) else ""
                    if not location and sds_code:
                        location = location_by_sds_code.get(sds_code, "")
                    values.append(f"{chapter}{f'（章节 {location}）' if location else ''}")
                return "\n".join(values)

            table_rows = []
            for row in rows or []:
                table_rows.append({
                    "srs_code": getattr(row, "srs_code", "") or "",
                    "sds_code": getattr(row, "sds_code", "") or "",
                    "chapter": build_chapter_cell(row),
                })
            return Table(headers=headers, rows=table_rows)

        async def __query_latest_sds_trace_tables(roots: List[SdsNodeForm]):
            resp = await sdstrace_serv.list_sds_trace(None, doc_id=id, page_size=10000)
            reqs: List[SdsTraceObj] = resp.data.rows or []
            location_by_sds_code = __build_sds_location_map_from_export_tree(roots)
            normal_rows = [row for row in reqs if not __is_change_trace_row(row)]
            change_rows = [row for row in reqs if __is_change_trace_row(row)]
            normal_table = __build_trace_table_from_rows(normal_rows, location_by_sds_code)
            change_table = __build_trace_table_from_rows(change_rows, location_by_sds_code)
            change_version = ""
            for row in change_rows:
                change_version = str(getattr(row, "product_version", "") or "").strip()
                if change_version:
                    break
            return normal_table, change_table, change_rows, change_version

        async def __sync_trace_table_nodes_for_export(roots: List[SdsNodeForm]):
            try:
                normal_table, change_table, change_rows, change_version = await __query_latest_sds_trace_tables(roots)
            except Exception:
                logger.exception("sync export sds trace table failed")
                return roots

            def is_old_trace_table_child(child: SdsNodeForm):
                title = __biz_title(getattr(child, "title", "") or "")
                label = __biz_title(getattr(child, "label", "") or "")
                has_table = getattr(child, "table", None) and child.table.headers
                return bool(has_table and (re.match(r"^导入表格\d*$", title or "") or label.endswith("变更需求") or title.endswith("变更需求")))

            def walk(items: List[SdsNodeForm]):
                for node in items or []:
                    if __is_trace_table_node(node):
                        node.ref_type = ""
                        node.table = normal_table
                        kept_children = [child for child in (getattr(node, "children", None) or []) if not is_old_trace_table_child(child)]
                        if change_rows:
                            kept_children.append(SdsNodeForm(
                                label=__make_trace_change_table_title(change_version),
                                table=change_table,
                                children=[],
                            ))
                        node.children = kept_children
                    walk(getattr(node, "children", None) or [])
            walk(roots or [])
            return roots

        def __major_of_text(value: str):
            nums, _ = __parse_heading(value)
            if not nums:
                return None
            return nums[0]

        def __first_major(nodes: List[SdsNodeForm]):
            for node in nodes or []:
                for val in [getattr(node, "title", ""), getattr(node, "label", "")]:
                    if __is_cover_section_title(val) or __is_catalog(val):
                        continue
                    major = __major_of_text(val)
                    if major and major > 0:
                        return major
                child_major = __first_major(getattr(node, "children", None) or [])
                if child_major:
                    return child_major
            return None

        def __shift_heading(value: str, major_offset: int):
            txt = __norm_title(value)
            if major_offset <= 0:
                return txt
            nums, rest = __parse_heading(txt)
            if not nums:
                return txt
            nums[0] = max(1, nums[0] - major_offset)
            prefix = ".".join(str(n) for n in nums)
            return f"{prefix} {rest}".rstrip()

        def __is_imported_placeholder_title(value: str):
            txt = (value or "").strip()
            return re.match(r"^导入(表格|图片)\d*$", txt) is not None

        def __is_imported_table_title(value: str):
            return re.match(r"^导入表格\d*$", (value or "").strip()) is not None

        def __is_imported_image_title(value: str):
            return re.match(r"^导入图片\d*$", (value or "").strip()) is not None

        def __is_table_caption_line(line: str):
            txt = (line or "").strip()
            if not txt:
                return False
            # JSON 键值行不是表题（如 "code":0, / "filename":"a.zip"）
            if re.match(r'^[\'"]\s*[^\'"]+\s*[\'"]\s*:\s*.+$', txt):
                return False
            if re.match(r"^\s*表\s*\d+\s*", txt):
                return True
            # 仅将“字段名: 值”这类无空格英文标识识别为表题，避免误判整句正文
            if re.match(r"^[A-Za-z][A-Za-z0-9_]{1,64}[:：]\s*.+$", txt):
                return True
            if re.search(r"[:：]", txt) and len(txt) <= 80 and re.search(r"[。！？]$", txt) is None:
                parts = [seg.strip() for seg in re.split(r"[:：]", txt)]
                left = parts[0] if parts else ""
                right = "".join(parts[1:]).strip() if len(parts) > 1 else ""
                left_is_identifier = re.match(r"^[A-Za-z][A-Za-z0-9_]{1,64}$", left or "") is not None
                if left and right and (left_is_identifier or "表" in left):
                    return True
                if left and not right and "表" in left:
                    return True
            if "表" in txt and re.match(r"^.+表\s*[:：]?$", txt):
                return True
            return False

        def __strip_chapter_prefix(value: str):
            txt = __norm_title(value)
            return re.sub(
                r"^([0-9]+(?:\.[0-9]+)*)(?:[\s、.．]+|(?=[\u4e00-\u9fffA-Za-z]))",
                "",
                txt,
            ).strip()

        def __is_data_table_title(value: str):
            nums, rest = __parse_heading(value)
            if nums and rest:
                # “2.4 设计与需求追溯表”这类带章节号的标题是章节，不是表题。
                # 只有明确的“表 N ...”或英文标识符表名才继续按表题处理。
                if not re.match(r"^(表|table)\s*\d+", rest, re.I) and not re.match(r"^[A-Za-z][A-Za-z0-9_]{1,64}\s*[:：]?$", rest):
                    return False
            txt = __strip_chapter_prefix(value)
            if not txt:
                return False
            if __is_table_caption_line(txt):
                return True
            if re.match(r"^[A-Za-z][A-Za-z0-9_]{1,64}\s*[:：]?$", txt):
                return True
            return False

        def __is_database_heading_title(value: str):
            txt = __strip_chapter_prefix(value)
            if not txt:
                return False
            return re.search(r"数据库\s*[:：]?$", txt) is not None

        def __looks_like_body_text_title(value: str):
            txt = __norm_title(value)
            if not txt:
                return False
            txt_no_mark = re.sub(r"^[\s\u3000•·▪■◆●○□◇\-–—]+", "", txt).strip()
            txt_body = re.sub(
                r"^([0-9]+(?:\.[0-9]+)*)(?:[\s、.．]+|(?=[\u4e00-\u9fffA-Za-z\"']))",
                "",
                txt_no_mark,
            ).strip()
            has_heading_prefix = txt_body != txt_no_mark
            probe = txt_body or txt_no_mark
            # JSON / 字典片段：不应作为章节标题
            if re.match(r'^[\'"]\s*[A-Za-z0-9_\-]+\s*[\'"]\s*:\s*.+$', probe):
                return True
            # JSON 标量值行（数组元素）也不应作为章节标题
            if (not has_heading_prefix) and re.match(r'^(?:".*"|-?\d+(?:\.\d+)?|true|false|null)\s*,?$', probe, re.I):
                return True
            if re.match(r'^[\{\[\}].*$', probe):
                return True
            if re.match(r'^.*[:：]\s*[\{\[]\s*$', probe):
                return True
            if re.match(r'^.*[,，]\s*$', probe) and (":" in probe or "：" in probe):
                return True
            # 句子型长文本（含中文标点）在导出中按正文处理，不作为章节
            if re.search(r"[，,。；;：:！？!?]", probe):
                return True
            if len(probe) > 24:
                return True
            return False

        def __is_image_caption_line(line: str):
            return re.match(r"^\s*图\s*\d+\s*", (line or "").strip()) is not None

        def __is_only_table_caption_text(text: str):
            lines = [(line or "").strip() for line in str(text or "").splitlines() if (line or "").strip()]
            if not lines:
                return False
            return all(__is_table_caption_line(line) for line in lines)

        def __is_only_image_caption_text(text: str):
            lines = [(line or "").strip() for line in str(text or "").splitlines() if (line or "").strip()]
            if not lines:
                return False
            return all(__is_image_caption_line(line) for line in lines)

        def __image_export_key(value: str):
            txt = str(value or "").strip()
            if not txt:
                return ""
            if txt.startswith("data:image/"):
                return txt
            return txt.split("?", 1)[0].lstrip("/")

        def __normalize_json_block_order(lines: List[str]) -> List[str]:
            clean_lines = [str(line or "").rstrip() for line in (lines or [])]
            if not clean_lines:
                return clean_lines
            first_kv_idx = next((idx for idx, line in enumerate(clean_lines) if re.match(r'^\s*[\'"]\s*[^\'"]+\s*[\'"]\s*:\s*.+$', line)), -1)
            if first_kv_idx < 0:
                return clean_lines
            version_idx = next((idx for idx, line in enumerate(clean_lines) if re.match(r'^\s*[\'"]\s*version\s*[\'"]\s*:\s*.+$', line, re.I)), -1)
            if version_idx < 0 or version_idx <= first_kv_idx:
                return clean_lines
            version_line = clean_lines.pop(version_idx)
            clean_lines.insert(first_kv_idx, version_line)
            return clean_lines

        def __is_json_kv_line(line: str):
            txt = str(line or "").strip()
            if not txt:
                return False
            if re.match(r'^\s*[\'"]\s*[^\'"]+\s*[\'"]\s*:\s*.+$', txt):
                return True
            txt_wo_chapter = re.sub(r'^(\d+(?:\.\d+)*)(?:[\s、.．]+|(?=[\u4e00-\u9fffA-Za-z"\']))', '', txt).strip()
            return re.match(r'^\s*[\'"]\s*[^\'"]+\s*[\'"]\s*:\s*.+$', txt_wo_chapter) is not None

        def __insert_json_line_before_first_kv(lines: List[str], json_line: str):
            items = [str(line or "").rstrip() for line in (lines or [])]
            line = str(json_line or "").strip()
            if not line:
                return items
            first_kv_idx = next((idx for idx, it in enumerate(items) if __is_json_kv_line(it)), -1)
            if first_kv_idx >= 0:
                items.insert(first_kv_idx, line)
                return items
            brace_idx = next((idx for idx, it in enumerate(items) if str(it).strip() == "{"), -1)
            if brace_idx >= 0:
                items.insert(brace_idx + 1, line)
                return items
            items.append(line)
            return items

        def __is_json_export_line(line: str):
            txt = str(line or "").strip()
            if not txt:
                return False
            if __is_json_kv_line(txt):
                return True
            if re.match(r'^\s*[\'"]\s*[^\'"]+\s*[\'"]\s*:\s*$', txt):
                return True
            return re.match(r'^[\{\}\[\]],?$', txt) is not None

        def __is_json_value_line(line: str):
            txt = str(line or "").strip()
            if not txt:
                return False
            return re.match(r'^(?:".*"|-?\d+(?:\.\d+)?|true|false|null)\s*,?$', txt, re.I) is not None

        def __format_json_like_lines(lines: List[str]) -> List[str]:
            raw_lines = [str(line or "").strip() for line in (lines or []) if str(line or "").strip()]
            if not raw_lines:
                return []
            if not any(__is_json_export_line(line) for line in raw_lines):
                return raw_lines
            out: List[str] = []
            indent = 0
            in_json_context = False
            for raw in raw_lines:
                line = raw
                if __is_json_kv_line(line):
                    line = __strip_chapter_prefix(line) or line
                opens_block = re.search(r'[\{\[]\s*,?$', line) is not None
                closes_block = re.match(r'^[\}\]],?$', line) is not None
                is_kv = re.match(r'^[\'"]([^\'"]+)[\'"]\s*:\s*(.+?)(,?)$', line) is not None
                is_scalar_value = re.match(r'^(?:".*"|-?\d+(?:\.\d+)?|true|false|null)\s*,?$', line, re.I) is not None

                if closes_block:
                    indent = max(0, indent - 1)

                formatted = line
                kv = re.match(r'^[\'"]([^\'"]+)[\'"]\s*:\s*(.+?)(,?)$', line)
                if kv:
                    key = kv.group(1)
                    val = (kv.group(2) or "").strip()
                    comma = kv.group(3) or ""
                    formatted = f'"{key}": {val}'
                    if comma and not formatted.endswith(","):
                        formatted += ","

                if __is_json_export_line(line) or (in_json_context and (is_scalar_value or is_kv)):
                    formatted = (" " * (4 * indent)) + formatted
                out.append(formatted)

                if __is_json_export_line(line) or is_kv or is_scalar_value:
                    in_json_context = True
                if opens_block and not closes_block:
                    indent += 1
            return out

        def __strip_explicit_bullet_prefix(text: str):
            raw = str(text or "")
            m = re.match(r"^\s*(?:[•●▪◦·\uf0b7]|\-|\*)\s+(.+?)\s*$", raw)
            if m:
                return m.group(1), True
            return raw, False

        def __is_bullet_intro_line(line: str):
            txt = str(line or "").strip()
            if not txt:
                return False
            if re.match(r"^[（(]\s*\d+\s*[）)]\s*(功能|步骤|流程|操作说明)\s*$", txt):
                return True
            return re.search(r"(如下|下列|包括|满足下列|技术要求|部署要求)\s*[:：]?$", txt) is not None

        def __is_numbered_section_line(line: str):
            txt = str(line or "").strip()
            if not txt:
                return False
            return re.match(r"^[（(]\s*\d+\s*[）)]\s*[^:：\n]{0,24}$", txt) is not None

        def __is_force_bullet_section_line(line: str):
            txt = str(line or "").strip()
            if not txt:
                return False
            return re.match(r"^[（(]\s*\d+\s*[）)]\s*(功能|步骤|流程|操作说明|实现|关键点|要点)\s*$", txt) is not None

        def __can_render_as_bullet_content(line: str):
            txt = str(line or "").strip()
            if not txt:
                return False
            if __is_numbered_section_line(txt):
                return False
            if __is_table_caption_line(txt) or __is_image_caption_line(txt):
                return False
            if __is_json_export_line(txt) or __is_json_value_line(txt):
                return False
            return True

        def __is_operation_step_line(line: str):
            txt = str(line or "").strip()
            if not txt:
                return False
            return re.match(
                r"^(点\s*击|单\s*击|双\s*击|输\s*入|选\s*择|填\s*写|若|如\s*果|然\s*后|最\s*后|确\s*认|保\s*存|删\s*除|编\s*辑|新\s*增|重\s*置)",
                txt,
            ) is not None

        def __is_bullet_candidate_line(line: str):
            txt = str(line or "").strip()
            if not txt:
                return False
            if len(txt) < 4 or len(txt) > 200:
                return False
            if txt.endswith(("：", ":")):
                return False
            if re.match(r"^\d+(?:\.\d+)+(?:[\s、.．]+|$)", txt):
                return False
            if re.match(r"^[（(]\s*\d+\s*[）)]", txt):
                return False
            if __is_table_caption_line(txt) or __is_image_caption_line(txt):
                return False
            if __is_json_export_line(txt) or __is_json_value_line(txt):
                return False
            if __is_operation_step_line(txt):
                return True
            return re.match(r"^(支持|提供|开启|定期|按照|根据|设置|接收|导入|对|可|需|必须|禁止|允许|具备|包含|检查|维护|升级)", txt) is not None

        def __save_line_txt(
            docx: Document,
            text: str,
            font_size: float = 10.5,
            is_json: bool = False,
            is_bullet: bool = False,
        ):
            txt = str(text or "")
            if not txt.strip():
                return
            leading_spaces = len(txt) - len(txt.lstrip(" "))
            json_level = max(0, leading_spaces // 4) if is_json else 0
            # JSON 内容去掉前导空格，改由段落缩进体现层级（可视化更稳定）
            render_txt = txt.strip() if is_json else txt.strip()
            if is_bullet:
                p = docx.add_paragraph(style="List Bullet")
            else:
                p = docx.add_paragraph()
            p.alignment = dox_enum.text.WD_ALIGN_PARAGRAPH.LEFT
            if not is_bullet:
                # 详细设计正文统一左对齐且不首行缩进
                p.paragraph_format.first_line_indent = Pt(0)
                # JSON 使用固定段落缩进表达层级，避免不同字体下空格缩进显示不稳定
                p.paragraph_format.left_indent = Pt(json_level * 10) if is_json else Pt(0)
            p.paragraph_format.right_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, render_txt, font_size)

        def __save_body_line_auto_bullet(
            docx: Document,
            text: str,
            font_size: float,
            is_json: bool,
            bullet_state: dict,
            allow_bullet: bool = True,
        ):
            raw = str(text or "")
            stripped = raw.strip()
            if not stripped:
                return
            is_bullet = False
            render_txt = raw
            if (not is_json) and allow_bullet:
                render_txt, explicit_bullet = __strip_explicit_bullet_prefix(raw)
                if __is_numbered_section_line(stripped):
                    # 遇到新的“小节标识”时重置列表状态；若是功能/步骤小节则开启强制列表模式
                    bullet_state["active"] = False
                    bullet_state["remain"] = 0
                    bullet_state["force_mode"] = __is_force_bullet_section_line(stripped)
                if explicit_bullet:
                    is_bullet = True
                    bullet_state["active"] = True
                    bullet_state["remain"] = 12
                elif __is_operation_step_line(stripped):
                    is_bullet = True
                    bullet_state["active"] = True
                    bullet_state["remain"] = 12
                elif bullet_state.get("force_mode") and __can_render_as_bullet_content(stripped):
                    is_bullet = True
                else:
                    if __is_bullet_intro_line(stripped):
                        bullet_state["active"] = True
                        bullet_state["remain"] = 12
                    elif (
                        bullet_state.get("active")
                        and int(bullet_state.get("remain", 0)) > 0
                        and __is_bullet_candidate_line(stripped)
                    ):
                        is_bullet = True
                        bullet_state["remain"] = int(bullet_state.get("remain", 0)) - 1
                    elif len(stripped) > 160 and not bullet_state.get("force_mode"):
                        bullet_state["active"] = False
                        bullet_state["remain"] = 0
            elif not allow_bullet:
                # 接口/JSON章节强制保持纯左对齐正文，避免列表样式带来的额外缩进
                bullet_state["active"] = False
                bullet_state["remain"] = 0
                bullet_state["force_mode"] = False
            __save_line_txt(docx, render_txt, font_size, is_json, is_bullet=is_bullet)

        def __split_interface_io_text(raw_text: str):
            lines = [str(line or "") for line in str(raw_text or "").replace("\r", "").split("\n")]
            output_idx = next(
                (idx for idx, line in enumerate(lines) if re.search(r'[（(]\s*2\s*[）)]\s*输出项', (line or "").strip())),
                -1
            )
            if output_idx <= 0:
                return None
            before = [line for line in lines[:output_idx]]
            after = [line for line in lines[output_idx:]]
            return before, after

        def __is_revision_table(table):
            if not table:
                return False
            header_txt = "".join((getattr(h, "name", "") or "").strip() for h in (getattr(table, "headers", None) or []))
            keys = ["修改日期", "版本号", "修订说明", "修订人", "批准人"]
            return sum(1 for key in keys if key in header_txt) >= 3

        def __insert_toc_field(docx: Document):
            if OxmlElement is None or qn is None:
                return
            p = docx.add_paragraph()
            run_begin = p.add_run()
            fld_begin = OxmlElement("w:fldChar")
            fld_begin.set(qn("w:fldCharType"), "begin")
            fld_begin.set(qn("w:dirty"), "true")
            run_begin._r.append(fld_begin)
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = ' TOC \\o "1-4" \\h \\z \\u '
            run_begin._r.append(instr)

            fld_separate = OxmlElement("w:fldChar")
            fld_separate.set(qn("w:fldCharType"), "separate")
            run_begin._r.append(fld_separate)
            p.add_run("目录将在打开文档后自动更新")

            run_end = p.add_run()
            fld_end = OxmlElement("w:fldChar")
            fld_end.set(qn("w:fldCharType"), "end")
            run_end._r.append(fld_end)

        def __write_center_section_title(docx: Document, title: str):
            p = docx.add_paragraph()
            p.alignment = dox_enum.text.WD_ALIGN_PARAGRAPH.CENTER
            font_size = 22.0 if __is_design_cover(title) else 16.0
            docx_util.fonted_txt(p, title, font_size=font_size)

        def __write_catalog_text(docx: Document, catalog_text: str):
            __write_center_section_title(docx, "目录")
            for raw in str(catalog_text or "").replace("\r", "").split("\n"):
                line = (raw or "").strip()
                if line:
                    docx_util.save_txt2docx(line, docx, 10.5)

        def __add_blank_lines(docx: Document, line_count: int):
            for _ in range(max(0, line_count)):
                docx.add_paragraph("")

        def __refresh_toc_with_libreoffice(output_stream):
            soffice = shutil.which("soffice") or shutil.which("libreoffice")
            if not soffice:
                logger.warning("LibreOffice not found, skip SDS TOC refresh")
                return False
            try:
                output_stream.seek(0)
                source_bytes = output_stream.read()
                with tempfile.TemporaryDirectory(prefix="sds_docx_refresh_") as tmpdir:
                    input_path = os.path.join(tmpdir, "input.docx")
                    output_path = os.path.join(tmpdir, "output.docx")
                    profile_dir = os.path.join(tmpdir, "lo_profile")
                    script_path = os.path.join(tmpdir, "refresh_toc.py")
                    with open(input_path, "wb") as f:
                        f.write(source_bytes)
                    script = r'''
import sys
import time
import uno
from com.sun.star.beans import PropertyValue

input_path = sys.argv[1]
output_path = sys.argv[2]

def prop(name, value):
    item = PropertyValue()
    item.Name = name
    item.Value = value
    return item

local_ctx = uno.getComponentContext()
resolver = local_ctx.ServiceManager.createInstanceWithContext(
    "com.sun.star.bridge.UnoUrlResolver",
    local_ctx,
)
ctx = None
last_error = None
for _ in range(60):
    try:
        ctx = resolver.resolve("uno:socket,host=127.0.0.1,port=2002;urp;StarOffice.ComponentContext")
        break
    except Exception as exc:
        last_error = exc
        time.sleep(0.5)
if ctx is None:
    raise RuntimeError(f"connect libreoffice failed: {last_error}")

desktop = ctx.ServiceManager.createInstanceWithContext("com.sun.star.frame.Desktop", ctx)
doc = desktop.loadComponentFromURL(
    uno.systemPathToFileUrl(input_path),
    "_blank",
    0,
    (
        prop("Hidden", True),
        prop("ReadOnly", False),
        prop("UpdateDocMode", 3),
    ),
)
if doc is None:
    raise RuntimeError("load docx failed")
try:
    indexes = doc.getDocumentIndexes()
    for idx in range(indexes.getCount()):
        indexes.getByIndex(idx).update()

    fields = doc.getTextFields()
    enum = fields.createEnumeration()
    while enum.hasMoreElements():
        field = enum.nextElement()
        try:
            field.update()
        except Exception:
            pass

    doc.storeAsURL(
        uno.systemPathToFileUrl(output_path),
        (
            prop("FilterName", "Office Open XML Text"),
            prop("Overwrite", True),
        ),
    )
finally:
    doc.close(True)
'''
                    with open(script_path, "w", encoding="utf-8") as f:
                        f.write(script)
                    server = subprocess.Popen(
                        [
                            soffice,
                            "--headless",
                            "--nologo",
                            "--nodefault",
                            "--nofirststartwizard",
                            "--nolockcheck",
                            f"-env:UserInstallation=file://{profile_dir}",
                            "--accept=socket,host=127.0.0.1,port=2002;urp;StarOffice.ServiceManager",
                        ],
                        stdout=subprocess.DEVNULL,
                        stderr=subprocess.DEVNULL,
                    )
                    try:
                        env = os.environ.copy()
                        dist_pkg = "/usr/lib/python3/dist-packages"
                        env["PYTHONPATH"] = f"{dist_pkg}:{env.get('PYTHONPATH', '')}"
                        subprocess.run(
                            ["python3", script_path, input_path, output_path],
                            check=True,
                            timeout=120,
                            env=env,
                        )
                    finally:
                        server.terminate()
                        try:
                            server.wait(timeout=10)
                        except subprocess.TimeoutExpired:
                            server.kill()
                    if not os.path.exists(output_path) or os.path.getsize(output_path) <= 0:
                        output_stream.seek(0)
                        return False
                    with open(output_path, "rb") as f:
                        refreshed = f.read()
                    output_stream.seek(0)
                    output_stream.truncate(0)
                    output_stream.write(refreshed)
                    output_stream.seek(0)
                    return True
            except Exception:
                logger.exception("refresh SDS TOC with LibreOffice failed")
                output_stream.seek(0)
                return False

        def __write_revision_body_title(docx: Document):
            p = docx.add_paragraph()
            p.alignment = dox_enum.text.WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.keep_with_next = True
            docx_util.fonted_txt(p, "文件修订记录", font_size=14.0, bold=True)

        def __node_has_revision_marker(node: SdsNodeForm):
            for val in [getattr(node, "title", ""), getattr(node, "label", ""), getattr(node, "text", "")]:
                if __is_revision_label(__norm_title(val or "")):
                    return True
            return False

        # 详细设计导出图片统一尺寸区间（像素）：
        # - 大图缩小，小图适度放大，最终视觉尺寸统一且可读
        IMG_MAX_W = 300
        IMG_MAX_H = 200
        IMG_MIN_W = 120
        IMG_MIN_H = 90
        IMG_TARGET_LONG = 200
        # 6章节“功能类/程序逻辑”图片再缩小一档，避免单图占据过大版面
        IMG_MAX_W_FUNC = 260
        IMG_MAX_H_FUNC = 180
        IMG_MIN_W_FUNC = 110
        IMG_MIN_H_FUNC = 80
        IMG_TARGET_LONG_FUNC = 180

        export_state = {"pending_rev_label": False, "pending_table_caps": [], "pending_image_caps": []}

        def __save_caption_txt(docx: Document, text: str, font_size: float = 10.5, align: str = "left"):
            txt = (text or "").strip()
            if not txt:
                return
            p = docx.add_paragraph()
            p.alignment = dox_enum.text.WD_ALIGN_PARAGRAPH.LEFT if align == "left" else dox_enum.text.WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, txt, font_size)

        def __save_table_caption_txt(docx: Document, text: str, font_size: float = 10.5):
            __save_caption_txt(docx, text, font_size, align="left")

        def __save_image_caption_txt(docx: Document, text: str, font_size: float = 10.5):
            __save_caption_txt(docx, text, font_size, align="center")

        def __flush_table_caption(docx: Document, font_size: float = 10.5):
            if export_state["pending_table_caps"]:
                cap = export_state["pending_table_caps"].pop(0)
                __save_table_caption_txt(docx, cap, font_size)

        def __flush_image_caption(docx: Document, font_size: float = 10.5):
            if export_state["pending_image_caps"]:
                cap = export_state["pending_image_caps"].pop(0)
                __save_image_caption_txt(docx, cap, font_size)
        async def __query_sds_traces_x():
            resp = await sdstrace_serv.list_sds_trace(None, doc_id=id, page_size=5000)
            reqs: List[SdsTraceObj] = resp.data.rows or []
            reqs_dict = dict()
            for req in reqs:
                reqs_dict.setdefault((req.type_code, req.type_name), []).append(req)
            
            results = []
            for (type_code, type_name), reqs in reqs_dict.items():
                headers = [TabHeader(code="srs_code", name="需求编号"), 
                       TabHeader(code="sds_code", name="设计编号"), 
                       TabHeader(code="chapter", name="需求/代码")]
                rows = []
                for req in reqs:
                    row = dict()
                    location = f"（章节 {req.location}） " if req.location else ""
                    row["srs_code"] = req.srs_code
                    row["sds_code"] = req.sds_code
                    row["chapter"] = (req.chapter or "") + location
                    rows.append(row)
                table = Table(headers=headers, rows=rows)
                results.append(SdsNodeForm(label=type_name, table=table))
            return results

        def __fix_chapter(p_title: str, nodes: List[SdsNodeForm]):
            chapter =re.search(r'(\d(\.\d)*)', p_title or "")
            chapter = chapter.group() if chapter else None
            chapter = f"{chapter}." if chapter else ""
            for idx, node in enumerate(nodes or []):
                if node.with_chapter == 1 and chapter and node.title:
                    node.title = f"{chapter}{idx+1} {node.title}"
                    __fix_chapter(node.title, node.children)


        def __query_sds_logics(reqd_ids):
            result_dict = dict()
            if not reqd_ids:
                return result_dict
            sql = select(Logic).where(Logic.reqd_id.in_(reqd_ids)).order_by(Logic.id)
            rows: List[Logic] = db.session.execute(sql).scalars().all()
            for row in rows:
                reqd_id = row.reqd_id
                logics = result_dict.get(reqd_id, [])
                logics.append(SdsNodeForm(img_url=row.img_url))
                logics.append(SdsNodeForm(text=row.txt))
                result_dict[reqd_id] = logics
            return result_dict

        async def __query_sds_reqds(p_title: str):
            resp = await sdstreqd_serv.list_sds_reqd(None, doc_id=id, page_size=2000)
            reqds: List[SdsReqdObj] = resp.data.rows or []
            reqd_ids = [reqd.id for reqd in reqds]
            sds_logics = __query_sds_logics(reqd_ids)
            parents = dict()
            for idx, reqd in enumerate(reqds):
                with_chapter = 1 if reqd.sub_function else 0
                title = reqd.name if reqd.sub_function else None

                node = SdsNodeForm(with_chapter=with_chapter, title=title, children=[])
                node.children.append(SdsNodeForm(label="（一）总体描述", text=reqd.overview))
                node.children.append(SdsNodeForm(label="（二）功能", text=reqd.func_detail))

                node.children.append(SdsNodeForm(label="（三）程序逻辑", text=reqd.logic_txt))
                logics = sds_logics.get(reqd.id, [])
                node.children.extend(logics)

                node.children.append(SdsNodeForm(label="（四）输入项", text=reqd.intput))
                node.children.append(SdsNodeForm(label="（五）输出项", text=reqd.output))
                node.children.append(SdsNodeForm(label="（六）接口", text=reqd.interface))
                p_node = find_parent(SdsNodeForm, [reqd.module, reqd.function], parents)
                p_node.children.append(node)
            p_nodes = [node for key, node in parents.items() if node.level == 0]
            __fix_chapter(p_title, p_nodes)
            return p_nodes

        def __is_compact_img_context(node: SdsNodeForm):
            title = __norm_title(getattr(node, "title", "") or "")
            label = __norm_title(getattr(node, "label", "") or "")
            text = __norm_title(getattr(node, "text", "") or "")
            is_ch6 = re.match(r"^6(?:\.\d+)*(?:[\s、.．]+|(?=[\u4e00-\u9fffA-Za-z]))", title) is not None
            is_func_logic = ("程序逻辑" in title) or ("程序逻辑" in label) or ("程序逻辑" in text)
            return is_ch6 or is_func_logic

        async def __writenodes(
            nodes: List[SdsNodeForm],
            docx: Document,
            level: int = 0,
            major_offset: int = 0,
            compact_img: bool = False,
            force_plain_context: bool = False,
        ):
            font_def = 10.5
            font_size = font_def
            if level == 0 :
                font_size = 16.0
            elif level == 1:
                font_size = 14.0
            for node in nodes or []:
                node_compact_img = compact_img or __is_compact_img_context(node)
                img_max_w = IMG_MAX_W_FUNC if node_compact_img else IMG_MAX_W
                img_max_h = IMG_MAX_H_FUNC if node_compact_img else IMG_MAX_H
                img_min_w = IMG_MIN_W_FUNC if node_compact_img else IMG_MIN_W
                img_min_h = IMG_MIN_H_FUNC if node_compact_img else IMG_MIN_H
                img_target_long = IMG_TARGET_LONG_FUNC if node_compact_img else IMG_TARGET_LONG
                written_child_ids = set()
                table_written = False
                image_written = False
                node_text_effective = str(node.text or "")
                pending_table_captions: List[str] = []
                pending_image_captions: List[str] = []
                imported_table_children = [
                    child for child in (node.children or [])
                    if getattr(child, "table", None) and child.table.headers
                ]
                imported_image_children = [
                    child for child in (node.children or [])
                    if getattr(child, "img_url", None)
                ]
                duplicate_child_image_keys = {
                    __image_export_key(getattr(child, "img_url", "") or "")
                    for child in imported_image_children
                    if __image_export_key(getattr(child, "img_url", "") or "")
                }
                suppress_own_duplicate_image = bool(
                    getattr(node, "img_url", None)
                    and __image_export_key(getattr(node, "img_url", "") or "") in duplicate_child_image_keys
                )
                imported_db_children = [
                    child for child in (node.children or [])
                    if __is_database_heading_title(getattr(child, "title", "") or "")
                ]
                node_title_norm = __norm_title(getattr(node, "title", "") or "")
                node_label_norm = __norm_title(getattr(node, "label", "") or "")
                node_text_hint = str(node_text_effective or "")
                is_interface_io_context = (
                    re.match(r"^\d+\.7\.[2-5](?:[\s、.．]+|$)", node_title_norm) is not None
                    or (
                        re.search(r"[（(]\s*1\s*[）)]\s*输入项", node_text_hint) is not None
                        and re.search(r"[（(]\s*2\s*[）)]\s*输出项", node_text_hint) is not None
                    )
                    or ("接口" in node_title_norm and ("输入项" in node_text_hint or "输出项" in node_text_hint))
                    or ("json示例" in node_text_hint.lower())
                    or ("json" in node_label_norm.lower())
                )
                current_plain_context = force_plain_context or is_interface_io_context
                program_logic_image_written = False

                def __image_caption_from_node(img_node: SdsNodeForm):
                    for value in [
                        getattr(img_node, "title", ""),
                        getattr(img_node, "label", ""),
                        getattr(img_node, "text", ""),
                    ]:
                        txt = __norm_title(value or "")
                        if __is_image_caption_line(txt):
                            return txt
                    return ""

                def __try_write_program_logic_image_after_line(line: str):
                    nonlocal program_logic_image_written
                    if program_logic_image_written:
                        return
                    if "程序逻辑" not in (line or ""):
                        return
                    candidates = [
                        child for child in imported_image_children
                        if builtins.id(child) not in written_child_ids
                    ]
                    matched = next(
                        (
                            child for child in candidates
                            if re.search(r"程序逻辑|逻辑图", __image_caption_from_node(child))
                        ),
                        None,
                    )
                    if matched is None:
                        return
                    docx_util.save_img2docx(
                        matched.img_url,
                        docx,
                        mw=img_max_w,
                        mh=img_max_h,
                        min_w=img_min_w,
                        min_h=img_min_h,
                        target_long=img_target_long,
                    )
                    caption = __image_caption_from_node(matched)
                    if caption:
                        __save_image_caption_txt(docx, caption, font_def)
                    written_child_ids.add(builtins.id(matched))
                    program_logic_image_written = True

                # 子节点中若存在“仅 JSON 键值标题（如 "version":4,）”，并入父节点正文，
                # 这样导出顺序可与编辑页保持一致，且不会单独漂移到后段。
                for child in (node.children or []):
                    c_title = __norm_title(getattr(child, "title", "") or "")
                    c_text = str(getattr(child, "text", "") or "").strip()
                    c_has_payload = bool(
                        (getattr(child, "table", None) and getattr(getattr(child, "table", None), "headers", None))
                        or getattr(child, "img_url", None)
                        or c_text
                        or (getattr(child, "children", None) or [])
                    )
                    if c_title and (__is_json_export_line(c_title) or __is_json_value_line(c_title)):
                        normalized_title = __strip_chapter_prefix(c_title) or c_title
                        if __is_json_kv_line(c_title):
                            node_text_effective = "\n".join(
                                __insert_json_line_before_first_kv(
                                    str(node_text_effective or "").replace("\r", "").split("\n"),
                                    normalized_title
                                )
                            )
                        else:
                            node_text_effective = "\n".join(
                                [*str(node_text_effective or "").replace("\r", "").split("\n"), normalized_title]
                            ).strip()
                        if c_has_payload:
                            # 保留子节点承载的图/表/正文结构，但清空 JSON 标题，避免重复输出
                            child.title = ""
                        else:
                            # 纯 JSON 占位子节点在并入父节点后可直接跳过
                            written_child_ids.add(builtins.id(child))
                is_catalog_root = level == 0 and __is_catalog(node.title)
                if node.title:
                    norm_title = __norm_title(node.title)
                    if not norm_title or __is_pure_punct_line(norm_title):
                        pass
                    elif __is_imported_placeholder_title(norm_title):
                        pass
                    elif __is_data_table_title(norm_title) and ((node.table and node.table.headers) or imported_table_children):
                        # 数据表标题不是章节：不写入Heading，避免进入Word目录
                        pending_table_captions.append(__strip_chapter_prefix(norm_title))
                    elif __is_image_caption_line(norm_title) and (node.img_url or imported_image_children):
                        # 图片标题不是章节：导出时放到图片下方作为题注
                        pending_image_captions.append(norm_title)
                    elif __is_database_heading_title(norm_title):
                        # 数据结构下的“库X数据库:”是章节标题，不是普通正文；需要走章节号偏移。
                        docx_util.save_title2docx(__shift_heading(norm_title, major_offset), docx, level+1, font_size)
                    elif __looks_like_body_text_title(norm_title):
                        # 形如 `"version": 4,` 的内容是正文，不是章节
                        __save_line_txt(
                            docx,
                            norm_title,
                            font_def,
                            __is_json_export_line(norm_title) or __is_json_value_line(norm_title),
                        )
                    elif is_catalog_root:
                        if str(getattr(node, "text", "") or "").strip():
                            __write_catalog_text(docx, node.text)
                        else:
                            __write_center_section_title(docx, "目录")
                            __insert_toc_field(docx)
                    elif level == 0 and __is_cover_section_title(norm_title):
                        if __is_design_cover(norm_title):
                            # 与SRS导出版式一致：封面标题上方保留10行
                            __add_blank_lines(docx, 10)
                        __write_center_section_title(docx, norm_title if __is_rev_title(norm_title) else "软件详细设计")
                        # 与SRS导出版式一致：封面/修订标题与下方内容保持固定留白
                        __add_blank_lines(docx, 9 if __is_design_cover(norm_title) else 2)
                    else:
                        docx_util.save_title2docx(__shift_heading(norm_title, major_offset), docx, level+1, font_size)
                if is_catalog_root:
                    # 目录页由TOC域生成，不再输出旧目录节点文本和子节点
                    continue
                if node.sds_code:
                    __save_line_txt(docx, "设计编号：" + node.sds_code, font_def, False)
                if node.label:
                    norm_label = __norm_title(node.label)
                    if not norm_label or __is_pure_punct_line(norm_label):
                        pass
                    elif __is_revision_label(norm_label):
                        export_state["pending_rev_label"] = True
                    elif __is_table_caption_line(norm_label) and ((node.table and node.table.headers) or imported_table_children):
                        pending_table_captions.append(norm_label)
                    elif __is_image_caption_line(norm_label) and (node.img_url or imported_image_children):
                        pending_image_captions.append(norm_label)
                    elif __is_table_caption_line(norm_label):
                        export_state["pending_table_caps"].append(norm_label)
                    elif __is_image_caption_line(norm_label):
                        export_state["pending_image_caps"].append(norm_label)
                    else:
                        __save_line_txt(
                            docx,
                            norm_label,
                            font_def,
                            __is_json_export_line(norm_label) or __is_json_value_line(norm_label),
                        )
                if node_text_effective:
                    raw_text_effective = str(node_text_effective or "")
                    norm_text = __norm_title(raw_text_effective)
                    if __is_revision_label(norm_text):
                        export_state["pending_rev_label"] = True
                    # “仅图题/仅表题”必须基于原始多行文本判断，不能用 __norm_title 压平成一行后判断
                    # 否则会把“图 15 ... + 正文”整段误判成图题，导致后续正文被居中输出。
                    elif __is_only_table_caption_text(raw_text_effective):
                        lines = [(line or "").strip() for line in raw_text_effective.splitlines() if (line or "").strip()]
                        if node.table and node.table.headers:
                            # 表题下置：先表后题
                            docx_util.save_tab2docx(node.table, docx)
                            table_written = True
                            for line in lines:
                                __save_table_caption_txt(docx, line, font_def)
                        elif imported_table_children:
                            for idx, line in enumerate(lines):
                                if idx < len(imported_table_children):
                                    tab_node = imported_table_children[idx]
                                    docx_util.save_tab2docx(tab_node.table, docx)
                                    __flush_table_caption(docx, font_def)
                                    table_written = True
                                    written_child_ids.add(builtins.id(tab_node))
                                __save_table_caption_txt(docx, line, font_def)
                        else:
                            export_state["pending_table_caps"].extend(lines)
                    elif __is_only_image_caption_text(raw_text_effective):
                        lines = [(line or "").strip() for line in raw_text_effective.splitlines() if (line or "").strip()]
                        if node.img_url:
                            # 图题下置：先图后题
                            docx_util.save_img2docx(
                                node.img_url,
                                docx,
                                mw=img_max_w,
                                mh=img_max_h,
                                min_w=img_min_w,
                                min_h=img_min_h,
                                target_long=img_target_long,
                            )
                            __flush_image_caption(docx, font_def)
                            image_written = True
                            for line in lines:
                                __save_image_caption_txt(docx, line, font_def)
                        elif imported_image_children:
                            for idx, line in enumerate(lines):
                                if idx < len(imported_image_children):
                                    img_node = imported_image_children[idx]
                                    docx_util.save_img2docx(
                                        img_node.img_url,
                                        docx,
                                        mw=img_max_w,
                                        mh=img_max_h,
                                        min_w=img_min_w,
                                        min_h=img_min_h,
                                        target_long=img_target_long,
                                    )
                                    __flush_image_caption(docx, font_def)
                                    image_written = True
                                    written_child_ids.add(builtins.id(img_node))
                                __save_image_caption_txt(docx, line, font_def)
                        else:
                            export_state["pending_image_caps"].extend(lines)
                    else:
                        split_io = __split_interface_io_text(node_text_effective)
                        has_own_table = bool(node.table and node.table.headers)
                        has_child_table = len(imported_table_children) > 0
                        if split_io and (has_own_table or has_child_table) and not table_written:
                            before_lines, after_lines = split_io
                            before_lines = __format_json_like_lines(__normalize_json_block_order(before_lines))
                            after_lines = __format_json_like_lines(__normalize_json_block_order(after_lines))
                            bullet_state = {"active": False, "remain": 0, "force_mode": False}

                            for raw_line in before_lines:
                                line_raw = str(raw_line or "")
                                line = line_raw.strip()
                                if not line:
                                    continue
                                if imported_db_children and __is_database_heading_title(line):
                                    continue
                                if __is_table_caption_line(line):
                                    pending_table_captions.append(line)
                                elif __is_image_caption_line(line):
                                    pending_image_captions.append(line)
                                else:
                                    __save_body_line_auto_bullet(
                                        docx,
                                        line_raw,
                                        font_def,
                                        __is_json_export_line(line_raw) or __is_json_value_line(line_raw),
                                        bullet_state,
                                        allow_bullet=not current_plain_context,
                                    )
                                    __try_write_program_logic_image_after_line(line)

                            if has_own_table:
                                docx_util.save_tab2docx(node.table, docx)
                            else:
                                first_tab_node = imported_table_children[0]
                                docx_util.save_tab2docx(first_tab_node.table, docx)
                                written_child_ids.add(builtins.id(first_tab_node))
                            __flush_table_caption(docx, font_def)
                            table_written = True
                            if pending_table_captions:
                                for cap in pending_table_captions:
                                    __save_table_caption_txt(docx, cap, font_def)
                                pending_table_captions = []

                            for raw_line in after_lines:
                                line_raw = str(raw_line or "")
                                line = line_raw.strip()
                                if not line:
                                    continue
                                if __is_table_caption_line(line):
                                    export_state["pending_table_caps"].append(line)
                                elif __is_image_caption_line(line):
                                    export_state["pending_image_caps"].append(line)
                                else:
                                    __save_body_line_auto_bullet(
                                        docx,
                                        line_raw,
                                        font_def,
                                        __is_json_export_line(line_raw) or __is_json_value_line(line_raw),
                                        bullet_state,
                                        allow_bullet=not current_plain_context,
                                    )
                        else:
                            normalized_lines = __format_json_like_lines(__normalize_json_block_order(str(node_text_effective or "").splitlines()))
                            tcp_anchor_table_children: List[SdsNodeForm] = []
                            if (not has_own_table) and len(imported_table_children) >= 2:
                                # 历史导入数据中，多张“端口表”常挂在同一父节点下，正文仍保留在父节点 text。
                                # 命中稳定语义锚点时，按原始阅读顺序内联写出，避免多表被连续挤在一起。
                                if re.search(r"TCP", str(node_text_effective or ""), re.I):
                                    tcp_anchor_table_children = list(imported_table_children)
                            first_tcp_table_written = False
                            second_tcp_table_written = False
                            bullet_state = {"active": False, "remain": 0, "force_mode": False}
                            for raw_line in normalized_lines:
                                line_raw = str(raw_line or "")
                                line = line_raw.strip()
                                if not line:
                                    continue
                                if imported_db_children and __is_database_heading_title(line):
                                    continue
                                if __is_table_caption_line(line):
                                    if (node.table and node.table.headers) or imported_table_children:
                                        pending_table_captions.append(line)
                                    else:
                                        export_state["pending_table_caps"].append(line)
                                elif __is_image_caption_line(line):
                                    if node.img_url or imported_image_children:
                                        pending_image_captions.append(line)
                                    else:
                                        export_state["pending_image_caps"].append(line)
                                else:
                                    __save_body_line_auto_bullet(
                                        docx,
                                        line_raw,
                                        font_def,
                                        __is_json_export_line(line_raw) or __is_json_value_line(line_raw),
                                        bullet_state,
                                        allow_bullet=not current_plain_context,
                                    )
                                    __try_write_program_logic_image_after_line(line)
                                    if tcp_anchor_table_children:
                                        if (not first_tcp_table_written) and re.search(r"提供下列\s*TCP\s*服务", line, re.I):
                                            first_tab_node = tcp_anchor_table_children[0]
                                            if getattr(first_tab_node, "table", None) and first_tab_node.table.headers:
                                                docx_util.save_tab2docx(first_tab_node.table, docx)
                                                __flush_table_caption(docx, font_def)
                                                written_child_ids.add(builtins.id(first_tab_node))
                                                first_tcp_table_written = True
                                        elif (
                                            first_tcp_table_written
                                            and (not second_tcp_table_written)
                                            and len(tcp_anchor_table_children) > 1
                                            and re.search(r"只能访问.*TCP\s*端口", line, re.I)
                                        ):
                                            second_tab_node = tcp_anchor_table_children[1]
                                            if getattr(second_tab_node, "table", None) and second_tab_node.table.headers:
                                                docx_util.save_tab2docx(second_tab_node.table, docx)
                                                __flush_table_caption(docx, font_def)
                                                written_child_ids.add(builtins.id(second_tab_node))
                                                second_tcp_table_written = True
                elif node.img_url and not image_written and not suppress_own_duplicate_image:
                    docx_util.save_img2docx(
                        node.img_url,
                        docx,
                        mw=img_max_w,
                        mh=img_max_h,
                        min_w=img_min_w,
                        min_h=img_min_h,
                        target_long=img_target_long,
                    )
                    __flush_image_caption(docx, font_def)
                    image_written = True
                if image_written and pending_image_captions:
                    for cap in pending_image_captions:
                        __save_image_caption_txt(docx, cap, font_def)
                    pending_image_captions = []
                if table_written and pending_table_captions:
                    for cap in pending_table_captions:
                        __save_table_caption_txt(docx, cap, font_def)
                    pending_table_captions = []

                if node.ref_type == RefTypes.sds_traces.value:
                    results = await __query_sds_traces_x()
                    await __writenodes(results, docx, level + 1, major_offset, node_compact_img, current_plain_context)
                elif node.ref_type == RefTypes.sds_reqds.value:
                    sds_reqds = await __query_sds_reqds(node.title)
                    await __writenodes(sds_reqds, docx, level + 1, major_offset, node_compact_img, current_plain_context)
                else:
                    if node.table and node.table.headers and not table_written:
                        if export_state["pending_rev_label"] and __is_revision_table(node.table):
                            __write_revision_body_title(docx)
                            export_state["pending_rev_label"] = False
                        docx_util.save_tab2docx(node.table, docx)
                        __flush_table_caption(docx, font_def)
                        table_written = True
                        for cap in pending_table_captions:
                            __save_table_caption_txt(docx, cap, font_def)
                        pending_table_captions = []
                        
                if node.children:
                    next_children = [child for child in node.children if builtins.id(child) not in written_child_ids]
                    logic_image_children = [
                        child for child in next_children
                        if getattr(child, "img_url", None)
                        and re.search(r"程序逻辑|逻辑图", __image_caption_from_node(child))
                    ]
                    if logic_image_children:
                        logic_image_ids = {builtins.id(child) for child in logic_image_children}
                        reordered_children = []
                        inserted_logic_images = False
                        for child in next_children:
                            if builtins.id(child) in logic_image_ids:
                                continue
                            reordered_children.append(child)
                            child_marker = "\n".join([
                                str(getattr(child, "title", "") or ""),
                                str(getattr(child, "label", "") or ""),
                                str(getattr(child, "text", "") or ""),
                            ])
                            if (not inserted_logic_images) and "程序逻辑" in child_marker:
                                reordered_children.extend(logic_image_children)
                                inserted_logic_images = True
                        if not inserted_logic_images:
                            reordered_children.extend(logic_image_children)
                        next_children = reordered_children
                    await __writenodes(next_children, docx, level + 1, major_offset, node_compact_img, current_plain_context)
                # 兜底：图/表在子节点时，将父节点标题下置到子节点图/表之后
                if pending_image_captions:
                    for cap in pending_image_captions:
                        __save_image_caption_txt(docx, cap, font_def)
                if pending_table_captions:
                    for cap in pending_table_captions:
                        __save_table_caption_txt(docx, cap, font_def)
                    pending_table_captions = []
            while export_state["pending_image_caps"]:
                __flush_image_caption(docx, font_def)
            while export_state["pending_table_caps"]:
                __flush_table_caption(docx, font_def)

        async def __writenodes_legacy(nodes: List[SdsNodeForm], docx: Document, level: int = 0):
            # 兜底导出：尽量保证可导出成功，避免接口直接报错
            font_def = 10.5
            font_size = 16.0 if level == 0 else (14.0 if level == 1 else font_def)
            for node in nodes or []:
                if node.title:
                    docx_util.save_title2docx(__norm_title(node.title), docx, level + 1, font_size)
                if node.sds_code:
                    docx_util.save_txt2docx("设计编号：" + node.sds_code, docx, font_def)
                if node.label:
                    docx_util.save_txt2docx(__norm_title(node.label), docx, font_def)
                if node.text:
                    docx_util.save_txt2docx(node.text, docx, font_def)
                if node.img_url:
                    docx_util.save_img2docx(
                        node.img_url,
                        docx,
                        mw=IMG_MAX_W,
                        mh=IMG_MAX_H,
                        min_w=IMG_MIN_W,
                        min_h=IMG_MIN_H,
                        target_long=IMG_TARGET_LONG,
                    )
                if node.table and node.table.headers:
                    docx_util.save_tab2docx(node.table, docx)
                if node.children:
                    await __writenodes_legacy(node.children, docx, level + 1)

        resp = await self.get_sds_doc(id=id, with_tree=True)
        sds_doc: SdsDocObj = resp.data
        if sds_doc:
            docx = Document()
            try:
                if OxmlElement is not None and qn is not None:
                    try:
                        update_fields = OxmlElement("w:updateFields")
                        update_fields.set(qn("w:val"), "true")
                        docx.settings.element.append(update_fields)
                    except Exception:
                        logger.exception("enable sds docx updateFields failed")

                header_para = docx.sections[0].header.add_paragraph()
                header_para.alignment = dox_enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                docx_util.fonted_txt(header_para, sds_doc.file_no)

                roots = await __sync_missing_reqd_nodes_for_export(sds_doc.content or [])
                roots = await __sync_trace_table_nodes_for_export(roots)
                design_root = next((n for n in roots if __is_design_cover(getattr(n, "title", ""))), None)
                rev_root = next((n for n in roots if __is_revision_label(getattr(n, "title", ""))), None)
                catalog_root = next((n for n in roots if __is_catalog(getattr(n, "title", ""))), None)
                used_ids = {builtins.id(node) for node in [design_root, rev_root, catalog_root] if node is not None}
                remaining_roots = [n for n in roots if builtins.id(n) not in used_ids]

                # 兼容历史导入数据：有些文档把“封面/修订记录/正文”都挂在“软件详细设计”根节点下
                # 这里做一次拆分，保证第一页/第二页样式稳定一致
                design_section_nodes = [design_root] if design_root else [SdsNodeForm(title="软件详细设计", children=[])]
                rev_section_nodes = [rev_root] if rev_root else []
                body_from_design = []
                if design_root:
                    cover_node = SdsNodeForm(title="软件详细设计", children=[])
                    cover_table_picked = False
                    rev_nodes_from_design = []
                    for child in (design_root.children or []):
                        if __node_has_revision_marker(child) or __is_revision_table(getattr(child, "table", None)):
                            rev_nodes_from_design.append(child)
                            continue
                        if (not cover_table_picked) and getattr(child, "table", None) and not __is_revision_table(child.table):
                            cover_node.children.append(child)
                            cover_table_picked = True
                            continue
                        body_from_design.append(child)

                    design_section_nodes = [cover_node]
                    if (not rev_section_nodes) and rev_nodes_from_design:
                        rev_section_nodes = [SdsNodeForm(title="文件修订记录", children=rev_nodes_from_design)]

                if not rev_section_nodes:
                    rev_section_nodes = [SdsNodeForm(title="文件修订记录", children=[])]

                remaining_roots = body_from_design + remaining_roots
                first_major = __first_major(remaining_roots)
                body_major_offset = (first_major - 1) if (first_major and first_major > 1) else 0

                export_sections = [
                    ("design", design_section_nodes),
                    ("rev", rev_section_nodes),
                    ("catalog", [catalog_root] if catalog_root else [SdsNodeForm(title="目录", children=[])]),
                    ("body", remaining_roots),
                ]
                first_section = True
                for section_name, section_nodes in export_sections:
                    if not section_nodes:
                        continue
                    if not first_section:
                        docx.add_page_break()
                    major_offset = body_major_offset if section_name == "body" else 0
                    await __writenodes(section_nodes, docx, level=0, major_offset=major_offset)
                    if section_name == "rev":
                        # 与SRS导出版式一致：修订记录页末保留5行
                        __add_blank_lines(docx, 5)
                    first_section = False
            except Exception:
                logger.exception("export_sds_doc styled-export failed, fallback to legacy exporter")
                # 重建文档，确保兜底导出不受前面失败状态影响
                docx = Document()
                header_para = docx.sections[0].header.add_paragraph()
                header_para.alignment = dox_enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                docx_util.fonted_txt(header_para, sds_doc.file_no)
                await __writenodes_legacy(sds_doc.content or [], docx, level=0)

            docx_util.fill_toc_cache(docx)
            docx.save(output)
            output.seek(0)
            __refresh_toc_with_libreoffice(output)

    async def get_sds_doc_txts(self, doc_id):
        def __gather_nodes(texts:List[str],nodes: List[SdsNodeForm]):
            for node in nodes:
                values = [node.title, node.text]
                values = [value for value in values if value]
                texts += values
                if node.children:
                    __gather_nodes(texts, node.children)
            return texts

        docdata: Resp[SdsDocObj] = (await self.get_sds_doc(doc_id, with_tree=True)).data
        content = docdata.content if docdata and docdata.content else []
        txts = __gather_nodes([], content)
        return Resp.resp_ok(data=txts)

    async def compare_sds_doc(self, id0: int, id1: int):
        def __feature_key(code: str, module: str, function: str):
            # 判定新增/减少时仅按功能编号，避免名称改动造成误判
            code = (code or "").strip()
            if code:
                return code
            module = (module or "").strip()
            function = (function or "").strip()
            return " - ".join([v for v in [module, function] if v])

        def __feature_display(code: str, module: str, function: str):
            code = (code or "").strip()
            module = (module or "").strip()
            function = (function or "").strip()
            name = " - ".join([v for v in [module, function] if v])
            if code and name:
                return f"{code} {name}"
            return code or name

        def __to_text(values: List[str]):
            return "；".join(values) if values else "无"

        def __query_feature_maps():
            feature_dict = {id0: set(), id1: set()}
            feature_name_dict = {id0: {}, id1: {}}
            sql = select(SdsReqd, SrsReq).join(SrsReq, SdsReqd.req_id == SrsReq.id)
            sql = sql.where(SdsReqd.doc_id.in_([id0, id1])).order_by(SdsReqd.doc_id, SrsReq.module, SrsReq.function, SrsReq.code)
            rows: List[Tuple[SdsReqd, SrsReq]] = db.session.execute(sql).all()
            for reqd, req in rows:
                key = __feature_key(req.code, req.module, req.function)
                if not key:
                    continue
                feature_dict.setdefault(reqd.doc_id, set()).add(key)
                feature_name_dict.setdefault(reqd.doc_id, {}).setdefault(
                    key, __feature_display(req.code, req.module, req.function) or key
                )

            trace_rows: List[SdsTrace] = db.session.execute(
                select(SdsTrace).where(SdsTrace.doc_id.in_([id0, id1]))
            ).scalars().all()
            for trace in trace_rows:
                key = __feature_key(trace.sds_code, "", "")
                if not key:
                    continue
                feature_dict.setdefault(trace.doc_id, set()).add(key)
                feature_name_dict.setdefault(trace.doc_id, {}).setdefault(key, key)
            return feature_dict, feature_name_dict

        sql = select(SdsDoc, SrsDoc, Product).join(SrsDoc, SdsDoc.srsdoc_id == SrsDoc.id).join(Product, SrsDoc.product_id == Product.id).where(SdsDoc.id.in_([id0, id1]))
        rows: List[Tuple[SdsDoc, SrsDoc, Product]] = db.session.execute(sql).all()
        if not rows:
            return Resp.resp_err(msg=ts("msg_obj_null"))

        feature_dict, feature_name_dict = __query_feature_maps()
        features0 = feature_dict.get(id0) or set()
        features1 = feature_dict.get(id1) or set()
        added0_keys = sorted(features0 - features1)
        added1_keys = sorted(features1 - features0)
        added0 = [feature_name_dict.get(id0, {}).get(key, key) for key in added0_keys]
        added1 = [feature_name_dict.get(id1, {}).get(key, key) for key in added1_keys]
        removed0 = added1
        removed1 = added0

        infos = {}
        for row_sdsdoc, row_srsdoc, row_prd in rows:
            infos[row_sdsdoc.id] = dict(
                product_name=row_prd.name,
                product_type_code=row_prd.type_code,
                product_version=row_prd.full_version,
                product_udi=row_prd.udi,
                product_scope=row_prd.scope,
                srs_version=row_srsdoc.version,
                sds_version=row_sdsdoc.version,
            )
        info0 = infos.get(id0) or {}
        info1 = infos.get(id1) or {}

        results = []
        for column in ["product_name", "product_type_code", "product_version", "product_udi", "product_scope", "srs_version", "sds_version"]:
            value0 = info0.get(column) or ""
            value1 = info1.get(column) or ""
            same_flag = 1 if value0 == value1 else 0
            results.append(CompareObj(column_code=column, column_name=ts(f"sdsdiff.{column}"), same_flag=same_flag, values=[value0, value1]))

        results += [
            CompareObj(
                column_code="feature_added",
                column_name="新增功能",
                same_flag=1 if not added0 and not added1 else 0,
                values=[__to_text(added0), __to_text(added1)],
            ),
            CompareObj(
                column_code="feature_removed",
                column_name="减少功能",
                same_flag=1 if not removed0 and not removed1 else 0,
                values=[__to_text(removed0), __to_text(removed1)],
            ),
        ]
        return Resp.resp_ok(data=results)
        