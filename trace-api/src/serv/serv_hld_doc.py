#!/usr/bin/env python
# encoding: utf-8

import base64
import io
import json
import logging
import os
import re
from collections import defaultdict
from datetime import datetime
from typing import List, Optional

from sqlalchemy import delete, desc, func, select
try:
    from docx import Document
    from docx.text.paragraph import Paragraph
except Exception:
    Document = None
    Paragraph = None

from ..model.doc_file import DocFile
from ..model.hld_doc import HldDoc, HldNode
from ..model.product import Product, UserProd
from ..model.sds_doc import SdsDoc
from ..obj import Page, Resp, c_ok
from ..obj.tobj_hld_doc import HldDocForm, HldNodeForm
from ..obj.tobj_srs_doc import Table
from ..obj.vobj_hld_doc import HldDocObj
from ..obj.vobj_user import UserObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from . import serv_review_util
from .serv_srs_doc import Server as ServSrsDoc
from .serv_utils import new_version, sync_file_no_version

logger = logging.getLogger(__name__)
srsdoc_serv = ServSrsDoc()


def _normalize_hld_img_url(url: Optional[str]) -> Optional[str]:
    if not url:
        return url
    s = str(url).strip()
    if "?" not in s:
        return s
    path, _ = s.split("?", 1)
    if path.startswith("/data.trace/"):
        return path
    return s


_HLD_MIME_BY_EXT = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".bmp": "image/bmp",
    ".webp": "image/webp",
}


def _guess_hld_image_mime(path_or_name: str, fallback: str = "image/png") -> str:
    ext = os.path.splitext(str(path_or_name or ""))[1].lower()
    return _HLD_MIME_BY_EXT.get(ext, fallback)


def _bytes_to_hld_data_url(bys: bytes, mime: str = "image/png") -> str:
    return f"data:{mime};base64,{base64.b64encode(bys).decode('ascii')}"


def _resolve_hld_local_img_path(url: str) -> Optional[str]:
    raw = str(url or "").strip()
    if not raw or raw.startswith("data:"):
        return None
    clean = raw.split("?", 1)[0].strip()
    if clean.startswith("/"):
        clean = clean.lstrip("/")
    candidates = [clean]
    if clean.startswith("data.trace/"):
        candidates.append(clean)
    else:
        candidates.append(os.path.join("data.trace", clean))
    for candidate in candidates:
        if candidate and os.path.isfile(candidate):
            return candidate
    return None


def _coerce_hld_img_url_to_data(url: Optional[str]) -> Optional[str]:
    """HLD 图片统一存库：img_url 使用 data URL，兼容历史本地路径。"""
    if not url:
        return url
    s = str(url).strip()
    if s.startswith("data:image/"):
        return _normalize_hld_img_url(s)
    local_path = _resolve_hld_local_img_path(s)
    if not local_path:
        return _normalize_hld_img_url(s)
    try:
        with open(local_path, "rb") as fs:
            bys = fs.read()
        if not bys:
            return _normalize_hld_img_url(s)
        return _bytes_to_hld_data_url(bys, _guess_hld_image_mime(local_path))
    except Exception:
        logger.exception("coerce hld img url failed: %s", s)
        return _normalize_hld_img_url(s)


def _hld_table_text(table) -> str:
    data = table if isinstance(table, dict) else (table.dict() if hasattr(table, "dict") else {})
    if not data:
        return ""
    headers = data.get("headers") or []
    header_txt = " ".join(str(h.get("name") or h.get("code") or "") for h in headers if isinstance(h, dict))
    row_txt = " ".join(
        " ".join(str(v or "") for v in (row or {}).values())
        for row in (data.get("rows") or [])
        if isinstance(row, dict)
    )
    return f"{header_txt} {row_txt}"


def _is_hld_change_log_table(table) -> bool:
    txt = re.sub(r"\s+", "", _hld_table_text(table))
    keys = ["修改日期", "版本号", "修订说明", "修订人", "批准人"]
    return sum(1 for key in keys if key in txt) >= 3


def _extract_hld_revision_change_desc(nodes, doc_version: str) -> str:
    version = str(doc_version or "").strip()
    fallback = ""
    for node in nodes or []:
        table = getattr(node, "table", None)
        if not table or not _is_hld_change_log_table(table):
            continue
        data = table if isinstance(table, dict) else (table.dict() if hasattr(table, "dict") else {})
        for row in data.get("rows") or []:
            if not isinstance(row, dict):
                continue
            desc = str(row.get("change_desc") or "").strip()
            if not desc:
                continue
            row_version = str(row.get("version_no") or "").strip()
            if version and row_version == version:
                return desc
            if not fallback:
                fallback = desc
    return fallback


class Server(object):

    @staticmethod
    def _serialize_table(table):
        if not table:
            return None
        if isinstance(table, Table):
            return table.dict()
        if isinstance(table, dict):
            return table
        if isinstance(table, str):
            try:
                return json.loads(table)
            except Exception:
                return None
        if hasattr(table, "dict"):
            return table.dict()
        return table

    @staticmethod
    def _parse_table(table):
        if not table:
            return None
        if isinstance(table, Table):
            return table
        if isinstance(table, str):
            return Table.parse_raw(table)
        if isinstance(table, dict):
            return Table.parse_obj(table)
        return None

    @staticmethod
    def _reset_tree_node_ids(nodes: List[HldNodeForm]):
        for node in nodes or []:
            node.n_id = 0
            Server._reset_tree_node_ids(getattr(node, "children", None) or [])

    def _update_nodes(self, doc: HldDoc, p_id, nodes: List[HldNodeForm]):
        for idx, node in enumerate(nodes or []):
            sql = select(HldNode).where(HldNode.doc_id == doc.id, HldNode.n_id == node.n_id) if node.n_id else None
            row = db.session.execute(sql).scalars().first() if sql is not None else None
            if not row:
                doc.n_id += 1
                table = self._serialize_table(node.table)
                row = HldNode(
                    doc_id=doc.id,
                    n_id=doc.n_id,
                    p_id=p_id,
                    priority=idx,
                    title=node.title,
                    label=node.label,
                    img_url=_normalize_hld_img_url(node.img_url),
                    text=node.text,
                    ref_type=node.ref_type,
                    table=table,
                )
                db.session.add(row)
                node.n_id = doc.n_id
            else:
                row.p_id = p_id
                for key, value in node.dict().items():
                    if key in ("doc_id", "n_id", "p_id", "children") or value is None:
                        continue
                    if key == "table":
                        value = self._serialize_table(value)
                    if key == "img_url":
                        value = _coerce_hld_img_url_to_data(value)
                    setattr(row, key, value)
                row.priority = idx
                node.n_id = row.n_id
            if node.children:
                self._update_nodes(doc, row.n_id, node.children)

    def _load_tree(self, doc_id: int) -> List[HldNodeForm]:
        nodes: List[HldNode] = db.session.execute(
            select(HldNode).where(HldNode.doc_id == doc_id).order_by(HldNode.priority)
        ).scalars().all()
        objs_dict = {}
        roots: List[HldNodeForm] = []
        for node in nodes:
            table = self._parse_table(node.table)
            obj = HldNodeForm(
                children=[],
                doc_id=node.doc_id,
                n_id=node.n_id,
                p_id=node.p_id,
                title=node.title,
                label=node.label,
                img_url=node.img_url,
                text=node.text,
                ref_type=node.ref_type,
                table=table,
            )
            objs_dict[obj.n_id] = obj
        for obj in objs_dict.values():
            if obj.p_id == 0:
                roots.append(obj)
            else:
                parent = objs_dict.get(obj.p_id)
                if parent:
                    parent.children.append(obj)
        return roots

    def __coerce_tree_img_urls_to_data(self, nodes: List[HldNodeForm]):
        for node in nodes or []:
            if getattr(node, "img_url", None):
                node.img_url = _coerce_hld_img_url_to_data(node.img_url)
            self.__coerce_tree_img_urls_to_data(getattr(node, "children", None) or [])

    def __query_imgs(self, product_id: int):
        subquery = select(DocFile.category, func.max(DocFile.id).label("max_id"))
        subquery = subquery.where(DocFile.product_id == product_id).group_by(DocFile.category).subquery()
        sql = select(DocFile).join(subquery, DocFile.id == subquery.c.max_id)
        rows: List[DocFile] = db.session.execute(sql).scalars().all()
        return {row.category: row.file_url for row in rows}

    def __hydrate_tree_product_images(self, nodes: List[HldNodeForm], prod_imgs: dict):
        if not nodes or not prod_imgs:
            return
        for node in nodes or []:
            ref_type = getattr(node, "ref_type", None)
            if ref_type in prod_imgs and prod_imgs.get(ref_type):
                node.img_url = prod_imgs[ref_type]
            self.__hydrate_tree_product_images(getattr(node, "children", None) or [], prod_imgs)

    def __upsert_product_doc_image(self, doc_id: int, ref_type: str, file_name: str, file_size: int, file_url: str):
        if ref_type not in {"img_flow", "img_topo", "img_struct"} or not file_url:
            return
        hld_doc = db.session.execute(select(HldDoc).where(HldDoc.id == doc_id)).scalars().first()
        if not hld_doc or not hld_doc.product_id:
            return
        row = db.session.execute(
            select(DocFile)
            .where(DocFile.product_id == hld_doc.product_id)
            .where(DocFile.category == ref_type)
            .order_by(desc(DocFile.id))
        ).scalars().first()
        if row is None:
            row = DocFile(product_id=hld_doc.product_id, category=ref_type)
            db.session.add(row)
        ext = os.path.splitext(str(file_name or ""))[1] or os.path.splitext(str(file_url or ""))[1] or ".png"
        row.file_name = f"{hld_doc.version or ref_type}_{ref_type}{ext}"
        row.file_size = file_size
        row.file_url = file_url
        row.update_time = datetime.now()
        db.session.flush()

    async def add_hld_doc(self, form: HldDocForm):
        try:
            count = db.session.execute(
                select(func.count(HldDoc.id)).where(HldDoc.product_id == form.product_id, HldDoc.version == form.version)
            ).scalar()
            if count > 0:
                version = (form.version or "").strip()
                return Resp.resp_err(msg=f"该产品下已经有{version}版本文档存在" if version else ts("msg_obj_exist"))
            row = HldDoc(
                product_id=form.product_id,
                version=form.version,
                change_log=form.change_log,
                n_id=0,
                file_no=serv_review_util.resolve_doc_file_no(form.product_id, form.file_no, form.version, "hld"),
            )
            db.session.add(row)
            db.session.flush()
            if form.content:
                self._reset_tree_node_ids(form.content)
                self._update_nodes(row, 0, form.content)
            db.session.commit()
            return Resp.resp_ok(data=HldDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_hld_doc(self, id: int, product_id: int = None):
        fromdoc: HldDocObj = (await self.get_hld_doc(id, with_tree=True)).data
        if not fromdoc:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        target_pid = product_id or fromdoc.product_id
        all_versions = db.session.execute(select(HldDoc.version).where(HldDoc.product_id == target_pid)).scalars().all()
        existing_set = {v for v in all_versions if v}
        if target_pid == fromdoc.product_id:
            version = new_version(fromdoc.version)
        else:
            def _version_seq(v):
                m = re.search(r"(\d+)(?!.*\d)", v or "")
                return int(m.group(1)) if m else -1
            valid = [v for v in all_versions if v]
            version = new_version(max(valid, key=_version_seq)) if valid else fromdoc.version
        while version in existing_set:
            version = new_version(version)
        newdoc = HldDoc(
            product_id=target_pid,
            version=version,
            file_no=serv_review_util.resolve_doc_file_no(target_pid, fromdoc.file_no, version, "hld"),
            change_log=fromdoc.change_log,
            n_id=0,
        )
        count = db.session.execute(
            select(func.count(HldDoc.id)).where(HldDoc.product_id == newdoc.product_id, HldDoc.version == newdoc.version)
        ).scalar()
        if count > 0:
            return Resp.resp_err(msg=ts("msg_obj_exist"))
        try:
            db.session.add(newdoc)
            db.session.flush()
            self._update_nodes(newdoc, 0, fromdoc.content or [])
            db.session.commit()
            return Resp.resp_ok(data=HldDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_hld_doc(self, id: int):
        try:
            db.session.execute(delete(HldNode).where(HldNode.doc_id == id))
            db.session.execute(delete(HldDoc).where(HldDoc.id == id))
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def add_hld_node(self, node: HldNodeForm):
        result = db.session.execute(
            select(HldNode, HldDoc)
            .join(HldDoc, HldNode.doc_id == HldDoc.id)
            .where(HldNode.doc_id == node.doc_id, HldNode.n_id == node.p_id)
        ).first()
        if not result:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        _, doc = result
        doc.n_id += 1
        table = self._serialize_table(node.table)
        row = HldNode(
            doc_id=doc.id,
            n_id=doc.n_id,
            p_id=node.p_id,
            priority=doc.n_id,
            title=node.title,
            label=node.label,
            img_url=_normalize_hld_img_url(node.img_url),
            text=node.text,
            ref_type=node.ref_type,
            table=table,
        )
        db.session.add(row)
        db.session.commit()
        return Resp.resp_ok(data=HldNodeForm(
            doc_id=row.doc_id,
            n_id=row.n_id,
            p_id=row.p_id,
            priority=row.priority,
            title=row.title,
            label=row.label,
            img_url=row.img_url,
            text=row.text,
            ref_type=row.ref_type,
            table=node.table,
        ))

    async def delete_hld_node(self, doc_id: int, n_id: int):
        db.session.execute(delete(HldNode).where(HldNode.doc_id == doc_id, HldNode.n_id == n_id))
        db.session.commit()
        return Resp.resp_ok()

    async def update_hld_doc(self, form: HldDocForm):
        try:
            count = db.session.execute(
                select(func.count(HldDoc.id)).where(
                    HldDoc.product_id == form.product_id,
                    HldDoc.version == form.version,
                    HldDoc.id != form.id,
                )
            ).scalar()
            if count > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row: HldDoc = db.session.execute(select(HldDoc).where(HldDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            if form.content is None:
                return Resp.resp_err(msg="保存失败：未收到文档结构内容，请刷新后重试")
            if isinstance(form.content, list) and len(form.content) == 0:
                return Resp.resp_err(msg="保存失败：文档结构为空，请刷新后重试")
            for key, value in form.dict().items():
                if key in ("id", "n_id", "content") or value is None:
                    continue
                setattr(row, key, value)
            row.file_no = sync_file_no_version(row.file_no, row.version)
            row.n_id = 0
            db.session.execute(delete(HldNode).where(HldNode.doc_id == row.id))
            self._reset_tree_node_ids(form.content or [])
            self._update_nodes(row, 0, form.content or [])
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_hld_doc_file_no(self, id: int, file_no: str):
        try:
            row: HldDoc = db.session.execute(select(HldDoc).where(HldDoc.id == id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            row.file_no = (file_no or "").strip() or None
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def get_hld_doc(self, id: int, with_tree: bool = False):
        row, row_prod = db.session.execute(
            select(HldDoc, Product).outerjoin(Product, HldDoc.product_id == Product.id).where(HldDoc.id == id)
        ).first() or (None, None)
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        tree = []
        if with_tree:
            tree = self._load_tree(row.id)
            if row.product_id:
                prod_imgs = self.__query_imgs(row.product_id)
                self.__hydrate_tree_product_images(tree, prod_imgs)
        doc_data = row.dict()
        if not (doc_data.get("file_no") or "").strip():
            doc_data["file_no"] = serv_review_util.resolve_doc_file_no(row.product_id, row.file_no, row.version, "hld") or doc_data.get("file_no")
        return Resp.resp_ok(data=HldDocObj(
            **doc_data,
            product_name=row_prod.name if row_prod else "",
            product_version=row_prod.full_version if row_prod else "",
            content=tree,
        ))

    async def list_hld_doc(
        self,
        op_user: UserObj,
        product_id: int = 0,
        version: str = None,
        page_index: int = 0,
        page_size: int = 10,
    ):
        page_index = page_index if page_index >= 0 else 0
        page_size = page_size if page_size > 0 else 10
        sql = select(HldDoc, Product).outerjoin(Product, HldDoc.product_id == Product.id)
        if product_id:
            sql = sql.where(HldDoc.product_id == product_id)
        if version:
            sql = sql.where(HldDoc.version.like(f"%{version}%"))
        if not product_id and op_user.id != 1:
            subquery = select(UserProd.product_id).where(UserProd.user_id == op_user.id).scalar_subquery()
            sql = sql.where(Product.id.in_(subquery))
        total = db.session.execute(select(func.count()).select_from(sql)).scalars().first()
        rows = db.session.execute(
            sql.offset(page_size * page_index).limit(page_size).order_by(desc(HldDoc.create_time))
        ).all()
        doc_ids = [row.id for row, _ in rows]
        nodes_by_doc = defaultdict(list)
        if doc_ids:
            node_rows = db.session.execute(select(HldNode).where(HldNode.doc_id.in_(doc_ids))).scalars().all()
            for node in node_rows:
                nodes_by_doc[node.doc_id].append(node)
        objs = []
        for row, row_prd in rows:
            obj = HldDocObj(**row.dict())
            if row_prd:
                obj.product_name = row_prd.name
                obj.product_version = row_prd.full_version
            if not (obj.file_no or "").strip():
                obj.file_no = serv_review_util.resolve_doc_file_no(row.product_id, row.file_no, row.version, "hld")
            rev_desc = _extract_hld_revision_change_desc(nodes_by_doc.get(row.id, []), row.version)
            if rev_desc:
                obj.change_log = rev_desc
            objs.append(obj)
        return Resp.resp_ok(data=Page(total=total, page_size=page_size, rows=objs, page_index=page_index))

    async def add_doc_file(self, doc_id: int, file, ref_type: str = None):
        if not file:
            return Resp.resp_err(msg="未收到图片文件")
        bys = await file.read()
        if not bys:
            return Resp.resp_err(msg="图片文件为空")
        filename = getattr(file, "filename", "") or "image.png"
        mime = getattr(file, "content_type", None) or _guess_hld_image_mime(filename)
        if not str(mime).startswith("image/"):
            mime = _guess_hld_image_mime(filename)
        data_url = _bytes_to_hld_data_url(bys, mime)
        self.__upsert_product_doc_image(
            doc_id,
            str(ref_type or "").strip(),
            filename,
            len(bys),
            data_url,
        )
        db.session.commit()
        return Resp.resp_ok(data=data_url)

    @staticmethod
    def _biz_title(value: str) -> str:
        txt = re.sub(r"^\s*\d+(?:\.\d+)*[\s、.．]*", "", str(value or "").strip())
        return re.sub(r"\s+", "", txt)

    @staticmethod
    def _heading_level(title: str) -> int:
        matched = re.match(r"^(\d+(?:\.\d+)*)", str(title or "").strip())
        if matched:
            return max(1, min(len(matched.group(1).split(".")), 5))
        return 1

    @staticmethod
    def _is_imported_image_title(value: str) -> bool:
        return re.match(r"^导入图片\d*$", str(value or "").strip()) is not None

    @staticmethod
    def _is_imported_table_title(value: str) -> bool:
        return re.match(r"^导入表格\d*$", str(value or "").strip()) is not None

    @staticmethod
    def _is_image_caption_line(value: str) -> bool:
        return re.match(r"^\s*图\s*\d+\s*", str(value or "").strip()) is not None

    @staticmethod
    def _is_table_caption_line(value: str) -> bool:
        return re.match(r"^\s*表\s*\d+\s*", str(value or "").strip()) is not None

    def _to_hld_node(self, node) -> HldNodeForm:
        table = getattr(node, "table", None)
        if table is not None and not isinstance(table, Table):
            try:
                table = Table.parse_obj(table.dict() if hasattr(table, "dict") else table)
            except Exception:
                table = None
        return HldNodeForm(
            title=getattr(node, "title", None),
            label=getattr(node, "label", None),
            img_url=getattr(node, "img_url", None),
            text=getattr(node, "text", None),
            ref_type=getattr(node, "ref_type", None),
            table=table,
            children=[self._to_hld_node(child) for child in (getattr(node, "children", None) or [])],
        )

    def __bind_imported_image_titles(self, nodes: List[HldNodeForm]):
        def walk(node_list: List[HldNodeForm]):
            for node in node_list or []:
                children = list(getattr(node, "children", None) or [])
                image_children = [child for child in children if str(getattr(child, "img_url", "") or "").strip()]
                if image_children:
                    lines = str(getattr(node, "text", "") or "").replace("\r", "").split("\n")
                    caption_entries = [
                        (idx, (line or "").strip())
                        for idx, line in enumerate(lines)
                        if self._is_image_caption_line(line)
                    ]
                    if caption_entries:
                        used_line_idx = set()
                        for idx, child in enumerate(image_children):
                            if idx >= len(caption_entries):
                                break
                            line_idx, caption = caption_entries[idx]
                            child_title = str(getattr(child, "title", "") or "").strip()
                            if not child_title or self._is_imported_image_title(child_title):
                                child.title = caption
                            used_line_idx.add(line_idx)
                        if used_line_idx:
                            node.text = "\n".join(
                                (line or "").strip()
                                for idx, line in enumerate(lines)
                                if idx not in used_line_idx and str(line or "").strip()
                            )
                walk(children)

        walk(nodes or [])

    def __bind_imported_table_titles(self, nodes: List[HldNodeForm]):
        def walk(node_list: List[HldNodeForm]):
            for node in node_list or []:
                children = list(getattr(node, "children", None) or [])
                table_children = [
                    child for child in children
                    if getattr(child, "table", None) is not None and getattr(getattr(child, "table", None), "headers", None)
                ]
                if table_children:
                    lines = str(getattr(node, "text", "") or "").replace("\r", "").split("\n")
                    caption_entries = [
                        (idx, (line or "").strip())
                        for idx, line in enumerate(lines)
                        if self._is_table_caption_line(line)
                    ]
                    if caption_entries:
                        used_line_idx = set()
                        for idx, child in enumerate(table_children):
                            if idx >= len(caption_entries):
                                break
                            line_idx, caption = caption_entries[idx]
                            child_title = str(getattr(child, "title", "") or "").strip()
                            if not child_title or self._is_imported_table_title(child_title):
                                child.title = caption
                            used_line_idx.add(line_idx)
                        if used_line_idx:
                            node.text = "\n".join(
                                (line or "").strip()
                                for idx, line in enumerate(lines)
                                if idx not in used_line_idx and str(line or "").strip()
                            )
                walk(children)

        walk(nodes or [])

    def __persist_data_url_images(self, nodes: List[HldNodeForm]):
        """图片以 data URL 存入 hld_node.img_url / doc_file.file_url，不再写入本地目录。"""

        def walk(node_list: List[HldNodeForm]):
            for node in node_list or []:
                img_url = (getattr(node, "img_url", None) or "").strip()
                if img_url:
                    node.img_url = _coerce_hld_img_url_to_data(img_url)
                walk(getattr(node, "children", None) or [])

        walk(nodes or [])

    def __apply_hld_ref_types(self, nodes: List[HldNodeForm]):
        def walk(node_list: List[HldNodeForm]):
            for node in node_list or []:
                title = str(getattr(node, "title", "") or "")
                norm_title = self._biz_title(title)
                if "物理拓扑" in norm_title or re.search(r"^图\s*1", title, re.I):
                    node.ref_type = "img_topo"
                elif re.search(r"^图\s*2", title, re.I) or ("系统结构" in norm_title and "图" in title):
                    node.ref_type = "img_struct"
                elif getattr(node, "ref_type", None) in {"img_topo", "img_struct"}:
                    pass
                walk(getattr(node, "children", None) or [])

        walk(nodes or [])

    def __sync_imported_hld_doc_images_to_doc_file(self, doc_id: int, nodes: List[HldNodeForm]):
        if not doc_id or not nodes:
            return

        def file_size_of(url: str) -> int:
            path = str(url or "").strip()
            if path.startswith("data:image/"):
                matched = re.match(r"^data:[^;]+;base64,(.+)$", path, re.S)
                if matched:
                    try:
                        return len(base64.b64decode(matched.group(1) or ""))
                    except Exception:
                        return 0
                return 0
            if path.startswith("/"):
                path = path[1:]
            try:
                return os.path.getsize(path) if path and os.path.exists(path) else 0
            except Exception:
                return 0

        for ref_type in ("img_topo", "img_struct"):
            best_url = ""
            best_score = -1

            def walk(items: List[HldNodeForm]):
                nonlocal best_url, best_score
                for node in items or []:
                    node_ref = str(getattr(node, "ref_type", "") or "").strip()
                    img_url = _normalize_hld_img_url(getattr(node, "img_url", "") or "")
                    if node_ref == ref_type and img_url:
                        title = f"{getattr(node, 'title', '') or ''} {getattr(node, 'text', '') or ''}"
                        score = 10
                        if ref_type == "img_topo" and "物理拓扑" in self._biz_title(title):
                            score = 100
                        if ref_type == "img_struct" and ("系统结构" in self._biz_title(title) or re.search(r"图\s*2", title, re.I)):
                            score = 100
                        if score > best_score:
                            best_url = img_url
                            best_score = score
                    walk(getattr(node, "children", None) or [])

            walk(nodes)
            if best_url:
                if best_url.startswith("/"):
                    best_url = best_url[1:]
                self.__upsert_product_doc_image(doc_id, ref_type, ref_type, file_size_of(best_url), best_url)

    async def import_hld_doc_word(self, product_id: int, version: str, change_log: str, file):
        if Document is None or Paragraph is None:
            return Resp.resp_err(msg="当前环境缺少 python-docx 依赖，暂不可用 Word 导入。")
        try:
            count = db.session.execute(
                select(func.count(HldDoc.id)).where(HldDoc.product_id == product_id, HldDoc.version == version)
            ).scalar()
            if count > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))

            bys = await file.read()
            docx = Document(io.BytesIO(bys))
            content, _heading_rows = srsdoc_serv._Server__parse_docx_content(docx)
            file_name = file.filename or ""
            _, file_no = srsdoc_serv._Server__extract_file_info(file_name)
            hld_content = [self._to_hld_node(node) for node in (content or [])]
            self.__bind_imported_image_titles(hld_content)
            self.__bind_imported_table_titles(hld_content)
            self.__apply_hld_ref_types(hld_content)
            self.__persist_data_url_images(hld_content)
            form = HldDocForm(
                product_id=product_id,
                version=version,
                file_no=file_no or None,
                change_log=change_log,
                content=hld_content,
            )
            resp = await self.add_hld_doc(form)
            if resp.code == c_ok and resp.data and resp.data.id:
                product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
                product_version = getattr(product, "full_version", "") or getattr(product, "name", "") or ""
                srsdoc_serv._Server__auto_sync_product_doc_images(
                    product_id,
                    content or [],
                    version,
                    docx,
                    product_version,
                    bys,
                    resp.data.id,
                )
                self.__sync_imported_hld_doc_images_to_doc_file(resp.data.id, hld_content)
                db.session.commit()
            return resp
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def export_hld_doc(self, output, id: int = 0, snapshot: HldDocForm = None):
        if Document is None:
            return
        from .serv_utils import docx_util

        if snapshot and snapshot.content:
            doc_obj = snapshot
            file_no = (snapshot.file_no or "").strip()
        else:
            resp = await self.get_hld_doc(id, with_tree=True)
            if resp.code != c_ok or not resp.data:
                return
            doc_obj = resp.data
            file_no = (doc_obj.file_no or "").strip()

        docx = Document()

        def __is_cover_title(title: str) -> bool:
            return self._biz_title(title) in ["软件概要设计", "软件概要设计说明书"]

        def __is_revision_title(title: str) -> bool:
            return self._biz_title(title) == "文件修订记录"

        def __write_nodes(nodes: List[HldNodeForm], default_level: int = 1):
            for node in nodes or []:
                title = str(getattr(node, "title", "") or "").strip()
                if not title and not getattr(node, "text", None) and not getattr(node, "table", None):
                    if getattr(node, "children", None):
                        __write_nodes(node.children, default_level)
                    continue
                if title and not self._is_imported_image_title(title):
                    level = self._heading_level(title) if re.match(r"^\d", title) else default_level
                    if __is_cover_title(title) or __is_revision_title(title):
                        level = 1
                    docx_util.save_title2docx(title, docx, level=level)
                text = str(getattr(node, "text", "") or "").strip()
                if text:
                    docx_util.save_txt2docx(text, docx)
                table = getattr(node, "table", None)
                if table and getattr(table, "headers", None):
                    docx_util.save_tab2docx(table, docx)
                img_url = _normalize_hld_img_url(getattr(node, "img_url", "") or "")
                if img_url:
                    if self._is_imported_image_title(title):
                        docx_util.save_txt2docx(title, docx)
                    docx_util.save_img2docx(img_url, docx, mw=420, mh=240)
                __write_nodes(getattr(node, "children", None) or [], default_level + 1)

        __write_nodes(doc_obj.content or [])
        if file_no:
            section = docx.sections[0] if docx.sections else None
            if section is not None and section.header.paragraphs:
                docx_util.fonted_txt(section.header.paragraphs[0], file_no)
        docx_util.add_page_number_footer(docx.sections[0], file_no=file_no, skip_first=True)
        docx.save(output)

    @staticmethod
    def __table_header_name(table, idx: int = 0) -> str:
        tbl = Server._parse_table(table)
        if not tbl or not tbl.headers or len(tbl.headers) <= idx:
            return ""
        return str(tbl.headers[idx].name or "").strip()

    @staticmethod
    def __detect_field_library(parent_titles: List[str]) -> str:
        chain = " ".join(parent_titles or [])
        if "库2" in chain or "库 2" in chain:
            return "lib2"
        return "lib1"

    async def sync_hld_from_sds(self, op_user: UserObj, product_id: int, version: str):
        if not product_id:
            return Resp.resp_err(msg="请选择产品")
        version = str(version or "").strip()
        if not version:
            return Resp.resp_err(msg="请填写版本")
        sds_row: SdsDoc = db.session.execute(
            select(SdsDoc)
            .where(SdsDoc.product_id == product_id, SdsDoc.version == version)
            .order_by(desc(SdsDoc.create_time))
        ).scalars().first()
        if not sds_row:
            return Resp.resp_err(msg="未找到同版本详细设计文档")
        from .serv_sds_doc import Server as SdsServer

        sds_resp = await SdsServer().get_sds_doc(sds_row.id, with_tree=True)
        if sds_resp.code != c_ok or not sds_resp.data:
            return Resp.resp_err(msg="读取详细设计失败")
        interface_table = None
        field_tables = []
        interface_details = []

        def walk(nodes, parent_titles: List[str] = None):
            nonlocal interface_table
            parent_titles = parent_titles or []
            for node in nodes or []:
                title = str(getattr(node, "title", "") or "").strip()
                chain = parent_titles + ([title] if title else [])
                tbl = self._parse_table(getattr(node, "table", None))
                if tbl and tbl.headers:
                    h0 = self.__table_header_name(tbl, 0)
                    if h0 == "接口设计编号" and interface_table is None:
                        interface_table = self._serialize_table(tbl)
                    elif h0 in ("字段ID", "Field ID"):
                        field_tables.append({
                            "title": title,
                            "library": self.__detect_field_library(parent_titles),
                            "table": self._serialize_table(tbl),
                        })
                detail_name = self.__parse_interface_detail_title(title)
                detail_text = str(getattr(node, "text", "") or "").strip()
                if detail_name and detail_text:
                    interface_details.append({
                        "name": detail_name,
                        "source_title": title,
                        "text": detail_text,
                    })
                walk(getattr(node, "children", None) or [], chain)

        walk(sds_resp.data.content or [])
        return Resp.resp_ok(data={
            "sds_doc_id": sds_row.id,
            "interface_table": interface_table,
            "interface_details": interface_details,
            "field_tables": field_tables,
        })

    @staticmethod
    def __parse_interface_detail_title(title: str) -> Optional[str]:
        title = str(title or "").strip()
        if not title:
            return None
        m = re.match(r"^\d+(?:\.\d+)+\s+(.+接口)\s*$", title)
        if not m:
            return None
        return m.group(1).strip()
