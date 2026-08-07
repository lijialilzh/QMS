#!/usr/bin/env python
# encoding: utf-8

# 软件配置管理计划服务层（开发文件，PDP 风格章节树）。默认内容取自 src-res/scm_default_content.json。

import base64
import copy
import json
import logging
import os
import re
from io import BytesIO
from typing import List

from sqlalchemy import delete, func, select
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.scm_doc import ScmDoc
from ..model.prod_dhf import ProdDhf
from ..model.project_member import ProjectMember
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_scm_doc import ScmDocForm
from ..obj.vobj_scm_doc import ScmDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from . import serv_review_util
from .serv_utils import new_version, sync_file_no_version
from .serv_utils import docx_util

logger = logging.getLogger(__name__)

_DEFAULT_FILE = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "src-res", "scm_default_content.json")
try:
    with open(_DEFAULT_FILE, encoding="utf-8") as _f:
        DEFAULT_SCM_CONTENT = json.load(_f)
except Exception:
    DEFAULT_SCM_CONTENT = {"sections": []}

# 模板基准产品名（用于全文替换为当前产品名）。
BASE_NAME = "InferOperate Suite"
DOC_NAME = "软件配置管理计划"
DOC_KEY = "scm"


class Server(object):

    def __normalize_node(self, node):
        if not isinstance(node, dict):
            return {"title": str(node or ""), "body": "", "tables": [], "children": []}
        result = dict(node)
        result["title"] = str(result.get("title") or "")
        result["body"] = str(result.get("body") or "")
        tables = result.get("tables")
        if not isinstance(tables, list):
            tables = []
        norm_tables = []
        for table in tables:
            if isinstance(table, list):
                norm_tables.append([[str(c) if c is not None else "" for c in (row or [])] for row in table if isinstance(row, list)])
        result["tables"] = norm_tables
        # 兼容旧数据：如果有 blocks（之前改造残留），把 blocks 中的 table/text 提取回 tables/body
        blocks = result.get("blocks")
        if isinstance(blocks, list) and blocks and not result["body"] and not norm_tables:
            parts = []
            tbls = []
            for blk in blocks:
                if not isinstance(blk, dict):
                    continue
                btype = blk.get("type")
                if btype == "text" and blk.get("text"):
                    parts.append(str(blk["text"]))
                elif btype == "table" and blk.get("table"):
                    tbl = blk["table"]
                    tbls.append([[str(c) if c is not None else "" for c in (row or [])] for row in tbl if isinstance(row, list)])
                    parts.append("见下表")
            if parts:
                result["body"] = "\n".join(parts)
            if tbls:
                result["tables"] = tbls
            result.pop("blocks", None)
        elif "blocks" in result:
            result.pop("blocks", None)
        children = result.get("children")
        result["children"] = [self.__normalize_node(c) for c in children] if isinstance(children, list) else []
        return result

    def __normalize_content(self, content):
        if not isinstance(content, dict) or not isinstance(content.get("sections"), list):
            return copy.deepcopy(DEFAULT_SCM_CONTENT)
        return {"sections": [self.__normalize_node(s) for s in content["sections"]]}

    def __replace_name(self, node, base, name, skip_titles=None):
        if not name or base == name:
            return
        if skip_titles and str(node.get("title") or "").strip() in skip_titles:
            return
        if node.get("body"):
            node["body"] = node["body"].replace(base, name)
        for tbl in (node.get("tables") or []):
            for row in tbl:
                for i in range(len(row)):
                    if isinstance(row[i], str) and base in row[i]:
                        row[i] = row[i].replace(base, name)
        for c in (node.get("children") or []):
            self.__replace_name(c, base, name, skip_titles=skip_titles)

    def __dhf_file_no(self, prod_id):
        if not prod_id:
            return ""
        row = db.session.execute(
            select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name == DOC_NAME).order_by(ProdDhf.id.asc())
        ).scalars().first()
        if not row:
            row = db.session.execute(
                select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name.like(f"%{DOC_NAME}%")).order_by(ProdDhf.id.asc())
            ).scalars().first()
        return (row.code or "").strip() if row and row.code else ""

    def __fill_revision(self, content, prod_id, version, force=False):
        """文件修订记录首行：修改日期(评审/封面日期)、版本号、首次发布、修订人(TPM)、批准人(研发负责人)。
        force=True（切换产品）时强制覆盖，无数据则置空。"""
        rev_date = serv_review_util.cover_date(prod_id, DOC_KEY) if prod_id else ""
        reviser = approver = ""
        if prod_id:
            members = db.session.execute(select(ProjectMember).where(ProjectMember.prod_id == prod_id)).scalars().all()
            reviser = next((m.name for m in members if "TPM" in str(m.role or "")), "")
            approver = next((m.name for m in members if "研发负责人" in str(m.role or "")), "")
        for s in (content.get("sections") or []):
            if s.get("ref_type") != "revision":
                continue
            for tbl in (s.get("tables") or []):
                if isinstance(tbl, list) and len(tbl) >= 2 and isinstance(tbl[1], list) and len(tbl[1]) >= 3:
                    row = tbl[1]
                    if force:
                        row[0] = rev_date or ""
                        if len(row) >= 2:
                            row[1] = str(version or "")
                        if len(row) >= 3:
                            row[2] = "首次发布"
                        if len(row) >= 4:
                            row[3] = reviser
                        if len(row) >= 5:
                            row[4] = approver
                    else:
                        if not str(row[0] or "").strip():
                            row[0] = rev_date
                        if version and len(row) >= 2 and not str(row[1] or "").strip():
                            row[1] = str(version)
                        if len(row) >= 3 and not str(row[2] or "").strip():
                            row[2] = "首次发布"
                        if len(row) >= 4 and reviser and not str(row[3] or "").strip():
                            row[3] = reviser
                        if len(row) >= 5 and approver and not str(row[4] or "").strip():
                            row[4] = approver
            break
        return content

    def __autofill(self, content, prod_id, product=None, version="", force=False):
        if not isinstance(content, dict):
            return content
        # 这些章节内容为固定模板（编号规则/构建配置/发布过程/测试环境），不替换产品名
        skip_titles = {"产品开发部软件构建配置项版本控制", "文件中涉及的编号命名规则", "发布过程", "软件测试环境的建立"}
        if product and product.name:
            for s in (content.get("sections") or []):
                self.__replace_name(s, BASE_NAME, product.name, skip_titles=skip_titles)
            # 更新完整版本号
            new_full_version = (product.full_version or "").strip()
            if new_full_version:
                def _update_version(node):
                    if node.get("body") and "完整版本" in str(node.get("body")):
                        node["body"] = re.sub(r"完整版本：[^\n]*", f"完整版本：{new_full_version}", str(node["body"]))
                    for c in (node.get("children") or []):
                        _update_version(c)
                for s in (content.get("sections") or []):
                    _update_version(s)
        self.__fill_revision(content, prod_id, version, force=force)
        serv_review_util.ensure_review(
            content, DOC_KEY,
            serv_review_util.review_date(prod_id, serv_review_util.REVIEW_DEFS[DOC_KEY]["name_keywords"]) if prod_id else "",
            prod_id,
        )
        serv_review_util.fill_cover_dates(content, serv_review_util.cover_date(prod_id, DOC_KEY) if prod_id else "", force=force)
        serv_review_util.fill_cover_signers(content, serv_review_util.cover_signers(prod_id, DOC_KEY) if prod_id else {}, force=force)
        return content

    def __to_obj(self, row: ScmDoc, product: Product = None):
        obj = ScmDocObj(**row.dict())
        obj.content = self.__autofill(self.__normalize_content(obj.content), row.product_id, product, row.version)
        if not (obj.file_no or "").strip():
            obj.file_no = self.__dhf_file_no(row.product_id)
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
        return obj

    async def add_scm_doc(self, form: ScmDocForm):
        try:
            sql = select(func.count(ScmDoc.id)).where(ScmDoc.product_id == form.product_id, ScmDoc.version == form.version)
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = ScmDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=ScmDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_scm_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: ScmDoc = db.session.execute(select(ScmDoc).where(ScmDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(ScmDoc.version).where(ScmDoc.product_id == target_pid)).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = ScmDoc(
                product_id=target_pid, version=version,
                file_no=sync_file_no_version((fromdoc.file_no or "").strip() or self.__dhf_file_no(target_pid), version) or None,
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            if target_pid != fromdoc.product_id:
                await self.rebind_product(newdoc.id, target_pid)
            return Resp.resp_ok(data=ScmDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_scm_doc(self, form: ScmDocForm):
        try:
            row: ScmDoc = db.session.execute(select(ScmDoc).where(ScmDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def rebind_product(self, id: int, product_id: int):
        """切换产品：更新 product_id 并强制用新产品信息重新获取封面/修订/产品信息后保存，返回新 obj。"""
        try:
            row: ScmDoc = db.session.execute(select(ScmDoc).where(ScmDoc.id == id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            old_product: Product = db.session.execute(select(Product).where(Product.id == row.product_id)).scalars().first() if row.product_id else None
            product: Product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
            if not product:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            db.session.execute(delete(ScmDoc).where(ScmDoc.product_id == product_id, ScmDoc.version == row.version, ScmDoc.id != id))
            # 重置为默认模板（和 imm 一样，避免旧产品污染），再用新产品信息填充
            content = copy.deepcopy(DEFAULT_SCM_CONTENT) if isinstance(DEFAULT_SCM_CONTENT, dict) else {"sections": []}
            # 完整版本号更新为新产品版本
            new_full_version = (product.full_version or "").strip()
            def update_version(node):
                if node.get("body") and "完整版本" in str(node.get("body")):
                    node["body"] = re.sub(r"完整版本：[^\n]*", f"完整版本：{new_full_version}", str(node["body"]))
                if node.get("body") and re.search(r"\d+\.\d+\.\d+\.\d+", str(node.get("body"))):
                    node["body"] = re.sub(r"\d+\.\d+\.\d+\.\d+", new_full_version, str(node["body"]))
                for c in (node.get("children") or []):
                    update_version(c)
            for s in (content.get("sections") or []):
                update_version(s)
            row.product_id = product_id
            content = self.__autofill(content, product_id, product, row.version, force=True)
            row.content = content
            db.session.commit()
            return Resp.resp_ok(data=self.__to_obj(row, product))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_scm_doc(self, id: int):
        db.session.execute(delete(ScmDoc).where(ScmDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_scm_doc(self, id: int):
        sql = select(ScmDoc, Product).join(Product, ScmDoc.product_id == Product.id).where(ScmDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_scm_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(ScmDoc.product_id == product_id)
        if version:
            wheres.append(ScmDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(ScmDoc.id)).join(Product, ScmDoc.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (select(ScmDoc, Product).join(Product, ScmDoc.product_id == Product.id).where(*wheres)
               .order_by(ScmDoc.id.desc()).offset(page_index * page_size).limit(page_size))
        rows: List[ScmDocObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    # ---------------- 导出 Word（PDP 风格章节树） ----------------
    async def export_scm_doc(self, output, id: int):
        resp = await self.get_scm_doc(id)
        obj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        c = self.__normalize_content(obj.content)
        sections = c.get("sections") or []
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        document.settings.element.append(update_fields)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")
        docx_util.add_page_number_footer(section, obj.file_no or "")

        def add_blank_lines(n):
            for _ in range(max(0, int(n or 0))):
                document.add_paragraph("")

        def write_center_title(text, size=22.0, bold=False):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            docx_util.fonted_txt(p, str(text or ""), font_size=size, bold=bold)

        def add_text(text):
            docx_util.save_txt2docx(str(text or ""), document)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            s = str(text or "")
            if s.startswith("data:image"):
                try:
                    b64 = s.split(",", 1)[1] if "," in s else ""
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = align
                    para.add_run().add_picture(BytesIO(base64.b64decode(b64)), height=Pt(33))
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    return
                except Exception:
                    pass
            cell.text = ""
            for i, line in enumerate(s.split("\n")):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER if align == WD_ALIGN_PARAGRAPH.CENTER else WD_CELL_VERTICAL_ALIGNMENT.TOP

        def _set_fixed_widths(table, widths_dxa):
            table.autofit = False
            table.allow_autofit = False
            tbl_pr = table._tbl.tblPr
            layout = tbl_pr.find(qn("w:tblLayout"))
            if layout is None:
                layout = OxmlElement("w:tblLayout")
                tbl_pr.append(layout)
            layout.set(qn("w:type"), "fixed")
            grid_el = table._tbl.find(qn("w:tblGrid"))
            if grid_el is not None:
                for gc in list(grid_el):
                    grid_el.remove(gc)
                for w in widths_dxa:
                    gc = OxmlElement("w:gridCol")
                    gc.set(qn("w:w"), str(w))
                    grid_el.append(gc)
            # 逐格写 tcW，确保固定列宽生效
            for row in table.rows:
                for i, w in enumerate(widths_dxa):
                    if i < len(row.cells):
                        tcpr = row.cells[i]._tc.get_or_add_tcPr()
                        tcw = tcpr.find(qn("w:tcW"))
                        if tcw is None:
                            tcw = OxmlElement("w:tcW")
                            tcpr.append(tcw)
                        tcw.set(qn("w:w"), str(w))
                        tcw.set(qn("w:type"), "dxa")

        def add_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for r_idx, row in enumerate(grid):
                cells = table.add_row().cells
                for c_idx in range(cols):
                    set_cell(cells[c_idx], row[c_idx] if c_idx < len(row) else "", bold=(r_idx == 0))
            # SCI 清单表：固定列宽，SCI名字/类型不换行、存储地点较宽
            first = str(grid[0][0]).strip() if grid and grid[0] else ""
            if first == "SCI名字":
                if cols == 5:
                    _set_fixed_widths(table, [1700, 1100, 2000, 1500, 3000])
                elif cols == 4:
                    _set_fixed_widths(table, [1700, 2200, 1500, 3600])
            elif first == "类别" and cols >= 3:
                # 配置项存储地址表：类别列收窄、存储路径较宽；首列连续相同「类别」纵向合并；文字左对齐
                if cols == 3:
                    _set_fixed_widths(table, [1300, 1300, 5000])
                rows = table.rows
                n = len(rows)
                r = 1
                while r < n:
                    val = str((grid[r][0] if r < len(grid) and grid[r] else "") or "").strip()
                    if not val:
                        r += 1
                        continue
                    r2 = r
                    while r2 + 1 < n and str((grid[r2 + 1][0] if grid[r2 + 1] else "") or "").strip() == val:
                        r2 += 1
                    if r2 > r:
                        merged = rows[r].cells[0].merge(rows[r2].cells[0])
                        set_cell(merged, val, align=WD_ALIGN_PARAGRAPH.LEFT)
                        merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    r = r2 + 1
            document.add_paragraph()

        def add_cover_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for row in grid:
                cells = table.add_row().cells
                for c_idx in range(cols):
                    text = row[c_idx] if c_idx < len(row) else ""
                    set_cell(cells[c_idx], text, bold=(c_idx % 2 == 0), align=WD_ALIGN_PARAGRAPH.CENTER)
                if (str(row[0]).strip() if row else "") == "生效日期" and cols > 2:
                    merged = cells[1]
                    for c_idx in range(2, cols):
                        merged = merged.merge(cells[c_idx])
                    set_cell(merged, row[1] if len(row) > 1 else "", align=WD_ALIGN_PARAGRAPH.CENTER)
            document.add_paragraph()

        def strip_num(title):
            return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()

        def add_body_heading(title, level):
            size = {1: 16.0, 2: 14.0, 3: 12.0}.get(level, 11.0)
            p = document.add_heading("", level=max(1, min(level, 9)))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, title, font_size=size, bold=True)

        def render_body_section(node, level, number=""):
            name = strip_num(node.get("title"))
            heading = f"{number} {name}".strip() if number else name
            add_body_heading(heading, level=max(1, min(level, 9)))
            if node.get("ref_type") == "review":
                for t_idx, table in enumerate(node.get("tables") or []):
                    serv_review_util.render_review_grid(document, table, set_cell, merge_col0=(t_idx == 0), merge_full=True)
            else:
                body_text = str(node.get("body") or "")
                tables = node.get("tables") or []
                # 按"见下表"/"表N"切分 body，交错输出正文段和表格
                if (("见下表" in body_text) or re.search(r"(?m)^表\s*\d", body_text)) and tables:
                    lines = body_text.split("\n")
                    buf = []
                    tbl_idx = 0
                    def flush_text():
                        t = "\n".join(buf).strip()
                        if t:
                            add_text(t)
                        buf.clear()
                    for ln in lines:
                        if ("见下表" in ln.strip() or re.match(r"^表\s*\d", ln.strip())) and tbl_idx < len(tables):
                            buf.append(ln)  # 保留"见下表"/"表N"行作为正文
                            flush_text()
                            add_grid(tables[tbl_idx])
                            tbl_idx += 1
                        else:
                            buf.append(ln)
                    flush_text()
                    for i in range(tbl_idx, len(tables)):
                        add_grid(tables[i])
                else:
                    if body_text.strip():
                        add_text(body_text)
                    for table in tables:
                        add_grid(table)
            idx = 0
            for child in (node.get("children") or []):
                idx += 1
                child_num = f"{number}.{idx}" if number else f"{idx}"
                render_body_section(child, level + 1, child_num)

        cover = next((s for s in sections if s.get("ref_type") == "cover"), None)
        revision = next((s for s in sections if s.get("ref_type") == "revision"), None)
        body = [s for s in sections if s.get("ref_type") not in ("cover", "revision")]

        add_blank_lines(6)
        write_center_title((strip_num(cover.get("title")) if cover else "") or DOC_NAME, size=22.0, bold=True)
        add_blank_lines(4)
        if cover:
            for table in (cover.get("tables") or []):
                add_cover_grid(table)

        document.add_page_break()
        write_center_title("文件修订记录", size=14.0, bold=True)
        add_blank_lines(2)
        if revision:
            for table in (revision.get("tables") or []):
                add_grid(table)

        document.add_page_break()
        write_center_title("目录", size=16.0, bold=True)
        docx_util.insert_toc_field(document)

        document.add_page_break()
        seq = 0
        for node in body:
            if node.get("ref_type") == "review":
                render_body_section(node, 1, "")
            else:
                seq += 1
                render_body_section(node, 1, str(seq))

        docx_util.fill_toc_cache(document)
        document.save(output)
        output.seek(0)
