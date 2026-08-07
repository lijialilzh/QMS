#!/usr/bin/env python
# encoding: utf-8

# 用户测试报告服务层（开发文件，PDP 风格章节树）。默认内容取自 src-res/utr_default_content.json。

import base64
import copy
import json
import logging
import os
import re
from io import BytesIO
from typing import List

from sqlalchemy import delete, func, select
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.utr_doc import UtrDoc
from ..model.doc_file import DocFile
from ..model.prod_dhf import ProdDhf
from ..model.project_member import ProjectMember
from ..model.test_set import TestSet
from ..model.test_case import TestCase
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_utr_doc import UtrDocForm
from ..obj.vobj_utr_doc import UtrDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from . import serv_review_util
from .serv_utils import new_version, sync_file_no_version
from .serv_utils import docx_util

logger = logging.getLogger(__name__)

_DEFAULT_FILE = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "src-res", "utr_default_content.json")
try:
    with open(_DEFAULT_FILE, encoding="utf-8") as _f:
        DEFAULT_UTR_CONTENT = json.load(_f)
except Exception:
    DEFAULT_UTR_CONTENT = {"sections": []}

# 模板基准产品名（用于全文替换为当前产品名）。
BASE_NAME = "InferOperate Suite"
DOC_NAME = "用户测试报告"
DOC_KEY = "utr"


class Server(object):

    def __normalize_node(self, node):
        if not isinstance(node, dict):
            return {"title": str(node or ""), "body": "", "tables": [], "children": []}
        result = dict(node)
        result["title"] = str(result.get("title") or "")
        result["body"] = str(result.get("body") or "")
        tables = result.get("tables")
        if not isinstance(tables, list):
            tables = []
        norm_tables = []
        for table in tables:
            if isinstance(table, list):
                norm_tables.append([[str(c) if c is not None else "" for c in (row or [])] for row in table if isinstance(row, list)])
        result["tables"] = norm_tables
        # 兼容旧数据：如果有 blocks（之前改造残留），把 blocks 中的 table/text 提取回 tables/body
        blocks = result.get("blocks")
        if isinstance(blocks, list) and blocks and not result["body"] and not norm_tables:
            parts = []
            tbls = []
            for blk in blocks:
                if not isinstance(blk, dict):
                    continue
                btype = blk.get("type")
                if btype == "text" and blk.get("text"):
                    parts.append(str(blk["text"]))
                elif btype == "table" and blk.get("table"):
                    tbl = blk["table"]
                    tbls.append([[str(c) if c is not None else "" for c in (row or [])] for row in tbl if isinstance(row, list)])
                    parts.append("见下表")
            if parts:
                result["body"] = "\n".join(parts)
            if tbls:
                result["tables"] = tbls
            result.pop("blocks", None)
        elif "blocks" in result:
            result.pop("blocks", None)
        children = result.get("children")
        result["children"] = [self.__normalize_node(c) for c in children] if isinstance(children, list) else []
        return result

    def __normalize_content(self, content):
        if not isinstance(content, dict) or not isinstance(content.get("sections"), list):
            return copy.deepcopy(DEFAULT_UTR_CONTENT)
        return {"sections": [self.__normalize_node(s) for s in content["sections"]]}

    def __replace_name(self, node, base, name, skip_titles=None):
        if not name or base == name:
            return
        if skip_titles and str(node.get("title") or "").strip() in skip_titles:
            return
        if node.get("body"):
            node["body"] = node["body"].replace(base, name)
        for tbl in (node.get("tables") or []):
            for row in tbl:
                for i in range(len(row)):
                    if isinstance(row[i], str) and base in row[i]:
                        row[i] = row[i].replace(base, name)
        for c in (node.get("children") or []):
            self.__replace_name(c, base, name, skip_titles=skip_titles)

    def __dhf_file_no(self, prod_id):
        if not prod_id:
            return ""
        row = db.session.execute(
            select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name == DOC_NAME).order_by(ProdDhf.id.asc())
        ).scalars().first()
        if not row:
            row = db.session.execute(
                select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name.like(f"%{DOC_NAME}%")).order_by(ProdDhf.id.asc())
            ).scalars().first()
        return (row.code or "").strip() if row and row.code else ""

    def __fill_revision(self, content, prod_id, version, force=False):
        """文件修订记录首行：修改日期(评审/封面日期)、版本号、首次发布、修订人(TPM)、批准人(研发负责人)。
        force=True（切换产品）时强制覆盖，无数据则置空。"""
        rev_date = serv_review_util.cover_date(prod_id, DOC_KEY) if prod_id else ""
        reviser = approver = ""
        if prod_id:
            members = db.session.execute(select(ProjectMember).where(ProjectMember.prod_id == prod_id)).scalars().all()
            reviser = next((m.name for m in members if "TPM" in str(m.role or "")), "")
            approver = next((m.name for m in members if "研发负责人" in str(m.role or "")), "")
        for s in (content.get("sections") or []):
            if s.get("ref_type") != "revision":
                continue
            tables = s.get("tables") or []
            if tables and isinstance(tables[0], list) and tables[0]:
                t = tables[0]
                cols = len(t[0]) if isinstance(t[0], list) and t[0] else 5
                while len(t) < 6:
                    t.append([""] * cols)
                if len(t) >= 2 and isinstance(t[1], list) and len(t[1]) >= 3:
                    row = t[1]
                    if force:
                        row[0] = rev_date or ""
                        if len(row) >= 2: row[1] = str(version or "")
                        if len(row) >= 3: row[2] = "首次发布"
                        if len(row) >= 4: row[3] = reviser
                        if len(row) >= 5: row[4] = approver
                    else:
                        if not str(row[0] or "").strip(): row[0] = rev_date
                        if version and len(row) >= 2 and not str(row[1] or "").strip(): row[1] = str(version)
                        if len(row) >= 3 and not str(row[2] or "").strip(): row[2] = "首次发布"
                        if len(row) >= 4 and reviser and not str(row[3] or "").strip(): row[3] = reviser
                        if len(row) >= 5 and approver and not str(row[4] or "").strip(): row[4] = approver
            break
        return content

    TESTITEM_HEADER = ["需求编号", "Case编号", "功能点概述", "用例个数", "是否通过", "备注"]

    def __stage_rows(self, prod_id, stage, prev):
        """按阶段(测试集)取用例，按功能点分组：需求编号/Case编号取段/功能点/用例个数/是否通过/备注。"""
        set_ids = db.session.execute(
            select(TestSet.id).where(TestSet.product_id == prod_id, TestSet.stage == stage)
        ).scalars().all()
        if not set_ids:
            return []
        cases = db.session.execute(
            select(TestCase).where(TestCase.set_id.in_(set_ids)).order_by(TestCase.id)
        ).scalars().all()
        order, gmap = [], {}
        for c in cases:
            func = (c.function or "").strip()
            key = func or (c.code or "").strip()
            if key not in gmap:
                gmap[key] = {"func": func, "codes": [], "srs": (c.srs_code or "").strip()}
                order.append(key)
            code = (c.code or "").strip()
            if code:
                gmap[key]["codes"].append(code)
            if not gmap[key]["srs"]:
                gmap[key]["srs"] = (c.srs_code or "").strip()
        rows = []
        for key in order:
            g = gmap[key]
            cs = sorted(g["codes"])
            case_no = "" if not cs else (cs[0] if len(cs) == 1 else f"{cs[0]} 至 {cs[-1]}")
            p = prev.get(g["func"], {})
            rows.append([g["srs"], case_no, g["func"], str(len(g["codes"])), p.get("pass", "是"), p.get("note", "")])
        return rows

    def __fill_test_items(self, content, prod_id):
        if not prod_id or not isinstance(content, dict):
            return content

        def walk(node):
            if node.get("ref_type") == "test_items":
                prev = {}
                for ch in (node.get("children") or []):
                    for tbl in (ch.get("tables") or []):
                        for r in (tbl or [])[1:]:
                            if isinstance(r, list) and len(r) >= 6 and str(r[2] or "").strip():
                                prev[str(r[2]).strip()] = {"pass": str(r[4] or "").strip() or "是", "note": str(r[5] or "").strip()}
                children = []
                for stage in (node.get("stages") or []):
                    grid = [list(self.TESTITEM_HEADER)] + self.__stage_rows(prod_id, stage, prev)
                    children.append({"title": stage, "body": "", "tables": [grid], "children": []})
                node["children"] = children
                node["tables"] = []
            for ch in (node.get("children") or []):
                walk(ch)

        for s in (content.get("sections") or []):
            walk(s)
        return content

    def __fill_workload(self, content, prod_id):
        """「测试工作量」表：按任务(测试计划/测试用例设计/测试执行/测试总结)从时间线取日期，填开始/结束时间。"""
        if not prod_id or not isinstance(content, dict):
            return content
        doc_date = serv_review_util.cover_date(prod_id, DOC_KEY) or ""
        is_user = DOC_KEY == "utr"
        if is_user:
            range_kw = {"计划": ["用户测试计划"], "用例": ["用户测试用例"],
                        "执行": ["用户测试记录"], "报告": ["用户测试报告"]}
        else:
            range_kw = {"计划": ["软件测试计划"],
                        "用例": ["单元测试用例", "集成测试用例", "系统测试用例"],
                        "执行": ["单元测试记录", "集成测试记录", "系统测试记录"],
                        "报告": ["软件测试报告"]}

        def kind_of(task):
            a = str(task or "")
            if "总结" in a or "报告" in a:
                return "报告"
            if "用例" in a or "设计" in a:
                return "用例"
            if "执行" in a:
                return "执行"
            if "计划" in a:
                return "计划"
            return ""

        def _range_days(a, b):
            m1 = re.match(r"(\d{4})\D+(\d{1,2})\D+(\d{1,2})", str(a or ""))
            m2 = re.match(r"(\d{4})\D+(\d{1,2})\D+(\d{1,2})", str(b or ""))
            if not (m1 and m2):
                return 0
            from datetime import date as _d
            try:
                return (_d(int(m2.group(1)), int(m2.group(2)), int(m2.group(3)))
                        - _d(int(m1.group(1)), int(m1.group(2)), int(m1.group(3)))).days + 1
            except Exception:
                return 0

        def walk(node):
            for tbl in (node.get("tables") or []):
                if not tbl or not isinstance(tbl[0], list):
                    continue
                header = [str(c or "") for c in tbl[0]]
                hj = " ".join(header)
                if "开始时间" in hj and "结束时间" in hj:
                    ci_s = next((i for i, h in enumerate(header) if "开始时间" in h), 2)
                    ci_e = next((i for i, h in enumerate(header) if "结束时间" in h), 3)
                    ci_days = next((i for i, h in enumerate(header) if "天数" in h), -1)
                    for r in tbl[1:]:
                        if not isinstance(r, list):
                            continue
                        task = ""
                        for i in range(min(ci_s, len(r))):
                            v = str(r[i] or "").strip()
                            if v and v != "实际":
                                task = v
                        kind = kind_of(task)
                        start, end = ("", "")
                        if kind and kind in range_kw:
                            start, end = serv_review_util.date_range(prod_id, range_kw[kind])
                        start = start or doc_date
                        end = end or start
                        if ci_s < len(r):
                            r[ci_s] = start
                        if ci_e < len(r):
                            r[ci_e] = end
                        if 0 <= ci_days < len(r) and kind and kind in range_kw:
                            _days = serv_review_util.date_count(prod_id, range_kw[kind])
                            if _days:
                                r[ci_days] = "%d天" % _days
            for ch in (node.get("children") or []):
                walk(ch)

        for s in (content.get("sections") or []):
            walk(s)
        return content

    def __fill_hr(self, content, prod_id, force=False):
        """人力资源表：按「角色」列自动匹配当前产品参与人员，填「资源数量/具体人员」列(N人/姓名)。
        force=False：匹配不到保留原值；force=True：强制覆盖，无人员则置空。"""
        if not isinstance(content, dict):
            return content
        doc_date = serv_review_util.cover_date(prod_id, DOC_KEY) if prod_id else ""
        members = db.session.execute(
            select(ProjectMember).where(ProjectMember.prod_id == prod_id)
        ).scalars().all() if prod_id else []
        primary = "宋月" if serv_review_util._before_202509(doc_date) else "孙家旭"

        def by_kw(pred):
            out = []
            for m in members:
                if pred(str(m.role or "")):
                    nm = (m.name or "").strip()
                    if nm and nm not in out:
                        out.append(nm)
            return out

        def names_for(label, duty):
            lab, d = str(label or ""), str(duty or "")
            if "用户" in lab and "测试" in lab:
                return by_kw(lambda r: "用户测试" in r)
            if "测试" in lab and "用户" not in lab:
                others = [n for n in by_kw(lambda r: ("测试" in r) and ("用户" not in r)) if n not in ("宋月", "孙家旭")]
                return ([primary] + others) if any(k in d for k in ("执行", "记录", "回归", "首轮")) else [primary]
            if "产品经理" in lab:
                return by_kw(lambda r: "产品经理" in r)
            if "开发" in lab and "负责" not in lab:
                return by_kw(lambda r: "开发人员" in r or "开发工程师" in r)
            if "QA" in lab:
                return by_kw(lambda r: "QA" in r or "质量" in r)
            if "RA" in lab:
                return by_kw(lambda r: "RA" in r or "法规" in r)
            if "临床" in lab:
                return by_kw(lambda r: "临床" in r)
            return by_kw(lambda r: bool(lab) and (lab in r or r in lab))

        def walk(node):
            for tbl in (node.get("tables") or []):
                if not tbl or not isinstance(tbl[0], list):
                    continue
                header = [str(c or "") for c in tbl[0]]
                if not (header and "角色" in header[0] and any(("具体人员" in h or "资源数量" in h) for h in header)):
                    continue
                ci_p = next((i for i, h in enumerate(header) if ("具体人员" in h or "资源数量" in h)), 1)
                ci_d = next((i for i, h in enumerate(header) if "职责" in h), 2)
                for r in tbl[1:]:
                    if not isinstance(r, list) or not r:
                        continue
                    nm = names_for(r[0], r[ci_d] if ci_d < len(r) else "")
                    if force:
                        if ci_p < len(r):
                            r[ci_p] = ("%d人/%s" % (len(nm), "、".join(nm))) if nm else ""
                    elif nm and ci_p < len(r):
                        r[ci_p] = "%d人/%s" % (len(nm), "、".join(nm))
            for ch in (node.get("children") or []):
                walk(ch)

        for s in (content.get("sections") or []):
            walk(s)
        return content

    def __fill_refs(self, content, prod_id):
        """引用文档/参考文件表：「作者或来源」列若为已知模板姓名，按角色替换为当前产品对应成员。"""
        if not prod_id or not isinstance(content, dict):
            return content
        doc_date = serv_review_util.cover_date(prod_id, DOC_KEY) or ""
        members = db.session.execute(select(ProjectMember).where(ProjectMember.prod_id == prod_id)).scalars().all()
        primary = "宋月" if serv_review_util._before_202509(doc_date) else "孙家旭"
        name2role = {"吴福乐": "pm", "余航": "pm", "杨静": "pm", "沈宏": "devlead", "宁随军": "tpm",
                     "林金贵": "qa", "夏晨": "pdir", "张淑芳": "ra", "宋月": "test", "王小敏": "test",
                     "孙家旭": "test", "徐冕": "utest", "罗鹏宇": "utest", "刘锦龙": "utest",
                     "徐学强": "utest", "刘娜": "utest"}

        def first(pred):
            for m in members:
                if pred(str(m.role or "")):
                    nm = (m.name or "").strip()
                    if nm:
                        return nm
            return ""

        def resolve(role):
            return {
                "test": primary,
                "utest": first(lambda r: "用户测试" in r),
                "pm": first(lambda r: "产品经理" in r),
                "devlead": first(lambda r: "研发负责人" in r),
                "tpm": first(lambda r: "TPM" in r),
                "qa": first(lambda r: "QA" in r or "质量" in r),
                "pdir": first(lambda r: "产品负责人" in r or "产品总监" in r),
                "ra": first(lambda r: "RA" in r or "法规" in r),
            }.get(role, "")

        def walk(node):
            for tbl in (node.get("tables") or []):
                if not tbl or not isinstance(tbl[0], list):
                    continue
                header = [str(c or "") for c in tbl[0]]
                if not (any("文档名称" in h for h in header) and any(("作者" in h or "来源" in h) for h in header)):
                    continue
                ci_a = next((i for i, h in enumerate(header) if ("作者" in h or "来源" in h)), 1)
                for r in tbl[1:]:
                    if not isinstance(r, list) or ci_a >= len(r):
                        continue
                    cur = str(r[ci_a] or "").strip()
                    role = name2role.get(cur)
                    if role:
                        nm = resolve(role)
                        if nm:
                            r[ci_a] = nm
            for ch in (node.get("children") or []):
                walk(ch)

        for s in (content.get("sections") or []):
            walk(s)
        return content

    def __autofill(self, content, prod_id, product=None, version="", force=False):
        if not isinstance(content, dict):
            return content
        if product and product.name:
            for s in (content.get("sections") or []):
                self.__replace_name(s, BASE_NAME, product.name)
            # 更新完整版本号
            new_full_version = (product.full_version or "").strip()
            if new_full_version:
                def _update_version(node):
                    if node.get("body") and "完整版本" in str(node.get("body")):
                        node["body"] = re.sub(r"完整版本：[^\n]*", f"完整版本：{new_full_version}", str(node["body"]))
                    # 表格中的版本号替换（如"测试版本"行的值含 1.1.0.1）
                    for tbl in (node.get("tables") or []):
                        for row in tbl:
                            for i in range(len(row)):
                                if isinstance(row[i], str) and re.search(r"\d+\.\d+\.\d+\.\d+", str(row[i])):
                                    row[i] = re.sub(r"\d+\.\d+\.\d+\.\d+", new_full_version, str(row[i]))
                    for c in (node.get("children") or []):
                        _update_version(c)
                for s in (content.get("sections") or []):
                    _update_version(s)
        self.__fill_test_items(content, prod_id)
        self.__fill_hr(content, prod_id, force=force)
        self.__fill_refs(content, prod_id)
        self.__fill_workload(content, prod_id)
        self.__fill_revision(content, prod_id, version, force=force)
        serv_review_util.ensure_review(
            content, DOC_KEY,
            serv_review_util.review_date(prod_id, serv_review_util.REVIEW_DEFS[DOC_KEY]["name_keywords"]) if prod_id else "",
            prod_id,
        )
        serv_review_util.fill_cover_dates(content, serv_review_util.cover_date(prod_id, DOC_KEY) if prod_id else "", force=force)
        serv_review_util.fill_cover_signers(content, serv_review_util.cover_signers(prod_id, DOC_KEY) if prod_id else {}, force=force)
        return content

    def __to_obj(self, row: UtrDoc, product: Product = None):
        obj = UtrDocObj(**row.dict())
        obj.content = self.__autofill(self.__normalize_content(obj.content), row.product_id, product, row.version)
        if not (obj.file_no or "").strip():
            obj.file_no = self.__dhf_file_no(row.product_id)
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
        return obj

    async def add_utr_doc(self, form: UtrDocForm):
        try:
            sql = select(func.count(UtrDoc.id)).where(UtrDoc.product_id == form.product_id, UtrDoc.version == form.version)
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = UtrDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=UtrDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_utr_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: UtrDoc = db.session.execute(select(UtrDoc).where(UtrDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(UtrDoc.version).where(UtrDoc.product_id == target_pid)).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = UtrDoc(
                product_id=target_pid, version=version,
                file_no=sync_file_no_version((fromdoc.file_no or "").strip() or self.__dhf_file_no(target_pid), version) or None,
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            if target_pid != fromdoc.product_id:
                await self.rebind_product(newdoc.id, target_pid)
            return Resp.resp_ok(data=UtrDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_utr_doc(self, form: UtrDocForm):
        try:
            row: UtrDoc = db.session.execute(select(UtrDoc).where(UtrDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def rebind_product(self, id: int, product_id: int):
        """切换产品：更新 product_id 并强制用新产品信息重新获取封面/修订/产品名后保存，返回新 obj。"""
        try:
            row: UtrDoc = db.session.execute(select(UtrDoc).where(UtrDoc.id == id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            old_product: Product = db.session.execute(select(Product).where(Product.id == row.product_id)).scalars().first() if row.product_id else None
            product: Product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
            if not product:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            db.session.execute(delete(UtrDoc).where(UtrDoc.product_id == product_id, UtrDoc.version == row.version, UtrDoc.id != id))
            content = self.__normalize_content(row.content)
            # 重置含产品名/完整版本的固定章节为模板原值（恢复基准名+原版本号，避免被旧产品污染）
            tpl = copy.deepcopy(DEFAULT_UTR_CONTENT) if isinstance(DEFAULT_UTR_CONTENT, dict) else {"sections": []}
            tpl_map = {}
            def collect_tpl(node):
                key = str(node.get("title") or "").strip()
                body = str(node.get("body") or "")
                caps = node.get("table_captions") or []
                tables = node.get("tables") or []
                if BASE_NAME in body or "完整版本" in body or any(BASE_NAME in str(r) for tbl in tables for r in tbl) or any("1.1.0.1" in str(r) for tbl in tables for r in tbl) or any("1.1.0.1" in str(c) for c in caps):
                    tpl_map[key] = {"body": node.get("body", ""), "tables": copy.deepcopy(tables), "table_captions": copy.deepcopy(caps)}
                for c in (node.get("children") or []):
                    collect_tpl(c)
            for s in (tpl.get("sections") or []):
                collect_tpl(s)
            def reset_fixed(node):
                key = str(node.get("title") or "").strip()
                if key in tpl_map:
                    node["body"] = tpl_map[key]["body"]
                    node["tables"] = copy.deepcopy(tpl_map[key]["tables"])
                    if tpl_map[key].get("table_captions"):
                        node["table_captions"] = copy.deepcopy(tpl_map[key]["table_captions"])
                for c in (node.get("children") or []):
                    reset_fixed(c)
            for s in (content.get("sections") or []):
                reset_fixed(s)
            # 完整版本号更新为新产品版本（body + tables + table_captions）
            new_full_version = (product.full_version or "").strip()
            def update_version(node):
                if node.get("body") and "完整版本" in str(node.get("body")):
                    node["body"] = re.sub(r"完整版本：[^\n]*", f"完整版本：{new_full_version}", str(node["body"]))
                for tbl in (node.get("tables") or []):
                    for r in tbl:
                        for i in range(len(r)):
                            if isinstance(r[i], str) and re.search(r"\d+\.\d+\.\d+\.\d+", r[i]):
                                r[i] = re.sub(r"\d+\.\d+\.\d+\.\d+", new_full_version, r[i])
                caps = node.get("table_captions") or []
                if caps:
                    node["table_captions"] = [re.sub(r"\d+\.\d+\.\d+\.\d+", new_full_version, str(c)) if re.search(r"\d+\.\d+\.\d+\.\d+", str(c)) else c for c in caps]
                for c in (node.get("children") or []):
                    update_version(c)
            for s in (content.get("sections") or []):
                update_version(s)
            row.product_id = product_id
            # 测试分布图：未获取时从该产品 doc_file（img_flow）获取，已获取的保留
            def fill_test_chart(nodes):
                for n in nodes:
                    if str(n.get("title") or "").strip() == "测试分布图":
                        imgs = n.get("images") or []
                        if not imgs:
                            row_df = db.session.execute(
                                select(DocFile).where(DocFile.product_id == product_id, DocFile.category == "img_flow").order_by(DocFile.id)
                            ).scalars().first()
                            if row_df and row_df.file_url:
                                url = str(row_df.file_url).strip()
                                if url.startswith("data.trace/"):
                                    url = "/" + url
                                n["images"] = [url]
                        break
                    for c in (n.get("children") or []):
                        fill_test_chart([c])
            fill_test_chart(content.get("sections") or [])
            content = self.__autofill(content, product_id, product, row.version, force=True)
            row.content = content
            db.session.commit()
            return Resp.resp_ok(data=self.__to_obj(row, product))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_utr_doc(self, id: int):
        db.session.execute(delete(UtrDoc).where(UtrDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_utr_doc(self, id: int):
        sql = select(UtrDoc, Product).join(Product, UtrDoc.product_id == Product.id).where(UtrDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_utr_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(UtrDoc.product_id == product_id)
        if version:
            wheres.append(UtrDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(UtrDoc.id)).join(Product, UtrDoc.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (select(UtrDoc, Product).join(Product, UtrDoc.product_id == Product.id).where(*wheres)
               .order_by(UtrDoc.id.desc()).offset(page_index * page_size).limit(page_size))
        rows: List[UtrDocObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    # ---------------- 导出 Word（PDP 风格章节树） ----------------
    async def export_utr_doc(self, output, id: int):
        resp = await self.get_utr_doc(id)
        obj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        c = self.__normalize_content(obj.content)
        sections = c.get("sections") or []
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        document.settings.element.append(update_fields)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")
        docx_util.add_page_number_footer(section, obj.file_no or "")

        def add_blank_lines(n):
            for _ in range(max(0, int(n or 0))):
                document.add_paragraph("")

        def write_center_title(text, size=22.0, bold=False):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            docx_util.fonted_txt(p, str(text or ""), font_size=size, bold=bold)

        def add_text(text):
            docx_util.save_txt2docx(str(text or ""), document)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            s = str(text or "")
            if s.startswith("data:image"):
                try:
                    b64 = s.split(",", 1)[1] if "," in s else ""
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = align
                    para.add_run().add_picture(BytesIO(base64.b64decode(b64)), height=Pt(33))
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    return
                except Exception:
                    pass
            cell.text = ""
            for i, line in enumerate(s.split("\n")):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER if align == WD_ALIGN_PARAGRAPH.CENTER else WD_CELL_VERTICAL_ALIGNMENT.TOP

        def _set_fixed_widths(table, widths_dxa):
            table.autofit = False
            table.allow_autofit = False
            tbl_pr = table._tbl.tblPr
            layout = tbl_pr.find(qn("w:tblLayout"))
            if layout is None:
                layout = OxmlElement("w:tblLayout")
                tbl_pr.append(layout)
            layout.set(qn("w:type"), "fixed")
            grid_el = table._tbl.find(qn("w:tblGrid"))
            if grid_el is not None:
                for gc in list(grid_el):
                    grid_el.remove(gc)
                for w in widths_dxa:
                    gc = OxmlElement("w:gridCol")
                    gc.set(qn("w:w"), str(w))
                    grid_el.append(gc)
            # 逐格写 tcW，确保固定列宽生效
            for row in table.rows:
                for i, w in enumerate(widths_dxa):
                    if i < len(row.cells):
                        tcpr = row.cells[i]._tc.get_or_add_tcPr()
                        tcw = tcpr.find(qn("w:tcW"))
                        if tcw is None:
                            tcw = OxmlElement("w:tcW")
                            tcpr.append(tcw)
                        tcw.set(qn("w:w"), str(w))
                        tcw.set(qn("w:type"), "dxa")

        def add_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for r_idx, row in enumerate(grid):
                cells = table.add_row().cells
                for c_idx in range(cols):
                    set_cell(cells[c_idx], row[c_idx] if c_idx < len(row) else "", bold=(r_idx == 0))
            # SCI 清单表：固定列宽，SCI名字/类型不换行、存储地点较宽
            first = str(grid[0][0]).strip() if grid and grid[0] else ""
            if first == "SCI名字":
                if cols == 5:
                    _set_fixed_widths(table, [1700, 1100, 2000, 1500, 3000])
                elif cols == 4:
                    _set_fixed_widths(table, [1700, 2200, 1500, 3600])
            elif first == "类别" and cols >= 3:
                # 配置项存储地址表：类别列收窄、存储路径较宽；首列连续相同「类别」纵向合并；文字左对齐
                if cols == 3:
                    _set_fixed_widths(table, [1300, 1300, 5000])
                rows = table.rows
                n = len(rows)
                r = 1
                while r < n:
                    val = str((grid[r][0] if r < len(grid) and grid[r] else "") or "").strip()
                    if not val:
                        r += 1
                        continue
                    r2 = r
                    while r2 + 1 < n and str((grid[r2 + 1][0] if grid[r2 + 1] else "") or "").strip() == val:
                        r2 += 1
                    if r2 > r:
                        merged = rows[r].cells[0].merge(rows[r2].cells[0])
                        set_cell(merged, val, align=WD_ALIGN_PARAGRAPH.LEFT)
                        merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    r = r2 + 1
            elif first in ("测试目标", "测试方法和技术") and cols == 2:
                _set_fixed_widths(table, [2200, 7100])
            elif first == "配置":
                if cols == 2:
                    _set_fixed_widths(table, [1900, 7400])
                elif cols == 3:
                    _set_fixed_widths(table, [1900, 3700, 3700])
            document.add_paragraph()

        def add_cover_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for row in grid:
                cells = table.add_row().cells
                for c_idx in range(cols):
                    text = row[c_idx] if c_idx < len(row) else ""
                    set_cell(cells[c_idx], text, bold=(c_idx % 2 == 0), align=WD_ALIGN_PARAGRAPH.CENTER)
                if (str(row[0]).strip() if row else "") == "生效日期" and cols > 2:
                    merged = cells[1]
                    for c_idx in range(2, cols):
                        merged = merged.merge(cells[c_idx])
                    set_cell(merged, row[1] if len(row) > 1 else "", align=WD_ALIGN_PARAGRAPH.CENTER)
            document.add_paragraph()

        def strip_num(title):
            return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()

        def add_body_heading(title, level):
            size = {1: 16.0, 2: 14.0, 3: 12.0}.get(level, 11.0)
            p = document.add_heading("", level=max(1, min(level, 9)))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, title, font_size=size, bold=True)

        def render_body_section(node, level, number=""):
            name = strip_num(node.get("title"))
            heading = f"{number} {name}".strip() if number else name
            add_body_heading(heading, level=max(1, min(level, 9)))
            if node.get("ref_type") == "review":
                for t_idx, table in enumerate(node.get("tables") or []):
                    serv_review_util.render_review_grid(document, table, set_cell, merge_col0=(t_idx == 0), merge_full=True)
            else:
                body_text = str(node.get("body") or "")
                tables = node.get("tables") or []
                caps = node.get("table_captions") or []
                def _emit_table(t_idx):
                    cap = caps[t_idx] if t_idx < len(caps) else ""
                    if str(cap or "").strip():
                        add_text(cap)
                    add_grid(tables[t_idx])
                # 按"见下表"/"表N"切分 body，交错输出正文段和表格
                if (("见下表" in body_text) or re.search(r"(?m)^表\s*\d", body_text)) and tables:
                    lines = body_text.split("\n")
                    buf = []
                    tbl_idx = 0
                    def flush_text():
                        t = "\n".join(buf).strip()
                        if t:
                            add_text(t)
                        buf.clear()
                    for ln in lines:
                        if ("见下表" in ln.strip() or re.match(r"^表\s*\d", ln.strip())) and tbl_idx < len(tables):
                            buf.append(ln)  # 保留"见下表"/"表N"行作为正文
                            flush_text()
                            _emit_table(tbl_idx)
                            tbl_idx += 1
                        else:
                            buf.append(ln)
                    flush_text()
                    for i in range(tbl_idx, len(tables)):
                        _emit_table(i)
                else:
                    if body_text.strip():
                        add_text(body_text)
                    for t_idx in range(len(tables)):
                        _emit_table(t_idx)
                # 分布图等节点图片（data URL）内嵌，排在表格之后、表格后正文之前
                for _img in (node.get("charts") or []):
                    if isinstance(_img, str) and _img.startswith("data:image"):
                        try:
                            _p = document.add_paragraph()
                            _p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            _p.add_run().add_picture(BytesIO(base64.b64decode(_img.split(",", 1)[1])), width=Inches(6.2))
                        except Exception:
                            pass
                if str(node.get("body_tail") or "").strip():
                    add_text(node.get("body_tail"))
            idx = 0
            for child in (node.get("children") or []):
                idx += 1
                child_num = f"{number}.{idx}" if number else f"{idx}"
                render_body_section(child, level + 1, child_num)

        cover = next((s for s in sections if s.get("ref_type") == "cover"), None)
        revision = next((s for s in sections if s.get("ref_type") == "revision"), None)
        body = [s for s in sections if s.get("ref_type") not in ("cover", "revision")]

        add_blank_lines(6)
        write_center_title((strip_num(cover.get("title")) if cover else "") or DOC_NAME, size=22.0, bold=True)
        add_blank_lines(4)
        if cover:
            for table in (cover.get("tables") or []):
                add_cover_grid(table)

        document.add_page_break()
        write_center_title("文件修订记录", size=14.0, bold=True)
        add_blank_lines(2)
        if revision:
            for table in (revision.get("tables") or []):
                add_grid(table)

        document.add_page_break()
        write_center_title("目录", size=16.0, bold=True)
        docx_util.insert_toc_field(document)

        document.add_page_break()
        seq = 0
        for node in body:
            if node.get("ref_type") in ("review", "attachment"):
                render_body_section(node, 1, "")
            else:
                seq += 1
                render_body_section(node, 1, str(seq))

        docx_util.fill_toc_cache(document)
        document.save(output)
        output.seek(0)
