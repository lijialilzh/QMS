#!/usr/bin/env python
# encoding: utf-8

# 软件测试计划服务层（开发文件，PDP 风格章节树）。默认内容取自 src-res/stp_default_content.json。

import base64
import copy
import json
import logging
import os
import re
from io import BytesIO
from typing import List

from sqlalchemy import delete, func, select
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.stp_doc import StpDoc
from ..model.prod_dhf import ProdDhf
from ..model.project_member import ProjectMember
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_stp_doc import StpDocForm
from ..obj.vobj_stp_doc import StpDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from . import serv_review_util
from .serv_utils import new_version, sync_file_no_version
from .serv_utils import docx_util

logger = logging.getLogger(__name__)

_DEFAULT_FILE = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "src-res", "stp_default_content.json")
try:
    with open(_DEFAULT_FILE, encoding="utf-8") as _f:
        DEFAULT_STP_CONTENT = json.load(_f)
except Exception:
    DEFAULT_STP_CONTENT = {"sections": []}

# 模板基准产品名（用于全文替换为当前产品名）。
BASE_NAME = "InferOperate Suite"
DOC_NAME = "软件测试计划"
DOC_KEY = "stp"


class Server(object):

    def __normalize_node(self, node):
        if not isinstance(node, dict):
            return {"title": str(node or ""), "body": "", "tables": [], "children": []}
        result = dict(node)
        result["title"] = str(result.get("title") or "")
        result["body"] = str(result.get("body") or "")
        tables = result.get("tables")
        if not isinstance(tables, list):
            tables = []
        norm_tables = []
        for table in tables:
            if isinstance(table, list):
                norm_tables.append([[str(c) if c is not None else "" for c in (row or [])] for row in table if isinstance(row, list)])
        result["tables"] = norm_tables
        children = result.get("children")
        result["children"] = [self.__normalize_node(c) for c in children] if isinstance(children, list) else []
        return result

    def __normalize_content(self, content):
        if not isinstance(content, dict) or not isinstance(content.get("sections"), list):
            return copy.deepcopy(DEFAULT_STP_CONTENT)
        return {"sections": [self.__normalize_node(s) for s in content["sections"]]}

    def __replace_name(self, node, base, name, skip_titles=None):
        if not name or base == name:
            return
        if skip_titles and str(node.get("title") or "").strip() in skip_titles:
            return
        if node.get("body"):
            node["body"] = node["body"].replace(base, name)
        for tbl in (node.get("tables") or []):
            for row in tbl:
                for i in range(len(row)):
                    if isinstance(row[i], str) and base in row[i]:
                        row[i] = row[i].replace(base, name)
        for c in (node.get("children") or []):
            self.__replace_name(c, base, name, skip_titles=skip_titles)

    def __dhf_file_no(self, prod_id):
        if not prod_id:
            return ""
        row = db.session.execute(
            select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name == DOC_NAME).order_by(ProdDhf.id.asc())
        ).scalars().first()
        if not row:
            row = db.session.execute(
                select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name.like(f"%{DOC_NAME}%")).order_by(ProdDhf.id.asc())
            ).scalars().first()
        return (row.code or "").strip() if row and row.code else ""

    def __fill_revision(self, content, prod_id, version):
        """文件修订记录首行：修改日期(评审/封面日期)、版本号、首次发布、修订人(TPM)、批准人(研发负责人)。"""
        rev_date = serv_review_util.cover_date(prod_id, DOC_KEY) if prod_id else ""
        reviser = approver = ""
        if prod_id:
            members = db.session.execute(select(ProjectMember).where(ProjectMember.prod_id == prod_id)).scalars().all()
            reviser = next((m.name for m in members if "TPM" in str(m.role or "")), "")
            approver = next((m.name for m in members if "研发负责人" in str(m.role or "")), "")
        for s in (content.get("sections") or []):
            if s.get("ref_type") != "revision":
                continue
            tables = s.get("tables") or []
            if tables and isinstance(tables[0], list) and tables[0]:
                t = tables[0]
                cols = len(t[0]) if isinstance(t[0], list) and t[0] else 5
                while len(t) < 6:  # 文件修订记录默认表头+5行
                    t.append([""] * cols)
                if len(t) >= 2 and isinstance(t[1], list) and len(t[1]) >= 3:
                    row = t[1]
                    if not str(row[0] or "").strip():
                        row[0] = rev_date
                    if version and len(row) >= 2 and not str(row[1] or "").strip():
                        row[1] = str(version)
                    if len(row) >= 3 and not str(row[2] or "").strip():
                        row[2] = "首次发布"
                    if len(row) >= 4 and reviser and not str(row[3] or "").strip():
                        row[3] = reviser
                    if len(row) >= 5 and approver and not str(row[4] or "").strip():
                        row[4] = approver
            break
        return content

    def __fill_test_plan(self, content, prod_id):
        """「测试计划」表：计划完成时间按行「测试活动」匹配时间线；负责人=主测试工程师(2025.09前宋月/之后孙家旭)+其他测试工程师。"""
        if not prod_id or not isinstance(content, dict):
            return content
        doc_date = serv_review_util.cover_date(prod_id, DOC_KEY) or ""
        members = db.session.execute(
            select(ProjectMember).where(ProjectMember.prod_id == prod_id)
        ).scalars().all()
        is_user = DOC_KEY == "utp"
        if is_user:
            date_kw = {"报告": ["用户测试报告"], "用例": ["用户测试用例", "用例"],
                       "计划": ["用户测试计划"], "执行": ["用户测试", "测试"]}

            def role_ok(r):
                return "用户测试" in r
        else:
            date_kw = {"报告": ["软件测试报告"],
                       "用例": ["软件测试用例", "单元测试用例", "集成测试用例", "系统测试用例", "测试用例"],
                       "计划": ["软件测试计划"],
                       "执行": ["软件测试", "系统测试", "集成测试", "单元测试", "测试"]}

            def role_ok(r):
                return ("测试" in r) and ("用户" not in r)
        # 该产品下的测试人员（软件测试计划=测试工程师；用户测试计划=用户测试），排除固定主测试工程师本身
        others = []
        for m in members:
            if role_ok(str(m.role or "")):
                nm = (m.name or "").strip()
                if nm and nm not in ("宋月", "孙家旭") and nm not in others:
                    others.append(nm)

        def to_month(d):
            mt = re.match(r"(\d{4})\D+(\d{1,2})", str(d or ""))
            return "%s.%02d" % (mt.group(1), int(mt.group(2))) if mt else str(d or "")

        def row_kind(activity):
            a = str(activity or "")
            if any(k in a for k in ("测试记录", "回归", "首轮")):
                return "执行"
            for kind in ("报告", "用例", "计划"):
                if kind in a:
                    return kind
            if "测试" in a:
                return "执行"
            return ""

        def row_date(kind):
            if kind in ("报告", "用例", "计划", "执行"):
                return serv_review_util.review_date(prod_id, date_kw[kind]) or doc_date
            return doc_date

        def owners(dt, kind):
            if is_user:
                names = list(others)
            else:
                primary = "宋月" if serv_review_util._before_202509(dt or doc_date) else "孙家旭"
                # 评审计划/测试用例/测试报告：只写主测试工程师一人；其余(执行)：主+其他测试人员
                if kind in ("报告", "用例", "计划"):
                    names = [primary]
                else:
                    names = [primary] + [n for n in others if n != primary]
            return "、".join([n for n in names if n])

        def walk(node):
            for tbl in (node.get("tables") or []):
                if not tbl or not isinstance(tbl[0], list):
                    continue
                header = [str(c or "") for c in tbl[0]]
                hj = " ".join(header)
                if "计划完成时间" in hj and "负责人" in hj:
                    ci_time = next((i for i, h in enumerate(header) if "计划完成时间" in h), 0)
                    ci_own = next((i for i, h in enumerate(header) if "负责人" in h), 1)
                    ci_act = next((i for i, h in enumerate(header) if "活动" in h), 2)
                    for r in tbl[1:]:
                        if not isinstance(r, list):
                            continue
                        act = r[ci_act] if ci_act < len(r) else ""
                        kind = row_kind(act)
                        dt = to_month(row_date(kind))
                        if ci_time < len(r):
                            r[ci_time] = dt
                        if ci_own < len(r):
                            r[ci_own] = owners(dt, kind)
            for ch in (node.get("children") or []):
                walk(ch)

        for s in (content.get("sections") or []):
            walk(s)
        return content

    def __fill_roles(self, content, prod_id):
        """「角色和职责」表：按行「角色」自动匹配该产品参与人员，填入「人员」列（匹配不到则保留原值）。"""
        if not prod_id or not isinstance(content, dict):
            return content
        doc_date = serv_review_util.cover_date(prod_id, DOC_KEY) or ""
        members = db.session.execute(
            select(ProjectMember).where(ProjectMember.prod_id == prod_id)
        ).scalars().all()

        def kind_of(duty):
            a = str(duty or "")
            if any(k in a for k in ("测试记录", "回归", "首轮")):
                return "执行"
            for k in ("报告", "用例", "计划"):
                if k in a:
                    return k
            return ""

        def names_for(role_label, kind):
            label = str(role_label or "")
            if ("测试" in label) and ("用户" not in label):
                primary = "宋月" if serv_review_util._before_202509(doc_date) else "孙家旭"
                others = []
                for m in members:
                    r = str(m.role or "")
                    if ("测试" in r) and ("用户" not in r):
                        nm = (m.name or "").strip()
                        if nm and nm not in ("宋月", "孙家旭") and nm not in others:
                            others.append(nm)
                names = [primary] if kind in ("报告", "用例", "计划") else [primary] + others
                return "、".join([n for n in names if n])
            if "用户测试" in label:
                names = [(m.name or "").strip() for m in members if "用户测试" in str(m.role or "")]
                names = [n for n in names if n]
                return "、".join(names) if names else None
            hit = []
            for m in members:
                r = str(m.role or "")
                if label and (label in r or r in label):
                    nm = (m.name or "").strip()
                    if nm and nm not in hit:
                        hit.append(nm)
            return "、".join(hit) if hit else None

        def walk(node):
            for tbl in (node.get("tables") or []):
                if not tbl or not isinstance(tbl[0], list):
                    continue
                header = [str(c or "") for c in tbl[0]]
                hj = " ".join(header)
                if "人员" in hj and "角色" in hj and "职责" in hj:
                    ci_p = next((i for i, h in enumerate(header) if "人员" in h), 0)
                    ci_r = next((i for i, h in enumerate(header) if "角色" in h), 1)
                    ci_d = next((i for i, h in enumerate(header) if "职责" in h), 2)
                    for r in tbl[1:]:
                        if not isinstance(r, list):
                            continue
                        role_label = r[ci_r] if ci_r < len(r) else ""
                        duty = r[ci_d] if ci_d < len(r) else ""
                        nm = names_for(role_label, kind_of(duty))
                        if nm and ci_p < len(r):
                            r[ci_p] = nm
            for ch in (node.get("children") or []):
                walk(ch)

        for s in (content.get("sections") or []):
            walk(s)
        return content

    def __autofill(self, content, prod_id, product=None, version=""):
        if not isinstance(content, dict):
            return content
        if product and product.name:
            for s in (content.get("sections") or []):
                self.__replace_name(s, BASE_NAME, product.name)
        self.__fill_test_plan(content, prod_id)
        self.__fill_roles(content, prod_id)
        self.__fill_revision(content, prod_id, version)
        serv_review_util.ensure_review(
            content, DOC_KEY,
            serv_review_util.review_date(prod_id, serv_review_util.REVIEW_DEFS[DOC_KEY]["name_keywords"]) if prod_id else "",
            prod_id,
        )
        serv_review_util.fill_cover_dates(content, serv_review_util.cover_date(prod_id, DOC_KEY) if prod_id else "")
        serv_review_util.fill_cover_signers(content, serv_review_util.cover_signers(prod_id, DOC_KEY) if prod_id else {})
        return content

    def __to_obj(self, row: StpDoc, product: Product = None):
        obj = StpDocObj(**row.dict())
        obj.content = self.__autofill(self.__normalize_content(obj.content), row.product_id, product, row.version)
        if not (obj.file_no or "").strip():
            obj.file_no = self.__dhf_file_no(row.product_id)
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
        return obj

    async def add_stp_doc(self, form: StpDocForm):
        try:
            sql = select(func.count(StpDoc.id)).where(StpDoc.product_id == form.product_id, StpDoc.version == form.version)
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = StpDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=StpDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_stp_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: StpDoc = db.session.execute(select(StpDoc).where(StpDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(StpDoc.version).where(StpDoc.product_id == target_pid)).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = StpDoc(
                product_id=target_pid, version=version,
                file_no=sync_file_no_version(fromdoc.file_no, version),
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            return Resp.resp_ok(data=StpDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_stp_doc(self, form: StpDocForm):
        try:
            row: StpDoc = db.session.execute(select(StpDoc).where(StpDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_stp_doc(self, id: int):
        db.session.execute(delete(StpDoc).where(StpDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_stp_doc(self, id: int):
        sql = select(StpDoc, Product).join(Product, StpDoc.product_id == Product.id).where(StpDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_stp_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(StpDoc.product_id == product_id)
        if version:
            wheres.append(StpDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(StpDoc.id)).join(Product, StpDoc.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (select(StpDoc, Product).join(Product, StpDoc.product_id == Product.id).where(*wheres)
               .order_by(StpDoc.id.desc()).offset(page_index * page_size).limit(page_size))
        rows: List[StpDocObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    # ---------------- 导出 Word（PDP 风格章节树） ----------------
    async def export_stp_doc(self, output, id: int):
        resp = await self.get_stp_doc(id)
        obj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        c = self.__normalize_content(obj.content)
        sections = c.get("sections") or []
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        document.settings.element.append(update_fields)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")
        docx_util.add_page_number_footer(section, obj.file_no or "")

        def add_blank_lines(n):
            for _ in range(max(0, int(n or 0))):
                document.add_paragraph("")

        def write_center_title(text, size=22.0, bold=False):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            docx_util.fonted_txt(p, str(text or ""), font_size=size, bold=bold)

        def add_text(text):
            docx_util.save_txt2docx(str(text or ""), document)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            s = str(text or "")
            if s.startswith("data:image"):
                try:
                    b64 = s.split(",", 1)[1] if "," in s else ""
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = align
                    para.add_run().add_picture(BytesIO(base64.b64decode(b64)), height=Pt(33))
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    return
                except Exception:
                    pass
            cell.text = ""
            for i, line in enumerate(s.split("\n")):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER if align == WD_ALIGN_PARAGRAPH.CENTER else WD_CELL_VERTICAL_ALIGNMENT.TOP

        def _set_fixed_widths(table, widths_dxa):
            table.autofit = False
            table.allow_autofit = False
            tbl_pr = table._tbl.tblPr
            layout = tbl_pr.find(qn("w:tblLayout"))
            if layout is None:
                layout = OxmlElement("w:tblLayout")
                tbl_pr.append(layout)
            layout.set(qn("w:type"), "fixed")
            grid_el = table._tbl.find(qn("w:tblGrid"))
            if grid_el is not None:
                for gc in list(grid_el):
                    grid_el.remove(gc)
                for w in widths_dxa:
                    gc = OxmlElement("w:gridCol")
                    gc.set(qn("w:w"), str(w))
                    grid_el.append(gc)
            # 逐格写 tcW，确保固定列宽生效
            for row in table.rows:
                for i, w in enumerate(widths_dxa):
                    if i < len(row.cells):
                        tcpr = row.cells[i]._tc.get_or_add_tcPr()
                        tcw = tcpr.find(qn("w:tcW"))
                        if tcw is None:
                            tcw = OxmlElement("w:tcW")
                            tcpr.append(tcw)
                        tcw.set(qn("w:w"), str(w))
                        tcw.set(qn("w:type"), "dxa")

        def add_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for r_idx, row in enumerate(grid):
                cells = table.add_row().cells
                for c_idx in range(cols):
                    set_cell(cells[c_idx], row[c_idx] if c_idx < len(row) else "", bold=(r_idx == 0))
            # SCI 清单表：固定列宽，SCI名字/类型不换行、存储地点较宽
            first = str(grid[0][0]).strip() if grid and grid[0] else ""
            if first == "SCI名字":
                if cols == 5:
                    _set_fixed_widths(table, [1700, 1100, 2000, 1500, 3000])
                elif cols == 4:
                    _set_fixed_widths(table, [1700, 2200, 1500, 3600])
            elif first == "类别" and cols >= 3:
                # 配置项存储地址表：类别列收窄、存储路径较宽；首列连续相同「类别」纵向合并；文字左对齐
                if cols == 3:
                    _set_fixed_widths(table, [1300, 1300, 5000])
                rows = table.rows
                n = len(rows)
                r = 1
                while r < n:
                    val = str((grid[r][0] if r < len(grid) and grid[r] else "") or "").strip()
                    if not val:
                        r += 1
                        continue
                    r2 = r
                    while r2 + 1 < n and str((grid[r2 + 1][0] if grid[r2 + 1] else "") or "").strip() == val:
                        r2 += 1
                    if r2 > r:
                        merged = rows[r].cells[0].merge(rows[r2].cells[0])
                        set_cell(merged, val, align=WD_ALIGN_PARAGRAPH.LEFT)
                        merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    r = r2 + 1
            elif first in ("测试目标", "测试方法和技术") and cols == 2:
                _set_fixed_widths(table, [2200, 7100])
            elif first == "配置":
                if cols == 2:
                    _set_fixed_widths(table, [1900, 7400])
                elif cols == 3:
                    _set_fixed_widths(table, [1900, 3700, 3700])
            document.add_paragraph()

        def add_cover_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for row in grid:
                cells = table.add_row().cells
                for c_idx in range(cols):
                    text = row[c_idx] if c_idx < len(row) else ""
                    set_cell(cells[c_idx], text, bold=(c_idx % 2 == 0), align=WD_ALIGN_PARAGRAPH.CENTER)
                if (str(row[0]).strip() if row else "") == "生效日期" and cols > 2:
                    merged = cells[1]
                    for c_idx in range(2, cols):
                        merged = merged.merge(cells[c_idx])
                    set_cell(merged, row[1] if len(row) > 1 else "", align=WD_ALIGN_PARAGRAPH.CENTER)
            document.add_paragraph()

        def strip_num(title):
            return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()

        def add_body_heading(title, level):
            size = {1: 16.0, 2: 14.0, 3: 12.0}.get(level, 11.0)
            p = document.add_heading("", level=max(1, min(level, 9)))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, title, font_size=size, bold=True)

        def render_body_section(node, level, number=""):
            name = strip_num(node.get("title"))
            heading = f"{number} {name}".strip() if number else name
            add_body_heading(heading, level=max(1, min(level, 9)))
            if (node.get("body") or "").strip():
                add_text(node.get("body"))
            if node.get("ref_type") == "review":
                for t_idx, table in enumerate(node.get("tables") or []):
                    serv_review_util.render_review_grid(document, table, set_cell, merge_col0=(t_idx == 0), merge_full=True)
            else:
                caps = node.get("table_captions") or []
                for t_idx, table in enumerate(node.get("tables") or []):
                    cap = caps[t_idx] if t_idx < len(caps) else ""
                    if str(cap or "").strip():
                        add_text(cap)
                    add_grid(table)
                if str(node.get("body_tail") or "").strip():
                    add_text(node.get("body_tail"))
            idx = 0
            for child in (node.get("children") or []):
                idx += 1
                child_num = f"{number}.{idx}" if number else f"{idx}"
                render_body_section(child, level + 1, child_num)

        cover = next((s for s in sections if s.get("ref_type") == "cover"), None)
        revision = next((s for s in sections if s.get("ref_type") == "revision"), None)
        body = [s for s in sections if s.get("ref_type") not in ("cover", "revision")]

        add_blank_lines(6)
        write_center_title((strip_num(cover.get("title")) if cover else "") or DOC_NAME, size=22.0, bold=True)
        add_blank_lines(4)
        if cover:
            for table in (cover.get("tables") or []):
                add_cover_grid(table)

        document.add_page_break()
        write_center_title("文件修订记录", size=14.0, bold=True)
        add_blank_lines(2)
        if revision:
            for table in (revision.get("tables") or []):
                add_grid(table)

        document.add_page_break()
        write_center_title("目录", size=16.0, bold=True)
        docx_util.insert_toc_field(document)

        document.add_page_break()
        seq = 0
        for node in body:
            if node.get("ref_type") == "review":
                render_body_section(node, 1, "")
            else:
                seq += 1
                render_body_section(node, 1, str(seq))

        document.save(output)
        output.seek(0)
