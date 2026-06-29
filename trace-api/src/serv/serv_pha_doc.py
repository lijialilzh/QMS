#!/usr/bin/env python
# encoding: utf-8

# 初步危害分析清单服务层。整份文档以 content(JSON) 存储；导出复用 docx_util.fonted_txt 生成 Word。
# 自动获取：
#   1) 全文产品名称（模板内置基准名 BASE_NAME，按所选产品名称全文替换）。
#   2) A.2/A.3/A.4(CFMEA/DFMEA/PFMEA) 表按「危害编号」从产品HAZ管理回填
#      潜在故障模式/故障的潜在原因/失效的潜在影响/分类。
#   3) 封面/修订记录日期按产品时间逻辑线更新。

import copy
import json
import logging
import os
import re
from typing import List
from sqlalchemy import delete, func, select

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.prod_dhf import ProdDhf
from ..model.haz import Haz
from ..model.prod_haz import ProdHaz
from ..model.pha_doc import PhaDoc
from ..model.project_timeline import ProjectTimelineRow, ProjectTimelineCell
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_pha_doc import PhaDocForm
from ..obj.vobj_pha_doc import PhaDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from .serv_utils import new_version
from .serv_utils import docx_util

logger = logging.getLogger(__name__)

# 本文档名（用于从产品 DHF 匹配文件编号）
DOC_NAME = "初步危害分析清单"
# 模板内置的基准产品名称，导出/编辑时按所选产品名称全文替换
BASE_NAME = "肿瘤CT图像随访与评估软件"
# 封面/修订日期从时间逻辑线匹配的关键字（按顺序取命中行里的最新日期）
DATE_KEYWORDS = ["初步危害分析", "危害分析", "风险管理"]

_TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "pha_template.json")


def _load_default():
    try:
        with open(_TEMPLATE_PATH, encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        logger.exception("load pha_template.json failed")
        return {"sections": []}


DEFAULT_PHA_CONTENT = _load_default()


class Server(object):

    def __normalize_node(self, node):
        if not isinstance(node, dict):
            return {"title": str(node or ""), "body": "", "tables": [], "children": []}
        result = dict(node)
        result["title"] = str(result.get("title") or "")
        result["body"] = str(result.get("body") or "")
        tables = result.get("tables")
        if not isinstance(tables, list):
            tables = []
        norm_tables = []
        for table in tables:
            if isinstance(table, list):
                norm_tables.append([[str(c) if c is not None else "" for c in (row or [])] for row in table if isinstance(row, list)])
        result["tables"] = norm_tables
        children = result.get("children")
        result["children"] = [self.__normalize_node(c) for c in children] if isinstance(children, list) else []
        return result

    def __normalize_content(self, content):
        if not isinstance(content, dict) or not isinstance(content.get("sections"), list):
            return copy.deepcopy(DEFAULT_PHA_CONTENT)
        return {"sections": [self.__normalize_node(s) for s in content["sections"]]}

    @staticmethod
    def __strip_num(title):
        return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()

    # ---------------- 自动获取 ----------------
    def __collect_autofill(self, prod_id, product, doc_version):
        name = (getattr(product, "name", "") or "").strip()

        def to_int(v):
            digits = re.sub(r"[^\d]", "", str(v or ""))
            return int(digits) if digits else None

        # 时间逻辑线日期
        tl_rows = db.session.execute(
            select(ProjectTimelineRow).where(ProjectTimelineRow.prod_id == prod_id)
        ).scalars().all()
        cell_map = {}
        if tl_rows:
            for c in db.session.execute(
                select(ProjectTimelineCell).where(ProjectTimelineCell.row_id.in_([r.id for r in tl_rows]))
            ).scalars().all():
                cell_map.setdefault(c.row_id, []).append(c.output_result or "")
        date_rows = [r for r in tl_rows if (r.row_type or "date") == "date" and to_int(r.year) and to_int(r.month)]

        def date_key(r):
            return to_int(r.year) * 10000 + to_int(r.month) * 100 + (to_int(r.day) or 0)

        def latest_date(keywords):
            rows = [r for r in date_rows if any(any(k in str(v or "") for k in keywords) for v in cell_map.get(r.id, []))]
            if not rows:
                return ""
            r = max(rows, key=date_key)
            return f"{to_int(r.year)}年{to_int(r.month)}月{to_int(r.day)}日"

        cover_date = latest_date(DATE_KEYWORDS)

        # 产品HAZ管理：code -> 内容
        haz_map = {}
        rows = db.session.execute(
            select(ProdHaz, Haz).outerjoin(Haz, ProdHaz.haz_id == Haz.id).where(ProdHaz.prod_id == prod_id)
        ).all()
        for ph, hz in rows:
            if not hz or not hz.code:
                continue
            code = str(hz.code).strip().upper()
            haz_map[code] = {
                "event": (hz.event or "").strip(),
                "situation": ((ph.situation if ph else None) or hz.situation or "").strip(),
                "damage": ((ph.damage if ph else None) or hz.damage or "").strip(),
                "category": (hz.category or "").strip(),
            }

        return {
            "name": name,
            "cover_date": cover_date,
            "version": doc_version,
            "haz_map": haz_map,
        }

    @staticmethod
    def __replace_name(node, old, new):
        if not old or not new or old == new:
            return
        node["title"] = str(node.get("title") or "").replace(old, new)
        node["body"] = str(node.get("body") or "").replace(old, new)
        for tbl in (node.get("tables") or []):
            for row in tbl:
                if not isinstance(row, list):
                    continue
                for i in range(len(row)):
                    if isinstance(row[i], str):
                        row[i] = row[i].replace(old, new)
        for ch in (node.get("children") or []):
            Server.__replace_name(ch, old, new)

    @staticmethod
    def __fill_fmea(tbl, haz_map):
        if not isinstance(tbl, list) or not tbl:
            return
        # 定位表头行（含「危害编号」与「潜在故障模式」）
        hidx = -1
        for i, row in enumerate(tbl):
            if isinstance(row, list) and any("危害编号" in str(c) for c in row) and any("潜在故障模式" in str(c) for c in row):
                hidx = i
                break
        if hidx < 0:
            return
        header = tbl[hidx]

        def col_of(keyword):
            for ci, c in enumerate(header):
                if keyword in str(c):
                    return ci
            return None

        code_idx = col_of("危害编号")
        field_idx = {
            "event": col_of("潜在故障模式"),
            "situation": col_of("故障的潜在原因"),
            "damage": col_of("失效的潜在影响"),
            "category": col_of("分类"),
        }
        if code_idx is None:
            return
        for row in tbl[hidx + 1:]:
            if not isinstance(row, list) or code_idx >= len(row):
                continue
            m = re.search(r"HAZ\d+", str(row[code_idx]).upper())
            if not m:
                continue
            info = haz_map.get(m.group(0))
            if not info:
                continue
            for key, ci in field_idx.items():
                if ci is None or ci >= len(row):
                    continue
                val = info.get(key) or ""
                if val:
                    row[ci] = val

    def __fill_node(self, node, info):
        ref = node.get("ref_type")
        title = self.__strip_num(node.get("title"))
        if ref == "cover" or title == DOC_NAME:
            for tbl in (node.get("tables") or []):
                for row in tbl:
                    if not isinstance(row, list):
                        continue
                    if row and str(row[0]).strip() == "生效日期":
                        continue  # 生效日期不自动回填
                    for ci in range(len(row)):
                        if str(row[ci]).strip() == "日期" and ci + 1 < len(row):
                            row[ci + 1] = info["cover_date"]
        if ref == "revision" or title == "文件修订记录":
            tables = node.get("tables") or []
            if tables and isinstance(tables[0], list):
                t = tables[0]
                cols = len(t[0]) if t and t[0] else 5
                while len(t) < 2:
                    t.append([""] * cols)
                row = t[1]
                while len(row) < 5:
                    row.append("")
                row[0] = info["cover_date"] or ""
                if info["version"]:
                    row[1] = info["version"]
                if not str(row[2] or "").strip():
                    row[2] = "首次发布"
        if ref == "pha_fmea" or any(k in title for k in ("CFMEA", "DFMEA", "PFMEA")):
            for tbl in (node.get("tables") or []):
                self.__fill_fmea(tbl, info["haz_map"])
        for child in (node.get("children") or []):
            self.__fill_node(child, info)

    def __autofill_for_export(self, content, obj: PhaDocObj):
        sections = (content or {}).get("sections") or []
        if not obj.product_id:
            return content
        product = db.session.execute(select(Product).where(Product.id == obj.product_id)).scalars().first()
        info = self.__collect_autofill(obj.product_id, product, obj.version)
        for node in sections:
            self.__replace_name(node, BASE_NAME, info["name"])
            self.__fill_node(node, info)
        return content

    def __dhf_file_no(self, prod_id):
        row = db.session.execute(
            select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name.like(f"%{DOC_NAME}%"))
        ).scalars().first()
        return (row.code or "").strip() if row and row.code else ""

    def __to_obj(self, row: PhaDoc, product: Product = None):
        obj = PhaDocObj(**row.dict())
        obj.content = self.__normalize_content(obj.content)
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
            dhf_no = self.__dhf_file_no(product.id)
            if dhf_no:
                obj.file_no = dhf_no
        return obj

    # ---------------- CRUD ----------------
    async def add_pha_doc(self, form: PhaDocForm):
        try:
            sql = select(func.count(PhaDoc.id)).where(
                PhaDoc.product_id == form.product_id,
                PhaDoc.version == form.version,
            )
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = PhaDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=PhaDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_pha_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: PhaDoc = db.session.execute(select(PhaDoc).where(PhaDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(PhaDoc.version).where(PhaDoc.product_id == target_pid)).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = PhaDoc(
                product_id=target_pid,
                version=version,
                file_no=fromdoc.file_no,
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            return Resp.resp_ok(data=PhaDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_pha_doc(self, form: PhaDocForm):
        try:
            row: PhaDoc = db.session.execute(select(PhaDoc).where(PhaDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_pha_doc(self, id: int):
        db.session.execute(delete(PhaDoc).where(PhaDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_pha_doc(self, id: int):
        sql = select(PhaDoc, Product).join(Product, PhaDoc.product_id == Product.id).where(PhaDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_pha_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(PhaDoc.product_id == product_id)
        if version:
            wheres.append(PhaDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(PhaDoc.id)).join(Product, PhaDoc.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (
            select(PhaDoc, Product)
            .join(Product, PhaDoc.product_id == Product.id)
            .where(*wheres)
            .order_by(PhaDoc.id.desc())
            .offset(page_index * page_size)
            .limit(page_size)
        )
        rows: List[PhaDocObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    # ---------------- 导出 Word ----------------
    async def export_pha_doc(self, output, id: int):
        resp = await self.get_pha_doc(id)
        obj: PhaDocObj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        c = self.__autofill_for_export(self.__normalize_content(obj.content), obj)
        sections = c.get("sections") or []
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")

        def write_center_title(text, size=22.0, bold=False):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, text, font_size=size, bold=bold)

        def add_blank_lines(count):
            for _ in range(max(0, int(count or 0))):
                document.add_paragraph("")

        def add_text(text):
            docx_util.save_txt2docx(str(text or ""), document)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            cell.text = ""
            lines = str(text or "").split("\n")
            for i, line in enumerate(lines):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER if align == WD_ALIGN_PARAGRAPH.CENTER else WD_CELL_VERTICAL_ALIGNMENT.TOP

        def merge_col_runs(table, col):
            rows = table.rows
            n = len(rows)
            r = 0
            while r < n:
                text = rows[r].cells[col].text
                r2 = r
                while r2 + 1 < n and rows[r2 + 1].cells[col].text == text:
                    r2 += 1
                if r2 > r and text.strip():
                    merged = rows[r].cells[col].merge(rows[r2].cells[col])
                    set_cell(merged, text, align=WD_ALIGN_PARAGRAPH.CENTER)
                    merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                r = r2 + 1

        def merge_row_full(table, r):
            cells = table.rows[r].cells
            if len(cells) < 2:
                return
            text = cells[0].text
            if text.strip() and all(c.text == text for c in cells):
                merged = cells[0].merge(cells[len(cells) - 1])
                set_cell(merged, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        def add_grid(grid, merge_col0=False, merge_banner=False, header_rows=1):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for r_idx, row in enumerate(grid):
                cells = table.add_row().cells
                for c_idx in range(cols):
                    set_cell(cells[c_idx], row[c_idx] if c_idx < len(row) else "", bold=(r_idx < header_rows))
            if merge_banner and len(table.rows) > 0:
                merge_row_full(table, 0)
            if merge_col0:
                merge_col_runs(table, 0)
            document.add_paragraph()

        def add_cover_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for row in grid:
                cells = table.add_row().cells
                for c_idx in range(cols):
                    text = row[c_idx] if c_idx < len(row) else ""
                    set_cell(cells[c_idx], text, bold=(c_idx % 2 == 0), align=WD_ALIGN_PARAGRAPH.CENTER)
                if (str(row[0]).strip() if row else "") == "生效日期" and cols > 2:
                    merged = cells[1]
                    for c_idx in range(2, cols):
                        merged = merged.merge(cells[c_idx])
                    set_cell(merged, row[1] if len(row) > 1 else "", align=WD_ALIGN_PARAGRAPH.CENTER)
            document.add_paragraph()

        def add_body_heading(title, level):
            size = {1: 16.0, 2: 14.0, 3: 12.0}.get(level, 11.0)
            p = document.add_heading("", level=max(1, min(level, 9)))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, title, font_size=size, bold=True)

        def render_body_section(node, level, number=""):
            name = self.__strip_num(node.get("title"))
            heading = f"{number} {name}".strip() if number else name
            add_body_heading(heading, level=max(1, min(level, 9)))
            if (node.get("body") or "").strip():
                add_text(node.get("body"))
            is_fmea = node.get("ref_type") == "pha_fmea" or any(k in name for k in ("CFMEA", "DFMEA", "PFMEA"))
            for table in (node.get("tables") or []):
                if is_fmea:
                    add_grid(table, merge_col0=True, merge_banner=True, header_rows=2)
                else:
                    add_grid(table)
            idx = 0
            for child in (node.get("children") or []):
                idx += 1
                child_num = f"{number}.{idx}" if number else f"{idx}"
                render_body_section(child, level + 1, child_num)

        cover = next((s for s in sections if s.get("ref_type") == "cover"), None)
        revision = next((s for s in sections if s.get("ref_type") == "revision"), None)
        body = [s for s in sections if s.get("ref_type") not in ("cover", "revision")]

        add_blank_lines(6)
        write_center_title((self.__strip_num(cover.get("title")) if cover else "") or DOC_NAME, size=22.0, bold=True)
        add_blank_lines(4)
        if cover:
            for table in (cover.get("tables") or []):
                add_cover_grid(table)

        document.add_page_break()
        write_center_title("文件修订记录", size=14.0, bold=True)
        add_blank_lines(2)
        if revision:
            for table in (revision.get("tables") or []):
                add_grid(table)

        document.add_page_break()
        for i, node in enumerate(body):
            render_body_section(node, 1, str(i + 1))

        document.save(output)
        output.seek(0)
