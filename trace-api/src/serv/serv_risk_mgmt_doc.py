#!/usr/bin/env python
# encoding: utf-8

import logging
import base64
import copy
import io
import os
import re
from typing import List
from sqlalchemy import delete, func, select
from docx import Document
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.haz import Haz
from ..model.prod_haz import ProdHaz
from ..model.risk_mgmt_doc import RiskAnalysis, RiskControl, RiskMgmtDoc, RiskParticipant
from ..model.project_member import ProjectMember
from ..model.prod_dhf import ProdDhf
from ..model.ptr_doc import PtrDoc
from ..model.pha_doc import PhaDoc
from ..model.project_timeline import ProjectTimelineRow
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_risk_mgmt_doc import RiskAnalysisForm, RiskControlForm, RiskMgmtDocForm, RiskParticipantForm
from ..obj.vobj_risk_mgmt_doc import RiskAnalysisObj, RiskControlObj, RiskMgmtDocObj, RiskParticipantObj
from ..utils.i18n import ts
from ..utils import get_uuid
from ..utils.sql_ctx import db
from . import msg_err_db
from .serv_utils import new_version, sync_file_no_version
from .serv_utils import docx_util
from . import serv_review_util

logger = logging.getLogger(__name__)


RISK_ACCEPTANCE_TABLE = [
    ["风险值", "", "", "严重度", "", "", "", ""],
    ["", "", "", "可忽略 A", "轻度 B", "严重 C", "危重的 D", "灾难性的 E"],
    ["发生概率", "经常", "5", "5A", "5B", "5C", "5D", "5E"],
    ["", "有时", "4", "4A", "4B", "4C", "4D", "4E"],
    ["", "偶然", "3", "3A", "3B", "3C", "3D", "3E"],
    ["", "很少", "2", "2A", "2B", "2C", "2D", "2E"],
    ["", "非常少", "1", "1A", "1B", "1C", "1D", "1E"],
    ["红色", "不可接受：这类风险本质上不可接受。必须寻求风险降低措施。", "", "", "", "", "", ""],
    ["橙色", "进一步降低的研究：这类风险必须降低到合理可行的最低限度才可视为可接受。", "", "", "", "", "", ""],
    ["绿色", "可忽略：这类风险实际上可接受，但只可挑选一步寻求风险降低措施。", "", "", "", "", "", ""],
]


DEFAULT_RISK_CONTENT = {
    "sections": [
        {"title": "风险管理报告", "ref_type": "cover", "children": [], "tables": [[
            ["编制部门", "", "文件版本", ""],
            ["编制人", "", "日期", ""],
            ["审核人", "", "日期", ""],
            ["批准人", "", "日期", ""],
            ["生效日期", "", "", ""],
        ]]},
        {"title": "文件修订记录", "ref_type": "revision", "children": [], "tables": [[
            ["修改日期", "版本号", "修订说明", "修订人", "批准人"],
            ["", "", "", "", ""],
        ]]},
        {"title": "1 目的", "children": []},
        {"title": "2 范围", "children": []},
        {"title": "3 产品描述", "children": [
            {"title": "3.1 产品预期用途", "children": []},
            {"title": "3.2 产品功能描述", "children": [
                {"title": "3.2.1 产品功能明细", "ref_type": "prod_func_detail", "children": []},
            ]},
        ]},
        {"title": "4 评审", "children": [
            {"title": "4.1 评审数据", "children": []},
            {"title": "4.2 风险分析参与人员", "ref_type": "participants", "children": []},
            {"title": "4.3 审评历史", "children": []},
        ]},
        {"title": "5 风险分析方式", "children": [
            {"title": "5.1 危害识别", "children": [
                {"title": "5.1.1 与合理可预见相关的环境相关的危害", "children": []},
                {"title": "5.1.2 考虑的危害包括", "children": []},
                {"title": "5.1.3 危害初步原因的考虑应包括", "children": []},
                {"title": "5.1.4 危害重点考虑的原因应包括", "children": []},
            ]},
            {"title": "5.2 风险评价准则", "children": [
                {"title": "5.2.1 严重度定义", "children": []},
                {"title": "5.2.2 发生概率定义", "children": []},
                {"title": "5.2.3 接受标准", "ref_type": "acceptance_standard", "children": []},
            ]},
        ]},
        {"title": "6 风险分析", "children": [
            {"title": "6.1 与安全有关特征的问题识别", "children": []},
            {"title": "6.2 已知或可预见的危险（源）识别", "children": []},
            {"title": "6.3 估计每个危险情况的风险", "children": []},
            {"title": "6.4 风险评价", "ref_type": "risk_analysis", "children": []},
            {"title": "6.5 风险控制", "ref_type": "risk_controls", "children": [
                {"title": "6.5.1 风险控制方案分析", "children": []},
                {"title": "6.5.2 风险控制措施的实施", "children": []},
                {"title": "6.5.3 剩余风险分析和风险/受益分析", "children": []},
                {"title": "6.5.4 由风险控制措施产生的风险", "children": [], "tables": [[
                    ["RCM编号", "引入的危害", "RCM引入的风险分析", "风险控制措施"],
                ]]},
            ]},
        ]},
        {"title": "7 风险的可接受性评价", "children": [
            {"title": "7.1 RCMs实施风险控制措施前/后的风险分布", "children": []},
            {"title": "7.2 综合剩余风险评价", "children": []},
            {"title": "7.3 软件安全级别判定", "children": []},
        ]},
        {"title": "8 生产和生产后活动", "children": []},
        {"title": "9 结论", "children": []},
        {"title": "10 参考标准", "children": []},
        {"title": "11 风险管理文件", "children": [], "tables": [[
            ["编号", "描述"],
            ["", "风险管理计划"],
            ["", "初步危害分析清单"],
            ["", "风险管理报告"],
            ["", "自研软件网络安全研究报告"],
            ["", "网络安全扫描报告"],
        ]]},
        {"title": "附录A 与安全有关特征的问题识别", "children": []},
        {"title": "附录B 风险分析矩阵", "ref_type": "risk_analysis", "children": []},
    ],
    "participants": [],
    "riskMatrix": [],
    "riskControls": [],
}


class Server(object):
    def __normalize_title_text(self, value):
        return re.sub(r"\s+", "", str(value or ""))

    def __is_cover_section(self, section):
        title = self.__normalize_title_text((section or {}).get("title", ""))
        return (section or {}).get("ref_type") == "cover" or title == "风险管理报告"

    def __is_revision_section(self, section):
        title = self.__normalize_title_text((section or {}).get("title", ""))
        return (section or {}).get("ref_type") == "revision" or title == "文件修订记录"

    def __normalize_cover_section(self, section):
        next_section = copy.deepcopy(section or copy.deepcopy(DEFAULT_RISK_CONTENT["sections"][0]))
        tables = next_section.get("tables") if isinstance(next_section.get("tables"), list) else []
        table = tables[0] if tables else []
        first_row = table[0] if table else []
        first_text = self.__normalize_title_text("".join([str(v or "") for v in first_row]))
        if "编制部门" in first_text and "文件版本" in first_text and len(table or []) >= 4:
            return next_section
        if "编制科室" in first_text or "编制部门" in first_text or "文件版本" in first_text:
            headers = table[0] if table else []
            values = table[1] if len(table or []) > 1 else []
            def get_value(label):
                for idx, header in enumerate(headers or []):
                    if label in self.__normalize_title_text(header):
                        return values[idx] if idx < len(values or []) else ""
                return ""
            next_section["tables"] = [[
                ["编制部门", get_value("编制") or "", "文件版本", get_value("文件版本") or ""],
                ["编制人", get_value("编制人") or "", "日期", ""],
                ["审核人", get_value("审核人") or "", "日期", ""],
                ["批准人", get_value("批准人") or "", "日期", ""],
                ["生效日期", get_value("生效日期") or "", "", ""],
            ]]
        return next_section

    def __ensure_front_matter_sections(self, content):
        next_content = copy.deepcopy(content or {})
        sections = next_content.get("sections") if isinstance(next_content.get("sections"), list) else []
        default_sections = copy.deepcopy(DEFAULT_RISK_CONTENT["sections"])
        default_cover = default_sections[0]
        default_revision = default_sections[1]
        cover = self.__normalize_cover_section(next((item for item in sections if self.__is_cover_section(item)), None) or default_cover)
        revision = next((item for item in sections if self.__is_revision_section(item)), None) or default_revision
        body_sections = [item for item in sections if not self.__is_cover_section(item) and not self.__is_revision_section(item)]
        next_content["sections"] = [cover, revision] + body_sections
        return next_content

    def __extract_participants_from_sections(self, sections):
        result = []

        def walk(items):
            nonlocal result
            for item in items or []:
                title = self.__normalize_title_text((item or {}).get("title", ""))
                if (item or {}).get("ref_type") == "participants" or "风险分析参与人员" in title:
                    for table_rows in (item or {}).get("tables", []) or []:
                        data_rows = table_rows[1:] if len(table_rows or []) > 1 else []
                        for row in data_rows:
                            role = str(row[0] if len(row or []) > 0 else "").strip()
                            name = str(row[1] if len(row or []) > 1 else "").strip()
                            if role or name:
                                result.append({"role": role, "name": name})
                        if result:
                            return
                walk((item or {}).get("children") or [])

        walk(sections or [])
        return result

    def __normalize_content(self, content):
        result = copy.deepcopy(DEFAULT_RISK_CONTENT)
        if isinstance(content, dict):
            result.update(content)
        result.setdefault("sections", copy.deepcopy(DEFAULT_RISK_CONTENT["sections"]))
        result.setdefault("participants", [])
        if not result.get("participants"):
            result["participants"] = self.__extract_participants_from_sections(result.get("sections") or [])
        result.setdefault("riskMatrix", [])
        result.setdefault("riskControls", [])
        return self.__ensure_front_matter_sections(result)

    def __member_name(self, prod_id, keywords):
        # 按项目角色关键字匹配项目人员姓名（按优先级取首个命中）
        for kw in keywords:
            member = db.session.execute(
                select(ProjectMember).where(ProjectMember.prod_id == prod_id, ProjectMember.role.like(f"%{kw}%"))
                .order_by(ProjectMember.sort_order.asc(), ProjectMember.id.asc())
            ).scalars().first()
            if member and (member.name or "").strip():
                return member.name.strip()
        return ""

    def __autofill_front_matter(self, content, product_id, version):
        # 封面、文件修订记录自动获取（仅填空，不覆盖已填内容），与「风险管理计划」保持一致
        if not product_id:
            return content
        reviser = self.__member_name(product_id, ("产品经理", "项目经理"))
        auditor = self.__member_name(product_id, ("QA负责人", "质量负责人", "QA"))
        approver = self.__member_name(product_id, ("产品负责人", "研发负责人", "管理者代表"))
        rev_date = serv_review_util.review_date(product_id, serv_review_util.REVIEW_DEFS["risk"]["name_keywords"]) or ""
        ver = str(version or "")
        for section in (content.get("sections") or []):
            if self.__is_cover_section(section):
                for table in (section.get("tables") or []):
                    for row in table:
                        if not isinstance(row, list) or not row:
                            continue
                        label = str(row[0]).strip()
                        if len(row) >= 4 and str(row[2]).strip() == "文件版本" and ver and not str(row[3] or "").strip():
                            row[3] = ver

                        def set_name(val):
                            if val and len(row) >= 2 and not str(row[1] or "").strip():
                                row[1] = val

                        def set_date(val):
                            if val and len(row) >= 4 and not str(row[3] or "").strip():
                                row[3] = val
                        # 编制/审核/批准人「姓名」列统一由签名规则填充（见文末 fill_cover_signers），此处仅填日期
                        if label in ("编制部门", "编写部门"):
                            set_name("产品部")
                        elif label in ("编制人", "审核人", "批准人"):
                            set_date(rev_date)
                        elif label == "生效日期":
                            set_name(rev_date)
            elif self.__is_revision_section(section):
                for table in (section.get("tables") or []):
                    if not isinstance(table, list) or not table:
                        continue
                    if len(table) >= 2 and isinstance(table[1], list) and len(table[1]) >= 5:
                        row = table[1]

                        def set_if(i, val):
                            if val and not str(row[i] if i < len(row) else "").strip():
                                row[i] = val
                        set_if(0, rev_date)
                        set_if(1, ver)
                        if not str(row[2] or "").strip():
                            row[2] = "首次发布"
                        set_if(3, reviser)
                        set_if(4, approver)
                    # 默认五行：补足到「表头 + 5 行」
                    col = len(table[0]) if isinstance(table[0], list) and table[0] else 5
                    while len(table) < 6:
                        table.append(["" for _ in range(col)])
        # 封面「编制/审核/批准人」按部门签名规则填充签名图（仅填空）
        serv_review_util.fill_cover_signers(content, serv_review_util.cover_signers(product_id, "risk"))
        return content

    def __dhf_code(self, prod_id, keyword):
        # 从产品 DHF 按名称关键字匹配「编号」
        row = db.session.execute(
            select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name.like(f"%{keyword}%"))
            .order_by(ProdDhf.id.asc())
        ).scalars().first()
        return (row.code or "").strip() if row and row.code else ""

    def __fill_risk_mgmt_files(self, content, product_id):
        # 「11 风险管理文件」表格：无表则补默认描述，编号列按描述从 DHF 自动获取（仅填空）
        for section in (content.get("sections") or []):
            if not self.__is_risk_mgmt_files_section(section):
                continue
            tables = section.get("tables") or []
            target = None
            for table in tables:
                if table and isinstance(table[0], list) \
                        and any("编号" in str(c) for c in table[0]) and any("描述" in str(c) for c in table[0]):
                    target = table
                    break
            if target is None:
                target = [
                    ["编号", "描述"],
                    ["", "风险管理计划"],
                    ["", "初步危害分析清单"],
                    ["", "风险管理报告"],
                    ["", "自研软件网络安全研究报告"],
                    ["", "网络安全扫描报告"],
                ]
                tables.append(target)
                section["tables"] = tables
            if product_id:
                for r in target[1:]:
                    if not isinstance(r, list) or len(r) < 2:
                        continue
                    desc = str(r[1] or "").strip()
                    if desc and not str(r[0] or "").strip():
                        code = self.__dhf_code(product_id, desc)
                        if code:
                            r[0] = code
            return content
        return content

    def __is_risk_mgmt_files_section(self, section):
        title = re.sub(r"\s+", "", str((section or {}).get("title") or ""))
        return title.startswith("11") and "风险管理文件" in title

    # ---------------- 正文章节自动获取（仅填空，不覆盖已填内容） ----------------
    @staticmethod
    def __to_int(value):
        m = re.search(r"\d+", str(value or ""))
        return int(m.group()) if m else 0

    @staticmethod
    def __strip_no(title):
        return re.sub(r"^[0-9．.、\s]+", "", str(title or "")).strip()

    @staticmethod
    def __level_number(depth, idx):
        if depth <= 1:
            return f"({idx})"
        if depth == 2:
            return f"{idx})"
        circled = "①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳"
        return circled[idx - 1] if 1 <= idx <= len(circled) else f"{idx}."

    def __node_outline(self, node, depth, lines, number=""):
        title = self.__strip_no(node.get("title"))
        body = str(node.get("body") or "").strip()
        if depth >= 1 and title:
            lines.append(f"{number} {title}" if number else title)
        if body:
            lines.append(body)
        idx = 0
        for child in node.get("children") or []:
            idx += 1
            self.__node_outline(child, depth + 1, lines, self.__level_number(depth + 1, idx))

    def __latest_source_doc(self, model, product_id):
        return db.session.execute(
            select(model)
            .where(model.product_id == product_id, ~model.version.like("__deleted%"))
            .order_by(model.id.desc())
        ).scalars().first()

    def __ptr_func_2_1(self, product_id):
        # 技术要求「2.1 性能指标→功能」完整内容（含各模块及功能点）
        doc = self.__latest_source_doc(PtrDoc, product_id)
        if not doc or not isinstance(doc.content, dict):
            return ""
        sections = doc.content.get("sections") or []
        body = [s for s in sections if s.get("ref_type") not in ("cover", "appendix")]
        target = next((s for s in body if self.__strip_no(s.get("title")) == "性能指标"), None)
        if target is None and len(body) >= 2:
            target = body[1]
        if not target:
            return ""
        children = target.get("children") or []
        child = next((c for c in children if self.__strip_no(c.get("title")) == "功能"), None) or (children[0] if children else None)
        if not child:
            return ""
        lines = []
        self.__node_outline(child, 0, lines)
        return "\n".join(lines)

    def __pha_appendix_a_table(self, product_id):
        # 附录A：初步危害分析清单「与安全有关特征的问题识别」ISO 问题表
        doc = self.__latest_source_doc(PhaDoc, product_id)
        if not doc or not isinstance(doc.content, dict):
            return None
        sections = doc.content.get("sections") or []
        target = next((s for s in sections if self.__strip_no(s.get("title")) == "与安全有关特征的问题识别"), None)
        if not target:
            return None

        def find_table(node):
            for t in node.get("tables") or []:
                if t and isinstance(t[0], list):
                    head = "".join(str(c) for c in t[0])
                    if "问题" in head and "考虑的内容" in head and "是否适用" in head:
                        return t
            for c in node.get("children") or []:
                found = find_table(c)
                if found:
                    return found
            return None

        return find_table(target)

    def __risk_dist_matrices(self, product_id):
        # 7.1 风险分布：初始/剩余（发生概率×严重度）5×5 数量矩阵
        rows = db.session.execute(
            select(ProdHaz, Haz).outerjoin(Haz, ProdHaz.haz_id == Haz.id)
            .where(ProdHaz.prod_id == product_id)
        ).all() if product_id else []
        # 无 HAZ 数据时也返回全 0 的两张矩阵（而非 None），确保 7.1 始终有表
        rate_to_row = {"5": 0, "4": 1, "3": 2, "2": 3, "1": 4}
        sev_to_col = {"A": 0, "B": 1, "C": 2, "D": 3, "E": 4}

        def build(rate_attr, degree_attr):
            counts = [[0] * 5 for _ in range(5)]
            for prod_haz, haz in rows:
                r = str(getattr(prod_haz, rate_attr, None) or getattr(haz, rate_attr, None) or "").strip()
                d = str(getattr(prod_haz, degree_attr, None) or getattr(haz, degree_attr, None) or "").strip().upper()
                if r in rate_to_row and d in sev_to_col:
                    counts[rate_to_row[r]][sev_to_col[d]] += 1
            return counts

        def grid(label, counts):
            prob_labels = ["经常(5)", "有时(4)", "偶然(3)", "很少(2)", "非常少(1)"]
            g = [[label, "可忽略(A)", "轻度(B)", "严重(C)", "危重的(D)", "灾难性的(E)"]]
            for i, pl in enumerate(prob_labels):
                g.append([pl] + [str(x) for x in counts[i]])
            return g

        return [
            grid("初始风险分布（措施前）", build("init_rate", "init_degree")),
            grid("剩余风险分布（措施后）", build("cur_rate", "cur_degree")),
        ]

    def __rcm_introduced_rows(self, section, product_id):
        # 6.5.4 由风险控制措施产生的风险：读取本节 RCM 表，按 RCM 编号匹配 HAZ 推导「引入的危害」
        tables = section.get("tables") or []
        table = tables[0] if tables and isinstance(tables[0], list) else []
        if not table:
            return []
        first = table[0] if table and isinstance(table[0], list) else []
        head_text = "".join(str(c) for c in (first or []))
        has_header = bool(re.search(r"RCM编号|引入的危害|风险分析|风险控制措施", head_text))
        data_rows = table[1:] if has_header else table

        def find_col(keywords, fallback):
            if has_header:
                for i, h in enumerate(first):
                    norm = re.sub(r"\s+", "", str(h))
                    if any(k in norm for k in keywords):
                        return i
            return fallback

        rcm_col = find_col(["RCM编号", "RCM"], 0)
        measure_col = find_col(["风险控制措施", "控制措施"], 3)
        haz_rows = db.session.execute(
            select(ProdHaz, Haz).outerjoin(Haz, ProdHaz.haz_id == Haz.id)
            .where(ProdHaz.prod_id == product_id)
        ).all() if product_id else []

        def rcm_codes(text):
            return [x.replace(" ", "") for x in re.findall(r"RCM\s*\d+", str(text or "").upper())]

        def matches(rcm):
            res = []
            for prod_haz, haz in haz_rows:
                related = f"{prod_haz.rcms or getattr(haz, 'rcms', '') or ''}\n{prod_haz.deal or getattr(haz, 'deal', '') or ''}"
                if rcm in rcm_codes(related):
                    analysis = prod_haz.situation or getattr(haz, "situation", "") or getattr(haz, "event", "") or getattr(haz, "source", "") or ""
                    res.append((getattr(haz, "code", "") or "", analysis))
            return res

        out = []
        for row in data_rows:
            if not isinstance(row, list):
                continue
            raw = str(row[rcm_col] if rcm_col < len(row) else "")
            search = raw if raw.strip() else str(row[measure_col] if measure_col < len(row) else "")
            codes = rcm_codes(search)
            rcm = codes[0] if codes else ""
            measure = (row[measure_col] if measure_col < len(row) else "") or (row[rcm_col] if rcm_col < len(row) else "")
            ms = matches(rcm) if rcm else []
            if not ms:
                out.append([raw, "未匹配到HAZ" if rcm else "", "", measure])
            else:
                for code, analysis in ms:
                    out.append([raw, code, analysis, measure])
        return out

    def __project_time_range(self, product_id):
        rows = db.session.execute(
            select(ProjectTimelineRow).where(ProjectTimelineRow.prod_id == product_id)
        ).scalars().all()
        rows = [r for r in rows if (r.row_type or "date") == "date" and self.__to_int(r.year) and self.__to_int(r.month)]
        if not rows:
            return ""

        def ym(r):
            return self.__to_int(r.year) * 100 + self.__to_int(r.month)

        s = min(rows, key=ym)
        e = max(rows, key=ym)
        return f"{self.__to_int(s.year)}年{self.__to_int(s.month)}月至{self.__to_int(e.year)}年{self.__to_int(e.month)}月"

    SEVERITY_TABLE = [
        ["严重度", "分类标准", "标记"],
        ["可忽略", "不便或者暂时不适", "A"],
        ["轻度", "导致不要求专业医疗介入的暂时伤害或损伤", "B"],
        ["严重", "导致要求专业医疗介入的暂时伤害或损伤", "C"],
        ["危重的", "导致永久性损伤或危及生命的危害", "D"],
        ["灾难性的", "导致患者死亡", "E"],
    ]
    PROBABILITY_TABLE = [
        ["概率", "事件频次/年/单位产品", "标记"],
        ["经常", "≥10-3", "5"],
        ["有时", "<10-3和≥10-4", "4"],
        ["偶然", "<10-4和≥10-5", "3"],
        ["很少", "<10-5和≥10-6", "2"],
        ["非常少", "＜10-6", "1"],
    ]

    def __body_text_defaults(self, pname, time_range):
        name = pname or "本产品"
        period = time_range or "产品研发周期内"
        return {
            "目的": (
                f"风险管理的目的是确保{name}的危害得到了定义，评估和评价了相关风险，"
                "控制了这些风险和在寿命周期中监控这些控制措施的有效性。"
                "本公司采用的主要方式和程序来自于GB/T 42062、ISO14971和YY/T 1406.1-2016。"
            ),
            "审评历史": (
                "按照评审记录的模板，在风险管理过程中，形成了以下风险相关文件（部分含评审记录）：\n"
                "《风险管理计划》及评审记录\n《初步危害分析清单》及评审记录\n《网络安全漏洞自评报告》\n"
                "《自研软件网络安全研究报告》\n《风险管理报告》及评审记录"
            ),
            "风险分析方式": (
                "根据YY/T 0316、ISO14971和风险管理控制程序，对于每个危害发生概率、危害程度的评估、"
                "综合考虑概率和危害程度的风险等级、风险可接受准则如下所示。"
            ),
            "危害识别": (
                "与合理可预见相关的环境相关的危害：\n正常使用\n不正确的使用\n人为恶意使用\n"
                "考虑的危害包括：\n对患者的危害\n对操作者的危害\n对信息资产的危害\n"
                "危害初步原因的考虑应包括:\n用户界面\n患者或者临床用户的忽视\n人因工程\n硬件故障\n软件故障\n集成错误\n环境条件\n网络安全\n"
                "危害重点考虑的原因应包括：\n网络工具；\n系统部件的集成，包括硬件和软件；\n用户界面，包括命令语言，警告和错误信息；\n"
                "在用户界面和用户手册中文字翻译的准确性；\n用户预期或非预期情况下数据的保护；\n第三方软件。"
            ),
            "与合理可预见相关的环境相关的危害": "与合理可预见相关的环境相关的危害：\n正常使用\n不正确的使用\n人为恶意使用",
            "考虑的危害包括": "考虑的危害包括：\n对患者的危害\n对操作者的危害\n对信息资产的危害",
            "危害初步原因的考虑应包括": (
                "危害初步原因的考虑应包括:\n用户界面\n患者或者临床用户的忽视\n人因工程\n硬件故障\n软件故障\n集成错误\n环境条件\n网络安全"
            ),
            "危害重点考虑的原因应包括": (
                "危害重点考虑的原因应包括：\n网络工具；\n系统部件的集成，包括硬件和软件；\n用户界面，包括命令语言，警告和错误信息；\n"
                "在用户界面和用户手册中文字翻译的准确性；\n用户预期或非预期情况下数据的保护；\n第三方软件。"
            ),
            "风险分析": (
                f"根据YY/T 0316、ISO14971和风险管理控制程序，{name}的风险分析过程应该定义可能的危险（源），"
                "评估每个危险情况，评估每个风险的可接受程度，降低风险的方式和评审由于采取风险控制措施带来的风险。"
                "在所有这些风险已经被分析后，这些程序和结果的记录见本报告。"
            ),
            "生产和生产后活动": (
                "在风险管理计划中，已经描述了生产和生产后信息收集的方式。\n"
                "通过对执行这些过程中记录的评审，来评审是否引入了风险和开始一个新的风险分析和管理过程。\n"
                "截至目前搜集到的所有信息，没有新的风险产生。"
            ),
            "参考标准": (
                "YY/T 0664-2020 医疗器械软件 软件生存周期过程\n"
                "GB/T 42062-2022 医疗器械 风险管理对医疗器械的应用\n"
                "YY/T 1406.1-2016 医疗器械软件 第1部分：YY/T 0316应用于医疗器械软件的指南\n"
                "ISO 14971-2019 医疗器械-风险管理对医疗器械的应用\n"
                "《医疗器械软件注册技术审查指导原则》（2022年第9号）\n"
                "《医疗器械网络安全注册审查指导原则》（2022年第7号）\n"
                "《人工智能医疗器械注册审查指导原则》（2022年第8号）\n"
                "FDA-Content of Premarket Submissions for Device Software Functions"
            ),
            "风险控制方案分析": "风险管理小组已经识别合理适用的风险控制措施来降低风险到可接受水平，具体风险控制措施的分析详见附录B。",
            "风险控制措施的实施": (
                f"通过对{name}产品风险分析和风险评价的结果的分析，所有的风险控制措施已经被识别并且所有风险控制措施已经在设计中实施。\n"
                "识别出的所有风险控制措施已经被验证，详见附录B的证据列，包括但不限于《软件测试报告》和《用户测试报告》。"
                "实施和验证的风险控制措施列表如下所示："
            ),
            "剩余风险分析和风险/受益分析": (
                f"{name}产品的所有单个剩余风险都已经控制在可接受的范围内，剩余风险可以接受，详见附录B，风险/受益分析评价列。"
            ),
            "由风险控制措施产生的风险": (
                "对采用的风险控制措施在评审过程中进行了分析，如果带来了新的风险，则进行分析，由风险控制措施带来的危害列表如下所示："
            ),
            "评审数据": (
                f"{name}的风险管理活动是由公司构成的联合风险管理小组开展的，这是一个跨专业的评审，至少应该包括风险管理小组的全部成员。\n"
                f"{name}的风险管理活动在{period}进行，风险管理小组按照风险管理计划制定风险分析的范围，对相关风险进行了分析和控制。\n"
                "风险管理活动的整个结果将包含在本风险管理报告中；风险管理活动在位于中国北京的公司进行。"
            ),
            "与安全有关特征的问题识别": (
                f"参照YY/T 0316的附录C，对{name}的预期用途和安全性特征进行了判定。"
                f"可能影响{name}安全性的预期用途和判定特征有关问题详见附录A。"
            ),
            "已知或可预见的危险（源）识别": (
                f"对可能影响{name}产品的已知或可预见危险（源）形成因素进行有关分析与判定，详见附录B。"
            ),
            "估计每个危险情况的风险": "基于产品需求的危险（源）识别来估计每个危险情况的风险，详见附录B。",
            "风险评价": f"{name}产品通过对每个风险点进行严重度和发生概率的联合评价，评价的结果详见附录B。",
            "RCMs实施风险控制措施前/后的风险分布": (
                "所有已识别的危险情况产生的一个或多个风险已经得到考虑，识别出不可接受的风险点已经通过有效的"
                "风险控制措施降低风险到可接受水平。表3和表4分别代表风险控制措施前和风险控制措施后的风险矩阵总表。\n"
                "通过风险水平分布的表格，我们可以确认所有风险都已经被分析和控制到可接受水平。"
            ),
            "综合剩余风险评价": (
                f"通过评审附录B的风险分析的结果和6.5.3的受益分析，{name}的主要风险来源于可用性，网络安全和模型及训练数据的正确性，"
                "应该注意的是，这些风险是医疗器械软件的共性，无法完全消除。\n"
                "此外，经过软件的验证和确认，验证了软件的临床功能满足预期用途，且综合剩余风险可接受。\n"
                f"基于以上计算和评估的总结，风险管理小组相信{name}的综合剩余风险是可接受的。"
            ),
            "软件安全级别判定": (
                f"根据YY/T 0664-2020 医疗器械软件 软件生存周期过程判定，{name}软件的安全级别为 C级。\n"
                "判断原因：\n"
                f"{name}用于辅助医师对多时相CT图像进行测量、比较及随访评估。尽管该软件的输出结果并非最终诊断依据，"
                "但其提供的测量、随访比较及提示可能在临床决策中对医师有所影响。如果软件在这些功能上存在不准确的结果，"
                "且医师未能识别和纠正这些偏差，可能会导致误判，从而引发漏检、误诊等情况，进而可能导致对患者造成严重伤害。\n"
                f"综上，根据YY/T 0664-2020的分类准则{name}的软件安全级别为C级。"
            ),
            "结论": (
                "通过评估风险管理文件，风险管理小组认为：\n"
                "1）风险管理计划已被实施，相关活动经过评审。\n"
                "2）经过临床评价，单个和综合剩余风险可接受。\n"
                "3）生产和生产后信息收集和评审活动已被安排。\n"
                f"{name}的研制阶段已对有关可能的危害及产生的风险进行了估计和评价，针对性地实施了降低风险的技术和管理方面的措施。"
                "产品测试对上述措施的有效性进行了验证，达到了通用和专用标准的要求。公司对所有剩余风险进行了评价，全部达到可接受的水平。\n"
                "经过生产及生产后信息收集和评审活动，同类产品未发现不良事件，产品不良事件监测及定期风险评价活动中，暂未发现新增风险。\n"
                "因此，产品风险管理过程满足标准要求，产品风险满足发行和上市要求。"
            ),
        }

    def __autofill_body_sections(self, content, product_id, version, product=None):
        if not product_id:
            return content
        if product is None:
            product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
        pname = (getattr(product, "name", "") or "").strip() if product else ""
        overall = (getattr(product, "overall_desc", "") or "").strip() if product else ""
        time_range = self.__project_time_range(product_id)
        defaults = self.__body_text_defaults(pname, time_range)
        lazy = {"ptr": None, "appx_a": None, "dist": None}

        def has_text(node):
            return bool(str(node.get("text") or node.get("content") or "").strip())

        def is_appendix_a(node):
            t = re.sub(r"\s+", "", str(node.get("title") or ""))
            return "附录A" in t and "安全有关特征" in t

        def product_desc():
            if not product:
                return ""
            return (
                f"产品名称：{getattr(product, 'name', '') or ''}\n"
                f"产品型号：{getattr(product, 'type_code', '') or ''}\n"
                f"完整版本：{getattr(product, 'full_version', '') or ''}"
            )

        def walk(node):
            title = self.__strip_no(node.get("title"))
            if title in defaults and not has_text(node):
                node["text"] = defaults[title]
            if title == "范围" and not has_text(node) and product and (getattr(product, "scope", "") or "").strip():
                node["text"] = product.scope.strip()
            if title == "产品描述" and not has_text(node) and product:
                node["text"] = product_desc()
            if title == "产品预期用途" and not has_text(node) and product and (getattr(product, "component", "") or "").strip():
                node["text"] = product.component.strip()
            if title == "严重度定义" and not (node.get("tables") or []):
                node["tables"] = [copy.deepcopy(self.SEVERITY_TABLE)]
            if title == "发生概率定义" and not (node.get("tables") or []):
                node["tables"] = [copy.deepcopy(self.PROBABILITY_TABLE)]
            if title == "产品功能描述":
                if overall and not has_text(node):
                    node["text"] = overall
                children = node.setdefault("children", [])
                if not any((c or {}).get("ref_type") == "prod_func_detail" for c in children):
                    children.insert(0, {"title": "3.2.1 产品功能明细", "ref_type": "prod_func_detail", "children": []})
            if (node or {}).get("ref_type") == "prod_func_detail":
                if pname:
                    node["title"] = f"3.2.1 {pname}"
                if not has_text(node):
                    if lazy["ptr"] is None:
                        lazy["ptr"] = self.__ptr_func_2_1(product_id) or ""
                    if lazy["ptr"]:
                        node["text"] = lazy["ptr"]
            if is_appendix_a(node) and not (node.get("tables") or []):
                if lazy["appx_a"] is None:
                    lazy["appx_a"] = self.__pha_appendix_a_table(product_id) or []
                if lazy["appx_a"]:
                    node["tables"] = [copy.deepcopy(lazy["appx_a"])]
            tnorm = re.sub(r"\s+", "", str(node.get("title") or ""))
            if "由风险控制措施产生的风险" in tnorm:
                tbs = node.get("tables") or []
                table0 = tbs[0] if tbs and isinstance(tbs[0], list) else None
                if not table0 or len(table0) <= 1:
                    node["tables"] = [[
                        ["RCM编号", "引入的危害", "RCM引入的风险分析", "风险控制措施"],
                        ["RCM017", "", "", "RCM017.在用户说明书中说明，在软件中无法找到所需图像时的故障排除步骤。"],
                        ["RCM115", "", "", "RCM115.在用户说明书中增加警告：请在随访前请确认所选病例的ID与医院系统匹配，以免造成错误。"],
                    ]]
            if title == "RCMs实施风险控制措施前/后的风险分布" and not (node.get("tables") or []):
                if lazy["dist"] is None:
                    lazy["dist"] = self.__risk_dist_matrices(product_id) or []
                if lazy["dist"]:
                    node["tables"] = copy.deepcopy(lazy["dist"])
            for child in node.get("children") or []:
                walk(child)

        for section in content.get("sections") or []:
            walk(section)
        return content

    def __fill_participants(self, content):
        # 4.2 风险分析参与人员：content.participants 为空时回填主表数据（与编辑页回退一致，保证导出也有内容）
        if not isinstance(content, dict) or content.get("participants"):
            return content
        rows = db.session.execute(select(RiskParticipant).order_by(RiskParticipant.id)).scalars().all()
        content["participants"] = [{"id": r.id, "role": r.role or "", "name": r.name or ""} for r in rows]
        return content

    def __to_obj(self, row: RiskMgmtDoc, product: Product = None, with_autofill=True):
        obj = RiskMgmtDocObj(**row.dict())
        obj.content = self.__normalize_content(obj.content)
        if with_autofill:
            serv_review_util.ensure_review(
                obj.content, "risk",
                serv_review_util.review_date(row.product_id, serv_review_util.REVIEW_DEFS["risk"]["name_keywords"]) if row.product_id else "",
                row.product_id,
            )
            obj.content = self.__autofill_front_matter(obj.content, row.product_id, row.version)
            obj.content = self.__fill_risk_mgmt_files(obj.content, row.product_id)
            obj.content = self.__autofill_body_sections(obj.content, row.product_id, row.version, product)
            obj.content = self.__fill_participants(obj.content)
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
            if not (obj.file_no or "").strip():
                resolved = serv_review_util.resolve_doc_file_no(product.id, obj.file_no, obj.version, "risk")
                if resolved:
                    obj.file_no = resolved
        return obj

    def __fill_list_obj(self, obj, product: Product = None, doc: RiskMgmtDoc = None):
        if product:
            obj.product_name = product.name
            obj.product_full_version = product.full_version
        if doc:
            obj.doc_version = doc.version
        return obj

    def __to_analysis_obj(self, row: RiskAnalysis, product: Product = None, doc: RiskMgmtDoc = None):
        return self.__fill_list_obj(RiskAnalysisObj(**row.dict()), product, doc)

    def __to_control_obj(self, row: RiskControl, product: Product = None, doc: RiskMgmtDoc = None):
        return self.__fill_list_obj(RiskControlObj(**row.dict()), product, doc)

    def __sync_product_id_from_doc(self, form):
        if form.doc_id and not form.product_id:
            doc = db.session.execute(select(RiskMgmtDoc).where(RiskMgmtDoc.id == form.doc_id)).scalars().first()
            if doc:
                form.product_id = doc.product_id
        return form

    def __normalize_section_title(self, value: str):
        txt = re.sub(r"\s+", "", str(value or "").strip())
        txt = re.sub(r"^[0-9０-９]+(?:[.．][0-9０-９]+)*(?:[、.．\s]+|(?=[\u4e00-\u9fffA-Za-z]))", "", txt)
        return txt

    def __is_acceptance_standard_section(self, section):
        title = self.__normalize_title_text((section or {}).get("title", ""))
        return (section or {}).get("ref_type") == "acceptance_standard" or ("5.2.3" in title and "接受标准" in title)

    def __iter_docx_blocks(self, docx: Document):
        for child in docx.element.body.iterchildren():
            tag = str(child.tag).lower()
            if tag.endswith("}p"):
                yield Paragraph(child, docx._body)
            elif tag.endswith("}tbl"):
                yield DocxTable(child, docx._body)

    def __extract_docx_paragraph_images(self, para: Paragraph):
        urls = []
        used_rids = set()
        mime_map = {
            "png": "image/png",
            "jpg": "image/jpeg",
            "jpeg": "image/jpeg",
            "gif": "image/gif",
            "bmp": "image/bmp",
            "webp": "image/webp",
        }
        blips = para._element.xpath(".//*[local-name()='blip']")
        for blip in blips:
            rid = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            if not rid or rid in used_rids:
                continue
            used_rids.add(rid)
            try:
                rel = para.part.rels[rid]
                target = getattr(rel, "target_ref", "")
                ext = (target.rsplit(".", 1)[-1].lower() if "." in target else "png")
                mime = mime_map.get(ext, "image/png")
                b64 = base64.b64encode(rel.target_part.blob).decode("ascii")
                urls.append(f"data:{mime};base64,{b64}")
            except Exception:
                logger.exception("解析风险管理 Word 图片失败")
        return urls

    def __find_acceptance_section(self, sections):
        for section in sections or []:
            if self.__is_acceptance_standard_section(section):
                return section
            found = self.__find_acceptance_section((section or {}).get("children") or [])
            if found:
                return found
        return None

    def __parse_docx_table(self, table: DocxTable):
        rows = []
        for row in table.rows:
            values = [re.sub(r"\s+", " ", (cell.text or "").strip()) for cell in row.cells]
            if any(values):
                rows.append(values)
        return rows

    def __extract_risk_content_from_word(self, docx: Document):
        content = self.__normalize_content(None)
        title_map = {}

        def walk(items):
            for section in items or []:
                title_map[self.__normalize_section_title(section.get("title", ""))] = section
                walk(section.get("children") or [])

        walk(content.get("sections") or [])
        current_section = None

        for block in self.__iter_docx_blocks(docx):
            if isinstance(block, Paragraph):
                txt = re.sub(r"\s+", " ", (block.text or "").strip())
                images = self.__extract_docx_paragraph_images(block)
                if txt:
                    key = self.__normalize_section_title(txt)
                    if key in title_map:
                        current_section = title_map[key]
                    elif current_section is not None:
                        current_text = str(current_section.get("text") or "").strip()
                        current_section["text"] = f"{current_text}\n{txt}".strip() if current_text else txt
                if images and current_section is not None:
                    current_section.setdefault("images", []).extend(images)
                    if self.__is_acceptance_standard_section(current_section) and not current_section.get("image_url"):
                        current_section["image_url"] = images[0]
            elif isinstance(block, DocxTable) and current_section is not None:
                rows = self.__parse_docx_table(block)
                if rows:
                    current_section.setdefault("tables", []).append(rows)
                    if current_section.get("ref_type") == "participants":
                        participants = []
                        data_rows = rows[1:] if len(rows) > 1 else []
                        for row in data_rows:
                            role = row[0] if len(row) > 0 else ""
                            name = row[1] if len(row) > 1 else ""
                            if not role and not name:
                                continue
                            participants.append({
                                "role": role,
                                "name": name,
                            })
                        if participants:
                            content["participants"] = participants
        return content

    def __save_imported_word_bytes(self, doc_id: int, filename: str, bys: bytes):
        suffix = (os.path.splitext(filename or "")[1] or ".docx").lower()
        path = os.path.join("data.trace", "risk_mgmt_doc_word", str(doc_id), get_uuid() + suffix)
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "wb") as fs:
            fs.write(bys or b"")
        return path

    async def add_risk_mgmt_doc(self, form: RiskMgmtDocForm):
        try:
            sql = select(func.count(RiskMgmtDoc.id)).where(
                RiskMgmtDoc.product_id == form.product_id,
                RiskMgmtDoc.version == form.version,
            )
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=f"版本 {form.version} 已存在，请更换文件版本号")
            row = RiskMgmtDoc(**form.dict(exclude_none=True))
            row.id = None
            row.file_no = serv_review_util.resolve_doc_file_no(form.product_id, form.file_no, form.version, "risk") or None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=RiskMgmtDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def import_risk_mgmt_doc_word(self, product_id: int, version: str, file_no: str = "", change_log: str = "", file=None):
        content = None
        bys = None
        filename = ""
        if file is not None:
            bys = await file.read()
            filename = file.filename or ""
            docx = Document(io.BytesIO(bys))
            content = self.__extract_risk_content_from_word(docx)
        form = RiskMgmtDocForm(product_id=product_id, version=version, file_no=file_no, change_log=change_log, content=content)
        resp = await self.add_risk_mgmt_doc(form)
        if resp.code == 1 and bys:
            row = db.session.execute(
                select(RiskMgmtDoc)
                .where(RiskMgmtDoc.product_id == product_id, RiskMgmtDoc.version == version)
            ).scalars().first()
            if row:
                next_content = self.__normalize_content(row.content)
                next_content["sourceWordPath"] = self.__save_imported_word_bytes(row.id, filename, bys)
                row.content = next_content
                db.session.commit()
        return resp

    async def duplicate_risk_mgmt_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: RiskMgmtDoc = db.session.execute(select(RiskMgmtDoc).where(RiskMgmtDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            # 复制目标产品：默认沿用原产品，跨产品复制时使用指定产品
            target_pid = product_id or fromdoc.product_id
            # 自动计算新版本号（不允许手动指定）
            all_versions = db.session.execute(select(RiskMgmtDoc.version).where(RiskMgmtDoc.product_id == target_pid)).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                # 同产品复制：在原文档版本号上递增
                version = new_version(fromdoc.version)
            else:
                # 跨产品复制：查目标产品现有最大版本后递增；目标产品无文档时沿用原版本
                def _version_seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_version_seq)) if valid else fromdoc.version
            # 兜底：确保目标产品下版本唯一
            while version in existing_set:
                version = new_version(version)
            base_file_no = (fromdoc.file_no or "").strip() or self.__dhf_code(target_pid, "风险管理报告")
            newdoc = RiskMgmtDoc(
                product_id=target_pid,
                version=version,
                file_no=sync_file_no_version(base_file_no, version) or None,
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.flush()

            for item in db.session.execute(select(RiskAnalysis).where(RiskAnalysis.doc_id == fromdoc.id)).scalars().all():
                newitem = RiskAnalysis(**item.dict())
                newitem.id = None
                newitem.doc_id = newdoc.id
                newitem.product_id = target_pid
                db.session.add(newitem)
            for item in db.session.execute(select(RiskControl).where(RiskControl.doc_id == fromdoc.id)).scalars().all():
                newitem = RiskControl(**item.dict())
                newitem.id = None
                newitem.doc_id = newdoc.id
                newitem.product_id = target_pid
                db.session.add(newitem)
            db.session.commit()
            return Resp.resp_ok(data=RiskMgmtDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_risk_mgmt_doc(self, form: RiskMgmtDocForm):
        try:
            row: RiskMgmtDoc = db.session.execute(select(RiskMgmtDoc).where(RiskMgmtDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = form.product_id if form.product_id is not None else row.product_id
            target_ver = form.version if form.version is not None else row.version
            dup = db.session.execute(
                select(func.count(RiskMgmtDoc.id)).where(
                    RiskMgmtDoc.product_id == target_pid,
                    RiskMgmtDoc.version == target_ver,
                    RiskMgmtDoc.id != form.id,
                )
            ).scalar()
            if dup and dup > 0:
                return Resp.resp_err(msg=f"版本 {target_ver} 已存在，请更换文件版本号")
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_risk_mgmt_doc(self, id: int):
        db.session.execute(delete(RiskAnalysis).where(RiskAnalysis.doc_id == id))
        db.session.execute(delete(RiskControl).where(RiskControl.doc_id == id))
        db.session.execute(delete(RiskMgmtDoc).where(RiskMgmtDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_risk_mgmt_doc(self, id: int):
        sql = select(RiskMgmtDoc, Product).join(Product, RiskMgmtDoc.product_id == Product.id).where(RiskMgmtDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def preview_risk_mgmt_content(self, product_id: int = 0, version: str = ""):
        # 编辑器切换产品时：基于模板按新产品跑一遍完整自动填充，返回全新 content（供前端整份刷新自动章节）
        product = None
        if product_id:
            product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
        content = self.__normalize_content(None)
        serv_review_util.ensure_review(
            content, "risk",
            serv_review_util.review_date(product_id, serv_review_util.REVIEW_DEFS["risk"]["name_keywords"]) if product_id else "",
            product_id,
        )
        content = self.__autofill_front_matter(content, product_id, version)
        content = self.__fill_risk_mgmt_files(content, product_id)
        content = self.__autofill_body_sections(content, product_id, version, product)
        return Resp.resp_ok(data=content)

    async def list_risk_mgmt_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(RiskMgmtDoc.product_id == product_id)
        if version:
            wheres.append(RiskMgmtDoc.version.like(f"%{version}%"))
        # 数据可见范围（与产品列表口径一致）：产品经理只看自己创建的产品对应的风险管理报告
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(RiskMgmtDoc.id)).join(Product, RiskMgmtDoc.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (
            select(RiskMgmtDoc, Product)
            .join(Product, RiskMgmtDoc.product_id == Product.id)
            .where(*wheres)
            .order_by(RiskMgmtDoc.id.desc())
            .offset(page_index * page_size)
            .limit(page_size)
        )
        rows: List[RiskMgmtDocObj] = [self.__to_obj(doc, product, with_autofill=False) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    async def add_risk_participant(self, form: RiskParticipantForm):
        try:
            role = (form.role or "").strip()
            name = (form.name or "").strip()
            if not role or not name:
                return Resp.resp_err(msg="请填写项目角色和姓名")
            sql = select(func.count(RiskParticipant.id)).where(RiskParticipant.role == role, RiskParticipant.name == name)
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = RiskParticipant(role=role, name=name)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_risk_participant(self, form: RiskParticipantForm):
        try:
            row: RiskParticipant = db.session.execute(select(RiskParticipant).where(RiskParticipant.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            role = (form.role or "").strip()
            name = (form.name or "").strip()
            if not role or not name:
                return Resp.resp_err(msg="请填写项目角色和姓名")
            sql = select(func.count(RiskParticipant.id)).where(
                RiskParticipant.role == role,
                RiskParticipant.name == name,
                RiskParticipant.id != form.id,
            )
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row.role = role
            row.name = name
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_risk_participant(self, id: int):
        db.session.execute(delete(RiskParticipant).where(RiskParticipant.id == id))
        db.session.commit()
        return Resp.resp_ok()

    def __seed_risk_participants_from_docs(self):
        if (db.session.execute(select(func.count(RiskParticipant.id))).scalar() or 0) > 0:
            return
        exists = set()
        rows = []
        docs = db.session.execute(select(RiskMgmtDoc)).scalars().all()
        for doc in docs:
            content = self.__normalize_content(doc.content)
            for item in content.get("participants") or []:
                role = str(item.get("role") or "").strip()
                name = str(item.get("name") or "").strip()
                key = (role, name)
                if role and name and key not in exists:
                    exists.add(key)
                    rows.append(RiskParticipant(role=role, name=name))
        if rows:
            db.session.add_all(rows)
            db.session.commit()

    async def list_risk_participant(self, keyword: str = None, page_index: int = 0, page_size: int = 10):
        self.__seed_risk_participants_from_docs()
        wheres = []
        if keyword:
            like = f"%{keyword}%"
            wheres.append((RiskParticipant.role.like(like)) | (RiskParticipant.name.like(like)))
        total = db.session.execute(select(func.count(RiskParticipant.id)).where(*wheres)).scalar() or 0
        rows = db.session.execute(
            select(RiskParticipant)
            .where(*wheres)
            .order_by(RiskParticipant.id.desc())
            .offset(page_index * page_size)
            .limit(page_size)
        ).scalars().all()
        return Resp.resp_ok(data=Page(
            total=total,
            rows=[RiskParticipantObj(**row.dict()) for row in rows],
            page_index=page_index,
            page_size=page_size,
        ))

    async def add_risk_analysis(self, form: RiskAnalysisForm):
        try:
            form = self.__sync_product_id_from_doc(form)
            sql = select(func.count(RiskAnalysis.id)).where(RiskAnalysis.doc_id == form.doc_id, RiskAnalysis.haz_code == form.haz_code)
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = RiskAnalysis(**form.dict(exclude_none=True))
            row.id = None
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_risk_analysis(self, form: RiskAnalysisForm):
        try:
            form = self.__sync_product_id_from_doc(form)
            row: RiskAnalysis = db.session.execute(select(RiskAnalysis).where(RiskAnalysis.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            sql = select(func.count(RiskAnalysis.id)).where(
                RiskAnalysis.doc_id == form.doc_id,
                RiskAnalysis.haz_code == form.haz_code,
                RiskAnalysis.id != form.id,
            )
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            for key, value in form.dict(exclude_none=True).items():
                if key != "id":
                    setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_risk_analysis(self, id: int):
        db.session.execute(delete(RiskAnalysis).where(RiskAnalysis.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def list_risk_analysis(self, product_id: int = 0, doc_id: int = 0, keyword: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(RiskAnalysis.product_id == product_id)
        if doc_id:
            wheres.append(RiskAnalysis.doc_id == doc_id)
        if keyword:
            like = f"%{keyword}%"
            wheres.append((RiskAnalysis.haz_code.like(like)) | (RiskAnalysis.source.like(like)) | (RiskAnalysis.hazard_situation.like(like)))
        total = db.session.execute(select(func.count(RiskAnalysis.id)).where(*wheres)).scalar() or 0
        sql = (
            select(RiskAnalysis, RiskMgmtDoc, Product)
            .join(RiskMgmtDoc, RiskAnalysis.doc_id == RiskMgmtDoc.id)
            .join(Product, RiskAnalysis.product_id == Product.id)
            .where(*wheres)
            .order_by(RiskAnalysis.id.desc())
            .offset(page_index * page_size)
            .limit(page_size)
        )
        rows = [self.__to_analysis_obj(row, product, doc) for row, doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    async def add_risk_control(self, form: RiskControlForm):
        try:
            form = self.__sync_product_id_from_doc(form)
            sql = select(func.count(RiskControl.id)).where(RiskControl.doc_id == form.doc_id, RiskControl.rcm_code == form.rcm_code)
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = RiskControl(**form.dict(exclude_none=True))
            row.id = None
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_risk_control(self, form: RiskControlForm):
        try:
            form = self.__sync_product_id_from_doc(form)
            row: RiskControl = db.session.execute(select(RiskControl).where(RiskControl.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            sql = select(func.count(RiskControl.id)).where(
                RiskControl.doc_id == form.doc_id,
                RiskControl.rcm_code == form.rcm_code,
                RiskControl.id != form.id,
            )
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            for key, value in form.dict(exclude_none=True).items():
                if key != "id":
                    setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_risk_control(self, id: int):
        db.session.execute(delete(RiskControl).where(RiskControl.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def list_risk_control(self, product_id: int = 0, doc_id: int = 0, keyword: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(RiskControl.product_id == product_id)
        if doc_id:
            wheres.append(RiskControl.doc_id == doc_id)
        if keyword:
            like = f"%{keyword}%"
            wheres.append((RiskControl.rcm_code.like(like)) | (RiskControl.description.like(like)) | (RiskControl.hazard_codes.like(like)))
        total = db.session.execute(select(func.count(RiskControl.id)).where(*wheres)).scalar() or 0
        sql = (
            select(RiskControl, RiskMgmtDoc, Product)
            .join(RiskMgmtDoc, RiskControl.doc_id == RiskMgmtDoc.id)
            .join(Product, RiskControl.product_id == Product.id)
            .where(*wheres)
            .order_by(RiskControl.id.desc())
            .offset(page_index * page_size)
            .limit(page_size)
        )
        rows = [self.__to_control_obj(row, product, doc) for row, doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    async def export_risk_mgmt_doc(self, output, id: int):
        resp = await self.get_risk_mgmt_doc(id)
        obj: RiskMgmtDocObj = resp.data
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        document.settings.element.append(update_fields)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")
        docx_util.add_page_number_footer(section, obj.file_no or "")
        risk_rates = {1: "1", 2: "2", 3: "3", 4: "4", 5: "5"}
        risk_degrees = {"A": "A", "B": "B", "C": "C", "D": "D", "E": "E"}
        risk_levels = {"1": "不可接受", "2": "进一步降低的研究", "3": "可忽略"}

        def normalized_title(value):
            return re.sub(r"\s+", "", str(value or ""))

        def is_appendix_b(section):
            title = normalized_title(section.get("title", ""))
            return "附录B" in title and "风险分析矩阵" in title

        def is_cover_section(section):
            title = normalized_title(section.get("title", ""))
            return section.get("ref_type") == "cover" or title == "风险管理报告"

        def is_revision_section(section):
            title = normalized_title(section.get("title", ""))
            return section.get("ref_type") == "revision" or title == "文件修订记录"

        def is_acceptance_standard_section(section):
            title = normalized_title(section.get("title", ""))
            return section.get("ref_type") == "acceptance_standard" or ("5.2.3" in title and "接受标准" in title)

        def write_center_section_title(title, font_size=16.0, bold=True):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, title, font_size=font_size, bold=bold)

        def add_blank_lines(count):
            for _ in range(max(0, int(count or 0))):
                document.add_paragraph("")

        def insert_toc_field():
            p = document.add_paragraph()
            run_begin = p.add_run()
            fld_begin = OxmlElement("w:fldChar")
            fld_begin.set(qn("w:fldCharType"), "begin")
            fld_begin.set(qn("w:dirty"), "true")
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = ' TOC \\o "1-3" \\h \\z \\u '
            fld_separate = OxmlElement("w:fldChar")
            fld_separate.set(qn("w:fldCharType"), "separate")
            run_end = p.add_run()
            fld_end = OxmlElement("w:fldChar")
            fld_end.set(qn("w:fldCharType"), "end")
            run_begin._r.append(fld_begin)
            run_begin._r.append(instr)
            run_begin._r.append(fld_separate)
            p.add_run("目录将在打开文档后自动更新")
            run_end._r.append(fld_end)

        def is_appendix_a_table(rows):
            first = normalized_title("".join([str(v or "") for v in ((rows or [[]])[0] or [])]))
            return all(x in first for x in ["问题", "考虑的内容", "是否适用", "可能的危险"])

        def is_appendix_b_table(rows):
            first = normalized_title("".join([str(v or "") for v in ((rows or [[]])[0] or [])]))
            return all(x in first for x in ["危害编号", "事件序列", "风险控制措施", "RCMID"])

        def is_appendix_a_section(section):
            title = normalized_title(section.get("title", ""))
            return "附录A" in title and "安全有关特征" in title

        def is_risk_mgmt_files_section(section):
            title = normalized_title(section.get("title", ""))
            return title.startswith("11") and "风险管理文件" in title

        def remove_appendix_reference_lines(text):
            lines = []
            for line in str(text or "").splitlines():
                value = normalized_title(line)
                if value.startswith("附录A") or value.startswith("附录B"):
                    continue
                lines.append(line)
            return "\n".join(lines).strip()

        def replace_all_text(value, old_text, new_text):
            if not old_text or not new_text or old_text == new_text:
                return str(value or "")
            return str(value or "").replace(old_text, new_text)

        def collect_content_text(content):
            chunks = []
            def walk(sections):
                for item in sections or []:
                    chunks.extend([str(item.get("text") or ""), str(item.get("content") or "")])
                    for table_rows in item.get("tables", []) or []:
                        for row in table_rows or []:
                            chunks.extend([str(cell or "") for cell in row or []])
                    walk(item.get("children") or [])
            walk((content or {}).get("sections") or [])
            return "\n".join(chunks)

        def infer_previous_product_name(content, current_name):
            all_text = collect_content_text(content)
            candidates = []
            for candidate in [
                re.sub(r"[0-9０-９]+$", "", current_name or ""),
                re.sub(r"[A-Za-z0-9０-９._\-（）()]+$", "", current_name or ""),
            ]:
                candidate = candidate.strip()
                if candidate and candidate != current_name and len(candidate) >= 4 and candidate not in candidates:
                    candidates.append(candidate)
            for candidate in candidates:
                if candidate in all_text:
                    return candidate
            return ""

        def prepare_export_content(content):
            next_content = copy.deepcopy(self.__normalize_content(content))
            sections = next_content.get("sections") or []
            risk_mgmt_files = None
            appendix_a = None
            appendix_b = None
            def walk(items):
                nonlocal risk_mgmt_files, appendix_a, appendix_b
                for item in items or []:
                    if is_risk_mgmt_files_section(item):
                        risk_mgmt_files = item
                    if is_appendix_a_section(item):
                        appendix_a = item
                    if is_appendix_b(item):
                        appendix_b = item
                    walk(item.get("children") or [])
            walk(sections)
            if risk_mgmt_files:
                risk_mgmt_files["text"] = remove_appendix_reference_lines(risk_mgmt_files.get("text"))
                source_tables = risk_mgmt_files.get("tables") or []
                appendix_a_tables = [t for t in source_tables if is_appendix_a_table(t)]
                appendix_b_tables = [t for t in source_tables if is_appendix_b_table(t)]
                if appendix_a and appendix_a_tables:
                    appendix_a.setdefault("tables", [])
                    appendix_a["tables"].extend([t for t in appendix_a_tables if t not in appendix_a["tables"]])
                if appendix_b and appendix_b_tables:
                    appendix_b.setdefault("tables", [])
                    appendix_b["tables"].extend([t for t in appendix_b_tables if t not in appendix_b["tables"]])
                risk_mgmt_files["tables"] = [t for t in source_tables if not is_appendix_a_table(t) and not is_appendix_b_table(t)]
            current_name = str(obj.product_name or "").strip()
            previous_name = str(next_content.get("productName") or "").strip() or infer_previous_product_name(next_content, current_name)
            if current_name and previous_name and previous_name != current_name:
                def sync_item(item):
                    if "text" in item:
                        item["text"] = replace_all_text(item.get("text"), previous_name, current_name)
                    if "content" in item:
                        item["content"] = replace_all_text(item.get("content"), previous_name, current_name)
                    if item.get("tables"):
                        item["tables"] = [
                            [[replace_all_text(cell, previous_name, current_name) for cell in (row or [])] for row in (table_rows or [])]
                            for table_rows in item.get("tables") or []
                        ]
                    for child in item.get("children") or []:
                        sync_item(child)
                for item in sections:
                    sync_item(item)
            next_content["productName"] = current_name
            acceptance_section = self.__find_acceptance_section(sections)
            source_word_path = str(next_content.get("sourceWordPath") or "").strip()
            if acceptance_section and not (acceptance_section.get("image_url") or acceptance_section.get("img_url")) and source_word_path and os.path.exists(source_word_path):
                try:
                    source_content = self.__extract_risk_content_from_word(Document(source_word_path))
                    source_acceptance = self.__find_acceptance_section(source_content.get("sections") or [])
                    source_image = (source_acceptance or {}).get("image_url") or (((source_acceptance or {}).get("images") or [None])[0])
                    if source_image:
                        acceptance_section["image_url"] = source_image
                except Exception:
                    logger.exception("从原始风险管理 Word 回捞接受标准图片失败")
            return next_content

        def find_section(sections, predicate):
            for item in sections or []:
                if predicate(item):
                    return item
            return None

        def set_cell_text(cell, text, bold=False):
            s = str(text or "")
            # 签名图（编制/审核/批准人）：等比嵌入图片，不渲染 base64 文本
            if s.startswith("data:image"):
                try:
                    b64 = s.split(",", 1)[1] if "," in s else ""
                    cell.text = ""
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.add_run().add_picture(io.BytesIO(base64.b64decode(b64)), height=Pt(33))
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    return
                except Exception:
                    pass
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.3
            docx_util.fonted_txt(paragraph, s, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        def set_cell_rcm(cell, text, bold=False):
            # RCM ID 列：每个 RCM 编号独占一行（用换行符 add_break，保证 Word 中真正换行）
            codes = re.findall(r"RCM\s*\d+", str(text or "").upper())
            codes = [c.replace(" ", "") for c in codes]
            uniq = list(dict.fromkeys(codes))
            if not uniq:
                set_cell_text(cell, text, bold=bold)
                return
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.3
            for i, code in enumerate(uniq):
                if i > 0:
                    paragraph.add_run().add_break()
                docx_util.fonted_txt(paragraph, code, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        def add_plain_table(rows):
            rows = rows or []
            col_count = max([len(row or []) for row in rows] or [0])
            if col_count <= 0:
                return
            table = document.add_table(rows=0, cols=col_count)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for row in rows:
                cells = table.add_row().cells
                for idx in range(col_count):
                    set_cell_text(cells[idx], row[idx] if idx < len(row or []) else "", bold=(len(table.rows) == 1))
            document.add_paragraph()

        def add_section_image(image_url):
            raw_url = str(image_url or "").strip()
            if not raw_url:
                return False
            try:
                if raw_url.startswith("data:image/"):
                    image_data = raw_url.split(",", 1)[1]
                    paragraph = document.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(io.BytesIO(base64.b64decode(image_data)), width=Inches(5.8))
                else:
                    image_path = raw_url.split("?", 1)[0]
                    if not os.path.isabs(image_path):
                        image_path = image_path.lstrip("/")
                    if not os.path.exists(image_path):
                        return False
                    paragraph = document.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(image_path, width=Inches(5.8))
                document.add_paragraph()
                return True
            except Exception:
                logger.exception("导出风险管理图片失败")
                return False

        def add_default_acceptance_image():
            image_path = os.path.join(os.path.dirname(__file__), "assets", "risk_acceptance_matrix.jpg")
            return add_section_image(image_path)

        def shade_cell(cell, hex_color):
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), hex_color)
            tc_pr.append(shd)

        def add_risk_dist_matrix(grid, caption):
            # 7.1 风险分布：合并表头 + 总计行列 + 红/橙/绿着色 + 图例（配色与编辑页 5.2.3 接受标准一致）
            RATE_ROWS = [("经常", "5"), ("有时", "4"), ("偶然", "3"), ("很少", "2"), ("非常少", "1")]
            SEV_LABELS = [("可忽略", "A"), ("轻度", "B"), ("严重", "C"), ("危重的", "D"), ("灾难性的", "E")]
            RISK_LEVELS = [
                ["bad", "bad", "bad", "bad", "bad"],
                ["bad", "bad", "bad", "bad", "bad"],
                ["ok", "warn", "bad", "bad", "bad"],
                ["ok", "warn", "warn", "bad", "bad"],
                ["ok", "ok", "warn", "warn", "warn"],
            ]
            LEVEL_FILL = {"bad": "FF0000", "warn": "FFC000", "ok": "92D050"}
            data_rows = (grid or [])[1:]
            counts = []
            for ri in range(5):
                src = data_rows[ri] if ri < len(data_rows) else []
                row = []
                for ci in range(5):
                    try:
                        row.append(int(str(src[ci + 1]).strip()))
                    except (ValueError, IndexError, TypeError):
                        row.append(0)
                counts.append(row)
            col_totals = [sum(counts[ri][ci] for ri in range(5)) for ci in range(5)]
            grand_total = sum(col_totals)

            if caption:
                write_center_section_title(caption, font_size=12.0, bold=True)
            table = document.add_table(rows=11, cols=9)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            def cell(r, c):
                return table.cell(r, c)

            # 表头：风险值(2×3) / 严重度(1×5) / 总计(2×1)
            cell(0, 0).merge(cell(1, 2))
            set_cell_text(cell(0, 0), "风险值", bold=True)
            cell(0, 3).merge(cell(0, 7))
            set_cell_text(cell(0, 3), "严重度", bold=True)
            cell(0, 8).merge(cell(1, 8))
            set_cell_text(cell(0, 8), "总计", bold=True)
            for ci, (label, letter) in enumerate(SEV_LABELS):
                set_cell_text(cell(1, 3 + ci), f"{label}\n{letter}", bold=True)

            # 数据 5 行
            cell(2, 0).merge(cell(6, 0))
            set_cell_text(cell(2, 0), "发生概率", bold=True)
            for ri, (rate, score) in enumerate(RATE_ROWS):
                r = 2 + ri
                set_cell_text(cell(r, 1), rate, bold=True)
                set_cell_text(cell(r, 2), score, bold=True)
                for ci in range(5):
                    target = cell(r, 3 + ci)
                    set_cell_text(target, str(counts[ri][ci]))
                    shade_cell(target, LEVEL_FILL[RISK_LEVELS[ri][ci]])
                set_cell_text(cell(r, 8), str(sum(counts[ri])), bold=True)

            # 总计行
            cell(7, 0).merge(cell(7, 2))
            set_cell_text(cell(7, 0), "总计", bold=True)
            for ci in range(5):
                set_cell_text(cell(7, 3 + ci), str(col_totals[ci]), bold=True)
            set_cell_text(cell(7, 8), str(grand_total), bold=True)

            # 图例
            legends = [
                ("红色", "FF0000", "不可接受：这类风险本质上不可接受。必须寻求风险降低措施。"),
                ("橙色", "FFC000", "进一步降低的研究：这类风险必须降低到合理可行的最低限度才可视为可接受。"),
                ("绿色", "92D050", "可忽略：这类风险实际上可接受，但仍应尽可能寻求风险降低措施。"),
            ]
            for li, (word, fill, desc) in enumerate(legends):
                r = 8 + li
                set_cell_text(cell(r, 0), word, bold=True)
                shade_cell(cell(r, 0), fill)
                cell(r, 1).merge(cell(r, 8))
                set_cell_text(cell(r, 1), desc)
            document.add_paragraph()

        def set_section_orientation(target_section, landscape=False):
            is_current_landscape = target_section.orientation == WD_ORIENT.LANDSCAPE
            if landscape != is_current_landscape:
                target_section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
                target_section.page_width, target_section.page_height = target_section.page_height, target_section.page_width
            if landscape:
                target_section.left_margin = Inches(0.4)
                target_section.right_margin = Inches(0.4)
            else:
                target_section.left_margin = Inches(0.7)
                target_section.right_margin = Inches(0.7)
            target_section.top_margin = Inches(0.8)
            target_section.bottom_margin = Inches(0.8)

        def add_product_haz_matrix():
            rows = db.session.execute(
                select(ProdHaz, Haz)
                .outerjoin(Haz, ProdHaz.haz_id == Haz.id)
                .where(ProdHaz.prod_id == obj.product_id)
                .order_by(Haz.code)
            ).all()
            table = document.add_table(rows=1, cols=16)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            headers = [
                "危害编号", "危险（源）", "事件序列", "危险情况", "伤害",
                "初始风险概率", "初始风险危害程度", "初始风险风险水平",
                "风险控制措施", "RCM ID", "证据，包括风险验证（详见软件测试报告）",
                "剩余风险概率", "剩余风险危害程度", "剩余风险风险水平",
                "收益是否大于风险（Y/N）", "分类",
            ]
            for idx, header in enumerate(headers):
                set_cell_text(table.rows[0].cells[idx], header, bold=True)
            for prod_haz, haz in rows:
                cells = table.add_row().cells
                values = [
                    getattr(haz, "code", "") or "",
                    getattr(haz, "source", "") or "",
                    getattr(haz, "event", "") or "",
                    prod_haz.situation or getattr(haz, "situation", "") or "",
                    prod_haz.damage or getattr(haz, "damage", "") or "",
                    risk_rates.get(prod_haz.init_rate or getattr(haz, "init_rate", None), prod_haz.init_rate or getattr(haz, "init_rate", "") or ""),
                    risk_degrees.get(prod_haz.init_degree or getattr(haz, "init_degree", None), prod_haz.init_degree or getattr(haz, "init_degree", "") or ""),
                    risk_levels.get(prod_haz.init_level or getattr(haz, "init_level", None), prod_haz.init_level or getattr(haz, "init_level", "") or ""),
                    prod_haz.deal or getattr(haz, "deal", "") or "",
                    prod_haz.rcms or getattr(haz, "rcms", "") or "",
                    prod_haz.evidence or getattr(haz, "evidence", "") or "",
                    risk_rates.get(prod_haz.cur_rate or getattr(haz, "cur_rate", None), prod_haz.cur_rate or getattr(haz, "cur_rate", "") or ""),
                    risk_degrees.get(prod_haz.cur_degree or getattr(haz, "cur_degree", None), prod_haz.cur_degree or getattr(haz, "cur_degree", "") or ""),
                    risk_levels.get(prod_haz.cur_level or getattr(haz, "cur_level", None), prod_haz.cur_level or getattr(haz, "cur_level", "") or ""),
                    "Y" if getattr(haz, "benefit_flag", None) else "N",
                    getattr(haz, "category", "") or "",
                ]
                for idx, value in enumerate(values):
                    if idx == 9:
                        set_cell_rcm(cells[idx], value)
                    else:
                    set_cell_text(cells[idx], value)
            document.add_paragraph()

        def add_section(section: dict, level: int = 1):
            title = section.get("title", "")
            if title and not is_cover_section(section):
                docx_util.save_title2docx(title, document, level=max(1, min(level, 9)))
            if section.get("text"):
                docx_util.save_txt2docx(str(section.get("text") or ""), document)
            elif section.get("content"):
                docx_util.save_txt2docx(str(section.get("content") or ""), document)
            image_added = add_section_image(section.get("image_url") or section.get("img_url"))
            if section.get("ref_type") == "participants":
                table = document.add_table(rows=1, cols=2)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                headers = ["项目角色", "姓名"]
                for idx, header in enumerate(headers):
                    set_cell_text(table.rows[0].cells[idx], header, bold=True)
                for item in (obj.content or {}).get("participants", []):
                    cells = table.add_row().cells
                    for cidx, value in enumerate([item.get("role", "") or "", item.get("name", "") or ""]):
                        set_cell_text(cells[cidx], value)
            elif section.get("ref_type") == "review":
                def _review_set_cell(cell, text, bold=False, align=None):
                    set_cell_text(cell, text, bold=bold)
                for t_idx, table_rows in enumerate(section.get("tables", []) or []):
                    serv_review_util.render_review_grid(document, table_rows, _review_set_cell, merge_col0=(t_idx == 0), merge_full=True)
            elif is_appendix_b(section):
                add_product_haz_matrix()
            elif is_acceptance_standard_section(section):
                if not image_added:
                    add_default_acceptance_image()
            elif section.get("ref_type") == "risk_analysis":
                rows = db.session.execute(select(RiskAnalysis).where(RiskAnalysis.doc_id == id).order_by(RiskAnalysis.id)).scalars().all()
                if rows:
                    table = document.add_table(rows=1, cols=7)
                    table.style = "Table Grid"
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    headers = ["HAZ编号", "危险源", "事件序列", "危险情况", "伤害", "初始风险", "分类"]
                    for idx, header in enumerate(headers):
                        set_cell_text(table.rows[0].cells[idx], header, bold=True)
                    for row in rows:
                        cells = table.add_row().cells
                        values = [row.haz_code or "", row.source or "", row.event_sequence or "", row.hazard_situation or "", row.harm or "", " / ".join([str(v) for v in [row.init_rate, row.init_degree, row.init_level] if v]), row.category or ""]
                        for idx, value in enumerate(values):
                            set_cell_text(cells[idx], value)
                for table_rows in section.get("tables", []) or []:
                    add_plain_table(table_rows)
            elif section.get("ref_type") == "risk_controls":
                rows = db.session.execute(select(RiskControl).where(RiskControl.doc_id == id).order_by(RiskControl.id)).scalars().all()
                if rows:
                    table = document.add_table(rows=1, cols=5)
                    table.style = "Table Grid"
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    headers = ["RCM编号", "控制措施描述", "关联HAZ编号", "验证证据", "是否引入新风险"]
                    for idx, header in enumerate(headers):
                        set_cell_text(table.rows[0].cells[idx], header, bold=True)
                    for row in rows:
                        cells = table.add_row().cells
                        values = [row.rcm_code or "", row.description or "", row.hazard_codes or "", row.verification_evidence or "", "是" if row.new_risk_flag else "否"]
                        for idx, value in enumerate(values):
                            set_cell_text(cells[idx], value)
                for table_rows in section.get("tables", []) or []:
                    add_plain_table(table_rows)
            elif "由风险控制措施产生的风险" in re.sub(r"\s+", "", str(section.get("title") or "")) \
                    or "RCM带来的危害" in re.sub(r"\s+", "", str(section.get("title") or "")):
                rcm_rows = self.__rcm_introduced_rows(section, obj.product_id)
                table = document.add_table(rows=1, cols=4)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for idx, header in enumerate(["RCM编号", "引入的危害", "RCM引入的风险分析", "风险控制措施"]):
                    set_cell_text(table.rows[0].cells[idx], header, bold=True)
                for r in rcm_rows:
                    cells = table.add_row().cells
                    for idx, value in enumerate(r[:4]):
                        set_cell_text(cells[idx], str(value or ""))
                document.add_paragraph()
            elif "风险分布" in normalized_title(section.get("title") or ""):
                titles = section.get("table_titles") if isinstance(section.get("table_titles"), list) else []
                dist_tables = section.get("tables", []) or [["初始风险分布（措施前）"], ["剩余风险分布（措施后）"]]
                for t_idx, table_rows in enumerate(dist_tables):
                    label = (table_rows[0][0] if (table_rows and table_rows[0]) else "") or ""
                    caption = (titles[t_idx] if t_idx < len(titles) and titles[t_idx] else f"表{t_idx + 3} {label}").strip()
                    add_risk_dist_matrix(table_rows, caption)
            else:
                for table_rows in section.get("tables", []) or []:
                    add_plain_table(table_rows)
            for child in section.get("children", []) or []:
                add_section(child, level + 1)

        def add_top_level_section(section: dict, is_first_body_section: bool):
            use_landscape = is_appendix_a_section(section) or is_appendix_b(section)
            current_section = document.sections[-1]
            current_landscape = current_section.orientation == WD_ORIENT.LANDSCAPE
            if is_first_body_section:
                set_section_orientation(current_section, use_landscape)
            elif use_landscape or use_landscape != current_landscape:
                current_section = document.add_section(WD_SECTION_START.NEW_PAGE)
                set_section_orientation(current_section, use_landscape)
            add_section(section)

        export_content = prepare_export_content(obj.content or {})
        serv_review_util.ensure_review(
            export_content, "risk",
            serv_review_util.review_date(obj.product_id, serv_review_util.REVIEW_DEFS["risk"]["name_keywords"]) if obj.product_id else "",
            obj.product_id,
        )
        export_sections = (export_content or {}).get("sections", [])
        cover_section = find_section(export_sections, is_cover_section)
        revision_section = find_section(export_sections, is_revision_section)
        body_sections = [section for section in export_sections if not is_cover_section(section) and not is_revision_section(section)]

        add_blank_lines(12)
        write_center_section_title("风险管理报告", font_size=22.0, bold=True)
        add_blank_lines(6)
        for table_rows in (cover_section or {}).get("tables", []) or []:
            add_plain_table(table_rows)

        document.add_page_break()
        write_center_section_title("文件修订记录", font_size=14.0, bold=True)
        add_blank_lines(2)
        for table_rows in (revision_section or {}).get("tables", []) or []:
            add_plain_table(table_rows)

        document.add_page_break()
        write_center_section_title("目录", font_size=16.0, bold=True)
        insert_toc_field()

        document.add_page_break()
        set_section_orientation(document.sections[-1], False)
        for index, section in enumerate(body_sections):
            add_top_level_section(section, index == 0)
        docx_util.fill_toc_cache(document)
        document.save(output)
        output.seek(0)
