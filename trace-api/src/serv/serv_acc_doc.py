#!/usr/bin/env python
# encoding: utf-8

# 产品验收记录服务层。
# 整份文档以 content(JSON) 的「章节树」存储；验收基本信息(acc_info)按产品自动获取（只读）。
# 导出：标题「产品验收记录」→ 各章节（标题 + 表格 + 正文）。

import copy
import json
import logging
import os
import re
from typing import List

from sqlalchemy import delete, func, select
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.acc_doc import AccDoc
from ..model.project_member import ProjectMember
from ..model.prod_dhf import ProdDhf
from ..model.project_timeline import ProjectTimelineRow, ProjectTimelineCell
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_acc_doc import AccDocForm
from ..obj.vobj_acc_doc import AccDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from .serv_utils import new_version, sync_file_no_version, docx_util

logger = logging.getLogger(__name__)

# 文档大标题（导出）
DOC_TITLE = "产品验收记录"
# 自动获取（只读）章节标识
AUTO_REFS = {"acc_info"}
# 验收时间：优先匹配含以下关键字的时间线行，取命中行最新日期
ACC_DATE_KEYWORDS = ["验收"]

DEFAULT_ACC_CONTENT = {"sections": []}
_DEFAULT_CONTENT_FILE = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "src-res", "acc_default_content.json"
)
try:
    with open(_DEFAULT_CONTENT_FILE, encoding="utf-8") as _f:
        _loaded = json.load(_f)
    if isinstance(_loaded, dict) and isinstance(_loaded.get("sections"), list) and _loaded["sections"]:
        DEFAULT_ACC_CONTENT = _loaded
except Exception:
    logger.exception("加载产品验收记录默认内容资源失败")


class Server(object):

    # ---------------- 归一化 ----------------
    def __normalize_node(self, node):
        if not isinstance(node, dict):
            return {"title": str(node or ""), "text": "", "tables": [], "table_titles": [], "images": [], "children": []}
        result = dict(node)
        result["title"] = str(result.get("title") or "")
        result["text"] = str(result.get("text") or "")
        tables = result.get("tables")
        norm_tables = []
        if isinstance(tables, list):
            for table in tables:
                if isinstance(table, list):
                    norm_tables.append([[str(c) if c is not None else "" for c in (row or [])] for row in table if isinstance(row, list)])
        result["tables"] = norm_tables
        titles = result.get("table_titles")
        result["table_titles"] = [str(x or "") for x in titles] if isinstance(titles, list) else []
        imgs = result.get("images")
        result["images"] = [str(x) for x in imgs if isinstance(x, str) and x] if isinstance(imgs, list) else []
        children = result.get("children")
        result["children"] = [self.__normalize_node(c) for c in children] if isinstance(children, list) else []
        # blocks 由渲染时动态生成，持久化时不强制保留
        if isinstance(result.get("blocks"), list):
            result.pop("blocks", None)
        return result

    def __normalize_content(self, content):
        if not isinstance(content, dict) or not isinstance(content.get("sections"), list):
            return copy.deepcopy(DEFAULT_ACC_CONTENT)
        return {"sections": [self.__normalize_node(s) for s in content["sections"]]}

    @staticmethod
    def __strip_name(title):
        return re.sub(r"^\s*[一二三四五六七八九十]+、\s*", "", str(title or "")).strip()

    # ---------------- 自动获取 ----------------
    def __product_manager(self, prod_id):
        # 产品经理：优先职能含「产品经理」，其次「项目经理」，按排序取第一位
        for kw in ("产品经理", "项目经理", "经理"):
            row = db.session.execute(
                select(ProjectMember).where(ProjectMember.prod_id == prod_id, ProjectMember.role.like(f"%{kw}%"))
                .order_by(ProjectMember.sort_order.asc(), ProjectMember.id.asc())
            ).scalars().first()
            if row and (row.name or "").strip():
                return row.name.strip()
        return ""

    def __accept_date(self, prod_id):
        # 验收时间：从项目时间线取日期。优先命中含「验收」关键字的行，取最新；否则取时间线最晚日期
        def to_int(v):
            digits = re.sub(r"[^\d]", "", str(v or ""))
            return int(digits) if digits else None

        tl_rows = db.session.execute(
            select(ProjectTimelineRow).where(ProjectTimelineRow.prod_id == prod_id)
        ).scalars().all()
        if not tl_rows:
            return ""
        cell_map = {}
        for c in db.session.execute(
            select(ProjectTimelineCell).where(ProjectTimelineCell.row_id.in_([r.id for r in tl_rows]))
        ).scalars().all():
            cell_map.setdefault(c.row_id, []).append(c.output_result or "")
        date_rows = [r for r in tl_rows if (r.row_type or "date") == "date" and to_int(r.year) and to_int(r.month)]
        if not date_rows:
            return ""

        def date_key(r):
            return to_int(r.year) * 10000 + to_int(r.month) * 100 + (to_int(r.day) or 0)

        hit = [r for r in date_rows if any(any(k in str(v or "") for k in ACC_DATE_KEYWORDS) for v in cell_map.get(r.id, []))]
        target = max(hit, key=date_key) if hit else max(date_rows, key=date_key)
        day = to_int(target.day)
        return f"{to_int(target.year)}年{to_int(target.month)}月{day}日" if day else f"{to_int(target.year)}年{to_int(target.month)}月"

    def __collect_autofill(self, product_id):
        product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
        if not product:
            return {}
        return {
            "product_name": (product.name or "").strip(),
            "type_code": (product.type_code or "").strip(),
            "full_version": (product.full_version or "").strip(),
            "manager": self.__product_manager(product_id),
            "accept_date": self.__accept_date(product_id),
        }

    def __apply_autofill(self, content, auto):
        auto = auto or {}
        for node in self.__iter_nodes(content.get("sections", [])):
            rt = node.get("ref_type")
            if not rt and self.__strip_name(node.get("title")) == "验收基本信息":
                rt = "acc_info"
                node["ref_type"] = rt
            if rt == "acc_info":
                self.__fill_info_table(node, auto)
            node["blocks"] = self.__content_blocks(node)
        return content

    @staticmethod
    def __iter_nodes(nodes):
        for n in nodes or []:
            yield n
            yield from Server.__iter_nodes(n.get("children", []))

    def __fill_info_table(self, node, auto):
        tables = node.get("tables") or []
        if not tables or not isinstance(tables[0], list):
            return
        mapping = {
            "产品名称": auto.get("product_name", ""),
            "产品经理": auto.get("manager", ""),
            "规格型号": auto.get("type_code", "") or "-",
            "版本号": auto.get("full_version", ""),
            "验收时间": auto.get("accept_date", ""),
        }
        for row in tables[0]:
            if not isinstance(row, list):
                continue
            # 形如 [键, 值, 键, 值]
            for i in range(0, len(row) - 1, 2):
                key = str(row[i] or "").strip()
                if key in mapping:
                    row[i + 1] = mapping[key]

    @staticmethod
    def __content_blocks(node):
        # 将 text + tables + images 组织为有序 blocks，供前端渲染与导出对齐
        blocks = []
        text = str(node.get("text") or "").strip()
        if text:
            blocks.append({"text": text})
        for url in node.get("images", []) or []:
            blocks.append({"type": "image", "url": url})
        titles = node.get("table_titles") or []
        no_merge = node.get("ref_type") in AUTO_REFS
        for idx, rows in enumerate(node.get("tables", []) or []):
            b = {"type": "table", "table_index": idx, "table": rows,
                 "title": titles[idx] if idx < len(titles) else ""}
            if no_merge:
                b["no_merge"] = True
            blocks.append(b)
        return blocks

    # ---------------- 文件编号（未手填时从产品 DHF 匹配） ----------------
    def __dhf_file_no(self, prod_id):
        row = db.session.execute(
            select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name.like("%验收记录%"))
        ).scalars().first()
        return (row.code or "").strip() if row and row.code else ""

    def __to_obj(self, row: AccDoc, product: Product = None, with_autofill=True):
        obj = AccDocObj(**row.dict())
        content = self.__normalize_content(obj.content)
        if with_autofill:
            auto = self.__collect_autofill(row.product_id)
            content = self.__apply_autofill(content, auto)
        if product:
            content["productName"] = product.name or ""
        obj.content = content
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
            if not (obj.file_no or "").strip():
                dhf_no = self.__dhf_file_no(product.id)
                if dhf_no:
                    obj.file_no = dhf_no
        return obj

    # ---------------- CRUD ----------------
    async def add_acc_doc(self, form: AccDocForm):
        try:
            sql = select(func.count(AccDoc.id)).where(
                AccDoc.product_id == form.product_id, AccDoc.version == form.version
            )
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = AccDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=AccDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_acc_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: AccDoc = db.session.execute(select(AccDoc).where(AccDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(AccDoc.version).where(AccDoc.product_id == target_pid)).scalars().all()
            existing = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing:
                version = new_version(version)
            newdoc = AccDoc(
                product_id=target_pid, version=version, file_no=sync_file_no_version((fromdoc.file_no or "").strip() or self.__dhf_file_no(target_pid), version) or None,
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            return Resp.resp_ok(data=AccDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_acc_doc(self, form: AccDocForm):
        try:
            row: AccDoc = db.session.execute(select(AccDoc).where(AccDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            if form.product_id is not None:
                row.product_id = form.product_id
            if form.version is not None:
                row.version = form.version
            if form.file_no is not None:
                row.file_no = (form.file_no or "").strip() or None
            if form.change_log is not None:
                row.change_log = form.change_log
            if form.content is not None:
                row.content = self.__normalize_content(form.content)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_acc_doc(self, id: int):
        db.session.execute(delete(AccDoc).where(AccDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_acc_doc(self, id: int):
        sql = select(AccDoc, Product).join(Product, AccDoc.product_id == Product.id).where(AccDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_acc_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(AccDoc.product_id == product_id)
        if version:
            wheres.append(AccDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        total = db.session.execute(
            select(func.count(AccDoc.id)).join(Product, AccDoc.product_id == Product.id).where(*wheres)
        ).scalar() or 0
        sql = (
            select(AccDoc, Product).join(Product, AccDoc.product_id == Product.id)
            .where(*wheres).order_by(AccDoc.id.desc()).offset(page_index * page_size).limit(page_size)
        )
        rows: List[AccDocObj] = [self.__to_obj(doc, product, with_autofill=False) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    async def acc_autofill(self, product_id: int):
        # 新增页 / 切换产品预览：返回应用了自动获取的默认内容
        content = self.__normalize_content(None)
        auto = self.__collect_autofill(product_id)
        content = self.__apply_autofill(content, auto)
        product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
        if product:
            content["productName"] = product.name or ""
        return Resp.resp_ok(data=content)

    # ---------------- 导出 ----------------
    async def export_acc_doc(self, output, id: int):
        resp = await self.get_acc_doc(id)
        obj: AccDocObj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")

        def set_cell_text(cell, text, bold=False, center=False):
            cell.text = ""
            align = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
            lines = str(text or "").split("\n")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = align
            paragraph.paragraph_format.line_spacing = 1.5
            docx_util.fonted_txt(paragraph, lines[0], font_size=10.5, bold=bold)
            for ln in lines[1:]:
                p = cell.add_paragraph()
                p.alignment = align
                p.paragraph_format.line_spacing = 1.5
                docx_util.fonted_txt(p, ln, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # ---- 整份=一张 6 列栅格表（与 Word 原版一致，合并单元格）----
        GRID_TWIPS = [700, 2100, 2100, 1700, 1400, 2200]  # 验收结果列加宽，避免折行
        INFO_SPANS = [2, 2, 1, 1]   # 键|值|键|值
        ITEM_SPANS = [1, 2, 2, 1]   # 序号|验收内容|描述|验收结果

        def set_no_wrap(cell):
            tcPr = cell._tc.get_or_add_tcPr()
            tcPr.append(OxmlElement("w:noWrap"))

        def set_symbol_cell(cell, text, center=False):
            # 勾选框/方框符号用 MS Gothic 渲染，避免被 Word 用 emoji 字体显示成黑块；中文仍用宋体
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.5
            for part in re.findall(r"[\u4e00-\u9fa5]+|[^\u4e00-\u9fa5]+", str(text or "")):
                run = p.add_run(part)
                run.font.size = Pt(10.5)
                name = "宋体" if re.match(r"[\u4e00-\u9fa5]", part) else "MS Gothic"
                run.font.name = name
                run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        def add_grid_row(table, texts, spans, bold=False, head=False, nowrap=None, symbol=None):
            cells = table.add_row().cells  # 6 个网格单元
            targets = []
            idx = 0
            for j, (text, span) in enumerate(zip(texts, spans)):
                target = cells[idx]
                for k in range(idx + 1, idx + span):
                    target = target.merge(cells[k])
                if symbol and j < len(symbol) and symbol[j]:
                    set_symbol_cell(target, text, center=head)
                else:
                    set_cell_text(target, text, bold=bold or head, center=head)
                if nowrap and j < len(nowrap) and nowrap[j]:
                    set_no_wrap(target)
                targets.append(target)
                idx += span
            return targets

        def norm4(row):
            texts = [str(c or "") for c in (row or [])][:4]
            while len(texts) < 4:
                texts.append("")
            return texts

        content = obj.content if isinstance(obj.content, dict) else self.__normalize_content(obj.content)
        sections = content.get("sections", [])

        def is_auto(s):
            return s.get("ref_type") in AUTO_REFS
        info_sec = next((s for s in sections if is_auto(s)), None)
        item_sec = next((s for s in sections if not is_auto(s) and (s.get("tables") or [])), None)
        conc_sec = next((s for s in sections if not is_auto(s) and not (s.get("tables") or [])), None)

        title_para = document.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.line_spacing = 1.5
        docx_util.fonted_txt(title_para, DOC_TITLE, font_size=18.0, bold=True)
        document.add_paragraph()

        table = document.add_table(rows=0, cols=6)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        tbl = table._tbl
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tbl.tblPr.append(layout)
        grid = tbl.find(qn("w:tblGrid"))
        if grid is not None:
            for gc, w in zip(grid.findall(qn("w:gridCol")), GRID_TWIPS):
                gc.set(qn("w:w"), str(w))

        # 1) 基本信息（各格均短，整行不折行）
        info_rows = (info_sec.get("tables") or [[]])[0] if info_sec else []
        for row in info_rows:
            add_grid_row(table, norm4(row), INFO_SPANS, nowrap=[True, True, True, True])
        # 2) 验收内容（首行表头）：序号/验收结果不折行，验收内容/描述可换行
        item_rows = (item_sec.get("tables") or [[]])[0] if item_sec else []
        seq_cells, name_cells = [], []  # 明细行的「序号」「验收内容」单元格，用于纵向合并
        for ri, row in enumerate(item_rows):
            targets = add_grid_row(
                table, norm4(row), ITEM_SPANS, head=(ri == 0),
                nowrap=[True, False, False, True],
                symbol=(None if ri == 0 else [False, False, False, True]),
            )
            if ri > 0:
                seq_cells.append(targets[0])
                name_cells.append(targets[1])
        # 重复项纵向合并（手动 vMerge：首行 restart，其余 continue 且清空文字）
        def set_vmerge(cell, restart):
            tcPr = cell._tc.get_or_add_tcPr()
            for e in tcPr.findall(qn("w:vMerge")):
                tcPr.remove(e)
            vm = OxmlElement("w:vMerge")
            if restart:
                vm.set(qn("w:val"), "restart")
            else:
                cell.text = ""  # 续行清空，避免文字重复
            tcPr.append(vm)

        def merge_group(cells, start, end):
            if end <= start:
                return
            set_vmerge(cells[start], True)
            for k in range(start + 1, end + 1):
                set_vmerge(cells[k], False)

        data = [norm4(r) for r in item_rows[1:]]
        n = len(data)
        i = 0
        while i < n:
            j = i
            while j + 1 < n and data[j + 1][0] == data[i][0]:
                j += 1
            merge_group(seq_cells, i, j)  # 序号列：相同序号合并
            a = i
            while a <= j:
                b = a
                while b + 1 <= j and data[b + 1][1] == data[a][1]:
                    b += 1
                merge_group(name_cells, a, b)  # 验收内容：同序号组内相同内容合并
                a = b + 1
            i = j + 1
        # 3) 验收结论（整行跨列）
        if conc_sec and str(conc_sec.get("text") or "").strip():
            add_grid_row(table, [str(conc_sec.get("text") or "")], [6])

        document.save(output)
        output.seek(0)
