#!/usr/bin/env python
# encoding: utf-8

# 软件配置状态报告服务层（开发文件，PDP 风格章节树）。默认内容取自 src-res/scs_default_content.json。

import base64
import copy
import json
import logging
import os
import re
from io import BytesIO
from typing import List

from sqlalchemy import delete, func, select
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.scs_doc import ScsDoc
from ..model.scm_doc import ScmDoc
from ..model.prod_dhf import ProdDhf
from ..model.project_member import ProjectMember
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_scs_doc import ScsDocForm
from ..obj.vobj_scs_doc import ScsDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from . import serv_review_util
from .serv_utils import new_version, sync_file_no_version
from .serv_utils import docx_util

logger = logging.getLogger(__name__)

_DEFAULT_FILE = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "src-res", "scs_default_content.json")
try:
    with open(_DEFAULT_FILE, encoding="utf-8") as _f:
        DEFAULT_SCS_CONTENT = json.load(_f)
except Exception:
    DEFAULT_SCS_CONTENT = {"sections": []}

# 模板基准产品名（用于全文替换为当前产品名）。
BASE_NAME = "InferOperate Suite"
DOC_NAME = "软件配置状态报告"
DOC_KEY = "scs"


class Server(object):

    def __normalize_node(self, node):
        if not isinstance(node, dict):
            return {"title": str(node or ""), "body": "", "tables": [], "children": []}
        result = dict(node)
        result["title"] = str(result.get("title") or "")
        result["body"] = str(result.get("body") or "")
        tables = result.get("tables")
        if not isinstance(tables, list):
            tables = []
        norm_tables = []
        for table in tables:
            if isinstance(table, list):
                norm_tables.append([[str(c) if c is not None else "" for c in (row or [])] for row in table if isinstance(row, list)])
        result["tables"] = norm_tables
        children = result.get("children")
        result["children"] = [self.__normalize_node(c) for c in children] if isinstance(children, list) else []
        return result

    def __normalize_content(self, content):
        if not isinstance(content, dict) or not isinstance(content.get("sections"), list):
            return copy.deepcopy(DEFAULT_SCS_CONTENT)
        return {"sections": [self.__normalize_node(s) for s in content["sections"]]}

    def __replace_name(self, node, base, name, skip_titles=None):
        if not name or base == name:
            return
        if skip_titles and str(node.get("title") or "").strip() in skip_titles:
            return
        if node.get("body"):
            node["body"] = node["body"].replace(base, name)
        for tbl in (node.get("tables") or []):
            for row in tbl:
                for i in range(len(row)):
                    if isinstance(row[i], str) and base in row[i]:
                        row[i] = row[i].replace(base, name)
        for c in (node.get("children") or []):
            self.__replace_name(c, base, name, skip_titles=skip_titles)

    def __dhf_file_no(self, prod_id):
        if not prod_id:
            return ""
        row = db.session.execute(
            select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name == DOC_NAME).order_by(ProdDhf.id.asc())
        ).scalars().first()
        if not row:
            row = db.session.execute(
                select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name.like(f"%{DOC_NAME}%")).order_by(ProdDhf.id.asc())
            ).scalars().first()
        return (row.code or "").strip() if row and row.code else ""

    def __fill_revision(self, content, prod_id, version, force=False):
        rev_date = serv_review_util.cover_date(prod_id, DOC_KEY) if prod_id else ""
        reviser = approver = ""
        if prod_id:
            members = db.session.execute(select(ProjectMember).where(ProjectMember.prod_id == prod_id)).scalars().all()
            reviser = next((m.name for m in members if "TPM" in str(m.role or "")), "")
            approver = next((m.name for m in members if "研发负责人" in str(m.role or "")), "")
        for s in (content.get("sections") or []):
            if s.get("ref_type") != "revision":
                continue
            for tbl in (s.get("tables") or []):
                if isinstance(tbl, list) and len(tbl) >= 2 and isinstance(tbl[1], list) and len(tbl[1]) >= 3:
                    row = tbl[1]
                    if force:
                        row[0] = rev_date or ""
                        if len(row) >= 2: row[1] = str(version or "")
                        if len(row) >= 3: row[2] = "首次发布"
                        if len(row) >= 4: row[3] = reviser
                        if len(row) >= 5: row[4] = approver
                    else:
                        if not str(row[0] or "").strip(): row[0] = rev_date
                        if version and len(row) >= 2 and not str(row[1] or "").strip(): row[1] = str(version)
                        if len(row) >= 3 and not str(row[2] or "").strip(): row[2] = "首次发布"
                        if len(row) >= 4 and reviser and not str(row[3] or "").strip(): row[3] = reviser
                        if len(row) >= 5 and approver and not str(row[4] or "").strip(): row[4] = approver
            break
        return content

    def __autofill(self, content, prod_id, product=None, version="", force=False):
        if not isinstance(content, dict):
            return content
        if product and product.name:
            for s in (content.get("sections") or []):
                self.__replace_name(s, BASE_NAME, product.name)
        self.__fill_revision(content, prod_id, version, force=force)
        # 从「软件配置管理计划(SCM)」获取两张配置项清单，注入状态报告对应章节
        # force（切换产品）时跳过，保留模板原值
        if not force:
            self.__pull_sci_from_scm(content, prod_id)
        serv_review_util.ensure_review(
            content, DOC_KEY,
            serv_review_util.review_date(prod_id, serv_review_util.REVIEW_DEFS[DOC_KEY]["name_keywords"]) if prod_id else "",
            prod_id,
        )
        serv_review_util.fill_cover_dates(content, serv_review_util.cover_date(prod_id, DOC_KEY) if prod_id else "", force=force)
        serv_review_util.fill_cover_signers(content, serv_review_util.cover_signers(prod_id, DOC_KEY) if prod_id else {}, force=force)
        return content

    @staticmethod
    def __find_node(sections, pred):
        for s in sections or []:
            if pred(str(s.get("title") or "")):
                return s
            hit = Server.__find_node(s.get("children") or [], pred)
            if hit:
                return hit
        return None

    def __pull_sci_from_scm(self, content, prod_id):
        """从同产品最新的「软件配置管理计划」取①软件配置项清单、②现成软件配置项清单，注入状态报告两个章节。"""
        if not prod_id:
            return content
        scm = db.session.execute(
            select(ScmDoc).where(ScmDoc.product_id == prod_id).order_by(ScmDoc.id.desc())
        ).scalars().first()
        if not scm:
            return content
        scm_content = scm.content if isinstance(scm.content, dict) else {}
        ident = self.__find_node(scm_content.get("sections") or [], lambda t: "标识配置" in t)
        tables = (ident.get("tables") if ident else None) or []
        sci_tbl = tables[0] if len(tables) > 0 else None
        ots_tbl = tables[1] if len(tables) > 1 else None
        secs = content.get("sections") or []
        n1 = self.__find_node(secs, lambda t: t.startswith("软件配置项状态"))
        n2 = self.__find_node(secs, lambda t: t.startswith("现成软件配置状态"))
        if n1 is not None and sci_tbl:
            n1["tables"] = [copy.deepcopy(sci_tbl)]
        if n2 is not None and ots_tbl:
            n2["tables"] = [copy.deepcopy(ots_tbl)]
        return content

    def __to_obj(self, row: ScsDoc, product: Product = None):
        obj = ScsDocObj(**row.dict())
        obj.content = self.__autofill(self.__normalize_content(obj.content), row.product_id, product, row.version)
        if not (obj.file_no or "").strip():
            obj.file_no = self.__dhf_file_no(row.product_id)
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
        return obj

    async def add_scs_doc(self, form: ScsDocForm):
        try:
            sql = select(func.count(ScsDoc.id)).where(ScsDoc.product_id == form.product_id, ScsDoc.version == form.version)
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = ScsDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=ScsDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_scs_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: ScsDoc = db.session.execute(select(ScsDoc).where(ScsDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(ScsDoc.version).where(ScsDoc.product_id == target_pid)).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = ScsDoc(
                product_id=target_pid, version=version,
                file_no=sync_file_no_version(fromdoc.file_no, version),
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            return Resp.resp_ok(data=ScsDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_scs_doc(self, form: ScsDocForm):
        try:
            row: ScsDoc = db.session.execute(select(ScsDoc).where(ScsDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def rebind_product(self, id: int, product_id: int):
        """切换产品：更新 product_id 并强制用新产品信息重新获取封面/修订/产品名后保存，返回新 obj。"""
        try:
            row: ScsDoc = db.session.execute(select(ScsDoc).where(ScsDoc.id == id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            old_product: Product = db.session.execute(select(Product).where(Product.id == row.product_id)).scalars().first() if row.product_id else None
            product: Product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
            if not product:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            content = self.__normalize_content(row.content)
            # 重置含产品名的固定章节为模板原值（恢复基准名 BASE_NAME，避免被旧产品名污染导致后续替换失效）
            # 同时重置软件配置状态表（3章节）为模板原值，切换产品不做替换
            tpl = copy.deepcopy(DEFAULT_SCS_CONTENT) if isinstance(DEFAULT_SCS_CONTENT, dict) else {"sections": []}
            tpl_map = {}
            fixed_titles = {"软件配置项状态(不包括现成软件)", "现成软件配置状态"}
            def collect_tpl(node):
                key = str(node.get("title") or "").strip()
                if BASE_NAME in str(node.get("body") or "") or any(BASE_NAME in str(r) for tbl in (node.get("tables") or []) for r in tbl) or key in fixed_titles:
                    tpl_map[key] = {"body": node.get("body", ""), "tables": copy.deepcopy(node.get("tables", []))}
                for c in (node.get("children") or []):
                    collect_tpl(c)
            for s in (tpl.get("sections") or []):
                collect_tpl(s)
            def reset_fixed(node):
                key = str(node.get("title") or "").strip()
                if key in tpl_map:
                    node["body"] = tpl_map[key]["body"]
                    node["tables"] = copy.deepcopy(tpl_map[key]["tables"])
                for c in (node.get("children") or []):
                    reset_fixed(c)
            for s in (content.get("sections") or []):
                reset_fixed(s)
            # 重置后，把完整版本号更新为新产品版本
            new_full_version = (product.full_version or "").strip()
            def update_version(node):
                if node.get("body") and "完整版本号" in str(node.get("body")):
                    node["body"] = re.sub(r"完整版本号：[^\n]*", f"完整版本号：{new_full_version}", str(node["body"]))
                for c in (node.get("children") or []):
                    update_version(c)
            for s in (content.get("sections") or []):
                update_version(s)
            row.product_id = product_id
            content = self.__autofill(content, product_id, product, row.version, force=True)
            row.content = content
            db.session.commit()
            return Resp.resp_ok(data=self.__to_obj(row, product))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_scs_doc(self, id: int):
        db.session.execute(delete(ScsDoc).where(ScsDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_scs_doc(self, id: int):
        sql = select(ScsDoc, Product).join(Product, ScsDoc.product_id == Product.id).where(ScsDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_scs_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(ScsDoc.product_id == product_id)
        if version:
            wheres.append(ScsDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(ScsDoc.id)).join(Product, ScsDoc.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (select(ScsDoc, Product).join(Product, ScsDoc.product_id == Product.id).where(*wheres)
               .order_by(ScsDoc.id.desc()).offset(page_index * page_size).limit(page_size))
        rows: List[ScsDocObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    # ---------------- 导出 Word（PDP 风格章节树） ----------------
    async def export_scs_doc(self, output, id: int):
        resp = await self.get_scs_doc(id)
        obj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        c = self.__normalize_content(obj.content)
        sections = c.get("sections") or []
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        document.settings.element.append(update_fields)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")
        docx_util.add_page_number_footer(section, obj.file_no or "")

        def add_blank_lines(n):
            for _ in range(max(0, int(n or 0))):
                document.add_paragraph("")

        def write_center_title(text, size=22.0, bold=False):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            docx_util.fonted_txt(p, str(text or ""), font_size=size, bold=bold)

        def add_text(text):
            docx_util.save_txt2docx(str(text or ""), document)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            s = str(text or "")
            if s.startswith("data:image"):
                try:
                    b64 = s.split(",", 1)[1] if "," in s else ""
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = align
                    para.add_run().add_picture(BytesIO(base64.b64decode(b64)), height=Pt(33))
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    return
                except Exception:
                    pass
            cell.text = ""
            for i, line in enumerate(s.split("\n")):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER if align == WD_ALIGN_PARAGRAPH.CENTER else WD_CELL_VERTICAL_ALIGNMENT.TOP

        def _set_fixed_widths(table, widths_dxa):
            table.autofit = False
            tbl_pr = table._tbl.tblPr
            layout = tbl_pr.find(qn("w:tblLayout"))
            if layout is None:
                layout = OxmlElement("w:tblLayout")
                tbl_pr.append(layout)
            layout.set(qn("w:type"), "fixed")
            grid_el = table._tbl.find(qn("w:tblGrid"))
            if grid_el is not None:
                for gc in list(grid_el):
                    grid_el.remove(gc)
                for w in widths_dxa:
                    gc = OxmlElement("w:gridCol")
                    gc.set(qn("w:w"), str(w))
                    grid_el.append(gc)

        def add_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for r_idx, row in enumerate(grid):
                cells = table.add_row().cells
                for c_idx in range(cols):
                    set_cell(cells[c_idx], row[c_idx] if c_idx < len(row) else "", bold=(r_idx == 0))
            first = str(grid[0][0]).strip() if grid and grid[0] else ""
            if first == "SCI名字":
                if cols == 5:
                    _set_fixed_widths(table, [1700, 1100, 2000, 1500, 3000])
                elif cols == 4:
                    _set_fixed_widths(table, [1700, 2200, 1500, 3600])
            document.add_paragraph()

        def add_cover_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for row in grid:
                cells = table.add_row().cells
                for c_idx in range(cols):
                    text = row[c_idx] if c_idx < len(row) else ""
                    set_cell(cells[c_idx], text, bold=(c_idx % 2 == 0), align=WD_ALIGN_PARAGRAPH.CENTER)
                if (str(row[0]).strip() if row else "") == "生效日期" and cols > 2:
                    merged = cells[1]
                    for c_idx in range(2, cols):
                        merged = merged.merge(cells[c_idx])
                    set_cell(merged, row[1] if len(row) > 1 else "", align=WD_ALIGN_PARAGRAPH.CENTER)
            document.add_paragraph()

        def strip_num(title):
            return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()

        def add_body_heading(title, level):
            size = {1: 16.0, 2: 14.0, 3: 12.0}.get(level, 11.0)
            p = document.add_heading("", level=max(1, min(level, 9)))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, title, font_size=size, bold=True)

        def render_body_section(node, level, number=""):
            name = strip_num(node.get("title"))
            heading = f"{number} {name}".strip() if number else name
            add_body_heading(heading, level=max(1, min(level, 9)))
            if (node.get("body") or "").strip():
                add_text(node.get("body"))
            if node.get("ref_type") == "review":
                for t_idx, table in enumerate(node.get("tables") or []):
                    serv_review_util.render_review_grid(document, table, set_cell, merge_col0=(t_idx == 0), merge_full=True)
            else:
                for table in (node.get("tables") or []):
                    add_grid(table)
            idx = 0
            for child in (node.get("children") or []):
                idx += 1
                child_num = f"{number}.{idx}" if number else f"{idx}"
                render_body_section(child, level + 1, child_num)

        cover = next((s for s in sections if s.get("ref_type") == "cover"), None)
        revision = next((s for s in sections if s.get("ref_type") == "revision"), None)
        body = [s for s in sections if s.get("ref_type") not in ("cover", "revision")]

        add_blank_lines(6)
        write_center_title((strip_num(cover.get("title")) if cover else "") or DOC_NAME, size=22.0, bold=True)
        add_blank_lines(4)
        if cover:
            for table in (cover.get("tables") or []):
                add_cover_grid(table)

        document.add_page_break()
        write_center_title("文件修订记录", size=14.0, bold=True)
        add_blank_lines(2)
        if revision:
            for table in (revision.get("tables") or []):
                add_grid(table)

        document.add_page_break()
        write_center_title("目录", size=16.0, bold=True)
        docx_util.insert_toc_field(document)

        document.add_page_break()
        seq = 0
        for node in body:
            if node.get("ref_type") == "review":
                render_body_section(node, 1, "")
            else:
                seq += 1
                render_body_section(node, 1, str(seq))

        document.save(output)
        output.seek(0)
