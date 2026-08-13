#!/usr/bin/env python
# encoding: utf-8

# 产品技术要求服务层，详见 docs/function_docs/56_产品技术要求.md。
# 整份文档以 content(JSON) 的「目录树」结构存储；导出：首页(编号+产品名称)→正文(含附录嵌图)，不含目录页。

import copy
import json
import logging
import os
import re
from typing import List
from sqlalchemy import delete, func, select
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from ..model.product import Product
from ..model.ptr_doc import PtrDoc
from ..model.version_rule import VersionRule
from ..model.prod_runtime_env import ProdRuntimeEnv
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_ptr_doc import PtrDocForm
from ..obj.vobj_ptr_doc import PtrDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from . import serv_review_util
from .serv_utils import new_version, sync_file_no_version
from .serv_utils import docx_util
from .serv_version_rule import DEFAULT_VERSION_RULE
from .serv_doc_file import pick_doc_image_file_row

logger = logging.getLogger(__name__)

# 默认模板内容（整份《产品技术要求》），优先从资源文件加载
DEFAULT_PTR_CONTENT = {"sections": []}
_DEFAULT_CONTENT_FILE = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "src-res", "ptr_default_content.json"
)
try:
    with open(_DEFAULT_CONTENT_FILE, encoding="utf-8") as _f:
        _loaded = json.load(_f)
    if isinstance(_loaded, dict) and isinstance(_loaded.get("sections"), list) and _loaded["sections"]:
        DEFAULT_PTR_CONTENT = _loaded
except Exception:
    logger.exception("加载产品技术要求默认内容资源失败")


class Server(object):

    def __to_obj(self, row: PtrDoc, product: Product = None):
        obj = PtrDocObj(**row.dict())
        obj.content = self.__normalize_content(obj.content)
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
            if not (obj.file_no or "").strip():
                resolved = serv_review_util.resolve_doc_file_no(product.id, obj.file_no, obj.version, "ptr")
                if resolved:
                    obj.file_no = resolved
        return obj

    def __normalize_node(self, node):
        if not isinstance(node, dict):
            return {"title": str(node or ""), "body": "", "tables": [], "children": []}
        result = dict(node)
        result["title"] = str(result.get("title") or "")
        result["body"] = str(result.get("body") or "")
        tables = result.get("tables")
        if not isinstance(tables, list):
            tables = []
        norm_tables = []
        for table in tables:
            if isinstance(table, list):
                norm_tables.append([[str(c) if c is not None else "" for c in (row or [])] for row in table if isinstance(row, list)])
        result["tables"] = norm_tables
        imgs = result.get("images")
        result["images"] = [str(x) for x in imgs if isinstance(x, str) and x] if isinstance(imgs, list) else []
        children = result.get("children")
        result["children"] = [self.__normalize_node(c) for c in children] if isinstance(children, list) else []
        return result

    def __normalize_content(self, content):
        if not isinstance(content, dict) or not isinstance(content.get("sections"), list):
            return copy.deepcopy(DEFAULT_PTR_CONTENT)
        return {"sections": [self.__normalize_node(s) for s in content["sections"]]}

    # ---------------- 自动获取 ----------------
    @staticmethod
    def __build_naming_body(c):
        c = c if isinstance(c, dict) else {}
        items = c.get("items") or []
        lines = [
            "软件版本命名规则为：",
            f"发布版本：{c.get('release_format', '')}",
            f"完整版本：{c.get('full_format', '')}",
            "软件完整版本及说明：",
        ]
        if str(c.get("note_top") or "").strip():
            lines.append(c["note_top"])
        for it in items:
            if not isinstance(it, dict):
                continue
            title = str(it.get("title") or "").strip()
            desc = str(it.get("desc") or "").strip()
            if title or desc:
                lines.append(f"{title}：{desc}")
        if str(c.get("note_bottom") or "").strip():
            lines.append(c["note_bottom"])
        return "\n".join(lines)

    def __collect_autofill(self, prod_id, product, doc_version):
        prod_name = (getattr(product, "name", "") or "").strip()
        full_version = (getattr(product, "full_version", "") or "").strip()
        release_version = (getattr(product, "release_version", "") or "").strip()
        overview = (getattr(product, "overall_desc", "") or "").strip()

        vr = db.session.execute(select(VersionRule).where(VersionRule.id == 1)).scalars().first()
        vr_content = vr.content if vr and isinstance(vr.content, dict) else DEFAULT_VERSION_RULE
        naming_body = self.__build_naming_body(vr_content)

        env = db.session.execute(select(ProdRuntimeEnv).where(ProdRuntimeEnv.prod_id == prod_id)).scalars().first()
        runtime = {}
        if env:
            for k in ("arch", "srv_cpu", "srv_memory", "srv_gpu", "srv_disk", "srv_nic", "srv_os", "srv_cuda",
                      "cli_cpu", "cli_memory", "cli_resolution", "cli_os", "cli_browser", "net_lan", "net_wan"):
                runtime[k] = (getattr(env, k, "") or "").strip()

        return {
            "prod_name": prod_name, "full_version": full_version, "release_version": release_version,
            "naming_body": naming_body, "runtime": runtime, "version": doc_version, "overview": overview,
        }

    @staticmethod
    def __strip_num(title):
        return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()

    @staticmethod
    def __overwrite_col1(table, label_map):
        # 表格按行首标签覆盖第 2 列（源值非空才覆盖）
        for row in table:
            if not isinstance(row, list) or len(row) < 2:
                continue
            key = str(row[0]).strip()
            if key in label_map and str(label_map[key] or "").strip():
                row[1] = label_map[key]

    def __fill_node(self, node, info):
        ref = node.get("ref_type")
        title = self.__strip_num(node.get("title"))
        rt = info.get("runtime") or {}
        if ref == "cover":
            if info["prod_name"]:
                node["body"] = info["prod_name"]
        elif ref == "prod_version" or title == "产品版本":
            if info["full_version"] or info["release_version"]:
                node["body"] = f"软件完整版本：{info['full_version']}\n软件发布版本：{info['release_version']}"
        elif ref == "naming_rule" or title == "版本命名规则":
            if info.get("naming_body"):
                node["body"] = info["naming_body"]
        elif ref == "ui_relation" or "用户界面关系图" in title:
            if info.get("overview"):
                node["body"] = f"图2 用户界面关系图\n{info['overview']}\n图3 主页面图示"
        elif ref == "runtime" or title == "运行环境":
            if rt.get("arch"):
                node["body"] = rt["arch"]
        elif ref == "rt_srv_hw":
            for t in (node.get("tables") or []):
                self.__overwrite_col1(t, {"CPU": rt.get("srv_cpu"), "内存": rt.get("srv_memory"), "GPU": rt.get("srv_gpu"), "硬盘": rt.get("srv_disk"), "网卡": rt.get("srv_nic")})
        elif ref == "rt_srv_sw":
            for t in (node.get("tables") or []):
                if len(t) >= 2 and len(t[1]) >= 3:
                    if str(rt.get("srv_os") or "").strip():
                        t[1][1] = rt["srv_os"]
                    if str(rt.get("srv_cuda") or "").strip():
                        t[1][2] = rt["srv_cuda"]
        elif ref == "rt_client":
            for t in (node.get("tables") or []):
                self.__overwrite_col1(t, {"CPU": rt.get("cli_cpu"), "内存": rt.get("cli_memory"), "显示器分辨率": rt.get("cli_resolution"), "操作系统": rt.get("cli_os"), "浏览器": rt.get("cli_browser")})
        elif ref == "rt_net":
            for t in (node.get("tables") or []):
                if len(t) >= 2 and len(t[1]) >= 3:
                    if str(rt.get("net_lan") or "").strip():
                        t[1][1] = rt["net_lan"]
                    if str(rt.get("net_wan") or "").strip():
                        t[1][2] = rt["net_wan"]
        for child in (node.get("children") or []):
            self.__fill_node(child, info)

    def __autofill_for_export(self, content, obj: PtrDocObj):
        sections = (content or {}).get("sections") or []
        if not obj.product_id:
            return content
        product = db.session.execute(select(Product).where(Product.id == obj.product_id)).scalars().first()
        info = self.__collect_autofill(obj.product_id, product, obj.version)
        for node in sections:
            self.__fill_node(node, info)
        return content

    # ---------------- CRUD ----------------
    async def add_ptr_doc(self, form: PtrDocForm):
        try:
            sql = select(func.count(PtrDoc.id)).where(PtrDoc.product_id == form.product_id, PtrDoc.version == form.version)
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = PtrDoc(**form.dict(exclude_none=True))
            row.id = None
            row.file_no = serv_review_util.resolve_doc_file_no(form.product_id, form.file_no, form.version, "ptr") or None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=PtrDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_ptr_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: PtrDoc = db.session.execute(select(PtrDoc).where(PtrDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(PtrDoc.version).where(PtrDoc.product_id == target_pid)).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = PtrDoc(
                product_id=target_pid,
                version=version,
                file_no=sync_file_no_version(
                    (fromdoc.file_no or "").strip()
                    or serv_review_util.resolve_doc_file_no(target_pid, "", version, "ptr")
                    or "",
                    version,
                ) or None,
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            return Resp.resp_ok(data=PtrDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_ptr_doc(self, form: PtrDocForm):
        try:
            row: PtrDoc = db.session.execute(select(PtrDoc).where(PtrDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_ptr_doc(self, id: int):
        db.session.execute(delete(PtrDoc).where(PtrDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_ptr_doc(self, id: int):
        sql = select(PtrDoc, Product).join(Product, PtrDoc.product_id == Product.id).where(PtrDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_ptr_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(PtrDoc.product_id == product_id)
        if version:
            wheres.append(PtrDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(PtrDoc.id)).join(Product, PtrDoc.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (
            select(PtrDoc, Product)
            .join(Product, PtrDoc.product_id == Product.id)
            .where(*wheres)
            .order_by(PtrDoc.id.desc())
            .offset(page_index * page_size)
            .limit(page_size)
        )
        rows: List[PtrDocObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    # ---------------- 导出 Word ----------------
    async def export_ptr_doc(self, output, id: int):
        resp = await self.get_ptr_doc(id)
        obj: PtrDocObj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        c = self.__autofill_for_export(self.__normalize_content(obj.content), obj)
        sections = c.get("sections") or []
        product_version = obj.product_full_version or ""

        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        # 技术要求首页即正文(无独立封面页)：所有页均加页码，从「第 1 页」起
        docx_util.add_page_number_footer(section, "", skip_first=False)
        def write_center_title(text, size=22.0, bold=False):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, text, font_size=size, bold=bold)

        def add_blank_lines(count):
            for _ in range(max(0, int(count or 0))):
                document.add_paragraph("")

        def add_text(text):
            docx_util.save_txt2docx(str(text or ""), document)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            cell.text = ""
            lines = str(text or "").split("\n")
            for i, line in enumerate(lines):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER if align == WD_ALIGN_PARAGRAPH.CENTER else WD_CELL_VERTICAL_ALIGNMENT.TOP

        def set_grid_widths(table, grid, cols):
            # 用「固定布局 + 按内容比例计算列宽」替代不稳定的 autofit，
            # 让列宽贴合内容、整表不超过页面可用宽度（参考原始 Word 模板）。
            def cell_units(text):
                m = 0
                for line in str(text or "").split("\n"):
                    w = sum(2 if ord(ch) > 127 else 1 for ch in line)
                    m = max(m, w)
                return max(m, 2)

            col_units = []
            for c in range(cols):
                u = 2
                for row in grid:
                    if c < len(row):
                        u = max(u, cell_units(row[c]))
                col_units.append(u)
            # 每个字符宽度单位约 120 dxa，外加单元格左右内边距约 440 dxa
            widths = [u * 120 + 440 for u in col_units]
            sect = document.sections[0]
            usable = int((sect.page_width - sect.left_margin - sect.right_margin) / 635)
            total = sum(widths)
            if total > usable and total > 0:
                scale = usable / total
                widths = [max(int(w * scale), 600) for w in widths]
                total = sum(widths)

            tbl_pr = table._tbl.tblPr
            layout = tbl_pr.find(qn("w:tblLayout"))
            if layout is None:
                layout = OxmlElement("w:tblLayout")
                tbl_pr.append(layout)
            layout.set(qn("w:type"), "fixed")
            tbl_w = tbl_pr.find(qn("w:tblW"))
            if tbl_w is None:
                tbl_w = OxmlElement("w:tblW")
                tbl_pr.append(tbl_w)
            tbl_w.set(qn("w:w"), str(total))
            tbl_w.set(qn("w:type"), "dxa")

            grid_el = table._tbl.find(qn("w:tblGrid"))
            if grid_el is not None:
                grid_cols = grid_el.findall(qn("w:gridCol"))
                for i, gc in enumerate(grid_cols):
                    if i < len(widths):
                        gc.set(qn("w:w"), str(widths[i]))
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    if i >= len(widths):
                        continue
                    tc_pr = cell._tc.get_or_add_tcPr()
                    tc_w = tc_pr.find(qn("w:tcW"))
                    if tc_w is None:
                        tc_w = OxmlElement("w:tcW")
                        tc_pr.append(tc_w)
                    tc_w.set(qn("w:w"), str(widths[i]))
                    tc_w.set(qn("w:type"), "dxa")

        def add_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            for r_idx, row in enumerate(grid):
                cells = table.add_row().cells
                for c_idx in range(cols):
                    set_cell(cells[c_idx], row[c_idx] if c_idx < len(row) else "", bold=(r_idx == 0))
            # 按内容比例设置列宽（不撑满页面，参考原始 Word 模板）
            set_grid_widths(table, grid, cols)
            document.add_paragraph()

        def add_body_heading(title_text, level):
            size = {1: 16.0, 2: 14.0, 3: 12.0}.get(level, 11.0)
            p = document.add_heading("", level=max(1, min(level, 9)))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, title_text, font_size=size, bold=True)

        def add_picture_by_category(category):
            try:
                row = pick_doc_image_file_row(obj.product_id, category, obj.version, product_version)
                if row and row.file_url:
                    docx_util.save_img2docx(row.file_url, document, mw=520, mh=520)
                    return True
            except Exception:
                logger.exception("ptr_add_picture_failed")
            return False

        def add_naming_diagram():
            img_path = os.path.join(os.path.dirname(__file__), "..", "..", "src-res", "assets", "version_naming_rule.png")
            if os.path.exists(img_path):
                try:
                    p = document.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(img_path, width=Inches(3.8))
                    document.add_paragraph()
                except Exception:
                    logger.exception("ptr_naming_diagram_failed")

        # 图题关键词 -> 图表文件管理分类（兼容新旧文档，按正文图题行匹配嵌图）
        caption_cats = [("用户界面关系图", "img_ui"), ("主页面图示", "img_home"), ("物理拓扑图", "img_topo"), ("体系结构图", "img_struct")]

        def render_appendix_body(body):
            # 逐行渲染；遇到「图N xxx」图题行且能匹配到分类时，在该行后插入对应图片
            buf = []

            def flush():
                if buf:
                    add_text("\n".join(buf))
                    buf.clear()

            for ln in (body or "").split("\n"):
                buf.append(ln)
                s = ln.strip()
                if re.match(r"^图\s*\d", s):
                    cat = next((c for kw, c in caption_cats if kw in s), None)
                    if cat:
                        flush()
                        add_picture_by_category(cat)
            flush()

        def render_section(node, level, number=""):
            name = self.__strip_num(node.get("title"))
            ref = node.get("ref_type")
            is_appendix = ref == "appendix"
            heading = name if (is_appendix or not number) else f"{number} {name}".strip()
            add_body_heading(heading, level=max(1, min(level, 9)))
            body = node.get("body") or ""
            if is_appendix:
                if body.strip():
                    render_appendix_body(body)
                elif node.get("img_category"):
                    add_picture_by_category(node["img_category"])
            elif ref == "naming_rule" or name == "版本命名规则":
                lines = body.split("\n")
                cut = next((i for i, ln in enumerate(lines) if ln.strip().startswith("软件完整版本及说明")), -1)
                if cut >= 0:
                    before = "\n".join(lines[:cut + 1])
                    after = "\n".join(lines[cut + 1:])
                    if before.strip():
                        add_text(before)
                    add_naming_diagram()
                    if after.strip():
                        add_text(after)
                else:
                    if body.strip():
                        add_text(body)
                    add_naming_diagram()
            elif body.strip():
                add_text(body)
            # 本章节手动上传的图片（base64），导出嵌入
            for img in (node.get("images") or []):
                try:
                    docx_util.save_img2docx(img, document, mw=520, mh=520)
                except Exception:
                    logger.exception("ptr_node_image_failed")
            for table in (node.get("tables") or []):
                add_grid(table)
            # 子章节编号
            idx = 0
            for child in (node.get("children") or []):
                idx += 1
                child_num = f"{number}.{idx}" if number else f"{idx}"
                render_section(child, level + 1, child_num)

        cover = next((s for s in sections if s.get("ref_type") == "cover"), None)
        body_sections = [s for s in sections if s.get("ref_type") != "cover"]

        # 首页头部：左侧「医疗器械产品技术要求编号：xxx」+ 居中产品名称（参考模板首页），随后紧接正文
        cover_title = (self.__strip_num(cover.get("title")) if cover else "") or "医疗器械产品技术要求"
        no_para = document.add_paragraph()
        no_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        no_para.paragraph_format.line_spacing = 1.5
        docx_util.fonted_txt(no_para, f"{cover_title}编号：", font_size=14.0, bold=True)
        add_blank_lines(1)
        product_name = ((cover.get("body") if cover else "") or obj.product_name or "").strip()
        if product_name:
            write_center_title(product_name, size=20.0, bold=True)
        add_blank_lines(1)

        # 正文（附录不编号）
        body_no = 0
        for node in body_sections:
            if node.get("ref_type") == "appendix":
                render_section(node, 1, "")
            else:
                body_no += 1
                render_section(node, 1, str(body_no))

        document.save(output)
        output.seek(0)
