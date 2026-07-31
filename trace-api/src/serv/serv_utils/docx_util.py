
import logging
import os
import re
import io
import base64
import shutil
import subprocess
import tempfile
from urllib.parse import unquote, urlparse
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import Document
from docx.shared import Inches, Pt
try:
    from PIL import Image
except Exception:
    Image = None
from docx import enum as dox_enum
from docx.shared import RGBColor

from ...obj.tobj_srs_doc import Table

def __fonted_cell(cell, text, font_size=10.5):
    # python-docx 合并单元格时会保留/拼接多个段落；写入前先清空，避免文本重复。
    cell.text = ""
    for paragraph in cell.paragraphs:
        paragraph.alignment = dox_enum.text.WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        fonted_txt(paragraph, text, font_size)

def __apply_table_border(tabx):
    tblBorders = OxmlElement('w:tblBorders')
    for pos in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{pos}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(8))
        border.set(qn('w:color'), '000000')
        border.set(qn('w:space'), '0')
        tblBorders.append(border)
    tblPr = tabx._tbl.tblPr
    tblPr.append(tblBorders)

def __apply_two_col_width(tabx, col_count: int):
    # 仅两列表格：左列窄，右列宽（约 1:2）
    if col_count != 2:
        return
    tabx.autofit = False
    left_w = Inches(1.6)
    right_w = Inches(4.8)
    for row in tabx.rows:
        if len(row.cells) < 2:
            continue
        row.cells[0].width = left_w
        row.cells[1].width = right_w

def __text_visual_len(value: str) -> float:
    txt = str(value or "").strip()
    if not txt:
        return 0.0
    # 中文按双宽估算，英文/数字按单宽估算；防止超长URL把比例拉爆
    score = 0.0
    for ch in txt:
        score += 2.0 if "\u4e00" <= ch <= "\u9fff" else 1.0
    return min(score, 80.0)

def __distribute_col_widths(scores, total_width: float, min_width: float, max_width: float):
    col_count = len(scores)
    if col_count == 0:
        return []
    if total_width <= col_count * min_width:
        return [total_width / col_count for _ in range(col_count)]
    safe_scores = [max(0.1, float(s or 0.1)) for s in scores]
    score_sum = sum(safe_scores) or 1.0
    widths = [total_width * (s / score_sum) for s in safe_scores]
    widths = [max(min_width, min(max_width, w)) for w in widths]
    cur_sum = sum(widths)
    # 迭代收敛到目标总宽
    for _ in range(6):
        diff = total_width - cur_sum
        if abs(diff) < 0.01:
            break
        if diff > 0:
            grow_idx = [i for i, w in enumerate(widths) if w < max_width - 1e-6]
            if not grow_idx:
                break
            base = sum(safe_scores[i] for i in grow_idx) or len(grow_idx)
            for i in grow_idx:
                inc = diff * ((safe_scores[i] / base) if base else (1 / len(grow_idx)))
                widths[i] = min(max_width, widths[i] + inc)
        else:
            shrink_idx = [i for i, w in enumerate(widths) if w > min_width + 1e-6]
            if not shrink_idx:
                break
            base = sum(safe_scores[i] for i in shrink_idx) or len(shrink_idx)
            for i in shrink_idx:
                dec = (-diff) * ((safe_scores[i] / base) if base else (1 / len(shrink_idx)))
                widths[i] = max(min_width, widths[i] - dec)
        cur_sum = sum(widths)
    return widths

def __apply_adaptive_col_width(tabx, headers, rows):
    col_count = len(headers or [])
    if col_count <= 0:
        return
    # 两列表格保留原有版式（封面信息/修订记录等）
    if col_count == 2:
        __apply_two_col_width(tabx, col_count)
        return
    # A4常规页边距下可用宽度约 6.5~6.9 英寸，这里取中值
    total_width = 6.7
    min_width = 0.78
    max_width = 3.6
    scores = []
    for ci, header in enumerate(headers or []):
        hname = str(getattr(header, "name", "") or "")
        header_score = max(2.0, __text_visual_len(hname) * 1.2)
        col_samples = []
        for row in (rows or [])[:40]:
            if isinstance(row, dict):
                col_samples.append(str(row.get(getattr(header, "code", ""), "") or ""))
        sample_max = max((__text_visual_len(x) for x in col_samples), default=0.0)
        sample_avg = (sum(__text_visual_len(x) for x in col_samples) / max(1, len(col_samples))) if col_samples else 0.0
        score = max(header_score, sample_max * 0.9, sample_avg * 1.1, 2.0)
        # URL/备注/描述这类文本列给更高权重，避免被压窄
        if re.search(r"(url|uri|http|路径|地址|链接|备注|说明|描述|内容|参数|详情)", hname, re.I):
            score *= 1.55
        # 编号/序号类列适度收窄
        if re.search(r"(编号|序号|id|编码)", hname, re.I):
            score *= 0.85
        scores.append(score)
    widths = __distribute_col_widths(scores, total_width=total_width, min_width=min_width, max_width=max_width)
    tabx.autofit = False
    for row in tabx.rows:
        for ci, width in enumerate(widths):
            if ci < len(row.cells):
                row.cells[ci].width = Inches(width)

def save_tab2docx(tab: Table,  docx: Document):
    # 优先使用 cells 导出（保留Word导入时的合并单元格结构）
    if tab.cells and len(tab.cells) > 0:
        row_count = len(tab.cells)
        col_count = max((len(row) for row in tab.cells), default=0)
        if row_count > 0 and col_count > 0:
            tabx = docx.add_table(rows=row_count, cols=col_count)
            for ri, row in enumerate(tab.cells):
                for ci, cell in enumerate(row):
                    if cell is None:
                        continue
                    rs = 1 if cell.row_span is None else int(cell.row_span)
                    cs = 1 if cell.col_span is None else int(cell.col_span)
                    if rs == 0 or cs == 0:
                        continue
                    text = str(cell.value or "")
                    end_r = min(row_count - 1, ri + max(1, rs) - 1)
                    end_c = min(col_count - 1, ci + max(1, cs) - 1)
                    target_cell = tabx.cell(ri, ci)
                    if end_r > ri or end_c > ci:
                        target_cell = target_cell.merge(tabx.cell(end_r, end_c))
                    __fonted_cell(target_cell, text)
            header_names = []
            if tab.show_header and row_count > 0:
                try:
                    header_names = [str(tab.cells[0][ci].value or "") for ci in range(col_count)]
                except Exception:
                    header_names = []
            pseudo_headers = tab.headers or [type("Header", (), {"code": f"c{idx}", "name": (header_names[idx] if idx < len(header_names) else f"列{idx+1}")}) for idx in range(col_count)]
            pseudo_rows = []
            data_start = 1 if (tab.show_header and row_count > 0) else 0
            for ri in range(data_start, row_count):
                row_dict = {}
                for ci in range(col_count):
                    code = getattr(pseudo_headers[ci], "code", f"c{ci}")
                    val = ""
                    try:
                        cell = tab.cells[ri][ci]
                        val = "" if cell is None else str(getattr(cell, "value", "") or "")
                    except Exception:
                        val = ""
                    row_dict[code] = val
                pseudo_rows.append(row_dict)
            __apply_adaptive_col_width(tabx, pseudo_headers, pseudo_rows)
            __apply_table_border(tabx)
            empty = docx.add_paragraph()
            empty.paragraph_format.space_after = Pt(20)
            return

    if not tab.headers:
        return

    tabx = docx.add_table(rows=0, cols=len(tab.headers))

    if tab.show_header:
        header_cells = tabx.add_row().cells
        for ci, header in enumerate(tab.headers):
            __fonted_cell(header_cells[ci], header.name)

    for row in tab.rows or []:
        row_cells = tabx.add_row().cells
        for ci, header in enumerate(tab.headers):
            cell_value = row.get(header.code)
            text = str(cell_value) if cell_value is not None else ""
            __fonted_cell(row_cells[ci], text)

    __apply_adaptive_col_width(tabx, tab.headers, tab.rows or [])
    __apply_table_border(tabx)

    empty = docx.add_paragraph()
    empty.paragraph_format.space_after = Pt(20)


def save_img2docx(
    path: str,
    docx: Document,
    mw: int = 600,
    mh: int = 600,
    min_w: int = 0,
    min_h: int = 0,
    target_long: int = 0,
):
    PIXELS_PER_INCH = 96
    SPACE_VALUE = Pt(20)
    image_source = None
    raw_path = str(path or "").strip()
    if raw_path and not raw_path.startswith("data:image/"):
        parsed = urlparse(raw_path)
        clean_path = unquote(parsed.path or raw_path.split("?", 1)[0])
        candidates = [clean_path]
        if clean_path.startswith("/data.trace/"):
            candidates.append(clean_path.lstrip("/"))
        elif clean_path.startswith("data.trace/"):
            candidates.append(clean_path)
        else:
            candidates.append(clean_path.lstrip("/"))
        for candidate in candidates:
            if candidate and os.path.exists(candidate):
                image_source = candidate
                break
    if image_source is None and path and str(path).startswith("data:image/"):
        matched = re.match(r"^data:image/[a-zA-Z0-9.+-]+;base64,(.+)$", str(path), re.S)
        if matched:
            try:
                image_source = io.BytesIO(base64.b64decode(matched.group(1)))
            except Exception:
                image_source = None

    if image_source is None:
        return

    # 结合页面可用宽高做硬限制，避免图片在一页内展示不下
    try:
        section = docx.sections[-1] if docx.sections else None
        if section is not None:
            page_w_in = float(section.page_width) / 914400.0
            page_h_in = float(section.page_height) / 914400.0
            margin_l_in = float(section.left_margin) / 914400.0
            margin_r_in = float(section.right_margin) / 914400.0
            margin_t_in = float(section.top_margin) / 914400.0
            margin_b_in = float(section.bottom_margin) / 914400.0
            avail_w_px = max(120.0, (page_w_in - margin_l_in - margin_r_in) * PIXELS_PER_INCH)
            avail_h_px = max(120.0, (page_h_in - margin_t_in - margin_b_in) * PIXELS_PER_INCH)
            # 预留上下正文与题注空间，避免“单图吃满一页”
            mw = int(min(float(mw), avail_w_px * 0.60))
            mh = int(min(float(mh), avail_h_px * 0.32))
    except Exception:
        pass

    node_para = docx.add_paragraph()
    node_para.alignment = dox_enum.text.WD_ALIGN_PARAGRAPH.CENTER
    node_para.paragraph_format.space_before = SPACE_VALUE
    node_para.paragraph_format.space_after = SPACE_VALUE
    if Image is None:
        # 本地环境未安装 Pillow 时，按最大宽度插入，避免导出失败
        node_para.add_run().add_picture(image_source, width=Inches(mw / PIXELS_PER_INCH))
        return

    try:
        if isinstance(image_source, io.BytesIO):
            image_source.seek(0)
            with Image.open(image_source) as img:
                ow, oh = img.size
            image_source.seek(0)
        else:
            with Image.open(image_source) as img:
                ow, oh = img.size
        max_scale = min(mw / ow, mh / oh)
        min_scale = 0.0
        if min_w > 0 or min_h > 0:
            min_scale = max(
                (min_w / ow) if min_w > 0 else 0.0,
                (min_h / oh) if min_h > 0 else 0.0,
            )
        if min_scale > max_scale:
            # 极端长宽比时，优先保证不超出最大边界
            min_scale = max_scale

        if target_long and target_long > 0:
            base_scale = target_long / max(ow, oh)
        else:
            base_scale = 1.0
        scale = min(max(base_scale, min_scale), max_scale)
        img_w = (ow * scale) / PIXELS_PER_INCH
        img_h = (oh * scale) / PIXELS_PER_INCH
        node_para.add_run().add_picture(image_source, width=Inches(img_w), height=Inches(img_h))
    except Exception:
        # 尺寸探测失败时退化为固定宽度，确保图片仍可导出
        if isinstance(image_source, io.BytesIO):
            image_source.seek(0)
        node_para.add_run().add_picture(image_source, width=Inches(mw / PIXELS_PER_INCH))

def save_title2docx(title: str, docx: Document, level: int = 1, font_size=10.5):
    # 按文档规范设置标题：一级三号加粗，二级四号加粗，三级及以下五号常规
    size_map = {1: 16.0, 2: 14.0, 3: 10.5}
    size = size_map.get(level, 10.5)
    is_bold = level <= 2
    # 使用Heading样式，便于Word目录域(TOC)识别并生成可点击目录
    node_para = docx.add_heading("", level=max(1, min(level, 9)))
    node_para.alignment = dox_enum.text.WD_ALIGN_PARAGRAPH.LEFT
    node_para.paragraph_format.first_line_indent = Pt(0)
    node_para.paragraph_format.left_indent = Pt(0)
    node_para.paragraph_format.right_indent = Pt(0)
    node_para.paragraph_format.line_spacing = 1.5
    node_para.paragraph_format.space_before = Pt(0)
    node_para.paragraph_format.space_after = Pt(0)
    fonted_txt(node_para, title, font_size=size, bold=is_bold)

def save_txt2docx(text: str, docx: Document, font_size=10.5):
    texts = (text or "").splitlines()
    for text in texts:
        text = text.strip()
        if not text:
            continue
        node_para = docx.add_paragraph()
        node_para.alignment = dox_enum.text.WD_ALIGN_PARAGRAPH.LEFT
        node_para.paragraph_format.first_line_indent = Pt(font_size*2)
        node_para.paragraph_format.left_indent = Pt(0)
        node_para.paragraph_format.right_indent = Pt(0)
        node_para.paragraph_format.line_spacing = 1.5
        node_para.paragraph_format.space_before = Pt(0)
        node_para.paragraph_format.space_after = Pt(0)
        fonted_txt(node_para, text, font_size)

def fonted_txt(node_para, text, font_size=10.5, bold=False):
    parts = re.findall(r'([\u4e00-\u9fa5]+|[^\u4e00-\u9fa5]+)', text or "")
    node_para.paragraph_format.space_before = Pt(0)
    node_para.paragraph_format.space_after = Pt(0)
    for part in parts:
        run = node_para.add_run(part)
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.italic = False
        run.font.bold = bool(bold)
        font_name="宋体" if re.match(r'[\u4e00-\u9fa5]', part) else "Times New Roman"
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_page_number_footer(section, file_no: str = "", skip_first: bool = True):
    """统一给正式报告类 Word 导出添加页码：
    - skip_first=True(默认，适用于有独立封面页的文档)：封面页眉页脚与其余页不同，
      封面不显示页码；页码从 0 起(封面为第0页不显示)，封面之后的第一页正好显示「第 1 页」；
      若提供 file_no，则同步写入首页页眉(右对齐)，避免封面丢失页眉。
    - skip_first=False(适用于首页即正文、无独立封面的文档)：所有页均显示页码，从「第 1 页」起。
    要求 skip_first=True 时调用方已给常规 section.header 写好文件编号(非首页页眉)。"""
    if OxmlElement is None or qn is None:
        return
    align_right = dox_enum.text.WD_ALIGN_PARAGRAPH.RIGHT
    align_center = dox_enum.text.WD_ALIGN_PARAGRAPH.CENTER

    if skip_first:
        section.different_first_page_header_footer = True

        if file_no:
            fp_header = section.first_page_header
            hp = fp_header.paragraphs[0] if fp_header.paragraphs else fp_header.add_paragraph()
            hp.alignment = align_right
            fonted_txt(hp, str(file_no or ""))

        pg_num_type = OxmlElement("w:pgNumType")
        pg_num_type.set(qn("w:start"), "1")
        section._sectPr.append(pg_num_type)

    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = align_center
    fonted_txt(fp, "第 ")
    page_run = fp.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    page_instr = OxmlElement("w:instrText")
    page_instr.set(qn("xml:space"), "preserve")
    page_instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    page_run._r.append(fld_begin)
    page_run._r.append(page_instr)
    page_run._r.append(fld_separate)
    fonted_txt(fp, "1")
    page_end_run = fp.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_end_run._r.append(fld_end)
    fonted_txt(fp, " 页")

logger = logging.getLogger(__name__)

def insert_toc_field(docx: Document, outline_levels: str = "1-4"):
    """插入 Word 目录域，打开文档后自动更新标题与页码。"""
    if OxmlElement is None or qn is None:
        return
    # 设置 updateFields=true，Word 打开时弹"更新域"对话框
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    docx.settings.element.append(update_fields)
    p = docx.add_paragraph()
    run_begin = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' TOC \\o "{outline_levels}" \\h \\u '
    run_begin._r.append(instr)
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    run_begin._r.append(fld_separate)
    p.add_run("请打开文档后右键目录，选择「更新域」以生成页码")
    run_end = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)


def enable_update_fields_on_open(docx: Document):
    if OxmlElement is None or qn is None:
        return
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    docx.settings.element.append(update_fields)


def fill_toc_cache(docx: Document):
    """在文档保存前，遍历所有 Heading 段落，把标题文本回填到目录域的缓存区。
    在 separate 和 end 之间插入带制表位（前导点）+ 页码占位的 run，用换行符分隔。
    不插入新段落（避免两遍目录），只修改目录域段落内部的 run。
    Word 打开时仍可右键目录更新为带真实页码的目录。"""
    if OxmlElement is None or qn is None:
        return
    # 收集所有标题段落（Heading 1-4）
    headings = []
    for para in docx.paragraphs:
        style_name = (para.style.name or "") if para.style else ""
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading", "").replace(" ", "").strip() or "1")
            except ValueError:
                level = 1
            if 1 <= level <= 4:
                text = para.text.strip()
                if text:
                    headings.append((level, text))
    if not headings:
        return
    # 找到目录域段落（含 TOC instrText 的段落）
    for para in docx.paragraphs:
        p_elem = para._p
        # 检查是否含 TOC instrText
        has_toc = False
        separate_fld = None  # fldChar separate 元素
        end_fld = None  # fldChar end 元素
        for fld in p_elem.iter(qn("w:fldChar")):
            ft = fld.get(qn("w:fldCharType"), "")
            if ft == "separate":
                separate_fld = fld
            elif ft == "end":
                end_fld = fld
        for instr in p_elem.iter(qn("w:instrText")):
            if instr.text and "TOC" in instr.text:
                has_toc = True
                break
        if not has_toc or separate_fld is None or end_fld is None:
            continue
        # separate 和 end 之间的 run 是占位文本，需要替换
        # 找到 separate 所在的 run，然后删除到 end 所在的 run 之间的所有 run
        sep_run = separate_fld.getparent()  # 所在的 w:r 元素
        end_run = end_fld.getparent()  # 所在的 w:r 元素
        # 获取段落的所有子元素
        children = list(p_elem)
        sep_idx = children.index(sep_run)
        end_idx = children.index(end_run)
        # 删除 separate run 和 end run 之间的元素（占位文本 run）
        for child in children[sep_idx + 1:end_idx]:
            p_elem.remove(child)
        # 在 separate run 后面插入标题条目
        insert_idx = sep_idx + 1
        for i, (level, text) in enumerate(headings):
            # 每个标题前加换行符（第一个不加）
            if i > 0:
                br_run = OxmlElement("w:r")
                br = OxmlElement("w:br")
                br_run.append(br)
                p_elem.insert(insert_idx, br_run)
                insert_idx += 1
            # 标题文本 run（带缩进）
            title_run = OxmlElement("w:r")
            title_t = OxmlElement("w:t")
            title_t.set(qn("xml:space"), "preserve")
            indent = "    " * (level - 1)
            title_t.text = f"{indent}{text}"
            title_run.append(title_t)
            p_elem.insert(insert_idx, title_run)
            insert_idx += 1
            # 制表符 run
            tab_run = OxmlElement("w:r")
            tab = OxmlElement("w:tab")
            tab_run.append(tab)
            p_elem.insert(insert_idx, tab_run)
            insert_idx += 1
            # 页码占位 run
            page_run = OxmlElement("w:r")
            page_t = OxmlElement("w:t")
            page_t.set(qn("xml:space"), "preserve")
            page_t.text = "1"
            page_run.append(page_t)
            p_elem.insert(insert_idx, page_run)
            insert_idx += 1
        # 给目录域段落加制表位（右对齐+前导点）
        ppr = p_elem.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            p_elem.insert(0, ppr)
        tabs = ppr.find(qn("w:tabs"))
        if tabs is None:
            tabs = OxmlElement("w:tabs")
            ppr.append(tabs)
            tab_def = OxmlElement("w:tab")
            tab_def.set(qn("w:val"), "right")
            tab_def.set(qn("w:leader"), "dot")
            tab_def.set(qn("w:pos"), "8306")
            tabs.append(tab_def)
        return


def refresh_docx_toc_with_libreoffice(output_stream) -> bool:
    """用 LibreOffice 刷新目录域，导出文件即可带正确页码。"""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        logger.warning("LibreOffice not found, skip docx TOC refresh")
        return False
    try:
        output_stream.seek(0)
        source_bytes = output_stream.read()
        with tempfile.TemporaryDirectory(prefix="docx_toc_refresh_") as tmpdir:
            input_path = os.path.join(tmpdir, "input.docx")
            output_path = os.path.join(tmpdir, "output.docx")
            profile_dir = os.path.join(tmpdir, "lo_profile")
            script_path = os.path.join(tmpdir, "refresh_toc.py")
            with open(input_path, "wb") as f:
                f.write(source_bytes)
            script = r'''
import sys
import time
import uno
from com.sun.star.beans import PropertyValue

input_path = sys.argv[1]
output_path = sys.argv[2]

def prop(name, value):
    item = PropertyValue()
    item.Name = name
    item.Value = value
    return item

local_ctx = uno.getComponentContext()
resolver = local_ctx.ServiceManager.createInstanceWithContext(
    "com.sun.star.bridge.UnoUrlResolver",
    local_ctx,
)
ctx = None
last_error = None
for _ in range(60):
    try:
        ctx = resolver.resolve("uno:socket,host=127.0.0.1,port=2002;urp;StarOffice.ComponentContext")
        break
    except Exception as exc:
        last_error = exc
        time.sleep(0.5)
if ctx is None:
    raise RuntimeError(f"connect libreoffice failed: {last_error}")

desktop = ctx.ServiceManager.createInstanceWithContext("com.sun.star.frame.Desktop", ctx)
doc = desktop.loadComponentFromURL(
    uno.systemPathToFileUrl(input_path),
    "_blank",
    0,
    (
        prop("Hidden", True),
        prop("ReadOnly", False),
        prop("UpdateDocMode", 3),
    ),
)
if doc is None:
    raise RuntimeError("load docx failed")
try:
    indexes = doc.getDocumentIndexes()
    for idx in range(indexes.getCount()):
        indexes.getByIndex(idx).update()
    fields = doc.getTextFields()
    enum = fields.createEnumeration()
    while enum.hasMoreElements():
        field = enum.nextElement()
        try:
            field.update()
        except Exception:
            pass
    doc.storeAsURL(
        uno.systemPathToFileUrl(output_path),
        (
            prop("FilterName", "Office Open XML Text"),
            prop("Overwrite", True),
        ),
    )
finally:
    doc.close(True)
'''
            with open(script_path, "w", encoding="utf-8") as f:
                f.write(script)
            server = subprocess.Popen(
                [
                    soffice,
                    "--headless",
                    "--nologo",
                    "--nodefault",
                    "--nofirststartwizard",
                    "--nolockcheck",
                    f"-env:UserInstallation=file://{profile_dir}",
                    "--accept=socket,host=127.0.0.1,port=2002;urp;StarOffice.ServiceManager",
                ],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
            try:
                env = os.environ.copy()
                dist_pkg = "/usr/lib/python3/dist-packages"
                env["PYTHONPATH"] = f"{dist_pkg}:{env.get('PYTHONPATH', '')}"
                subprocess.run(
                    ["python3", script_path, input_path, output_path],
                    check=True,
                    timeout=120,
                    env=env,
                )
            finally:
                server.terminate()
                try:
                    server.wait(timeout=10)
                except subprocess.TimeoutExpired:
                    server.kill()
            if not os.path.exists(output_path) or os.path.getsize(output_path) <= 0:
                output_stream.seek(0)
                return False
            with open(output_path, "rb") as f:
                refreshed = f.read()
            output_stream.seek(0)
            output_stream.truncate(0)
            output_stream.write(refreshed)
            output_stream.seek(0)
            return True
    except Exception:
        logger.exception("refresh docx TOC with LibreOffice failed")
        output_stream.seek(0)
        return False
