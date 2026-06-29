#!/usr/bin/env python
# encoding: utf-8

# 产品标签样稿服务层，详见 docs/function_docs/58_产品标签样稿.md。
# 整份文档以 content(JSON) 的「章节树」结构存储；标签（图1/图2/图3）以文本表格表示。
# 导出：封面(标题+封面信息表)→修订记录→标签表格→技术要求，参考原始 Word 版式。

import copy
import json
import logging
import os
import re
from typing import List
from sqlalchemy import delete, func, select
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from ..model.product import Product
from ..model.label_doc import LabelDoc
from ..model.company_info import CompanyInfo
from ..model.project_timeline import ProjectTimelineRow, ProjectTimelineCell
from ..model.project_member import ProjectMember
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_label_doc import LabelDocForm
from ..obj.vobj_label_doc import LabelDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from .serv_utils import new_version
from .serv_utils import docx_util

logger = logging.getLogger(__name__)

# 默认模板内容（整份《产品标签样稿》），优先从资源文件加载
DEFAULT_LABEL_CONTENT = {"sections": []}
_DEFAULT_CONTENT_FILE = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "src-res", "label_default_content.json"
)
try:
    with open(_DEFAULT_CONTENT_FILE, encoding="utf-8") as _f:
        _loaded = json.load(_f)
    if isinstance(_loaded, dict) and isinstance(_loaded.get("sections"), list) and _loaded["sections"]:
        DEFAULT_LABEL_CONTENT = _loaded
except Exception:
    logger.exception("加载产品标签样稿默认内容资源失败")


class Server(object):

    def __to_obj(self, row: LabelDoc, product: Product = None):
        obj = LabelDocObj(**row.dict())
        obj.content = self.__normalize_content(obj.content)
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
        return obj

    def __normalize_node(self, node):
        if not isinstance(node, dict):
            return {"title": str(node or ""), "body": "", "tables": [], "children": []}
        result = dict(node)
        result["title"] = str(result.get("title") or "")
        result["body"] = str(result.get("body") or "")
        tables = result.get("tables")
        if not isinstance(tables, list):
            tables = []
        norm_tables = []
        for table in tables:
            if isinstance(table, list):
                norm_tables.append([[str(c) if c is not None else "" for c in (row or [])] for row in table if isinstance(row, list)])
        result["tables"] = norm_tables
        children = result.get("children")
        result["children"] = [self.__normalize_node(c) for c in children] if isinstance(children, list) else []
        return result

    def __normalize_content(self, content):
        if not isinstance(content, dict) or not isinstance(content.get("sections"), list):
            return copy.deepcopy(DEFAULT_LABEL_CONTENT)
        return {"sections": [self.__normalize_node(s) for s in content["sections"]]}

    # ---------------- 自动获取 ----------------
    @staticmethod
    def __strip_num(title):
        return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()

    @staticmethod
    def __overwrite_col1(table, label_map):
        # 表格按行首标签覆盖第 2 列（源值非空才覆盖）
        for row in table:
            if not isinstance(row, list) or len(row) < 2:
                continue
            key = str(row[0]).strip()
            if key in label_map and str(label_map[key] or "").strip():
                row[1] = label_map[key]

    def __collect_label_map(self, product):
        type_code = (getattr(product, "type_code", "") or "").strip()
        prod_name = (getattr(product, "name", "") or "").strip()
        full_version = (getattr(product, "full_version", "") or "").strip()
        release_version = (getattr(product, "release_version", "") or "").strip()
        # 公司信息取「公司基本信息」首条记录（按 id 升序）作为注册人主体
        company = db.session.execute(select(CompanyInfo).order_by(CompanyInfo.id)).scalars().first()
        label_map = {
            "产品型号": type_code,
            "英文名称": type_code,
            "完整版本": full_version,
            "发布版本": release_version,
            "软件名称": prod_name,
        }
        if company:
            label_map.update({
                "注册人": (company.registrant or "").strip(),
                "住所": (company.address or "").strip(),
                "受托生产企业": (company.manufacturer or "").strip(),
                "生产地址": (company.production_address or "").strip(),
                "生产许可证编号": (company.production_license_no or "").strip(),
                "联系电话": (company.contact_phone or "").strip(),
            })
        return label_map

    def __fill_node(self, node, label_map):
        # 封面表（编写部门/文件版本等）不自动覆盖；其余章节按标签映射覆盖表格第 2 列
        if node.get("ref_type") != "cover":
            for table in (node.get("tables") or []):
                self.__overwrite_col1(table, label_map)
        for child in (node.get("children") or []):
            self.__fill_node(child, label_map)

    def __compute_revision_info(self, prod_id, doc_version):
        # 修改日期=时间线中含「标签」的最早日期行；修订人/批准人取参与人员（产品经理/产品部负责人）
        def to_int(v):
            digits = re.sub(r"[^\d]", "", str(v or ""))
            return int(digits) if digits else None

        tl_rows = db.session.execute(
            select(ProjectTimelineRow).where(ProjectTimelineRow.prod_id == prod_id)
        ).scalars().all()
        cell_map = {}
        if tl_rows:
            for c in db.session.execute(
                select(ProjectTimelineCell).where(ProjectTimelineCell.row_id.in_([r.id for r in tl_rows]))
            ).scalars().all():
                cell_map.setdefault(c.row_id, []).append(c.output_result or "")
        date_rows = [r for r in tl_rows if (r.row_type or "date") == "date" and to_int(r.year) and to_int(r.month)]

        def date_key(r):
            return to_int(r.year) * 10000 + to_int(r.month) * 100 + (to_int(r.day) or 0)

        file_rows = [r for r in date_rows if any("标签" in str(v or "") for v in cell_map.get(r.id, []))]
        file_date = ""
        if file_rows:
            fr = min(file_rows, key=date_key)
            file_date = f"{to_int(fr.year)}年{to_int(fr.month)}月{to_int(fr.day)}日"

        members = db.session.execute(select(ProjectMember).where(ProjectMember.prod_id == prod_id)).scalars().all()

        def find_member(pred):
            for mem in members:
                if pred(str(mem.role or "")):
                    return (mem.name or "").strip()
            return ""

        pm = find_member(lambda r: "产品经理" in r)
        approver = find_member(lambda r: "负责人" in r and "产品" in r)
        return {"file_date": file_date, "version": doc_version, "pm": pm, "approver": approver}

    def __fill_revision(self, node, info):
        # 文件修订记录首行：修改日期/版本号/修订说明/修订人/批准人（仅填空、不覆盖已填）
        ref = node.get("ref_type")
        title = self.__strip_num(node.get("title"))
        if ref == "revision" or title == "文件修订记录":
            tables = node.get("tables") or []
            if tables and isinstance(tables[0], list):
                t = tables[0]
                cols = len(t[0]) if t and t[0] else 5
                while len(t) < 6:
                    t.append([""] * cols)
                row = t[1]
                while len(row) < 5:
                    row.append("")

                def set_if(i, val):
                    if val and not str(row[i] or "").strip():
                        row[i] = val

                set_if(0, info.get("file_date"))
                set_if(1, info.get("version"))
                if not str(row[2] or "").strip():
                    row[2] = "首次发布"
                set_if(3, info.get("pm"))
                set_if(4, info.get("approver"))
        for child in (node.get("children") or []):
            self.__fill_revision(child, info)

    def __autofill_for_export(self, content, obj: LabelDocObj):
        sections = (content or {}).get("sections") or []
        if not obj.product_id:
            return content
        product = db.session.execute(select(Product).where(Product.id == obj.product_id)).scalars().first()
        label_map = self.__collect_label_map(product)
        rev_info = self.__compute_revision_info(obj.product_id, obj.version)
        for node in sections:
            self.__fill_node(node, label_map)
            self.__fill_revision(node, rev_info)
        return content

    # ---------------- CRUD ----------------
    async def add_label_doc(self, form: LabelDocForm):
        try:
            sql = select(func.count(LabelDoc.id)).where(LabelDoc.product_id == form.product_id, LabelDoc.version == form.version)
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = LabelDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=LabelDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_label_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: LabelDoc = db.session.execute(select(LabelDoc).where(LabelDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(LabelDoc.version).where(LabelDoc.product_id == target_pid)).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = LabelDoc(
                product_id=target_pid,
                version=version,
                file_no=fromdoc.file_no,
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            return Resp.resp_ok(data=LabelDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_label_doc(self, form: LabelDocForm):
        try:
            row: LabelDoc = db.session.execute(select(LabelDoc).where(LabelDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_label_doc(self, id: int):
        db.session.execute(delete(LabelDoc).where(LabelDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_label_doc(self, id: int):
        sql = select(LabelDoc, Product).join(Product, LabelDoc.product_id == Product.id).where(LabelDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_label_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(LabelDoc.product_id == product_id)
        if version:
            wheres.append(LabelDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(LabelDoc.id)).join(Product, LabelDoc.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (
            select(LabelDoc, Product)
            .join(Product, LabelDoc.product_id == Product.id)
            .where(*wheres)
            .order_by(LabelDoc.id.desc())
            .offset(page_index * page_size)
            .limit(page_size)
        )
        rows: List[LabelDocObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    # ---------------- 导出 Word ----------------
    async def export_label_doc(self, output, id: int):
        resp = await self.get_label_doc(id)
        obj: LabelDocObj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        c = self.__autofill_for_export(self.__normalize_content(obj.content), obj)
        sections = c.get("sections") or []

        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        # 页眉：文件编号右对齐（参考产品开发计划）
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")

        def write_center_title(text, size=22.0, bold=True):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.5
            docx_util.fonted_txt(p, text, font_size=size, bold=bold)

        def add_text(text):
            docx_util.save_txt2docx(str(text or ""), document)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            cell.text = ""
            lines = str(text or "").split("\n")
            for i, line in enumerate(lines):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER if align == WD_ALIGN_PARAGRAPH.CENTER else WD_CELL_VERTICAL_ALIGNMENT.TOP

        def set_grid_widths(table, grid, cols):
            # 固定布局 + 按内容比例计算列宽（不撑满页面，参考原始 Word）
            def cell_units(text):
                m = 0
                for line in str(text or "").split("\n"):
                    w = sum(2 if ord(ch) > 127 else 1 for ch in line)
                    m = max(m, w)
                return max(m, 2)

            col_units = []
            for ci in range(cols):
                u = 2
                for row in grid:
                    if ci < len(row):
                        u = max(u, cell_units(row[ci]))
                col_units.append(u)
            widths = [u * 120 + 440 for u in col_units]
            sect = document.sections[0]
            usable = int((sect.page_width - sect.left_margin - sect.right_margin) / 635)
            total = sum(widths)
            if total > usable and total > 0:
                scale = usable / total
                widths = [max(int(w * scale), 600) for w in widths]
                total = sum(widths)

            tbl_pr = table._tbl.tblPr
            layout = tbl_pr.find(qn("w:tblLayout"))
            if layout is None:
                layout = OxmlElement("w:tblLayout")
                tbl_pr.append(layout)
            layout.set(qn("w:type"), "fixed")
            tbl_w = tbl_pr.find(qn("w:tblW"))
            if tbl_w is None:
                tbl_w = OxmlElement("w:tblW")
                tbl_pr.append(tbl_w)
            tbl_w.set(qn("w:w"), str(total))
            tbl_w.set(qn("w:type"), "dxa")

            grid_el = table._tbl.find(qn("w:tblGrid"))
            if grid_el is not None:
                for i, gc in enumerate(grid_el.findall(qn("w:gridCol"))):
                    if i < len(widths):
                        gc.set(qn("w:w"), str(widths[i]))
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    if i >= len(widths):
                        continue
                    tc_pr = cell._tc.get_or_add_tcPr()
                    tc_w = tc_pr.find(qn("w:tcW"))
                    if tc_w is None:
                        tc_w = OxmlElement("w:tcW")
                        tc_pr.append(tc_w)
                    tc_w.set(qn("w:w"), str(widths[i]))
                    tc_w.set(qn("w:type"), "dxa")

        def add_grid(grid, header=True):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            for r_idx, row in enumerate(grid):
                cells = table.add_row().cells
                for c_idx in range(cols):
                    set_cell(cells[c_idx], row[c_idx] if c_idx < len(row) else "", bold=(header and r_idx == 0))
            set_grid_widths(table, grid, cols)
            document.add_paragraph()

        def add_heading(title_text, level):
            size = {1: 15.0, 2: 13.0, 3: 12.0}.get(level, 11.0)
            p = document.add_heading("", level=max(1, min(level, 9)))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, title_text, font_size=size, bold=True)

        def add_seal_image():
            # 封口贴合格标签为固定图（QC PASS），写死嵌入
            img_path = os.path.join(os.path.dirname(__file__), "..", "..", "src-res", "assets", "qc_pass.png")
            if os.path.exists(img_path):
                try:
                    p = document.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(img_path, width=Inches(1.2))
                    document.add_paragraph()
                except Exception:
                    logger.exception("label_seal_image_failed")

        def render_section(node, level):
            name = self.__strip_num(node.get("title"))
            header = node.get("ref_type") in ("revision",) or not node.get("ref_type")
            add_heading(name, level=level)
            body = node.get("body") or ""
            if body.strip():
                add_text(body)
            for table in (node.get("tables") or []):
                add_grid(table, header=header)
            for child in (node.get("children") or []):
                render_section(child, level + 1)

        # ---------- 版式辅助（参考产品开发计划 + 图1 方框样式） ----------
        def add_blank_lines(count):
            for _ in range(max(0, int(count or 0))):
                document.add_paragraph("")

        def usable_dxa():
            sect = document.sections[0]
            return int((sect.page_width - sect.left_margin - sect.right_margin) / 635)

        def set_table_borders(table, val="single", sz=8, color="000000"):
            tbl_pr = table._tbl.tblPr
            borders = tbl_pr.find(qn("w:tblBorders"))
            if borders is None:
                borders = OxmlElement("w:tblBorders")
                tbl_pr.append(borders)
            for pos in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = borders.find(qn(f"w:{pos}"))
                if el is None:
                    el = OxmlElement(f"w:{pos}")
                    borders.append(el)
                el.set(qn("w:val"), val)
                el.set(qn("w:sz"), str(sz))
                el.set(qn("w:color"), color)
                el.set(qn("w:space"), "0")

        def set_fixed_width(table, widths):
            total = sum(widths)
            tbl_pr = table._tbl.tblPr
            layout = tbl_pr.find(qn("w:tblLayout"))
            if layout is None:
                layout = OxmlElement("w:tblLayout")
                tbl_pr.append(layout)
            layout.set(qn("w:type"), "fixed")
            tbl_w = tbl_pr.find(qn("w:tblW"))
            if tbl_w is None:
                tbl_w = OxmlElement("w:tblW")
                tbl_pr.append(tbl_w)
            tbl_w.set(qn("w:w"), str(total))
            tbl_w.set(qn("w:type"), "dxa")
            grid_el = table._tbl.find(qn("w:tblGrid"))
            if grid_el is not None:
                for i, gc in enumerate(grid_el.findall(qn("w:gridCol"))):
                    if i < len(widths):
                        gc.set(qn("w:w"), str(widths[i]))
            for row in table.rows:
                for i, c in enumerate(row.cells):
                    if i >= len(widths):
                        continue
                    tc_pr = c._tc.get_or_add_tcPr()
                    tc_w = tc_pr.find(qn("w:tcW"))
                    if tc_w is None:
                        tc_w = OxmlElement("w:tcW")
                        tc_pr.append(tc_w)
                    tc_w.set(qn("w:w"), str(widths[i]))
                    tc_w.set(qn("w:type"), "dxa")

        def fig_caption(title):
            write_center_title(title, size=14.0, bold=True)

        def kv_of(node):
            kv = {}
            for table in (node.get("tables") or []):
                for row in table:
                    if isinstance(row, list) and len(row) >= 2:
                        kv[str(row[0]).strip()] = str(row[1] or "").strip()
            return kv

        def add_bracket_para(container, label, value, size=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            p = container.add_paragraph()
            p.alignment = align
            p.paragraph_format.line_spacing = 1.4
            docx_util.fonted_txt(p, f"【{label}】{value}", font_size=size, bold=bold)

        def add_cover_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for row in grid:
                cells = table.add_row().cells
                for c_idx in range(cols):
                    text = row[c_idx] if c_idx < len(row) else ""
                    set_cell(cells[c_idx], text, bold=(c_idx % 2 == 0), align=WD_ALIGN_PARAGRAPH.CENTER)
                if (str(row[0]).strip() if row else "") == "生效日期" and cols > 2:
                    merged = cells[1]
                    for c_idx in range(2, cols):
                        merged = merged.merge(cells[c_idx])
                    set_cell(merged, row[1] if len(row) > 1 else "", align=WD_ALIGN_PARAGRAPH.CENTER)
            document.add_paragraph()

        def render_label_product(node):
            # 图1：实线方框，内含 软件名称标题 +【字段】值 + 虚线 UDI 条形码框 + 结尾文案
            kv = kv_of(node)
            total = int(usable_dxa() * 0.7)
            box = document.add_table(rows=1, cols=1)
            box.alignment = WD_TABLE_ALIGNMENT.CENTER
            box.autofit = False
            set_table_borders(box, "single", 12)
            set_fixed_width(box, [total])
            cell = box.cell(0, 0)
            cell.text = ""
            tp = cell.paragraphs[0]
            tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tp.paragraph_format.line_spacing = 1.6
            docx_util.fonted_txt(tp, kv.get("软件名称", ""), font_size=15.0, bold=True)
            # 三行两列字段
            pair_w = max(total - 200, 400)
            ptbl = cell.add_table(rows=0, cols=2)
            ptbl.autofit = False
            set_table_borders(ptbl, "none", 0)
            # 清零嵌套表格的缩进与单元格边距，使左列与下方单列字段左对齐
            tblPr = ptbl._tbl.tblPr
            ind = OxmlElement("w:tblInd")
            ind.set(qn("w:w"), "0")
            ind.set(qn("w:type"), "dxa")
            tblPr.append(ind)
            mar = OxmlElement("w:tblCellMar")
            for side in ("top", "left", "bottom", "right"):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:w"), "0")
                el.set(qn("w:type"), "dxa")
                mar.append(el)
            tblPr.append(mar)
            for l, r in [("产品型号", "完整版本"), ("成品序列号", "发布版本"), ("生产日期", "使用期限")]:
                cells = ptbl.add_row().cells
                set_cell(cells[0], f"【{l}】{kv.get(l, '')}")
                set_cell(cells[1], f"【{r}】{kv.get(r, '')}")
            set_fixed_width(ptbl, [pair_w // 2, pair_w - pair_w // 2])
            # 单列字段
            for label in ["产品注册证号", "注册人", "住所", "受托生产企业", "生产地址", "生产许可证编号", "联系电话", "UDI"]:
                add_bracket_para(cell, label, kv.get(label, ""))
            # 虚线 UDI 条形码占位框
            cell.add_paragraph()
            udi = cell.add_table(rows=1, cols=1)
            udi.alignment = WD_TABLE_ALIGNMENT.CENTER
            udi.autofit = False
            set_table_borders(udi, "dashed", 8)
            set_fixed_width(udi, [int(total * 0.5)])
            uc = udi.cell(0, 0)
            uc.text = ""
            for k in range(3):
                pp = uc.paragraphs[0] if k == 0 else uc.add_paragraph()
                pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pp.paragraph_format.line_spacing = 1.5
                docx_util.fonted_txt(pp, "UDI 条形码" if k == 1 else "", font_size=12.0)
            cell.add_paragraph()
            fp = cell.add_paragraph()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            docx_util.fonted_txt(fp, "其他内容详见说明书", font_size=12.0, bold=True)
            document.add_paragraph()

        def render_label_udisk(node):
            # 图3：小方框 + 软件名称标题 +【英文名称/完整版本】
            kv = kv_of(node)
            total = int(usable_dxa() * 0.6)
            box = document.add_table(rows=1, cols=1)
            box.alignment = WD_TABLE_ALIGNMENT.CENTER
            box.autofit = False
            set_table_borders(box, "single", 12)
            set_fixed_width(box, [total])
            cell = box.cell(0, 0)
            cell.text = ""
            tp = cell.paragraphs[0]
            tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tp.paragraph_format.line_spacing = 1.5
            docx_util.fonted_txt(tp, kv.get("软件名称", ""), font_size=13.0, bold=True)
            add_bracket_para(cell, "产品型号", kv.get("产品型号", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
            add_bracket_para(cell, "完整版本", kv.get("完整版本", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
            document.add_paragraph()

        def render_body_node(node):
            ref = node.get("ref_type")
            title = node.get("title") or ""
            if ref == "label_product":
                fig_caption(title)
                render_label_product(node)
            elif ref == "label_seal":
                fig_caption(title)
                add_seal_image()
            elif ref == "label_udisk":
                fig_caption(title)
                render_label_udisk(node)
            else:
                render_section(node, 1)

        cover = next((s for s in sections if s.get("ref_type") == "cover"), None)
        revision = next((s for s in sections if s.get("ref_type") == "revision"), None)
        body_sections = [s for s in sections if s.get("ref_type") not in ("cover", "revision")]

        # 首页（参考产品开发计划）：留白 + 居中大标题 + 封面信息表
        add_blank_lines(6)
        cover_title = (self.__strip_num(cover.get("title")) if cover else "") or "产品标签样稿"
        write_center_title(cover_title, size=22.0, bold=True)
        add_blank_lines(4)
        if cover:
            for table in (cover.get("tables") or []):
                add_cover_grid(table)

        # 第二页（参考产品开发计划）：文件修订记录
        document.add_page_break()
        write_center_title("文件修订记录", size=14.0, bold=True)
        add_blank_lines(2)
        if revision:
            for table in (revision.get("tables") or []):
                add_grid(table, header=True)

        # 第三页起：标签图（图1/图2/图3）与技术要求
        document.add_page_break()
        for node in body_sections:
            render_body_node(node)

        document.save(output)
        output.seek(0)
