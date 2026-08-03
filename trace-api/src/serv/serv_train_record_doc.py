#!/usr/bin/env python
# encoding: utf-8

# 培训记录表服务层（测试文件）。
# 默认内容取自 src-res/train_record_default_content.json（一张大表格）。

import copy
import io
import json
import logging
import os
import re
from typing import Any, List

from sqlalchemy import delete, func, select
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.train_record_doc import TrainRecordDoc
from ..model.prod_dhf import ProdDhf
from ..model.project_member import ProjectMember
from ..model.project_timeline import ProjectTimelineRow
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_train_record_doc import TrainRecordDocForm
from ..obj.vobj_train_record_doc import TrainRecordDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from .serv_utils import new_version, sync_file_no_version
from .serv_utils import docx_util

logger = logging.getLogger(__name__)

DOC_NAME = "培训记录表"
DOC_KEY = "train_record"
BASE_NAME = "InferOperate Suite"

_DEFAULT_FILE = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "src-res", "train_record_default_content.json"
)
try:
    with open(_DEFAULT_FILE, encoding="utf-8") as _f:
        DEFAULT_CONTENT = json.load(_f)
except Exception:
    DEFAULT_CONTENT = {"sections": []}


class Server(object):

    def __normalize_content(self, content):
        if not isinstance(content, dict) or not isinstance(content.get("sections"), list):
            return copy.deepcopy(DEFAULT_CONTENT)
        result = {"sections": []}
        for s in content["sections"]:
            node = dict(s)
            node["title"] = str(node.get("title") or "")
            node["body"] = str(node.get("body") or "")
            tables = node.get("tables")
            if not isinstance(tables, list):
                tables = []
            norm = []
            for t in tables:
                if isinstance(t, list):
                    norm.append([[str(c) if c is not None else "" for c in (r or [])] for r in t if isinstance(r, list)])
            node["tables"] = norm
            node["children"] = []
            result["sections"].append(node)
        return result

    def __replace_name(self, node, base, name):
        if not name or base == name:
            return
        if node.get("body"):
            node["body"] = node["body"].replace(base, name)
        for tbl in (node.get("tables") or []):
            for row in tbl:
                for i in range(len(row)):
                    if isinstance(row[i], str) and base in row[i]:
                        row[i] = row[i].replace(base, name)

    def __dhf_file_no(self, prod_id):
        if not prod_id:
            return ""
        row = db.session.execute(
            select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name == DOC_NAME).order_by(ProdDhf.id.asc())
        ).scalars().first()
        if not row:
            row = db.session.execute(
                select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name.like(f"%{DOC_NAME}%")).order_by(ProdDhf.id.asc())
            ).scalars().first()
        return (row.code or "").strip() if row and row.code else ""

    def __autofill(self, content, prod_id, product=None, version="", force=False):
        if not isinstance(content, dict):
            return content
        if product and product.name:
            for s in (content.get("sections") or []):
                self.__replace_name(s, BASE_NAME, product.name)
        if product:
            new_fv = (product.full_version or "").strip()
            for s in (content.get("sections") or []):
                for tbl in (s.get("tables") or []):
                    for r in tbl:
                        for i in range(len(r)):
                            if isinstance(r[i], str) and re.search(r"\d+\.\d+\.\d+\.\d+", r[i]):
                                r[i] = re.sub(r"\d+\.\d+\.\d+\.\d+", new_fv, r[i])
        # 自动填充培训时间、人员名单、培训内容中的产品名称/版本
        self.__autofill_training(content, prod_id, product)
        return content

    def __autofill_training(self, content, prod_id, product=None):
        """填充培训记录表专属字段：培训内容(产品名称/版本)、培训时间、培训人员名单。"""
        if not isinstance(content, dict) or not prod_id:
            return
        sections = content.get("sections") or []
        if not sections:
            return
        tables = sections[0].get("tables") or []
        if not tables:
            return
        grid = tables[0]
        if not isinstance(grid, list) or len(grid) < 1:
            return

        # —— 培训内容（行0）：产品名称 / 版本号 ——
        pname = (product.name if product else "") or ""
        pver = (product.full_version if product else "") or ""
        if pname or pver:
            parts = []
            if pname:
                parts.append(f"产品名称：{pname}")
            if pver:
                parts.append(f"版本号：{pver}")
            parts.append("产品的使用和《用户测试计划》、《用户测试用例》")
            scope_text = "\n".join(parts)
            row0 = grid[0]
            # 内容存在 [1..5]（合并），写入 [1]，其余清空
            if len(row0) >= 6:
                row0[1] = scope_text
                for c in range(2, 6):
                    row0[c] = ""

        # 确保至少 12 行（兼容旧版 10 行：补空白填写行）
        if len(grid) == 10:
            grid.insert(5, ["", "", "", "", "", ""])
            grid.insert(7, ["", "", "", "", "", ""])
        while len(grid) < 12:
            grid.append(["", "", "", "", "", ""])

        # —— 培训时间（行1 [3]）：取项目时间线第一个日期行 ——
        try:
            trow = db.session.execute(
                select(ProjectTimelineRow)
                .where(ProjectTimelineRow.prod_id == prod_id, ProjectTimelineRow.row_type == "date")
                .order_by(ProjectTimelineRow.sort_order.asc(), ProjectTimelineRow.id.asc())
                .limit(1)
            ).scalars().first()
            if trow:
                y = (trow.year or "").strip()
                m = (trow.month or "").strip()
                d = (trow.day or "").strip()
                # 规范化：确保带"年/月/日"后缀
                if y and not y.endswith("年"):
                    y = y + "年"
                if m and not m.endswith("月"):
                    m = m + "月"
                if d and not d.endswith("日"):
                    d = d + "日"
                date_str = "".join([x for x in [y, m, d] if x])
                if len(grid[1]) >= 6:
                    grid[1][2] = date_str
        except Exception:
            logger.exception("autofill timeline failed")

        # —— 培训人员名单（行5 [0]）：取该产品「用户测试」职能人员 ——
        ut_names: list = []
        try:
            members = db.session.execute(
                select(ProjectMember.name)
                .where(ProjectMember.prod_id == prod_id, ProjectMember.role == "用户测试")
                .order_by(ProjectMember.sort_order.asc(), ProjectMember.id.asc())
            ).scalars().all()
            ut_names = [n for n in members if n and n.strip()]
            if ut_names:
                roster = "、".join(ut_names)
                if len(grid[5]) >= 1:
                    grid[5][0] = roster
        except Exception:
            logger.exception("autofill members failed")

        # —— 授课老师 / 考核人员：取该产品「产品经理」 ——
        pm_name = ""
        try:
            pm_name = (db.session.execute(
                select(ProjectMember.name)
                .where(ProjectMember.prod_id == prod_id, ProjectMember.role == "产品经理")
                .order_by(ProjectMember.sort_order.asc(), ProjectMember.id.asc())
                .limit(1)
            ).scalars().first() or "").strip()
        except Exception:
            logger.exception("autofill pm failed")
        # 行3: [2]=授课老师值
        if pm_name and len(grid[3]) >= 6 and not (grid[3][2] or "").strip():
            grid[3][2] = pm_name
        # 行9: [2]=考核人员值
        if pm_name and len(grid[9]) >= 3 and not (grid[9][2] or "").strip():
            grid[9][2] = pm_name

        # —— 默认值：培训地点 / 培训人数 / 培训方式 / 培训学时 / 考核方式 / 考核结果 / 培训评价 ——
        # 行1: [2]=培训时间(已填), [3]=培训地点标签, [4]=培训地点值
        if len(grid[1]) >= 6 and not (grid[1][4] or "").strip():
            grid[1][4] = "公司"
        # 行2: [2]=培训人数值, [3]=培训方式标签, [4]=培训方式值
        if len(grid[2]) >= 6 and not (grid[2][2] or "").strip() and ut_names:
            grid[2][2] = str(len(ut_names))
        if len(grid[2]) >= 6 and not (grid[2][4] or "").strip():
            grid[2][4] = "现场操作及线上"
        # 行3: [2]=授课老师值, [3]=培训学时标签, [4]=培训学时值
        if len(grid[3]) >= 6 and not (grid[3][4] or "").strip():
            grid[3][4] = "2"
        # 行7: 培训内容摘要 [0]（产品名称/版本 + 培训说明）
        if (pname or pver) and len(grid[7]) >= 1 and not (grid[7][0] or "").strip():
            summary_parts = []
            if pname:
                summary_parts.append(f"产品名称：{pname}")
            if pver:
                summary_parts.append(f"版本号：{pver}")
            summary_parts.append("产品的使用、《用户测试计划》、《用户测试用例》的培训")
            grid[7][0] = "\n".join(summary_parts)
        # 行8: 考核方式 [2]
        if len(grid[8]) >= 3 and not (grid[8][2] or "").strip():
            grid[8][2] = "实操，进行用户测试"
        # 行10: 考核结果 [2]
        if len(grid[10]) >= 3 and not (grid[10][2] or "").strip():
            grid[10][2] = "全部考核通过。\n\n考核人签字：            日期："
        # 行11: 培训评价 [2]
        if len(grid[11]) >= 3 and not (grid[11][2] or "").strip():
            grid[11][2] = "经过培训，测试人员已经了解产品的使用，熟悉《用户测试计划》及《用户测试用例》。\n\n评价人签字：            日期："

    def __to_obj(self, row: TrainRecordDoc, product: Product = None):
        obj = TrainRecordDocObj(**row.dict())
        obj.content = self.__autofill(self.__normalize_content(obj.content), row.product_id, product, row.version)
        if not (obj.file_no or "").strip():
            obj.file_no = self.__dhf_file_no(row.product_id)
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
        return obj

    async def add_train_record_doc(self, form: TrainRecordDocForm):
        try:
            sql = select(func.count(TrainRecordDoc.id)).where(TrainRecordDoc.product_id == form.product_id, TrainRecordDoc.version == form.version)
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = TrainRecordDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            # 自动填充产品名称/版本、培训时间、培训人员名单
            if form.product_id:
                product: Product = db.session.execute(select(Product).where(Product.id == form.product_id)).scalars().first()
                if product:
                    row.content = self.__autofill(row.content, form.product_id, product, form.version, force=True)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=TrainRecordDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_train_record_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: TrainRecordDoc = db.session.execute(select(TrainRecordDoc).where(TrainRecordDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(TrainRecordDoc.version).where(TrainRecordDoc.product_id == target_pid)).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = TrainRecordDoc(
                product_id=target_pid, version=version,
                file_no=sync_file_no_version(fromdoc.file_no, version),
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            if target_pid != fromdoc.product_id:
                await self.rebind_product(newdoc.id, target_pid)
            return Resp.resp_ok(data=TrainRecordDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_train_record_doc(self, form: TrainRecordDocForm):
        try:
            row: TrainRecordDoc = db.session.execute(select(TrainRecordDoc).where(TrainRecordDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def rebind_product(self, id: int, product_id: int):
        try:
            row: TrainRecordDoc = db.session.execute(select(TrainRecordDoc).where(TrainRecordDoc.id == id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            product: Product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
            if not product:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            db.session.execute(delete(TrainRecordDoc).where(TrainRecordDoc.product_id == product_id, TrainRecordDoc.version == row.version, TrainRecordDoc.id != id))
            content = copy.deepcopy(DEFAULT_CONTENT) if isinstance(DEFAULT_CONTENT, dict) else {"sections": []}
            content = self.__autofill(content, product_id, product, row.version, force=True)
            row.product_id = product_id
            row.content = content
            db.session.commit()
            return Resp.resp_ok(data=self.__to_obj(row, product))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_train_record_doc(self, id: int):
        db.session.execute(delete(TrainRecordDoc).where(TrainRecordDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_train_record_doc(self, id: int):
        sql = select(TrainRecordDoc, Product).join(Product, TrainRecordDoc.product_id == Product.id).where(TrainRecordDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_train_record_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(TrainRecordDoc.product_id == product_id)
        if version:
            wheres.append(TrainRecordDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(TrainRecordDoc.id)).join(Product, TrainRecordDoc.product_id == Product.id)
        if wheres:
            sql_total = sql_total.where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = select(TrainRecordDoc, Product).join(Product, TrainRecordDoc.product_id == Product.id)
        if wheres:
            sql = sql.where(*wheres)
        sql = sql.order_by(TrainRecordDoc.id.desc()).offset(page_index * page_size).limit(page_size)
        rows: List[TrainRecordDocObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    async def export_train_record_doc(self, output, id: int):
        resp = await self.get_train_record_doc(id)
        obj: TrainRecordDocObj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        content = obj.content if isinstance(obj.content, dict) else self.__normalize_content(obj.content)
        sections = content.get("sections") or []

        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")
        docx_util.add_page_number_footer(section, obj.file_no or "")

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            s = str(text or "")
            cell.text = ""
            for i, line in enumerate(s.split("\n")):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        for s in sections:
            title = str(s.get("title") or "").strip()
            if title:
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                docx_util.fonted_txt(p, title, font_size=18.0, bold=True)
                document.add_paragraph()
            for grid in (s.get("tables") or []):
                grid = [r for r in (grid or []) if isinstance(r, list)]
                cols = max((len(r) for r in grid), default=0)
                if cols <= 0:
                    continue
                # 培训记录表固定 6 列 12 行，按模板合并单元格
                is_train_table = (cols == 6 and len(grid) == 12)
                # 兼容旧版 10 行结构：在行4/5后补两个空行
                if cols == 6 and len(grid) == 10:
                    grid = grid[:4] + [grid[4], [""] * 6, grid[5], [""] * 6] + grid[6:]
                    is_train_table = True
                table = document.add_table(rows=0, cols=cols)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = True
                for r_idx, r in enumerate(grid):
                    cells = table.add_row().cells
                    for c in cells:
                        c.text = ""
                    if not is_train_table:
                        for c_idx in range(cols):
                            set_cell(cells[c_idx], r[c_idx] if c_idx < len(r) else "")
                        continue
                    # 按模板合并单元格
                    if r_idx == 0:
                        # 培训内容：标签1列 + 内容5列合并
                        cells[1].merge(cells[5])
                        set_cell(cells[0], r[0], bold=True)
                        set_cell(cells[1], r[1])
                    elif r_idx in (1, 2, 3):
                        # 培训时间/人数/老师：标签2列合并 + 值 + 标签 + 值2列合并
                        cells[0].merge(cells[1])
                        set_cell(cells[0], r[0], bold=True)
                        set_cell(cells[2], r[2])
                        set_cell(cells[3], r[3], bold=True)
                        cells[4].merge(cells[5])
                        set_cell(cells[4], r[4])
                    elif r_idx == 4:
                        # 培训人员名单标签：6列合并
                        cells[0].merge(cells[5])
                        set_cell(cells[0], r[0], bold=True)
                    elif r_idx == 5:
                        # 培训人员名单填写区：6列合并
                        cells[0].merge(cells[5])
                        set_cell(cells[0], r[0])
                    elif r_idx == 6:
                        # 培训内容摘要标签：6列合并
                        cells[0].merge(cells[5])
                        set_cell(cells[0], r[0], bold=True)
                    elif r_idx == 7:
                        # 培训内容摘要填写区：6列合并
                        cells[0].merge(cells[5])
                        set_cell(cells[0], r[0])
                    elif r_idx in (8, 9):
                        # 考核方式/人员：标签2列合并 + 内容4列合并
                        cells[0].merge(cells[1])
                        set_cell(cells[0], r[0], bold=True)
                        cells[2].merge(cells[5])
                        set_cell(cells[2], r[2])
                    elif r_idx in (10, 11):
                        # 考核结果/培训评价：标签2列合并 + 内容4列合并
                        cells[0].merge(cells[1])
                        set_cell(cells[0], r[0], bold=True)
                        cells[2].merge(cells[5])
                        set_cell(cells[2], r[2])

        document.save(output)
        output.seek(0)