#!/usr/bin/env python
# encoding: utf-8

# 测试环境维护记录服务层（测试文件）。
# 结构：说明正文(固定模板) + 资产表(固定模板，产品名称/完整版本自动获取) + 各资产周检查表。
# 周检查表：日期行按产品时间线「开始测试~测试结束」每周一条(周一~周五)自动生成；每格 是/否 复选框。

import base64
import copy
import logging
import re
from datetime import date, timedelta
from io import BytesIO
from typing import List

from sqlalchemy import delete, func, select
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.tem_doc import TemDoc
from ..model.teq_doc import TeqDoc
from ..model.prod_dhf import ProdDhf
from ..model.project_member import ProjectMember
from ..model.project_timeline import ProjectTimelineRow, ProjectTimelineCell
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_tem_doc import TemDocForm
from ..obj.vobj_tem_doc import TemDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from . import serv_review_util
from .serv_utils import new_version, sync_file_no_version
from .serv_utils import docx_util

logger = logging.getLogger(__name__)

# 说明正文（固定模板，取自《测试环境维护记录》）。
_DESC = (
    "开发/测试环境定期验证\n"
    "开发/测试活动开始前，需提前对开发/测试环境进行验证并形成记录；在开发/测试期间每周对测试环境进行验证并形成记录；\n"
    "具体开发/测试环境验证项为：\n"
    "服务器：\n"
    "验证硬件环境：CPU、GPU、内存、网卡；\n"
    "验证软件环境：操作系统、数据库、应用服务；\n"
    "开发/测试环境更新升级、病毒防护；\n"
    "网络环境验证；\n"
    "服务器备份、日志监控；\n"
    "验证测试工具运行、更新升级\n"
    "开发/测试机：\n"
    "验证硬件环境：CPU、GPU、内存、网卡；\n"
    "验证软件环境：操作系统、应用服务；\n"
    "开发/测试环境更新升级、病毒防护；\n"
    "网络环境验证；\n"
    "验证开发/测试工具运行、更新升级；\n"
    "\n"
    "一：硬件环境\n"
    "设备正常开机后：\n"
    "服务器，在“设置>系统信息”界面，查看CPU、GPU、内存信息是否准确无误，是否满足开发、测试需求；\n"
    "(1) 执行top命令，查看CPU和内存的使用情况是否正常；\n"
    "(2) 执行nvidia-smi命令，查看GPU使用情况是否正常\n"
    "开发/测试机，“设置>系统信息”或“右键计算机>属性”界面查看CPU、GPU、内存信息是否准确无误，是否满足开发、测试需求；\n"
    "(1) 开发/测试机：打开任务管理器，查看CPU、GPU、内存使用是否正常；\n"
    "服务器或测试机正常开机后能ping 192.168.111.1 网关，表示网卡正常无误；\n"
    "*注：服务器、开发/测试机需求同时满足开发及测试环境需求；\n"
    "测试环境维护记录上的硬件环境同时包括了服务器和开发/测试机。\n"
    "\n"
    "二：软件环境\n"
    "操作系统运行正常：服务器或设备能正常开机并且能流程的操作；\n"
    "数据库运行正常：数据库能正常访问、查询，无报错日志；\n"
    "应用服务器运行：服务正常运行、正常访问、无严重报错日志；\n"
    "应用服务运行正常：浏览器/客户端能正常访问产品，操作所有功能。\n"
    "*开发/测试环境维护记录上的软件环境：\n"
    "服务器包含了操作系统、数据库和应用服务器；\n"
    "开发/测试机包含了操作系统、浏览器/客户端。\n"
    "\n"
    "三：更新升级\n"
    "针对开发/测试环境，服务器、开发/测试机需要与运行环境确定的设备信息保持一致，不进行更新升级操作；\n"
    "\n"
    "四：服务器、开发/测试机是否杀毒\n"
    "软件必须安装官方正版软件，严禁安装非官方软件；\n"
    "运行杀毒软件进行杀毒，服务器、开发/测试机在本项目使用时间内一周杀毒一次。\n"
    "每周更新病毒库，并重新执行杀毒软件进行杀毒；\n"
    "\n"
    "五：网络环境是否正常\n"
    "服务器或开发/测试机能正常在百兆及以上的局域网使用表示网络正常；\n"
    "\n"
    "六：开发工具/测试工具\n"
    "是否正常运行：实际操作一遍验证是否能正常运行；\n"
    "是否更新升级：不进行更新升级。\n"
    "注：服务器和开发/测试机需要实际操作开发/测试工具。\n"
    "\n"
    "七：服务器备份\n"
    "对服务器系统及环境进行备份；\n"
    "\n"
    "八：服务器日志是否错误\n"
    "查询系统日志保证系统、服务正常运行无严重报错"
)

# 资产表（固定模板：资产编码 + 设备信息）。
_ASSET_HEADER = ["资产编码", "设备信息", "产品名称", "完整版本"]
_UBUNTU_INFO = (
    "操作系统： Ubuntu 24.04 LTS（64位）\nCPU：主频：2GHz\n核心数：10核\n"
    "指令集：x86指令集\n内存：64G\n网卡：千兆PCI-E网卡"
)
_MAC_INFO = "操作系统：macOS Monterey 12.0.1\n内存：16 GB LPDDR4\n网卡：Broadcom 57762-A0"

# 各资产的周检查表。
_ASSET_CHECKS = [
    ("SER01405", "server"),
    ("SER01358", "dev"),
    ("SER01268", "dev"),
    ("SER01100", "dev"),
    ("ISER21120027", "dev"),
    ("ISER21120029", "dev"),
    ("ISER21120023", "dev"),
]

# 检查表分组表头（组标签, [叶列标签]）；叶列为空表示该组即单列(纵向跨表头两行)。
_SERVER_GROUPS = [
    ("日期", []),
    ("硬件环境", ["CPU", "GPU", "内存", "网卡"]),
    ("软件环境", ["操作系统\n运行是否正常", "数据库\n运行是否正常", "应用服务\n运行是否正常"]),
    ("测试环境\n是否更新升级", []),
    ("服务器\n是否杀毒", []),
    ("网络环境\n是否正常", []),
    ("测试工具", ["是否正常运行", "是否更新升级"]),
    ("服务器\n是否备份", []),
    ("服务器\n日志是否错误", []),
    ("出现的问题及处理方式", []),
    ("检查人", []),
]
_DEV_GROUPS = [
    ("日期", []),
    ("硬件环境", ["CPU", "GPU", "内存", "网卡"]),
    ("软件环境", ["操作系统\n运行是否正常", "浏览器\n运行是否正常"]),
    ("测试环境\n是否更新升级", []),
    ("测试机\n是否杀毒", []),
    ("网络环境\n是否正常", []),
    ("测试工具", ["是否正常运行", "是否更新升级"]),
    ("出现的问题及处理方式", []),
    ("检查人", []),
]
_GROUPS = {"server": _SERVER_GROUPS, "dev": _DEV_GROUPS}
_SPECIAL = ("日期", "出现的问题及处理方式", "检查人")


def _leaf_columns(kind):
    """展开为叶列：[{label, type: date/check/problem/checker}]。"""
    cols = []
    for gl, leaves in _GROUPS[kind]:
        if leaves:
            for lf in leaves:
                cols.append({"label": lf, "type": "check"})
        else:
            if gl == "日期":
                t = "date"
            elif gl.startswith("出现的问题"):
                t = "problem"
            elif gl == "检查人":
                t = "checker"
            else:
                t = "check"
            cols.append({"label": gl, "type": t})
    return cols


def _check_count(kind):
    return sum(1 for c in _leaf_columns(kind) if c["type"] == "check")


def _default_mark(label):
    """默认勾选「通过状态」：更新升级类、日志是否错误的通过状态为「否」，其余为「是」。"""
    lb = str(label or "")
    if ("更新升级" in lb) or ("日志是否错误" in lb):
        return "否"
    return "是"


def _check_defaults(kind):
    return [_default_mark(c["label"]) for c in _leaf_columns(kind) if c["type"] == "check"]


def _default_content():
    checks = []
    for code, kind in _ASSET_CHECKS:
        checks.append({"asset": code, "kind": kind, "rows": []})
    return {"desc": _DESC, "assets": [], "checks": checks}


class Server(object):

    def __normalize_content(self, content):
        base = _default_content()
        if not isinstance(content, dict):
            return base
        if isinstance(content.get("desc"), str) and content.get("desc").strip():
            base["desc"] = content["desc"]
        if isinstance(content.get("assets"), list) and content["assets"]:
            base["assets"] = [list(r) if isinstance(r, list) else [r] for r in content["assets"]]
        if isinstance(content.get("checks"), list) and content["checks"]:
            base["checks"] = content["checks"]
        return base

    # ---------------- 时间线：单元测试开始 ~ 用户测试结束 ----------------
    def __dev_test_range(self, prod_id):
        if not prod_id:
            return (None, None)
        rows = db.session.execute(
            select(ProjectTimelineRow).where(ProjectTimelineRow.prod_id == prod_id)
        ).scalars().all()
        if not rows:
            return (None, None)
        cell_map = {}
        for c in db.session.execute(
            select(ProjectTimelineCell).where(ProjectTimelineCell.row_id.in_([r.id for r in rows]))
        ).scalars().all():
            cell_map.setdefault(c.row_id, []).append(c.output_result or "")

        def to_int(v):
            digits = re.sub(r"[^\d]", "", str(v or ""))
            return int(digits) if digits else None

        def as_date(r):
            y, m, d = to_int(r.year), to_int(r.month), to_int(r.day)
            if not y or not m:
                return None
            try:
                return date(y, m, d or 1)
            except ValueError:
                return None

        test_dates = []
        for r in rows:
            if (r.row_type or "date") != "date":
                continue
            dt = as_date(r)
            if not dt:
                continue
            vals = cell_map.get(r.id, [])
            if any(any(phase in str(v) for phase in ("单元测试", "集成测试", "系统测试", "用户测试"))
                   and "用例" not in str(v) for v in vals):
                test_dates.append(dt)
        start_d = min(test_dates) if test_dates else None
        end_d = max(test_dates) if test_dates else None
        return (start_d, end_d)

    def __week_ranges(self, prod_id):
        start_d, end_d = self.__dev_test_range(prod_id)
        if not start_d or not end_d or start_d > end_d:
            return []

        def fmt(d):
            return f"{d.year}.{d.month:02d}.{d.day:02d}"

        ranges = []
        cur = start_d - timedelta(days=start_d.weekday())  # 起始周的周一
        while cur <= end_d:
            monday = cur
            friday = monday + timedelta(days=4)
            ws = max(monday, start_d)
            we = min(friday, end_d)
            if ws.weekday() >= 5:
                cur = monday + timedelta(days=7)
                continue
            if we.weekday() >= 5:
                we = friday
            if ws <= we:
                ranges.append(f"{fmt(ws)}- {fmt(we)}")
            cur = monday + timedelta(days=7)
        return ranges

    def __teq_assets(self, prod_id):
        """从「测试设备清单」取设备：返回 [(资产编码, 用途)]（按清单顺序，去重）。"""
        if not prod_id:
            return []
        doc = db.session.execute(
            select(TeqDoc).where(TeqDoc.product_id == prod_id).order_by(TeqDoc.id.desc())
        ).scalars().first()
        if not doc:
            return []
        content = doc.content if isinstance(doc.content, dict) else {}
        rows = content.get("rows") or []
        out, seen = [], set()
        for r in (rows[1:] if rows else []):
            if not isinstance(r, list) or len(r) < 9:
                continue
            code = str(r[4] or "").strip()
            usage = str(r[6] or "").strip()
            if code and code not in seen:
                seen.add(code)
                out.append((code, usage))
        return out

    def __autofill(self, content, prod_id, product=None):
        if not isinstance(content, dict):
            return content
        # 资产编码/检查表：从「测试设备清单」自动获取
        teq_assets = self.__teq_assets(prod_id)
        if teq_assets:
            content["assets"] = [list(_ASSET_HEADER)] + [
                [code, (_MAC_INFO if code.upper().startswith("ISER") else _UBUNTU_INFO), "", ""]
                for code, _u in teq_assets
            ]
            old_checks = {str(ch.get("asset")): ch for ch in (content.get("checks") or []) if isinstance(ch, dict)}
            content["checks"] = [
                {
                    "asset": code,
                    "kind": ("server" if "共用" in str(usage) else "dev"),
                    "rows": (old_checks.get(code, {}).get("rows") or []),
                }
                for code, usage in teq_assets
            ]
        # 资产表：首个数据行填产品名称/完整版本
        assets = content.get("assets") or []
        if product and len(assets) >= 2 and isinstance(assets[1], list):
            row = assets[1]
            while len(row) < 4:
                row.append("")
            if not str(row[2] or "").strip():
                row[2] = product.name or ""
            if not str(row[3] or "").strip():
                row[3] = product.full_version or ""
        # 检查人：基于时间线起始日期，2025年9月前 宋月，之后 孙家旭
        start_d, _end_d = self.__dev_test_range(prod_id) if prod_id else (None, None)
        checker_name = "宋月" if serv_review_util._before_202509(start_d) else "孙家旭"
        checker_val = serv_review_util._sign_by_name(checker_name) or checker_name
        # 各检查表：按时间线周区间生成日期行（保留已填勾选/问题/检查人）
        weeks = self.__week_ranges(prod_id) if prod_id else []
        for chk in content.get("checks") or []:
            if not isinstance(chk, dict):
                continue
            kind = chk.get("kind", "dev")
            defaults = _check_defaults(kind)
            n = len(defaults)
            old = {str(r.get("date")): r for r in (chk.get("rows") or []) if isinstance(r, dict)}
            new_rows = []
            for wk in weeks:
                prev = old.get(wk)
                prev_marks = (prev.get("marks") if prev else None) or []
                marks = []
                for i in range(n):
                    m = str(prev_marks[i]) if i < len(prev_marks) else ""
                    marks.append(m if m.strip() else defaults[i])
                prob = (prev.get("problem") if prev else "") or ""
                chker = (prev.get("checker") if prev else "") or ""
                new_rows.append({
                    "date": wk,
                    "marks": marks,
                    "problem": prob if prob.strip() else "无",
                    "checker": chker if chker.strip() else checker_val,
                })
            chk["rows"] = new_rows if weeks else (chk.get("rows") or [])
        return content

    def __dhf_file_no(self, prod_id):
        """从产品 DHF 按名称「测试环境维护说明」精确匹配取文件编号。"""
        if not prod_id:
            return ""
        row = db.session.execute(
            select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name == "测试环境维护说明")
            .order_by(ProdDhf.id.asc())
        ).scalars().first()
        if not row:
            row = db.session.execute(
                select(ProdDhf).where(
                    ProdDhf.prod_id == prod_id,
                    ProdDhf.name.like("%测试环境维护说明%"),
                ).order_by(ProdDhf.id.asc())
            ).scalars().first()
        return (row.code or "").strip() if row and row.code else ""

    def __to_obj(self, row: TemDoc, product: Product = None):
        obj = TemDocObj(**row.dict())
        obj.content = self.__autofill(self.__normalize_content(obj.content), row.product_id, product)
        if not (obj.file_no or "").strip():
            obj.file_no = self.__dhf_file_no(row.product_id)
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
        return obj

    async def add_tem_doc(self, form: TemDocForm):
        try:
            sql = select(func.count(TemDoc.id)).where(
                TemDoc.product_id == form.product_id,
                TemDoc.version == form.version,
            )
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = TemDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            if not (row.file_no or "").strip():
                row.file_no = self.__dhf_file_no(form.product_id)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=TemDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_tem_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: TemDoc = db.session.execute(select(TemDoc).where(TemDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(TemDoc.version).where(TemDoc.product_id == target_pid)).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = TemDoc(
                product_id=target_pid,
                version=version,
                file_no=sync_file_no_version(fromdoc.file_no, version),
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            return Resp.resp_ok(data=TemDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_tem_doc(self, form: TemDocForm):
        try:
            row: TemDoc = db.session.execute(select(TemDoc).where(TemDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            if not (row.file_no or "").strip():
                row.file_no = self.__dhf_file_no(row.product_id)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_tem_doc(self, id: int):
        db.session.execute(delete(TemDoc).where(TemDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def refresh_content(self, product_id: int):
        """根据产品 ID 重新生成自动填充内容（用于编辑页切换产品时刷新）。"""
        product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
        content = self.__autofill(self.__normalize_content({}), product_id, product=product)
        return Resp.resp_ok(data=content)

    async def get_tem_doc(self, id: int):
        sql = select(TemDoc, Product).join(Product, TemDoc.product_id == Product.id).where(TemDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_tem_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(TemDoc.product_id == product_id)
        if version:
            wheres.append(TemDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(TemDoc.id)).join(Product, TemDoc.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (
            select(TemDoc, Product)
            .join(Product, TemDoc.product_id == Product.id)
            .where(*wheres)
            .order_by(TemDoc.id.desc())
            .offset(page_index * page_size)
            .limit(page_size)
        )
        rows: List[TemDocObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    # ---------------- 导出 Word ----------------
    async def export_tem_doc(self, output, id: int):
        resp = await self.get_tem_doc(id)
        obj: TemDocObj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        c = self.__normalize_content(obj.content)
        document = Document()
        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")
        docx_util.add_page_number_footer(section, obj.file_no or "", skip_first=False)

        def _png_size(data):
            try:
                if data[:8] == b"\x89PNG\r\n\x1a\n":
                    w = int.from_bytes(data[16:20], "big")
                    h = int.from_bytes(data[20:24], "big")
                    if w and h:
                        return w, h
            except Exception:
                pass
            return 0, 0

        def _embed_sign(para, data, max_w_pt=40.0, max_h_pt=20.0):
            """签名图按宽高上限等比缩放，避免宽签名超出检查人列被裁切。"""
            w_px, h_px = _png_size(data)
            if w_px and h_px:
                scale = min(max_w_pt / w_px, max_h_pt / h_px)
                para.add_run().add_picture(
                    BytesIO(data), width=Pt(w_px * scale), height=Pt(h_px * scale))
            else:
                para.add_run().add_picture(BytesIO(data), height=Pt(max_h_pt))

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=10.0,
                     sign_max_w=40.0, sign_max_h=20.0):
            s = str(text or "")
            if s.startswith("data:image"):
                try:
                    b64 = s.split(",", 1)[1] if "," in s else ""
                    raw = base64.b64decode(b64)
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    _embed_sign(para, raw, max_w_pt=sign_max_w, max_h_pt=sign_max_h)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    return
                except Exception:
                    pass
            cell.text = ""
            for i, line in enumerate(s.split("\n")):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.2
                docx_util.fonted_txt(para, line, font_size=size, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        def _mono(run):
            run.font.name = "宋体"
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            for _attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                rfonts.set(qn(_attr), "宋体")

        def set_yesno(cell, mark):
            cell.text = ""
            yes = (str(mark or "").strip() == "是")
            no = (str(mark or "").strip() == "否")
            p1 = cell.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p1.add_run(("\u2611\ufe0e" if yes else "\u2610") + " 是")
            r1.font.size = Pt(9)
            _mono(r1)
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(("\u2611\ufe0e" if no else "\u2610") + " 否")
            r2.font.size = Pt(9)
            _mono(r2)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # 标题
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        docx_util.fonted_txt(title, "测试环境维护记录", font_size=18.0, bold=True)

        # 资产表
        assets = [r for r in (c.get("assets") or []) if isinstance(r, list)]
        if assets:
            acols = max(len(r) for r in assets)
            at = document.add_table(rows=0, cols=acols)
            at.style = "Table Grid"
            at.alignment = WD_TABLE_ALIGNMENT.CENTER
            for r_idx, row in enumerate(assets):
                cells = at.add_row().cells
                for ci in range(acols):
                    set_cell(cells[ci], row[ci] if ci < len(row) else "", bold=(r_idx == 0),
                             align=WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT)
            document.add_paragraph()

        # 说明正文（位于资产表下方；标题加粗：开发/测试环境定期验证 + 一~八 编号小节）
        for line in str(c.get("desc") or "").split("\n"):
            lt = line.strip()
            is_head = (lt == "开发/测试环境定期验证") or bool(re.match(r"^[一二三四五六七八九十]：", lt))
            p = document.add_paragraph()
            p.paragraph_format.line_spacing = 1.3
            docx_util.fonted_txt(p, line, font_size=(12.0 if is_head else 10.5), bold=is_head)
        document.add_paragraph()

        # 各资产周检查表
        for chk in (c.get("checks") or []):
            if not isinstance(chk, dict):
                continue
            kind = chk.get("kind", "dev")
            leaves = _leaf_columns(kind)
            ncols = len(leaves)
            groups = _GROUPS[kind]
            code = chk.get("asset", "")
            for _a in assets[1:] if assets else []:
                if isinstance(_a, list) and _a and str(_a[0]).strip() == str(code).strip():
                    code = _a[0]
                    break
            title_txt = f"测试共用-{'服务器' if kind == 'server' else '测试机'}检查表（{code}）"

            tb = document.add_table(rows=0, cols=ncols)
            tb.style = "Table Grid"
            tb.alignment = WD_TABLE_ALIGNMENT.CENTER
            # 标题行（整行合并）
            trow = tb.add_row().cells
            tmerge = trow[0]
            for i in range(1, ncols):
                tmerge = tmerge.merge(trow[i])
            set_cell(tmerge, title_txt, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            # 分组行 + 叶列行
            grow = tb.add_row().cells
            lrow = tb.add_row().cells
            ci = 0
            for gl, gleaves in groups:
                if gleaves:
                    gm = grow[ci]
                    for k in range(1, len(gleaves)):
                        gm = gm.merge(grow[ci + k])
                    set_cell(gm, gl, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9.0)
                    for k, lf in enumerate(gleaves):
                        set_cell(lrow[ci + k], lf, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9.0)
                    ci += len(gleaves)
                else:
                    vm = grow[ci].merge(lrow[ci])
                    set_cell(vm, gl, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9.0)
                    ci += 1
            # 周记录行
            for row in (chk.get("rows") or []):
                if not isinstance(row, dict):
                    continue
                cells = tb.add_row().cells
                marks = row.get("marks") or []
                j = 0
                for idx, col in enumerate(leaves):
                    t = col["type"]
                    if t == "date":
                        set_cell(cells[idx], str(row.get("date", "")).replace("- ", "-\n"), align=WD_ALIGN_PARAGRAPH.CENTER, size=9.0)
                    elif t == "check":
                        set_yesno(cells[idx], marks[j] if j < len(marks) else "")
                        j += 1
                    elif t == "problem":
                        set_cell(cells[idx], row.get("problem", ""), align=WD_ALIGN_PARAGRAPH.CENTER, size=9.0)
                    elif t == "checker":
                        # 服务器表列多、总宽易超页宽被压缩，签名上限略小
                        sw, sh = (36.0, 18.0) if kind == "server" else (44.0, 22.0)
                        set_cell(cells[idx], row.get("checker", ""), align=WD_ALIGN_PARAGRAPH.CENTER,
                                 size=9.0, sign_max_w=sw, sign_max_h=sh)
            tb.autofit = False
            _tblPr = tb._tbl.tblPr
            _layout = _tblPr.find(qn("w:tblLayout"))
            if _layout is None:
                _layout = OxmlElement("w:tblLayout")
                _tblPr.append(_layout)
            _layout.set(qn("w:type"), "fixed")
            _grid = tb._tbl.find(qn("w:tblGrid"))
            if _grid is not None:
                for _gc in list(_grid):
                    _grid.remove(_gc)
                check_w = 560 if kind == "server" else 620
                for col in leaves:
                    _w = 1300 if col["type"] == "date" else 2200 if col["type"] == "problem" else 2800 if col["type"] == "checker" else check_w
                    _gce = OxmlElement("w:gridCol")
                    _gce.set(qn("w:w"), str(_w))
                    _grid.append(_gce)
            document.add_paragraph()

        document.save(output)
        output.seek(0)