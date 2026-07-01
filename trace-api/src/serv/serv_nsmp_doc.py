#!/usr/bin/env python
# encoding: utf-8

# 网络安全维护计划服务层。
# 整份文档以 content(JSON) 的「章节树」存储；产品信息（名称/版本）自动获取注入「维护范围」章节，其余模板化。
# 导出：标题「网络安全维护计划」→ 封面信息表 → 文件修订记录 → 正文章节（带章节号）。

import copy
import json
import logging
import os
import re
from typing import List

from sqlalchemy import delete, func, select
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from ..model.product import Product
from ..model.nsmp_doc import NsmpDoc
from ..model.prod_dhf import ProdDhf
from ..model.company_info import CompanyInfo
from ..model.project_member import ProjectMember
from ..model.project_timeline import ProjectTimelineRow, ProjectTimelineCell
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_nsmp_doc import NsmpDocForm
from ..obj.vobj_nsmp_doc import NsmpDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from .serv_utils import new_version, docx_util

logger = logging.getLogger(__name__)

DOC_TITLE = "网络安全维护计划"
# 修订日期从时间线「输出结果」匹配的文档名关键字（命中项取最早日期，参考产品开发计划）
DATE_KEYWORDS = ["网络安全维护计划", "维护计划"]

DEFAULT_NSMP_CONTENT = {"sections": []}
_DEFAULT_CONTENT_FILE = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "src-res", "nsmp_default_content.json"
)
try:
    with open(_DEFAULT_CONTENT_FILE, encoding="utf-8") as _f:
        _loaded = json.load(_f)
    if isinstance(_loaded, dict) and isinstance(_loaded.get("sections"), list) and _loaded["sections"]:
        DEFAULT_NSMP_CONTENT = _loaded
except Exception:
    logger.exception("加载网络安全维护计划默认内容资源失败")


class Server(object):

    # ---------------- 归一化 ----------------
    def __normalize_node(self, node):
        if not isinstance(node, dict):
            return {"title": str(node or ""), "body": "", "tables": [], "images": [], "children": []}
        result = dict(node)
        result["title"] = str(result.get("title") or "")
        result["body"] = str(result.get("body") or "")
        if "body_after" in result:
            result["body_after"] = str(result.get("body_after") or "")
        tables = result.get("tables")
        norm_tables = []
        if isinstance(tables, list):
            for table in tables:
                if isinstance(table, list):
                    norm_tables.append([[str(c) if c is not None else "" for c in (row or [])] for row in table if isinstance(row, list)])
        result["tables"] = norm_tables
        imgs = result.get("images")
        result["images"] = [str(x) for x in imgs if isinstance(x, str) and x] if isinstance(imgs, list) else []
        children = result.get("children")
        result["children"] = [self.__normalize_node(c) for c in children] if isinstance(children, list) else []
        return result

    def __normalize_content(self, content):
        if not isinstance(content, dict) or not isinstance(content.get("sections"), list):
            return copy.deepcopy(DEFAULT_NSMP_CONTENT)
        return {"sections": [self.__normalize_node(s) for s in content["sections"]]}

    @staticmethod
    def __strip_num(title):
        return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()

    # ---------------- 自动获取 ----------------
    def __company_name(self, product):
        # 公司名称：优先按产品注册人匹配「公司基本信息」，未匹配再取首条
        prod_registrant = (getattr(product, "registrant", "") or "").strip() if product else ""
        company = None
        if prod_registrant:
            company = db.session.execute(
                select(CompanyInfo).where(CompanyInfo.registrant == prod_registrant)
            ).scalars().first()
        if not company:
            company = db.session.execute(select(CompanyInfo).order_by(CompanyInfo.id.asc())).scalars().first()
        return (getattr(company, "registrant", "") or "").strip() if company else prod_registrant

    def __member_name(self, prod_id, keywords):
        # 按角色关键字（按优先级）取项目人员姓名
        for kw in keywords:
            row = db.session.execute(
                select(ProjectMember).where(ProjectMember.prod_id == prod_id, ProjectMember.role.like(f"%{kw}%"))
                .order_by(ProjectMember.sort_order.asc(), ProjectMember.id.asc())
            ).scalars().first()
            if row and (row.name or "").strip():
                return row.name.strip()
        return ""

    def __release_date(self, prod_id):
        # 修订日期：从项目时间线取「输出结果含文档名」的日期行，取最早一条（参考产品开发计划）；取不到留空
        tl_rows = db.session.execute(
            select(ProjectTimelineRow).where(ProjectTimelineRow.prod_id == prod_id)
        ).scalars().all()
        if not tl_rows:
            return ""
        cell_map = {}
        for c in db.session.execute(
            select(ProjectTimelineCell).where(ProjectTimelineCell.row_id.in_([r.id for r in tl_rows]))
        ).scalars().all():
            cell_map.setdefault(c.row_id, []).append(c.output_result or "")

        def to_int(v):
            digits = re.sub(r"[^\d]", "", str(v or ""))
            return int(digits) if digits else None

        date_rows = [r for r in tl_rows if (r.row_type or "date") == "date" and to_int(r.year) and to_int(r.month)]

        def date_key(r):
            return to_int(r.year) * 10000 + to_int(r.month) * 100 + (to_int(r.day) or 0)

        file_rows = [r for r in date_rows if any(k in str(v or "") for k in DATE_KEYWORDS for v in cell_map.get(r.id, []))]
        if not file_rows:
            return ""
        fr = min(file_rows, key=date_key)
        return f"{to_int(fr.year)}年{to_int(fr.month)}月{to_int(fr.day)}日"

    def __collect_autofill(self, product_id):
        product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
        if not product:
            return {}
        return {
            "prod_name": (product.name or "").strip(),
            "full_version": (product.full_version or "").strip(),
            "company": self.__company_name(product),
            "reviser": self.__member_name(product_id, ("产品经理", "项目经理")),
            "approver": self.__member_name(product_id, ("产品负责人", "研发负责人", "管理者代表")),
            "rev_date": self.__release_date(product_id),
        }

    def __fill_node(self, node, info, version):
        ref = node.get("ref_type")
        title = self.__strip_num(node.get("title"))
        if ref == "maint_scope" or title == "维护范围":
            company = info.get("company") or "北京推想健康医疗科技有限公司"
            node["body"] = (
                f"本文档适用于由{company}开发的下列软件产品的网络安全相关内容：\n"
                f"产品名称：{info.get('prod_name', '')}\n"
                f"产品版本：{info.get('full_version', '')}"
            )
        elif ref == "cover":
            # 封面信息表：文件版本填文档版本
            for table in (node.get("tables") or []):
                for row in table:
                    if isinstance(row, list) and len(row) >= 4 and str(row[2]).strip() == "文件版本" and version:
                        row[3] = version
        elif ref == "revision" or title == "文件修订记录":
            # 修订记录首行：修改日期/版本号/修订说明/修订人/批准人 自动获取（仅填空，不覆盖已填）
            for table in (node.get("tables") or []):
                if len(table) >= 2 and isinstance(table[1], list) and len(table[1]) >= 5:
                    row = table[1]

                    def set_if(i, val):
                        if val and not str(row[i] if i < len(row) else "").strip():
                            row[i] = val
                    set_if(0, info.get("rev_date", ""))
                    set_if(1, version)
                    if not str(row[2] or "").strip():
                        row[2] = "首次发布"
                    set_if(3, info.get("reviser", ""))
                    set_if(4, info.get("approver", ""))
        for child in (node.get("children") or []):
            self.__fill_node(child, info, version)

    def __apply_autofill(self, content, info, version):
        for node in (content.get("sections") or []):
            self.__fill_node(node, info, version)
        return content

    # ---------------- 文件编号（未手填时从产品 DHF 匹配） ----------------
    def __dhf_file_no(self, prod_id):
        row = db.session.execute(
            select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name.like("%网络安全维护计划%"))
        ).scalars().first()
        return (row.code or "").strip() if row and row.code else ""

    def __to_obj(self, row: NsmpDoc, product: Product = None, with_autofill=True):
        obj = NsmpDocObj(**row.dict())
        content = self.__normalize_content(obj.content)
        if with_autofill and row.product_id:
            info = self.__collect_autofill(row.product_id)
            content = self.__apply_autofill(content, info, obj.version)
        obj.content = content
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
            if not (obj.file_no or "").strip():
                dhf_no = self.__dhf_file_no(product.id)
                if dhf_no:
                    obj.file_no = dhf_no
        return obj

    # ---------------- CRUD ----------------
    async def add_nsmp_doc(self, form: NsmpDocForm):
        try:
            sql = select(func.count(NsmpDoc.id)).where(NsmpDoc.product_id == form.product_id, NsmpDoc.version == form.version)
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = NsmpDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=NsmpDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_nsmp_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: NsmpDoc = db.session.execute(select(NsmpDoc).where(NsmpDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(NsmpDoc.version).where(NsmpDoc.product_id == target_pid)).scalars().all()
            existing = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing:
                version = new_version(version)
            newdoc = NsmpDoc(
                product_id=target_pid, version=version, file_no=fromdoc.file_no,
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            return Resp.resp_ok(data=NsmpDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_nsmp_doc(self, form: NsmpDocForm):
        try:
            row: NsmpDoc = db.session.execute(select(NsmpDoc).where(NsmpDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            if form.product_id is not None:
                row.product_id = form.product_id
            if form.version is not None:
                row.version = form.version
            if form.file_no is not None:
                row.file_no = (form.file_no or "").strip() or None
            if form.change_log is not None:
                row.change_log = form.change_log
            if form.content is not None:
                row.content = self.__normalize_content(form.content)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_nsmp_doc(self, id: int):
        db.session.execute(delete(NsmpDoc).where(NsmpDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_nsmp_doc(self, id: int):
        sql = select(NsmpDoc, Product).join(Product, NsmpDoc.product_id == Product.id).where(NsmpDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_nsmp_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(NsmpDoc.product_id == product_id)
        if version:
            wheres.append(NsmpDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        total = db.session.execute(
            select(func.count(NsmpDoc.id)).join(Product, NsmpDoc.product_id == Product.id).where(*wheres)
        ).scalar() or 0
        sql = (
            select(NsmpDoc, Product).join(Product, NsmpDoc.product_id == Product.id)
            .where(*wheres).order_by(NsmpDoc.id.desc()).offset(page_index * page_size).limit(page_size)
        )
        rows: List[NsmpDocObj] = [self.__to_obj(doc, product, with_autofill=False) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    async def nsmp_autofill(self, product_id: int, version: str = ""):
        content = self.__normalize_content(None)
        info = self.__collect_autofill(product_id)
        content = self.__apply_autofill(content, info, version or "")
        return Resp.resp_ok(data=content)

    # ---------------- 导出 Word ----------------
    async def export_nsmp_doc(self, output, id: int):
        resp = await self.get_nsmp_doc(id)
        obj: NsmpDocObj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        content = obj.content if isinstance(obj.content, dict) else self.__normalize_content(obj.content)
        sections = content.get("sections") or []

        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")

        def add_text(text):
            if str(text or "").strip():
                docx_util.save_txt2docx(str(text or ""), document)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            cell.text = ""
            for i, line in enumerate(str(text or "").split("\n")):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        def set_grid_widths(table, grid, cols):
            def cell_units(text):
                m = 0
                for line in str(text or "").split("\n"):
                    w = sum(2 if ord(ch) > 127 else 1 for ch in line)
                    m = max(m, w)
                return max(m, 2)
            col_units = []
            for c in range(cols):
                u = 2
                for row in grid:
                    if c < len(row):
                        u = max(u, cell_units(row[c]))
                col_units.append(u)
            widths = [u * 120 + 440 for u in col_units]
            sect = document.sections[0]
            usable = int((sect.page_width - sect.left_margin - sect.right_margin) / 635)
            total = sum(widths)
            if total > usable and total > 0:
                scale = usable / total
                widths = [max(int(w * scale), 600) for w in widths]
                total = sum(widths)
            tbl_pr = table._tbl.tblPr
            layout = tbl_pr.find(qn("w:tblLayout"))
            if layout is None:
                layout = OxmlElement("w:tblLayout")
                tbl_pr.append(layout)
            layout.set(qn("w:type"), "fixed")
            tbl_w = tbl_pr.find(qn("w:tblW"))
            if tbl_w is None:
                tbl_w = OxmlElement("w:tblW")
                tbl_pr.append(tbl_w)
            tbl_w.set(qn("w:w"), str(total))
            tbl_w.set(qn("w:type"), "dxa")
            grid_el = table._tbl.find(qn("w:tblGrid"))
            if grid_el is not None:
                for i, gc in enumerate(grid_el.findall(qn("w:gridCol"))):
                    if i < len(widths):
                        gc.set(qn("w:w"), str(widths[i]))
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    if i >= len(widths):
                        continue
                    tc_pr = cell._tc.get_or_add_tcPr()
                    tc_w = tc_pr.find(qn("w:tcW"))
                    if tc_w is None:
                        tc_w = OxmlElement("w:tcW")
                        tc_pr.append(tc_w)
                    tc_w.set(qn("w:w"), str(widths[i]))
                    tc_w.set(qn("w:type"), "dxa")

        def add_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            for r_idx, row in enumerate(grid):
                cells = table.add_row().cells
                for c_idx in range(cols):
                    set_cell(cells[c_idx], row[c_idx] if c_idx < len(row) else "", bold=(r_idx == 0))
            set_grid_widths(table, grid, cols)
            document.add_paragraph()

        def write_center_title(text, size=22.0, bold=True):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, text, font_size=size, bold=bold)

        def add_blank_lines(count):
            for _ in range(max(0, int(count or 0))):
                document.add_paragraph("")

        def add_cover_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for row in grid:
                cells = table.add_row().cells
                for c_idx in range(cols):
                    text = row[c_idx] if c_idx < len(row) else ""
                    # 偶数列为标签（编写部门/文件版本/日期等）加粗，奇数列为填写值
                    set_cell(cells[c_idx], text, bold=(c_idx % 2 == 0))
                    cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 「生效日期」行：合并后面的填写格，跨满整行
                if (str(row[0]).strip() if row else "") == "生效日期" and cols > 2:
                    merged = cells[1]
                    for c_idx in range(2, cols):
                        merged = merged.merge(cells[c_idx])
                    set_cell(merged, row[1] if len(row) > 1 else "", align=WD_ALIGN_PARAGRAPH.CENTER)
            document.add_paragraph()

        def add_heading(text, level):
            size = {1: 16.0, 2: 14.0, 3: 12.0}.get(level, 11.0)
            p = document.add_heading("", level=max(1, min(level, 9)))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, text, font_size=size, bold=True)

        def render_section(node, level, number=""):
            name = self.__strip_num(node.get("title"))
            heading = f"{number} {name}".strip() if number else name
            add_heading(heading, level)
            add_text(node.get("body"))
            for table in (node.get("tables") or []):
                add_grid(table)
            add_text(node.get("body_after"))
            idx = 0
            for child in (node.get("children") or []):
                idx += 1
                child_num = f"{number}.{idx}" if number else f"{idx}"
                render_section(child, level + 1, child_num)

        cover = next((s for s in sections if s.get("ref_type") == "cover"), None)
        revision = next((s for s in sections if s.get("ref_type") == "revision"), None)
        body = [s for s in sections if s.get("ref_type") not in ("cover", "revision")]

        # 封面：上方留白 + 居中大标题 + 留白 + 封面表（参考产品开发计划）
        add_blank_lines(6)
        write_center_title((self.__strip_num(cover.get("title")) if cover else "") or DOC_TITLE, size=22.0, bold=True)
        add_blank_lines(4)
        if cover:
            for table in (cover.get("tables") or []):
                add_cover_grid(table)

        # 文件修订记录：独立页、居中标题
        document.add_page_break()
        write_center_title("文件修订记录", size=14.0, bold=True)
        add_blank_lines(2)
        if revision:
            for table in (revision.get("tables") or []):
                add_grid(table)

        # 正文：另起一页，按 1/2/3 编号
        document.add_page_break()
        for i, node in enumerate(body):
            render_section(node, 1, str(i + 1))

        document.save(output)
        output.seek(0)
