#!/usr/bin/env python
# encoding: utf-8

# 自研软件研究报告服务层。
# 默认内容来自 src-res/research_default_content.json（章节树+表格+内置流程图）。
# 自动获取章节通过节点 ref_type / img_category 标记，在 get/export/预览时注入产品数据；
# 数据源：Product/CompanyInfo（软件标识）、Product.overall_desc（总体描述）、PtrDoc 2.1功能、
#         SrsDoc 2.3章节、ProdRuntimeEnv（运行环境）、ProjectTimeline（发布日期）、DocFile（产品图）。

import base64
import copy
import datetime
import io
import json
import logging
import os
import re
from typing import List

from sqlalchemy import func, select, delete
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.company_info import CompanyInfo
from ..model.research_doc import ResearchDoc
from ..model.ptr_doc import PtrDoc
from ..model.srs_doc import SrsDoc, SrsNode
from ..model.prod_runtime_env import ProdRuntimeEnv
from ..model.doc_file import DocFile
from ..model.prod_dhf import ProdDhf
from ..model.project_timeline import ProjectTimelineRow, ProjectTimelineCell
from ..model.project_member import ProjectMember
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_research_doc import ResearchDocForm
from ..obj.vobj_research_doc import ResearchDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from .serv_utils import new_version, sync_file_no_version, docx_util
from .serv_prod_runtime_env import DEFAULT_RUNTIME_ENV

logger = logging.getLogger(__name__)

# 取发布日期时的时间线关键字（命中项取最新日期）
DATE_KEYWORDS = ["自研软件研究报告", "软件研究报告", "软件发布", "发布"]

# 模板占位产品名（含可能的空格变体），自动获取时全文统一替换为实际产品名
_TEMPLATE_PRODUCT_NAME_RE = re.compile(r"肿瘤\s*CT\s*图像随访与评估软件")

DEFAULT_RESEARCH_CONTENT = {"productName": "", "sections": []}

_DEFAULT_CONTENT_FILE = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "src-res", "research_default_content.json"
)
try:
    with open(_DEFAULT_CONTENT_FILE, encoding="utf-8") as _f:
        _loaded = json.load(_f)
    if isinstance(_loaded, dict) and isinstance(_loaded.get("sections"), list) and _loaded["sections"]:
        DEFAULT_RESEARCH_CONTENT = _loaded
except Exception:
    logger.exception("加载自研软件研究报告默认内容资源失败")

# 模板表名（标题 -> 各表表名），用于给旧文档补全表上方显示的表名（未人工编辑时）
_DEFAULT_TABLE_TITLES = {}


def _build_title_map(nodes):
    for n in nodes or []:
        if isinstance(n, dict):
            if n.get("table_titles") and n.get("title"):
                _DEFAULT_TABLE_TITLES[n["title"]] = n["table_titles"]
            _build_title_map(n.get("children"))


_build_title_map(DEFAULT_RESEARCH_CONTENT.get("sections"))


class Server(object):
    # ---------------- 内容归一 ----------------
    def __fill_template_titles(self, nodes):
        # 旧文档若节点有表格但缺表名，按标题从模板补全（不覆盖人工已填的表名）
        for n in nodes or []:
            if isinstance(n, dict):
                if n.get("tables") and not n.get("table_titles"):
                    titles = _DEFAULT_TABLE_TITLES.get(n.get("title"))
                    if titles:
                        n["table_titles"] = list(titles)
                        # 旧文档：正文里若仍残留这些表名行（已上移为表名），移除避免重复
                        text = n.get("text")
                        if isinstance(text, str) and text:
                            title_set = {re.sub(r"\s+", "", t) for t in titles if t}
                            n["text"] = "\n".join(
                                ln for ln in text.split("\n")
                                if re.sub(r"\s+", "", ln.strip()) not in title_set
                            )
                self.__fill_template_titles(n.get("children"))

    def __normalize_content(self, content):
        result = copy.deepcopy(DEFAULT_RESEARCH_CONTENT)
        if isinstance(content, dict):
            result.update(content)
        result.setdefault("sections", copy.deepcopy(DEFAULT_RESEARCH_CONTENT.get("sections", [])))
        result.setdefault("productName", "")
        self.__strip_blocks(result.get("sections"))
        self.__fill_template_titles(result.get("sections"))
        return result

    def __strip_blocks(self, nodes):
        # blocks 为读取时由 autofill 派生的展示结构，不持久化；保留原始 text/images/tables 以便重算
        for n in nodes or []:
            if isinstance(n, dict):
                n.pop("blocks", None)
                self.__strip_blocks(n.get("children"))

    # ---------------- 自动获取数据源 ----------------
    def __latest_doc(self, model, product_id):
        # 排除软删除（版本号以 __deleted 前缀标记）的文档，取最新一版
        return db.session.execute(
            select(model)
            .where(model.product_id == product_id, ~model.version.like("__deleted%"))
            .order_by(model.id.desc())
        ).scalars().first()

    def __strip_name(self, title):
        return re.sub(r"^[0-9．.、\s]+", "", str(title or "")).strip()

    @staticmethod
    def __level_number(depth, idx):
        # 不同层级使用不同序号样式：第一级 (1)(2)(3)，第二级 1)2)3)，第三级及以下 ①②③ / a) 兜底
        if depth <= 1:
            return f"({idx})"
        if depth == 2:
            return f"{idx})"
        circled = "①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳"
        return circled[idx - 1] if 1 <= idx <= len(circled) else f"{idx}."

    def __node_outline(self, node, depth, lines, number=""):
        # 将一个节点（含其全部子孙）按「序号 + 标题 + 正文」逐行展开；
        # depth==0 的根节点只取正文不重复标题，每一层在其父级下独立编号，且按层级使用不同序号样式。
        title = self.__strip_name(node.get("title"))
        body = str(node.get("body") or "").strip()
        if depth >= 1 and title:
            lines.append(f"{number} {title}" if number else title)
        if body:
            lines.append(body)
        idx = 0
        for child in node.get("children") or []:
            idx += 1
            self.__node_outline(child, depth + 1, lines, self.__level_number(depth + 1, idx))

    def __ptr_func_2_1(self, product_id):
        # 技术要求 2.1 功能：正文第 2 章（性能指标）下「功能」子节的完整内容（含各模块及功能点）
        doc = self.__latest_doc(PtrDoc, product_id)
        if not doc or not isinstance(doc.content, dict):
            return ""
        sections = doc.content.get("sections") or []
        body = [s for s in sections if (s.get("ref_type") not in ("cover", "appendix"))]
        target_sec = next((s for s in body if self.__strip_name(s.get("title")) == "性能指标"), None)
        if target_sec is None and len(body) >= 2:
            target_sec = body[1]
        if not target_sec:
            return ""
        children = target_sec.get("children") or []
        child = next((c for c in children if self.__strip_name(c.get("title")) == "功能"), None) or (children[0] if children else None)
        if not child:
            return ""
        lines = []
        self.__node_outline(child, 0, lines)
        return "\n".join(lines)

    def __srs_2_3(self, product_id):
        # 需求规格说明 2.3 章节完整内容：2.3 本体 + 其全部子节（2.3.1/2.3.2…体系结构各模块功能与用途）
        doc = self.__latest_doc(SrsDoc, product_id)
        if not doc:
            return ""
        nodes = db.session.execute(
            select(SrsNode).where(SrsNode.doc_id == doc.id).order_by(SrsNode.priority, SrsNode.n_id)
        ).scalars().all()

        def num_of(n):
            matched = re.match(r"^(\d+(?:\.\d+)*)", str(n.title or "").strip())
            return matched.group(1) if matched else ""

        picked = []
        for n in nodes:
            num = num_of(n)
            if num == "2.3" or num.startswith("2.3."):
                picked.append((num, n))
        picked.sort(key=lambda it: [int(p) for p in it[0].split(".")])

        lines = []
        sub_idx = 0
        for num, n in picked:
            title = self.__strip_name(n.title)
            text = str(n.text or "").strip()
            if num == "2.3":
                if text:
                    lines.append(text)
            else:
                sub_idx += 1
                if title:
                    lines.append(f"({sub_idx}) {title}")
                if text:
                    lines.append(text)
        return "\n".join(lines)

    def __runtime_env(self, product_id):
        row = db.session.execute(select(ProdRuntimeEnv).where(ProdRuntimeEnv.prod_id == product_id)).scalars().first()
        env = dict(DEFAULT_RUNTIME_ENV)
        if row:
            for key in DEFAULT_RUNTIME_ENV.keys():
                val = getattr(row, key, None)
                if val is not None and str(val).strip():
                    env[key] = val
        return env

    def __release_date(self, product_id):
        rows = db.session.execute(
            select(ProjectTimelineRow).where(ProjectTimelineRow.prod_id == product_id)
        ).scalars().all()
        if not rows:
            return ""
        cells = db.session.execute(
            select(ProjectTimelineCell).where(ProjectTimelineCell.row_id.in_([r.id for r in rows]))
        ).scalars().all()
        cell_map = {}
        for c in cells:
            cell_map.setdefault(c.row_id, []).append(str(c.output_result or ""))

        def to_int(v):
            try:
                return int(str(v).strip())
            except Exception:
                return None

        hits = []
        for r in rows:
            if str(getattr(r, "row_type", "")) != "date":
                continue
            y, m, d = to_int(r.year), to_int(r.month), to_int(r.day)
            if y is None or m is None:
                continue
            texts = cell_map.get(r.id, [])
            if any(k in t for k in DATE_KEYWORDS for t in texts):
                hits.append((y, m, d or 1, r))
        if not hits:
            return ""
        y, m, d, _ = max(hits, key=lambda x: x[0] * 10000 + x[1] * 100 + x[2])
        return f"{y}年{m}月{d}日"

    def __dev_amount(self, product_id):
        # 表5 开发量：
        #   开发人员数量 = 参与人员中角色含「开发」的人数；
        #   开发时间 = 时间线中标注「产品开发」阶段的日期行首末跨度（含首尾天数），不是整张时间线首末；
        #   工作量 = 人数 × 天数。
        members = db.session.execute(
            select(ProjectMember).where(ProjectMember.prod_id == product_id)
        ).scalars().all()
        headcount = sum(1 for m in members if "开发" in str(m.role or ""))

        rows = db.session.execute(
            select(ProjectTimelineRow).where(
                ProjectTimelineRow.prod_id == product_id,
                ProjectTimelineRow.row_type == "date",
            )
        ).scalars().all()
        row_map = {r.id: r for r in rows}
        cells = db.session.execute(
            select(ProjectTimelineCell).where(ProjectTimelineCell.row_id.in_(list(row_map.keys())))
        ).scalars().all() if row_map else []
        # 标注「产品开发」阶段的日期行（单元格中存在独立一行恰为「产品开发」，排除「产品开发计划」等）
        dev_row_ids = set()
        for c in cells:
            for ln in str(c.output_result or "").split("\n"):
                if ln.strip() == "产品开发":
                    dev_row_ids.add(c.row_id)
                    break

        def to_int(v):
            # 时间线字段可能带单位（如「5月」「25日」「2025年」），只取数字
            digits = re.sub(r"[^\d]", "", str(v or ""))
            return int(digits) if digits else None

        def parse_date(r):
            y, m = to_int(r.year), to_int(r.month)
            if y is None or m is None:
                return None
            try:
                return datetime.date(y, m, to_int(r.day) or 1)
            except Exception:
                return None

        # 优先取「产品开发」阶段；若未标注则回退到全部日期行，保证有值
        target_rows = [row_map[i] for i in dev_row_ids] if dev_row_ids else rows
        dates = [d for d in (parse_date(r) for r in target_rows) if d]
        if len(dates) >= 2:
            days = (max(dates) - min(dates)).days + 1
        elif len(dates) == 1:
            days = 1
        else:
            days = 0
        workload = headcount * days
        return {"headcount": headcount, "days": days, "workload": workload}

    def __doc_file_url(self, product_id, category):
        row = db.session.execute(
            select(DocFile).where(DocFile.product_id == product_id, DocFile.category == category).order_by(DocFile.id)
        ).scalars().first()
        return (row.file_url or "") if row else ""

    def __collect_autofill(self, product_id):
        product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
        if not product:
            return None
        company = None
        if product.registrant:
            company = db.session.execute(
                select(CompanyInfo).where(CompanyInfo.registrant == product.registrant)
            ).scalars().first()
        company_name = (product.registrant or (company.registrant if company else "") or "").strip()
        address = (company.address if company else "") or ""
        env = self.__runtime_env(product_id)
        release_date = self.__release_date(product_id)

        sw_ident_text = (
            f"产品名称：{product.name or ''}\n"
            f"软件发布版本: {product.release_version or ''}\n"
            f"注册申请人：{company_name}\n"
            f"设计开发地址：{address}\n"
            f"现成软件的标识信息详见1.3.2.1章节组件信息。"
        )
        update_text = (
            f"本软件当前发布版本{product.release_version or ''}，完整版本为{product.full_version or ''}，"
            f"发布日期是{release_date or ''}，版本命名规则和更新历史详见提交文件《版本更新历史》。"
        )
        return {
            "product_name": (product.name or "").strip(),
            "release_version": (product.release_version or "").strip(),
            "full_version": (product.full_version or "").strip(),
            "sw_ident_text": sw_ident_text,
            "overall_desc": (product.overall_desc or "").strip(),
            "ptr_2_1": self.__ptr_func_2_1(product_id),
            "srs_2_3": self.__srs_2_3(product_id),
            "update_text": update_text,
            "runtime": env,
            "dev_amount": self.__dev_amount(product_id),
            "images": {
                "img_ui": self.__doc_file_url(product_id, "img_ui"),
                "img_struct": self.__doc_file_url(product_id, "img_struct"),
                "img_home": self.__doc_file_url(product_id, "img_home"),
                "img_topo": self.__doc_file_url(product_id, "img_topo"),
            },
        }

    def __runtime_tables(self, kind, env):
        if kind == "rt_hw":
            return [
                [["配置", "要求"], ["CPU", env.get("srv_cpu", "")], ["内存", env.get("srv_memory", "")],
                 ["GPU", env.get("srv_gpu", "")], ["硬盘", env.get("srv_disk", "")], ["网卡", env.get("srv_nic", "")]],
                [["配置", "要求"], ["CPU", env.get("cli_cpu", "")], ["内存", env.get("cli_memory", "")],
                 ["显示器分辨率", env.get("cli_resolution", "")]],
            ]
        if kind == "rt_sw":
            return [
                [["类别", "操作系统", "其他"],
                 ["服务器", env.get("srv_os", ""), f"CUDA {env.get('srv_cuda', '')}"],
                 ["用户端", env.get("cli_os", ""), f"浏览器 {env.get('cli_browser', '')}"]],
            ]
        if kind == "rt_net":
            return [
                [["网络", "要求"], ["架构", env.get("arch", "")],
                 ["局域网", env.get("net_lan", "")], ["广域网", env.get("net_wan", "")]],
            ]
        return []

    @staticmethod
    def __split_image_blocks(text, image_urls, img_category=None):
        # 统一图文版式，使图与正文位置和原 Word 一致：
        #   1) 正文中含占位符行「{{IMG}}」时，图片按占位符出现的位置插入（精确锚定）；
        #   2) 否则按「图N …」图题行定位：文字 → 图片 → 图题(图下居中) → 文字；
        #   3) 都没有时：正文在前、图片在后。
        urls = [u for u in (image_urls or []) if u]
        lines = str(text or "").split("\n")
        blocks = []
        buf = []

        def flush():
            t = "\n".join(buf).strip()
            if t:
                blocks.append({"text": t})
            buf.clear()

        if any(ln.strip() == "{{IMG}}" for ln in lines):
            ui = 0
            for ln in lines:
                s = ln.strip()
                if s == "{{IMG}}":
                    flush()
                    if ui < len(urls):
                        blocks.append({"type": "image", "url": urls[ui], "img_category": img_category})
                        ui += 1
                elif re.match(r"^图\s*\d", s):
                    flush()
                    blocks.append({"type": "caption", "text": s})
                else:
                    buf.append(ln)
            flush()
            for u in urls[ui:]:
                blocks.append({"type": "image", "url": u, "img_category": img_category})
            return blocks

        cap_i = next((i for i, ln in enumerate(lines) if re.match(r"^图\s*\d", ln.strip())), None)
        imgs = [{"type": "image", "url": u, "img_category": img_category} for u in urls]
        if cap_i is not None:
            before = "\n".join(lines[:cap_i]).strip()
            caption = lines[cap_i].strip()
            after = "\n".join(lines[cap_i + 1:]).strip()
            if before:
                blocks.append({"text": before})
            blocks.extend(imgs)
            if caption:
                blocks.append({"type": "caption", "text": caption})
            if after:
                blocks.append({"text": after})
        else:
            if str(text or "").strip():
                blocks.append({"text": str(text)})
            blocks.extend(imgs)
        return blocks

    def __apply_autofill(self, content, auto):
        if not auto:
            return content
        images = auto.get("images", {})

        def img_url(cat):
            url = images.get(cat, "")
            return f"/{url}" if url else ""

        def walk(nodes):
            for node in nodes or []:
                rt = node.get("ref_type")
                cat = node.get("img_category")
                # 兼容旧文档：缺 ref_type 时按标题兜底识别版本命名规则章节
                if not rt and self.__strip_name(node.get("title")) == "软件版本命名规则":
                    rt = "version_rule"
                if rt == "sw_ident":
                    node["text"] = auto["sw_ident_text"]
                elif rt == "func_module":
                    node["text"] = ""
                    node["images"] = []
                    overall = auto.get("overall_desc", "") or "（产品总体描述未填写，请在产品信息中维护）"
                    blocks = [{"text": overall + "\n详见图1 用户关系图。"}]
                    blocks.append({"type": "image", "url": img_url("img_ui"), "img_category": "img_ui"})
                    ptr = auto.get("ptr_2_1", "")
                    blocks.append({"text": "各功能模块的主要功能见下：\n" + ptr if ptr else "各功能模块的主要功能见下：\n（未获取到技术要求2.1功能，请先维护产品技术要求）"})
                    node["blocks"] = blocks
                elif rt == "arch_func":
                    srs = auto.get("srs_2_3", "")
                    node["text"] = ""
                    node["images"] = []
                    blocks = [{"type": "image", "url": img_url("img_struct"), "img_category": "img_struct"}]
                    blocks.append({"text": srs or "（未获取到需求规格说明2.3章节，请先维护需求规格说明）"})
                    for tbl in node.get("tables") or []:
                        blocks.append({"type": "table", "table": tbl})
                    node["blocks"] = blocks
                elif rt in ("rt_hw", "rt_sw", "rt_net"):
                    # 运行环境：按原文「正文 → 表格」交替排版（每段正文紧跟其对应表格）
                    tables = self.__runtime_tables(rt, auto.get("runtime", {}))
                    lines = str(node.get("text") or "").split("\n")
                    titles = node.get("table_titles") or []
                    blocks = []
                    for i, tbl in enumerate(tables):
                        if i < len(lines) and lines[i].strip():
                            blocks.append({"text": lines[i].strip()})
                        blocks.append({"type": "table", "table": tbl, "title": titles[i] if i < len(titles) else ""})
                    for j in range(len(tables), len(lines)):
                        if lines[j].strip():
                            blocks.append({"text": lines[j].strip()})
                    node["text"] = ""
                    node["tables"] = []
                    node["table_titles"] = []
                    node["blocks"] = blocks
                elif rt == "update_history":
                    node["text"] = auto["update_text"]
                elif rt == "version_rule":
                    # 版本命名规则：发布版本、完整版本按产品实际版本自动获取，其余规则说明保留；
                    # 版本结构图锚定在「软件完整版本及说明:」之后，与原 Word 位置一致。
                    rv = auto.get("release_version", "")
                    fv = auto.get("full_version", "")
                    out = []
                    has_img_holder = any(ln.strip() == "{{IMG}}" for ln in str(node.get("text") or "").split("\n"))
                    for ln in str(node.get("text") or "").split("\n"):
                        s = ln.strip()
                        if s.startswith("发布版本") and rv:
                            out.append(f"发布版本：V{rv}")
                        elif s.startswith("完整版本") and fv:
                            out.append(f"完整版本：V{fv}")
                        else:
                            out.append(ln)
                        if not has_img_holder and s.startswith("软件完整版本及说明"):
                            out.append("{{IMG}}")
                    node["text"] = "\n".join(out)
                elif cat:
                    # 自动获取图类章节：图片来自图表文件管理，按图题切分排版（保留原 text 以幂等重算）
                    node["blocks"] = self.__split_image_blocks(node.get("text"), [img_url(cat)], cat)
                # 通用：模板内嵌固定图（base64 存于 images）的章节也统一为「正文 → 图 → 图题 → 正文」
                if not node.get("blocks") and (node.get("images") or []):
                    node["blocks"] = self.__split_image_blocks(node.get("text"), node.get("images") or [], None)
                # 表5 开发量：按表名定位（含「开发量」），仅重算「软件名称/开发人员数量/开发时间/工作量」列，保留代码总行数
                titles = node.get("table_titles") or []
                da_idx = next((i for i, t in enumerate(titles) if "开发量" in str(t)), None)
                da = auto.get("dev_amount")
                if da_idx is not None and da is not None:
                    tbls = node.get("tables") or []
                    if 0 <= da_idx < len(tbls) and len(tbls[da_idx]) >= 2:
                        data_row = tbls[da_idx][1]
                        if len(data_row) >= 4:
                            data_row[0] = auto.get("product_name", "") or data_row[0]
                            data_row[1] = str(da["headcount"])
                            data_row[2] = f'{da["days"]}天'
                            data_row[3] = str(da["workload"])
                walk(node.get("children"))

        walk(content.get("sections"))
        product_name = (auto.get("product_name") or "").strip()
        if product_name and re.sub(r"\s", "", product_name) != "肿瘤CT图像随访与评估软件":
            self.__sync_product_name(content, product_name)
        return content

    def __sync_product_name(self, content, product_name):
        # 全文统一替换：把模板占位产品名（含空格变体）替换为实际产品名
        def rep(s):
            return _TEMPLATE_PRODUCT_NAME_RE.sub(product_name, s) if isinstance(s, str) else s

        def walk(nodes):
            for n in nodes or []:
                for k in ("text", "body"):
                    if isinstance(n.get(k), str):
                        n[k] = rep(n[k])
                for tbl in n.get("tables") or []:
                    for row in tbl:
                        for i, c in enumerate(row):
                            if isinstance(c, str):
                                row[i] = rep(c)
                for b in n.get("blocks") or []:
                    for k in ("text", "title"):
                        if isinstance(b.get(k), str):
                            b[k] = rep(b[k])
                    for row in b.get("table") or []:
                        for i, c in enumerate(row):
                            if isinstance(c, str):
                                row[i] = rep(c)
                walk(n.get("children"))

        walk(content.get("sections"))

    # ---------------- 转换 ----------------
    # 从产品 DHF 按文档名匹配文件编号（文件编号未手填时自动获取）
    def __dhf_file_no(self, prod_id):
        row = db.session.execute(
            select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name.like("%自研软件研究报告%"))
        ).scalars().first()
        return (row.code or "").strip() if row and row.code else ""

    def __to_obj(self, row: ResearchDoc, product: Product = None, with_autofill=True):
        obj = ResearchDocObj(**row.dict())
        content = self.__normalize_content(obj.content)
        if with_autofill:
            auto = self.__collect_autofill(row.product_id)
            content = self.__apply_autofill(content, auto)
        if product:
            content["productName"] = product.name or ""
        obj.content = content
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
            if not (obj.file_no or "").strip():
                dhf_no = self.__dhf_file_no(product.id)
                if dhf_no:
                    obj.file_no = dhf_no
        return obj

    # ---------------- CRUD ----------------
    async def add_research_doc(self, form: ResearchDocForm):
        try:
            sql = select(func.count(ResearchDoc.id)).where(
                ResearchDoc.product_id == form.product_id, ResearchDoc.version == form.version
            )
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = ResearchDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=ResearchDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_research_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: ResearchDoc = db.session.execute(select(ResearchDoc).where(ResearchDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(ResearchDoc.version).where(ResearchDoc.product_id == target_pid)).scalars().all()
            existing = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing:
                version = new_version(version)
            newdoc = ResearchDoc(
                product_id=target_pid, version=version, file_no=sync_file_no_version((fromdoc.file_no or "").strip() or self.__dhf_file_no(target_pid), version) or None,
                change_log=fromdoc.change_log, content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            return Resp.resp_ok(data=ResearchDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_research_doc(self, form: ResearchDocForm):
        try:
            row: ResearchDoc = db.session.execute(select(ResearchDoc).where(ResearchDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_research_doc(self, id: int):
        db.session.execute(delete(ResearchDoc).where(ResearchDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_research_doc(self, id: int):
        sql = select(ResearchDoc, Product).join(Product, ResearchDoc.product_id == Product.id).where(ResearchDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_research_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(ResearchDoc.product_id == product_id)
        if version:
            wheres.append(ResearchDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        total = db.session.execute(
            select(func.count(ResearchDoc.id)).join(Product, ResearchDoc.product_id == Product.id).where(*wheres)
        ).scalar() or 0
        sql = (
            select(ResearchDoc, Product).join(Product, ResearchDoc.product_id == Product.id)
            .where(*wheres).order_by(ResearchDoc.id.desc()).offset(page_index * page_size).limit(page_size)
        )
        rows: List[ResearchDocObj] = [self.__to_obj(doc, product, with_autofill=False) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    async def research_autofill(self, product_id: int):
        # 新增页预览：返回应用了自动获取的默认内容
        content = self.__normalize_content(None)
        auto = self.__collect_autofill(product_id)
        content = self.__apply_autofill(content, auto)
        product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
        if product:
            content["productName"] = product.name or ""
        return Resp.resp_ok(data=content)

    # ---------------- 导出 ----------------
    async def export_research_doc(self, output, id: int):
        resp = await self.get_research_doc(id)
        obj: ResearchDocObj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        document.settings.element.append(update_fields)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")

        def normalized(value):
            return re.sub(r"\s+", "", str(value or ""))

        def is_cover(sec):
            return sec.get("ref_type") == "cover" or normalized(sec.get("title")) == "自研软件研究报告"

        def is_revision(sec):
            return sec.get("ref_type") == "revision" or normalized(sec.get("title")) == "文件修订记录"

        def set_cell_text(cell, text, bold=False):
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.line_spacing = 1.5
            docx_util.fonted_txt(paragraph, str(text or ""), font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        def add_table_title(title):
            title = str(title or "").strip()
            if not title:
                return
            para = document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            docx_util.fonted_txt(para, title, font_size=10.5, bold=True)

        def add_plain_table(rows, title=""):
            rows = rows or []
            col_count = max([len(row or []) for row in rows] or [0])
            if col_count <= 0:
                return
            add_table_title(title)
            table = document.add_table(rows=0, cols=col_count)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for row in rows:
                cells = table.add_row().cells
                for idx in range(col_count):
                    set_cell_text(cells[idx], row[idx] if idx < len(row or []) else "", bold=(len(table.rows) == 1))
            document.add_paragraph()

        def add_image(url):
            raw = str(url or "").strip()
            if not raw:
                return
            try:
                docx_util.save_img2docx(raw, document)
            except Exception:
                logger.exception("导出研究报告图片失败")

        def add_section(sec, level=1):
            title = sec.get("title", "")
            if title and not is_cover(sec):
                docx_util.save_title2docx(title, document, level=max(1, min(level, 9)))
            blocks = sec.get("blocks")
            if isinstance(blocks, list) and blocks:
                for block in blocks:
                    if not isinstance(block, dict):
                        continue
                    if block.get("type") == "image":
                        add_image(block.get("url"))
                    elif block.get("type") == "caption":
                        cap = document.add_paragraph()
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        docx_util.fonted_txt(cap, str(block.get("text") or ""), font_size=10.5)
                    elif block.get("type") == "table":
                        add_plain_table(block.get("table") or [], block.get("title") or "")
                    elif block.get("text"):
                        docx_util.save_txt2docx(str(block.get("text") or ""), document)
            else:
                if sec.get("text"):
                    docx_util.save_txt2docx(str(sec.get("text") or ""), document)
                for url in sec.get("images", []) or []:
                    add_image(url)
                titles = sec.get("table_titles") or []
                for idx, rows in enumerate(sec.get("tables", []) or []):
                    add_plain_table(rows, titles[idx] if idx < len(titles) else "")
            for child in sec.get("children", []) or []:
                add_section(child, level + 1)

        content = obj.content if isinstance(obj.content, dict) else self.__normalize_content(obj.content)
        sections = content.get("sections", [])
        cover = next((s for s in sections if is_cover(s)), None)
        revision = next((s for s in sections if is_revision(s)), None)
        body = [s for s in sections if not is_cover(s) and not is_revision(s)]

        # 封面：上方留白，使大标题落在封面页垂直中部
        for _ in range(10):
            document.add_paragraph("")
        title_para = document.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.line_spacing = 1.5
        docx_util.fonted_txt(title_para, "自研软件研究报告", font_size=22.0, bold=False)
        for _ in range(6):
            document.add_paragraph("")
        for rows in (cover or {}).get("tables", []) or []:
            add_plain_table(rows)
        document.add_page_break()
        # 文件修订记录
        if revision:
            docx_util.save_title2docx("文件修订记录", document, level=1)
            for rows in revision.get("tables", []) or []:
                add_plain_table(rows)
            document.add_page_break()
        # 目录
        docx_util.insert_toc_field(document, "1-4")
        document.add_page_break()
        # 正文
        for sec in body:
            add_section(sec, 1)

        document.save(output)
        output.seek(0)
