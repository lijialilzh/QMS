#!/usr/bin/env python
# encoding: utf-8

# 现场测试规程服务层（测试文件 VV-006），结构与软件测试计划(stp)一致。
# 默认内容取自 src-res/ftr_default_content.json。

import copy
import json
import logging
import os
import re
from typing import List

from sqlalchemy import delete, func, select
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.ftr_doc import FtrDoc
from ..model.prod_dhf import ProdDhf
from ..model.project_member import ProjectMember
from ..model.prod_runtime_env import ProdRuntimeEnv
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_ftr_doc import FtrDocForm
from ..obj.vobj_ftr_doc import FtrDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from .serv_utils import new_version, sync_file_no_version
from .serv_utils import docx_util
from . import serv_review_util
from .serv_prod_runtime_env import DEFAULT_RUNTIME_ENV

logger = logging.getLogger(__name__)

DOC_NAME = "现场测试规程"
DOC_KEY = "ftr"
BASE_NAME = "InferOperate Suite"

_DEFAULT_FILE = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "src-res", "ftr_default_content.json"
)
try:
    with open(_DEFAULT_FILE, encoding="utf-8") as _f:
        DEFAULT_FTR_CONTENT = json.load(_f)
except Exception:
    DEFAULT_FTR_CONTENT = {"sections": []}


class Server(object):

    def __normalize_node(self, node):
        if not isinstance(node, dict):
            return {"title": str(node or ""), "body": "", "tables": [], "children": []}
        result = dict(node)
        result["title"] = str(result.get("title") or "")
        result["body"] = str(result.get("body") or "")
        tables = result.get("tables")
        if not isinstance(tables, list):
            tables = []
        norm_tables = []
        for table in tables:
            if isinstance(table, list):
                norm_tables.append([[str(c) if c is not None else "" for c in (row or [])] for row in table if isinstance(row, list)])
        result["tables"] = norm_tables
        children = result.get("children")
        result["children"] = [self.__normalize_node(c) for c in children] if isinstance(children, list) else []
        return result

    def __normalize_content(self, content):
        if not isinstance(content, dict) or not isinstance(content.get("sections"), list):
            return copy.deepcopy(DEFAULT_FTR_CONTENT)
        return {"sections": [self.__normalize_node(s) for s in content["sections"]]}

    def __replace_name(self, node, base, name, skip_titles=None):
        if not name or base == name:
            return
        if skip_titles and str(node.get("title") or "").strip() in skip_titles:
            return
        if node.get("body"):
            node["body"] = node["body"].replace(base, name)
        for tbl in (node.get("tables") or []):
            for row in tbl:
                for i in range(len(row)):
                    if isinstance(row[i], str) and base in row[i]:
                        row[i] = row[i].replace(base, name)
        for c in (node.get("children") or []):
            self.__replace_name(c, base, name, skip_titles=skip_titles)

    def __dhf_file_no(self, prod_id):
        if not prod_id:
            return ""
        row = db.session.execute(
            select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name == DOC_NAME).order_by(ProdDhf.id.asc())
        ).scalars().first()
        if not row:
            row = db.session.execute(
                select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name.like(f"%{DOC_NAME}%")).order_by(ProdDhf.id.asc())
            ).scalars().first()
        return (row.code or "").strip() if row and row.code else ""

    def __fill_revision(self, content, prod_id, version, force=False):
        rev_date = serv_review_util.cover_date(prod_id, DOC_KEY) if prod_id else ""
        reviser = approver = ""
        if prod_id:
            members = db.session.execute(select(ProjectMember).where(ProjectMember.prod_id == prod_id)).scalars().all()
            reviser = next((m.name for m in members if "TPM" in str(m.role or "")), "")
            approver = next((m.name for m in members if "研发负责人" in str(m.role or "")), "")
        for s in (content.get("sections") or []):
            if s.get("ref_type") != "revision":
                continue
            tables = s.get("tables") or []
            if tables and isinstance(tables[0], list) and tables[0]:
                t = tables[0]
                cols = len(t[0]) if isinstance(t[0], list) and t[0] else 5
                while len(t) < 6:
                    t.append([""] * cols)
                if len(t) >= 2 and isinstance(t[1], list) and len(t[1]) >= 3:
                    row = t[1]
                    if force:
                        row[0] = rev_date or ""
                        if len(row) >= 2: row[1] = str(version or "")
                        if len(row) >= 3: row[2] = "首次发布"
                        if len(row) >= 4: row[3] = reviser
                        if len(row) >= 5: row[4] = approver
                    else:
                        if not str(row[0] or "").strip(): row[0] = rev_date
                        if version and len(row) >= 2 and not str(row[1] or "").strip(): row[1] = str(version)
                        if len(row) >= 3 and not str(row[2] or "").strip(): row[2] = "首次发布"
                        if len(row) >= 4 and reviser and not str(row[3] or "").strip(): row[3] = reviser
                        if len(row) >= 5 and approver and not str(row[4] or "").strip(): row[4] = approver
            break
        return content

    def __fill_runtime_env(self, content, prod_id, force=False):
        if not isinstance(content, dict):
            return content
        env = dict(DEFAULT_RUNTIME_ENV)
        if prod_id:
            row = db.session.execute(select(ProdRuntimeEnv).where(ProdRuntimeEnv.prod_id == prod_id)).scalars().first()
            if row:
                for key in DEFAULT_RUNTIME_ENV.keys():
                    val = getattr(row, key, None)
                    if val is not None and str(val).strip():
                        env[key] = val
        def strip_title(title):
            return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()
        def overwrite_col1(table, label_map):
            for r in table:
                if not isinstance(r, list) or len(r) < 2:
                    continue
                key = str(r[0]).strip()
                if force:
                    r[1] = label_map.get(key, "") or ""
                elif key in label_map and str(label_map[key] or "").strip():
                    r[1] = label_map[key]
        def fill_node(node):
            title = str(node.get("title") or "")
            plain = strip_title(title)
            tables = node.get("tables") or []
            # 运行环境节点：4张表按顺序 — 表1服务器硬件/表2服务器软件/表3用户端/表4网络
            if "运行环境" in title and len(tables) >= 4:
                overwrite_col1(tables[0], {"CPU": env.get("srv_cpu"), "内存": env.get("srv_memory"), "GPU": env.get("srv_gpu"), "硬盘": env.get("srv_disk"), "网卡": env.get("srv_nic")})
                tbl2 = tables[1]
                if len(tbl2) >= 2 and isinstance(tbl2[1], list) and len(tbl2[1]) >= 3:
                    if force or str(env.get("srv_os") or "").strip(): tbl2[1][1] = env.get("srv_os", "")
                    if force or str(env.get("srv_cuda") or "").strip(): tbl2[1][2] = env.get("srv_cuda", "")
                overwrite_col1(tables[2], {"CPU": env.get("cli_cpu"), "内存": env.get("cli_memory"), "显示器分辨率": env.get("cli_resolution"), "操作系统": env.get("cli_os"), "浏览器": env.get("cli_browser")})
                for r in tables[3]:
                    if not isinstance(r, list) or str(r[0]).strip() != "带宽" or len(r) < 3:
                        continue
                    if force:
                        r[1] = env.get("net_lan", "")
                        r[2] = env.get("net_wan", "")
                    else:
                        if str(env.get("net_lan") or "").strip(): r[1] = env["net_lan"]
                        if str(env.get("net_wan") or "").strip(): r[2] = env["net_wan"]
            elif "表1" in title or plain.startswith("服务器硬件"):
                for tbl in tables:
                    overwrite_col1(tbl, {"CPU": env.get("srv_cpu"), "内存": env.get("srv_memory"), "GPU": env.get("srv_gpu"), "硬盘": env.get("srv_disk"), "网卡": env.get("srv_nic")})
            elif "表2" in title or "服务器软件" in plain:
                for tbl in tables:
                    if len(tbl) >= 2 and isinstance(tbl[1], list) and len(tbl[1]) >= 3:
                        if force or str(env.get("srv_os") or "").strip(): tbl[1][1] = env.get("srv_os", "")
                        if force or str(env.get("srv_cuda") or "").strip(): tbl[1][2] = env.get("srv_cuda", "")
            elif "表3" in title or plain.startswith("用户端"):
                for tbl in tables:
                    overwrite_col1(tbl, {"CPU": env.get("cli_cpu"), "内存": env.get("cli_memory"), "显示器分辨率": env.get("cli_resolution"), "操作系统": env.get("cli_os"), "浏览器": env.get("cli_browser")})
            elif "表4" in title or "网络" in plain:
                for tbl in tables:
                    for r in tbl:
                        if not isinstance(r, list) or str(r[0]).strip() != "带宽" or len(r) < 3:
                            continue
                        if force:
                            r[1] = env.get("net_lan", "")
                            r[2] = env.get("net_wan", "")
                        else:
                            if str(env.get("net_lan") or "").strip(): r[1] = env["net_lan"]
                            if str(env.get("net_wan") or "").strip(): r[2] = env["net_wan"]
            for c in (node.get("children") or []):
                fill_node(c)
        for s in (content.get("sections") or []):
            if "运行环境" in str(s.get("title") or ""):
                fill_node(s)
            else:
                for c in (s.get("children") or []):
                    if "运行环境" in str(c.get("title") or ""):
                        fill_node(c)
        return content

    def __autofill(self, content, prod_id, product=None, version="", force=False):
        if not isinstance(content, dict):
            return content
        if product and product.name:
            for s in (content.get("sections") or []):
                self.__replace_name(s, BASE_NAME, product.name)
        self.__fill_revision(content, prod_id, version, force=force)
        self.__fill_runtime_env(content, prod_id, force=force)
        serv_review_util.ensure_review(
            content, DOC_KEY,
            serv_review_util.review_date(prod_id, serv_review_util.REVIEW_DEFS.get(DOC_KEY, {}).get("name_keywords", [DOC_NAME])) if prod_id else "",
            prod_id,
        )
        serv_review_util.fill_cover_dates(content, serv_review_util.cover_date(prod_id, DOC_KEY) if prod_id else "", force=force)
        serv_review_util.fill_cover_signers(content, serv_review_util.cover_signers(prod_id, DOC_KEY) if prod_id else {}, force=force)
        return content

    def __to_obj(self, row: FtrDoc, product: Product = None):
        obj = FtrDocObj(**row.dict())
        obj.content = self.__autofill(self.__normalize_content(obj.content), row.product_id, product, row.version)
        if not (obj.file_no or "").strip():
            obj.file_no = self.__dhf_file_no(row.product_id)
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
        return obj

    async def add_ftr_doc(self, form: FtrDocForm):
        try:
            sql = select(func.count(FtrDoc.id)).where(FtrDoc.product_id == form.product_id, FtrDoc.version == form.version)
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = FtrDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=FtrDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_ftr_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: FtrDoc = db.session.execute(select(FtrDoc).where(FtrDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(FtrDoc.version).where(FtrDoc.product_id == target_pid)).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = FtrDoc(
                product_id=target_pid, version=version,
                file_no=sync_file_no_version(fromdoc.file_no, version),
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            if target_pid != fromdoc.product_id:
                await self.rebind_product(newdoc.id, target_pid)
            return Resp.resp_ok(data=FtrDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_ftr_doc(self, form: FtrDocForm):
        try:
            row: FtrDoc = db.session.execute(select(FtrDoc).where(FtrDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def rebind_product(self, id: int, product_id: int):
        try:
            row: FtrDoc = db.session.execute(select(FtrDoc).where(FtrDoc.id == id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            product: Product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
            if not product:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            db.session.execute(delete(FtrDoc).where(FtrDoc.product_id == product_id, FtrDoc.version == row.version, FtrDoc.id != id))
            content = copy.deepcopy(DEFAULT_FTR_CONTENT) if isinstance(DEFAULT_FTR_CONTENT, dict) else {"sections": []}
            new_full_version = (product.full_version or "").strip()
            def update_version(node):
                if node.get("body") and "完整版本" in str(node.get("body")):
                    node["body"] = re.sub(r"完整版本：[^\n]*", f"完整版本：{new_full_version}", str(node["body"]))
                if node.get("body") and re.search(r"\d+\.\d+\.\d+\.\d+", str(node.get("body"))):
                    node["body"] = re.sub(r"\d+\.\d+\.\d+\.\d+", new_full_version, str(node["body"]))
                for c in (node.get("children") or []):
                    update_version(c)
            for s in (content.get("sections") or []):
                update_version(s)
            row.product_id = product_id
            content = self.__autofill(content, product_id, product, row.version, force=True)
            row.content = content
            db.session.commit()
            return Resp.resp_ok(data=self.__to_obj(row, product))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_ftr_doc(self, id: int):
        db.session.execute(delete(FtrDoc).where(FtrDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_ftr_doc(self, id: int):
        sql = select(FtrDoc, Product).join(Product, FtrDoc.product_id == Product.id).where(FtrDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_ftr_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(FtrDoc.product_id == product_id)
        if version:
            wheres.append(FtrDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(FtrDoc.id)).join(Product, FtrDoc.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (select(FtrDoc, Product).join(Product, FtrDoc.product_id == Product.id).where(*wheres)
               .order_by(FtrDoc.id.desc()).offset(page_index * page_size).limit(page_size))
        rows: List[FtrDocObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    async def export_ftr_doc(self, output, id: int):
        resp = await self.get_ftr_doc(id)
        obj: FtrDocObj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        content = obj.content if isinstance(obj.content, dict) else self.__normalize_content(obj.content)
        sections = content.get("sections") or []

        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        document.settings.element.append(update_fields)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")
        docx_util.add_page_number_footer(section, obj.file_no or "")

        def strip_num(title):
            return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()

        cover = next((s for s in sections if s.get("ref_type") == "cover"), None)
        revision = next((s for s in sections if s.get("ref_type") == "revision"), None)
        body = [s for s in sections if s.get("ref_type") not in ("cover", "revision")]

        def add_blank_lines(n):
            for _ in range(max(0, int(n or 0))):
                document.add_paragraph("")

        def write_center_title(text, size=22.0, bold=False):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            docx_util.fonted_txt(p, str(text or ""), font_size=size, bold=bold)

        def add_text(text):
            docx_util.save_txt2docx(str(text or ""), document)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            s = str(text or "")
            cell.text = ""
            for i, line in enumerate(s.split("\n")):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        def add_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for r_idx, row in enumerate(grid):
                cells = table.add_row().cells
                for c_idx in range(cols):
                    set_cell(cells[c_idx], row[c_idx] if c_idx < len(row) else "", bold=(r_idx == 0))
            document.add_paragraph()

        def add_cover_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for row in grid:
                cells = table.add_row().cells
                for c_idx in range(cols):
                    text = row[c_idx] if c_idx < len(row) else ""
                    set_cell(cells[c_idx], text, bold=(c_idx % 2 == 0), align=WD_ALIGN_PARAGRAPH.CENTER)
            document.add_paragraph()

        def add_body_heading(title, level):
            size = {1: 16.0, 2: 14.0, 3: 12.0}.get(level, 11.0)
            p = document.add_heading("", level=max(1, min(level, 9)))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            docx_util.fonted_txt(p, title, font_size=size, bold=True)

        def render_body_section(node, level, number="", numbered=True):
            name = strip_num(node.get("title"))
            heading = f"{number} {name}".strip() if number else name
            if heading:
                add_body_heading(heading, level=max(1, min(level, 9)))
            if (node.get("body") or "").strip():
                add_text(node.get("body"))
            for table in (node.get("tables") or []):
                add_grid(table)
            idx = 0
            for child in (node.get("children") or []):
                idx += 1
                if numbered and number:
                    child_num = f"{number}.{idx}"
                elif numbered:
                    child_num = str(idx)
                else:
                    child_num = ""
                render_body_section(child, level + 1, child_num, numbered=numbered)

        add_blank_lines(6)
        write_center_title((strip_num(cover.get("title")) if cover else "") or DOC_NAME, size=22.0, bold=True)
        add_blank_lines(4)
        if cover:
            for table in (cover.get("tables") or []):
                add_cover_grid(table)

        document.add_page_break()
        write_center_title("文件修订记录", size=14.0, bold=True)
        add_blank_lines(2)
        if revision:
            for table in (revision.get("tables") or []):
                add_grid(table)

        document.add_page_break()
        write_center_title("目录", size=16.0, bold=True)
        docx_util.insert_toc_field(document)

        document.add_page_break()
        seq = 0
        for node in body:
            seq += 1
            render_body_section(node, 1, str(seq), numbered=True)

        docx_util.fill_toc_cache(document)
        document.save(output)
        output.seek(0)