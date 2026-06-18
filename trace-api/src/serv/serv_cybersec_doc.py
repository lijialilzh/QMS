#!/usr/bin/env python
# encoding: utf-8

# 网络安全管理服务层，严格对应 docs/function_docs/47_网络安全管理.md。
# 与风险管理报告（serv_risk_mgmt_doc）零耦合：独立模型、独立常量 CYBERSEC_SCORE_TABLE、独立内容树。

import logging
import base64
import copy
import io
import json
import os
import re
from typing import List
from sqlalchemy import delete, func, select
from docx import Document
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.cybersec_doc import (
    CybersecDoc,
    CybersecThreat,
    CybersecControlInternal,
    CybersecControlSbom,
    CybersecControlScan,
)
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_cybersec_doc import (
    CybersecDocForm,
    CybersecThreatForm,
    CybersecControlInternalForm,
    CybersecControlSbomForm,
    CybersecControlScanForm,
)
from ..obj.vobj_cybersec_doc import (
    CybersecDocObj,
    CybersecThreatObj,
    CybersecControlInternalObj,
    CybersecControlSbomObj,
    CybersecControlScanObj,
)
from ..utils.i18n import ts
from ..utils import get_uuid
from ..utils.sql_ctx import db
from . import msg_err_db
from .serv_utils import new_version
from .serv_utils import docx_util

logger = logging.getLogger(__name__)


# 5×5 likelihood × severity 评分标准（合规内容，红线不可随意改动；与 RISK_ACCEPTANCE_TABLE 同图例但独立常量）
CYBERSEC_SCORE_TABLE = [
    ["风险值", "", "", "严重度", "", "", "", ""],
    ["", "", "", "可忽略 A", "轻度 B", "严重 C", "危重的 D", "灾难性的 E"],
    ["发生概率", "经常", "5", "5A", "5B", "5C", "5D", "5E"],
    ["", "有时", "4", "4A", "4B", "4C", "4D", "4E"],
    ["", "偶然", "3", "3A", "3B", "3C", "3D", "3E"],
    ["", "很少", "2", "2A", "2B", "2C", "2D", "2E"],
    ["", "非常少", "1", "1A", "1B", "1C", "1D", "1E"],
    ["红色", "不可接受：这类网络安全风险本质上不可接受，必须寻求风险降低措施。", "", "", "", "", "", ""],
    ["橙色", "可控：需进一步降低到合理可行的最低限度才可视为可接受。", "", "", "", "", "", ""],
    ["绿色", "可接受：这类风险实际上可接受。", "", "", "", "", "", ""],
]


# 内容树（文档 47 第 2.3 节 1~7 章），ref_type 取值见文档第 5 节
DEFAULT_CYBERSEC_CONTENT = {
    "sections": [
        {"title": "网络安全风险管理报告", "ref_type": "cover", "children": [], "tables": [[
            ["编制部门", "", "文件版本", ""],
            ["编制人", "", "日期", ""],
            ["审核人", "", "日期", ""],
            ["批准人", "", "日期", ""],
            ["生效日期", "", "", ""],
        ]]},
        {"title": "文件修订记录", "ref_type": "revision", "children": [], "tables": [[
            ["修改日期", "版本号", "修订说明", "修订人", "批准人"],
            ["", "", "", "", ""],
            ["", "", "", "", ""],
            ["", "", "", "", ""],
            ["", "", "", "", ""],
            ["", "", "", "", ""],
        ]]},
        {"title": "1 概述", "children": [
            {"title": "1.1 目的", "children": []},
            {"title": "1.2 产品描述", "children": []},
            {"title": "1.3 适用范围", "children": []},
            {"title": "1.4 系统架构和安全实现", "children": [
                {"title": "1.4.1 系统总体架构", "ref_type": "flow_diagram", "children": []},
            ]},
        ]},
        {"title": "2 阶段活动", "ref_type": "stage_activity", "children": [], "tables": [[
            ["阶段", "开始", "结束", "结果"],
            ["", "", "", ""],
        ]]},
        {"title": "3 关联文件", "children": []},
        {"title": "4 视图分析与威胁建模", "ref_type": "view_analysis", "children": [
            {"title": "4.1 系统全局视图", "children": [
                {"title": "4.1.1 核心组件", "children": []},
                {"title": "4.1.2 数据流和网络安全控制", "children": []},
                {"title": "4.1.3 外部连接与安全影响", "children": []},
                {"title": "4.1.4 外部节点的控制与信任边界", "children": []},
                {"title": "4.1.5 安全控制映射", "children": []},
                {"title": "4.1.6 威胁评估", "children": []},
            ]},
            {"title": "4.2 多患者危害视图", "children": [
                {"title": "4.2.1 安全措施", "children": []},
                {"title": "4.2.2 非医疗功能引发的多患者危害场景", "children": []},
                {"title": "4.2.3 攻击链与多患者风险路径图", "children": []},
                {"title": "4.2.4 多患者攻击的响应与控制", "children": []},
                {"title": "4.2.5 威胁分类", "children": []},
            ]},
            {"title": "4.3 安全用例视图", "children": [
                {"title": "4.3.1 目标作用", "children": []},
                {"title": "4.3.2 用例分类与范围", "children": []},
                {"title": "4.3.3 用例及威胁缓解", "children": []},
            ]},
            {"title": "4.4 可更新性视图", "children": [
                {"title": "4.4.1 更新类型与目标", "children": []},
                {"title": "4.4.2 更新策略与流程", "children": []},
                {"title": "4.4.3 更新频率与计划", "children": []},
                {"title": "4.4.4 回退机制", "children": []},
                {"title": "4.4.5 兼容性与依赖关系", "children": []},
                {"title": "4.4.6 安全性与合规性", "children": []},
            ]},
            {"title": "4.5 威胁建模（STRIDE）", "children": [
                {"title": "4.5.1 威胁分类", "children": []},
                {"title": "4.5.2 威胁评估", "ref_type": "stride_threats", "children": []},
            ]},
        ]},
        {"title": "5 风险评估", "children": [
            {"title": "5.1 风险评估评分标准", "children": [], "tables": [copy.deepcopy(CYBERSEC_SCORE_TABLE)]},
            {"title": "5.2 风险评估及控制措施（RCM）", "children": [
                {"title": "5.2.1 内部风险评估及控制措施", "ref_type": "cybersec_controls_internal", "children": []},
                {"title": "5.2.2 SBOM风险评估及控制措施", "ref_type": "cybersec_controls_sbom", "children": []},
                {"title": "5.2.3 网络安全扫描风险评估及控制措施", "ref_type": "cybersec_controls_scan", "children": []},
            ]},
            {"title": "5.3 残余风险评估", "children": [
                {"title": "5.3.1 残留风险结论", "ref_type": "residual_risk", "children": []},
            ]},
        ]},
        {"title": "6 维护更新", "ref_type": "maintenance", "children": [
            {"title": "6.1 设计保证", "children": []},
            {"title": "6.2 异常情况响应", "children": []},
            {"title": "6.3 安全更新策略", "children": []},
            {"title": "6.4 用户指导", "children": []},
        ]},
        {"title": "7 威胁缓解措施追溯", "ref_type": "traceability", "children": []},
    ],
    "productName": "",
}


# 默认内容优先从资源文件加载（含各章节默认正文/表格/图片；自动获取章节为模板态）；
# 文件缺失时回退到上面的内联结构。资源文件由编辑页参考文档生成，前端新增页用同一份 JSON。
_DEFAULT_CONTENT_FILE = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "src-res", "cybersec_default_content.json"
)
try:
    with open(_DEFAULT_CONTENT_FILE, encoding="utf-8") as _f:
        _loaded = json.load(_f)
    if isinstance(_loaded, dict) and isinstance(_loaded.get("sections"), list) and _loaded["sections"]:
        DEFAULT_CYBERSEC_CONTENT = _loaded
except Exception:
    logging.getLogger(__name__).exception("加载网络安全默认内容资源失败，使用内联默认模版")


class Server(object):
    def __normalize_title_text(self, value):
        return re.sub(r"\s+", "", str(value or ""))

    def __is_cover_section(self, section):
        title = self.__normalize_title_text((section or {}).get("title", ""))
        return (section or {}).get("ref_type") == "cover" or title == "网络安全风险管理报告"

    def __is_revision_section(self, section):
        title = self.__normalize_title_text((section or {}).get("title", ""))
        return (section or {}).get("ref_type") == "revision" or title == "文件修订记录"

    def __ensure_front_matter_sections(self, content):
        next_content = copy.deepcopy(content or {})
        sections = next_content.get("sections") if isinstance(next_content.get("sections"), list) else []
        default_sections = copy.deepcopy(DEFAULT_CYBERSEC_CONTENT["sections"])
        default_cover = default_sections[0]
        default_revision = default_sections[1]
        cover = next((item for item in sections if self.__is_cover_section(item)), None) or default_cover
        revision = next((item for item in sections if self.__is_revision_section(item)), None) or default_revision
        body_sections = [item for item in sections if not self.__is_cover_section(item) and not self.__is_revision_section(item)]
        next_content["sections"] = [cover, revision] + body_sections
        return next_content

    def __normalize_content(self, content):
        result = copy.deepcopy(DEFAULT_CYBERSEC_CONTENT)
        if isinstance(content, dict):
            result.update(content)
        result.setdefault("sections", copy.deepcopy(DEFAULT_CYBERSEC_CONTENT["sections"]))
        result.setdefault("productName", "")
        return self.__ensure_front_matter_sections(result)

    def __to_obj(self, row: CybersecDoc, product: Product = None):
        obj = CybersecDocObj(**row.dict())
        obj.content = self.__normalize_content(obj.content)
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
        return obj

    def __fill_list_obj(self, obj, product: Product = None, doc: CybersecDoc = None):
        if product:
            obj.product_name = product.name
            obj.product_full_version = product.full_version
        if doc:
            obj.doc_version = doc.version
        return obj

    def __sync_product_id_from_doc(self, form):
        if form.doc_id and not form.product_id:
            doc = db.session.execute(select(CybersecDoc).where(CybersecDoc.id == form.doc_id)).scalars().first()
            if doc:
                form.product_id = doc.product_id
        return form

    # ---------------- Word 导入解析 ----------------
    def __normalize_section_title(self, value: str):
        txt = re.sub(r"\s+", "", str(value or "").strip())
        txt = re.sub(r"^[0-9０-９]+(?:[.．][0-9０-９]+)*(?:[、.．\s]+|(?=[\u4e00-\u9fffA-Za-z]))", "", txt)
        return txt

    def __iter_docx_blocks(self, docx: Document):
        for child in docx.element.body.iterchildren():
            tag = str(child.tag).lower()
            if tag.endswith("}p"):
                yield Paragraph(child, docx._body)
            elif tag.endswith("}tbl"):
                yield DocxTable(child, docx._body)

    def __extract_docx_paragraph_images(self, para: Paragraph):
        urls = []
        used_rids = set()
        mime_map = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg", "gif": "image/gif", "bmp": "image/bmp", "webp": "image/webp"}
        blips = para._element.xpath(".//*[local-name()='blip']")
        for blip in blips:
            rid = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            if not rid or rid in used_rids:
                continue
            used_rids.add(rid)
            try:
                rel = para.part.rels[rid]
                target = getattr(rel, "target_ref", "")
                ext = (target.rsplit(".", 1)[-1].lower() if "." in target else "png")
                mime = mime_map.get(ext, "image/png")
                b64 = base64.b64encode(rel.target_part.blob).decode("ascii")
                urls.append(f"data:{mime};base64,{b64}")
            except Exception:
                logger.exception("解析网络安全 Word 图片失败")
        return urls

    def __parse_docx_table(self, table: DocxTable):
        rows = []
        for row in table.rows:
            values = []
            prev_tc = None
            for cell in row.cells:
                tc = cell._tc
                if prev_tc is not None and tc is prev_tc:
                    # 横向合并的延续单元格：python-docx 会重复返回同一文本，置空避免重复
                    values.append("")
                else:
                    values.append(re.sub(r"\s+", " ", (cell.text or "").strip()))
                prev_tc = tc
            if any(values):
                rows.append(values)
        return rows

    def __extract_cybersec_content_from_word(self, docx: Document):
        content = self.__normalize_content(None)
        # 按模板 DFS 顺序记录章节，序列匹配以正确处理重复标题（如「威胁分类」「威胁评估」各出现两次）
        ordered = []

        def walk(items):
            for section in items or []:
                ordered.append((self.__normalize_section_title(section.get("title", "")), section))
                walk(section.get("children") or [])

        walk(content.get("sections") or [])
        cursor = 0
        current_section = None
        # 记录本次导入中已落过表的章节：首次遇到导入表时先清空模板默认表，避免「默认空表 + 导入表」重复
        table_cleared_sections = set()

        for block in self.__iter_docx_blocks(docx):
            if isinstance(block, Paragraph):
                txt = re.sub(r"\s+", " ", (block.text or "").strip())
                images = self.__extract_docx_paragraph_images(block)
                if txt:
                    key = self.__normalize_section_title(txt)
                    # 优先从游标向后找同名章节（按出现顺序消费重复标题），找不到再全局兜底
                    idx = next((i for i in range(cursor, len(ordered)) if ordered[i][0] == key), None)
                    if idx is None:
                        idx = next((i for i in range(0, len(ordered)) if ordered[i][0] == key), None)
                    if idx is not None:
                        current_section = ordered[idx][1]
                        cursor = idx + 1
                    elif current_section is not None:
                        current_text = str(current_section.get("text") or "").strip()
                        current_section["text"] = f"{current_text}\n{txt}".strip() if current_text else txt
                if images and current_section is not None:
                    current_section.setdefault("images", []).extend(images)
            elif isinstance(block, DocxTable) and current_section is not None:
                rows = self.__parse_docx_table(block)
                if rows:
                    sid = id(current_section)
                    if sid not in table_cleared_sections:
                        current_section["tables"] = []
                        table_cleared_sections.add(sid)
                    current_section["tables"].append(rows)
        return content

    def __save_imported_word_bytes(self, doc_id: int, filename: str, bys: bytes):
        suffix = (os.path.splitext(filename or "")[1] or ".docx").lower()
        path = os.path.join("data.trace", "cybersec_doc_word", str(doc_id), get_uuid() + suffix)
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "wb") as fs:
            fs.write(bys or b"")
        return path

    # ---------------- 文档 CRUD ----------------
    async def add_cybersec_doc(self, form: CybersecDocForm):
        try:
            sql = select(func.count(CybersecDoc.id)).where(
                CybersecDoc.product_id == form.product_id,
                CybersecDoc.version == form.version,
            )
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = CybersecDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def import_cybersec_doc_word(self, product_id: int, version: str, file_no: str = "", change_log: str = "", file=None):
        content = None
        bys = None
        filename = ""
        if file is not None:
            bys = await file.read()
            filename = file.filename or ""
            docx = Document(io.BytesIO(bys))
            content = self.__extract_cybersec_content_from_word(docx)
        form = CybersecDocForm(product_id=product_id, version=version, file_no=file_no, change_log=change_log, content=content)
        resp = await self.add_cybersec_doc(form)
        if resp.code == 1 and bys:
            row = db.session.execute(
                select(CybersecDoc).where(CybersecDoc.product_id == product_id, CybersecDoc.version == version)
            ).scalars().first()
            if row:
                next_content = self.__normalize_content(row.content)
                next_content["sourceWordPath"] = self.__save_imported_word_bytes(row.id, filename, bys)
                row.content = next_content
                db.session.commit()
        return resp

    async def duplicate_cybersec_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: CybersecDoc = db.session.execute(select(CybersecDoc).where(CybersecDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(CybersecDoc.version).where(CybersecDoc.product_id == target_pid)).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _version_seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_version_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = CybersecDoc(
                product_id=target_pid,
                version=version,
                file_no=fromdoc.file_no,
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.flush()

            for model in (CybersecThreat, CybersecControlInternal, CybersecControlSbom, CybersecControlScan):
                for item in db.session.execute(select(model).where(model.doc_id == fromdoc.id)).scalars().all():
                    newitem = model(**item.dict())
                    newitem.id = None
                    newitem.doc_id = newdoc.id
                    newitem.product_id = target_pid
                    db.session.add(newitem)
            db.session.commit()
            return Resp.resp_ok(data=CybersecDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_cybersec_doc(self, form: CybersecDocForm):
        try:
            row: CybersecDoc = db.session.execute(select(CybersecDoc).where(CybersecDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_cybersec_doc(self, id: int):
        db.session.execute(delete(CybersecThreat).where(CybersecThreat.doc_id == id))
        db.session.execute(delete(CybersecControlInternal).where(CybersecControlInternal.doc_id == id))
        db.session.execute(delete(CybersecControlSbom).where(CybersecControlSbom.doc_id == id))
        db.session.execute(delete(CybersecControlScan).where(CybersecControlScan.doc_id == id))
        db.session.execute(delete(CybersecDoc).where(CybersecDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_cybersec_doc(self, id: int):
        sql = select(CybersecDoc, Product).join(Product, CybersecDoc.product_id == Product.id).where(CybersecDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_cybersec_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(CybersecDoc.product_id == product_id)
        if version:
            wheres.append(CybersecDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(CybersecDoc.id)).join(Product, CybersecDoc.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (
            select(CybersecDoc, Product)
            .join(Product, CybersecDoc.product_id == Product.id)
            .where(*wheres)
            .order_by(CybersecDoc.id.desc())
            .offset(page_index * page_size)
            .limit(page_size)
        )
        rows: List[CybersecDocObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    # ---------------- 威胁 CRUD ----------------
    async def add_cybersec_threat(self, form: CybersecThreatForm):
        try:
            form = self.__sync_product_id_from_doc(form)
            sql = select(func.count(CybersecThreat.id)).where(CybersecThreat.doc_id == form.doc_id, CybersecThreat.threat_code == form.threat_code)
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = CybersecThreat(**form.dict(exclude_none=True))
            row.id = None
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_cybersec_threat(self, form: CybersecThreatForm):
        try:
            form = self.__sync_product_id_from_doc(form)
            row: CybersecThreat = db.session.execute(select(CybersecThreat).where(CybersecThreat.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            sql = select(func.count(CybersecThreat.id)).where(
                CybersecThreat.doc_id == form.doc_id,
                CybersecThreat.threat_code == form.threat_code,
                CybersecThreat.id != form.id,
            )
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            for key, value in form.dict(exclude_none=True).items():
                if key != "id":
                    setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_cybersec_threat(self, id: int):
        db.session.execute(delete(CybersecThreat).where(CybersecThreat.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def list_cybersec_threat(self, product_id: int = 0, doc_id: int = 0, view_type: str = None, keyword: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(CybersecThreat.product_id == product_id)
        if doc_id:
            wheres.append(CybersecThreat.doc_id == doc_id)
        if view_type:
            wheres.append(CybersecThreat.view_type == view_type)
        if keyword:
            like = f"%{keyword}%"
            wheres.append((CybersecThreat.threat_code.like(like)) | (CybersecThreat.asset.like(like)) | (CybersecThreat.description.like(like)))
        total = db.session.execute(select(func.count(CybersecThreat.id)).where(*wheres)).scalar() or 0
        sql = (
            select(CybersecThreat, CybersecDoc, Product)
            .join(CybersecDoc, CybersecThreat.doc_id == CybersecDoc.id)
            .join(Product, CybersecThreat.product_id == Product.id)
            .where(*wheres)
            .order_by(CybersecThreat.id.desc())
            .offset(page_index * page_size)
            .limit(page_size)
        )
        rows = [self.__fill_list_obj(CybersecThreatObj(**row.dict()), product, doc) for row, doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    # ---------------- 三类 RCM 控制措施 CRUD（结构同构，独立持久化） ----------------
    async def __add_control(self, model, obj_cls, form):
        try:
            form = self.__sync_product_id_from_doc(form)
            sql = select(func.count(model.id)).where(model.doc_id == form.doc_id, model.rcm_code == form.rcm_code)
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = model(**form.dict(exclude_none=True))
            row.id = None
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def __update_control(self, model, form):
        try:
            form = self.__sync_product_id_from_doc(form)
            row = db.session.execute(select(model).where(model.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            sql = select(func.count(model.id)).where(model.doc_id == form.doc_id, model.rcm_code == form.rcm_code, model.id != form.id)
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            for key, value in form.dict(exclude_none=True).items():
                if key != "id":
                    setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def __delete_control(self, model, id: int):
        db.session.execute(delete(model).where(model.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def __list_control(self, model, obj_cls, product_id: int = 0, doc_id: int = 0, keyword: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(model.product_id == product_id)
        if doc_id:
            wheres.append(model.doc_id == doc_id)
        if keyword:
            like = f"%{keyword}%"
            wheres.append((model.rcm_code.like(like)) | (model.description.like(like)) | (model.threat_codes.like(like)))
        total = db.session.execute(select(func.count(model.id)).where(*wheres)).scalar() or 0
        sql = (
            select(model, CybersecDoc, Product)
            .join(CybersecDoc, model.doc_id == CybersecDoc.id)
            .join(Product, model.product_id == Product.id)
            .where(*wheres)
            .order_by(model.id.desc())
            .offset(page_index * page_size)
            .limit(page_size)
        )
        rows = [self.__fill_list_obj(obj_cls(**row.dict()), product, doc) for row, doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    async def add_cybersec_control_internal(self, form: CybersecControlInternalForm):
        return await self.__add_control(CybersecControlInternal, CybersecControlInternalObj, form)

    async def update_cybersec_control_internal(self, form: CybersecControlInternalForm):
        return await self.__update_control(CybersecControlInternal, form)

    async def delete_cybersec_control_internal(self, id: int):
        return await self.__delete_control(CybersecControlInternal, id)

    async def list_cybersec_control_internal(self, product_id: int = 0, doc_id: int = 0, keyword: str = None, page_index: int = 0, page_size: int = 10):
        return await self.__list_control(CybersecControlInternal, CybersecControlInternalObj, product_id, doc_id, keyword, page_index, page_size)

    async def add_cybersec_control_sbom(self, form: CybersecControlSbomForm):
        return await self.__add_control(CybersecControlSbom, CybersecControlSbomObj, form)

    async def update_cybersec_control_sbom(self, form: CybersecControlSbomForm):
        return await self.__update_control(CybersecControlSbom, form)

    async def delete_cybersec_control_sbom(self, id: int):
        return await self.__delete_control(CybersecControlSbom, id)

    async def list_cybersec_control_sbom(self, product_id: int = 0, doc_id: int = 0, keyword: str = None, page_index: int = 0, page_size: int = 10):
        return await self.__list_control(CybersecControlSbom, CybersecControlSbomObj, product_id, doc_id, keyword, page_index, page_size)

    async def add_cybersec_control_scan(self, form: CybersecControlScanForm):
        return await self.__add_control(CybersecControlScan, CybersecControlScanObj, form)

    async def update_cybersec_control_scan(self, form: CybersecControlScanForm):
        return await self.__update_control(CybersecControlScan, form)

    async def delete_cybersec_control_scan(self, id: int):
        return await self.__delete_control(CybersecControlScan, id)

    async def list_cybersec_control_scan(self, product_id: int = 0, doc_id: int = 0, keyword: str = None, page_index: int = 0, page_size: int = 10):
        return await self.__list_control(CybersecControlScan, CybersecControlScanObj, product_id, doc_id, keyword, page_index, page_size)

    # ---------------- Word 导出 ----------------
    def __build_traceability_rows(self, doc_id: int):
        # 7. 威胁缓解措施追溯运行时计算：合并三张 control_* 表，按 (threat_code ↔ threat_codes / rcm_codes) 关联，带来源列
        threats = db.session.execute(select(CybersecThreat).where(CybersecThreat.doc_id == doc_id).order_by(CybersecThreat.id)).scalars().all()
        sources = [
            ("内部", db.session.execute(select(CybersecControlInternal).where(CybersecControlInternal.doc_id == doc_id)).scalars().all()),
            ("SBOM", db.session.execute(select(CybersecControlSbom).where(CybersecControlSbom.doc_id == doc_id)).scalars().all()),
            ("扫描", db.session.execute(select(CybersecControlScan).where(CybersecControlScan.doc_id == doc_id)).scalars().all()),
        ]
        rows = []
        for threat in threats:
            tcode = (threat.threat_code or "").strip()
            for source_name, controls in sources:
                for ctrl in controls:
                    linked_by_threat = tcode and tcode in str(ctrl.threat_codes or "")
                    linked_by_rcm = (ctrl.rcm_code or "").strip() and (ctrl.rcm_code or "").strip() in str(threat.rcm_codes or "")
                    if linked_by_threat or linked_by_rcm:
                        rows.append([tcode, threat.description or "", ctrl.rcm_code or "", ctrl.description or "", source_name])
        return rows

    async def export_cybersec_doc(self, output, id: int):
        resp = await self.get_cybersec_doc(id)
        obj: CybersecDocObj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        document.settings.element.append(update_fields)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")

        def normalized_title(value):
            return re.sub(r"\s+", "", str(value or ""))

        def is_cover_section(sec):
            title = normalized_title(sec.get("title", ""))
            return sec.get("ref_type") == "cover" or title == "网络安全风险管理报告"

        def is_revision_section(sec):
            title = normalized_title(sec.get("title", ""))
            return sec.get("ref_type") == "revision" or title == "文件修订记录"

        def write_center_section_title(title, font_size=16.0, bold=True):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.5
            docx_util.fonted_txt(p, title, font_size=font_size, bold=bold)

        def add_blank_lines(count):
            for _ in range(max(0, int(count or 0))):
                document.add_paragraph("")

        def insert_toc_field():
            p = document.add_paragraph()
            run_begin = p.add_run()
            fld_begin = OxmlElement("w:fldChar")
            fld_begin.set(qn("w:fldCharType"), "begin")
            fld_begin.set(qn("w:dirty"), "true")
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = ' TOC \\o "1-3" \\h \\z \\u '
            fld_separate = OxmlElement("w:fldChar")
            fld_separate.set(qn("w:fldCharType"), "separate")
            run_end = p.add_run()
            fld_end = OxmlElement("w:fldChar")
            fld_end.set(qn("w:fldCharType"), "end")
            run_begin._r.append(fld_begin)
            run_begin._r.append(instr)
            run_begin._r.append(fld_separate)
            p.add_run("目录将在打开文档后自动更新")
            run_end._r.append(fld_end)

        def set_cell_text(cell, text, bold=False):
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.line_spacing = 1.5
            docx_util.fonted_txt(paragraph, str(text or ""), font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        def add_plain_table(rows):
            rows = rows or []
            col_count = max([len(row or []) for row in rows] or [0])
            if col_count <= 0:
                return
            table = document.add_table(rows=0, cols=col_count)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for row in rows:
                cells = table.add_row().cells
                for idx in range(col_count):
                    set_cell_text(cells[idx], row[idx] if idx < len(row or []) else "", bold=(len(table.rows) == 1))
            document.add_paragraph()

        def set_cell_bg(cell, color_hex):
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), color_hex)
            tc_pr.append(shd)

        SCORE_GREEN = "00B050"
        SCORE_YELLOW = "FFFF00"
        SCORE_RED = "FF0000"

        def is_score_matrix(rows):
            has_rv = False
            has_sev = False
            for row in rows or []:
                for cell in row or []:
                    v = str(cell or "").strip()
                    if v == "风险值":
                        has_rv = True
                    elif v == "严重度":
                        has_sev = True
            return has_rv and has_sev

        def score_zone_color(likelihood, severity):
            # likelihood / severity 均为 1..5；与《风险评估评分标准》矩阵配色一致
            if likelihood <= 1:
                return SCORE_GREEN
            if likelihood in (2, 3):
                return SCORE_GREEN if severity <= 3 else SCORE_YELLOW
            if severity <= 1:
                return SCORE_GREEN
            if severity in (2, 3):
                return SCORE_YELLOW
            return SCORE_RED

        def add_score_matrix_table(rows):
            rows = rows or []
            n = len(rows)
            width = max([len(row or []) for row in rows] or [0])
            if width <= 0 or n <= 0:
                return

            def txt(r, c):
                row = rows[r] or []
                return str(row[c]).strip() if c < len(row) and row[c] is not None else ""

            skip = [[False] * width for _ in range(n)]
            span = [[None] * width for _ in range(n)]
            for r in range(n):
                for c in range(width):
                    if skip[r][c]:
                        continue
                    colspan = 1
                    while c + colspan < width and txt(r, c + colspan) == "":
                        skip[r][c + colspan] = True
                        colspan += 1
                    rowspan = 1
                    if txt(r, c) != "":
                        while r + rowspan < n and txt(r + rowspan, c) == txt(r, c):
                            for cc in range(c, c + colspan):
                                skip[r + rowspan][cc] = True
                            rowspan += 1
                    span[r][c] = (colspan, rowspan)

            likelihood_map = {}
            li = 0
            for r in range(n):
                if txt(r, 0) == "可利用性":
                    li += 1
                    likelihood_map[r] = li

            label_colors = {"红色": SCORE_RED, "黄色": SCORE_YELLOW, "绿色": SCORE_GREEN}
            table = document.add_table(rows=n, cols=width)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for r in range(n):
                for c in range(width):
                    if skip[r][c]:
                        continue
                    colspan, rowspan = span[r][c]
                    top_left = table.cell(r, c)
                    if colspan > 1 or rowspan > 1:
                        merged = top_left.merge(table.cell(r + rowspan - 1, c + colspan - 1))
                    else:
                        merged = top_left
                    value = txt(r, c)
                    bold = r <= 1 or value == "可利用性" or value in label_colors
                    set_cell_text(merged, value, bold=bold)
                    if value in label_colors:
                        set_cell_bg(merged, label_colors[value])
                    elif r in likelihood_map and c >= 2:
                        severity = c - 1
                        if 1 <= severity <= 5:
                            set_cell_bg(merged, score_zone_color(likelihood_map[r], severity))
            document.add_paragraph()

        def add_header_table(headers, value_rows):
            if not value_rows:
                return
            table = document.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for idx, header in enumerate(headers):
                set_cell_text(table.rows[0].cells[idx], header, bold=True)
            for value_row in value_rows:
                cells = table.add_row().cells
                for idx in range(len(headers)):
                    set_cell_text(cells[idx], value_row[idx] if idx < len(value_row) else "")
            document.add_paragraph()

        def add_stride_threats():
            rows = db.session.execute(
                select(CybersecThreat).where(CybersecThreat.doc_id == id, CybersecThreat.view_type == "STRIDE").order_by(CybersecThreat.id)
            ).scalars().all()
            headers = ["威胁编号", "STRIDE类别", "资产/对象", "威胁描述", "攻击路径", "影响", "可能性", "严重度", "风险水平", "控制措施"]
            value_rows = [[
                r.threat_code or "", r.stride_category or "", r.asset or "", r.description or "", r.attack_path or "",
                r.impact or "", str(r.likelihood) if r.likelihood is not None else "", r.severity or "", r.risk_level or "", r.control_measures or "",
            ] for r in rows]
            add_header_table(headers, value_rows)

        def add_control_table(model):
            rows = db.session.execute(select(model).where(model.doc_id == id).order_by(model.id)).scalars().all()
            headers = ["RCM编号", "控制措施描述", "关联威胁编号", "验证证据", "是否引入新风险", "备注"]
            value_rows = [[
                r.rcm_code or "", r.description or "", r.threat_codes or "", r.verification_evidence or "",
                "是" if r.new_risk_flag else "否", r.note or "",
            ] for r in rows]
            add_header_table(headers, value_rows)

        def add_traceability():
            headers = ["威胁编号", "威胁描述", "关联RCM编号", "控制措施描述", "来源"]
            add_header_table(headers, self.__build_traceability_rows(id))

        def add_section_image(image_url):
            raw_url = str(image_url or "").strip()
            if not raw_url:
                return
            try:
                if raw_url.startswith("data:image/"):
                    image_data = raw_url.split(",", 1)[1]
                    paragraph = document.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pic = paragraph.add_run().add_picture(io.BytesIO(base64.b64decode(image_data)))
                    # 等比缩放，限制最大宽/高，避免竖向流程图撑出页面；不放大小图
                    max_w = Inches(5.5)
                    max_h = Inches(7.0)
                    if pic.width and pic.height:
                        ratio = min(max_w / pic.width, max_h / pic.height, 1)
                        pic.width = int(pic.width * ratio)
                        pic.height = int(pic.height * ratio)
                    document.add_paragraph()
            except Exception:
                logger.exception("导出网络安全图片失败")

        def slim_trace_rcm_cells(rows):
            # 仅用于追溯表导出：含 RCM 编号的单元格（即 RCMID 列）只保留编号，去重并按数字排序、每个一行；
            # 不含 RCM 编号的单元格（威胁编号/SDS/测试/备注/表头）保持原样。
            out = []
            for row in rows or []:
                new_row = []
                for cell in (row or []):
                    codes = re.findall(r"RCM\d+", str(cell or ""))
                    if codes:
                        uniq = list(dict.fromkeys(codes))
                        uniq.sort(key=lambda c: int(re.sub(r"\D", "", c) or 0))
                        new_row.append("\n".join(uniq))
                    else:
                        new_row.append(cell)
                out.append(new_row)
            return out

        def add_one_table(table_rows, ref_type=None):
            if is_score_matrix(table_rows):
                add_score_matrix_table(table_rows)
            else:
                add_plain_table(slim_trace_rcm_cells(table_rows) if ref_type == "traceability" else table_rows)

        def add_section(sec: dict, level: int = 1):
            title = sec.get("title", "")
            if title and not is_cover_section(sec):
                docx_util.save_title2docx(title, document, level=max(1, min(level, 9)))
            # 有序内容块（text/table 交错）：按块顺序输出，忽略平铺 text/tables
            blocks = sec.get("blocks")
            if isinstance(blocks, list) and blocks:
                for image_url in sec.get("images", []) or []:
                    add_section_image(image_url)
                for block in blocks:
                    if not isinstance(block, dict):
                        continue
                    if block.get("type") == "table":
                        add_one_table(block.get("table") or [])
                    elif block.get("text"):
                        docx_util.save_txt2docx(str(block.get("text") or ""), document)
                for child in sec.get("children", []) or []:
                    add_section(child, level + 1)
                return
            if sec.get("text"):
                docx_util.save_txt2docx(str(sec.get("text") or ""), document)
            for image_url in sec.get("images", []) or []:
                add_section_image(image_url)
            ref_type = sec.get("ref_type")
            if ref_type == "stride_threats":
                add_stride_threats()
            elif ref_type == "cybersec_controls_internal":
                add_control_table(CybersecControlInternal)
            elif ref_type == "cybersec_controls_sbom":
                add_control_table(CybersecControlSbom)
            elif ref_type == "cybersec_controls_scan":
                add_control_table(CybersecControlScan)
            elif ref_type == "traceability":
                add_traceability()
            for table_rows in sec.get("tables", []) or []:
                add_one_table(table_rows, ref_type)
            for child in sec.get("children", []) or []:
                add_section(child, level + 1)

        export_content = self.__normalize_content(obj.content or {})
        export_sections = (export_content or {}).get("sections", [])
        cover_section = next((s for s in export_sections if is_cover_section(s)), None)
        revision_section = next((s for s in export_sections if is_revision_section(s)), None)
        body_sections = [s for s in export_sections if not is_cover_section(s) and not is_revision_section(s)]

        write_center_section_title("网络安全风险管理报告", font_size=22.0, bold=False)
        add_blank_lines(10)
        for table_rows in (cover_section or {}).get("tables", []) or []:
            add_plain_table(table_rows)

        document.add_page_break()
        write_center_section_title("文件修订记录", font_size=14.0, bold=True)
        add_blank_lines(2)
        for table_rows in (revision_section or {}).get("tables", []) or []:
            add_plain_table(table_rows)

        document.add_page_break()
        write_center_section_title("目录", font_size=16.0, bold=True)
        insert_toc_field()

        document.add_page_break()
        for sec in body_sections:
            add_section(sec)
        document.save(output)
        output.seek(0)
