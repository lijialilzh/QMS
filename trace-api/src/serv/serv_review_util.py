#!/usr/bin/env python
# encoding: utf-8

# 「评审记录」章节共享工具：
#   - 各文档模块在正文末尾追加一个「评审记录」章节（内容模板化）。
#   - 评审时间从产品时间线按「文档名关键字 + 评审」自动获取，格式 yyyy.MM.dd。
#   - 提供导出 Word 时评审内容表/参评人员表的合并渲染（类别列纵向合并、整行横向合并）。
# 说明：内容取自各文档对应的《XXX 附：评审记录》模板；勾选统一用「■通过 □存在问题」。
#      选中标记用实心方块 ■(U+25A0)，与空心 □(U+25A1) 同族且非 emoji，避免 Word 渲染成彩色 emoji。

import re
import base64
import contextvars
from io import BytesIO

from sqlalchemy import select
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE

from ..model.project_timeline import ProjectTimelineRow, ProjectTimelineCell
from ..model.project_member import ProjectMember
from ..model.person_sign import PersonSign
from ..model.prod_dhf import ProdDhf
from ..utils.sql_ctx import db

# 签名开关：导出时若设为 False，封面和评审记录的签名图将被清空（仅保留空单元格）。
# 默认 True（带签名）。通过 set_export_sign_mode(False) / restore_export_sign_mode() 控制。
_sign_mode_var: contextvars.ContextVar[bool] = contextvars.ContextVar("_sign_mode_var", default=True)


def set_export_sign_mode(with_sign: bool):
    """设置当前导出请求的签名模式：True=带签名，False=不带签名（清空签名）。返回 token 用于恢复。"""
    return _sign_mode_var.set(with_sign)


def restore_export_sign_mode(token):
    """恢复签名模式到之前的状态。"""
    _sign_mode_var.reset(token)


def sign_mode_enabled() -> bool:
    """当前是否启用签名填充。"""
    return _sign_mode_var.get()


# 选中用 ☑(框内对号)，加文本呈现选择符(U+FE0E)避免渲染成彩色 emoji；未选空框 □。
CHECK = "\u2611\ufe0e通过 □存在问题"

# 各模块的评审记录模板：items 为 [类别, 评审项] 列表，persons 为参评人员 6 列行
REVIEW_DEFS = {
    "pha": {
        "name_keywords": ["初步危害分析", "危害分析"],
        "items": [
            ["初步危害分析清单", "是否按照ISO14971、GB/T 42062的附录识别了产品的安全特性"],
            ["初步危害分析清单", "已知或可预见的危险（源）是否识别？"],
            ["初步危害分析清单", "软件功能初步危害是否识别？"],
            ["初步危害分析清单", "模型相关初步危害是否识别？"],
            ["初步危害分析清单", "数据标注初步危害是否识别？"],
            ["初步危害分析清单", "网络安全初步危害是否识别？"],
            ["初步危害分析清单", "初步危害分析表是否可追溯？"],
        ],
        "conclusion": (
            "评审结论：\n通过，已经按照ISO14971、GB/T 42062的附录识别了产品的安全特性，"
            "已知或可预见的危险（源）已识别，网络安全初步危害已识别，初步危害分析表可追溯，"
            "产品的初步危害分析活动已完成。"
        ),
        "persons": [
            ["研发总监", "沈宏", "", "产品部经理", "夏晨", ""],
            ["开发负责人", "宁随军", "", "测试负责人", "王小敏", ""],
            ["RA", "张淑芳", "", "QA", "林金贵", ""],
            ["临床人员", "齐济", "", "产品经理", "杨静", ""],
            ["其他参评人员", "", "", "", "", ""],
            ["批准人员签字/日期", "", "", "", "", ""],
        ],
    },
    "pdp": {
        "name_keywords": ["产品开发计划", "开发计划"],
        "items": [
            ["产品简介", "是否包含产品介绍"],
            ["产品简介", "是否包含明确的产品描述"],
            ["产品简介", "是否明确产品开发周期"],
            ["产品简介", "是否明确产品经理"],
            ["资源", "是否明确每个部门的人员资源"],
            ["资源", "是否明确设计开发过程中的设备资源"],
            ["资源", "是否人员资源合理"],
            ["资源", "是否设备资源齐全"],
            ["计划及里程碑", "是否明确每个阶段任务划分和主要活动"],
            ["计划及里程碑", "是否明确每个阶段的责任部门"],
            ["计划及里程碑", "是否明确每个阶段的评审部门"],
            ["计划及里程碑", "是否明确每个阶段的完成时间"],
            ["计划及里程碑", "是否明确每个阶段的交付物"],
            ["计划及里程碑", "是否包含相关所有计划"],
            ["计划及里程碑", "是否各个阶段时间划分合理"],
        ],
        "conclusion": (
            "评审结论：\n通过，产品开发计划中包含了产品介绍、产品描述、产品开发周期，"
            "明确了每个部门的人员资源，明确了设计开发过程中的设备资源，人员资源安排合理，"
            "设备资源齐全，计划中明确了每个阶段任务划分和主要活动，明确了每个阶段的责任部门、"
            "评审部门、完成时间以及交付物，每个阶段时间划分合理。"
        ),
        "persons": [
            ["研发总监", "沈宏", "", "产品部经理", "夏晨", ""],
            ["产品开发部经理", "沈宏", "", "开发负责人", "宁随军", ""],
            ["产品经理", "杨静", "", "测试负责人", "王小敏", ""],
            ["QA", "林金贵", "", "临床人员", "齐济", ""],
            ["RA", "张淑芳", "", "", "", ""],
            ["其他参会人员", "", "", "", "", ""],
            ["批准人员签字/日期", "", "", "", "", ""],
        ],
    },
    "sd": {
        "name_keywords": ["软件开发计划", "开发计划"],
        "items": [
            ["文档完整程度", "项目定义清晰"],
            ["文档完整程度", "有明确且合理的时间计划"],
            ["文档完整程度", "明确人员要求"],
            ["文档完整程度", "明确开发要求"],
            ["文档完整程度", "明确设备资源要求"],
            ["文档完整程度", "符合设计开发流程"],
        ],
        "conclusion": (
            "评审结论：\n通过，项目定义清晰，明确合理的时间计划、人员要求、设备资源要求，"
            "设计开发流程符合要求。"
        ),
        "persons": [
            ["产品经理", "杨静", "", "产品开发部经理", "沈宏", ""],
            ["开发负责人", "宁随军", "", "QA", "林金贵", ""],
            ["其他参评人员", "", "", "", "", ""],
            ["批准人员签字/日期", "", "", "", "", ""],
        ],
    },
    "sds": {
        "name_keywords": ["软件详细设计", "详细设计"],
        "items": [
            ["法规标准引用", "是否明确"],
            ["法规标准引用", "是否合理"],
            ["法规标准引用", "是否完整"],
            ["法规标准引用", "是否符合法规"],
            ["文档完整程度", "文档结构清楚、内容详尽"],
            ["文档完整程度", "包含架构设计"],
            ["文档完整程度", "包含模块设计"],
            ["文档完整程度", "包含接口设计"],
            ["文档完整程度", "包含功能详细设计"],
            ["文档完整程度", "包含必要的数据结构"],
            ["文档完整程度", "软件整体输入、输出接口清晰"],
            ["文档完整程度", "是否可追溯"],
            ["功能覆盖程度", "设计中考虑了整体功能需求"],
            ["功能覆盖程度", "性能要求清晰、明确"],
            ["功能覆盖程度", "接口定义清晰、明确"],
            ["功能覆盖程度", "模块设计覆盖所有功能要求"],
            ["功能覆盖程度", "针对每一项功能都有详细设计"],
            ["功能覆盖程度", "功能设计中具备输入、输出项明确"],
            ["功能覆盖程度", "功能设计中具备逻辑或结构图"],
            ["功能覆盖程度", "能实现软件系统结构"],
            ["功能覆盖程度", "设计的内容不与软件系统结构互相矛盾"],
        ],
        "conclusion": (
            "评审结论：\n通过，详细设计包含架构设计、包含模块设计、包含接口设计、包含必要的数据结构，"
            "输入、输出接口清晰，模块设计覆盖了所有功能要求，针对需求完成了可追溯。"
        ),
        "persons": [
            ["产品经理", "杨静", "", "产品开发部经理", "沈宏", ""],
            ["开发负责人", "宁随军", "", "QA", "林金贵", ""],
            ["其他参评人员", "", "", "", "", ""],
            ["批准人员签字/日期", "", "", "", "", ""],
        ],
    },
    "srs": {
        "name_keywords": ["需求规格说明", "需求规格"],
        "items": [
            ["风险、法规标准引用", "是否明确"],
            ["风险、法规标准引用", "是否合理"],
            ["风险、法规标准引用", "是否完整"],
            ["风险、法规标准引用", "是否符合法规"],
            ["风险、法规标准引用", "是否包含风险控制措施"],
            ["需求撰写", "需求设计前是否具备有效的调研"],
            ["需求撰写", "是否包含完整需求编号"],
            ["需求撰写", "是否包含需求背景及意义"],
            ["需求撰写", "需求撰写是否细化"],
            ["需求撰写", "是否具有一致性"],
            ["需求撰写", "任务是否可被具体拆分"],
            ["需求撰写", "是否考虑异常情况"],
            ["需求撰写", "语义是否清晰明确"],
            ["需求撰写", "是否考虑边界情况"],
            ["需求撰写", "是否满足《需求设计和操作规范》"],
        ],
        "conclusion": (
            "评审结论：\n通过。风险、法规标准引用明确合理，需求撰写内容完整清晰，满足规范要求."
        ),
        "persons": [
            ["研发总监", "沈宏", "", "产品部经理", "夏晨", ""],
            ["开发负责人", "宁随军", "", "测试负责人", "孙家旭", ""],
            ["QA", "林金贵", "", "RA", "杨冰", ""],
            ["临床人员", "齐济", "", "产品经理", "杨静", ""],
            ["其他参会人员", "", "", "", "", ""],
            ["批准人员签字/日期", "", "", "", "", ""],
        ],
    },
    "scm": {
        "name_keywords": ["软件配置管理计划", "配置管理计划"],
        "items": [
            ["软件配置管理计划", "配置管理职责、资源、过程是否明确？"],
            ["软件配置管理计划", "配置项标识与清单是否完整？"],
            ["软件配置管理计划", "版本控制与变更控制是否定义？"],
            ["软件配置管理计划", "发布过程与配置审核是否明确？"],
        ],
        "conclusion": (
            "评审结论：\n通过，软件配置管理计划中配置管理职责、资源、过程明确，配置项标识与清单完整，"
            "版本控制与变更控制已定义，发布过程与配置审核明确。"
        ),
        "persons": [
            ["产品开发部经理", "沈宏", "", "产品经理", "杨静", ""],
            ["开发负责人", "宁随军", "", "QA", "林金贵", ""],
            ["其他参评人员", "", "", "", "", ""],
            ["批准人员签字/日期", "", "", "", "", ""],
        ],
    },
    "scs": {
        "name_keywords": ["软件配置状态报告", "配置管理状态报告", "配置状态报告"],
        "items": [
            ["软件配置状态报告", "产品定义是否明确？"],
            ["软件配置状态报告", "产品范围是否明确？"],
            ["软件配置状态报告", "配置项版本信息是否清晰？"],
            ["软件配置状态报告", "配置项可追溯信息是否完整？"],
        ],
        "conclusion": (
            "评审结论：\n通过，软件配置状态报告明确了定义和范围。配置项具备清晰的版本信息和完整的可追溯信息。"
        ),
        "persons": [
            ["产品经理", "杨静", "", "产品开发部经理", "沈宏", ""],
            ["开发负责人", "宁随军", "", "QA", "林金贵", ""],
            ["其他参评人员", "", "", "", "", ""],
            ["批准人员签字/日期", "", "", "", "", ""],
        ],
    },
    "stp": {
        "name_keywords": ["软件测试计划", "软件测试"],
        "items": [
            ["文档完整程度", "测试目的明确"],
            ["文档完整程度", "报告描述/范围严谨准确"],
            ["文档完整程度", "确认测试设备信息"],
            ["文档完整程度", "确认人力资源"],
            ["文档完整程度", "确认测试工作量"],
            ["文档完整程度", "每一项功能都具备Case编号"],
            ["文档完整程度", "测试项输出清单及测试结果清晰明确"],
            ["文档完整程度", "每项测试都具备测试用例分析"],
            ["文档完整程度", "缺陷统计及遗留bug分析"],
            ["文档完整程度", "测试结论合理性"],
            ["文档完整程度", "是否可追溯"],
        ],
        "conclusion": (
            "评审结论：\n通过，测试目的明确，报告描述/范围严谨准确，均已确认人力资源及测试工作量，"
            "测试项输出清单及测试结果清晰明确。"
        ),
        "persons": [
            ["产品经理", "吴福乐", "", "产品开发部经理", "沈宏", ""],
            ["开发负责人", "宁随军", "", "测试负责人", "宋月", ""],
            ["QA", "林金贵", "", "", "", ""],
            ["其他参评人员", "", "", "", "", ""],
            ["批准人员签字/日期", "", "", "", "", ""],
        ],
    },
    "utp": {
        "name_keywords": ["用户测试计划", "用户测试"],
        "items": [
            ["文档完整程度", "测试定义清晰"],
            ["文档完整程度", "有明确且合理的时间计划"],
            ["文档完整程度", "明确人员要求"],
            ["文档完整程度", "明确软硬件设备要求"],
            ["文档完整程度", "符合测试流程"],
            ["文档完整程度", "明确测试项通过准则"],
            ["文档完整程度", "确定测试的管理工具及测试工具"],
            ["文档完整程度", "明确测试步骤"],
            ["文档完整程度", "明确测试方法"],
        ],
        "conclusion": (
            "评审结论：\n通过，测试定义清晰，明确且合理的时间计划、人员要求，测试流程符合法规，"
            "具有清晰的测试项目输入/输出清单。"
        ),
        "persons": [
            ["产品经理", "吴福乐", "", "产品部经理", "夏晨", ""],
            ["QA", "林金贵", "", "", "", ""],
            ["其他参评人员", "", "", "", "", ""],
            ["批准人员签字/日期", "", "", "", "", ""],
        ],
    },
    "utr": {
        "name_keywords": ["用户测试报告", "用户测试"],
        "items": [
            ["文档完整程度", "测试目的明确"],
            ["文档完整程度", "报告描述/范围严谨准确"],
            ["文档完整程度", "确认测试设备信息"],
            ["文档完整程度", "确认人力资源"],
            ["文档完整程度", "确认测试工作量"],
            ["文档完整程度", "每一项功能都具备Case编号"],
            ["文档完整程度", "测试项输出清单及测试结果清晰明确"],
            ["文档完整程度", "每项测试都具备测试用例分析"],
            ["文档完整程度", "测试结论合理性"],
        ],
        "conclusion": (
            "评审结论：\n通过，测试目的明确，报告描述/范围严谨准确，均已确认人力资源及测试工作量，"
            "测试结论与建议均合理。"
        ),
        "persons": [
            ["产品经理", "吴福乐", "", "产品部经理", "夏晨", ""],
            ["QA", "林金贵", "", "", "", ""],
            ["其他参评人员", "", "", "", "", ""],
            ["批准人员签字/日期", "", "", "", "", ""],
        ],
    },
    "str": {
        "name_keywords": ["软件测试报告"],
        "items": [
            ["文档完整程度", "测试定义清晰"],
            ["文档完整程度", "有明确且合理的时间计划"],
            ["文档完整程度", "明确人员要求"],
            ["文档完整程度", "明确软硬件设备要求"],
            ["文档完整程度", "符合测试流程"],
            ["文档完整程度", "明确测试项通过准则"],
            ["文档完整程度", "确定测试的管理工具及测试工具"],
            ["文档完整程度", "明确测试步骤"],
            ["文档完整程度", "明确测试方法"],
        ],
        "conclusion": (
            "评审结论：\n通过，测试定义清晰，明确且合理的时间计划，测试流程符合法规，"
            "具有清晰的测试项目输入/输出清单。"
        ),
        "persons": [
            ["产品经理", "吴福乐", "", "产品开发部经理", "沈宏", ""],
            ["开发负责人", "宁随军", "", "测试负责人", "宋月", ""],
            ["QA", "林金贵", "", "", "", ""],
            ["其他参评人员", "", "", "", "", ""],
            ["批准人员签字/日期", "", "", "", "", ""],
        ],
    },
    "risk": {
        "name_keywords": ["风险管理报告", "风险管理"],
        "items": [
            ["风险管理报告", "产品定义是否明确？"],
            ["风险管理报告", "产品范围是否明确？"],
            ["风险管理报告", "风险管理固定的内容是否完成？"],
            ["风险管理报告", "风险可接受准则是否定义？"],
            ["风险管理报告", "与安全有关特征的问题是否识别？"],
            ["风险管理报告", "危险（源）是否识别？"],
            ["风险管理报告", "风险控制措施是否合理？"],
            ["风险管理报告", "剩余风险是否可以接受？"],
            ["风险管理报告", "风险控制措施是否引入新的风险？"],
            ["风险管理报告", "风险控制措施的引入是否影响以前识别的危险情况所估计的风险"],
            ["风险管理报告", "综合剩余风险是否合理？"],
            ["风险管理报告", "生产和生产后的风险管理活动是否定义？"],
        ],
        "conclusion": (
            "评审结论：\n通过，风险管理报告中产品定义和产品范围明确，风险可接受准则已定义，"
            "与安全有关特性的问题已识别，危险源已识别，风险控制措施合理，剩余风险可接受，"
            "综合剩余风险可接受，生产和生产后的风险管理活动已定义。"
        ),
        "persons": [
            ["研发总监", "沈宏", "", "产品部经理", "夏晨", ""],
            ["开发负责人", "宁随军", "", "测试负责人", "王小敏", ""],
            ["RA", "张淑芳", "", "QA", "林金贵", ""],
            ["临床人员", "齐济", "", "产品经理", "杨静", ""],
            ["其他参评人员", "/", "/", "/", "/", "/"],
            ["批准人员签字/日期", "", "", "", "", ""],
        ],
    },
}


# 「封面/评审」签署人按部门规则解析（编制人、审核人、批准人）。
#   ("member_role", 角色关键字)：从本产品参与人员中按角色关键字找到姓名，再取其签名；
#   ("name", 姓名)：固定姓名（公司层面固定签署人）。
#   规则来源（用户约定）：
#     - 产品文件：编制人=产品经理；审核/批准=产品总监(夏晨)
#     - 开发文件：编制人=TPM；审核/批准=研发负责人
#     - 测试文件：编制人=测试人员；审核/批准=研发负责人
DEPT_SIGNERS = {
    "product": {
        "编制人": ("member_role", "产品经理"),
        "审核人": ("name", "夏晨"),
        "批准人": ("name", "夏晨"),
    },
    "dev": {
        "编制人": ("member_role", "TPM"),
        "审核人": ("member_role", "研发负责人"),
        "批准人": ("member_role", "研发负责人"),
    },
    "test": {
        "编制人": ("member_role", "测试人员"),
        "审核人": ("member_role", "研发负责人"),
        "批准人": ("member_role", "研发负责人"),
    },
    "model": {
        "编制人": ("member_role", ["模型部负责人", "模型负责人", "模型"]),
        "审核人": ("member_role", ["算法", "研发负责人"]),
        "批准人": ("member_role", "研发负责人"),
    },
    "data": {
        "编制人": ("member_role", "数据"),
        "审核人": ("member_role", "模型"),
        "批准人": ("member_role", "研发负责人"),
    },
}
DEFAULT_DEPT = "product"

# 各文档模块所属部门（决定封面/评审签署人规则）。如需调整只改此表即可。
DOC_DEPT = {
    # 产品线文件：编制人=产品经理，审核/批准=产品总监(夏晨)
    "pdp": "product", "pir": "product", "label": "product",
    "release_note": "product", "vuh": "product",
    "risk": "product", "rmp": "product", "pha": "product",
    # 开发文件：编制人=TPM，审核/批准=研发负责人
    "sd": "dev", "srs": "dev", "sds": "dev", "cybersec": "dev",
    "nsmp": "dev", "nsr": "dev", "research": "dev", "crr": "dev",
    "scm": "dev", "scs": "dev",
    # 测试文件：编制人=测试人员，审核/批准=研发负责人
    "stp": "test", "utp": "test", "utr": "test", "str": "test", "bug": "test", "imm": "test",
    # 模型文件：编制人=模型部负责人（无则模型负责人/角色含模型），审核人=算法（无则研发负责人），批准人=研发负责人（无匹配则空）
    "md_001": "model",
    "md_004": "model", "md_005": "model", "md_006": "model", "md_007": "model",
    "md_014": "model", "md_017": "model", "md_019": "model", "md_020": "model",
    "md_021": "model", "md_022": "model", "pd_003": "model",
    "md_008_01": "model", "md_008_02": "model", "md_009_01": "model", "md_009_02": "model",
    "md_010_01": "model", "md_010_02": "model", "md_011_01": "model", "md_011_02": "model",
    "md_012_01": "model", "md_012_02": "model", "md_013_01": "model", "md_013_02": "model",
    "md_015_01": "model", "md_015_02": "model", "md_016": "model", "md_018": "model",
    "md_019_qr": "model", "md_020_qr": "model", "md_deq": "model", "md_teq": "model", "md_eq": "model",
    "dd_001": "data", "md_002_01": "data", "md_002_02": "data", "md_003": "data",
    "dd_006": "data", "dd_007": "data", "dd_016": "data", "dd_017": "data",
    "dd_002": "data", "dd_003": "data", "dd_004": "data", "dd_005_01": "data", "dd_005_02": "data",
    "dd_008_01": "data", "dd_008_02": "data", "dd_009_01": "data", "dd_009_02": "data", "dd_009_03": "data",
    "dd_010": "data", "dd_011": "data", "dd_012": "data",
    "dd_013_01": "data", "dd_013_02": "data", "dd_013_03": "data", "dd_013_04": "data",
    "dd_013_05": "data", "dd_013_06": "data", "dd_013_07": "data",
    "dd_014": "data", "dd_015_01": "data", "dd_015_02": "data", "dd_015_03": "data",
    "dd_016_qr": "data", "dd_017_qr": "data", "dd_eq": "data",
}


def _before_202509(rev_date):
    """判断评审/文档日期是否在 2025 年 9 月之前（用于测试人员签名统一取宋月的规则）。"""
    m = re.match(r"(\d{4})\D+(\d{1,2})", str(rev_date or ""))
    if not m:
        return False
    return int(m.group(1)) * 100 + int(m.group(2)) < 202509


def _parse_timeline_year_value(value):
    """从时间线「年」列或里程碑文本解析四位年份。"""
    s = str(value or "").strip()
    if not s:
        return None
    m = re.search(r"(20\d{2})", s)
    if m:
        return int(m.group(1))
    digits = re.sub(r"[^\d]", "", s)
    if len(digits) == 4 and digits.startswith("20"):
        return int(digits)
    return None


def _timeline_date_rows_with_year(tl_rows):
    """按 sort_order 为日期行补全年份（向下填充；模板缺年列时用年份行/兜底当前年）。"""
    import datetime

    sorted_rows = sorted(tl_rows or [], key=lambda r: ((r.sort_order or 0), (r.id or 0)))
    default_year = datetime.date.today().year
    for r in sorted_rows:
        y = _parse_timeline_year_value(getattr(r, "year", None)) or _parse_timeline_year_value(getattr(r, "milestone_text", None))
        if y:
            default_year = y
    last_year = None
    dated = []
    for r in sorted_rows:
        rtype = (getattr(r, "row_type", None) or "date")
        if rtype == "year":
            y = _parse_timeline_year_value(getattr(r, "milestone_text", None) or getattr(r, "year", None))
            if y:
                last_year = y
            continue
        if rtype != "date":
            continue
        explicit = _parse_timeline_year_value(getattr(r, "year", None))
        if explicit:
            last_year = explicit
        year = explicit or last_year or default_year
        dated.append((r, year))
    return dated


def review_date(prod_id, name_keywords):
    """从时间线取评审日期（命中「文档名关键字」且优先命中「评审」的行，取最新），格式 yyyy.MM.dd。"""
    if not prod_id:
        return ""

    def to_int(v):
        digits = re.sub(r"[^\d]", "", str(v or ""))
        return int(digits) if digits else None

    tl_rows = db.session.execute(
        select(ProjectTimelineRow).where(ProjectTimelineRow.prod_id == prod_id)
    ).scalars().all()
    if not tl_rows:
        return ""
    cell_map = {}
    for c in db.session.execute(
        select(ProjectTimelineCell).where(ProjectTimelineCell.row_id.in_([r.id for r in tl_rows]))
    ).scalars().all():
        cell_map.setdefault(c.row_id, []).append(c.output_result or "")
    date_rows = [(r, y) for r, y in _timeline_date_rows_with_year(tl_rows) if to_int(r.month)]

    def date_key(item):
        r, year = item
        return year * 10000 + to_int(r.month) * 100 + (to_int(r.day) or 0)

    def match(r, need_review):
        vals = cell_map.get(r.id, [])
        hit_name = any(any(k in str(v or "") for k in name_keywords) for v in vals)
        hit_review = any("评审" in str(v or "") for v in vals)
        return hit_name and (hit_review if need_review else True)

    rows = [item for item in date_rows if match(item[0], True)] or [item for item in date_rows if match(item[0], False)]
    if not rows:
        return ""
    r, year = max(rows, key=date_key)
    return f"{year}.{to_int(r.month):02d}.{(to_int(r.day) or 1):02d}"


def date_range(prod_id, name_keywords):
    """时间线里命中任一关键字的活动，返回其（最早, 最晚）日期区间，形如 (yyyy.MM.dd, yyyy.MM.dd)。
    关键字需足够具体以排除其它轨道（如用『单元测试记录』而非『测试记录』以排除『模型测试记录』）。"""
    if not prod_id or not name_keywords:
        return ("", "")

    def to_int(v):
        digits = re.sub(r"[^\d]", "", str(v or ""))
        return int(digits) if digits else None

    tl_rows = db.session.execute(
        select(ProjectTimelineRow).where(ProjectTimelineRow.prod_id == prod_id)
    ).scalars().all()
    if not tl_rows:
        return ("", "")
    cell_map = {}
    for c in db.session.execute(
        select(ProjectTimelineCell).where(ProjectTimelineCell.row_id.in_([r.id for r in tl_rows]))
    ).scalars().all():
        cell_map.setdefault(c.row_id, []).append(c.output_result or "")

    def date_key(item):
        r, year = item
        return year * 10000 + to_int(r.month) * 100 + (to_int(r.day) or 0)

    matched = []
    for r, year in _timeline_date_rows_with_year(tl_rows):
        if not to_int(r.month):
            continue
        vals = cell_map.get(r.id, [])
        if any(any(k in str(v or "") for k in name_keywords) for v in vals):
            matched.append((r, year))
    if not matched:
        return ("", "")
    lo = min(matched, key=date_key)
    hi = max(matched, key=date_key)

    def fmt(item):
        r, year = item
        return f"{year}.{to_int(r.month):02d}.{(to_int(r.day) or 1):02d}"

    return (fmt(lo), fmt(hi))


def date_count(prod_id, name_keywords):
    """时间线里命中任一关键字的『不同日期行』数量（= 实际工作日天数，时间线只在工作日建行，已排除周末/节假日）。"""
    if not prod_id or not name_keywords:
        return 0

    def to_int(v):
        digits = re.sub(r"[^\d]", "", str(v or ""))
        return int(digits) if digits else None

    tl_rows = db.session.execute(
        select(ProjectTimelineRow).where(ProjectTimelineRow.prod_id == prod_id)
    ).scalars().all()
    if not tl_rows:
        return 0
    cell_map = {}
    for c in db.session.execute(
        select(ProjectTimelineCell).where(ProjectTimelineCell.row_id.in_([r.id for r in tl_rows]))
    ).scalars().all():
        cell_map.setdefault(c.row_id, []).append(c.output_result or "")
    days = set()
    for r in tl_rows:
        if (r.row_type or "date") != "date" or not to_int(r.year) or not to_int(r.month):
            continue
        vals = cell_map.get(r.id, [])
        if any(any(k in str(v or "") for k in name_keywords) for v in vals):
            days.add((to_int(r.year), to_int(r.month), to_int(r.day) or 0))
    return len(days)


# 各文档「封面日期」取值用的时间线关键字（供 review_date 使用）。
# 规则：编制/审核/批准日期 + 生效日期 统一取该文档在时间线里评审/最后一天的日期。
COVER_KEYWORDS = {
    "sd": ["软件开发计划", "开发计划"],
    "crr": ["代码审查", "代码评审"],
    "scm": ["软件配置管理计划", "配置管理计划"],
    "scs": ["软件配置状态报告", "配置管理状态报告", "配置状态报告"],
    "pdp": ["产品开发计划", "开发计划"],
    "cybersec": ["网络安全风险管理报告", "网络安全风险管理", "网络安全风险管理计划", "风险管理报告"],
    "risk": ["风险管理报告", "风险管理"],
    "rmp": ["风险管理计划"],
    "pha": ["初步危害分析", "危害分析"],
    "srs": ["需求规格说明", "需求规格"],
    "sds": ["软件详细设计", "详细设计"],
    "hld": ["软件概要设计", "概要设计"],
    "pir": ["产品立项报告", "立项报告"],
    "ptr": ["产品技术要求", "技术要求"],
    "acc": ["产品验收记录", "验收记录"],
    "label": ["产品标签样稿", "标签"],
    "vuh": ["版本更新历史"],
    "nsmp": ["网络安全维护计划", "维护计划"],
    "release_note": ["产品发布说明", "发布说明"],
    "nsr": ["自研软件网络安全研究报告", "网络安全研究报告"],
    "research": ["自研软件研究报告", "软件研究报告"],
    "stp": ["软件测试计划", "软件测试"],
    "utp": ["用户测试计划", "用户测试"],
    "utr": ["用户测试报告", "用户测试"],
    "str": ["软件测试报告"],
    "bug": ["Bug管理及回归测试", "回归测试", "Bug管理", "缺陷"],
    "imm": ["安装维护手册", "安装维护", "MD5"],
    "ftr": ["现场测试规程", "现场测试"],
    "cyber_cap": ["网络安全能力分析", "能力分析"],
    "train_record": ["培训记录表", "培训记录"],
    "md_001": ["模型配置管理计划"],
    "md_002_01": ["数据标注规则", "肺栓塞分割数据标注规则"],
    "md_002_02": ["标记规则", "肺叶分割标记规则"],
    "md_003": ["数据标注需求"],
    "md_004": ["算法方案概要设计"],
    "md_005": ["模型测试方案设计"],
    "md_006": ["模型开发计划"],
    "md_007": ["算法方案详细设计"],
    "md_014": ["模型测试报告"],
    "md_017": ["模型性能测试报告"],
    "md_019": ["开发环境维护记录说明", "模型开发环境"],
    "md_020": ["测试环境维护记录说明", "模型测试环境"],
    "md_021": ["模型配置管理报告"],
    "md_022": ["模型可追溯性分析报告", "可追溯性分析"],
    "pd_003": ["模型需求规格说明"],
    "md_008_01": ["代码审查记录", "肺栓塞分割代码审查"],
    "md_008_02": ["代码审查记录", "肺叶分割代码审查"],
    "md_009_01": ["训练集构建记录", "肺栓塞分割模型训练集"],
    "md_009_02": ["训练集构建记录", "肺叶分割模型训练集"],
    "md_010_01": ["调优集构建记录", "肺栓塞分割模型调优集"],
    "md_010_02": ["调优集构建记录", "肺叶分割模型调优集"],
    "md_011_01": ["测试集构建记录", "肺栓塞分诊模型测试集"],
    "md_011_02": ["测试集构建记录", "肺叶分割模型测试集"],
    "md_012_01": ["模型训练记录", "肺栓塞分割模型训练"],
    "md_012_02": ["模型训练记录", "肺叶分割模型训练"],
    "md_013_01": ["模型测试记录", "肺栓塞分诊模型测试记录"],
    "md_013_02": ["模型测试记录", "肺叶分割模型测试记录"],
    "md_015_01": ["封装需求", "肺栓塞分割封装"],
    "md_015_02": ["封装需求", "肺叶分割封装"],
    "md_016": ["模型工程封装记录"],
    "md_018": ["模型服务提交记录"],
    "md_019_qr": ["开发环境维护记录"],
    "md_020_qr": ["测试环境维护记录"],
    "md_deq": ["开发设备清单"],
    "md_teq": ["测试设备清单"],
    "md_eq": ["设备清单"],
    "dd_001": ["数据采集需求"],
    "dd_006": ["数据标注质量评估方法"],
    "dd_007": ["人员考核评价方法"],
    "dd_016": ["开发环境维护记录说明", "数据开发环境"],
    "dd_017": ["标注环境维护记录说明"],
    "dd_002": ["多中心数据回传记录", "数据回传"],
    "dd_003": ["数据整理记录"],
    "dd_004": ["数据采集需求反馈"],
    "dd_005_01": ["人员培训记录", "肺栓塞分割人员培训"],
    "dd_005_02": ["人员培训记录", "肺叶分割人员培训"],
    "dd_008_01": ["试标注记录", "肺栓塞分割试标注"],
    "dd_008_02": ["试标注记录", "肺叶分割试标注"],
    "dd_009_01": ["数据标注记录", "肺栓塞分割标注记录"],
    "dd_009_02": ["数据标注记录", "肺叶分割标注记录"],
    "dd_009_03": ["数据标注记录", "肺栓塞分诊标注记录"],
    "dd_010": ["数据库上传记录"],
    "dd_011": ["数据标注需求反馈"],
    "dd_012": ["查重记录", "训练集测试集查重"],
    "dd_013_01": ["初次考核记录"],
    "dd_013_02": ["初次考核记录"],
    "dd_013_03": ["定期考核记录"],
    "dd_013_04": ["定期考核记录"],
    "dd_013_05": ["日常考核", "肺栓塞分割日常考核"],
    "dd_013_06": ["日常考核", "肺叶分割日常考核"],
    "dd_013_07": ["日常考核", "肺栓塞分诊日常考核"],
    "dd_014": ["数据库维护记录"],
    "dd_015_01": ["原始数据库统计表"],
    "dd_015_02": ["基础数据库统计表"],
    "dd_015_03": ["标注数据库统计表"],
    "dd_016_qr": ["开发环境维护记录"],
    "dd_017_qr": ["标注环境维护记录"],
    "dd_eq": ["设备清单"],
}


# DHF 名称模糊匹配时需排除的片段（避免「需求规格」误命中「需求规格附件」等）
DHF_NAME_EXCLUDES = {
    "srs": ("附件", "附录"),
    "sds": ("附件", "附录"),
    "hld": ("附件", "附录"),
    "md_019_qr": ("说明",),
    "md_020_qr": ("说明",),
    "dd_016_qr": ("说明",),
    "dd_017_qr": ("说明",),
}


def _dhf_name_excluded(name, exclude_fragments=()):
    txt = str(name or "").strip()
    return any(frag and frag in txt for frag in (exclude_fragments or ()))


def _dhf_match_score(name, keyword):
    txt = str(name or "").strip()
    kw = str(keyword or "").strip()
    if txt == kw:
        return 0
    if txt.endswith("说明") and kw in txt:
        return 1
    if "说明" in txt and kw in txt:
        return 2
    return 10


def _normalize_dhf_code(code):
    """编号字段偶发混入括号说明，只保留编码本体。"""
    txt = str(code or "").strip()
    if not txt:
        return ""
    for sep in ("(", "（"):
        if sep in txt:
            txt = txt.split(sep, 1)[0].strip()
    return txt


def dhf_file_no(prod_id, name_keywords=None, exclude_fragments=()):
    """从产品 DHF 清单按文件名称关键字匹配文件编号 code。"""
    if not prod_id:
        return ""
    keywords = sorted(
        [str(k or "").strip() for k in (name_keywords or []) if str(k or "").strip()],
        key=len,
        reverse=True,
    )
    if not keywords:
        return ""
    excludes = tuple(exclude_fragments or ())
    for kw in keywords:
        row = db.session.execute(
            select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name == kw).order_by(ProdDhf.id.asc())
        ).scalars().first()
        if row and (row.code or "").strip() and not _dhf_name_excluded(row.name, excludes):
            return _normalize_dhf_code(row.code)

        rows = db.session.execute(
            select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name.like(f"%{kw}%")).order_by(ProdDhf.id.asc())
        ).scalars().all()
        candidates = []
        for item in rows:
            if not (item.code or "").strip() or _dhf_name_excluded(item.name, excludes):
                continue
            candidates.append((_dhf_match_score(item.name, kw), item))
        if candidates:
            candidates.sort(key=lambda pair: (pair[0], pair[1].id or 0))
            return _normalize_dhf_code(candidates[0][1].code)
    return ""


def resolve_doc_file_no(prod_id, file_no="", version="", doc_key=""):
    """未填写 file_no 时从产品 DHF 匹配；并按 version 同步末尾版本段。"""
    from .serv_utils import sync_file_no_version

    base = _normalize_dhf_code((file_no or "").strip())
    if not base and prod_id:
        base = dhf_file_no(
            prod_id,
            COVER_KEYWORDS.get(doc_key, []),
            DHF_NAME_EXCLUDES.get(doc_key, ()),
        )
    if not base:
        return None
    if version:
        synced = sync_file_no_version(base, version)
        return _normalize_dhf_code(synced or base)
    return base


def fill_cover_dates(content, rev_date, force=False):
    """统一填充封面表日期：
      - 「编制人/审核人/批准人」行的「日期」列(第4列, index 3) 填 rev_date；
      - 「生效日期」行的值(第2列, index 1) 填 rev_date。
    rev_date 为该文档在时间线里的评审/最后一天日期（见 review_date）。
    force=False（默认）：仅填空，不覆盖已填（保留用户手动编辑）。
    force=True（切换产品）：强制覆盖，rev_date 为空则置空。
    通用实现：扫描所有 section 的所有表格，按行首标签匹配，兼容各模块封面结构。"""
    if not isinstance(content, dict):
        return content
    for section in (content.get("sections") or []):
        if not isinstance(section, dict):
            continue
        for table in (section.get("tables") or []):
            if not isinstance(table, list):
                continue
            for row in table:
                if not isinstance(row, list) or not row:
                    continue
                label = str(row[0] or "").strip()
                if label in ("编制人", "审核人", "批准人"):
                    if len(row) >= 4:
                        if force or not str(row[3] or "").strip():
                            row[3] = rev_date or ""
                elif label == "生效日期":
                    if len(row) >= 2:
                        if force or not str(row[1] or "").strip():
                            row[1] = rev_date or ""
    return content


def cover_date(prod_id, key):
    """按模块 key 取封面日期（复用 review_date + COVER_KEYWORDS）。"""
    kws = COVER_KEYWORDS.get(key)
    if not prod_id or not kws:
        return ""
    return review_date(prod_id, kws)


def _sign_by_name(name):
    """按姓名从「人员签名管理(person_sign)」取签名图（sign_img，data URL）。"""
    name = (name or "").strip()
    if not name:
        return ""
    row = db.session.execute(
        select(PersonSign).where(PersonSign.name == name)
    ).scalars().first()
    return (getattr(row, "sign_img", "") or "") if row else ""


# 个别模块如需覆盖部门默认规则，可在此登记 key -> {label:(kind,arg)}；留空则一律按部门规则。
COVER_SIGNERS = {}


def _signer_config(key):
    """返回某模块的签署人配置（优先 COVER_SIGNERS 覆盖，否则按部门 DEPT_SIGNERS）。"""
    if key in COVER_SIGNERS:
        return COVER_SIGNERS[key]
    return DEPT_SIGNERS.get(DOC_DEPT.get(key, DEFAULT_DEPT), {})


def _resolve_signer_name(spec, members, rev_date=""):
    """按签署人规则解析姓名。测试人员在 2025.09 之前统一取宋月。"""
    kind, arg = spec
    if kind == "name":
        return arg
    keywords = arg if isinstance(arg, (list, tuple)) else [arg]
    if any("测试" in str(k) for k in keywords) and _before_202509(rev_date):
        return "宋月"
    for k in keywords:
        hit = next((m.name for m in members if str(k) in str(m.role or "")), "")
        if hit:
            return hit
    return ""


def cover_signers(prod_id, key="pdp", rev_date=""):
    """按模块 key 返回封面「编制人/审核人/批准人」应填的签名图 data URL 字典（无图则回退姓名）。
    rev_date 缺省时自动按模块取评审/最后日期（用于测试人员宋月规则）。"""
    cfg = _signer_config(key)
    signers = {}
    if not prod_id or not cfg:
        return signers
    if not rev_date:
        rev_date = cover_date(prod_id, key)
    members = db.session.execute(
        select(ProjectMember).where(ProjectMember.prod_id == prod_id)
    ).scalars().all()
    for label, spec in cfg.items():
        name = (_resolve_signer_name(spec, members, rev_date) or "").strip()
        if not name:
            continue
        # 有签名图就放图；没有签名图则回退显示姓名文字（保证签署人不为空）
        signers[label] = _sign_by_name(name) or name
    return signers


def cover_signer_names(prod_id, key="pdp", rev_date=""):
    """按模块 key 返回封面「编制人/审核人/批准人」应填的『姓名文本』字典（不取签名图）。
    用于非 JSON 内容（如上传的 xlsx 封面）按部门规则回填姓名。"""
    cfg = _signer_config(key)
    out = {}
    if not prod_id or not cfg:
        return out
    if not rev_date:
        rev_date = cover_date(prod_id, key)
    members = db.session.execute(
        select(ProjectMember).where(ProjectMember.prod_id == prod_id)
    ).scalars().all()
    for label, spec in cfg.items():
        name = (_resolve_signer_name(spec, members, rev_date) or "").strip()
        if name:
            out[label] = name
    return out


def review_approver(key, prod_id=None, rev_date=""):
    """评审记录「批准人」姓名：按部门规则解析（产品=夏晨；开发/测试=研发负责人）。"""
    spec = _signer_config(key).get("批准人")
    if not spec:
        return ""
    if spec[0] == "name":
        return spec[1]
    if not prod_id:
        return ""
    members = db.session.execute(
        select(ProjectMember).where(ProjectMember.prod_id == prod_id)
    ).scalars().all()
    return _resolve_signer_name(spec, members, rev_date) or ""


def fill_cover_signers(content, signers, force=False):
    """把封面「编制人/审核人/批准人」行的姓名列(第2列, index 1) 填成签名图 data URL。
    force=False（默认）：仅填空，不覆盖已填（已是签名图则保留，避免重复覆盖用户已放的图）。
    force=True（切换产品）：强制覆盖，signers 无对应人则置空。
    当全局签名模式关闭（sign_mode_enabled()==False）时：清空所有签署人单元格（置空）。
    通用实现：扫描所有 section 的所有表格，按行首标签匹配，兼容各模块封面结构。"""
    if not isinstance(content, dict):
        return content
    # 不带签名模式：清空封面签署人单元格
    if not sign_mode_enabled():
        for section in (content.get("sections") or []):
            if not isinstance(section, dict):
                continue
            for table in (section.get("tables") or []):
                if not isinstance(table, list):
                    continue
                for row in table:
                    if not isinstance(row, list) or not row:
                        continue
                    label = str(row[0] or "").strip()
                    if label in ("编制人", "审核人", "批准人") and len(row) >= 2:
                        row[1] = ""
        return content
    signers = signers or {}
    for section in (content.get("sections") or []):
        if not isinstance(section, dict):
            continue
        for table in (section.get("tables") or []):
            if not isinstance(table, list):
                continue
            for row in table:
                if not isinstance(row, list) or not row:
                    continue
                label = str(row[0] or "").strip()
                if label not in ("编制人", "审核人", "批准人"):
                    continue
                if len(row) < 2:
                    continue
                val = signers.get(label, "") if force else signers.get(label)
                if force:
                    row[1] = val or ""
                else:
                    # 覆盖空值或旧姓名文本；已是签名图则保留（避免重复覆盖用户已放的图）
                    if val and not str(row[1] or "").startswith("data:image"):
                        row[1] = val
    return content


# 整行横向合并的行首标记（这些行内容跨满整行）
_BANNERS = ("参评人员签字", "评审时间", "评审结论", "批准人员签字")


def _is_banner(text):
    t = str(text or "")
    return any(t.startswith(b) for b in _BANNERS)


def _is_other_row(text):
    """「其他参会人员/其他参评人员」行：标签后单元格合并为一格并填 /。"""
    t = str(text or "").strip()
    return t.startswith("其他参会人员") or t.startswith("其他参评人员")


# 评审记录「人员角色」→ 产品参与人员角色关键字（按优先级匹配当前产品成员）
REVIEW_ROLE_KW = {
    "产品经理": ["产品经理"],
    "产品开发部经理": ["研发负责人", "产品开发部经理"],
    "研发总监": ["研发负责人", "研发总监"],
    "产品部经理": ["产品负责人", "产品部经理", "产品总监"],
    "开发负责人": ["开发负责人", "TPM"],
    "QA": ["QA", "质量"],
    "RA": ["RA", "法规"],
    "临床人员": ["临床"],
    "管理者代表": ["管理者代表"],
}


def _resolve_review_name(role_label, members, rev_date):
    """按评审记录的「人员角色」取当前产品参与人员姓名；匹配不到返回 None（保留模板名）。
    测试角色按日期规则：2025.09 前=宋月，之后=孙家旭。"""
    label = str(role_label or "").strip()
    if not label or _is_banner(label) or label.startswith("其他"):
        return None
    if ("测试" in label) and ("用户" not in label):
        return "宋月" if _before_202509(rev_date) else "孙家旭"
    if "用户测试" in label:
        for m in members:
            if "用户测试" in str(m.role or ""):
                return (m.name or "").strip()
        return None
    for kw in REVIEW_ROLE_KW.get(label, [label]):
        for m in members:
            if kw and kw in str(m.role or ""):
                return (m.name or "").strip()
    return None


def build_review_section(key, rev_date="", prod_id=None):
    d = REVIEW_DEFS.get(key)
    if not d:
        return None
    # 评审内容表：类别仅在每段首行保留，其余留空（合并前的自然形态，编辑页不重复）
    content_tbl = [["评审内容", "评审项", "评审结论"]]
    prev_cat = None
    for cat, q in d["items"]:
        content_tbl.append(["" if cat == prev_cat else cat, q, CHECK])
        prev_cat = cat
    content_tbl.append([d["conclusion"], "", ""])
    # 参评人员表：整行标记行只在首格保留文字
    person_tbl = [
        ["参评人员签字", "", "", "", "", ""],
        [f"评审时间：{rev_date}", "", "", "", "", ""],
        ["人员角色", "姓名", "签字", "人员角色", "姓名", "签字"],
    ] + [list(r) for r in d["persons"]]
    # 「签字」列按「姓名」列自动取签名图（第2列姓名→第3列签字；第5列姓名→第6列签字），仅填空
    old_test = _before_202509(rev_date)
    # 测试文件：评审记录按「人员角色」获取当前产品参与人员姓名（其余模块保持模板名+宋月规则）
    role_based = key in ("stp", "utp", "utr", "str", "sds", "srs")
    rb_members = []
    if role_based and prod_id:
        rb_members = db.session.execute(
            select(ProjectMember).where(ProjectMember.prod_id == prod_id)
        ).scalars().all()
    for row in person_tbl[3:]:
        if not isinstance(row, list) or _is_banner(str(row[0] or "")):
            continue
        if role_based:
            if len(row) >= 2:
                nm = _resolve_review_name(row[0], rb_members, rev_date)
                if nm:
                    row[1] = nm
            if len(row) >= 5:
                nm = _resolve_review_name(row[3] if len(row) >= 4 else "", rb_members, rev_date)
                if nm:
                    row[4] = nm
        elif old_test:
            # 测试角色：2025.09 之前签名人统一为宋月
            if "测试" in str(row[0] or "") and len(row) >= 2:
                row[1] = "宋月"
            if "测试" in str(row[3] or "") and len(row) >= 5:
                row[4] = "宋月"
        if len(row) >= 3 and str(row[1] or "").strip() and not str(row[2] or "").strip():
            row[2] = _sign_by_name(str(row[1]).strip()) if sign_mode_enabled() else ""
        if len(row) >= 6 and str(row[4] or "").strip() and not str(row[5] or "").strip():
            row[5] = _sign_by_name(str(row[4]).strip()) if sign_mode_enabled() else ""
    # 「批准人员签字/日期」行：自动填批准人签名图(第2列，无图回退姓名) 与评审日期(第3列)
    approver = review_approver(key, prod_id, rev_date)
    for row in person_tbl:
        if isinstance(row, list) and row and str(row[0] or "").startswith("批准人员签字"):
            if approver and len(row) >= 2 and not str(row[1] or "").strip():
                row[1] = (_sign_by_name(approver) or approver) if sign_mode_enabled() else ""
            if len(row) >= 3 and not str(row[2] or "").strip():
                row[2] = rev_date
            break
    # 「其他参会人员/其他参评人员」行：标签后单元格合并为一格并填 /
    for row in person_tbl:
        if isinstance(row, list) and row and _is_other_row(row[0]):
            for i in range(1, len(row)):
                row[i] = ""
            if len(row) >= 2:
                row[1] = "/"
            break
    return {
        "title": "评审记录",
        "ref_type": "review",
        "body": "",
        "tables": [content_tbl, person_tbl],
        "children": [],
    }


def autofill_review_person_table(tbl, key="", rev_date="", prod_id=None):
    """对「已存在的参评人员签字表」就地自动填充（用于评审记录内置于默认内容、不经 build_review_section 的模块，如 rmp）：
      - 「签字」列按「姓名」列取签名图（仅填空）；测试角色 2025.09 前统一取宋月；
      - 「批准人员签字/日期」行填批准人签名(按部门) + 评审日期；
      - 「其他参评人员/其他参会人员」行标签后合并为一格并填 /。
    仅处理含「参评人员签字」标记的人员表，避免误伤评审内容表。"""
    if not isinstance(tbl, list):
        return tbl
    if not any(isinstance(r, list) and r and str(r[0] or "").startswith("参评人员签字") for r in tbl):
        return tbl
    passed_date = bool(rev_date)
    # 评审日期缺省时，取表内「评审时间」行已有的日期
    if not rev_date:
        for row in tbl:
            if isinstance(row, list) and row and str(row[0] or "").startswith("评审时间"):
                parts = str(row[0]).split("：", 1)
                if len(parts) > 1:
                    rev_date = parts[1].strip()
                break
    old_test = _before_202509(rev_date)
    approver = review_approver(key, prod_id, rev_date) if key else ""
    for row in tbl:
        if not isinstance(row, list) or not row:
            continue
        lab = str(row[0] or "").strip()
        if lab.startswith("评审时间"):
            if passed_date and rev_date:
                row[0] = f"评审时间：{rev_date}"
            continue
        if lab.startswith("批准人员签字"):
            if approver and len(row) >= 2 and not str(row[1] or "").strip():
                row[1] = (_sign_by_name(approver) or approver) if sign_mode_enabled() else ""
            if len(row) >= 3 and not str(row[2] or "").strip():
                row[2] = rev_date
            continue
        if _is_banner(lab) or lab == "人员角色":
            continue
        if _is_other_row(lab):
            for i in range(1, len(row)):
                row[i] = ""
            if len(row) >= 2:
                row[1] = "/"
            continue
        # 参评人员数据行：测试角色宋月 + 签字列取签名图
        if old_test:
            if "测试" in lab and len(row) >= 2:
                row[1] = "宋月"
            if len(row) >= 4 and "测试" in str(row[3] or ""):
                row[4] = "宋月"
        if len(row) >= 3 and str(row[1] or "").strip() and not str(row[2] or "").strip():
            row[2] = _sign_by_name(str(row[1]).strip()) if sign_mode_enabled() else ""
        if len(row) >= 6 and str(row[4] or "").strip() and not str(row[5] or "").strip():
            row[5] = _sign_by_name(str(row[4]).strip()) if sign_mode_enabled() else ""
    return tbl


def ensure_review(content, key, rev_date="", prod_id=None):
    """在 content.sections 末尾放置「评审记录」章节。内容为模板化，每次按最新格式重建，
    保证样式一致并带最新评审时间（同时清理历史遗留的旧格式）。"""
    sections = (content or {}).get("sections")
    if not isinstance(sections, list):
        return content
    sec = build_review_section(key, rev_date, prod_id)
    if not sec:
        return content
    sections[:] = [s for s in sections if s.get("ref_type") != "review"]
    sections.append(sec)
    return content


def render_review_grid(document, grid, set_cell, header_rows=1, **_ignore):
    """评审表渲染：整行标记行横向合并；首列非空单元格纵向合并其后的空单元格。
    set_cell 由各模块传入以保持字体一致，签名 set_cell(cell, text, bold=False, align=...)。"""
    grid = [row for row in (grid or []) if isinstance(row, list)]
    cols = max((len(row) for row in grid), default=0)
    if cols <= 0:
        return
    table = document.add_table(rows=0, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, row in enumerate(grid):
        cells = table.add_row().cells
        for c_idx in range(cols):
            val = row[c_idx] if c_idx < len(row) else ""
            s = str(val or "")
            if s.startswith("data:image"):
                # 签字列签名图：等比缩放渲染，不变形
                try:
                    b64 = s.split(",", 1)[1] if "," in s else ""
                    pic = base64.b64decode(b64)
                    cell = cells[c_idx]
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.add_run().add_picture(BytesIO(pic), height=Pt(28))
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    continue
                except Exception:
                    pass
            set_cell(cells[c_idx], val, bold=(r_idx < header_rows))
    rows = table.rows
    n = len(rows)
    # 表头行横向合并：相邻列内容相同的合并为一格（如"评审内容"两列相同）
    for r in range(min(header_rows, n)):
        c = 0
        while c < cols:
            cval = str(rows[r].cells[c].text or "").strip()
            if not cval:
                c += 1
                continue
            c2 = c
            while c2 + 1 < cols and str(rows[r].cells[c2 + 1].text or "").strip() == cval:
                c2 += 1
            if c2 > c:
                merged = rows[r].cells[c].merge(rows[r].cells[c2])
                set_cell(merged, cval, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            c = c2 + 1
    # 整行标记行横向合并
    for r in range(n):
        first = rows[r].cells[0].text
        if _is_banner(first) and len(rows[r].cells) > 1:
            merged = rows[r].cells[0].merge(rows[r].cells[-1])
            if first.startswith("批准人员签字"):
                # 批准人签字行：左对齐，渲染「标签 + 批准人签名图/姓名 + 日期」
                gr = grid[r] if r < len(grid) else []
                sign = str(gr[1] or "") if len(gr) > 1 else ""
                date = str(gr[2] or "") if len(gr) > 2 else ""
                label = first if first.rstrip().endswith("：") else first + "："
                set_cell(merged, label, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
                para = merged.paragraphs[0]
                if sign.startswith("data:image"):
                    try:
                        b64 = sign.split(",", 1)[1] if "," in sign else ""
                        para.add_run("  ")
                        para.add_run().add_picture(BytesIO(base64.b64decode(b64)), height=Pt(28))
                    except Exception:
                        pass
                elif sign.strip():
                    rn = para.add_run("  " + sign)
                    rn.font.size = Pt(10.5)
                if date.strip():
                    rd = para.add_run("      " + date)
                    rd.font.size = Pt(10.5)
                merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            elif first.startswith("评审结论"):
                # 评审结论为长段落，左对齐
                set_cell(merged, first, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)
            else:
                set_cell(merged, first, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    # 「其他参会人员/其他参评人员」行：标签后单元格合并为一格，居中显示「/」
    for r in range(n):
        first = rows[r].cells[0].text
        if _is_other_row(first) and len(rows[r].cells) > 2:
            merged = rows[r].cells[1].merge(rows[r].cells[-1])
            set_cell(merged, "/", align=WD_ALIGN_PARAGRAPH.CENTER)
            merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # 首列纵向合并：非空且非标记的单元格向下合并其后的空单元格或相同内容单元格
    r = 0
    while r < n:
        text = rows[r].cells[0].text
        if text.strip() and not _is_banner(text):
            r2 = r
            # 向下合并：后续单元格为空，或与当前单元格内容相同（评审内容表首列相同类别合并）
            while r2 + 1 < n:
                nxt = rows[r2 + 1].cells[0].text
                if _is_banner(nxt):
                    break
                if nxt.strip() == text.strip() or not nxt.strip():
                    r2 += 1
                else:
                    break
            if r2 > r:
                merged = rows[r].cells[0].merge(rows[r2].cells[0])
                set_cell(merged, text, align=WD_ALIGN_PARAGRAPH.CENTER)
                merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            r = r2 + 1
        else:
            r += 1
    # 加大行高，避免签名图/文字压到单元格边框
    for _row in table.rows:
        _row.height = Pt(40)
        _row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    document.add_paragraph()
