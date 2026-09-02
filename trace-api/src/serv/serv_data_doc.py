#!/usr/bin/env python
# encoding: utf-8

# 数据文件服务层，详见 docs/function_docs/100_数据文件管理.md。
# 单表 data_doc + doc_type；导出结构与产品立项报告一致：封面→分页→修订记录→分页→目录→分页→正文。

import base64
import copy
import logging
import re
from datetime import date, timedelta
from io import BytesIO
from typing import List
from sqlalchemy import delete, func, select
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.data_doc import DataDoc
from ..model.project_timeline import ProjectTimelineRow, ProjectTimelineCell
from ..model.project_member import ProjectMember
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_data_doc import DataDocForm
from ..obj.vobj_data_doc import DataDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from . import serv_review_util
from .serv_utils import new_version, sync_file_no_version
from .serv_utils import docx_util
from .data_doc_templates import DOC_META, DEFAULT_CONTENTS, REVIEW_TABLES

logger = logging.getLogger(__name__)

COVER_DEPT = "数据部"
ENV_DOC_TYPES = ("dd_016", "dd_017")
ENV_CHECK_GROUPS = {
    ("dd_016", "server"): [
        ("日期", []),
        ("硬件环境", ["CPU", "GPU", "内存", "网卡"]),
        ("软件环境", ["操作系统\n运行是否正常", "数据库\n运行是否正常", "应用服务\n运行是否正常"]),
        ("开发环境\n是否更新升级", []),
        ("服务器\n是否杀毒", []),
        ("网络环境\n是否正常", []),
        ("开发工具", ["是否正常运行", "是否更新升级"]),
        ("服务器\n是否备份", []),
        ("服务器\n日志是否错误", []),
        ("出现的问题及处理方式", []),
        ("检查人", []),
    ],
    ("dd_016", "dev"): [
        ("日期", []),
        ("硬件环境", ["CPU", "GPU", "内存", "网卡"]),
        ("软件环境", ["操作系统\n运行是否正常", "浏览器\n运行是否正常"]),
        ("开发环境\n是否更新升级", []),
        ("开发机\n是否杀毒", []),
        ("网络环境\n是否正常", []),
        ("开发工具", ["是否正常运行", "是否更新升级"]),
        ("出现的问题及处理方式", []),
        ("检查人", []),
    ],
    ("dd_017", "server"): [
        ("日期", []),
        ("硬件环境", ["CPU", "GPU", "内存", "网卡"]),
        ("软件环境", ["操作系统\n运行是否正常", "数据库\n运行是否正常", "应用服务\n运行是否正常"]),
        ("标注环境\n是否更新升级", []),
        ("服务器\n是否杀毒", []),
        ("网络环境\n是否正常", []),
        ("标注工具", ["是否正常运行", "是否更新升级"]),
        ("服务器\n是否备份", []),
        ("服务器\n日志是否错误", []),
        ("出现的问题及处理方式", []),
        ("检查人", []),
    ],
    ("dd_017", "dev"): [
        ("日期", []),
        ("硬件环境", ["CPU", "GPU", "内存", "网卡"]),
        ("软件环境", ["操作系统\n运行是否正常", "浏览器\n运行是否正常"]),
        ("标注环境\n是否更新升级", []),
        ("标注机\n是否杀毒", []),
        ("网络环境\n是否正常", []),
        ("标注工具", ["是否正常运行", "是否更新升级"]),
        ("出现的问题及处理方式", []),
        ("检查人", []),
    ],
}


def doc_title(doc_type):
    return (DOC_META.get(doc_type) or {}).get("title") or "数据文件"


def doc_keywords(doc_type):
    return list((DOC_META.get(doc_type) or {}).get("keywords") or [])


def doc_format(doc_type):
    return (DOC_META.get(doc_type) or {}).get("format") or "docx"


def _empty_template(doc_type):
    title = doc_title(doc_type)
    return {
        "sections": [
            {
                "title": title, "ref_type": "cover", "body": "", "children": [],
                "tables": [[
                    ["编制部门", COVER_DEPT, "文件版本", "A0"],
                    ["编制人", "", "日期", ""],
                    ["审核人", "", "日期", ""],
                    ["批准人", "", "日期", ""],
                    ["生效日期", "", "", ""],
                ]],
            },
            {
                "title": "文件修订记录", "ref_type": "revision", "body": "", "children": [],
                "tables": [[
                    ["修改日期", "版本号", "修订说明", "修订人", "批准人"],
                    ["", "", "首次发布", "", ""],
                    ["", "", "", "", ""],
                    ["", "", "", "", ""],
                    ["", "", "", "", ""],
                    ["", "", "", "", ""],
                ]],
            },
        ]
    }


class Server(object):

    def __default_content(self, doc_type):
        raw = DEFAULT_CONTENTS.get(doc_type)
        content = copy.deepcopy(raw) if raw else _empty_template(doc_type)
        self.__drop_product_info(content)
        self.__ensure_review_annex(content, doc_type)
        if doc_type in ENV_DOC_TYPES:
            self.__complete_env_maint_chapter(content, doc_type)
        return content

    def __to_obj(self, row: DataDoc, product: Product = None):
        obj = DataDocObj(**row.dict())
        obj.content = self.__normalize_content(obj.content, row.doc_type)
        self.__fill_cover_meta(obj.content, obj.version)
        key = row.doc_type or ""
        serv_review_util.fill_cover_dates(
            obj.content, serv_review_util.cover_date(row.product_id, key) if row.product_id else ""
        )
        serv_review_util.fill_cover_signers(
            obj.content, serv_review_util.cover_signers(row.product_id, key) if row.product_id else {}
        )
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
            if not (obj.file_no or "").strip():
                resolved = serv_review_util.resolve_doc_file_no(product.id, obj.file_no, obj.version, key)
                if resolved:
                    obj.file_no = resolved
        return obj

    @staticmethod
    def __migrate_cover_table(rows):
        """把旧版 2 列或「使用部门/版本号」封面迁移为 4 列：编制部门 / 数据部。"""
        rows = [r for r in (rows or []) if isinstance(r, list)]
        if rows and len(rows[0]) >= 4 and str(rows[0][0]).strip() == "编制部门":
            if not str(rows[0][1] if len(rows[0]) > 1 else "").strip():
                rows[0][1] = COVER_DEPT
            return rows
        items = [(str(r[0]).strip(), str(r[1]).strip() if len(r) > 1 else "") for r in rows if r]
        def val(*labels):
            for want in labels:
                for l, v in items:
                    if l == want:
                        return v
            return ""
        dates = [v for l, v in items if l == "日期"]
        d = lambda i: dates[i] if i < len(dates) else ""
        dept = val("编制部门", "使用部门", "编写部门") or COVER_DEPT
        ver = val("文件版本", "版本号") or "A0"
        return [
            ["编制部门", dept, "文件版本", ver],
            ["编制人", val("编制人"), "日期", d(0)],
            ["审核人", val("审核人"), "日期", d(1)],
            ["批准人", val("批准人"), "日期", d(2)],
            ["生效日期", val("生效日期"), "", ""],
        ]

    def __normalize_node(self, node):
        if not isinstance(node, dict):
            return {"title": str(node or ""), "body": "", "tables": [], "children": []}
        result = dict(node)
        result["title"] = str(result.get("title") or "")
        result["body"] = str(result.get("body") or "")
        tables = result.get("tables")
        if not isinstance(tables, list):
            tables = []
        norm_tables = []
        for table in tables:
            if isinstance(table, list):
                norm_tables.append([[str(c) if c is not None else "" for c in (row or [])] for row in table if isinstance(row, list)])
        if result.get("ref_type") == "cover" and norm_tables:
            norm_tables = [self.__migrate_cover_table(norm_tables[0])] + norm_tables[1:]
        result["tables"] = norm_tables
        children = result.get("children")
        result["children"] = [self.__normalize_node(c) for c in children] if isinstance(children, list) else []
        return result

    def __normalize_content(self, content, doc_type=None):
        if not isinstance(content, dict) or not isinstance(content.get("sections"), list):
            return self.__default_content(doc_type)
        out = {"sections": [self.__normalize_node(s) for s in content["sections"]]}
        self.__drop_product_info(out)
        self.__ensure_review_annex(out, doc_type)
        if doc_type in ENV_DOC_TYPES:
            self.__complete_env_maint_chapter(out, doc_type)
        return out

    @classmethod
    def __drop_product_info(cls, content):
        """原样例没有独立「产品信息」章；新增/打开/导出时去掉。"""
        def drop(ns):
            out = []
            for n in ns or []:
                t = cls.__strip_num(n.get("title"))
                if n.get("ref_type") == "basic_info" or t == "产品信息":
                    continue
                n["children"] = drop(n.get("children") or [])
                out.append(n)
            return out
        content["sections"] = drop((content or {}).get("sections") or [])

    @classmethod
    def __is_annex_title(cls, title):
        t = cls.__strip_num(title).replace(" ", "")
        return t.startswith("附件") and "评审记录" in t

    @staticmethod
    def __is_review_grid(tb):
        return isinstance(tb, list) and tb and isinstance(tb[0], list) and "评审记录" in str(tb[0][0] or "")

    @staticmethod
    def __cell_eq_nonempty(a, b):
        sa, sb = str(a or ""), str(b or "")
        return sa == sb and sa.strip() != ""

    def __grid_span_origins(self, grid):
        rows = [list(r) for r in grid]
        r_n = len(rows)
        c_n = max((len(r) for r in rows), default=0)
        for r in rows:
            while len(r) < c_n:
                r.append("")
        skip = [[False] * c_n for _ in range(r_n)]
        colspan = [[1] * c_n for _ in range(r_n)]
        rowspan = [[1] * c_n for _ in range(r_n)]
        for r in range(r_n):
            c = 0
            while c < c_n:
                if skip[r][c]:
                    c += 1
                    continue
                c2 = c
                while c2 + 1 < c_n and self.__cell_eq_nonempty(rows[r][c], rows[r][c2 + 1]):
                    c2 += 1
                colspan[r][c] = c2 - c + 1
                for k in range(c + 1, c2 + 1):
                    skip[r][k] = True
                c = c2 + 1
        for r in range(r_n):
            for c in range(c_n):
                if skip[r][c]:
                    continue
                if c != 0:
                    continue
                if not str(rows[r][c] or "").strip():
                    continue
                cs = colspan[r][c]
                r2 = r
                while r2 + 1 < r_n:
                    if skip[r2 + 1][c] or colspan[r2 + 1][c] != cs:
                        break
                    if not self.__cell_eq_nonempty(rows[r][c], rows[r2 + 1][c]):
                        break
                    r2 += 1
                rs = r2 - r + 1
                if rs > 1:
                    rowspan[r][c] = rs
                    for k in range(r + 1, r2 + 1):
                        skip[k][c] = True
        origins = []
        for r in range(r_n):
            for c in range(c_n):
                if skip[r][c]:
                    continue
                origins.append((r, c, rowspan[r][c], colspan[r][c]))
        return rows, origins

    @classmethod
    def __ensure_review_annex(cls, content, doc_type=None):
        """原 Word 有评审表的，附件为空则补上；误挂在其它章节的评审表挪回附件。"""
        sections = (content or {}).get("sections")
        if not isinstance(sections, list):
            return

        def strip_annex_line(n):
            body = str(n.get("body") or "")
            lines = [ln for ln in body.split("\n") if ln.strip() not in ("附件 1 评审记录", "附件1 评审记录")]
            n["body"] = "\n".join(lines).rstrip()

        annex_nodes = []

        def find(ns):
            for n in ns or []:
                if not isinstance(n, dict):
                    continue
                strip_annex_line(n)
                if cls.__is_annex_title(n.get("title")):
                    annex_nodes.append(n)
                find(n.get("children") or [])

        find(sections)
        annex = annex_nodes[0] if annex_nodes else None
        annex_ids = {id(n) for n in annex_nodes}
        pulled = []

        def pull(ns):
            for n in ns or []:
                if not isinstance(n, dict):
                    continue
                if id(n) in annex_ids:
                    pull(n.get("children") or [])
                    continue
                keep = []
                for tb in (n.get("tables") or []):
                    if cls.__is_review_grid(tb):
                        pulled.append(tb)
                    else:
                        keep.append(tb)
                n["tables"] = keep
                pull(n.get("children") or [])

        pull(sections)
        src = copy.deepcopy(REVIEW_TABLES.get(doc_type) or []) if doc_type else []
        table = None
        if annex:
            existing = [tb for tb in (annex.get("tables") or []) if cls.__is_review_grid(tb)]
            table = existing[0] if existing else (pulled[0] if pulled else (src or None))
            if table:
                annex["tables"] = [table]
                annex["title"] = "附件 1 评审记录"
        elif pulled or src:
            table = pulled[0] if pulled else src
            sections.append({"title": "附件 1 评审记录", "body": "", "tables": [table], "children": []})

    @staticmethod
    def __fill_cover_meta(content, version):
        """封面编制部门 / 文件版本：仅填空。"""
        for section in (content or {}).get("sections") or []:
            if not isinstance(section, dict):
                continue
            for table in (section.get("tables") or []):
                if not isinstance(table, list):
                    continue
                for row in table:
                    if not isinstance(row, list) or not row:
                        continue
                    label = str(row[0] or "").strip()
                    if label in ("编制部门", "使用部门", "编写部门") and len(row) >= 2:
                        if not str(row[1] or "").strip():
                            row[1] = COVER_DEPT
                    if label in ("文件版本", "版本号") and len(row) >= 4:
                        if version and not str(row[3] or "").strip():
                            row[3] = version

    def __autofill_for_export(self, content, obj: DataDocObj):
        sections = (content or {}).get("sections") or []
        prod_id = obj.product_id
        if not prod_id:
            return content
        product = db.session.execute(select(Product).where(Product.id == prod_id)).scalars().first()
        info = self.__collect_autofill(prod_id, product, obj.version, obj.doc_type)
        for node in sections:
            self.__fill_node(node, info)
        if obj.doc_type in ENV_DOC_TYPES:
            def walk(ns):
                for n in ns or []:
                    if isinstance(n, dict):
                        self.__fill_env_maint_node(n, info)
                        walk(n.get("children") or [])
            walk(sections)
            self.__rebuild_env_checks(content, info)
        self.__fill_cover_meta(content, obj.version)
        key = obj.doc_type or ""
        serv_review_util.fill_cover_dates(content, serv_review_util.cover_date(prod_id, key))
        serv_review_util.fill_cover_signers(content, serv_review_util.cover_signers(prod_id, key))
        return content

    def __collect_autofill(self, prod_id, product, doc_version, doc_type):
        prod_name = (getattr(product, "name", "") or "").strip()
        full_version = (getattr(product, "full_version", "") or "").strip()
        product_code = (getattr(product, "product_code", "") or "").strip()
        scope = (getattr(product, "scope", "") or "").strip()

        tl_rows = db.session.execute(
            select(ProjectTimelineRow).where(ProjectTimelineRow.prod_id == prod_id)
        ).scalars().all()
        cell_map = {}
        if tl_rows:
            for c in db.session.execute(
                select(ProjectTimelineCell).where(ProjectTimelineCell.row_id.in_([r.id for r in tl_rows]))
            ).scalars().all():
                cell_map.setdefault(c.row_id, []).append(c.output_result or "")

        def to_int(v):
            digits = re.sub(r"[^\d]", "", str(v or ""))
            return int(digits) if digits else None

        date_rows = [r for r in tl_rows if (r.row_type or "date") == "date" and to_int(r.year) and to_int(r.month)]

        def date_key(r):
            return to_int(r.year) * 10000 + to_int(r.month) * 100 + (to_int(r.day) or 0)

        kws = doc_keywords(doc_type) or [doc_title(doc_type)]
        file_rows = [
            r for r in date_rows
            if any(any(k in str(v or "") for k in kws) for v in cell_map.get(r.id, []))
        ]
        file_date = ""
        if file_rows:
            fr = min(file_rows, key=date_key)
            file_date = f"{to_int(fr.year)}年{to_int(fr.month)}月{to_int(fr.day)}日"

        members = db.session.execute(select(ProjectMember).where(ProjectMember.prod_id == prod_id)).scalars().all()
        def find_member(pred):
            for m in members:
                if pred(str(m.role or "")):
                    return (m.name or "").strip()
            return ""
        def member_names(pred):
            return [(m.name or "").strip() for m in members if (m.name or "").strip() and pred(str(m.role or ""))]
        modeler = find_member(lambda r: "模型" in r)
        algo = find_member(lambda r: "算法" in r)
        approver = find_member(lambda r: "负责人" in r)

        env_weeks = []
        env_assets = None
        env_checker = env_title = ""
        if doc_type in ENV_DOC_TYPES:
            env_title = "开发环境维护记录" if doc_type == "dd_016" else "标注环境维护记录"
            usage_kw = "标注" if doc_type == "dd_017" else "开发"
            dev_ds, test_ds = [], []
            for r in date_rows:
                y, m, d = to_int(r.year), to_int(r.month), to_int(r.day) or 1
                try:
                    dt = date(y, m, d)
                except Exception:
                    continue
                vals = cell_map.get(r.id, [])
                if any(("产品开发" in str(v)) and ("计划" not in str(v)) for v in vals):
                    dev_ds.append(dt)
                if any("测试" in str(v) for v in vals):
                    test_ds.append(dt)
            if dev_ds and test_ds:
                env_weeks = self.__week_ranges_from_dates([min(dev_ds), max(test_ds)])
            eq_doc = db.session.execute(
                select(DataDoc).where(DataDoc.product_id == prod_id, DataDoc.doc_type == "dd_eq").order_by(DataDoc.id.desc())
            ).scalars().first()
            if eq_doc:
                env_assets = self.__parse_eq_codes(eq_doc.content if isinstance(eq_doc.content, dict) else {}, usage_kw)
            checkers = member_names(lambda r: r == "数据部负责人") or member_names(lambda r: r == "数据负责人") or member_names(lambda r: "数据" in r)
            env_checker = serv_review_util._sign_by_name(checkers[0] if checkers else "") or (checkers[0] if checkers else "")

        return {
            "prod_name": prod_name, "full_version": full_version, "product_code": product_code,
            "scope": scope, "file_date": file_date, "version": doc_version,
            "reviser": modeler or algo, "approver": approver,
            "doc_type": doc_type, "env_weeks": env_weeks, "env_assets": env_assets,
            "env_checker": env_checker, "env_title": env_title,
        }

    @staticmethod
    def __strip_num(title):
        return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()

    def __fill_node(self, node, info):
        ref = node.get("ref_type")
        title = self.__strip_num(node.get("title"))
        if ref == "revision" or title == "文件修订记录":
            tables = node.get("tables") or []
            if tables and isinstance(tables[0], list):
                t = tables[0]
                cols = len(t[0]) if t and t[0] else 5
                while len(t) < 6:
                    t.append([""] * cols)
                row = t[1]
                def set_if(i, val):
                    if val and not str(row[i] if i < len(row) else "").strip():
                        row[i] = val
                set_if(0, info["file_date"])
                set_if(1, info["version"])
                if not str(row[2] if len(row) > 2 else "").strip():
                    row[2] = "首次发布"
                set_if(3, info["reviser"])
                set_if(4, info["approver"])
        if ref == "basic_info" or title == "产品信息":
            label_map = {
                "产品名称": info["prod_name"],
                "软件版本": info["full_version"],
                "完整版本": info["full_version"],
                "产品标识": info["product_code"],
                "产品代码": info["product_code"],
                "适用范围": info["scope"],
                "预期用途": info["scope"],
                "项目名称": info["prod_name"],
            }
            for table in (node.get("tables") or []):
                for row in table:
                    if not isinstance(row, list) or len(row) < 2:
                        continue
                    key = str(row[0]).strip()
                    if key in label_map and label_map[key] and not str(row[1] or "").strip():
                        row[1] = label_map[key]
        for child in (node.get("children") or []):
            self.__fill_node(child, info)

    def __exists(self, product_id, doc_type, version, exclude_id=None):
        sql = select(func.count(DataDoc.id)).where(
            DataDoc.product_id == product_id,
            DataDoc.doc_type == doc_type,
            DataDoc.version == version,
        )
        if exclude_id:
            sql = sql.where(DataDoc.id != exclude_id)
        return (db.session.execute(sql).scalar() or 0) > 0

    @staticmethod
    def __is_env_check_grid(tb):
        return isinstance(tb, list) and tb and isinstance(tb[0], list) and str(tb[0][0] or "").strip() == "env_check"

    @staticmethod
    def __is_asset_grid(tb):
        if not (isinstance(tb, list) and tb and isinstance(tb[0], list) and tb[0]):
            return False
        hdr = [str(c or "") for c in tb[0]]
        return str(hdr[0] or "").strip() == "资产编码" and any("设备信息" in h for h in hdr)

    @staticmethod
    def __env_check_leaves(doc_type, kind):
        cols = []
        for gl, leaves in ENV_CHECK_GROUPS.get((doc_type, kind), ENV_CHECK_GROUPS[("dd_016", "dev")]):
            if leaves:
                for lf in leaves:
                    cols.append({"label": lf, "type": "check"})
            else:
                t = "date" if gl == "日期" else "problem" if gl.startswith("出现的问题") else "checker" if gl == "检查人" else "check"
                cols.append({"label": gl, "type": t})
        return cols

    @classmethod
    def __env_check_defaults(cls, doc_type, kind):
        out = []
        for c in cls.__env_check_leaves(doc_type, kind):
            if c["type"] != "check":
                continue
            lb = c["label"]
            out.append("否" if ("更新升级" in lb or "日志是否错误" in lb) else "是")
        return out

    def __complete_env_maint_chapter(self, content, doc_type):
        if doc_type not in ENV_DOC_TYPES:
            return
        want = "开发环境维护记录" if doc_type == "dd_016" else "标注环境维护记录"
        after = "开发环境定期检查" if doc_type == "dd_016" else "标注环境定期检查"
        old = "开发环境定期验证" if doc_type == "dd_016" else "标注环境定期验证"
        sections = (content or {}).get("sections")
        if not isinstance(sections, list):
            return

        def rename(ns):
            for n in ns or []:
                if not isinstance(n, dict):
                    continue
                if self.__strip_num(n.get("title")) == old:
                    n["title"] = after
                rename(n.get("children") or [])

        rename(sections)

        def find_title(ns, name):
            for n in ns or []:
                if not isinstance(n, dict):
                    continue
                if self.__strip_num(n.get("title")) == name:
                    return n
                hit = find_title(n.get("children") or [], name)
                if hit:
                    return hit
            return None

        if find_title(sections, want):
            inspect = find_title(sections, after)
            self.__ensure_asset_header(inspect)
            return
        new_node = {"title": want, "body": "", "tables": [], "children": []}
        idx = next((i for i, n in enumerate(sections) if isinstance(n, dict) and self.__strip_num(n.get("title")) == after), -1)
        if idx >= 0:
            sections.insert(idx + 1, new_node)
        else:
            sections.append(new_node)
        self.__ensure_asset_header(find_title(sections, after))

    @staticmethod
    def __ensure_asset_header(node):
        if not isinstance(node, dict):
            return
        tables = node.get("tables") or []
        for i, tb in enumerate(tables):
            if not (isinstance(tb, list) and tb and isinstance(tb[0], list)):
                continue
            hdr = [str(c or "") for c in tb[0]]
            if str(hdr[0] or "").strip() == "资产编码":
                continue
            if len(hdr) >= 2:
                tables[i] = [["资产编码", "设备信息", "产品名称", "完整版本"]] + tb
        node["tables"] = tables

    @staticmethod
    def __week_ranges_from_dates(dates):
        if not dates:
            return []
        start_d, end_d = min(dates), max(dates)
        if start_d > end_d:
            return []

        def fmt(d):
            return f"{d.year}.{d.month:02d}.{d.day:02d}"

        ranges = []
        cur = start_d - timedelta(days=start_d.weekday())
        while cur <= end_d:
            monday = cur
            friday = monday + timedelta(days=4)
            ws = max(monday, start_d)
            we = min(friday, end_d)
            if ws.weekday() >= 5:
                cur = monday + timedelta(days=7)
                continue
            if we.weekday() >= 5:
                we = friday
            if ws <= we:
                ranges.append(f"{fmt(ws)}- {fmt(we)}")
            cur = monday + timedelta(days=7)
        return ranges

    def __parse_eq_table(self, tb, usage_contains=""):
        out, seen = [], set()
        if not isinstance(tb, list):
            return out
        hi = brand_i = code_i = name_i = usage_i = -1
        for i, row in enumerate(tb):
            if not isinstance(row, list):
                continue
            cells = [str(c or "").strip() for c in row]
            if any(c == "品牌" for c in cells) and any("资产编码" in c for c in cells):
                hi = i
                brand_i = next(j for j, c in enumerate(cells) if c == "品牌")
                code_i = next(j for j, c in enumerate(cells) if "资产编码" in c)
                name_i = next((j for j, c in enumerate(cells) if c == "名称"), -1)
                usage_i = next((j for j, c in enumerate(cells) if c == "用途"), -1)
                break
        if hi < 0:
            return out
        for row in tb[hi + 1:]:
            if not isinstance(row, list):
                continue
            brand = str(row[brand_i] if brand_i < len(row) else "").strip()
            code = str(row[code_i] if code_i < len(row) else "").strip()
            name = str(row[name_i] if 0 <= name_i < len(row) else "").strip()
            usage = str(row[usage_i] if 0 <= usage_i < len(row) else "").strip()
            if name == "显示器":
                continue
            if usage_contains and usage_contains not in usage:
                continue
            if brand in ("组装机", "Apple") and code and code not in seen:
                seen.add(code)
                out.append((code, usage))
        return out

    def __parse_eq_codes(self, content, usage_contains=""):
        if isinstance(content, dict) and isinstance(content.get("rows"), list):
            return self.__parse_eq_table(content.get("rows"), usage_contains)
        out, seen = [], set()
        for code, usage in self.__parse_eq_table_walk((content or {}).get("sections") or [], usage_contains):
            if code not in seen:
                seen.add(code)
                out.append((code, usage))
        return out

    def __parse_eq_table_walk(self, ns, usage_contains=""):
        out = []
        for n in ns or []:
            if not isinstance(n, dict):
                continue
            for tb in n.get("tables") or []:
                out.extend(self.__parse_eq_table(tb, usage_contains))
            out.extend(self.__parse_eq_table_walk(n.get("children") or [], usage_contains))
        return out

    def __fill_env_maint_node(self, node, info):
        eq_assets = info.get("env_assets")
        prod_name = info.get("prod_name") or ""
        full_version = info.get("full_version") or ""
        tables = node.get("tables") or []
        for ti, tb in enumerate(tables):
            if not isinstance(tb, list) or not self.__is_asset_grid(tb):
                continue
            old = {}
            for row in tb[1:]:
                if isinstance(row, list) and row:
                    code = str(row[0] or "").strip()
                    if code:
                        old[code] = str(row[1] if len(row) > 1 else "")
            hdr = [str(c or "") for c in tb[0]]
            hdr = hdr[:4] if len(hdr) >= 4 else ["资产编码", "设备信息", "产品名称", "完整版本"]
            if eq_assets is not None:
                body = [[code, old.get(code, ""), prod_name, full_version] for code, _u in eq_assets]
                tables[ti] = [hdr] + body
            else:
                new_tb = [hdr]
                for row in tb[1:]:
                    if not isinstance(row, list):
                        continue
                    next_row = list(row)
                    while len(next_row) < 4:
                        next_row.append("")
                    next_row[2] = prod_name
                    next_row[3] = full_version
                    new_tb.append(next_row)
                tables[ti] = new_tb
        node["tables"] = tables

    def __collect_asset_codes(self, content):
        out, seen = [], set()

        def walk(ns):
            for n in ns or []:
                if not isinstance(n, dict):
                    continue
                for tb in n.get("tables") or []:
                    if not self.__is_asset_grid(tb):
                        continue
                    for row in tb[1:]:
                        if not isinstance(row, list) or not row:
                            continue
                        code = str(row[0] or "").strip()
                        if code and code not in seen:
                            seen.add(code)
                            out.append((code, ""))
                walk(n.get("children") or [])

        walk((content or {}).get("sections") or [])
        return out

    def __rebuild_env_checks(self, content, info):
        want = info.get("env_title") or ""
        doc_type = info.get("doc_type") or "dd_016"
        weeks = info.get("env_weeks") or []
        checker = info.get("env_checker") or ""
        assets = info.get("env_assets")
        if assets is None:
            assets = self.__collect_asset_codes(content)

        def find_title(ns, name):
            for n in ns or []:
                if not isinstance(n, dict):
                    continue
                if self.__strip_num(n.get("title")) == name:
                    return n
                hit = find_title(n.get("children") or [], name)
                if hit:
                    return hit
            return None

        node = find_title((content or {}).get("sections") or [], want)
        if not node:
            return
        old = {}
        for tb in node.get("tables") or []:
            if not self.__is_env_check_grid(tb):
                continue
            code = str(tb[0][2] if len(tb[0]) > 2 else "")
            by_date = {}
            for row in tb[1:]:
                if isinstance(row, list) and row:
                    by_date[str(row[0] or "")] = [str(c or "") for c in row]
            old[code] = by_date
        tables = []
        for code, usage in assets or []:
            kind = "server" if "共用" in str(usage or "") else "dev"
            defaults = self.__env_check_defaults(doc_type, kind)
            prev = old.get(code) or {}
            rows = [["env_check", kind, code]]
            for w in weeks:
                p = prev.get(w) or []
                marks = []
                for i, d in enumerate(defaults):
                    v = str(p[i + 1] if i + 1 < len(p) else "").strip()
                    marks.append(v if v in ("是", "否") else d)
                problem = str(p[len(defaults) + 1] if len(p) > len(defaults) + 1 else "").strip() or "无"
                rows.append([w] + marks + [problem, checker])
            tables.append(rows)
        node["tables"] = tables

    def __apply_env_eq_assets(self, obj: DataDocObj, product: Product = None):
        if obj.doc_type not in ENV_DOC_TYPES or not obj.product_id:
            return
        info = self.__collect_autofill(obj.product_id, product, obj.version, obj.doc_type)

        def walk(ns):
            for n in ns or []:
                if isinstance(n, dict):
                    self.__fill_env_maint_node(n, info)
                    walk(n.get("children") or [])

        walk((obj.content or {}).get("sections") or [])
        self.__rebuild_env_checks(obj.content, info)

    async def add_data_doc(self, form: DataDocForm):
        try:
            doc_type = (form.doc_type or "").strip()
            if doc_type not in DOC_META:
                return Resp.resp_err(msg=ts("msg_err_param"))
            if self.__exists(form.product_id, doc_type, form.version):
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            payload = form.dict(exclude_none=True)
            payload["doc_type"] = doc_type
            row = DataDoc(**payload)
            row.id = None
            row.file_no = serv_review_util.resolve_doc_file_no(form.product_id, form.file_no, form.version, doc_type) or None
            row.content = self.__normalize_content(row.content, doc_type)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=DataDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_data_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: DataDoc = db.session.execute(select(DataDoc).where(DataDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(
                select(DataDoc.version).where(DataDoc.product_id == target_pid, DataDoc.doc_type == fromdoc.doc_type)
            ).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = DataDoc(
                product_id=target_pid,
                doc_type=fromdoc.doc_type,
                version=version,
                file_no=sync_file_no_version(
                    (fromdoc.file_no or "").strip()
                    or serv_review_util.resolve_doc_file_no(target_pid, "", version, fromdoc.doc_type)
                    or "",
                    version,
                ) or None,
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content, fromdoc.doc_type)),
            )
            db.session.add(newdoc)
            db.session.commit()
            return Resp.resp_ok(data=DataDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_data_doc(self, form: DataDocForm):
        try:
            row: DataDoc = db.session.execute(select(DataDoc).where(DataDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            payload = form.dict(exclude_none=True)
            payload.pop("doc_type", None)
            next_pid = payload.get("product_id", row.product_id)
            next_ver = payload.get("version", row.version)
            if next_pid != row.product_id or next_ver != row.version:
                if self.__exists(next_pid, row.doc_type, next_ver, exclude_id=row.id):
                    return Resp.resp_err(msg=ts("msg_obj_exist"))
            for key, value in payload.items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value, row.doc_type)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_data_doc(self, id: int):
        db.session.execute(delete(DataDoc).where(DataDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_data_doc(self, id: int):
        sql = select(DataDoc, Product).join(Product, DataDoc.product_id == Product.id).where(DataDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        obj = self.__to_obj(doc, product)
        self.__apply_env_eq_assets(obj, product)
        return Resp.resp_ok(data=obj)

    def parse_stats_excel(self, raw: bytes):
        """解析统计脚本输出的 xlsx：每个工作表一章，供编辑页展示。不落库。"""
        from openpyxl import load_workbook
        if not raw:
            return Resp.resp_err(msg="请选择 Excel 文件")
        try:
            wb = load_workbook(BytesIO(raw), data_only=True)
        except Exception:
            return Resp.resp_err(msg="无法读取 Excel")
        sections = []
        for name in wb.sheetnames:
            ws = wb[name]
            grid = []
            max_col = ws.max_column or 1
            max_row = ws.max_row or 1
            for row in ws.iter_rows(min_row=1, max_row=max_row, max_col=max_col, values_only=True):
                cells = ["" if v is None else str(v) for v in row]
                if not any(str(c).strip() for c in cells):
                    continue
                while cells and not str(cells[-1]).strip():
                    cells.pop()
                grid.append(cells)
            sections.append({
                "title": (name or "Sheet")[:64],
                "body": "",
                "tables": [grid] if grid else [],
                "children": [],
            })
        if not sections or not any((s.get("tables") or [None])[0] for s in sections):
            return Resp.resp_err(msg="Excel 无有效表格")
        return Resp.resp_ok(data={"sections": sections})

    async def list_data_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None,
                             doc_type: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if doc_type:
            wheres.append(DataDoc.doc_type == doc_type)
        if product_id:
            wheres.append(DataDoc.product_id == product_id)
        if version:
            wheres.append(DataDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(DataDoc.id)).join(Product, DataDoc.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (
            select(DataDoc, Product)
            .join(Product, DataDoc.product_id == Product.id)
            .where(*wheres)
            .order_by(DataDoc.id.desc())
            .offset(page_index * page_size)
            .limit(page_size)
        )
        rows: List[DataDocObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    def __export_xlsx(self, output, obj: DataDocObj, content):
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, Border, Side

        wb = Workbook()
        default = wb.active
        used_names = set()
        thin = Border(
            left=Side(style="thin"), right=Side(style="thin"),
            top=Side(style="thin"), bottom=Side(style="thin"),
        )

        def sheet_name(title):
            raw = self.__strip_num(title) or "Sheet"
            name = re.sub(r'[:\\/\?\*\[\]]', "_", raw)[:31] or "Sheet"
            base = name
            idx = 2
            while name.lower() in used_names:
                suffix = str(idx)
                name = (base[: 31 - len(suffix)] + suffix)
                idx += 1
            used_names.add(name.lower())
            return name

        def write_table(ws, grid, start_row=1):
            r_idx = start_row
            for r_i, row in enumerate(grid or []):
                if not isinstance(row, list):
                    continue
                for c_i, val in enumerate(row, 1):
                    s = str(val or "")
                    if s.startswith("data:image"):
                        s = "[签名]"
                    cell = ws.cell(r_idx, c_i, s)
                    cell.alignment = Alignment(wrap_text=True, vertical="center")
                    cell.border = thin
                    if r_i == 0:
                        cell.font = Font(bold=True, name="宋体")
                    else:
                        cell.font = Font(name="宋体")
                r_idx += 1
            return r_idx

        first = True
        for node in (content or {}).get("sections") or []:
            ws = default if first else wb.create_sheet()
            first = False
            ws.title = sheet_name(node.get("title"))
            title = self.__strip_num(node.get("title"))
            ws.cell(1, 1, title).font = Font(bold=True, size=14, name="宋体")
            row = 3
            if (node.get("body") or "").strip():
                ws.cell(row, 1, node.get("body"))
                row += 2
            for table in (node.get("tables") or []):
                row = write_table(ws, table, row) + 2
            for child in (node.get("children") or []):
                ws.cell(row, 1, self.__strip_num(child.get("title"))).font = Font(bold=True, name="宋体")
                row += 1
                if (child.get("body") or "").strip():
                    ws.cell(row, 1, child.get("body"))
                    row += 1
                for table in (child.get("tables") or []):
                    row = write_table(ws, table, row) + 2
        if first:
            default.title = "数据文件"
        wb.save(output)
        output.seek(0)

    async def export_data_doc(self, output, id: int):
        resp = await self.get_data_doc(id)
        obj: DataDocObj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return "数据文件", "docx"
        c = self.__autofill_for_export(self.__normalize_content(obj.content, obj.doc_type), obj)
        title = doc_title(obj.doc_type)
        if doc_format(obj.doc_type) == "xlsx":
            self.__export_xlsx(output, obj, c)
            return title, "xlsx"
        sections = c.get("sections") or []
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        document.settings.element.append(update_fields)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")
        docx_util.add_page_number_footer(section, obj.file_no or "")

        def write_center_title(text, size=22.0, bold=False):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.right_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, text, font_size=size, bold=bold)

        def add_blank_lines(count):
            for _ in range(max(0, int(count or 0))):
                document.add_paragraph("")

        def add_text(text):
            docx_util.save_txt2docx(str(text or ""), document)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            s = str(text or "")
            if s.startswith("data:image"):
                try:
                    b64 = s.split(",", 1)[1] if "," in s else ""
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = align
                    para.add_run().add_picture(BytesIO(base64.b64decode(b64)), height=Pt(33))
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    return
                except Exception:
                    pass
            cell.text = ""
            lines = s.split("\n")
            for i, line in enumerate(lines):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER if align == WD_ALIGN_PARAGRAPH.CENTER else WD_CELL_VERTICAL_ALIGNMENT.TOP

        def set_yesno(cell, mark):
            cell.text = ""
            yes = str(mark or "").strip() == "是"
            no = str(mark or "").strip() == "否"
            p1 = cell.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p1.add_run(("\u2611\ufe0e" if yes else "\u2610") + " 是")
            r1.font.size = Pt(9)
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(("\u2610" if not no else "\u2611\ufe0e") + " 否")
            r2.font.size = Pt(9)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        def add_env_check_grid(grid):
            kind = str(grid[0][1] if len(grid[0]) > 1 else "dev") or "dev"
            code = str(grid[0][2] if len(grid[0]) > 2 else "")
            doc_t = obj.doc_type or "dd_016"
            leaves = self.__env_check_leaves(doc_t, kind)
            groups = ENV_CHECK_GROUPS.get((doc_t, kind), ENV_CHECK_GROUPS[("dd_016", "dev")])
            ncols = len(leaves)
            if doc_t == "dd_017":
                title_txt = "标注共用-%s检查表（%s）" % ("服务器" if kind == "server" else "标注机", code)
            else:
                title_txt = "开发共用-%s检查表（%s）" % ("服务器" if kind == "server" else "开发机", code)
            tb = document.add_table(rows=0, cols=ncols)
            tb.style = "Table Grid"
            tb.alignment = WD_TABLE_ALIGNMENT.CENTER
            trow = tb.add_row().cells
            tmerge = trow[0]
            for i in range(1, ncols):
                tmerge = tmerge.merge(trow[i])
            set_cell(tmerge, title_txt, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            grow = tb.add_row().cells
            lrow = tb.add_row().cells
            ci = 0
            for gl, gleaves in groups:
                if gleaves:
                    gm = grow[ci]
                    for k in range(1, len(gleaves)):
                        gm = gm.merge(grow[ci + k])
                    set_cell(gm, gl, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                    for k, lf in enumerate(gleaves):
                        set_cell(lrow[ci + k], lf, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                    ci += len(gleaves)
                else:
                    vm = grow[ci].merge(lrow[ci])
                    set_cell(vm, gl, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                    ci += 1
            for row in grid[1:]:
                cells = tb.add_row().cells
                j = 0
                for idx, col in enumerate(leaves):
                    t = col["type"]
                    if t == "date":
                        set_cell(cells[idx], str(row[0] if row else "").replace("- ", "-\n"), align=WD_ALIGN_PARAGRAPH.CENTER)
                    elif t == "check":
                        set_yesno(cells[idx], row[j + 1] if j + 1 < len(row) else "")
                        j += 1
                    elif t == "problem":
                        set_cell(cells[idx], row[j + 1] if j + 1 < len(row) else "", align=WD_ALIGN_PARAGRAPH.CENTER)
                    elif t == "checker":
                        set_cell(cells[idx], row[j + 2] if j + 2 < len(row) else "", align=WD_ALIGN_PARAGRAPH.CENTER)
            document.add_paragraph()

        def add_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            if self.__is_env_check_grid(grid):
                add_env_check_grid(grid)
                return
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            if self.__is_review_grid(grid):
                padded, origins = self.__grid_span_origins(grid)
                for _ in padded:
                    table.add_row()
                for r, c, rs, cs in origins:
                    cell = table.cell(r, c)
                    if rs > 1 or cs > 1:
                        cell = cell.merge(table.cell(r + rs - 1, c + cs - 1))
                    set_cell(cell, padded[r][c], bold=(r == 0))
            else:
                for r_idx, row in enumerate(grid):
                    cells = table.add_row().cells
                    for c_idx in range(cols):
                        set_cell(cells[c_idx], row[c_idx] if c_idx < len(row) else "", bold=(r_idx == 0))
            document.add_paragraph()

        def add_cover_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for row in grid:
                cells = table.add_row().cells
                for c_idx in range(cols):
                    text = row[c_idx] if c_idx < len(row) else ""
                    set_cell(cells[c_idx], text, bold=(c_idx % 2 == 0), align=WD_ALIGN_PARAGRAPH.CENTER)
                if (str(row[0]).strip() if row else "") == "生效日期" and cols > 2:
                    merged = cells[1]
                    for c_idx in range(2, cols):
                        merged = merged.merge(cells[c_idx])
                    set_cell(merged, row[1] if len(row) > 1 else "", align=WD_ALIGN_PARAGRAPH.CENTER)
            document.add_paragraph()

        def add_body_heading(title, level):
            size = {1: 16.0, 2: 14.0, 3: 12.0}.get(level, 11.0)
            p = document.add_heading("", level=max(1, min(level, 9)))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, title, font_size=size, bold=True)

        def render_body_section(node, level, number=""):
            name = self.__strip_num(node.get("title"))
            heading = f"{number} {name}".strip() if number else name
            add_body_heading(heading, level=max(1, min(level, 9)))
            if (node.get("body") or "").strip():
                add_text(node.get("body"))
            for table in (node.get("tables") or []):
                add_grid(table)
            idx = 0
            for child in (node.get("children") or []):
                idx += 1
                child_num = f"{number}.{idx}" if number else f"{idx}"
                render_body_section(child, level + 1, child_num)

        cover = next((s for s in sections if s.get("ref_type") == "cover"), None)
        revision = next((s for s in sections if s.get("ref_type") == "revision"), None)
        body = [s for s in sections if s.get("ref_type") not in ("cover", "revision")]
        title = (self.__strip_num(cover.get("title")) if cover else "") or doc_title(obj.doc_type)

        add_blank_lines(6)
        write_center_title(title, size=22.0, bold=True)
        add_blank_lines(4)
        if cover:
            for table in (cover.get("tables") or []):
                add_cover_grid(table)

        document.add_page_break()
        write_center_title("文件修订记录", size=14.0, bold=True)
        add_blank_lines(2)
        if revision:
            for table in (revision.get("tables") or []):
                add_grid(table)

        document.add_page_break()
        write_center_title("目录", size=16.0, bold=True)
        docx_util.insert_toc_field(document)

        document.add_page_break()
        for i, node in enumerate(body):
            render_body_section(node, 1, str(i + 1))

        docx_util.fill_toc_cache(document)
        document.save(output)
        output.seek(0)
        return title, "docx"
