#!/usr/bin/env python
# encoding: utf-8

# 版本更新历史服务层，详见 docs/function_docs/54_版本更新历史.md。
# 整份文档以 content(JSON) 的「目录树」结构存储；导出结构与产品开发计划一致：
# 封面→分页→修订记录→分页→目录→分页→正文。

import base64
import copy
import logging
import os
import re
from io import BytesIO
from typing import List
from sqlalchemy import delete, func, select
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.vuh_doc import VuhDoc
from ..model.version_rule import VersionRule
from ..model.project_timeline import ProjectTimelineRow, ProjectTimelineCell
from ..model.project_member import ProjectMember
from .serv_version_rule import DEFAULT_VERSION_RULE
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_vuh_doc import VuhDocForm
from ..obj.vobj_vuh_doc import VuhDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from . import serv_review_util
from .serv_utils import new_version, sync_file_no_version
from .serv_utils import docx_util

logger = logging.getLogger(__name__)


# 标准模板默认内容（取自《版本更新历史》模板），新增文档时预填、可改。
DEFAULT_VUH_CONTENT = {
    "sections": [
        {
            "title": "版本更新历史", "ref_type": "cover", "body": "", "children": [],
            "tables": [[
                ["编制部门", "产品部", "文件版本", "A0"],
                ["编制人", "", "日期", ""],
                ["审核人", "", "日期", ""],
                ["批准人", "", "日期", ""],
                ["生效日期", "", "", ""],
            ]],
        },
        {
            "title": "文件修订记录", "ref_type": "revision", "body": "", "children": [],
            "tables": [[
                ["修改日期", "版本号", "修订说明", "修订人", "批准人"],
                ["", "", "首次发布", "", ""],
                ["", "", "", "", ""],
                ["", "", "", "", ""],
                ["", "", "", "", ""],
                ["", "", "", "", ""],
            ]],
        },
        {"title": "版本信息", "ref_type": "version_info", "body": "", "tables": [], "children": []},
        {"title": "软件版本命名规则", "body": "", "tables": [], "children": []},
        {
            "title": "软件开发阶段更新历史", "ref_type": "update_history", "body": "", "children": [],
            "tables": [[
                ["完整版本", "发布版本", "类型", "发布日期", "具体更新内容"],
                ["", "", "首次发布", "", "首次发布"],
            ]],
        },
    ]
}


class Server(object):

    def __to_obj(self, row: VuhDoc, product: Product = None):
        obj = VuhDocObj(**row.dict())
        obj.content = self.__normalize_content(obj.content)
        serv_review_util.fill_cover_dates(
            obj.content, serv_review_util.cover_date(row.product_id, "vuh") if row.product_id else ""
        )
        serv_review_util.fill_cover_signers(
            obj.content, serv_review_util.cover_signers(row.product_id, "vuh") if row.product_id else {}
        )
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
        return obj

    @staticmethod
    def __migrate_cover_table(rows):
        rows = [r for r in (rows or []) if isinstance(r, list)]
        if rows and len(rows[0]) >= 4 and str(rows[0][0]).strip() in ("编制部门", "编写部门"):
            return rows
        items = [(str(r[0]).strip(), str(r[1]).strip() if len(r) > 1 else "") for r in rows if r]
        def val(label):
            for l, v in items:
                if l == label:
                    return v
            return ""
        dates = [v for l, v in items if l == "日期"]
        d = lambda i: dates[i] if i < len(dates) else ""
        return [
            ["编制部门", val("编制部门") or val("编写部门") or "产品部", "文件版本", val("文件版本") or "A0"],
            ["编制人", val("编制人"), "日期", d(0)],
            ["审核人", val("审核人"), "日期", d(1)],
            ["批准人", val("批准人"), "日期", d(2)],
            ["生效日期", val("生效日期"), "", ""],
        ]

    def __normalize_node(self, node):
        if not isinstance(node, dict):
            return {"title": str(node or ""), "body": "", "tables": [], "children": []}
        result = dict(node)
        result["title"] = str(result.get("title") or "")
        result["body"] = str(result.get("body") or "")
        tables = result.get("tables")
        if not isinstance(tables, list):
            tables = []
        norm_tables = []
        for table in tables:
            if isinstance(table, list):
                norm_tables.append([[str(c) if c is not None else "" for c in (row or [])] for row in table if isinstance(row, list)])
        if result.get("ref_type") == "cover" and norm_tables:
            norm_tables = [self.__migrate_cover_table(norm_tables[0])] + norm_tables[1:]
        result["tables"] = norm_tables
        children = result.get("children")
        result["children"] = [self.__normalize_node(c) for c in children] if isinstance(children, list) else []
        return result

    def __normalize_content(self, content):
        if not isinstance(content, dict) or not isinstance(content.get("sections"), list):
            return copy.deepcopy(DEFAULT_VUH_CONTENT)
        return {"sections": [self.__normalize_node(s) for s in content["sections"]]}

    def __autofill_for_export(self, content, obj: VuhDocObj):
        sections = (content or {}).get("sections") or []
        prod_id = obj.product_id
        if not prod_id:
            return content
        product = db.session.execute(select(Product).where(Product.id == prod_id)).scalars().first()
        info = self.__collect_autofill(prod_id, product, obj.version)
        for node in sections:
            self.__fill_node(node, info)
        serv_review_util.fill_cover_dates(content, serv_review_util.cover_date(prod_id, "vuh"))
        serv_review_util.fill_cover_signers(content, serv_review_util.cover_signers(prod_id, "vuh"))
        return content

    def __collect_autofill(self, prod_id, product, doc_version):
        prod_name = (getattr(product, "name", "") or "").strip()
        full_version = (getattr(product, "full_version", "") or "").strip()
        release_version = (getattr(product, "release_version", "") or "").strip()

        tl_rows = db.session.execute(
            select(ProjectTimelineRow).where(ProjectTimelineRow.prod_id == prod_id)
        ).scalars().all()
        cell_map = {}
        if tl_rows:
            for c in db.session.execute(
                select(ProjectTimelineCell).where(ProjectTimelineCell.row_id.in_([r.id for r in tl_rows]))
            ).scalars().all():
                cell_map.setdefault(c.row_id, []).append(c.output_result or "")

        def to_int(v):
            digits = re.sub(r"[^\d]", "", str(v or ""))
            return int(digits) if digits else None

        date_rows = [r for r in tl_rows if (r.row_type or "date") == "date" and to_int(r.year) and to_int(r.month)]

        def date_key(r):
            return to_int(r.year) * 10000 + to_int(r.month) * 100 + (to_int(r.day) or 0)

        def file_date_for(keyword, fmt_dot=False):
            rows = [r for r in date_rows if any(keyword in str(v or "") for v in cell_map.get(r.id, []))]
            if not rows:
                return ""
            fr = min(rows, key=date_key)
            if fmt_dot:
                return f"{to_int(fr.year)}.{to_int(fr.month):02d}.{to_int(fr.day) or 1:02d}"
            return f"{to_int(fr.year)}年{to_int(fr.month)}月{to_int(fr.day)}日"

        file_date = file_date_for("版本更新历史")
        release_date = file_date_for("版本更新历史", fmt_dot=True)

        members = db.session.execute(select(ProjectMember).where(ProjectMember.prod_id == prod_id)).scalars().all()
        def find_member(pred):
            for m in members:
                if pred(str(m.role or "")):
                    return (m.name or "").strip()
            return ""
        pm = find_member(lambda r: "产品经理" in r)
        approver = find_member(lambda r: "负责人" in r and "产品" in r)

        vr = db.session.execute(select(VersionRule).where(VersionRule.id == 1)).scalars().first()
        vr_content = vr.content if vr and isinstance(vr.content, dict) else DEFAULT_VERSION_RULE
        naming_body = self.__build_naming_body(vr_content)

        return {
            "prod_name": prod_name, "full_version": full_version, "release_version": release_version,
            "file_date": file_date, "release_date": release_date, "version": doc_version,
            "pm": pm, "approver": approver, "naming_body": naming_body,
        }

    @staticmethod
    def __strip_num(title):
        return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()

    @staticmethod
    def __build_naming_body(c):
        # 由「基础数据-版本命名规则」全局配置生成「软件版本命名规则」章节正文
        c = c if isinstance(c, dict) else {}
        items = c.get("items") or []
        lines = [
            "软件版本命名规则为：",
            f"发布版本：{c.get('release_format', '')}",
            f"完整版本：{c.get('full_format', '')}",
            "软件完整版本及说明：",
        ]
        if str(c.get("note_top") or "").strip():
            lines.append(c["note_top"])
        for it in items:
            if not isinstance(it, dict):
                continue
            title = str(it.get("title") or "").strip()
            desc = str(it.get("desc") or "").strip()
            if title or desc:
                lines.append(f"{title}：{desc}")
        if str(c.get("note_bottom") or "").strip():
            lines.append(c["note_bottom"])
        return "\n".join(lines)

    def __fill_node(self, node, info):
        ref = node.get("ref_type")
        title = self.__strip_num(node.get("title"))
        # 版本信息：自动获取产品名/发布版本/完整版本（始终取最新）
        if ref == "version_info" or title == "版本信息":
            if info["full_version"] or info["release_version"] or info["prod_name"]:
                node["body"] = (
                    f"本次软件为首次注册，软件完整版本为{info['full_version']}，发布版本为{info['release_version']}。\n"
                    f"产品名称：{info['prod_name']}\n"
                    f"发布版本：{info['release_version']}\n"
                    f"完整版本：{info['full_version']}"
                )
        elif title == "软件版本命名规则":
            # 从全局「版本命名规则」配置始终取最新覆盖
            if info.get("naming_body"):
                node["body"] = info["naming_body"]
        # 修订记录首行
        if ref == "revision" or title == "文件修订记录":
            tables = node.get("tables") or []
            if tables and isinstance(tables[0], list):
                t = tables[0]
                cols = len(t[0]) if t and t[0] else 5
                while len(t) < 6:
                    t.append([""] * cols)
                row = t[1]
                def set_if(i, val):
                    if val and not str(row[i] if i < len(row) else "").strip():
                        row[i] = val
                set_if(0, info["file_date"])
                set_if(1, info["version"])
                if not str(row[2] if len(row) > 2 else "").strip():
                    row[2] = "首次发布"
                set_if(3, info["pm"])
                set_if(4, info["approver"])
        # 软件开发阶段更新历史表首行：完整版本/发布版本/首次发布/发布日期/首次发布（仅填空）
        if ref == "update_history" or title == "软件开发阶段更新历史":
            tables = node.get("tables") or []
            if tables and isinstance(tables[0], list) and len(tables[0]) >= 2:
                t = tables[0]
                row = t[1]
                def set_uh(i, val):
                    if val and i < len(row) and not str(row[i]).strip():
                        row[i] = val
                set_uh(0, info["full_version"])
                set_uh(1, info["release_version"])
                if len(row) > 2 and not str(row[2]).strip():
                    row[2] = "首次发布"
                set_uh(3, info["release_date"])
                if len(row) > 4 and not str(row[4]).strip():
                    row[4] = "首次发布"
        for child in (node.get("children") or []):
            self.__fill_node(child, info)

    async def add_vuh_doc(self, form: VuhDocForm):
        try:
            sql = select(func.count(VuhDoc.id)).where(
                VuhDoc.product_id == form.product_id,
                VuhDoc.version == form.version,
            )
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = VuhDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=VuhDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_vuh_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: VuhDoc = db.session.execute(select(VuhDoc).where(VuhDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(VuhDoc.version).where(VuhDoc.product_id == target_pid)).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = VuhDoc(
                product_id=target_pid,
                version=version,
                file_no=sync_file_no_version(fromdoc.file_no, version),
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            return Resp.resp_ok(data=VuhDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_vuh_doc(self, form: VuhDocForm):
        try:
            row: VuhDoc = db.session.execute(select(VuhDoc).where(VuhDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_vuh_doc(self, id: int):
        db.session.execute(delete(VuhDoc).where(VuhDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_vuh_doc(self, id: int):
        sql = select(VuhDoc, Product).join(Product, VuhDoc.product_id == Product.id).where(VuhDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_vuh_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(VuhDoc.product_id == product_id)
        if version:
            wheres.append(VuhDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(VuhDoc.id)).join(Product, VuhDoc.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (
            select(VuhDoc, Product)
            .join(Product, VuhDoc.product_id == Product.id)
            .where(*wheres)
            .order_by(VuhDoc.id.desc())
            .offset(page_index * page_size)
            .limit(page_size)
        )
        rows: List[VuhDocObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    # ---------------- 导出 Word ----------------
    async def export_vuh_doc(self, output, id: int):
        resp = await self.get_vuh_doc(id)
        obj: VuhDocObj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        c = self.__autofill_for_export(self.__normalize_content(obj.content), obj)
        sections = c.get("sections") or []
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        document.settings.element.append(update_fields)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")
        docx_util.add_page_number_footer(section, obj.file_no or "")

        def write_center_title(text, size=22.0, bold=False):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.right_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, text, font_size=size, bold=bold)

        def add_blank_lines(count):
            for _ in range(max(0, int(count or 0))):
                document.add_paragraph("")

        def add_text(text):
            docx_util.save_txt2docx(str(text or ""), document)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            s = str(text or "")
            if s.startswith("data:image"):
                try:
                    b64 = s.split(",", 1)[1] if "," in s else ""
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = align
                    para.add_run().add_picture(BytesIO(base64.b64decode(b64)), height=Pt(33))
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    return
                except Exception:
                    pass
            cell.text = ""
            lines = s.split("\n")
            for i, line in enumerate(lines):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER if align == WD_ALIGN_PARAGRAPH.CENTER else WD_CELL_VERTICAL_ALIGNMENT.TOP

        def add_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for r_idx, row in enumerate(grid):
                cells = table.add_row().cells
                for c_idx in range(cols):
                    set_cell(cells[c_idx], row[c_idx] if c_idx < len(row) else "", bold=(r_idx == 0))
            document.add_paragraph()

        def add_cover_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for row in grid:
                cells = table.add_row().cells
                for c_idx in range(cols):
                    text = row[c_idx] if c_idx < len(row) else ""
                    set_cell(cells[c_idx], text, bold=(c_idx % 2 == 0), align=WD_ALIGN_PARAGRAPH.CENTER)
                if (str(row[0]).strip() if row else "") == "生效日期" and cols > 2:
                    merged = cells[1]
                    for c_idx in range(2, cols):
                        merged = merged.merge(cells[c_idx])
                    set_cell(merged, row[1] if len(row) > 1 else "", align=WD_ALIGN_PARAGRAPH.CENTER)
            document.add_paragraph()

        def add_body_heading(title, level):
            size = {1: 16.0, 2: 14.0, 3: 12.0}.get(level, 11.0)
            p = document.add_heading("", level=max(1, min(level, 9)))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, title, font_size=size, bold=True)

        def add_naming_diagram():
            # 嵌入版本命名规则示意图（与模板一致的箭头图）；缺图时回落为对应关系表格
            img_path = os.path.join(os.path.dirname(__file__), "..", "..", "src-res", "assets", "version_naming_rule.png")
            if os.path.exists(img_path):
                try:
                    p = document.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(img_path, width=Inches(5.3))
                    document.add_paragraph()
                    return
                except Exception:
                    logger.exception("add_naming_diagram_picture_failed")
            add_grid([
                ["软件完整版本", "V X . Y . Z . B"],
                ["主版本号", "X"],
                ["次版本号", "Y"],
                ["修订版本号", "Z"],
                ["上市后软件升级次数号", "B"],
            ])

        def render_body_section(node, level, number=""):
            name = self.__strip_num(node.get("title"))
            heading = f"{number} {name}".strip() if number else name
            add_body_heading(heading, level=max(1, min(level, 9)))
            body = node.get("body") or ""
            if name == "软件版本命名规则":
                # 示意图紧跟在「软件完整版本及说明：」行之后，再接剩余正文
                lines = body.split("\n")
                cut = next((i for i, ln in enumerate(lines) if ln.strip().startswith("软件完整版本及说明")), -1)
                if cut >= 0:
                    before = "\n".join(lines[:cut + 1])
                    after = "\n".join(lines[cut + 1:])
                    if before.strip():
                        add_text(before)
                    add_naming_diagram()
                    if after.strip():
                        add_text(after)
                else:
                    if body.strip():
                        add_text(body)
                    add_naming_diagram()
            elif body.strip():
                add_text(body)
            for table in (node.get("tables") or []):
                add_grid(table)
            idx = 0
            for child in (node.get("children") or []):
                idx += 1
                child_num = f"{number}.{idx}" if number else f"{idx}"
                render_body_section(child, level + 1, child_num)

        cover = next((s for s in sections if s.get("ref_type") == "cover"), None)
        revision = next((s for s in sections if s.get("ref_type") == "revision"), None)
        body = [s for s in sections if s.get("ref_type") not in ("cover", "revision")]

        add_blank_lines(6)
        write_center_title((self.__strip_num(cover.get("title")) if cover else "") or "版本更新历史", size=22.0, bold=True)
        add_blank_lines(4)
        if cover:
            for table in (cover.get("tables") or []):
                add_cover_grid(table)

        document.add_page_break()
        write_center_title("文件修订记录", size=14.0, bold=True)
        add_blank_lines(2)
        if revision:
            for table in (revision.get("tables") or []):
                add_grid(table)

        document.add_page_break()
        write_center_title("目录", size=16.0, bold=True)
        docx_util.insert_toc_field(document)

        document.add_page_break()
        for i, node in enumerate(body):
            render_body_section(node, 1, str(i + 1))

        document.save(output)
        output.seek(0)
