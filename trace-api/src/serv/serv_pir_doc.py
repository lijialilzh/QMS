#!/usr/bin/env python
# encoding: utf-8

# 产品立项报告服务层，详见 docs/function_docs/53_产品立项报告.md。
# 整份文档以 content(JSON) 的「目录树」结构存储；导出复用 docx_util 生成 Word，
# 结构与产品开发计划一致：封面→分页→修订记录→分页→目录→分页→正文。

import copy
import logging
import re
from typing import List
from sqlalchemy import delete, func, select
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.project import Project
from ..model.pir_doc import PirDoc
from ..model.project_timeline import ProjectTimelineRow, ProjectTimelineCell
from ..model.project_member import ProjectMember
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_pir_doc import PirDocForm
from ..obj.vobj_pir_doc import PirDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from .serv_utils import new_version, sync_file_no_version
from .serv_utils import docx_util

logger = logging.getLogger(__name__)


# 标准模板默认内容（取自《产品立项报告》模板），新增文档时预填、可改。
# 采用「目录树」结构：content.sections 为可递归章节树，节点 {title, body, tables, children, ref_type}。
# ref_type：cover=封面（不编号）、revision=修订记录（不编号）、basic_info=项目概述基本信息表（自动填充）。

DEFAULT_PIR_CONTENT = {
    "sections": [
        {
            "title": "产品立项报告", "ref_type": "cover", "body": "", "children": [],
            "tables": [[
                ["编写部门", "产品部", "文件版本", "A0"],
                ["编制人", "", "日期", ""],
                ["审核人", "", "日期", ""],
                ["批准人", "", "日期", ""],
                ["生效日期", "", "", ""],
            ]],
        },
        {
            "title": "文件修订记录", "ref_type": "revision", "body": "", "children": [],
            "tables": [[
                ["修改日期", "版本号", "修订说明", "修订人", "批准人"],
                ["", "", "首次发布", "", ""],
                ["", "", "", "", ""],
                ["", "", "", "", ""],
                ["", "", "", "", ""],
                ["", "", "", "", ""],
            ]],
        },
        {
            "title": "项目概述", "ref_type": "basic_info", "body": "", "children": [],
            "tables": [[
                ["基本信息", "描述"],
                ["项目", ""],
                ["产品名称", ""],
                ["软件版本", ""],
                ["产品标识", ""],
                ["预期用途", ""],
                ["项目周期", ""],
            ]],
        },
        {
            "title": "立项目的", "body": "", "tables": [], "children": [
                {"title": "产品的背景及意义", "body": "", "tables": [], "children": []},
                {"title": "国内外研究现状分析及存在问题", "body": "", "tables": [], "children": []},
                {"title": "应用前景", "body": "", "tables": [], "children": []},
            ],
        },
        {
            "title": "产品技术可行性分析", "body": "", "tables": [], "children": [
                {"title": "产品描述", "ref_type": "prod_overview", "body": "", "tables": [], "children": []},
                {"title": "产品系统结构", "body": "", "tables": [], "children": []},
                {"title": "核心技术", "body": "", "tables": [], "children": []},
            ],
        },
        {
            "title": "经济可行性", "body": "", "tables": [], "children": [
                {"title": "投资", "body": "", "tables": [], "children": []},
                {"title": "市场预测", "body": "", "tables": [], "children": []},
            ],
        },
        {"title": "法律可行性", "body": "", "tables": [], "children": []},
        {"title": "产品可提供成果", "body": "", "tables": [], "children": []},
    ]
}


# 默认正文（写死、可改）：仅在该章节正文为空时填入，按章节标题匹配
FIXED_PIR_BODY = {
    "产品系统结构": "本软件应包含4个模块：图像接收、图像处理、图像预测和图像显示。",
    "核心技术": "1）高性能和稳定的数据接收和传输机制\n2）多租户的用户管理系统",
    "投资": "基础投资包括房屋和办公设施、计算机软硬件设备、安全与保密设备以及直接投资，包括开发及维护期的直接工资成本，质量管理体系建立费用，注册检测费用。",
    "市场预测": "产品上市后将在美国医疗市场推广使用，包括医院、实验室、医疗单位、卫生院、健康体检中心等，市场空间巨大。",
    "法律可行性": "符合现行的法律或行政机构规定的文件",
}


class Server(object):

    def __to_obj(self, row: PirDoc, product: Product = None):
        obj = PirDocObj(**row.dict())
        obj.content = self.__normalize_content(obj.content)
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
        return obj

    @staticmethod
    def __migrate_cover_table(rows):
        """把旧版 2 列封面表（项目/内容）迁移为新版 4 列结构，保留已填值。"""
        rows = [r for r in (rows or []) if isinstance(r, list)]
        if rows and len(rows[0]) >= 4 and str(rows[0][0]).strip() == "编写部门":
            return rows
        items = [(str(r[0]).strip(), str(r[1]).strip() if len(r) > 1 else "") for r in rows if r]
        def val(label):
            for l, v in items:
                if l == label:
                    return v
            return ""
        dates = [v for l, v in items if l == "日期"]
        d = lambda i: dates[i] if i < len(dates) else ""
        return [
            ["编写部门", val("编写部门") or "产品部", "文件版本", val("文件版本") or "A0"],
            ["编制人", val("编制人"), "日期", d(0)],
            ["审核人", val("审核人"), "日期", d(1)],
            ["批准人", val("批准人"), "日期", d(2)],
            ["生效日期", val("生效日期"), "", ""],
        ]

    def __normalize_node(self, node):
        if not isinstance(node, dict):
            return {"title": str(node or ""), "body": "", "tables": [], "children": []}
        result = dict(node)
        result["title"] = str(result.get("title") or "")
        result["body"] = str(result.get("body") or "")
        tables = result.get("tables")
        if not isinstance(tables, list):
            tables = []
        norm_tables = []
        for table in tables:
            if isinstance(table, list):
                norm_tables.append([[str(c) if c is not None else "" for c in (row or [])] for row in table if isinstance(row, list)])
        if result.get("ref_type") == "cover" and norm_tables:
            norm_tables = [self.__migrate_cover_table(norm_tables[0])] + norm_tables[1:]
        result["tables"] = norm_tables
        children = result.get("children")
        result["children"] = [self.__normalize_node(c) for c in children] if isinstance(children, list) else []
        return result

    def __normalize_content(self, content):
        if not isinstance(content, dict) or not isinstance(content.get("sections"), list):
            return copy.deepcopy(DEFAULT_PIR_CONTENT)
        return {"sections": [self.__normalize_node(s) for s in content["sections"]]}

    def __autofill_for_export(self, content, obj: PirDocObj):
        """导出时按产品/时间线/参与人员实时填充，与「新增/编辑」一致（仅填空、不覆盖已填）。"""
        sections = (content or {}).get("sections") or []
        prod_id = obj.product_id
        if not prod_id:
            return content
        product = db.session.execute(select(Product).where(Product.id == prod_id)).scalars().first()
        info = self.__collect_autofill(prod_id, product, obj.version)
        for node in sections:
            self.__fill_node(node, info)
        return content

    def __collect_autofill(self, prod_id, product, doc_version):
        prod_name = (getattr(product, "name", "") or "").strip()
        full_version = (getattr(product, "full_version", "") or "").strip()
        product_code = (getattr(product, "product_code", "") or "").strip()
        scope = (getattr(product, "scope", "") or "").strip()
        overall_desc = (getattr(product, "overall_desc", "") or "").strip()

        # 国家认证：从产品绑定的产品线(project.country)获取；项目 = 产品名称 + 国家认证 + 注册
        country = ""
        if product is not None and getattr(product, "project_id", None):
            proj = db.session.execute(select(Project).where(Project.id == product.project_id)).scalars().first()
            country = (getattr(proj, "country", "") or "").strip() if proj else ""
        project_field = (" ".join([p for p in [prod_name, country] if p]) + "注册") if prod_name else ""

        tl_rows = db.session.execute(
            select(ProjectTimelineRow).where(ProjectTimelineRow.prod_id == prod_id)
        ).scalars().all()
        cell_map = {}
        if tl_rows:
            for c in db.session.execute(
                select(ProjectTimelineCell).where(ProjectTimelineCell.row_id.in_([r.id for r in tl_rows]))
            ).scalars().all():
                cell_map.setdefault(c.row_id, []).append(c.output_result or "")

        def to_int(v):
            digits = re.sub(r"[^\d]", "", str(v or ""))
            return int(digits) if digits else None

        date_rows = [r for r in tl_rows if (r.row_type or "date") == "date" and to_int(r.year) and to_int(r.month)]

        def date_key(r):
            return to_int(r.year) * 10000 + to_int(r.month) * 100 + (to_int(r.day) or 0)

        period = ""
        if date_rows:
            start = min(date_rows, key=date_key)
            out_rows = [r for r in date_rows if any(str(v or "").strip() for v in cell_map.get(r.id, []))]
            pool = out_rows or date_rows
            end = max(pool, key=date_key)
            period = f"{to_int(start.year)}.{to_int(start.month):02d}-{to_int(end.year)}.{to_int(end.month):02d}"

        file_rows = [r for r in date_rows if any("立项报告" in str(v or "") for v in cell_map.get(r.id, []))]
        file_date = ""
        if file_rows:
            fr = min(file_rows, key=date_key)
            file_date = f"{to_int(fr.year)}年{to_int(fr.month)}月{to_int(fr.day)}日"

        members = db.session.execute(select(ProjectMember).where(ProjectMember.prod_id == prod_id)).scalars().all()
        def find_member(pred):
            for m in members:
                if pred(str(m.role or "")):
                    return (m.name or "").strip()
            return ""
        pm = find_member(lambda r: "产品经理" in r)
        approver = find_member(lambda r: "负责人" in r and "产品" in r)

        return {
            "prod_name": prod_name, "full_version": full_version, "product_code": product_code,
            "scope": scope, "overall_desc": overall_desc, "period": period,
            "file_date": file_date, "version": doc_version, "pm": pm, "approver": approver,
            "project_field": project_field, "country": country,
        }

    @staticmethod
    def __strip_num(title):
        return re.sub(r"^\s*\d+(?:\.\d+)*[\.、\s]*", "", str(title or "")).strip()

    def __fill_node(self, node, info):
        ref = node.get("ref_type")
        title = self.__strip_num(node.get("title"))
        # 章节默认/自动获取正文
        cur = str(node.get("body") or "").strip()
        if ref == "prod_overview" or title == "产品描述":
            # 产品描述：自动获取总体描述，始终取最新（源为空时保留原值）
            if info["overall_desc"]:
                node["body"] = info["overall_desc"]
        elif title == "产品可提供成果":
            # 自动获取句：产品名 + 国家认证；只要是这句模板就按数据源刷新，自定义内容保留
            fresh = f"本项目目标是申请并获得{info['prod_name'] or 'InferCare'}产品的 {info['country'] or 'FDA'}注册认可"
            if not cur or re.match(r"^本项目目标是申请并获得.*产品的\s*\S*注册认可$", cur):
                node["body"] = fresh
        elif not cur and title in FIXED_PIR_BODY:
            # 写死默认（可改，仅填空不覆盖）
            node["body"] = FIXED_PIR_BODY[title]
        # 修订记录首行
        if ref == "revision" or title == "文件修订记录":
            tables = node.get("tables") or []
            if tables and isinstance(tables[0], list):
                t = tables[0]
                cols = len(t[0]) if t and t[0] else 5
                while len(t) < 6:
                    t.append([""] * cols)
                row = t[1]
                def set_if(i, val):
                    if val and not str(row[i] if i < len(row) else "").strip():
                        row[i] = val
                set_if(0, info["file_date"])
                set_if(1, info["version"])
                if not str(row[2] if len(row) > 2 else "").strip():
                    row[2] = "首次发布"
                set_if(3, info["pm"])
                set_if(4, info["approver"])
        # 项目概述-基本信息表
        if ref == "basic_info" or title == "项目概述":
            label_map = {
                "项目": info["project_field"],
                "产品名称": info["prod_name"],
                "软件版本": info["full_version"],
                "产品标识": info["product_code"],
                "预期用途": info["scope"],
                "项目周期": info["period"],
            }
            for table in (node.get("tables") or []):
                for row in table:
                    if not isinstance(row, list) or len(row) < 2:
                        continue
                    key = str(row[0]).strip()
                    # 自动获取字段：始终取最新数据源覆盖（源为空时保留原值）
                    if key in label_map and label_map[key]:
                        row[1] = label_map[key]
        for child in (node.get("children") or []):
            self.__fill_node(child, info)

    async def add_pir_doc(self, form: PirDocForm):
        try:
            sql = select(func.count(PirDoc.id)).where(
                PirDoc.product_id == form.product_id,
                PirDoc.version == form.version,
            )
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = PirDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=PirDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_pir_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: PirDoc = db.session.execute(select(PirDoc).where(PirDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(PirDoc.version).where(PirDoc.product_id == target_pid)).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = PirDoc(
                product_id=target_pid,
                version=version,
                file_no=sync_file_no_version(fromdoc.file_no, version),
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            return Resp.resp_ok(data=PirDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_pir_doc(self, form: PirDocForm):
        try:
            row: PirDoc = db.session.execute(select(PirDoc).where(PirDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_pir_doc(self, id: int):
        db.session.execute(delete(PirDoc).where(PirDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_pir_doc(self, id: int):
        sql = select(PirDoc, Product).join(Product, PirDoc.product_id == Product.id).where(PirDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_pir_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(PirDoc.product_id == product_id)
        if version:
            wheres.append(PirDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(PirDoc.id)).join(Product, PirDoc.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (
            select(PirDoc, Product)
            .join(Product, PirDoc.product_id == Product.id)
            .where(*wheres)
            .order_by(PirDoc.id.desc())
            .offset(page_index * page_size)
            .limit(page_size)
        )
        rows: List[PirDocObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    # ---------------- 导出 Word ----------------
    async def export_pir_doc(self, output, id: int):
        resp = await self.get_pir_doc(id)
        obj: PirDocObj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        c = self.__autofill_for_export(self.__normalize_content(obj.content), obj)
        sections = c.get("sections") or []
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        document.settings.element.append(update_fields)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")

        def write_center_title(text, size=22.0, bold=False):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.right_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, text, font_size=size, bold=bold)

        def add_blank_lines(count):
            for _ in range(max(0, int(count or 0))):
                document.add_paragraph("")

        def add_text(text):
            docx_util.save_txt2docx(str(text or ""), document)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            cell.text = ""
            lines = str(text or "").split("\n")
            for i, line in enumerate(lines):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER if align == WD_ALIGN_PARAGRAPH.CENTER else WD_CELL_VERTICAL_ALIGNMENT.TOP

        def add_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for r_idx, row in enumerate(grid):
                cells = table.add_row().cells
                for c_idx in range(cols):
                    set_cell(cells[c_idx], row[c_idx] if c_idx < len(row) else "", bold=(r_idx == 0))
            document.add_paragraph()

        def add_cover_grid(grid):
            grid = [row for row in (grid or []) if isinstance(row, list)]
            cols = max((len(row) for row in grid), default=0)
            if cols <= 0:
                return
            table = document.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for row in grid:
                cells = table.add_row().cells
                for c_idx in range(cols):
                    text = row[c_idx] if c_idx < len(row) else ""
                    set_cell(cells[c_idx], text, bold=(c_idx % 2 == 0), align=WD_ALIGN_PARAGRAPH.CENTER)
                if (str(row[0]).strip() if row else "") == "生效日期" and cols > 2:
                    merged = cells[1]
                    for c_idx in range(2, cols):
                        merged = merged.merge(cells[c_idx])
                    set_cell(merged, row[1] if len(row) > 1 else "", align=WD_ALIGN_PARAGRAPH.CENTER)
            document.add_paragraph()

        def add_body_heading(title, level):
            size = {1: 16.0, 2: 14.0, 3: 12.0}.get(level, 11.0)
            p = document.add_heading("", level=max(1, min(level, 9)))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, title, font_size=size, bold=True)

        def render_body_section(node, level, number=""):
            name = self.__strip_num(node.get("title"))
            heading = f"{number} {name}".strip() if number else name
            add_body_heading(heading, level=max(1, min(level, 9)))
            if (node.get("body") or "").strip():
                add_text(node.get("body"))
            for table in (node.get("tables") or []):
                add_grid(table)
            idx = 0
            for child in (node.get("children") or []):
                idx += 1
                child_num = f"{number}.{idx}" if number else f"{idx}"
                render_body_section(child, level + 1, child_num)

        cover = next((s for s in sections if s.get("ref_type") == "cover"), None)
        revision = next((s for s in sections if s.get("ref_type") == "revision"), None)
        body = [s for s in sections if s.get("ref_type") not in ("cover", "revision")]

        add_blank_lines(6)
        write_center_title((self.__strip_num(cover.get("title")) if cover else "") or "产品立项报告", size=22.0, bold=True)
        add_blank_lines(4)
        if cover:
            for table in (cover.get("tables") or []):
                add_cover_grid(table)

        document.add_page_break()
        write_center_title("文件修订记录", size=14.0, bold=True)
        add_blank_lines(2)
        if revision:
            for table in (revision.get("tables") or []):
                add_grid(table)

        document.add_page_break()
        write_center_title("目录", size=16.0, bold=True)
        docx_util.insert_toc_field(document)

        document.add_page_break()
        for i, node in enumerate(body):
            render_body_section(node, 1, str(i + 1))

        document.save(output)
        output.seek(0)
