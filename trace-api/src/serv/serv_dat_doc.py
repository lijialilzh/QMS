#!/usr/bin/env python
# encoding: utf-8

# 数据申请单服务层（开发文件）。单张表单，字段化存储于 content(JSON)。

import base64
import copy
import logging
import re
from io import BytesIO
from typing import List

from sqlalchemy import delete, func, select
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.dat_doc import DatDoc
from ..model.prod_dhf import ProdDhf
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_dat_doc import DatDocForm
from ..obj.vobj_dat_doc import DatDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from . import serv_review_util
from .serv_utils import new_version, sync_file_no_version
from .serv_utils import docx_util

logger = logging.getLogger(__name__)

DOC_NAME = "数据申请单"

DEFAULT_DAT_CONTENT = {
    "project": "",
    "provider": "数据部",
    "quantity": "100",
    "apply_dept": "产品开发部",
    "applicant": "宋月",
    "apply_date": "",
    "check_way": "胸部、腹部（肝胆、泌尿）",
    "data_source": "国内数据",
    "other": "无",
    "deliver_date": "",
    "reason": "",
    "sign_img": "",
    "sign_date": "",
    "approve": "",
}


class Server(object):

    def __normalize_content(self, content):
        base = copy.deepcopy(DEFAULT_DAT_CONTENT)
        if isinstance(content, dict):
            for k in base.keys():
                if content.get(k) is not None:
                    base[k] = content.get(k)
        return base

    def __dhf_file_no(self, prod_id):
        if not prod_id:
            return ""
        row = db.session.execute(
            select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name == DOC_NAME).order_by(ProdDhf.id.asc())
        ).scalars().first()
        if not row:
            row = db.session.execute(
                select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name.like(f"%{DOC_NAME}%")).order_by(ProdDhf.id.asc())
            ).scalars().first()
        return (row.code or "").strip() if row and row.code else ""

    def __autofill(self, content, prod_id, product=None, force=False):
        """产品信息自动获取：project/reason←产品名+版本；日期←时间线(数据申请)；签字←申请人签名章。
        force=False（默认）：仅填空，不覆盖已填。
        force=True（切换产品）：强制覆盖，无数据则置空。"""
        if not isinstance(content, dict):
            return content
        name = (product.name or "").strip() if product else ""
        ver = (product.full_version or "").strip() if product else ""
        project_val = f"{name}（{ver}）" if ver else name
        reason_val = f"“{name} {ver}”项目测试时应有的数据支持。" if name else ""
        rev_date = serv_review_util.review_date(prod_id, ["数据申请单", "数据申请"]) if prod_id else ""
        applicant = str(content.get("applicant") or "").strip()
        sign_val = (serv_review_util._sign_by_name(applicant) if applicant else "") or ""
        if force:
            content["project"] = project_val
            content["reason"] = reason_val
            for k in ("apply_date", "deliver_date", "sign_date"):
                content[k] = rev_date or ""
            content["sign_img"] = sign_val
        else:
            if name and not str(content.get("project") or "").strip():
                content["project"] = project_val
            if name and not str(content.get("reason") or "").strip():
                content["reason"] = reason_val
            for k in ("apply_date", "deliver_date", "sign_date"):
                if rev_date and not str(content.get(k) or "").strip():
                    content[k] = rev_date
            if applicant and not str(content.get("sign_img") or "").startswith("data:image"):
                content["sign_img"] = sign_val
        return content

    def __to_obj(self, row: DatDoc, product: Product = None):
        obj = DatDocObj(**row.dict())
        obj.content = self.__autofill(self.__normalize_content(obj.content), row.product_id, product)
        if not (obj.file_no or "").strip():
            obj.file_no = self.__dhf_file_no(row.product_id)
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
        return obj

    async def add_dat_doc(self, form: DatDocForm):
        try:
            sql = select(func.count(DatDoc.id)).where(DatDoc.product_id == form.product_id, DatDoc.version == form.version)
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = DatDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=DatDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_dat_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: DatDoc = db.session.execute(select(DatDoc).where(DatDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(DatDoc.version).where(DatDoc.product_id == target_pid)).scalars().all()
            existing_set = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing_set:
                version = new_version(version)
            newdoc = DatDoc(
                product_id=target_pid, version=version,
                file_no=sync_file_no_version((fromdoc.file_no or "").strip() or self.__dhf_file_no(target_pid), version) or None,
                change_log=fromdoc.change_log,
                content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            if target_pid != fromdoc.product_id:
                await self.rebind_product(newdoc.id, target_pid)
            return Resp.resp_ok(data=DatDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_dat_doc(self, form: DatDocForm):
        try:
            row: DatDoc = db.session.execute(select(DatDoc).where(DatDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def rebind_product(self, id: int, product_id: int):
        """切换产品：更新 product_id 并强制用新产品信息重新获取后保存，返回新 obj。"""
        try:
            row: DatDoc = db.session.execute(select(DatDoc).where(DatDoc.id == id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            product: Product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
            if not product:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            db.session.execute(delete(DatDoc).where(DatDoc.product_id == product_id, DatDoc.version == row.version, DatDoc.id != id))
            row.product_id = product_id
            content = self.__normalize_content(row.content)
            content = self.__autofill(content, product_id, product, force=True)
            row.content = content
            db.session.commit()
            return Resp.resp_ok(data=self.__to_obj(row, product))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_dat_doc(self, id: int):
        db.session.execute(delete(DatDoc).where(DatDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_dat_doc(self, id: int):
        sql = select(DatDoc, Product).join(Product, DatDoc.product_id == Product.id).where(DatDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=self.__to_obj(doc, product))

    async def list_dat_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(DatDoc.product_id == product_id)
        if version:
            wheres.append(DatDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        sql_total = select(func.count(DatDoc.id)).join(Product, DatDoc.product_id == Product.id).where(*wheres)
        total = db.session.execute(sql_total).scalar() or 0
        sql = (select(DatDoc, Product).join(Product, DatDoc.product_id == Product.id).where(*wheres)
               .order_by(DatDoc.id.desc()).offset(page_index * page_size).limit(page_size))
        rows: List[DatDocObj] = [self.__to_obj(doc, product) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    # ---------------- 导出 Word ----------------
    async def export_dat_doc(self, output, id: int):
        resp = await self.get_dat_doc(id)
        obj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        c = self.__normalize_content(obj.content)
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")
        docx_util.add_page_number_footer(section, obj.file_no or "", skip_first=False)

        def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            s = str(text or "")
            if s.startswith("data:image"):
                try:
                    b64 = s.split(",", 1)[1] if "," in s else ""
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = align
                    para.add_run().add_picture(BytesIO(base64.b64decode(b64)), height=Pt(30))
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    return
                except Exception:
                    pass
            cell.text = ""
            for i, line in enumerate(s.split("\n")):
                para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                para.alignment = align
                para.paragraph_format.line_spacing = 1.3
                docx_util.fonted_txt(para, line, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        docx_util.fonted_txt(title, DOC_NAME, font_size=18.0, bold=True)
        document.add_paragraph()

        tb = document.add_table(rows=10, cols=4)
        tb.style = "Table Grid"
        tb.alignment = WD_TABLE_ALIGNMENT.CENTER
        rows = tb.rows

        def lbl(cell, text):
            set_cell(cell, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        def merge_val(row_idx, start, text, align=WD_ALIGN_PARAGRAPH.LEFT):
            cells = rows[row_idx].cells
            m = cells[start]
            for i in range(start + 1, 4):
                m = m.merge(cells[i])
            set_cell(m, text, align=align)
            return m

        # 项目名称
        lbl(rows[0].cells[0], "项目名称")
        merge_val(0, 1, c.get("project", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
        # 提供部门 / 数据数量
        lbl(rows[1].cells[0], "提供部门")
        set_cell(rows[1].cells[1], c.get("provider", ""))
        lbl(rows[1].cells[2], "数据数量")
        set_cell(rows[1].cells[3], c.get("quantity", ""))
        # 申请部门 / 申请人/日期
        lbl(rows[2].cells[0], "申请部门")
        set_cell(rows[2].cells[1], c.get("apply_dept", ""))
        lbl(rows[2].cells[2], "申请人/日期")
        set_cell(rows[2].cells[3], "/".join([x for x in [c.get("applicant", ""), c.get("apply_date", "")] if x]))
        # 数据需求（纵向合并3行）
        lbl(rows[3].cells[0], "数据需求")
        lbl(rows[4].cells[0], "数据需求")
        lbl(rows[5].cells[0], "数据需求")
        dm = rows[3].cells[0].merge(rows[4].cells[0]).merge(rows[5].cells[0])
        set_cell(dm, "数据需求", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        lbl(rows[3].cells[1], "检查方式")
        merge_val(3, 2, c.get("check_way", ""))
        lbl(rows[4].cells[1], "数据来源")
        merge_val(4, 2, c.get("data_source", ""))
        lbl(rows[5].cells[1], "其他需求")
        merge_val(5, 2, c.get("other", ""))
        # 交付日期
        lbl(rows[6].cells[0], "交付日期")
        merge_val(6, 1, c.get("deliver_date", ""))
        # 申请原因
        lbl(rows[7].cells[0], "申请原因")
        merge_val(7, 1, c.get("reason", ""))
        # 申请人签字/日期
        lbl(rows[8].cells[0], "申请人签字/日期")
        cells8 = rows[8].cells
        m8 = cells8[1].merge(cells8[2]).merge(cells8[3])
        m8.text = ""
        p8 = m8.paragraphs[0]
        p8.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sign = str(c.get("sign_img") or "")
        if sign.startswith("data:image"):
            try:
                p8.add_run().add_picture(BytesIO(base64.b64decode(sign.split(",", 1)[1])), height=Pt(30))
            except Exception:
                pass
        elif c.get("applicant"):
            docx_util.fonted_txt(p8, str(c.get("applicant")), font_size=10.5)
        if c.get("sign_date"):
            r = p8.add_run("    " + str(c.get("sign_date")))
            r.font.size = Pt(10.5)
        m8.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        rows[8].height = Pt(44)
        # 批准意见（整行合并）
        am = rows[9].cells[0]
        for i in range(1, 4):
            am = am.merge(rows[9].cells[i])
        set_cell(am, c.get("approve", "") or "批准意见：                                     签字/日期：")
        rows[9].height = Pt(50)

        document.save(output)
        output.seek(0)
