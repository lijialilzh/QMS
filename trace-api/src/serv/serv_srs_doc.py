import logging
import base64
import io
import json
import hashlib
import os
import re
import sys
import zipfile
from copy import deepcopy
from datetime import datetime
from enum import Enum
from typing import Dict, List, Tuple
from sqlalchemy import select, delete, func
from sqlalchemy.dialects.postgresql import insert as pg_insert
from sqlalchemy.sql import desc
try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph
    from docx.shared import Inches, Pt
    from docx import enum as dox_enum
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
except Exception:
    Document = None
    OxmlElement = None
    DocxTable = None
    Paragraph = None
    Inches = None
    Pt = None
    dox_enum = None
    qn = None
    RGBColor = None
from openpyxl import load_workbook
from openpyxl.styles import Alignment
from ..obj.vobj_user import UserObj
from ..model.srs_type import SrsType
from ..model.sds_trace import SdsTrace
from ..obj.vobj_srs_reqd import SrsReqdObj
from ..model.doc_file import DocFile
from ..model.sds_doc import SdsDoc, SdsNode
from ..model.sds_reqd import SdsReqd, Logic
from ..model.test_set import TestSet
from ..model.test_case import TestCase
from ..model.rcm import Rcm
from ..obj.tobj_srs_doc import Table, TabHeader, TableCell
from ..model.product import Product, UserProd
from ..model.srs_req import ReqRcm, SrsReq
from ..model.srs_reqd import SrsReqd
from ..obj.vobj_srs_doc import SrsDocObj
from ..obj.vobj_sds_doc import CompareObj
from ..model.srs_doc import SrsDoc, SrsNode
from ..obj.tobj_srs_doc import SrsDocForm, SrsNodeForm
from ..utils.sql_ctx import db
from ..utils.i18n import ts
from ..utils import get_uuid
from ..obj import Page, Resp
from .serv_srs_req import Server as ServSrsReq
from .serv_srs_reqd import Server as ServSrsReqd
from .serv_sds_trace import NAME_DICT
from .serv_utils import new_version
from .serv_utils.tree_util import find_parent, iter_tree
from .serv_doc_file import build_doc_image_file_name, pick_doc_image_file_row, sanitize_doc_image_token
from . import msg_err_db, save_file

logger = logging.getLogger(__name__)
srsreq_serv = ServSrsReq()
srsreqd_serv = ServSrsReqd()
DELETED_SRS_VERSION_PREFIX = "__deleted_srs__"

DEF_SRS = [
    ("SRS-XUS00-001", "数据库要求"),
    ("SRS-XUS00-002", "性能要求"),
    ("SRS-XUS00-003", "基本要求"),
    ("SRS-XUS00-004", "图像接收"),
    ("SRS-XUS00-005", "图像存储"),
    ("SRS-XUS00-006", "图像预测"),
    ("SRS-XUS00-007", "图像显示"),
    ("SRS-XUS00-008", "文档需求"),
    ("SRS-XUS00-009", "法规符合需求"),
    ("SRS-XUS00-010", "外部连接"),
]

class RefTypes(Enum):
    img_struct = "img_struct"
    img_flow = "img_flow"
    img_topo = "img_topo"
    srs_reqs = "srs_reqs"
    srs_reqs_1 = "srs_reqs_1"
    srs_reqs_2 = "srs_reqs_2"
    srs_reqs_x = "srs_reqs_x"
    srs_reqds = "srs_reqds"

class Server(object):
    DOC_IMG_KEYWORDS = {
        "img_topo": ["物理拓扑图", "拓扑图"],
        "img_struct": ["系统结构图", "结构图"],
        "img_flow": ["网络安全流程图", "安全流程图", "流程图"],
    }
    TRACE_FIXED_NOTE_CODE = "SRS-RCN300-009"
    TRACE_FIXED_NOTE_PRODUCT_CODE_FALLBACK = "RCN3V2000"

    @classmethod
    def __build_trace_fixed_note_text(cls, product_code: str):
        code = str(product_code or "").strip() or cls.TRACE_FIXED_NOTE_PRODUCT_CODE_FALLBACK
        return (
            f"TX-TF-{code}-RD-009-A0 IEC62304《医疗器械软件 软件生存周期过程》符合性核查表"
            "、TX-TF-SD-001-A0 《DICOM一致性声明》"
            f"、TX-TF-{code}-RD-014-A0 《网络安全漏洞自评报告》"
        )

    @staticmethod
    def __extract_data_url_blob(data_url: str):
        if not data_url or not str(data_url).startswith("data:"):
            return None, None
        matched = re.match(r"^data:(image/[a-zA-Z0-9.+-]+);base64,(.+)$", data_url, re.S)
        if not matched:
            return None, None
        mime = matched.group(1).lower()
        b64 = matched.group(2)
        ext_map = {
            "image/png": ".png",
            "image/jpeg": ".jpg",
            "image/jpg": ".jpg",
            "image/gif": ".gif",
            "image/bmp": ".bmp",
            "image/webp": ".webp",
        }
        ext = ext_map.get(mime, ".png")
        try:
            blob = base64.b64decode(b64)
        except Exception:
            return None, None
        return blob, ext

    @staticmethod
    def __extract_image_blob_and_ext(img_url: str):
        if not img_url:
            return None, None
        img_url = str(img_url).strip()
        if img_url.startswith("data:"):
            return Server.__extract_data_url_blob(img_url)
        path = img_url
        if not os.path.exists(path):
            return None, None
        ext = os.path.splitext(path)[1] or ".png"
        try:
            with open(path, "rb") as fs:
                blob = fs.read()
            return blob, ext
        except Exception:
            return None, None

    @staticmethod
    def __heading_no_from_title(title: str):
        matched = re.match(r"^(\d+(?:\.\d+)*)", str(title or "").strip())
        return matched.group(1) if matched else ""

    @staticmethod
    def __blob_to_data_url(blob: bytes, ext_hint: str = ".png") -> str:
        ext = (ext_hint or ".png").lower()
        if not ext.startswith("."):
            ext = f".{ext}"
        mime_map = {
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".gif": "image/gif",
            ".bmp": "image/bmp",
            ".webp": "image/webp",
            ".tif": "image/tiff",
            ".tiff": "image/tiff",
            ".emf": "image/emf",
            ".wmf": "image/wmf",
        }
        mime = mime_map.get(ext, "image/png")
        b64 = base64.b64encode(blob).decode("ascii")
        return f"data:{mime};base64,{b64}"

    @staticmethod
    def __extract_images_from_ooxml_element(part, element):
        urls = []
        if element is None or part is None:
            return urls
        used_rids = set()
        rel_embed = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
        rel_link = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link"
        rel_id = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"

        def read_rid(rid):
            if not rid or rid in used_rids:
                return
            used_rids.add(rid)
            try:
                rel = part.rels[rid]
            except Exception:
                return
            target = getattr(rel, "target_ref", "")
            ext = (target.rsplit(".", 1)[-1].lower() if "." in target else "png")
            blob = getattr(rel.target_part, "blob", None)
            if blob:
                urls.append(Server.__blob_to_data_url(blob, f".{ext}"))

        for blip in element.xpath(".//*[local-name()='blip']"):
            read_rid(blip.get(rel_embed) or blip.get(rel_link))
        for imagedata in element.xpath(".//*[local-name()='imagedata']"):
            read_rid(imagedata.get(rel_id) or imagedata.get(rel_embed) or imagedata.get(rel_link))
        return urls

    def __pick_doc_images_from_docx_zip(self, docx_bytes: bytes, docx: Document = None):
        picked = {}
        if not docx_bytes:
            return picked
        try:
            with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
                if "word/document.xml" not in zf.namelist():
                    return picked
                doc_xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
                rid_to_media = {}
                rels_name = "word/_rels/document.xml.rels"
                if rels_name in zf.namelist():
                    rels_xml = zf.read(rels_name).decode("utf-8", errors="ignore")
                    for matched in re.finditer(r'Id="([^"]+)"[^>]*Target="([^"]+)"', rels_xml):
                        rid, target = matched.group(1), matched.group(2)
                        if "media/" in target.replace("\\", "/"):
                            normalized = target.replace("\\", "/").lstrip("/")
                            rid_to_media[rid] = normalized if normalized.startswith("word/") else f"word/{normalized}"
                text_only = re.sub(r"<[^>]+>", "", doc_xml)
                pos_22_candidates = [p for p in (text_only.find("2.2"), text_only.find("物理拓扑")) if p >= 0]
                pos_23_candidates = [p for p in (text_only.find("2.3"), text_only.find("系统结构")) if p >= 0]
                pos_22 = min(pos_22_candidates) if pos_22_candidates else -1
                pos_23 = min(pos_23_candidates) if pos_23_candidates else -1
                heading_category_map = {
                    "2.2": RefTypes.img_topo.value,
                    "2.3": RefTypes.img_struct.value,
                }
                embeds = []
                for matched in re.finditer(r'r:embed="([^"]+)"', doc_xml):
                    rid = matched.group(1)
                    media_path = rid_to_media.get(rid)
                    if media_path and media_path in zf.namelist():
                        embeds.append((matched.start(), media_path))
                embeds.sort(key=lambda item: item[0])
                for pos, media_path in embeds:
                    section = None
                    if pos_22 >= 0 and pos >= pos_22 and (pos_23 < 0 or pos < pos_23):
                        section = "2.2"
                    elif pos_23 >= 0 and pos >= pos_23:
                        section = "2.3"
                    if not section:
                        continue
                    blob = zf.read(media_path)
                    ext = os.path.splitext(media_path)[1] or ".png"
                    picked[heading_category_map[section]] = self.__blob_to_data_url(blob, ext)
        except Exception:
            logger.exception("pick doc images from docx zip failed")
        return picked

    def __pick_doc_images_from_doc_nodes(self, doc_id: int):
        picked = {}
        if not doc_id:
            return picked
        heading_category_map = {
            "2.2": RefTypes.img_topo.value,
            "2.3": RefTypes.img_struct.value,
        }
        nodes = db.session.execute(
            select(SrsNode).where(SrsNode.doc_id == doc_id).order_by(SrsNode.priority, SrsNode.n_id)
        ).scalars().all()
        if not nodes:
            return picked
        node_map = {row.n_id: row for row in nodes}

        def node_context_text(node: SrsNode) -> str:
            texts = []
            cur = node
            safety = 0
            while cur and safety < 100:
                for value in [cur.title, cur.label, cur.text]:
                    if value:
                        texts.append(self.__normalize_text(value))
                p_id = cur.p_id or 0
                if p_id == 0:
                    break
                cur = node_map.get(p_id)
                safety += 1
            texts.reverse()
            return " ".join(texts)

        for node in nodes:
            img_url = getattr(node, "img_url", None)
            if not img_url:
                continue
            img_str = str(img_url).strip()
            if not img_str.startswith("data:") and not os.path.exists(img_str):
                continue
            heading_no = self.__heading_no_from_title(node.title or "")
            if heading_no in heading_category_map:
                picked[heading_category_map[heading_no]] = img_str
                continue
            ref_type = getattr(node, "ref_type", None)
            if ref_type in heading_category_map.values():
                picked[ref_type] = img_str
                continue
            ctx_text = node_context_text(node)
            for category, keywords in self.DOC_IMG_KEYWORDS.items():
                if category not in heading_category_map.values():
                    continue
                if any(word in ctx_text for word in keywords):
                    picked[category] = img_str
        return picked

    def __pick_doc_images_from_docx(self, docx: Document):
        if docx is None:
            return {}
        picked = {}
        heading_category_map = {
            "2.2": RefTypes.img_topo.value,
            "2.3": RefTypes.img_struct.value,
        }
        active_target = None
        context_titles: List[str] = []

        def update_active_target(heading_no: str):
            nonlocal active_target
            if not heading_no:
                return
            if heading_no == "2.2" or heading_no.startswith("2.2."):
                active_target = "2.2"
                return
            if heading_no == "2.3" or heading_no.startswith("2.3."):
                active_target = "2.3"
                return
            if heading_no in ("2.4", "2.5", "2.6", "2.7") or not heading_no.startswith("2."):
                if active_target in heading_category_map:
                    active_target = None

        def assign_image(img_url: str):
            if not img_url:
                return
            if active_target in heading_category_map:
                picked[heading_category_map[active_target]] = img_url
                return
            ctx_text = " ".join(context_titles[-12:])
            for category, keywords in self.DOC_IMG_KEYWORDS.items():
                if category not in (RefTypes.img_topo.value, RefTypes.img_struct.value):
                    continue
                if any(word in ctx_text for word in keywords):
                    picked[category] = img_url

        for child in docx.element.body.iterchildren():
            tag = str(child.tag).lower()
            if tag.endswith("}p"):
                para = Paragraph(child, docx._body)
                if self.__is_toc_paragraph(para):
                    continue
                txt = self.__normalize_text(para.text)
                if txt:
                    level = self.__guess_heading_level(para)
                    heading_no = self.__extract_heading_number(txt) or self.__heading_no_from_title(txt)
                    if level is not None:
                        context_titles.append(txt)
                        update_active_target(heading_no)
                    elif heading_no in heading_category_map:
                        context_titles.append(txt)
                        update_active_target(heading_no)
                    elif "物理拓扑图" in txt or ("拓扑图" in txt and "物理" in txt):
                        context_titles.append(txt)
                        active_target = "2.2"
                    elif "系统结构图" in txt or ("结构图" in txt and "系统" in txt):
                        context_titles.append(txt)
                        active_target = "2.3"
                    elif level is None and heading_no:
                        update_active_target(heading_no)
                for img_url in self.__extract_images_from_ooxml_element(para.part, child):
                    assign_image(img_url)
            elif tag.endswith("}tbl"):
                tab = DocxTable(child, docx._body)
                for row in tab.rows:
                    for cell in row.cells:
                        cell_text = self.__normalize_text(cell.text)
                        if cell_text:
                            heading_no = self.__extract_heading_number(cell_text) or self.__heading_no_from_title(cell_text)
                            if heading_no:
                                context_titles.append(cell_text)
                                update_active_target(heading_no)
                        for img_url in self.__extract_images_from_ooxml_element(docx.part, cell._tc):
                            assign_image(img_url)
                for img_url in self.__extract_images_from_ooxml_element(docx.part, child):
                    assign_image(img_url)
        return picked

    def __pick_doc_images_from_tree(self, nodes: List[SrsNodeForm]):
        picked = {}
        heading_category_map = {
            "2.2": RefTypes.img_topo.value,
            "2.3": RefTypes.img_struct.value,
        }

        def collect_image_url(node: SrsNodeForm):
            img_url = getattr(node, "img_url", None)
            if not img_url:
                return None
            img_str = str(img_url).strip()
            if img_str.startswith("data:") or os.path.exists(img_str):
                return img_str
            return None

        def assign_by_heading(heading_no: str, img_url: str):
            category = heading_category_map.get(heading_no)
            if category and img_url:
                picked[category] = img_url

        def walk(items: List[SrsNodeForm], ctx_titles: List[str], active_heading: str = ""):
            for node in items or []:
                title = self.__normalize_text(getattr(node, "title", "") or "")
                heading_no = self.__heading_no_from_title(title)
                section_heading = heading_no if heading_no in heading_category_map else active_heading
                next_ctx = [*ctx_titles]
                if title:
                    next_ctx.append(title)

                img_url = collect_image_url(node)
                if img_url:
                    if section_heading in heading_category_map:
                        assign_by_heading(section_heading, img_url)
                    else:
                        ctx_text = " ".join(next_ctx)
                        for category, keywords in self.DOC_IMG_KEYWORDS.items():
                            if category in picked:
                                continue
                            if any(k in ctx_text for k in keywords):
                                picked[category] = img_url

                for child in getattr(node, "children", None) or []:
                    child_img = collect_image_url(child)
                    child_title = self.__normalize_text(getattr(child, "title", "") or "")
                    child_heading = heading_no if heading_no in heading_category_map else section_heading
                    if child_img and child_heading in heading_category_map:
                        assign_by_heading(child_heading, child_img)

                walk(getattr(node, "children", None) or [], next_ctx, section_heading)

        walk(nodes or [], [])
        return picked

    def __upsert_product_doc_image(
        self,
        product_id: int,
        category: str,
        img_url: str,
        doc_version: str = None,
        product_version: str = None,
    ):
        blob, ext = self.__extract_image_blob_and_ext(img_url)
        if not blob:
            return
        normalized_doc_version = sanitize_doc_image_token(doc_version or "")
        normalized_product_version = sanitize_doc_image_token(product_version or "")
        if not normalized_product_version:
            product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
            normalized_product_version = sanitize_doc_image_token(
                getattr(product, "full_version", "") or getattr(product, "name", "") or ""
            )
        row = pick_doc_image_file_row(
            product_id,
            category,
            normalized_doc_version,
            normalized_product_version,
        )
        if not row:
            row = DocFile(product_id=product_id, category=category)
            db.session.add(row)
            db.session.flush()
        path = os.path.join("data.trace", category, f"{row.id}{ext}")
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "wb") as fs:
            fs.write(blob)
        row.file_name = build_doc_image_file_name(
            normalized_product_version,
            normalized_doc_version,
            category,
            ext,
        )
        row.file_size = len(blob)
        row.file_url = path
        db.session.commit()
        logger.info(
            "sync product doc image: product_id=%s category=%s file=%s size=%s",
            product_id,
            category,
            row.file_name,
            row.file_size,
        )

    def __auto_sync_product_doc_images(
        self,
        product_id: int,
        nodes: List[SrsNodeForm],
        doc_version: str = None,
        docx: Document = None,
        product_version: str = None,
        docx_bytes: bytes = None,
        doc_id: int = None,
    ):
        picked = {}
        if docx_bytes:
            picked.update(self.__pick_doc_images_from_docx_zip(docx_bytes, docx))
        if docx is not None:
            for category, img_url in self.__pick_doc_images_from_docx(docx).items():
                if img_url:
                    picked[category] = img_url
        for category, img_url in self.__pick_doc_images_from_tree(nodes).items():
            if category not in picked and img_url:
                picked[category] = img_url
        if doc_id:
            for category, img_url in self.__pick_doc_images_from_doc_nodes(doc_id).items():
                if category not in picked and img_url:
                    picked[category] = img_url
        if not picked:
            logger.warning(
                "word import picked no doc images: product_id=%s doc_version=%s doc_id=%s",
                product_id,
                doc_version,
                doc_id,
            )
        else:
            logger.info(
                "word import picked doc images: product_id=%s doc_version=%s categories=%s",
                product_id,
                doc_version,
                list(picked.keys()),
            )
        for category, img_url in picked.items():
            self.__upsert_product_doc_image(
                product_id,
                category,
                img_url,
                doc_version,
                product_version,
            )

    @staticmethod
    def __guess_numpr_level(para):
        """读取 Word 自动编号层级（numPr/ilvl 或 outlineLvl）。"""
        def _level_from_ppr(p_pr):
            if p_pr is None:
                return None
            try:
                num_pr = getattr(p_pr, "numPr", None)
                ilvl = getattr(num_pr, "ilvl", None) if num_pr is not None else None
                val = getattr(ilvl, "val", None) if ilvl is not None else None
                if val is not None:
                    return max(1, min(int(str(val)) + 1, 5))
            except Exception:
                pass
            try:
                outline = getattr(p_pr, "outlineLvl", None)
                oval = getattr(outline, "val", None) if outline is not None else None
                if oval is not None:
                    return max(1, min(int(str(oval)) + 1, 5))
            except Exception:
                pass
            return None

        try:
            p_pr = getattr(getattr(para, "_element", None), "pPr", None)
            level = _level_from_ppr(p_pr)
            if level is not None:
                return level
            style = getattr(para, "style", None)
            hops = 0
            while style is not None and hops < 8:
                style_ppr = getattr(getattr(style, "_element", None), "pPr", None)
                level = _level_from_ppr(style_ppr)
                if level is not None:
                    return level
                style = getattr(style, "base_style", None)
                hops += 1
        except Exception:
            pass
        return None

    @staticmethod
    def __paragraph_num_info(para):
        """读取段落的 Word 自动编号 numId/ilvl，用于还原正文列表编号。"""
        def _num_info_from_ppr(p_pr):
            if p_pr is None:
                return None
            try:
                num_pr = Server.__xml_child(p_pr, "numPr")
                if num_pr is None:
                    num_pr = getattr(p_pr, "numPr", None)
                if num_pr is None:
                    return None
                num_id_el = Server.__xml_child(num_pr, "numId")
                if num_id_el is None:
                    num_id_el = getattr(num_pr, "numId", None)
                ilvl_el = Server.__xml_child(num_pr, "ilvl")
                if ilvl_el is None:
                    ilvl_el = getattr(num_pr, "ilvl", None)
                num_id = Server.__xml_attr(num_id_el, "val") or (getattr(num_id_el, "val", None) if num_id_el is not None else None)
                ilvl = Server.__xml_attr(ilvl_el, "val") or (getattr(ilvl_el, "val", None) if ilvl_el is not None else 0)
                if num_id is None:
                    return None
                return str(num_id), int(str(ilvl or 0))
            except Exception:
                return None

        try:
            p_pr = getattr(getattr(para, "_element", None), "pPr", None)
            info = _num_info_from_ppr(p_pr)
            if info is not None:
                return info
            style = getattr(para, "style", None)
            hops = 0
            while style is not None and hops < 8:
                style_ppr = getattr(getattr(style, "_element", None), "pPr", None)
                info = _num_info_from_ppr(style_ppr)
                if info is not None:
                    return info
                style = getattr(style, "base_style", None)
                hops += 1
        except Exception:
            pass
        return None

    @staticmethod
    def __xml_child(el, local_name: str):
        try:
            for child in el:
                if str(getattr(child, "tag", "")).endswith(f"}}{local_name}"):
                    return child
        except Exception:
            pass
        return None

    @staticmethod
    def __xml_attr(el, local_name: str):
        if el is None or qn is None:
            return None
        try:
            return el.get(qn(f"w:{local_name}"))
        except Exception:
            return None

    def __build_numbering_definitions(self, docx: Document):
        """解析 numbering.xml，建立 numId 到列表格式的映射。"""
        result = {}
        try:
            numbering_part = getattr(getattr(docx, "part", None), "numbering_part", None)
            numbering = getattr(numbering_part, "element", None)
            if numbering is None:
                return result

            abstract_levels = {}
            for abstract in numbering.xpath("./*[local-name()='abstractNum']"):
                abstract_id = self.__xml_attr(abstract, "abstractNumId")
                if abstract_id is None:
                    continue
                levels = {}
                for lvl in abstract.xpath("./*[local-name()='lvl']"):
                    ilvl = self.__xml_attr(lvl, "ilvl")
                    if ilvl is None:
                        continue
                    start_el = self.__xml_child(lvl, "start")
                    num_fmt_el = self.__xml_child(lvl, "numFmt")
                    lvl_text_el = self.__xml_child(lvl, "lvlText")
                    try:
                        start = int(str(self.__xml_attr(start_el, "val") or "1"))
                    except Exception:
                        start = 1
                    levels[int(str(ilvl))] = {
                        "start": start,
                        "num_fmt": self.__xml_attr(num_fmt_el, "val") or "decimal",
                        "lvl_text": self.__xml_attr(lvl_text_el, "val") or f"%{int(str(ilvl)) + 1}.",
                    }
                abstract_levels[str(abstract_id)] = levels

            for num in numbering.xpath("./*[local-name()='num']"):
                num_id = self.__xml_attr(num, "numId")
                abstract_id_el = self.__xml_child(num, "abstractNumId")
                abstract_id = self.__xml_attr(abstract_id_el, "val")
                if num_id is None or abstract_id is None:
                    continue
                levels = deepcopy(abstract_levels.get(str(abstract_id), {}))
                for override in num.xpath("./*[local-name()='lvlOverride']"):
                    ilvl = self.__xml_attr(override, "ilvl")
                    start_el = self.__xml_child(override, "startOverride")
                    if ilvl is None or start_el is None:
                        continue
                    try:
                        levels.setdefault(int(str(ilvl)), {})["start"] = int(str(self.__xml_attr(start_el, "val") or "1"))
                    except Exception:
                        pass
                result[str(num_id)] = levels
        except Exception:
            logger.exception("parse word numbering failed")
        return result

    def __build_body_numbering_prefix(self, para, numbering_defs: Dict[str, dict], counters: Dict[str, dict]):
        num_info = self.__paragraph_num_info(para)
        if not num_info:
            return None
        num_id, ilvl = num_info
        levels = numbering_defs.get(str(num_id)) or {}
        level_def = levels.get(int(ilvl)) or {}
        num_fmt = str(level_def.get("num_fmt") or "decimal").lower()
        if num_fmt == "bullet":
            return None

        item_counters = counters.setdefault(str(num_id), {})
        for key in list(item_counters.keys()):
            if int(key) > int(ilvl):
                item_counters.pop(key, None)
        start = int(level_def.get("start") or 1)
        item_counters[int(ilvl)] = int(item_counters.get(int(ilvl), start - 1)) + 1

        lvl_text = str(level_def.get("lvl_text") or f"%{int(ilvl) + 1}.")
        prefix = lvl_text
        for idx in range(9):
            value = item_counters.get(idx)
            if value is None:
                value = item_counters.get(int(ilvl), start)
            prefix = prefix.replace(f"%{idx + 1}", str(value))
        return prefix.strip()

    def __paragraph_text_with_numbering(self, para, numbering_defs: Dict[str, dict], counters: Dict[str, dict]):
        txt = self.__normalize_text(getattr(para, "text", "") or "")
        if not txt:
            return ""
        numbering_prefix = self.__build_body_numbering_prefix(para, numbering_defs or {}, counters)
        if numbering_prefix and not re.match(r"^\s*(?:\d+(?:\.\d+)*|[一二三四五六七八九十]+)[）).、．]\s*", txt):
            spacer = "" if re.search(r"[）)]$", numbering_prefix) else " "
            txt = f"{numbering_prefix}{spacer}{txt}".strip()
        return txt

    @staticmethod
    def __is_bold_paragraph(para):
        # 优先按 run 判断；若 run 未显式设置，再回退到样式链 bold。
        if any(run.bold for run in para.runs if (run.text or "").strip()):
            return True
        try:
            style = getattr(para, "style", None)
            hops = 0
            while style is not None and hops < 8:
                font = getattr(style, "font", None)
                if getattr(font, "bold", None) is True:
                    return True
                style = getattr(style, "base_style", None)
                hops += 1
        except Exception:
            pass
        return False

    @staticmethod
    def __is_req_table_title_text(text: str) -> bool:
        """需求表标题（如 2.0 变更需求 / 产品需求列表）不是文档章节。"""
        clean = Server.__normalize_text(text).replace("：", ":").rstrip(":").strip()
        if not clean:
            return False
        compact = re.sub(r"\s+", "", clean)
        return bool(
            re.search(r"(产品需求|标准需求|其他需求|变更需求|变更列表)", compact)
            or re.match(r"^表\s*\d+", clean)
        )

    @staticmethod
    def __guess_heading_level(para):
        txt = (para.text or "").strip()
        if not txt:
            return None
        if Server.__is_req_table_title_text(txt):
            return None
        is_bold = Server.__is_bold_paragraph(para)
        # JSON 键值行（如 "version":4,）按正文处理，不能当作章节标题
        if re.match(r'^\s*[\'"]\s*[^\'"]+\s*[\'"]\s*:\s*.+$', txt):
            return None
        # 带章节号前缀的 JSON 行（如 5.7.1.1 "version":4,）也按正文处理
        txt_wo_chapter = re.sub(r'^(\d+(?:\.\d+)*)(?:[\s、.．]+|(?=[\u4e00-\u9fffA-Za-z"\']))', '', txt).strip()
        if txt_wo_chapter and re.match(r'^\s*[\'"]\s*[^\'"]+\s*[\'"]\s*:\s*.+$', txt_wo_chapter):
            return None
        # 放宽：只要标题文本带明确章节号（如 5.7.1 / 5.6.1），即使未加粗也按章节识别。
        # 但排除“1.参数文件;”这类枚举项（单数字+点号+句末标点）。
        numbering = re.match(r"^(\d+(?:\.\d+){0,4})([\s、.．]+|(?=[\u4e00-\u9fffA-Za-z]))(.*)$", txt)
        if numbering:
            chapter_no = numbering.group(1) or ""
            sep = numbering.group(2) or ""
            tail = (numbering.group(3) or "").strip()
            if not tail:
                return None
            # 只有单一数字编号时（如 "1 xxx"），要求使用空白分隔；"1.xxx" 视为枚举项。
            if chapter_no.count(".") == 0 and not re.search(r"\s", sep):
                return None
            # 单级编号（如 "7 xxx"）时，进一步过滤“句子型正文项”，避免误识别为一级标题。
            # 例如：7 默认的科室不允许删除，删除时，系统提示：...
            if chapter_no.count(".") == 0:
                if len(tail) > 40:
                    return None
                if re.search(r"[，,。；;：:！？!?]", tail):
                    return None
            # 句末为分号/句号等更像正文项，不识别为标题。
            # 多级章节标题在 Word 中常写成“2.3.1 图像接收：”，这类冒号结尾仍应按章节处理。
            if re.search(r"[;；,，。！？!?]$", tail) or (chapter_no.count(".") < 2 and re.search(r"[:：]$", tail)):
                return None
            # 非粗体的两级编号（如 7.1 xxx）误识别概率高，增加约束：
            # 仅当尾部很短且无正文标点时才作为标题。
            if (not is_bold) and chapter_no.count(".") == 1:
                if len(tail) > 24:
                    return None
                if re.search(r"[，,。；;：:！？!?]", tail):
                    return None
            return max(1, min(chapter_no.count(".") + 1, 5))
        # 无明确章节号时，仍保持“粗体优先”约束，降低正文误判为标题
        if not is_bold:
            return None
        # 文本无显式编号但为粗体标题时，尝试读取 Word 编号层级（numPr/outlineLvl）
        numpr_level = Server.__guess_numpr_level(para)
        if numpr_level is not None:
            # 编号列表中的句子型文本（常见于正文要点）不应识别为章节标题
            if len(txt) > 24:
                return None
            if re.search(r"[，,。；;：:！？!?]", txt):
                return None
        return numpr_level

    @staticmethod
    def __normalize_text(value):
        return (value or "").replace("\xa0", " ").strip()

    @staticmethod
    def __is_toc_paragraph(para):
        txt = (para.text or "").replace("\xa0", " ").strip()
        if not txt:
            return False
        try:
            style_name = str(getattr(getattr(para, "style", None), "name", "") or "").lower()
            if "toc" in style_name or "目录" in style_name:
                return True
        except Exception:
            pass
        try:
            instr_text = " ".join([str(item.text or "") for item in para._element.xpath(".//*[local-name()='instrText']")])
            if re.search(r"\bTOC\b", instr_text, re.I):
                return True
        except Exception:
            pass
        # 目录项常见形态：“4. 图像接收 12”，最后一段数字是页码，不能当作正文标题导入。
        return bool(re.match(r"^\d+(?:\.\d+){0,4}[\s、.．]+.{1,80}\s+\d{1,4}$", txt))

    @staticmethod
    def __normalize_rcm_code(code: str):
        txt = re.sub(r"\s+", "", (code or "")).strip().upper()
        txt = txt.replace("＿", "_")
        txt = re.sub(r"[，。；;、,.]+$", "", txt)
        return txt

    @staticmethod
    def __normalize_srs_code(code: str):
        txt = (code or "").strip().upper()
        # 兼容“ SRS- RCN306-003 ”这类带空格/不可见字符的编号
        txt = re.sub(r"\s+", "", txt)
        txt = re.sub(r"[，。；;、,.]+$", "", txt)
        return txt

    def __clean_req_title(self, txt: str):
        value = self.__normalize_text(txt or "")
        value = re.sub(r"^\s*\d+(?:\.\d+)*[\s、.．:：\-]*", "", value).strip()
        value = re.sub(r"\bSRS[-_\sA-Za-z0-9.]+\b", "", value, flags=re.I).strip()
        return value

    def __clean_req_table_field(self, txt: str):
        value = self.__normalize_text(txt or "")
        if not value:
            return value
        value = re.sub(r"\s+", "", value)
        # Word 纵向合并单元格在部分解析路径中会把锚点文本按跨行次数拼接，
        # 例如“系统管理系统管理系统管理”应恢复为“系统管理”。
        for size in range(1, max(1, len(value) // 2) + 1):
            prefix = value[:size]
            pos = 0
            repeat_count = 0
            while value.startswith(prefix, pos):
                repeat_count += 1
                pos += size
            if repeat_count < 2:
                continue
            rest = value[pos:]
            if not rest:
                return prefix
            # 兼容“登录登录登录登录需求”这类重复主词后跟固定后缀的情况。
            if len(rest) <= 8 and not value.startswith(prefix, pos + len(rest)):
                return f"{prefix}{rest}"
        return value

    def __keep_req_field(self, txt: str):
        # 导出阶段：数据库需求字段已是干净值，仅做基础空白规整，
        # 不能像 __clean_req_table_field 那样压缩重复字符（会误伤 AAA、AAAA1 等合法内容）。
        return re.sub(r"\s+", "", self.__normalize_text(txt or ""))

    def __normalize_rcm_codes(self, codes):
        result = []
        for code in codes or []:
            c = self.__normalize_rcm_code(code)
            if c and c not in result:
                result.append(c)
        return result

    def __extract_rcm_codes_from_text(self, text: str):
        result = []
        for hit in re.findall(r"\bRCM[\s\-_]*[A-Z0-9]+(?:[\-_][A-Z0-9]+)*\b", str(text or ""), flags=re.I):
            code = self.__normalize_rcm_code(hit)
            if code and code not in result:
                result.append(code)
        return result

    def __table_search_text(self, table):
        payload = table
        if isinstance(payload, str):
            try:
                payload = json.loads(payload)
            except Exception:
                return payload
        if not isinstance(payload, dict):
            return str(payload or "")
        values = []
        for header in payload.get("headers") or []:
            values.append(str((header or {}).get("name") or ""))
        for row in payload.get("rows") or []:
            for value in (row or {}).values():
                values.append(str(value or ""))
        for row in payload.get("cells") or []:
            for cell in row or []:
                values.append(str((cell or {}).get("value") or ""))
        return "\n".join(values)

    @staticmethod
    def __normalize_header(value: str):
        return re.sub(r"[\s_:/（）()]+", "", (value or "").lower())

    @staticmethod
    def __extract_heading_number(title: str):
        matched = re.match(r"^(\d+(?:\.\d+)*)(?:[\s、.．]+|(?=[\u4e00-\u9fffA-Za-z]))", (title or "").strip())
        return matched.group(1) if matched else None

    def __validate_heading_numbers(self, heading_rows: List[dict]):
        rows = [row for row in heading_rows if row.get("level") in [1, 2, 3, 4, 5]]
        if not rows:
            return None
        level1_rows = [r for r in rows if r.get("level") == 1]
        if not level1_rows:
            return None

        # 按业务约定：目录/修订记录等前置章节不校验，从一级标题“介绍”开始校验。
        intro_idx = next(
            (
                idx
                for idx, row in enumerate(level1_rows)
                if "介绍" in self.__normalize_text(row.get("title") or "")
            ),
            None,
        )
        if intro_idx is not None:
            level1_rows = level1_rows[intro_idx:]
            if not level1_rows:
                return None

        level1_nums = [r for r in level1_rows if r.get("number")]
        if not level1_nums:
            return "Word导入校验失败：未识别到带编号的一级标题，请检查Word标题样式与编号。"

        # 放宽规则：从“介绍”开始，仅校验“有编号”的一级标题；未编号一级标题跳过。

        valid_level1_numbers = {self.__normalize_text(r.get("number") or "") for r in level1_rows if r.get("number")}
        if intro_idx is not None:
            rows = [
                r
                for r in rows
                if int(r.get("level") or 1) != 1
                or self.__normalize_text(r.get("number") or "") in valid_level1_numbers
            ]

        # Validate numbering continuity under each level/parent path.
        parent_last: Dict[Tuple[int, str], int] = {}
        for ridx, row in enumerate(rows, start=1):
            level = int(row.get("level") or 1)
            num = row.get("number") or ""
            if not num:
                continue
            parts = num.split(".")
            if len(parts) < level:
                # 放宽校验：当样式层级与编号层级不一致时，以编号层级为准继续解析，不阻断导入。
                level = len(parts)
            if level <= 0:
                continue
            try:
                cur = int(parts[level - 1])
            except Exception:
                if level == 1:
                    return f"Word导入校验失败：第{ridx}个一级标题编号格式错误（{row.get('title')}）"
                return f"Word导入校验失败：标题编号格式错误（{row.get('title')}）"

            parent_key = ".".join(parts[: level - 1]) if level > 1 else "_root_"
            key = (level, parent_key)
            last = parent_last.get(key, 0)
            expected = last + 1
            if cur != expected:
                level_text = "一级" if level == 1 else f"{level}级"
                return (
                    "Word导入校验失败："
                    f"{level_text}标题编号应为 {expected}，实际为 {cur}（{row.get('title')}）"
                )
            parent_last[key] = cur
        return None

    def __extract_file_info(self, file_name: str):
        base_name = os.path.splitext(os.path.basename(file_name or ""))[0]
        if not base_name:
            return None, None
        # 常见命名：TX-TF-RCN3V2000-PD-003-A0需求规格说明.docx
        # 规则：
        # - 文件名称（folder_name）优先取中文标题部分（如“需求规格说明”）
        # - 文件编号（file_no）优先取中文标题前的编码串
        title_cn = "".join(re.findall(r"[\u4e00-\u9fff]+", base_name)).strip()
        prefix = re.split(r"[\u4e00-\u9fff]+", base_name, maxsplit=1)[0].strip(" _-")
        name_for_parse = prefix or base_name
        tokens = [tok for tok in re.split(r"[_\-\s]+", name_for_parse) if tok]
        if not tokens and not title_cn:
            return None, None

        folder_name = title_cn or (tokens[0] if tokens else None)
        file_no = name_for_parse if name_for_parse else None
        if not file_no:
            for token in tokens:
                if re.match(r"^[A-Za-z]{1,6}\d{2,}$", token) or re.match(r"^[A-Za-z0-9]+-\d+$", token):
                    file_no = token
                    break
            if not file_no and len(tokens) > 1:
                file_no = tokens[1]
        return folder_name, file_no or folder_name

    def __is_product_req_context(self, context_text: str):
        normalized = self.__normalize_header(context_text or "")
        keywords = [
            "产品需求表",
            "产品需求",
            "需求列表",
            "software requirement",
            "product requirement",
            "srs",
        ]
        return any(self.__normalize_header(word) in normalized for word in keywords)

    def __resolve_req_columns(self, headers_norm: List[str]):
        col_idx: Dict[str, int] = {}
        for idx, h in enumerate(headers_norm):
            if ("需求编号" in h or "需求列表" in h or "srscode" in h or h in ["code", "srs_code"]) and "code" not in col_idx:
                col_idx["code"] = idx
            if ("模块" in h or h == "module") and "module" not in col_idx:
                col_idx["module"] = idx
            if ("子功能" in h or "subfunction" in h) and "sub_function" not in col_idx:
                col_idx["sub_function"] = idx
            if ("功能" in h or h == "function") and "function" not in col_idx:
                col_idx["function"] = idx
            if ("章节" in h or "位置" in h or "location" in h) and "location" not in col_idx:
                col_idx["location"] = idx
            if ("rcm" in h or "风险控制" in h) and "rcm" not in col_idx:
                col_idx["rcm"] = idx
        return col_idx

    def __fill_req_table_merged_values(self, values: List[str], col_idx: Dict[str, int], last_values: Dict[str, str]):
        """Word 纵向合并单元格在续行常为空，需求表字段需要继承上一条非空值。"""
        next_values = list(values or [])
        for field in ["module", "function", "sub_function", "location"]:
            idx = col_idx.get(field)
            if idx is None or idx >= len(next_values):
                continue
            current = self.__normalize_text(next_values[idx])
            if current:
                last_values[field] = current
                next_values[idx] = current
            elif last_values.get(field):
                next_values[idx] = last_values[field]
        return next_values

    def __extract_srs_reqs_from_tables(self, docx: Document):
        req_rows = []
        req_rcm_map: Dict[str, set] = {}
        code_pattern = re.compile(r"^SRS[-_A-Za-z0-9.]+$", re.I)
        rcm_pattern = re.compile(r"\bRCM[-_A-Za-z0-9]+\b", re.I)
        current_context = ""
        for child in docx.element.body.iterchildren():
            tag = str(child.tag).lower()
            if tag.endswith("}p"):
                para = Paragraph(child, docx._body)
                txt = self.__normalize_text(para.text)
                if txt:
                    current_context = txt
                continue
            if not tag.endswith("}tbl"):
                continue
            tab = DocxTable(child, docx._body)
            if not tab.rows:
                continue
            parsed_table = self.__parse_docx_table(tab, self.__build_numbering_definitions(docx))
            headers = [self.__normalize_text(getattr(h, "name", "") or "") for h in (getattr(parsed_table, "headers", None) or [])]
            raw_rows = getattr(parsed_table, "rows", None) or []
            header_codes = [getattr(h, "code", "") for h in (getattr(parsed_table, "headers", None) or [])]
            if not headers:
                headers = [self.__normalize_text(cell.text) for cell in tab.rows[0].cells]
            headers_norm = [self.__normalize_header(item) for item in headers]
            if not headers_norm:
                continue

            col_idx = self.__resolve_req_columns(headers_norm)
            if "code" not in col_idx:
                continue
            # 强规则：需满足“上下文命中产品需求”或“表头至少含编号+模块+功能”。
            has_core_cols = ("module" in col_idx and "function" in col_idx)
            if not (self.__is_product_req_context(current_context) or has_core_cols):
                continue

            # 需求分类：
            # - 产品需求：需求编号/模块/功能(/子功能)
            # - 其他需求：需求编号/模块/对应章节（无功能与子功能）
            has_location_col = "location" in col_idx
            has_function_col = "function" in col_idx
            has_sub_function_col = "sub_function" in col_idx
            type_code = "2" if has_location_col and not has_function_col and not has_sub_function_col else "1"
            last_values: Dict[str, str] = {}
            source_rows = raw_rows if raw_rows else tab.rows[1:]
            for row in source_rows:
                if isinstance(row, dict):
                    values = [self.__normalize_text(str((row or {}).get(code, "") or "")) for code in header_codes]
                else:
                    values = [self.__normalize_text(cell.text) for cell in row.cells]
                code = values[col_idx["code"]] if col_idx["code"] < len(values) else ""
                code = self.__normalize_srs_code(code)
                if not code_pattern.match(code or ""):
                    continue
                values = self.__fill_req_table_merged_values(values, col_idx, last_values)
                code_upper = code.upper()
                req_rows.append(
                    dict(
                        code=code_upper,
                        type_code=type_code,
                        module=(self.__clean_req_table_field(values[col_idx["module"]]) if "module" in col_idx and col_idx["module"] < len(values) else None),
                        function=(self.__clean_req_table_field(values[col_idx["function"]]) if "function" in col_idx and col_idx["function"] < len(values) else None),
                        sub_function=(self.__clean_req_table_field(values[col_idx["sub_function"]]) if "sub_function" in col_idx and col_idx["sub_function"] < len(values) else None),
                        location=(self.__clean_req_table_field(values[col_idx["location"]]) if "location" in col_idx and col_idx["location"] < len(values) else None),
                    )
                )
                if "rcm" in col_idx and col_idx["rcm"] < len(values):
                    rcm_codes = {self.__normalize_rcm_code(item) for item in rcm_pattern.findall(values[col_idx["rcm"]] or "")}
                    rcm_codes = {code for code in rcm_codes if code}
                    if rcm_codes:
                        req_rcm_map.setdefault(code_upper, set()).update(rcm_codes)
        return req_rows, req_rcm_map

    def __normalize_change_req_type_name(self, value: str):
        name = self.__normalize_text(str(value or "")).replace("：", ":").rstrip(":").strip()
        return name or "变更需求"

    def __normalize_change_req_type_key(self, value: str):
        return re.sub(r"\s+", "", self.__normalize_change_req_type_name(value))

    def __strip_change_table_title_heading(self, value: str) -> str:
        txt = self.__normalize_change_req_type_name(value)
        return re.sub(
            r"^\d+(?:\.\d+)*(?:[\s、.．]+|(?=[\u4e00-\u9fffA-Za-z]))",
            "",
            txt,
        ).strip()

    def __change_table_title_matches(self, table_title: str, type_name: str) -> bool:
        """仅用于识别同一张变更表，不会把表名统一改成 type_name。"""
        left_key = self.__normalize_change_req_type_key(table_title)
        right_key = self.__normalize_change_req_type_key(type_name)
        if not left_key or not right_key:
            return False
        if left_key == right_key:
            return True
        left_body = self.__normalize_change_req_type_key(
            self.__strip_change_table_title_heading(table_title)
        )
        right_body = self.__normalize_change_req_type_key(
            self.__strip_change_table_title_heading(type_name)
        )
        return bool(left_body and right_body and left_body == right_body)

    def __remap_change_req_types_by_existing_names(self, doc_id: int, req_rows: List[dict]):
        if not doc_id or not req_rows:
            return req_rows
        type_rows: List[SrsType] = db.session.execute(
            select(SrsType).where(SrsType.doc_id == doc_id)
        ).scalars().all()
        type_code_by_name = {}
        for row in type_rows:
            key = self.__normalize_change_req_type_key(row.type_name)
            if key and key not in type_code_by_name:
                type_code_by_name[key] = row.type_code
        for item in req_rows:
            type_code = item.get("type_code")
            type_name = item.get("type_name")
            if type_code in ["1", "2", "", None] or not type_name:
                continue
            name_key = self.__normalize_change_req_type_key(type_name)
            existing_type_code = type_code_by_name.get(name_key)
            if not existing_type_code:
                for row in type_rows:
                    if self.__change_table_title_matches(type_name, row.type_name):
                        existing_type_code = row.type_code
                        break
            if existing_type_code:
                item["type_code"] = existing_type_code
            elif name_key:
                type_code_by_name[name_key] = type_code
        return req_rows

    def __merge_duplicate_change_req_types_by_name(self, doc_id: int):
        if not doc_id:
            return
        type_rows: List[SrsType] = db.session.execute(
            select(SrsType).where(SrsType.doc_id == doc_id).order_by(SrsType.id)
        ).scalars().all()
        groups: Dict[str, List[SrsType]] = {}
        for row in type_rows:
            type_name = self.__normalize_change_req_type_name(row.type_name)
            if "变更" not in type_name:
                continue
            key = self.__normalize_change_req_type_key(type_name)
            if key:
                groups.setdefault(key, []).append(row)
        for rows in groups.values():
            if len(rows) <= 1:
                continue
            rows = sorted(rows, key=lambda row: (str(row.type_code or "").startswith("change_"), row.id or 0))
            keep_type = rows[0]
            for dup_type in rows[1:]:
                dup_reqs: List[SrsReq] = db.session.execute(
                    select(SrsReq).where(SrsReq.doc_id == doc_id, SrsReq.type_code == dup_type.type_code)
                ).scalars().all()
                for dup_req in dup_reqs:
                    keep_req = db.session.execute(
                        select(SrsReq).where(
                            SrsReq.doc_id == doc_id,
                            SrsReq.type_code == keep_type.type_code,
                            SrsReq.code == dup_req.code,
                        )
                    ).scalars().first()
                    if not keep_req:
                        dup_req.type_code = keep_type.type_code
                        continue
                    for trace_row in db.session.execute(select(SdsTrace).where(SdsTrace.req_id == dup_req.id)).scalars().all():
                        existed_trace = db.session.execute(
                            select(SdsTrace).where(SdsTrace.doc_id == trace_row.doc_id, SdsTrace.req_id == keep_req.id)
                        ).scalars().first()
                        if existed_trace:
                            db.session.delete(trace_row)
                        else:
                            trace_row.req_id = keep_req.id
                    db.session.execute(delete(ReqRcm).where(ReqRcm.req_id == dup_req.id))
                    db.session.execute(delete(SrsReqd).where(SrsReqd.req_id == dup_req.id))
                    db.session.execute(delete(SdsReqd).where(SdsReqd.req_id == dup_req.id))
                    db.session.delete(dup_req)
                db.session.delete(dup_type)

    def __resolve_doc_table_req_type(self, node: SrsNodeForm, has_other_cols: bool, has_function_col: bool, has_sub_function_col: bool):
        table = getattr(node, "table", None)
        table_name = self.__normalize_change_req_type_name(getattr(table, "name", "") or getattr(node, "title", ""))
        if "变更" in table_name:
            digest = hashlib.md5(table_name.encode("utf-8")).hexdigest()[:12]
            return f"change_{digest}", table_name
        if has_other_cols and not has_function_col and not has_sub_function_col:
            return "2", None
        return "1", None

    def __extract_srs_reqs_from_nodes(self, nodes: List[SrsNodeForm], include_node_codes: bool = True):
        req_rows = []
        req_rcm_map: Dict[str, set] = {}
        # 放宽编号格式，兼容 SRS-XXX / CNXXX / 其他编码串
        code_pattern = re.compile(r"^[A-Z]{2,}(?:[-_][A-Z0-9.]+)+$", re.I)
        rcm_pattern = re.compile(r"\bRCM[-_A-Za-z0-9]+\b", re.I)
        seen = set()
        heading_no_re = re.compile(r"^\s*\d+(?:\.\d+)*[\s、.．:：\-]*")

        def clean_title(txt: str):
            return self.__clean_req_title(txt)

        def walk(items: List[SrsNodeForm], parent_titles: List[str] = None):
            parent_titles = parent_titles or []
            for node in items or []:
                table = getattr(node, "table", None)
                headers = getattr(table, "headers", None) if table else None
                rows = getattr(table, "rows", None) if table else None
                if headers and rows:
                    header_names = [self.__normalize_text(getattr(h, "name", "") or "") for h in headers]
                    header_norm = [self.__normalize_header(h) for h in header_names]
                    col_idx = self.__resolve_req_columns(header_norm)
                    has_product_cols = ("code" in col_idx and "module" in col_idx and "function" in col_idx)
                    has_other_cols = ("code" in col_idx and "module" in col_idx and "location" in col_idx)
                    if has_product_cols or has_other_cols:
                        type_code, type_name = self.__resolve_doc_table_req_type(
                            node,
                            has_other_cols,
                            "function" in col_idx,
                            "sub_function" in col_idx,
                        )
                        col_codes = [getattr(h, "code", "") for h in headers]
                        last_values: Dict[str, str] = {}
                        for row in rows or []:
                            values = [self.__normalize_text(str((row or {}).get(code, "") or "")) for code in col_codes]
                            code = values[col_idx["code"]] if col_idx["code"] < len(values) else ""
                            code = self.__normalize_srs_code(code)
                            if not code_pattern.match(code or ""):
                                continue
                            values = self.__fill_req_table_merged_values(values, col_idx, last_values)
                            code_upper = code.upper()
                            key = (type_code, code_upper)
                            if key in seen:
                                continue
                            seen.add(key)
                            req_rows.append(
                                dict(
                                    code=code_upper,
                                    type_code=type_code,
                                    type_name=type_name,
                                    module=(self.__clean_req_table_field(values[col_idx["module"]]) if "module" in col_idx and col_idx["module"] < len(values) else None),
                                    function=(self.__clean_req_table_field(values[col_idx["function"]]) if "function" in col_idx and col_idx["function"] < len(values) else None),
                                    sub_function=(self.__clean_req_table_field(values[col_idx["sub_function"]]) if "sub_function" in col_idx and col_idx["sub_function"] < len(values) else None),
                                    location=(self.__clean_req_table_field(values[col_idx["location"]]) if "location" in col_idx and col_idx["location"] < len(values) else None),
                                )
                            )
                            if "rcm" in col_idx and col_idx["rcm"] < len(values):
                                rcm_codes = {self.__normalize_rcm_code(item) for item in rcm_pattern.findall(values[col_idx["rcm"]] or "")}
                                rcm_codes = {code for code in rcm_codes if code}
                                if rcm_codes:
                                    req_rcm_map.setdefault(code_upper, set()).update(rcm_codes)
                if include_node_codes:
                    # 兜底：从章节节点上的 srs_code 直接生成需求，避免因表格格式变化导致 SRS 管理为空
                    node_srs_code = self.__normalize_srs_code(str(getattr(node, "srs_code", "") or ""))
                    if node_srs_code and code_pattern.match(node_srs_code) and not re.match(r"^SRS-RCN300-\d+$", node_srs_code, re.I):
                        key = ("1", node_srs_code.upper())
                        if key not in seen:
                            seen.add(key)
                            title_txt = clean_title(getattr(node, "title", "") or "")
                            parent_txt = clean_title(parent_titles[-1] if parent_titles else "")
                            req_rows.append(
                                dict(
                                    code=node_srs_code.upper(),
                                    type_code="1",
                                    module=parent_txt or None,
                                    function=title_txt or None,
                                    sub_function=None,
                                    location=None,
                                )
                            )

                next_parents = parent_titles + [getattr(node, "title", "") or ""]
                walk(getattr(node, "children", None) or [], next_parents)

        walk(nodes or [], [])
        return req_rows, req_rcm_map

    def __map_reqd_field(self, label: str):
        norm = self.__normalize_header(label or "")
        if not norm:
            return None
        if "需求编号" in norm or norm in ["srscode", "code"]:
            return "code"
        if "需求名称" in norm or norm == "name":
            return "name"
        if "需求概述" in norm or "概述" in norm or norm == "overview":
            return "overview"
        if "主参加者" in norm or "参与人" in norm or norm in ["participant"]:
            return "participant"
        if "前置条件" in norm or norm in ["precondition", "pre_condition"]:
            return "pre_condition"
        if "触发器" in norm or "触发条件" in norm or norm in ["trigger"]:
            return "trigger"
        if "工作流" in norm or "工作流程" in norm or norm in ["workflow", "work_flow"]:
            return "work_flow"
        if "后置条件" in norm or norm in ["postcondition", "post_condition"]:
            return "post_condition"
        if "异常情况" in norm or "异常" in norm or norm in ["exception"]:
            return "exception"
        if "约束" in norm or "限制" in norm or norm in ["constraint"]:
            return "constraint"
        return None

    def __extract_srs_reqds_from_nodes(self, nodes: List[SrsNodeForm]):
        code_pattern = re.compile(r"^SRS[-_A-Za-z0-9.]+$", re.I)
        reqd_dict: Dict[str, dict] = {}

        def merge_row(code: str, data: dict):
            code_up = (code or "").strip().upper()
            if not code_up:
                return
            item = reqd_dict.setdefault(code_up, {"code": code_up})
            for key in ["name", "overview", "participant", "pre_condition", "trigger", "work_flow", "post_condition", "exception", "constraint"]:
                val = (data.get(key) or "").strip()
                if val and not item.get(key):
                    item[key] = val

        def walk(items: List[SrsNodeForm]):
            for node in items or []:
                node_code = self.__normalize_srs_code(str(getattr(node, "srs_code", "") or ""))
                node_text = self.__normalize_text(str(getattr(node, "text", "") or ""))
                if node_code and code_pattern.match(node_code) and node_text:
                    merge_row(node_code, {
                        "name": self.__normalize_text(str(getattr(node, "title", "") or "")),
                        "overview": node_text,
                    })

                table = getattr(node, "table", None)
                headers = getattr(table, "headers", None) if table else None
                rows = getattr(table, "rows", None) if table else None
                if headers and rows and len(headers) >= 2:
                    col_codes = [getattr(h, "code", "") for h in headers]
                    pairs = []
                    # 两列表格常把“需求编号|SRS-XXX”解析为表头，先作为首行键值对处理
                    h_left = self.__normalize_text(getattr(headers[0], "name", "") or "")
                    h_right = self.__normalize_text(getattr(headers[1], "name", "") or "")
                    if h_left or h_right:
                        pairs.append((h_left, h_right))
                    for row in rows or []:
                        left = self.__normalize_text(str((row or {}).get(col_codes[0], "") or ""))
                        right = self.__normalize_text(str((row or {}).get(col_codes[1], "") or ""))
                        if left or right:
                            pairs.append((left, right))

                    payload = {}
                    req_code = ""
                    for left, right in pairs:
                        field_key = self.__map_reqd_field(left)
                        if not field_key:
                            continue
                        if field_key == "code":
                            req_code = self.__normalize_srs_code(right or "")
                        else:
                            payload[field_key] = right
                    if not req_code and getattr(node, "srs_code", None):
                        req_code = self.__normalize_srs_code(str(getattr(node, "srs_code") or ""))
                    if req_code and code_pattern.match(req_code):
                        merge_row(req_code, payload)

                walk(getattr(node, "children", None) or [])

        walk(nodes or [])
        return list(reqd_dict.values())

    def __upsert_imported_srs_reqds(self, doc_id: int, reqd_rows: List[dict]):
        if not reqd_rows:
            return
        req_codes = [self.__normalize_srs_code(str((item or {}).get("code") or "")) for item in reqd_rows]
        req_codes = [code for code in req_codes if code]
        if not req_codes:
            return
        reqs = db.session.execute(
            select(SrsReq).where(SrsReq.doc_id == doc_id, SrsReq.type_code != "2", SrsReq.code.in_(req_codes))
        ).scalars().all()
        req_map = {row.code: row for row in reqs}
        if not req_map:
            return
        req_ids = [row.id for row in reqs]
        reqd_exists = db.session.execute(select(SrsReqd).where(SrsReqd.req_id.in_(req_ids))).scalars().all()
        reqd_map = {row.req_id: row for row in reqd_exists}

        for item in reqd_rows:
            code = self.__normalize_srs_code(str((item or {}).get("code") or ""))
            req_row = req_map.get(code)
            if not req_row:
                continue
            reqd_row = reqd_map.get(req_row.id)
            if not reqd_row:
                reqd_row = SrsReqd(req_id=req_row.id)
                db.session.add(reqd_row)
                reqd_map[req_row.id] = reqd_row
            for key in ["name", "overview", "participant", "pre_condition", "trigger", "work_flow", "post_condition", "exception", "constraint"]:
                val = str((item or {}).get(key) or "").strip()
                if val:
                    setattr(reqd_row, key, val)
        db.session.commit()

    def __sync_srs_req_names_from_doc_nodes(self, doc_id: int, nodes: List[SrsNodeForm]):
        sync_map: Dict[str, dict] = {}

        def extract_req_detail_from_table(table):
            headers = getattr(table, "headers", None) if table else None
            rows = getattr(table, "rows", None) if table else None
            if not headers or not rows or len(headers) < 2:
                return {}
            col_codes = [getattr(h, "code", "") for h in headers]
            pairs = []
            h_left = self.__normalize_text(getattr(headers[0], "name", "") or "")
            h_right = self.__normalize_text(getattr(headers[1], "name", "") or "")
            if h_left or h_right:
                pairs.append((h_left, h_right))
            for row in rows or []:
                left = self.__normalize_text(str((row or {}).get(col_codes[0], "") or ""))
                right = self.__normalize_text(str((row or {}).get(col_codes[1], "") or ""))
                if left or right:
                    pairs.append((left, right))
            result = {}
            for left, right in pairs:
                if self.__map_reqd_field(left) == "code":
                    result["code"] = self.__normalize_srs_code(right or "")
                elif self.__map_reqd_field(left) == "name":
                    result["name"] = self.__normalize_text(right or "")
            return result

        def put_entry(code: str, titles: List[str]):
            code = self.__normalize_srs_code(code or "")
            if not code:
                return
            clean_titles = [self.__clean_req_title(title) for title in titles or []]
            clean_titles = [title for title in clean_titles if title]
            if not clean_titles:
                return
            # 第一层通常是“7 图像显示”这类大章；需求管理使用其下的模块/功能/子功能。
            parts = clean_titles[1:] if len(clean_titles) > 1 else clean_titles
            item = sync_map.setdefault(code, {})
            item["name"] = parts[-1]
            if len(parts) >= 1:
                item["module"] = self.__clean_req_table_field(parts[0])
            if len(parts) >= 2:
                item["function"] = self.__clean_req_table_field(parts[1])
            if len(parts) >= 3:
                item["sub_function"] = self.__clean_req_table_field(parts[2])

        def walk(items: List[SrsNodeForm], path: List[str] = None):
            path = path or []
            for node in items or []:
                title = getattr(node, "title", "") or ""
                next_path = path + [title]
                node_code = self.__normalize_srs_code(str(getattr(node, "srs_code", "") or ""))
                table_detail = extract_req_detail_from_table(getattr(node, "table", None))
                table_code = table_detail.get("code") or ""
                table_path = path if table_code and re.match(r"^导入表格\d*$", title.strip()) else next_path
                put_entry(node_code or table_code, table_path)
                if table_code and table_detail.get("name"):
                    sync_map.setdefault(table_code, {})["name"] = table_detail.get("name")
                walk(getattr(node, "children", None) or [], next_path)

        walk(nodes or [], [])
        if not sync_map:
            return
        rows: List[SrsReq] = db.session.execute(
            select(SrsReq).where(SrsReq.doc_id == doc_id, SrsReq.code.in_(list(sync_map.keys())), SrsReq.type_code != "2")
        ).scalars().all()
        if not rows:
            return
        req_ids = [row.id for row in rows]
        reqd_rows: List[SrsReqd] = db.session.execute(select(SrsReqd).where(SrsReqd.req_id.in_(req_ids))).scalars().all()
        reqd_map = {row.req_id: row for row in reqd_rows}
        for row in rows:
            item = sync_map.get(row.code) or {}
            # 标准 SRS 表是模块/功能/子功能的权威来源；详情章节标题只用于补充需求名称，
            # 不能反向覆盖用户在标准表中编辑的功能名。
            if item.get("name"):
                reqd_row = reqd_map.get(row.id)
                if not reqd_row:
                    reqd_row = SrsReqd(req_id=row.id)
                    db.session.add(reqd_row)
                    reqd_map[row.id] = reqd_row
                reqd_row.name = item.get("name")
        db.session.commit()

    def __sync_doc_srs_tables_from_doc_nodes(self, nodes: List[SrsNodeForm]):
        sync_map: Dict[str, dict] = {}

        def extract_req_code_from_table(table):
            headers = getattr(table, "headers", None) if table else None
            rows = getattr(table, "rows", None) if table else None
            if not headers or not rows or len(headers) < 2:
                return ""
            col_codes = [getattr(h, "code", "") for h in headers]
            pairs = []
            h_left = self.__normalize_text(getattr(headers[0], "name", "") or "")
            h_right = self.__normalize_text(getattr(headers[1], "name", "") or "")
            if h_left or h_right:
                pairs.append((h_left, h_right))
            for row in rows or []:
                left = self.__normalize_text(str((row or {}).get(col_codes[0], "") or ""))
                right = self.__normalize_text(str((row or {}).get(col_codes[1], "") or ""))
                if left or right:
                    pairs.append((left, right))
            for left, right in pairs:
                if self.__map_reqd_field(left) == "code":
                    return self.__normalize_srs_code(right or "")
            return ""

        def put_entry(code: str, titles: List[str]):
            code = self.__normalize_srs_code(code or "")
            if not code:
                return
            clean_titles = [self.__clean_req_title(title) for title in titles or []]
            clean_titles = [title for title in clean_titles if title]
            if not clean_titles:
                return
            parts = clean_titles[1:] if len(clean_titles) > 1 else clean_titles
            item = sync_map.setdefault(code, {})
            if len(parts) >= 1:
                item["module"] = self.__clean_req_table_field(parts[0])
            if len(parts) >= 2:
                item["function"] = self.__clean_req_table_field(parts[1])
            if len(parts) >= 3:
                item["sub_function"] = self.__clean_req_table_field(parts[2])

        def collect(items: List[SrsNodeForm], path: List[str] = None):
            path = path or []
            for node in items or []:
                title = getattr(node, "title", "") or ""
                next_path = path + [title]
                node_code = self.__normalize_srs_code(str(getattr(node, "srs_code", "") or ""))
                table_code = extract_req_code_from_table(getattr(node, "table", None))
                table_path = path if table_code and re.match(r"^导入表格\d*$", title.strip()) else next_path
                put_entry(node_code or table_code, table_path)
                collect(getattr(node, "children", None) or [], next_path)

        def apply_tables(items: List[SrsNodeForm]):
            for node in items or []:
                table = getattr(node, "table", None)
                headers = getattr(table, "headers", None) if table else None
                rows = getattr(table, "rows", None) if table else None
                if headers and rows:
                    header_norm = [self.__normalize_header(getattr(h, "name", "") or "") for h in headers]
                    col_idx = self.__resolve_req_columns(header_norm)
                    if "code" in col_idx and ("module" in col_idx or "function" in col_idx or "sub_function" in col_idx):
                        col_codes = [getattr(h, "code", "") for h in headers]
                        for row in rows or []:
                            code_col = col_codes[col_idx["code"]]
                            code = self.__normalize_srs_code(str((row or {}).get(code_col, "") or ""))
                            item = sync_map.get(code)
                            if not item:
                                continue
                            if item.get("module") and "module" in col_idx:
                                row[col_codes[col_idx["module"]]] = item.get("module")
                            if item.get("function") and "function" in col_idx:
                                row[col_codes[col_idx["function"]]] = item.get("function")
                            if item.get("sub_function") and "sub_function" in col_idx:
                                row[col_codes[col_idx["sub_function"]]] = item.get("sub_function")
                apply_tables(getattr(node, "children", None) or [])

        collect(nodes or [], [])
        if sync_map:
            apply_tables(nodes or [])

    def __sync_saved_doc_srs_tables_from_req_rows(self, doc_id: int):
        req_rows: List[SrsReq] = db.session.execute(
            select(SrsReq).where(SrsReq.doc_id == doc_id)
        ).scalars().all()
        req_map = {
            self.__normalize_srs_code(row.code or ""): {
                "module": row.module,
                "function": row.function,
                "sub_function": row.sub_function,
                "location": row.location,
            }
            for row in req_rows
            if self.__normalize_srs_code(row.code or "")
        }

        nodes: List[SrsNode] = db.session.execute(
            select(SrsNode).where(SrsNode.doc_id == doc_id, SrsNode.table.isnot(None))
        ).scalars().all()
        changed = False
        for node in nodes:
            try:
                table = json.loads(node.table) if isinstance(node.table, str) else (node.table or {})
            except Exception:
                continue
            headers = table.get("headers") or []
            rows = table.get("rows") or []
            if not headers or not rows:
                continue
            header_norm = [self.__normalize_header((h or {}).get("name") or "") for h in headers]
            col_idx = self.__resolve_req_columns(header_norm)
            if "code" not in col_idx or not any(key in col_idx for key in ["module", "function", "sub_function", "location"]):
                continue
            header_codes = [(h or {}).get("code") or "" for h in headers]
            table_changed = False
            cells = table.get("cells") or []
            keep_indexes = []
            next_rows = []
            for row_idx, table_row in enumerate(rows):
                code = self.__normalize_srs_code(str((table_row or {}).get(header_codes[col_idx["code"]], "") or ""))
                req_item = req_map.get(code)
                if not req_item:
                    keep_indexes.append(row_idx)
                    next_rows.append(table_row)
                    continue
                keep_indexes.append(row_idx)
                next_rows.append(table_row)
                for field in ["module", "function", "sub_function", "location"]:
                    if field not in col_idx:
                        continue
                    value = req_item.get(field)
                    if value is None:
                        continue
                    col_code = header_codes[col_idx[field]]
                    if table_row.get(col_code) != value:
                        table_row[col_code] = value
                        table_changed = True
                    cell_row_idx = row_idx + 1
                    cell_col_idx = col_idx[field]
                    if (
                        isinstance(cells, list) and
                        cell_row_idx < len(cells) and
                        isinstance(cells[cell_row_idx], list) and
                        cell_col_idx < len(cells[cell_row_idx]) and
                        isinstance(cells[cell_row_idx][cell_col_idx], dict) and
                        cells[cell_row_idx][cell_col_idx].get("value") != value
                    ):
                        cells[cell_row_idx][cell_col_idx]["value"] = value
                        table_changed = True
            if table_changed:
                table["rows"] = next_rows
                if isinstance(cells, list) and cells:
                    table["cells"] = [cells[0], *[cells[idx + 1] for idx in keep_indexes if idx + 1 < len(cells)]]
            if table_changed:
                node.table = json.dumps(table, ensure_ascii=False)
                changed = True
        if changed:
            db.session.commit()

    def __parse_docx_table(self, tab, numbering_defs: Dict[str, dict] = None):
        # Parse table content and merged-cell structure from Word XML.
        tr_list = list(tab._tbl.tr_lst)  # type: ignore[attr-defined]
        if not tr_list:
            return None

        def grid_span(tc):
            try:
                gs = tc.tcPr.gridSpan  # type: ignore[attr-defined]
                if gs is not None and gs.val:
                    return int(gs.val)
            except Exception:
                pass
            return 1

        def v_merge(tc):
            try:
                vm = tc.tcPr.vMerge  # type: ignore[attr-defined]
                if vm is None:
                    return None
                val = vm.val
                return (str(val).lower() if val is not None else "continue")
            except Exception:
                return None

        def h_align(tc):
            try:
                first_p = tc.p_lst[0] if tc.p_lst else None  # type: ignore[attr-defined]
                jc = first_p.pPr.jc if first_p is not None and first_p.pPr is not None else None
                if jc is None or jc.val is None:
                    return "left"
                val = str(jc.val).lower()
                if "center" in val:
                    return "center"
                if "right" in val:
                    return "right"
                return "left"
            except Exception:
                return "left"

        def v_align(tc):
            try:
                va = tc.tcPr.vAlign  # type: ignore[attr-defined]
                if va is None or va.val is None:
                    return "top"
                val = str(va.val).lower()
                if "center" in val:
                    return "middle"
                if "bottom" in val:
                    return "bottom"
                return "top"
            except Exception:
                return "top"

        def cell_text(tc):
            counters: Dict[str, dict] = {}
            lines = []
            for p in getattr(tc, "p_lst", []) or []:  # type: ignore[attr-defined]
                try:
                    para = Paragraph(p, getattr(tab, "_parent", None))
                    txt = self.__paragraph_text_with_numbering(para, numbering_defs or {}, counters)
                except Exception:
                    txt = self.__normalize_text(getattr(p, "text", "") or "")
                if txt:
                    lines.append(txt)
            return self.__normalize_text("\n".join(lines))

        col_count = 0
        for tr in tr_list:
            count = 0
            for tc in tr.tc_lst:
                count += grid_span(tc)
            col_count = max(col_count, count)
        if col_count <= 0:
            return None

        cells: List[List[TableCell]] = []
        anchors: Dict[int, Tuple[int, int]] = {}
        for r_idx, tr in enumerate(tr_list):
            row_cells = [TableCell(value="", row_span=1, col_span=1) for _ in range(col_count)]
            c_idx = 0
            for tc in tr.tc_lst:
                while c_idx < col_count and row_cells[c_idx].row_span == 0:
                    c_idx += 1
                if c_idx >= col_count:
                    break
                span = max(1, grid_span(tc))
                text = cell_text(tc)
                cell_h_align = h_align(tc)
                cell_v_align = v_align(tc)
                vm = v_merge(tc)
                if vm == "continue":
                    touched = set()
                    for k in range(c_idx, min(col_count, c_idx + span)):
                        anchor = anchors.get(k)
                        if anchor and anchor not in touched:
                            ar, ac = anchor
                            cells[ar][ac].row_span = (cells[ar][ac].row_span or 1) + 1
                            touched.add(anchor)
                        row_cells[k] = TableCell(value="", row_span=0, col_span=0, h_align=cell_h_align, v_align=cell_v_align)
                else:
                    row_cells[c_idx] = TableCell(value=text, row_span=1, col_span=span, h_align=cell_h_align, v_align=cell_v_align)
                    for k in range(c_idx + 1, min(col_count, c_idx + span)):
                        row_cells[k] = TableCell(value="", row_span=0, col_span=0, h_align=cell_h_align, v_align=cell_v_align)
                    if vm == "restart":
                        for k in range(c_idx, min(col_count, c_idx + span)):
                            anchors[k] = (r_idx, c_idx)
                    else:
                        for k in range(c_idx, min(col_count, c_idx + span)):
                            anchors.pop(k, None)
                c_idx += span
            cells.append(row_cells)

        if not cells:
            return None

        header_row = cells[0]
        headers = [TabHeader(code=f"col_{idx+1}", name=(header_row[idx].value or f"列{idx+1}")) for idx in range(col_count)]
        rows = []
        for body_row in cells[1:]:
            row_obj = {}
            for idx in range(col_count):
                row_obj[f"col_{idx+1}"] = body_row[idx].value or ""
            if any(v for v in row_obj.values()):
                rows.append(row_obj)
        return Table(headers=headers, rows=rows, cells=cells)

    def __upsert_imported_srs_reqs(self, doc_id: int, req_rows: List[dict]):
        if not req_rows:
            return
        req_rows = self.__remap_change_req_types_by_existing_names(doc_id, req_rows)
        sql = select(SrsReq).where(SrsReq.doc_id == doc_id)
        exists = db.session.execute(sql).scalars().all()
        exists_dict = {(row.type_code, row.code): row for row in exists}

        change_type_names = {
            item.get("type_code"): self.__normalize_change_req_type_name(item.get("type_name") or "变更需求")
            for item in req_rows
            if item.get("type_code") not in ["1", "2", "", None] and item.get("type_name")
        }
        for type_code, type_name in change_type_names.items():
            type_row = db.session.execute(
                select(SrsType).where(SrsType.doc_id == doc_id, SrsType.type_code == type_code)
            ).scalars().first()
            if type_row:
                type_row.type_name = type_name
            else:
                db.session.add(SrsType(doc_id=doc_id, type_code=type_code, type_name=type_name))

        seen = set()
        added_change = 0
        skipped_change = 0
        for item in req_rows:
            key = (item.get("type_code") or "1", item.get("code") or "")
            type_code_for_log = key[0]
            is_change_type = type_code_for_log not in ["1", "2", ""]
            if not key[1] or key in seen:
                if is_change_type:
                    skipped_change += 1
                continue
            seen.add(key)
            row = exists_dict.get(key)
            if row:
                row.module = item.get("module")
                row.function = item.get("function")
                row.sub_function = item.get("sub_function")
                row.location = item.get("location")
                if row.type_code != "2":
                    chapter = row.sub_function or row.function or row.module or "/"
                    for trace_row in db.session.execute(select(SdsTrace).where(SdsTrace.req_id == row.id)).scalars().all():
                        trace_row.chapter = chapter
            else:
                row_data = {key: value for key, value in item.items() if key != "type_name"}
                db.session.add(SrsReq(doc_id=doc_id, **row_data))
                if is_change_type:
                    added_change += 1
        if any(str(item.get("type_code") or "") not in ["1", "2", ""] for item in req_rows or []):
            logger.info(
                "__upsert_imported_srs_reqs change: doc=%s req_rows=%d added=%d skipped=%d",
                doc_id, len(req_rows or []), added_change, skipped_change,
            )
        self.__merge_duplicate_change_req_types_by_name(doc_id)
        db.session.commit()

    def __upsert_imported_change_req_tables_from_tree(self, doc_id: int, nodes: List[SrsNodeForm]):
        """Word 导入兜底：扫描已入库的文档树，把所有"table.name 含变更"的需求表入到 SrsType + SrsReq。

        - 已存在 SrsType / SrsReq 原则上保留；但如果发现 srs_req 里同 code 行被错存为
          type_code='1'/'2'（标准/其他需求），会迁移到正确的变更类型，避免变更需求被当成
          产品需求展示和自动生成章节缺失。
        - 与主流程 `__extract_srs_reqs_from_nodes` 并行存在；用于覆盖前者因表头/路径异常漏识别的场景。
        - 不参与"标准需求(1) / 其他需求(2)"分类下的产品需求行；只针对"变更需求表"里出现的行。
        """
        if not doc_id or not nodes:
            return
        code_pattern = re.compile(r"^[A-Z]{2,}(?:[-_][A-Z0-9.]+)+$", re.I)
        existing_types: List[SrsType] = db.session.execute(
            select(SrsType).where(SrsType.doc_id == doc_id)
        ).scalars().all()
        type_code_by_name: Dict[str, str] = {}
        for type_row in existing_types:
            key = self.__normalize_change_req_type_key(type_row.type_name or "")
            if key and key not in type_code_by_name:
                type_code_by_name[key] = type_row.type_code

        existing_reqs: List[SrsReq] = db.session.execute(
            select(SrsReq).where(SrsReq.doc_id == doc_id)
        ).scalars().all()
        existing_req_set = {
            (row.type_code or "", self.__normalize_srs_code(row.code or "").upper())
            for row in existing_reqs
        }
        # 按 code 索引现有行（不区分 type_code），用于检测 type_code 错分类。
        existing_reqs_by_code: Dict[str, List[SrsReq]] = {}
        for row in existing_reqs:
            code_key = self.__normalize_srs_code(row.code or "").upper()
            if code_key:
                existing_reqs_by_code.setdefault(code_key, []).append(row)

        added_type = 0
        added_req = 0
        migrated_req = 0

        def migrate_misclassified(target_type_code: str, code_upper: str) -> bool:
            """把同 code 错存为 type_code='1'/'2' 的 srs_req 行迁移到 target_type_code。

            返回 True 表示已迁移（之后无需再 INSERT 新行）。
            """
            nonlocal migrated_req
            for candidate in existing_reqs_by_code.get(code_upper, []) or []:
                if (candidate.type_code or "") not in ("1", "2"):
                    continue
                # 同 (doc_id, target_type_code, code) 已经存在就不能再迁移，避免唯一约束冲突。
                if (target_type_code, code_upper) in existing_req_set:
                    return False
                old_type_code = candidate.type_code or ""
                candidate.type_code = target_type_code
                # SdsTrace 不依赖 type_code，但更新章节文本以反映新分类下的功能字段。
                chapter = candidate.sub_function or candidate.function or candidate.module or "/"
                for trace_row in db.session.execute(
                    select(SdsTrace).where(SdsTrace.req_id == candidate.id)
                ).scalars().all():
                    trace_row.chapter = chapter
                existing_req_set.discard((old_type_code, code_upper))
                existing_req_set.add((target_type_code, code_upper))
                migrated_req += 1
                logger.info(
                    "import change_req fallback migrate: doc_id=%s code=%s %s -> %s",
                    doc_id, code_upper, old_type_code, target_type_code,
                )
                return True
            return False

        def walk(items: List[SrsNodeForm]):
            nonlocal added_type, added_req
            for node in items or []:
                table = getattr(node, "table", None)
                raw_name = getattr(table, "name", "") if table else ""
                table_name = self.__normalize_change_req_type_name(raw_name or "")
                if table and "变更" in table_name:
                    headers = getattr(table, "headers", None) or []
                    rows = getattr(table, "rows", None) or []
                    if headers and rows:
                        header_names = [self.__normalize_text(getattr(h, "name", "") or "") for h in headers]
                        header_norm = [self.__normalize_header(item) for item in header_names]
                        col_idx = self.__resolve_req_columns(header_norm)
                        if "code" in col_idx:
                            type_key = self.__normalize_change_req_type_key(table_name)
                            type_code = type_code_by_name.get(type_key)
                            if not type_code:
                                digest = hashlib.md5(table_name.encode("utf-8")).hexdigest()[:12]
                                type_code = f"change_{digest}"
                                db.session.add(SrsType(doc_id=doc_id, type_code=type_code, type_name=table_name))
                                type_code_by_name[type_key] = type_code
                                added_type += 1
                            col_codes = [getattr(h, "code", "") for h in headers]
                            last_values: Dict[str, str] = {}
                            for row in rows:
                                values = [self.__normalize_text(str((row or {}).get(code, "") or "")) for code in col_codes]
                                code = values[col_idx["code"]] if col_idx["code"] < len(values) else ""
                                code = self.__normalize_srs_code(code)
                                if not code_pattern.match(code or ""):
                                    continue
                                code_upper = code.upper()
                                if (type_code, code_upper) in existing_req_set:
                                    continue
                                if migrate_misclassified(type_code, code_upper):
                                    continue
                                values = self.__fill_req_table_merged_values(values, col_idx, last_values)
                                db.session.add(SrsReq(
                                    doc_id=doc_id,
                                    code=code_upper,
                                    type_code=type_code,
                                    module=(self.__clean_req_table_field(values[col_idx["module"]]) if "module" in col_idx and col_idx["module"] < len(values) else None),
                                    function=(self.__clean_req_table_field(values[col_idx["function"]]) if "function" in col_idx and col_idx["function"] < len(values) else None),
                                    sub_function=(self.__clean_req_table_field(values[col_idx["sub_function"]]) if "sub_function" in col_idx and col_idx["sub_function"] < len(values) else None),
                                    location=(self.__clean_req_table_field(values[col_idx["location"]]) if "location" in col_idx and col_idx["location"] < len(values) else None),
                                ))
                                existing_req_set.add((type_code, code_upper))
                                added_req += 1
                walk(getattr(node, "children", None) or [])

        walk(nodes)
        # 兜底建表后清理“变更类型空壳”：模板占位/无合规需求行的“变更”表节点会被建出 type 但 0 条
        # srs_req，这类 change_ 前缀空壳会被前端渲染成“暂无数据”空表，且每次保存都重建。
        # 这里统一清掉本文档下无任何 srs_req 的 change_ 空壳；有数据的导入表(req>0)不受影响，
        # 用户手动新增的变更表 type_code 为 uuid（不以 change_ 开头）也不受影响。
        db.session.flush()
        removed_shell_ids: List[int] = []
        shell_types: List[SrsType] = db.session.execute(
            select(SrsType).where(
                SrsType.doc_id == doc_id,
                SrsType.type_code.like("change_%"),
            )
        ).scalars().all()
        for shell in shell_types:
            shell_req_count = db.session.execute(
                select(func.count()).select_from(SrsReq).where(
                    SrsReq.doc_id == doc_id,
                    SrsReq.type_code == shell.type_code,
                )
            ).scalars().first()
            if not shell_req_count and shell.id:
                db.session.delete(shell)
                removed_shell_ids.append(shell.id)
        if added_type or added_req or migrated_req or removed_shell_ids:
            db.session.commit()
            logger.info(
                "import change_req fallback: doc_id=%s added_type=%d added_req=%d migrated_req=%d removed_shells=%s",
                doc_id, added_type, added_req, migrated_req, sorted(removed_shell_ids),
            )

    def __sync_srs_reqs_from_doc_tables(self, doc_id: int, nodes: List[SrsNodeForm]):
        req_rows, req_rcm_map = self.__extract_srs_reqs_from_nodes(nodes or [], include_node_codes=False)

        def collect_managed_type_codes(items: List[SrsNodeForm]):
            type_codes = set()
            for node in items or []:
                table = getattr(node, "table", None)
                headers = getattr(table, "headers", None) if table else None
                rows = getattr(table, "rows", None) if table else None
                if headers is not None and rows is not None:
                    header_names = [self.__normalize_text(getattr(h, "name", "") or "") for h in headers]
                    header_norm = [self.__normalize_header(h) for h in header_names]
                    col_idx = self.__resolve_req_columns(header_norm)
                    has_product_cols = ("code" in col_idx and "module" in col_idx and "function" in col_idx)
                    has_other_cols = ("code" in col_idx and "module" in col_idx and "location" in col_idx)
                    if has_product_cols or has_other_cols:
                        type_code, _ = self.__resolve_doc_table_req_type(
                            node,
                            has_other_cols,
                            "function" in col_idx,
                            "sub_function" in col_idx,
                        )
                        if type_code in ["1", "2"]:
                            type_codes.add(type_code)
                type_codes.update(collect_managed_type_codes(getattr(node, "children", None) or []))
            return type_codes

        normalized_rows = []
        seen = set()
        for item in req_rows:
            type_code = str((item or {}).get("type_code") or "1")
            # 变更需求以 srs_req 表编辑入口为准，文档树里的展示表不能反向覆盖已入库数据。
            if type_code not in ["1", "2"]:
                continue
            code = self.__normalize_srs_code(str((item or {}).get("code") or "")).upper()
            if not code:
                continue
            key = (type_code, code)
            if key in seen:
                continue
            seen.add(key)
            normalized_rows.append({
                **item,
                "type_code": type_code,
                "code": code,
            })

        managed_type_codes = sorted({*collect_managed_type_codes(nodes or []), *{item["type_code"] for item in normalized_rows}})
        if not managed_type_codes:
            return

        change_type_names = {
            item["type_code"]: self.__normalize_change_req_type_name(item.get("type_name") or "变更需求")
            for item in normalized_rows
            if item["type_code"] not in ["1", "2", ""] and item.get("type_name")
        }
        for type_code, type_name in change_type_names.items():
            type_row = db.session.execute(
                select(SrsType).where(SrsType.doc_id == doc_id, SrsType.type_code == type_code)
            ).scalars().first()
            if type_row:
                type_row.type_name = type_name
            else:
                db.session.add(SrsType(doc_id=doc_id, type_code=type_code, type_name=type_name))
        if change_type_names:
            # 兼容已被旧逻辑误保存为标准需求(type_code=1)的变更表行，保留原 req_id 和追溯关系。
            managed_type_codes = sorted({*managed_type_codes, "1"})

        existing: List[SrsReq] = db.session.execute(
            select(SrsReq).where(SrsReq.doc_id == doc_id, SrsReq.type_code.in_(managed_type_codes))
        ).scalars().all()
        exists_dict = {(row.type_code, row.code): row for row in existing}
        current_keys = {(item["type_code"], item["code"]) for item in normalized_rows}

        def logical_key(obj):
            if isinstance(obj, dict):
                type_code = str(obj.get("type_code") or "1")
                values = [obj.get("module"), obj.get("function"), obj.get("sub_function")]
                if type_code == "2":
                    values.append(obj.get("location"))
            else:
                type_code = str(getattr(obj, "type_code", "") or "1")
                values = [getattr(obj, "module", None), getattr(obj, "function", None), getattr(obj, "sub_function", None)]
                if type_code == "2":
                    values.append(getattr(obj, "location", None))
            normalized = tuple(self.__normalize_text(str(value or "")) for value in values)
            return (type_code, *normalized)

        # 编号变更不代表需求换位置：相同模块/功能/子功能的行视为同一需求，保留 req_id 及 SDS 追溯关系。
        logical_candidates: Dict[tuple, List[SrsReq]] = {}
        for row in existing:
            logical_candidates.setdefault(logical_key(row), []).append(row)

        def sync_sds_trace_and_nodes(req_row: SrsReq, old_sds_code: str = None):
            desired_sds_code = (req_row.code or "").replace("SRS", "SDS")
            chapter = req_row.sub_function or req_row.function or req_row.module or "/"
            old_sds_code = (old_sds_code or "").strip()
            for trace_row in db.session.execute(select(SdsTrace).where(SdsTrace.req_id == req_row.id)).scalars().all():
                previous_code = str(trace_row.sds_code or "").strip()
                if desired_sds_code:
                    trace_row.sds_code = desired_sds_code
                trace_row.chapter = chapter
                match_codes = {
                    re.sub(r"\s+", "", old_sds_code.upper()),
                    re.sub(r"\s+", "", previous_code.upper()),
                    re.sub(r"\s+", "", desired_sds_code.upper()),
                }
                match_codes = {code for code in match_codes if code}
                nodes = db.session.execute(
                    select(SdsNode).where(SdsNode.doc_id == trace_row.doc_id, SdsNode.sds_code.isnot(None))
                ).scalars().all()
                for node in nodes:
                    node_code = re.sub(r"\s+", "", str(getattr(node, "sds_code", "") or "").strip().upper())
                    if node_code not in match_codes:
                        continue
                    node.sds_code = desired_sds_code
                    if trace_row.location and chapter:
                        node.title = f"{trace_row.location} {chapter}"

        used_renamed_ids = set()
        for item in normalized_rows:
            key = (item["type_code"], item["code"])
            if key in exists_dict:
                continue
            same_code_candidates = [
                row for row in existing
                if row.id not in used_renamed_ids
                and row.code == item["code"]
                and row.type_code != item["type_code"]
                and (row.type_code, row.code) not in current_keys
            ]
            if len(same_code_candidates) == 1:
                row = same_code_candidates[0]
                old_key = (row.type_code, row.code)
                row.type_code = item["type_code"]
                row.module = item.get("module")
                row.function = item.get("function")
                row.sub_function = item.get("sub_function")
                row.location = item.get("location")
                exists_dict.pop(old_key, None)
                exists_dict[key] = row
                used_renamed_ids.add(row.id)
                sync_sds_trace_and_nodes(row)
                continue
            candidates = logical_candidates.get(logical_key(item)) or []
            candidates = [
                row for row in candidates
                if row.id not in used_renamed_ids and (row.type_code, row.code) not in current_keys
            ]
            if len(candidates) != 1:
                continue
            row = candidates[0]
            old_sds_code = (row.code or "").replace("SRS", "SDS")
            new_sds_code = (item["code"] or "").replace("SRS", "SDS")
            row.code = item["code"]
            row.module = item.get("module")
            row.function = item.get("function")
            row.sub_function = item.get("sub_function")
            row.location = item.get("location")
            exists_dict[key] = row
            used_renamed_ids.add(row.id)
            sync_sds_trace_and_nodes(row, old_sds_code=old_sds_code)

        delete_ids = [row.id for row in existing if (row.type_code, row.code) not in current_keys]

        if delete_ids:
            db.session.execute(delete(ReqRcm).where(ReqRcm.req_id.in_(delete_ids)))
            db.session.execute(delete(SrsReqd).where(SrsReqd.req_id.in_(delete_ids)))
            db.session.execute(delete(SdsReqd).where(SdsReqd.req_id.in_(delete_ids)))
            db.session.execute(delete(SdsTrace).where(SdsTrace.req_id.in_(delete_ids)))
            db.session.execute(delete(SrsReq).where(SrsReq.id.in_(delete_ids)))

        for item in normalized_rows:
            key = (item["type_code"], item["code"])
            row = exists_dict.get(key)
            if row:
                row.module = item.get("module")
                row.function = item.get("function")
                row.sub_function = item.get("sub_function")
                row.location = item.get("location")
                if row.type_code != "2":
                    sync_sds_trace_and_nodes(row)
            else:
                row_data = {key: value for key, value in item.items() if key != "type_name"}
                db.session.add(SrsReq(doc_id=doc_id, **row_data))
        self.__merge_duplicate_change_req_types_by_name(doc_id)
        db.session.commit()
        self.__sync_imported_req_rcms(doc_id, req_rcm_map)

    def __sync_imported_req_rcms(self, doc_id: int, req_rcm_map: Dict[str, set]):
        if not req_rcm_map:
            return
        normalized_req_rcm_map = {}
        for req_code, rcm_codes in (req_rcm_map or {}).items():
            normalized_req_code = self.__normalize_srs_code(req_code)
            normalized_rcm_codes = {
                self.__normalize_rcm_code(rcm_code)
                for rcm_code in (rcm_codes or set())
                if self.__normalize_rcm_code(rcm_code)
            }
            if normalized_req_code and normalized_rcm_codes:
                normalized_req_rcm_map.setdefault(normalized_req_code, set()).update(normalized_rcm_codes)
        req_codes = [code for code in normalized_req_rcm_map.keys() if code]
        if not req_codes:
            return
        req_rows = db.session.execute(select(SrsReq).where(SrsReq.doc_id == doc_id, SrsReq.code.in_(req_codes))).scalars().all()
        if not req_rows:
            return
        req_ids = [row.id for row in req_rows]
        db.session.execute(delete(ReqRcm).where(ReqRcm.req_id.in_(req_ids)))

        all_rcm_codes = sorted({code for codes in normalized_req_rcm_map.values() for code in codes})
        if not all_rcm_codes:
            db.session.commit()
            return
        rcm_rows = db.session.execute(select(Rcm)).scalars().all()
        rcm_id_dict = {self.__normalize_rcm_code(row.code): row.id for row in rcm_rows}
        insert_values = []
        for req_row in req_rows:
            for rcm_code in sorted(normalized_req_rcm_map.get(req_row.code, set())):
                rcm_id = rcm_id_dict.get(rcm_code)
                if rcm_id:
                    insert_values.append(dict(req_id=req_row.id, rcm_id=rcm_id))
        if insert_values:
            db.session.execute(pg_insert(ReqRcm).values(insert_values).on_conflict_do_nothing())
        db.session.commit()

    def __parse_docx_content(self, docx: Document, bind_block_lead: bool = False):
        roots: List[SrsNodeForm] = []
        stack: List[Tuple[int, SrsNodeForm]] = []
        current: SrsNodeForm = None
        heading_rows = []
        srs_pattern = re.compile(r"\bSRS[-_A-Za-z0-9.]+\b", re.I)
        rcm_pattern = re.compile(r"\bRCM[-_A-Za-z0-9]+\b", re.I)
        mime_map = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg", "gif": "image/gif", "bmp": "image/bmp", "webp": "image/webp"}
        img_idx = 0
        table_idx = 0
        heading_counters = [0, 0, 0, 0, 0]
        numbering_defs = self.__build_numbering_definitions(docx)
        body_numbering_counters: Dict[str, dict] = {}
        pending_table_title = ""
        # 仅 SDS 导入启用：暂存「引导句段落」（如「功能清单：」），绑定到紧跟其后的表/图节点上方，
        # 而不并入父章节正文。SRS 导入不开启，行为完全不变。
        pending_block_lead = None  # Optional[Tuple[holder_node, lead_text]]

        def ensure_text_holder():
            nonlocal current
            if current is None:
                current = SrsNodeForm(title="导入正文", text="", children=[])
                roots.append(current)
            if current.children is None:
                current.children = []
            return current

        def attach_to_current(node: SrsNodeForm):
            if current:
                current.children = current.children or []
                current.children.append(node)
                img_url = getattr(node, "img_url", None)
                if img_url and current:
                    heading_no = self.__heading_no_from_title(getattr(current, "title", "") or "")
                    if heading_no == "2.2":
                        current.ref_type = RefTypes.img_topo.value
                        current.img_url = img_url
                    elif heading_no == "2.3":
                        current.ref_type = RefTypes.img_struct.value
                        current.img_url = img_url
            else:
                roots.append(node)

        def attach_node(level: int, node: SrsNodeForm):
            nonlocal current
            while stack and stack[-1][0] >= level:
                stack.pop()
            if stack:
                parent = stack[-1][1]
                parent.children = parent.children or []
                parent.children.append(node)
            else:
                roots.append(node)
            stack.append((level, node))
            current = node

        def extract_images_from_para(para: Paragraph):
            return self.__extract_images_from_ooxml_element(para.part, para._element)

        def sync_counters_with_number(number_text: str):
            try:
                parts = [int(p) for p in str(number_text or "").split(".") if str(p).strip()]
            except Exception:
                return None
            if not parts:
                return None
            depth = min(len(parts), 5)
            for idx in range(depth):
                heading_counters[idx] = max(0, parts[idx])
            for idx in range(depth, 5):
                heading_counters[idx] = 0
            return ".".join(str(v) for v in heading_counters[:depth] if v > 0)

        def build_number_from_level(level: int):
            depth = max(1, min(int(level or 1), 5))
            for idx in range(depth - 1):
                if heading_counters[idx] <= 0:
                    heading_counters[idx] = 1
            heading_counters[depth - 1] = heading_counters[depth - 1] + 1 if heading_counters[depth - 1] > 0 else 1
            for idx in range(depth, 5):
                heading_counters[idx] = 0
            return ".".join(str(v) for v in heading_counters[:depth] if v > 0)

        def is_table_title_text(text: str):
            return self.__is_req_table_title_text(text)

        def is_block_lead(text: str) -> bool:
            # 引导句：短、以冒号结尾、非完整句、非表标题（如「功能清单：」「衔接处理任务流程：」）。
            if not bind_block_lead:
                return False
            s = self.__normalize_text(text or "").strip()
            if not s or not (s.endswith("：") or s.endswith(":")):
                return False
            core = s[:-1].strip()
            if not core or len(core) > 15:
                return False
            if re.search(r"[。！？；.!?;]$", core):
                return False
            if re.search(r"https?://|/[\w\-]", core, re.I):
                return False
            if is_table_title_text(s):
                return False
            return True

        def consume_block_lead(target_node: SrsNodeForm):
            # 把暂存引导句从父正文末尾移除，并置于目标表/图节点正文最前（显示在块上方）。
            nonlocal pending_block_lead
            if not pending_block_lead:
                return
            holder, lead = pending_block_lead
            pending_block_lead = None
            if holder is not None and getattr(holder, "text", ""):
                lines = holder.text.split("\n")
                while lines and lines[-1].strip() == "":
                    lines.pop()
                if lines and lines[-1].strip() == str(lead).strip():
                    lines.pop()
                    holder.text = "\n".join(lines).strip()
            existing = getattr(target_node, "text", "") or ""
            target_node.text = f"{lead}\n{existing}".strip() if existing else str(lead)

        def extract_table_titles_from_text(text: str):
            titles = []
            for line in str(text or "").replace("\r", "").split("\n"):
                clean = self.__normalize_text(line).replace("：", ":").rstrip(":").strip()
                if clean and is_table_title_text(clean):
                    titles.append(clean)
            return titles

        def apply_table_titles_from_parent_text(nodes: List[SrsNodeForm]):
            for node in nodes or []:
                children = getattr(node, "children", None) or []
                table_children = [
                    child for child in children
                    if getattr(child, "table", None) is not None and re.match(r"^导入表格\d*$", getattr(child, "title", "") or "")
                ]
                titles = extract_table_titles_from_text(getattr(node, "text", "") or "")
                for child, title in zip(table_children, titles):
                    child.table.name = title
                apply_table_titles_from_parent_text(children)

        for child in docx.element.body.iterchildren():
            tag = str(child.tag).lower()
            if tag.endswith("}p"):
                para = Paragraph(child, docx._body)
                txt = self.__normalize_text(para.text)
                if txt and self.__is_toc_paragraph(para):
                    continue
                numpr_level = self.__guess_numpr_level(para) if txt else None
                level = self.__guess_heading_level(para) if txt else None
                # 在已进入任一章节（1/2/3...级）后，"1. xxx / 2. xxx" 这类枚举项按正文处理，不识别为标题。
                if txt and level is not None and stack:
                    is_enum_item = bool(
                        re.match(r"^\d+[.．、]\s+\S+", txt)
                        and not re.match(r"^\d+\.\d+", txt)
                    )
                    if is_enum_item:
                        heading_number_for_enum = self.__extract_heading_number(txt)
                        is_next_root_heading = False
                        if heading_number_for_enum and "." not in heading_number_for_enum:
                            try:
                                heading_num = int(heading_number_for_enum)
                                # Word 中一级标题有时写成“4. 图像接收”。若序号正好承接当前一级章节，
                                # 应保留为章节，否则才按正文枚举项处理。
                                is_next_root_heading = heading_num == heading_counters[0] + 1
                            except Exception:
                                is_next_root_heading = False
                        if not is_next_root_heading:
                            level = None
                # 兼容“接口章节下的无编号三级标题”：
                # 在父级为“x.x 接口”时，将“以‘接口’结尾”的短行识别为下一层级标题（如：数据上传接口、创建处理任务接口）。
                if txt and level is None and stack:
                    parent_level, parent_node = stack[-1]
                    parent_title = self.__normalize_text(getattr(parent_node, "title", ""))
                    is_interface_parent = bool(
                        re.match(r"^\d+(?:\.\d+)+\s*接口$", parent_title)
                        or parent_title == "接口"
                    )
                    is_interface_subtitle = bool(
                        re.search(r"接口$", txt)
                        and len(txt) <= 80
                        and not re.search(r"[。！？；;]$", txt)
                        and not re.search(r"https?://|/[\w\-]+", txt, re.I)
                    )
                    if is_interface_parent and is_interface_subtitle:
                        level = min(parent_level + 1, 5)
                # 引导句后若紧跟的是标题/表标题而非表/图，则放弃绑定，保持其留在父正文（行为不变）。
                if bind_block_lead and txt and level is not None and pending_block_lead is not None:
                    pending_block_lead = None
                # 需求表标题（如 2.0 变更需求）仅作后续表格名，不生成文档章节节点。
                if txt and level is not None and is_table_title_text(txt):
                    pending_table_title = txt
                    level = None
                if txt and level is not None:
                    heading_number = self.__extract_heading_number(txt)
                    title_with_number = txt
                    if heading_number:
                        synced = sync_counters_with_number(heading_number)
                        heading_number = synced or heading_number
                    elif numpr_level is not None:
                        generated_number = build_number_from_level(level)
                        if generated_number:
                            heading_number = generated_number
                            title_with_number = f"{generated_number} {txt}".strip()
                    node = SrsNodeForm(title=title_with_number, text="", children=[])
                    heading_rows.append(dict(level=level, title=title_with_number, number=heading_number))
                    if heading_number == "2.2":
                        node.ref_type = RefTypes.img_topo.value
                    elif heading_number == "2.3":
                        node.ref_type = RefTypes.img_struct.value
                    srs_hit = srs_pattern.search(txt)
                    if srs_hit:
                        node.srs_code = srs_hit.group(0).upper()
                    attach_node(level, node)
                elif txt:
                    body_txt = self.__paragraph_text_with_numbering(para, numbering_defs, body_numbering_counters)
                    if is_table_title_text(body_txt):
                        pending_table_title = body_txt
                        pending_block_lead = None
                    else:
                        pending_table_title = ""
                        holder = ensure_text_holder()
                        holder.text = f"{holder.text}\n{body_txt}".strip() if holder.text else body_txt
                        rcm_codes = {self.__normalize_rcm_code(item) for item in rcm_pattern.findall(body_txt)}
                        rcm_codes = {code for code in rcm_codes if code}
                        if rcm_codes:
                            existed = set(self.__normalize_rcm_codes(holder.rcm_codes or []))
                            holder.rcm_codes = sorted(existed.union(rcm_codes))
                        srs_hit = srs_pattern.search(body_txt)
                        if srs_hit and not holder.srs_code:
                            holder.srs_code = srs_hit.group(0).upper()
                        # 引导句：暂存，等待紧跟的表/图来消费；否则保持在正文。
                        pending_block_lead = (holder, body_txt) if is_block_lead(body_txt) else None

                for img_url in extract_images_from_para(para):
                    img_idx += 1
                    img_node = SrsNodeForm(title=f"导入图片{img_idx}", img_url=img_url, children=[])
                    attach_to_current(img_node)
                    consume_block_lead(img_node)
            elif tag.endswith("}tbl"):
                tab = DocxTable(child, docx._body)
                for row in tab.rows:
                    for cell in row.cells:
                        for img_url in self.__extract_images_from_ooxml_element(docx.part, cell._tc):
                            img_idx += 1
                            attach_to_current(SrsNodeForm(title=f"导入图片{img_idx}", img_url=img_url, children=[]))
                table = self.__parse_docx_table(tab, numbering_defs)
                if table is None or not table.headers:
                    continue
                table_idx += 1
                table_title = pending_table_title or f"导入表格{table_idx}"
                pending_table_title = ""
                table.name = table_title if not re.match(r"^导入表格\d*$", table_title or "") else None
                table_node = SrsNodeForm(title=f"导入表格{table_idx}", table=table, children=[])
                consume_block_lead(table_node)
                attach_to_current(table_node)
        apply_table_titles_from_parent_text(roots)
        return roots, heading_rows

    async def import_srs_doc_word(self, product_id: int, version: str, change_log: str, file):
        if Document is None or DocxTable is None or Paragraph is None:
            return Resp.resp_err(msg="当前环境缺少 python-docx 依赖，暂不可用 Word 导入。")
        try:
            bys = await file.read()
            docx = Document(io.BytesIO(bys))
            file_name = file.filename or ""
            folder_name, file_no = self.__extract_file_info(file_name)
            content, heading_rows = self.__parse_docx_content(docx)
            heading_err = self.__validate_heading_numbers(heading_rows)
            if heading_err:
                logger.warning("word heading validation warning (ignored): %s", heading_err)

            # 先基于“已解析节点表格”抽取，避免不同Word表格结构导致漏识别；再与原始docx抽取结果合并去重
            srs_req_rows_nodes, req_rcm_map_nodes = self.__extract_srs_reqs_from_nodes(content)
            srs_req_rows_docx, req_rcm_map_docx = self.__extract_srs_reqs_from_tables(docx)
            srs_reqd_rows_nodes = self.__extract_srs_reqds_from_nodes(content)
            srs_req_rows = []
            req_row_dict = {}

            def req_row_score(item: dict):
                return sum(
                    1
                    for field in ["module", "function", "sub_function", "location"]
                    if self.__normalize_text(str((item or {}).get(field) or "")).strip() not in ["", "/", "\\", "／", "＼"]
                )

            for item in [*(srs_req_rows_nodes or []), *(srs_req_rows_docx or [])]:
                key = ((item or {}).get("type_code") or "1", (item or {}).get("code") or "")
                if not key[1]:
                    continue
                existed = req_row_dict.get(key)
                if not existed:
                    req_row_dict[key] = dict(item)
                    continue
                if req_row_score(item) > req_row_score(existed):
                    req_row_dict[key] = {**existed, **item}
                else:
                    for field in ["module", "function", "sub_function", "location"]:
                        if not self.__normalize_text(str(existed.get(field) or "")) and self.__normalize_text(str((item or {}).get(field) or "")):
                            existed[field] = item.get(field)
            srs_req_rows = list(req_row_dict.values())
            req_rcm_map = {}
            for req_map in [req_rcm_map_nodes or {}, req_rcm_map_docx or {}]:
                for req_code, rcm_set in req_map.items():
                    if not req_code:
                        continue
                    req_rcm_map.setdefault(req_code, set()).update(rcm_set or set())
            form = SrsDocForm(
                product_id=product_id,
                version=version,
                folder_name=folder_name or None,
                file_no=file_no or None,
                change_log=change_log,
                content=content,
            )
            def _is_change_type_code(value) -> bool:
                return str(value or "") not in ["1", "2", "", "None", "none", "null"]
            change_req_row_count = sum(1 for item in srs_req_rows if _is_change_type_code((item or {}).get("type_code")))
            change_from_nodes = sum(1 for item in srs_req_rows_nodes or [] if _is_change_type_code((item or {}).get("type_code")))
            change_from_docx = sum(1 for item in srs_req_rows_docx or [] if _is_change_type_code((item or {}).get("type_code")))
            logger.info(
                "import_srs_doc_word: req_rows=%d (change=%d, change_from_nodes=%d, change_from_docx=%d), reqd_rows=%d",
                len(srs_req_rows or []),
                change_req_row_count,
                change_from_nodes,
                change_from_docx,
                len(srs_reqd_rows_nodes or []),
            )
            for item in srs_req_rows:
                if _is_change_type_code((item or {}).get("type_code")):
                    logger.info(
                        "import change row: type=%s code=%s type_name=%s module=%s function=%s",
                        (item or {}).get("type_code"),
                        (item or {}).get("code"),
                        (item or {}).get("type_name"),
                        (item or {}).get("module"),
                        (item or {}).get("function"),
                    )

            resp = await self.add_srs_doc(form)
            if resp.code == 200 and resp.data and resp.data.id:
                self.__upsert_imported_srs_reqs(resp.data.id, srs_req_rows)
                # 兜底：再次按"已入库的文档树节点"扫描一次变更需求表，确保 Word 自带的
                # "X.X 变更需求"表行被入到 SrsType+SrsReq，不依赖前面解析路径是否漏识别。
                self.__upsert_imported_change_req_tables_from_tree(resp.data.id, content)
                self.__sync_saved_doc_srs_tables_from_req_rows(resp.data.id)
                self.__sync_imported_req_rcms(resp.data.id, req_rcm_map)
                self.__upsert_imported_srs_reqds(resp.data.id, srs_reqd_rows_nodes)
                # 根据导入 Word 中的 2.2/2.3 图片，按产品完整版本+文档版本覆盖图表文件库
                product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
                product_version = getattr(product, "full_version", "") or getattr(product, "name", "") or ""
                self.__auto_sync_product_doc_images(
                    product_id,
                    content,
                    version,
                    docx,
                    product_version,
                    bys,
                    resp.data.id,
                )
                row = db.session.execute(select(SrsDoc).where(SrsDoc.id == resp.data.id)).scalars().first()
                if row:
                    self.__fix_rcms(row)
            return resp
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    def __update_nodes(self, doc: SrsDoc, p_id, nodes: List[SrsNodeForm]):
        for idx, node in enumerate(nodes or []):
            sql = select(SrsNode).where(SrsNode.doc_id == doc.id, SrsNode.n_id == node.n_id) if node.n_id else None
            row = db.session.execute(sql).scalars().first() if sql is not None else None
            if not row:
                doc.n_id += 1
                table = json.loads(node.table.json()) if node.table else None
                row = SrsNode(doc_id=doc.id, n_id=doc.n_id, p_id=p_id, priority=idx, title=node.title, label=node.label, img_url=node.img_url, text=node.text, ref_type=node.ref_type,
                            table=table, srs_code=node.srs_code)
                row.rcm_codes = ",".join(node.rcm_codes) if node.rcm_codes is not None else None
                db.session.add(row)
                logger.info("add_node: %s, %s, %s", p_id, doc.n_id, node.title)
            else:
                for key, value in node.dict().items():
                    if key == "doc_id" or key == "n_id" or key == "p_id" or value is None:
                        continue
                    if key == "table":
                        value = json.dumps(value) if value else None
                    setattr(row, key, value)
                row.priority = idx
                logger.info("alt_node: %s, %s, %s", p_id, doc.n_id, node.title)
            if node.children:
                self.__update_nodes(doc, row.n_id, node.children)

    def __reset_tree_node_ids(self, nodes: List[SrsNodeForm]):
        # update_srs_doc 采用“全量重建”策略，需清空前端携带的旧 n_id，
        # 否则新节点与旧节点 n_id 冲突时会在同一轮重建中被覆盖。
        for node in nodes or []:
            node.n_id = None
            if node.children:
                self.__reset_tree_node_ids(node.children)

    def __sync_change_req_tables_from_db(
        self, doc_id: int, nodes: List[SrsNodeForm], *, restructure: bool = False
    ):
        """将已有变更需求表同步为数据库最新行；restructure=True 时才补齐缺失表并调整节点位置。"""
        req_rows: List[SrsReq] = db.session.execute(
            select(SrsReq).where(
                SrsReq.doc_id == doc_id,
                SrsReq.type_code.isnot(None),
                SrsReq.type_code.notin_(["1", "2"]),
            ).order_by(SrsReq.id)
        ).scalars().all()
        if not req_rows:
            return False

        type_rows: List[SrsType] = db.session.execute(
            select(SrsType).where(SrsType.doc_id == doc_id)
        ).scalars().all()
        type_name_by_code = {row.type_code: row.type_name for row in type_rows if row.type_code}
        reqs_by_type: Dict[str, List[SrsReq]] = {}
        for row in req_rows:
            reqs_by_type.setdefault(row.type_code or "", []).append(row)
        changed = False

        def normalize_change_table_node_titles(items: List[SrsNodeForm]):
            nonlocal changed
            for item in items or []:
                table = getattr(item, "table", None)
                node_title = str(getattr(item, "title", "") or "").strip()
                table_name = getattr(table, "name", None) if table else None
                resolved_title = table_name or node_title or getattr(item, "label", "") or ""
                if (
                    table is not None
                    and "变更" in str(resolved_title or "")
                    and node_title
                    and not re.match(r"^导入表格\d*$", node_title)
                    and (
                        self.__change_table_title_matches(node_title, resolved_title)
                        or re.match(r"^\d+(?:\.\d+)*\s+\S*变更", node_title)
                    )
                ):
                    item.title = ""
                    if not table_name:
                        table.name = resolved_title
                    if not getattr(item, "label", None):
                        item.label = resolved_title
                    changed = True
                normalize_change_table_node_titles(getattr(item, "children", None) or [])

        normalize_change_table_node_titles(nodes or [])

        def norm_text(value):
            return re.sub(r"\s+", "", str(value or "").replace("：", ":").rstrip(":").strip())

        def map_field(value):
            txt = norm_text(value).lower()
            if "需求编号" in txt or txt in ["srscode", "code"]:
                return "code"
            if "子功能" in txt:
                return "sub_function"
            if "功能" in txt:
                return "function"
            if "模块" in txt:
                return "module"
            return ""

        def node_change_title(item: SrsNodeForm):
            table = getattr(item, "table", None)
            table_name = getattr(table, "name", None) if table else None
            return table_name or getattr(item, "title", "") or getattr(item, "label", "") or ""

        def find_preferred_change_table_parent(items: List[SrsNodeForm]):
            for item in items or []:
                for child in getattr(item, "children", None) or []:
                    title = node_change_title(child)
                    if "变更" in str(title or "") and re.match(r"^导入表格\d*$", str(getattr(child, "title", "") or "").strip()):
                        return item
                found = find_preferred_change_table_parent(getattr(item, "children", None) or [])
                if found is not None:
                    return found
            return None

        preferred_target = find_preferred_change_table_parent(nodes or [])

        def is_generated_change_table_node(item: SrsNodeForm):
            if re.match(r"^导入表格\d*$", str(getattr(item, "title", "") or "").strip()):
                return False
            title = node_change_title(item)
            if "变更" not in str(title or ""):
                return False
            for type_code in reqs_by_type.keys():
                type_name = type_name_by_code.get(type_code) or "变更需求"
                if self.__change_table_title_matches(title, type_name):
                    return True
            return False

        if restructure and preferred_target is not None:
            moved_nodes: List[SrsNodeForm] = []

            def detach_misplaced(items: List[SrsNodeForm], under_preferred: bool = False):
                result = []
                for item in items or []:
                    is_preferred = item is preferred_target
                    item.children = detach_misplaced(
                        getattr(item, "children", None) or [],
                        under_preferred or is_preferred,
                    )
                    if is_generated_change_table_node(item) and not (under_preferred or is_preferred):
                        moved_nodes.append(item)
                        continue
                    result.append(item)
                return result

            nodes[:] = detach_misplaced(nodes or [])
            if moved_nodes:
                preferred_target.children = list(getattr(preferred_target, "children", None) or [])
                existing_under_target = {
                    norm_text(node_change_title(child))
                    for child in preferred_target.children or []
                }
                for moved in moved_nodes:
                    moved_title = norm_text(node_change_title(moved))
                    if moved_title not in existing_under_target:
                        preferred_target.children.append(moved)
                        existing_under_target.add(moved_title)
                changed = True

        existing_titles = set()

        def collect_change_table_titles(items: List[SrsNodeForm]):
            for item in items or []:
                table = getattr(item, "table", None)
                table_name = getattr(table, "name", None) if table else None
                title = table_name or getattr(item, "title", "") or getattr(item, "label", "") or ""
                if "变更" in str(title or ""):
                    existing_titles.add(norm_text(title))
                collect_change_table_titles(getattr(item, "children", None) or [])

        collect_change_table_titles(nodes or [])

        def build_change_table_rows(rows: List[SrsReq]):
            return [
                dict(
                    srs_code=req.code or "",
                    module=req.module or "",
                    function=req.function or "",
                    sub_function=req.sub_function or "",
                )
                for req in rows or []
            ]

        def apply_change_rows_to_node(item: SrsNodeForm, rows: List[SrsReq]) -> bool:
            table = getattr(item, "table", None)
            if table is None:
                return False
            preserved_name = getattr(table, "name", None)
            headers = [
                TabHeader(code="srs_code", name="需求编号"),
                TabHeader(code="module", name="模块"),
                TabHeader(code="function", name="功能"),
                TabHeader(code="sub_function", name="子功能"),
            ]
            table.headers = headers
            table.rows = build_change_table_rows(rows)
            if preserved_name:
                table.name = preserved_name
            if hasattr(table, "cells"):
                table.cells = None
            return True

        def sync_type_name_from_tree_item(item: SrsNodeForm, type_code: str):
            nonlocal changed
            tree_name = self.__normalize_change_req_type_name(node_change_title(item))
            if not tree_name or "变更" not in tree_name:
                return
            type_row = db.session.execute(
                select(SrsType).where(SrsType.doc_id == doc_id, SrsType.type_code == type_code)
            ).scalars().first()
            if not type_row:
                return
            current_name = self.__normalize_change_req_type_name(type_row.type_name or "")
            if tree_name != current_name:
                type_row.type_name = tree_name
                changed = True

        def sync_existing_change_tables(items: List[SrsNodeForm]):
            nonlocal changed
            for item in items or []:
                table = getattr(item, "table", None)
                if table is not None:
                    title = node_change_title(item)
                    if "变更" in str(title or ""):
                        for type_code, rows in reqs_by_type.items():
                            type_name = type_name_by_code.get(type_code) or "变更需求"
                            if self.__change_table_title_matches(title, type_name):
                                if apply_change_rows_to_node(item, rows):
                                    changed = True
                                sync_type_name_from_tree_item(item, type_code)
                                break
                sync_existing_change_tables(getattr(item, "children", None) or [])

        sync_existing_change_tables(nodes or [])

        def has_existing_change_table(type_name: str) -> bool:
            for title in existing_titles:
                if self.__change_table_title_matches(title, type_name):
                    return True
            return False

        def build_change_table_node(type_code: str, type_name: str, rows: List[SrsReq]):
            headers = [
                TabHeader(code="srs_code", name="需求编号"),
                TabHeader(code="module", name="模块"),
                TabHeader(code="function", name="功能"),
                TabHeader(code="sub_function", name="子功能"),
            ]
            table_rows = build_change_table_rows(rows)
            # 表名只放 table.name；节点 title 不能带 2.1 变更需求，否则前端会当成章节号。
            return SrsNodeForm(
                title="",
                label=type_name,
                table=Table(name=type_name, headers=headers, rows=table_rows),
                children=[],
            )

        missing_nodes = []
        if restructure:
            for type_code, rows in reqs_by_type.items():
                type_name = type_name_by_code.get(type_code) or "变更需求"
                if has_existing_change_table(type_name):
                    continue
                missing_nodes.append(build_change_table_node(type_code, type_name, rows))
        if restructure and missing_nodes:
            target = preferred_target

            def score_node(item: SrsNodeForm):
                text = f"{getattr(item, 'title', '') or ''}\n{getattr(item, 'label', '') or ''}\n{getattr(item, 'text', '') or ''}"
                score = 0
                if "产品需求" in text:
                    score += 5
                if "其他需求" in text:
                    score += 4
                if "变更需求" in text:
                    score += 3
                if getattr(item, "children", None):
                    score += 1
                return score

            def walk_for_target(items: List[SrsNodeForm]):
                nonlocal target
                for item in items or []:
                    if target is None or score_node(item) > score_node(target):
                        target = item
                    walk_for_target(getattr(item, "children", None) or [])

            walk_for_target(nodes or [])
            if target is None and nodes:
                target = nodes[0]
            if target is not None:
                target.children = list(getattr(target, "children", None) or [])
                for node in missing_nodes:
                    node_title = node_change_title(node)
                    if has_existing_change_table(node_title):
                        continue
                    target.children.append(node)
            else:
                nodes.extend(missing_nodes)
            changed = True

        def is_change_table_node(item: SrsNodeForm):
            title = node_change_title(item)
            if "变更" not in str(title or ""):
                return False
            table = getattr(item, "table", None)
            if table is None:
                return False
            headers = getattr(table, "headers", None) or []
            rows = getattr(table, "rows", None) or []
            cells = getattr(table, "cells", None) or []
            table_name = getattr(table, "name", None) or ""
            return bool(headers or rows or cells or table_name)

        def dedupe_change_table_nodes(items: List[SrsNodeForm]):
            nonlocal changed
            groups: Dict[str, List[SrsNodeForm]] = {}

            def collect(nodes: List[SrsNodeForm]):
                for node in nodes or []:
                    if is_change_table_node(node):
                        title = node_change_title(node)
                        group_key = None
                        for key in groups.keys():
                            if self.__change_table_title_matches(key, title):
                                group_key = key
                                break
                        if group_key is None:
                            group_key = title
                        groups.setdefault(group_key, []).append(node)
                    collect(getattr(node, "children", None) or [])

            collect(items or [])
            remove_ids = set()
            for group in groups.values():
                if len(group) <= 1:
                    continue
                group.sort(
                    key=lambda node: (
                        0 if re.match(r"^导入表格\d*$", str(getattr(node, "title", "") or "").strip()) else 1,
                        0 if getattr(getattr(node, "table", None), "name", None) else 1,
                    )
                )
                for extra in group[1:]:
                    remove_ids.add(id(extra))
            if not remove_ids:
                return

            def filter_removed(nodes: List[SrsNodeForm]):
                result = []
                for node in nodes or []:
                    if id(node) in remove_ids:
                        continue
                    node.children = filter_removed(getattr(node, "children", None) or [])
                    result.append(node)
                return result

            items[:] = filter_removed(items or [])
            changed = True

        dedupe_change_table_nodes(nodes or [])
        return changed

    def __fix_rcms(self, doc: SrsDoc):
        objs_dict, tree = self.__tree(doc)
        req_rows = db.session.execute(select(SrsReq).where(SrsReq.doc_id == doc.id)).scalars().all()

        def normalize_match_text(value: str):
            txt = self.__clean_req_title(str(value or ""))
            return re.sub(r"[\s\u3000、，。；;：:（）()【】\[\]_\-]+", "", txt).upper()

        def build_req_name_index(rows: List[SrsReq]):
            index = {}
            for req in rows or []:
                names = [
                    getattr(req, "sub_function", "") or "",
                    getattr(req, "function", "") or "",
                    getattr(req, "module", "") or "",
                ]
                for name in names:
                    key = normalize_match_text(name)
                    if key:
                        codes = index.setdefault(key, [])
                        if req.code not in codes:
                            codes.append(req.code)
            return index

        req_name_index = build_req_name_index(req_rows)

        def node_search_text(node):
            parts = [
                getattr(node, "title", "") or "",
                getattr(node, "label", "") or "",
                getattr(node, "text", "") or "",
            ]
            table = getattr(node, "table", None)
            if table:
                parts.append(self.__table_search_text(table))
            return "\n".join(parts)

        def extract_rcm_codes_from_node(node):
            existed = self.__normalize_rcm_codes(getattr(node, "rcm_codes", None) or [])
            picked = self.__extract_rcm_codes_from_text(node_search_text(node))
            return self.__normalize_rcm_codes([*existed, *picked])

        def extract_srs_codes_from_node(node):
            result = []
            for hit in re.findall(r"SRS[\s\-_]*[A-Z0-9.]+(?:\s*-\s*[A-Z0-9.]+)*", node_search_text(node), flags=re.IGNORECASE):
                code = self.__normalize_srs_code(hit)
                if code and code not in result:
                    result.append(code)
            return result

        def resolve_srs_code_by_function_context(node):
            titles = []
            cur = node
            safety = 0
            while cur and safety < 50:
                title = getattr(cur, "title", "") or ""
                if title:
                    titles.append(title)
                cur = objs_dict.get(getattr(cur, "p_id", 0))
                safety += 1
            for title in titles:
                key = normalize_match_text(title)
                if key and key in req_name_index and len(req_name_index[key]) == 1:
                    return req_name_index[key][0]
            return ""

        all_reqs = []
        all_rcms = []
        all_pairs = []
        for node in iter_tree(tree):
            rcm_codes = extract_rcm_codes_from_node(node)
            if rcm_codes:
                srs_code = self.__normalize_srs_code(node.srs_code)
                direct_srs_codes = extract_srs_codes_from_node(node)
                if direct_srs_codes:
                    for direct_srs_code in direct_srs_codes:
                        all_reqs.append(direct_srs_code)
                        all_rcms.extend(rcm_codes)
                        all_pairs.append((direct_srs_code, rcm_codes))
                    continue
                if not srs_code:
                    p_node = objs_dict.get(node.p_id)
                    while p_node:
                        srs_code = self.__normalize_srs_code(p_node.srs_code)
                        if srs_code:
                            break
                        p_node = objs_dict.get(p_node.p_id)
                if not srs_code:
                    srs_code = self.__normalize_srs_code(resolve_srs_code_by_function_context(node))
                if not srs_code:
                    continue
                all_reqs.append(srs_code)
                all_rcms.extend(rcm_codes)
                all_pairs.append((srs_code, rcm_codes))

        reqs_dict = dict()
        for req in req_rows:
            reqs_dict.setdefault(req.code, []).append(req.id)
        
        rcms = db.session.execute(select(Rcm)).scalars().all()
        rcms_dict = {self.__normalize_rcm_code(rcm.code): rcm.id for rcm in rcms}
        
        delete_values = []
        insert_values = []
        for srs_code, rcm_codes in all_pairs:
            req_ids = reqs_dict.get(srs_code, [])
            rcm_ids = [rcms_dict.get(rcm_code) for rcm_code in rcm_codes if rcm_code in rcms_dict]
            if not req_ids or not rcm_ids:
                continue
            delete_values.extend(req_ids)
            for req_id in req_ids:
                for rcm_id in rcm_ids:
                    insert_values.append(dict(req_id=req_id, rcm_id=rcm_id))
        if delete_values:
            db.session.execute(delete(ReqRcm).where(ReqRcm.req_id.in_(delete_values)))
        if insert_values:
            db.session.execute(pg_insert(ReqRcm).values(insert_values).on_conflict_do_nothing())
        db.session.commit()

    async def add_srs_doc(self, form: SrsDocForm):
        try:
            sql = select(func.count(SrsDoc.id)).where(SrsDoc.product_id == form.product_id, SrsDoc.version == form.version)
            count = db.session.execute(sql).scalar()
            if count > 0:
                version = (form.version or "").strip()
                return Resp.resp_err(msg=f"该产品下已经有{version}版本文档存在" if version else ts("msg_obj_exist"))
            
            row = SrsDoc(
                product_id=form.product_id,
                version=form.version,
                folder_name=form.folder_name,
                change_log=form.change_log,
                n_id=0,
                file_no=form.file_no,
            )
            db.session.add(row)
            db.session.flush()
            if form.content:
                self.__update_nodes(row, 0, form.content)
            db.session.commit()
            self.__fix_rcms(row)
            return Resp.resp_ok(data=SrsDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))
    
    async def duplicate_srs_doc(self, id: int):
        fromdoc:SrsDocObj = (await self.get_srs_doc(id, with_tree=True)).data
        if not fromdoc:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        version = new_version(fromdoc.version)
        newdoc = SrsDoc(
            product_id=fromdoc.product_id,
            version=version,
            folder_name=fromdoc.folder_name,
            file_no=fromdoc.file_no,
            change_log=fromdoc.change_log,
            n_id=0,
        )
        sql = select(func.count(SrsDoc.id)).where(SrsDoc.product_id == newdoc.product_id, SrsDoc.version == newdoc.version)
        count = db.session.execute(sql).scalar()
        if count > 0:
            return Resp.resp_err(msg=ts("msg_obj_exist"))
        try:
            db.session.add(newdoc)
            db.session.flush()
            self.__update_nodes(newdoc, 0, fromdoc.content)

            # 变更需求表的 type_code 是跨文档业务标识，复制时必须重新生成，否则新旧文档共享
            # 同一 type_code，导致在副本里操作变更需求会串改到原文档。标准/其他需求(1/2)语义固定，保持不变。
            type_code_map: Dict[str, str] = {}
            srstypes: List[SrsType] = db.session.execute(select(SrsType).where(SrsType.doc_id == fromdoc.id).order_by(SrsType.id)).scalars().all()
            for srstype in srstypes:
                new_type_code = srstype.type_code
                if srstype.type_code and srstype.type_code not in ("1", "2"):
                    new_type_code = get_uuid()
                    type_code_map[srstype.type_code] = new_type_code
                newtype = SrsType(doc_id=newdoc.id, type_code=new_type_code, type_name=srstype.type_name)
                db.session.add(newtype)

            sql = select(SrsReq, SrsReqd).outerjoin(SrsReqd, SrsReq.id == SrsReqd.req_id).where(SrsReq.doc_id == fromdoc.id)
            srsreqs: List[Tuple[SrsReq, SrsReqd]] = db.session.execute(sql).all()
            for srsreq, reqd in srsreqs:
                new_req_type_code = type_code_map.get(srsreq.type_code, srsreq.type_code)
                newreq = SrsReq(doc_id=newdoc.id, code=srsreq.code, module=srsreq.module, 
                            function=srsreq.function, sub_function=srsreq.sub_function,
                            location=srsreq.location, type_code=new_req_type_code)
                db.session.add(newreq)
                db.session.flush()
                if reqd:
                    newreqd = SrsReqd(req_id=newreq.id, name=reqd.name, overview=reqd.overview,
                                    participant=reqd.participant, pre_condition=reqd.pre_condition,
                                    trigger=reqd.trigger, work_flow=reqd.work_flow,
                                    post_condition=reqd.post_condition, exception=reqd.exception,
                                    constraint=reqd.constraint)
                    db.session.add(newreqd)

                reqrcms: List[ReqRcm] = db.session.execute(select(ReqRcm).where(ReqRcm.req_id == srsreq.id)).scalars().all()
                for reqrcm in reqrcms:
                    newreqrcm = ReqRcm(req_id=newreq.id, rcm_id=reqrcm.rcm_id)
                    db.session.add(newreqrcm)
            db.session.commit()
            return Resp.resp_ok(data=SrsDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))
   
    async def delete_srs_doc(self, id):
        try:
            sql = select(func.count(SdsDoc.id)).where(SdsDoc.srsdoc_id == id)
            count = db.session.execute(sql).scalar()

            # 清理SRS管理数据
            req_ids = [req_id for req_id, in db.session.query(SrsReq.id).filter(SrsReq.doc_id == id).all()]
            if req_ids:
                db.session.execute(delete(ReqRcm).where(ReqRcm.req_id.in_(req_ids)))
                db.session.execute(delete(SrsReqd).where(SrsReqd.req_id.in_(req_ids)))
            db.session.execute(delete(SrsReq).where(SrsReq.doc_id == id))
            db.session.execute(delete(SrsType).where(SrsType.doc_id == id))
            db.session.execute(delete(SrsNode).where(SrsNode.doc_id == id))
            if count > 0:
                # 若已绑定详细设计：保留SRS主记录用于维持产品绑定，但标记为“已删除”并从可选列表隐藏
                row = db.session.execute(select(SrsDoc).where(SrsDoc.id == id)).scalars().first()
                if not row:
                    return Resp.resp_err(msg=ts("msg_obj_null"))
                stamp = datetime.now().strftime("%y%m%d%H%M%S")
                row.version = f"{DELETED_SRS_VERSION_PREFIX}{id}_{stamp}"
                row.change_log = "已删除需求规格说明（保留绑定占位）"
                row.n_id = 0
            else:
                db.session.execute(delete(SrsDoc).where(SrsDoc.id == id))
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))
    
    async def add_srs_node(self, node: SrsNodeForm):
        sql = select(SrsNode, SrsDoc).join(SrsDoc, SrsNode.doc_id == SrsDoc.id)
        sql = sql.where(SrsNode.doc_id == node.doc_id, SrsNode.n_id == node.p_id)
        result = db.session.execute(sql).first()
        if not result:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        _, doc = result
        doc.n_id += 1
        table = json.loads(node.table.json()) if node.table else None
        row = SrsNode(doc_id=doc.id, n_id=doc.n_id, p_id=node.p_id, priority=doc.n_id, 
                            title=node.title, text=node.text, table=table)
        row.rcm_codes = ",".join(node.rcm_codes) if node.rcm_codes is not None else None
        db.session.add(row)
        db.session.commit()
        data = dict(doc_id=row.doc_id, n_id=row.n_id, p_id=row.p_id, priority=row.priority, **node.dict())
        return Resp.resp_ok(data=SrsNodeForm(**data))
    
    async def delete_srs_node(self, doc_id, n_id):
        db.session.execute(delete(SrsNode).where(SrsNode.doc_id == doc_id, SrsNode.n_id == n_id))
        db.session.commit()
        return Resp.resp_ok()
   
    async def update_srs_doc(self, form: SrsDocForm):
        try:
            sql = select(func.count(SrsDoc.id)).where(SrsDoc.product_id == form.product_id, SrsDoc.version == form.version, SrsDoc.id != form.id)
            count = db.session.execute(sql).scalar()
            if count > 0:
                version = (form.version or "").strip()
                return Resp.resp_err(msg=f"该产品下已经有{version}版本文档存在" if version else ts("msg_obj_exist"))
            sql = select(SrsDoc).where(SrsDoc.id == form.id)
            row:SrsDoc = db.session.execute(sql).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            if form.content is None:
                logger.warning("update_srs_doc missing content: doc_id=%s", form.id)
                return Resp.resp_err(msg="保存失败：未收到文档结构内容，请刷新后重试")
            if isinstance(form.content, list) and len(form.content) == 0:
                logger.warning("update_srs_doc empty content: doc_id=%s", form.id)
                return Resp.resp_err(msg="保存失败：文档结构为空，请刷新后重试")
            logger.info("update_srs_doc content_count: doc_id=%s count=%s", form.id, len(form.content or []))
            for key, value in form.dict().items():
                if key == "id" or key == "n_id" or value is None:
                    continue
                setattr(row, key, value)
            row.n_id = 0
            db.session.execute(delete(SrsNode).where(SrsNode.doc_id == row.id))
            self.__sync_change_req_tables_from_db(row.id, form.content or [])
            self.__reset_tree_node_ids(form.content or [])
            self.__update_nodes(row, 0, form.content)
            db.session.commit()
            self.__sync_srs_reqs_from_doc_tables(row.id, form.content or [])
            # 兜底：保存时也扫描树上的变更需求表，把误存为 type_code='1'/'2' 的行
            # 自动迁移到正确的变更类型，避免 Word 导入历史遗留的错分类继续影响。
            self.__upsert_imported_change_req_tables_from_tree(row.id, form.content or [])
            self.__sync_srs_req_names_from_doc_nodes(row.id, form.content or [])
            self.__upsert_imported_srs_reqds(row.id, self.__extract_srs_reqds_from_nodes(form.content or []))
            self.__fix_rcms(row)
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_srs_doc_file_no(self, id: int, file_no: str):
        try:
            sql = select(SrsDoc).where(SrsDoc.id == id)
            row: SrsDoc = db.session.execute(sql).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            row.file_no = (file_no or "").strip() or None
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))
    
    async def get_srs_doc_txts(self, doc_id):
        def __gather_nodes(texts:List[str],nodes: List[SrsNodeForm]):
            for node in nodes:
                values = [node.title, node.text]
                values = [value for value in values if value]
                texts += values
                if node.children:
                    __gather_nodes(texts, node.children)
            return texts

        docdata: Resp[SrsDocObj] = (await self.get_srs_doc(doc_id, with_tree=True)).data
        content = docdata.content if docdata and docdata.content else []
        txts = __gather_nodes([], content)
        return Resp.resp_ok(data=txts)
   
    def __pick_doc_file_row(self, product_id: int, category: str, doc_version: str = None, product_version: str = None):
        return pick_doc_image_file_row(product_id, category, doc_version or "", product_version or "")

    def __query_imgs(self, product_id: int, doc_version: str = None, product_version: str = None):
        if not product_id:
            return {}
        result = {}
        for category in (RefTypes.img_topo.value, RefTypes.img_struct.value, RefTypes.img_flow.value):
            row = self.__pick_doc_file_row(product_id, category, doc_version, product_version)
            if row and row.file_url:
                result[category] = row.file_url
        return result

    def __resolve_product_bound_doc_image_ref_type(self, node: SrsNodeForm):
        ref_type = getattr(node, "ref_type", None)
        if ref_type in (RefTypes.img_topo.value, RefTypes.img_struct.value):
            return ref_type
        heading_no = self.__heading_no_from_title(getattr(node, "title", "") or "")
        if heading_no == "2.2":
            return RefTypes.img_topo.value
        if heading_no == "2.3":
            return RefTypes.img_struct.value
        return None

    def __hydrate_tree_product_doc_images(self, nodes: List[SrsNodeForm], prod_imgs: dict):
        if not nodes or not prod_imgs:
            return
        for node in nodes or []:
            bound_type = self.__resolve_product_bound_doc_image_ref_type(node)
            if bound_type:
                if not getattr(node, "ref_type", None):
                    node.ref_type = bound_type
                if prod_imgs.get(bound_type):
                    node.img_url = prod_imgs[bound_type]
            if node.children:
                self.__hydrate_tree_product_doc_images(node.children, prod_imgs)

    def __tree(self, doc: SrsDoc):
        tree = []
        sql = select(SrsNode).where(SrsNode.doc_id == doc.id).order_by(SrsNode.priority)
        nodes: List[SrsNode] = db.session.execute(sql).scalars().all()
        objs_dict = dict()
        objs = []
        product = db.session.execute(select(Product).where(Product.id == doc.product_id)).scalars().first()
        product_version = getattr(product, "full_version", "") or getattr(product, "name", "") or ""
        prod_imgs = self.__query_imgs(doc.product_id, doc.version, product_version)
        for node in nodes:
            table = None
            if node.table:
                try:
                    if isinstance(node.table, Table):
                        table = node.table
                    elif isinstance(node.table, (dict, list)):
                        table = Table.parse_obj(node.table)
                    elif isinstance(node.table, str):
                        table = Table.parse_raw(node.table)
                    else:
                        table = Table.parse_obj(node.table)
                except Exception:
                    logger.warning("parse srs_node.table failed: doc_id=%s n_id=%s", node.doc_id, node.n_id)
                    table = None

            obj = SrsNodeForm(children=[], doc_id=node.doc_id, n_id=node.n_id, p_id=node.p_id,
                            title=node.title, label=node.label, img_url=node.img_url, text=node.text, ref_type=node.ref_type, table=table, srs_code=node.srs_code)
            obj.rcm_codes = self.__normalize_rcm_codes(node.rcm_codes.split(",")) if node.rcm_codes is not None else None

            objs_dict[obj.n_id] = obj
            objs.append(obj)
        for obj in objs:
            if obj.p_id == 0:
                tree.append(obj)
            else:
                p_obj = objs_dict.get(obj.p_id)
                if not p_obj:
                    logger.warning("ignoreNode:: %s %s %s", obj.doc_id, obj.p_id, obj.n_id)
                    continue
                p_obj.children.append(obj)
        self.__hydrate_tree_product_doc_images(tree, prod_imgs)
        return objs_dict, tree

    async def get_srs_doc(self, id:str, with_tree: bool = False):
        sql = select(SrsDoc, Product).outerjoin(Product, SrsDoc.product_id == Product.id).where(SrsDoc.id == id)
        row, row_prod = db.session.execute(sql).first() or (None, None)
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        if (row.version or "").startswith(DELETED_SRS_VERSION_PREFIX):
            return Resp.resp_err(msg=ts("msg_obj_null"))
        objs_dict, tree = self.__tree(row) if with_tree else (None, [])
        if with_tree and self.__sync_change_req_tables_from_db(row.id, tree):
            row.n_id = 0
            db.session.execute(delete(SrsNode).where(SrsNode.doc_id == row.id))
            self.__reset_tree_node_ids(tree)
            self.__update_nodes(row, 0, tree)
            db.session.commit()
            objs_dict, tree = self.__tree(row)
        product_name = row_prod.name if row_prod else ""
        product_version = row_prod.full_version if row_prod else ""

        # srs_nodes = dict()
        # if tree:
        #     all_srscodes = []
        #     for node in iter_tree(tree):
        #         if node.rcm_codes is None:
        #             continue
        #         srs_code = node.srs_code
        #         if not srs_code:
        #             p_node = objs_dict.get(node.p_id)
        #             while p_node:
        #                 srs_code = p_node.srs_code
        #                 if srs_code:
        #                     break
        #                 p_node = objs_dict.get(p_node.p_id)
        #         if not srs_code:
        #             continue
        #         srs_nodes.setdefault(srs_code, []).append(node)
        #         all_srscodes.append(srs_code)

        #     sql = select(ReqRcm, SrsReq, Rcm).outerjoin(SrsReq, ReqRcm.req_id == SrsReq.id)
        #     sql = sql.outerjoin(Rcm, ReqRcm.rcm_id == Rcm.id)
        #     sql = sql.where(SrsReq.code.in_(all_srscodes))
        #     sql = sql.distinct(SrsReq.code, Rcm.code).order_by(SrsReq.code, Rcm.code)
        #     rows: list[Tuple[ReqRcm, SrsReq, Rcm]] = db.session.execute(sql).all()
        #     srs_rcms = dict()
        #     for _, req, rcm in rows:
        #         srs_rcms.setdefault(req.code, []).append(rcm.code)
        #     for srs_code, nodes in srs_nodes.items():
        #         rcms = srs_rcms.get(srs_code) or []
        #         for node in nodes:
        #             node.rcm_codes = rcms
        return Resp.resp_ok(data=SrsDocObj(**row.dict(), product_name=product_name, product_version=product_version, content=tree))

    async def list_srs_doc(self, op_user: UserObj, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        page_index = page_index if page_index >= 0 else 0
        page_size = page_size if page_size > 0 else 10 
    
        sql = select(SrsDoc, Product).outerjoin(Product, SrsDoc.product_id == Product.id)
        sql = sql.where(~SrsDoc.version.like(f"{DELETED_SRS_VERSION_PREFIX}%"))
        if product_id:
            sql = sql.where(SrsDoc.product_id == product_id)
        if version:
            sql = sql.where(SrsDoc.version.like(f"%{version}%"))
        if not product_id and op_user.id != 1:
            subquery = select(UserProd.product_id).where(UserProd.user_id == op_user.id).scalar_subquery()
            sql = sql.where(Product.id.in_(subquery))
        
        sql_count = select(func.count()).select_from(sql)
        total = db.session.execute(sql_count).scalars().first()
        sql = sql.offset(page_size * page_index).limit(page_size).order_by(desc(SrsDoc.create_time))
        rows: List[SrsDoc] = db.session.execute(sql).all()

        objs = []
        for row, row_prd in rows:
            obj = SrsDocObj(**row.dict())
            if row_prd:
                obj.product_name = row_prd.name
                obj.product_version = row_prd.full_version
            objs.append(obj)
        return Resp.resp_ok(data=Page(total=total, page_size=page_size, rows=objs, page_index=page_index))

    async def compare_srs_doc(self, id0: int, id1: int):
        def __feature_key(req: SrsReq):
            # 判定新增/减少时仅按功能编号，避免名称改动造成误判
            code = (req.code or "").strip()
            if code:
                return code
            module = (req.module or "").strip()
            function = (req.function or "").strip()
            return " - ".join([v for v in [module, function] if v])

        def __feature_display(req: SrsReq):
            code = (req.code or "").strip()
            module = (req.module or "").strip()
            function = (req.function or "").strip()
            name = " - ".join([v for v in [module, function] if v])
            if code and name:
                return f"{code} {name}"
            return code or name

        def __to_text(values: List[str]):
            return "；".join(values) if values else "无"

        def __query_feature_maps():
            feature_dict = {id0: set(), id1: set()}
            feature_name_dict = {id0: {}, id1: {}}
            rows: List[SrsReq] = db.session.execute(
                select(SrsReq).where(SrsReq.doc_id.in_([id0, id1])).order_by(SrsReq.doc_id, SrsReq.module, SrsReq.function, SrsReq.code)
            ).scalars().all()
            for req in rows:
                key = __feature_key(req)
                if not key:
                    continue
                feature_dict.setdefault(req.doc_id, set()).add(key)
                feature_name_dict.setdefault(req.doc_id, {}).setdefault(key, __feature_display(req) or key)
            return feature_dict, feature_name_dict

        sql = select(SrsDoc, Product).join(Product, SrsDoc.product_id == Product.id).where(SrsDoc.id.in_([id0, id1]))
        rows: List[Tuple[SrsDoc, Product]] = db.session.execute(sql).all()
        if not rows:
            return Resp.resp_err(msg=ts("msg_obj_null"))

        feature_dict, feature_name_dict = __query_feature_maps()
        features0 = feature_dict.get(id0) or set()
        features1 = feature_dict.get(id1) or set()
        added0_keys = sorted(features0 - features1)
        added1_keys = sorted(features1 - features0)
        added0 = [feature_name_dict.get(id0, {}).get(key, key) for key in added0_keys]
        added1 = [feature_name_dict.get(id1, {}).get(key, key) for key in added1_keys]
        removed0 = added1
        removed1 = added0

        infos = {}
        for row_srsdoc, row_prd in rows:
            infos[row_srsdoc.id] = dict(
                product_name=row_prd.name,
                product_type_code=row_prd.type_code,
                product_version=row_prd.full_version,
                product_udi=row_prd.udi,
                product_scope=row_prd.scope,
                srs_version=row_srsdoc.version,
            )
        info0 = infos.get(id0) or {}
        info1 = infos.get(id1) or {}

        results = []
        for column in ["product_name", "product_type_code", "product_version", "product_udi", "product_scope", "srs_version"]:
            value0 = info0.get(column) or ""
            value1 = info1.get(column) or ""
            same_flag = 1 if value0 == value1 else 0
            results.append(CompareObj(column_code=column, column_name=ts(f"sdsdiff.{column}"), same_flag=same_flag, values=[value0, value1]))

        results += [
            CompareObj(
                column_code="feature_added",
                column_name="新增功能",
                same_flag=1 if not added0 and not added1 else 0,
                values=[__to_text(added0), __to_text(added1)],
            ),
            CompareObj(
                column_code="feature_removed",
                column_name="减少功能",
                same_flag=1 if not removed0 and not removed1 else 0,
                values=[__to_text(removed0), __to_text(removed1)],
            ),
        ]
        return Resp.resp_ok(data=results)

    async def export_srs_doc(self, output, doc_id, snapshot: SrsDocForm = None, *args, **kwargs):
        if Document is None or Pt is None or dox_enum is None:
            return
        from .serv_utils import docx_util
        def __norm_title(value: str):
            return re.sub(r"\s+", "", value or "")

        def __is_cover_section_title(title: str):
            txt = __norm_title(title)
            return txt in ["需求规格说明", "文件修订记录"]

        def __is_spec_title(title: str):
            return __norm_title(title) == "需求规格说明"

        def __is_rev_title(title: str):
            return __norm_title(title) == "文件修订记录"

        def __is_revision_label(value: str):
            return __norm_title(value) == "文件修订记录"

        def __is_imported_catalog_title(value: str):
            txt = (value or "").strip()
            if not txt:
                return False
            if __norm_title(txt) == "目录":
                return True
            # Word 原目录项常被导入成“1 介绍 1”“2.2 物理拓扑图 6”
            return re.match(r"^\d+(?:\.\d+)*\.?\s+\S.*\s+\d+$", txt) is not None

        def __is_imported_catalog_line(value: str):
            txt = (value or "").strip()
            if not txt:
                return False
            if __norm_title(txt) in ["需求规格说明", "文件修订记录", "目录"]:
                return True
            if __is_imported_catalog_title(txt):
                return True
            # 兼容带点线的目录行
            return re.match(r"^\d+(?:\.\d+)*\s+\S.*[.·…]{3,}\s*\d+$", txt) is not None

        def __is_imported_catalog_root(node: SrsNodeForm):
            title = __norm_title(getattr(node, "title", "") or "")
            text = str(getattr(node, "text", "") or "")
            return title == "导入正文" and "目录" in text and any(
                __is_imported_catalog_line(line) for line in text.splitlines()
            )

        def __strip_imported_catalog_suffix(value: str):
            txt = (value or "").strip()
            matched = re.match(r"^(\d+(?:\.\d+)*\.?\s+\S.*\S)\s+\d+$", txt)
            return matched.group(1).strip() if matched else txt

        def __strip_imported_catalog_lines(value: str):
            lines = [
                line for line in str(value or "").splitlines()
                if not __is_imported_catalog_line(line)
            ]
            return "\n".join(lines).strip()

        def __is_imported_placeholder_title(title: str):
            txt = (title or "").strip()
            return re.match(r"^导入(表格|图片)\d*$", txt) is not None

        def __is_imported_table_title(title: str):
            return re.match(r"^导入表格\d*$", (title or "").strip()) is not None

        def __is_table_caption_line(line: str):
            return re.match(r"^\s*表\s*\d+\s*", (line or "").strip()) is not None

        def __is_image_caption_line(line: str):
            return re.match(r"^\s*图\s*\d+\s*", (line or "").strip()) is not None

        def __save_image_caption_txt(docx: Document, text: str, font_size: float = 10.5):
            txt = (text or "").strip()
            if not txt:
                return
            p = docx.add_paragraph()
            p.alignment = dox_enum.text.WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            docx_util.fonted_txt(p, txt, font_size)

        def __strip_heading_no(value: str):
            return re.sub(r"^\s*\d+(?:\.\d+)*\s*", "", value or "").strip()

        def __insert_toc_field(docx: Document):
            docx_util.insert_toc_field(docx, outline_levels="1-4")

        def __write_catalog_page(docx: Document, catalog_text: str):
            __write_center_section_title(docx, "目录")
            # 始终插入 TOC 域，由 Word/LibreOffice 按 Heading 样式计算真实页码
            __insert_toc_field(docx)

        def __extract_imported_catalog_text(*nodes: SrsNodeForm):
            lines = []
            def walk(node: SrsNodeForm):
                if not node:
                    return
                title = str(getattr(node, "title", "") or "").strip()
                if __is_imported_catalog_title(title) and __norm_title(title) != "目录":
                    lines.append(title)
                for raw in str(getattr(node, "text", "") or "").splitlines():
                    line = (raw or "").strip()
                    if not line:
                        continue
                    if __norm_title(line) in ["需求规格说明", "文件修订记录", "目录"]:
                        continue
                    if __is_imported_catalog_line(line):
                        lines.append(line)
                for child in (getattr(node, "children", None) or []):
                    walk(child)
            for node in nodes:
                walk(node)
            return "\n".join(lines).strip()

        def __clean_srs_table_for_export(table):
            if not table:
                return table
            headers = getattr(table, "headers", None) or []
            header_names = [self.__normalize_text(getattr(h, "name", "") or "") for h in headers]
            cells = getattr(table, "cells", None) or []
            if cells and (not header_names):
                try:
                    header_names = [self.__normalize_text(str(getattr(cell, "value", "") or "")) for cell in cells[0]]
                except Exception:
                    header_names = []
            col_idx = self.__resolve_req_columns([self.__normalize_header(name) for name in header_names])
            clean_cols = [col_idx[key] for key in ["module", "function", "sub_function", "location"] if key in col_idx]
            if not clean_cols:
                if cells:
                    table.cells = None
                return table

            rows = getattr(table, "rows", None) or []
            header_codes = [getattr(h, "code", "") for h in headers]
            if (not rows) and cells and len(cells) > 1:
                rows = []
                active_spans = {}
                for cell_row in cells[1:]:
                    row = {}
                    for col_index, header in enumerate(headers):
                        code = getattr(header, "code", "") or (header_codes[col_index] if col_index < len(header_codes) else "")
                        active = active_spans.get(col_index)
                        cell = cell_row[col_index] if col_index < len(cell_row) else None
                        row_span = int(getattr(cell, "row_span", 1) or 1) if cell is not None else 1
                        col_span = int(getattr(cell, "col_span", 1) or 1) if cell is not None else 1
                        if row_span == 0 or col_span == 0:
                            row[code] = active.get("value", "") if active else ""
                            if active:
                                active["remaining"] -= 1
                                if active["remaining"] <= 0:
                                    active_spans.pop(col_index, None)
                            continue
                        value = self.__clean_req_table_field(getattr(cell, "value", "") if cell is not None else "")
                        row[code] = value
                        if row_span > 1:
                            active_spans[col_index] = {"value": value, "remaining": row_span - 1}
                    rows.append(row)
                table.rows = rows
            field_codes = {
                field: header_codes[col_idx[field]]
                for field in ["module", "function", "sub_function", "location"]
                if field in col_idx and col_idx[field] < len(header_codes)
            }
            last_values: Dict[str, str] = {}
            for row in rows:
                if not isinstance(row, dict):
                    continue
                current_values = {
                    field: re.sub(r"\s+", "", self.__normalize_text(row.get(code, "") or ""))
                    for field, code in field_codes.items()
                }
                module_changed = bool(current_values.get("module"))
                function_changed = bool(current_values.get("function"))
                for field, code in field_codes.items():
                    value = current_values.get(field) or ""
                    if value:
                        last_values[field] = value
                        row[code] = value
                        if field == "module":
                            last_values.pop("function", None)
                            last_values.pop("sub_function", None)
                        elif field == "function":
                            last_values.pop("sub_function", None)
                    elif field == "module" and last_values.get("module"):
                        row[code] = last_values["module"]
                    elif field == "function" and not module_changed and last_values.get("function"):
                        row[code] = last_values["function"]
                    elif field == "sub_function" and not (module_changed or function_changed) and last_values.get("sub_function"):
                        row[code] = last_values["sub_function"]
                    elif field == "location" and last_values.get("location"):
                        row[code] = last_values["location"]
                    else:
                        row[code] = value

            # 行数据已补全，合并结构由 __build_srs_table_export_cells 重建
            return table

        def __srs_req_code_group(code: str):
            normalized = self.__normalize_srs_code(str(code or ""))
            matched = re.match(r"^(SRS-[A-Z]+\d+)-\d+$", normalized or "")
            return matched.group(1) if matched else normalized

        def __is_main_req_export_table(table):
            if not table or not getattr(table, "headers", None):
                return False
            headers = table.headers or []
            header_names = [
                self.__normalize_header(getattr(header, "name", "") or "")
                for header in headers
            ]
            header_codes = [
                self.__normalize_header(getattr(header, "code", "") or "")
                for header in headers
            ]
            col_idx = self.__resolve_req_columns(header_names + header_codes)
            return "code" in col_idx and "function" in col_idx

        def __walk_doc_nodes(nodes):
            for node in nodes or []:
                yield node
                for child in __walk_doc_nodes(getattr(node, "children", None) or []):
                    yield child

        def __find_doc_req_table_in_nodes(nodes, req_kind: str):
            for node in __walk_doc_nodes(nodes):
                table = getattr(node, "table", None)
                if not table or not getattr(table, "headers", None):
                    continue
                if req_kind == "main" and __is_main_req_export_table(table):
                    return table
                if req_kind == "other" and __is_other_req_export_table(table):
                    return table
            return None

        def __build_srs_table_export_cells(table: Table):
            headers = getattr(table, "headers", None) or []
            rows = getattr(table, "rows", None) or []
            if not headers or not rows:
                return None
            header_codes = [getattr(header, "code", "") for header in headers]
            headers_norm = [self.__normalize_header(getattr(header, "name", "") or "") for header in headers]
            col_idx = self.__resolve_req_columns(headers_norm)
            module_idx = col_idx.get("module")
            function_idx = col_idx.get("function")
            location_idx = col_idx.get("location")
            cells = [
                [TableCell(value=getattr(header, "name", "") or "", row_span=1, col_span=1) for header in headers],
                *[
                    [TableCell(value=str(row.get(getattr(header, "code", ""), "") or ""), row_span=1, col_span=1) for header in headers]
                    for row in rows
                ],
            ]

            def cell_value(row_idx, col_index):
                if col_index is None or col_index >= len(header_codes):
                    return ""
                return str(rows[row_idx].get(header_codes[col_index], "") or "")

            def merge_column(col_index, parent_col_indexes):
                if col_index is None:
                    return
                start = 0
                while start < len(rows):
                    start_value = cell_value(start, col_index)
                    if not start_value:
                        start += 1
                        continue
                    end = start + 1
                    while end < len(rows):
                        if cell_value(end, col_index) != start_value:
                            break
                        same_parent = all(
                            parent_col_idx is None or cell_value(end, parent_col_idx) == cell_value(start, parent_col_idx)
                            for parent_col_idx in parent_col_indexes
                        )
                        if not same_parent:
                            break
                        end += 1
                    span = end - start
                    cells[start + 1][col_index].value = start_value
                    if span > 1:
                        cells[start + 1][col_index].row_span = span
                        for row_idx in range(start + 1, end):
                            cells[row_idx + 1][col_index].value = ""
                            cells[row_idx + 1][col_index].row_span = 0
                    else:
                        cells[start + 1][col_index].row_span = 1
                    start = end

            merge_column(module_idx, [])
            merge_column(function_idx, [module_idx] if module_idx is not None else [])
            merge_column(location_idx, [module_idx] if module_idx is not None else [])
            return cells

        NUMBERED_REQ_DETAIL_FIELDS = ("事件流", "工作流", "工作流程", "前置条件", "触发器", "后置条件", "异常情况", "约束")

        def __renumber_req_detail_lines(value):
            # 与前端 normalizeReqDetailNumberedText 保持一致：把每行行首的数字序号重排为 1,2,3...
            next_no = 1
            pat = re.compile(r"^(\s*)(\d{1,4})([）)、．]|[.](?!\d))\s*(.*)$")
            out = []
            for line in str(value or "").split("\n"):
                m = pat.match(line or "")
                if not m:
                    out.append(line)
                    continue
                prefix = m.group(1) or ""
                sep = "." if m.group(3) == "．" else m.group(3)
                rest = m.group(4) or ""
                out.append(f"{prefix}{next_no}{sep} {rest}".rstrip())
                next_no += 1
            return "\n".join(out)

        def __renumber_req_detail_table(table):
            # 仅对“字段|值”详细需求表中、字段为事件流/工作流/前置条件等的行，重排其值的行首序号；其他表/行不动
            def _is_numbered_field(label):
                lab = re.sub(r"\s+", "", str(label or ""))
                return any(re.sub(r"\s+", "", f) in lab for f in NUMBERED_REQ_DETAIL_FIELDS)
            cells = getattr(table, "cells", None)
            if isinstance(cells, list):
                for row in cells:
                    if not isinstance(row, list) or not row:
                        continue
                    if _is_numbered_field(getattr(row[0], "value", "")):
                        for c in row[1:]:
                            if getattr(c, "value", None):
                                c.value = __renumber_req_detail_lines(c.value)
            rows = getattr(table, "rows", None)
            headers = getattr(table, "headers", None) or []
            if isinstance(rows, list) and headers:
                codes = [getattr(h, "code", "") for h in headers]
                if codes:
                    for row in rows:
                        if not isinstance(row, dict):
                            continue
                        if _is_numbered_field(row.get(codes[0], "")):
                            for code in codes[1:]:
                                if row.get(code):
                                    row[code] = __renumber_req_detail_lines(row[code])

        def __prepare_srs_table_for_word_export(table):
            if not table:
                return table
            export_table = deepcopy(table)
            __renumber_req_detail_table(export_table)
            # 表格本身带横向合并(col_span>1)时（如导入Word的合并表头），
            # 直接保留原合并结构导出，不清空、不重建（重建无法还原横向合并）。
            # 需求表的合并为纵向(col_span 全为1)，不受影响，仍走下方重建逻辑。
            orig_cells = getattr(export_table, "cells", None)
            has_h_merge = isinstance(orig_cells, list) and any(
                isinstance(r, list) and any((getattr(c, "col_span", 1) or 1) > 1 for c in r)
                for r in orig_cells
            )
            if has_h_merge:
                return export_table
            export_table = __clean_srs_table_for_export(export_table)
            built_cells = __build_srs_table_export_cells(export_table)
            if built_cells:
                export_table.cells = built_cells
            else:
                export_table.cells = None
            return export_table

        def __find_req_table_in_export_sources(req_kind: str):
            for roots in req_table_sources:
                found = __find_doc_req_table_in_nodes(roots, req_kind)
                if found:
                    return found
            return None

        async def __export_db_req_list_table(req_kind: str, docx, font_def, is_other: bool = False):
            db_table = __find_req_table_in_export_sources(req_kind)
            if db_table:
                prepared = __prepare_srs_table_for_word_export(db_table)
                if is_other:
                    await __save_tab_and_export_change_if_other(prepared, docx, font_def, show_name=False, prepared=True)
                else:
                    __save_tab2docx(prepared, docx, font_def, show_name=False, prepared=True)
                return True
            fallback = await __query_srs_reqs("2" if is_other else "1")
            if is_other:
                await __save_tab_and_export_change_if_other(fallback, docx, font_def, show_name=False)
            else:
                __save_tab2docx(fallback, docx, font_def, show_name=False)
            return True

        exported_req_labels = set()
        exported_req_tables = set()
        change_req_export_done = {"value": False}
        other_req_title_written = {"value": False}
        written_table_caption_norms = set()

        def __write_table_caption_once(table, docx, font_def, context_text: str = ""):
            # 补写表格名称（如「表1 服务器硬件配置要求」「产品需求列表如下：」）。
            # 仅当表名为真实名（非「导入表格N」占位）、未写过、且不在节点正文中时才写，避免重复。
            name = str(getattr(table, "name", "") or "").strip() if table else ""
            if not name or re.match(r"^导入表格\d*$", name):
                return
            norm = __norm_title(name)
            if not norm or norm in written_table_caption_norms:
                return
            if context_text and norm in __norm_title(context_text):
                return
            docx_util.save_txt2docx(name, docx, font_def)
            written_table_caption_norms.add(norm)

        def __is_other_req_export_table(table):
            if not table or not getattr(table, "headers", None):
                return False
            headers = table.headers or []
            header_names = [
                self.__normalize_header(getattr(header, "name", "") or "")
                for header in headers
            ]
            header_codes = [
                self.__normalize_header(getattr(header, "code", "") or "")
                for header in headers
            ]
            table_name = str(getattr(table, "name", "") or "")
            if "其他需求" in table_name:
                col_idx = self.__resolve_req_columns(header_names + header_codes)
                return "code" in col_idx
            col_idx = self.__resolve_req_columns(header_names)
            if "code" not in col_idx:
                col_idx = self.__resolve_req_columns(header_codes)
            if "code" not in col_idx:
                col_idx = self.__resolve_req_columns(header_names + header_codes)
            if "code" not in col_idx:
                return False
            if "function" in col_idx:
                return False
            return "module" in col_idx or "location" in col_idx

        def __mark_other_req_title_from_text(line: str):
            text = str(line or "")
            if "其他需求列表" in text or text.strip() in ["其他需求", "其他需求:"]:
                other_req_title_written["value"] = True
                exported_req_labels.add(__norm_title("其他需求列表"))

        def __other_req_table_title(table):
            table_name = str(getattr(table, "name", "") or "").strip()
            if table_name and "其他需求" in table_name and not re.match(r"^导入表格\d*$", table_name):
                return table_name if table_name.endswith(":") or table_name.endswith("：") else f"{table_name}:"
            return "其他需求列表:"

        def __write_other_req_title_if_needed(docx, font_def, title: str = ""):
            if other_req_title_written["value"]:
                return False
            label = str(title or "其他需求列表:").strip()
            if label and not label.endswith(":") and not label.endswith("："):
                label = f"{label}:"
            if not label:
                label = "其他需求列表:"
            norm_label = __norm_title(label.replace(":", "").replace("：", ""))
            if not norm_label:
                return False
            docx_util.save_txt2docx(label, docx, font_def)
            other_req_title_written["value"] = True
            exported_req_labels.add(norm_label)
            return True

        def __save_tab2docx(table, docx, font_def=10.5, show_name: bool = True, prepared: bool = False):
            table_name = (getattr(table, "name", "") or "").strip() if table else ""
            if table and __is_other_req_export_table(table):
                __write_other_req_title_if_needed(docx, font_def, __other_req_table_title(table))
                show_name = False
            elif show_name and table_name and not re.match(r"^导入表格\d*$", table_name):
                docx_util.save_txt2docx(table_name, docx, font_def)
            export_table = table if prepared else __prepare_srs_table_for_word_export(table)
            docx_util.save_tab2docx(export_table, docx)

        def __write_center_section_title(docx: Document, title: str):
            p = docx.add_paragraph()
            p.alignment = dox_enum.text.WD_ALIGN_PARAGRAPH.CENTER
            # 需求规格说明使用二号字（22pt）
            font_size = 22.0 if __norm_title(title) == "需求规格说明" else 16.0
            docx_util.fonted_txt(p, title, font_size=font_size)

        def __add_blank_lines(docx: Document, line_count: int):
            for _ in range(max(0, line_count)):
                docx.add_paragraph("")

        def __is_revision_table(table):
            if not table:
                return False
            header_txt = "".join((getattr(h, "name", "") or "").strip() for h in (getattr(table, "headers", None) or []))
            keys = ["修改日期", "版本号", "修订说明", "修订人", "批准人"]
            return sum(1 for key in keys if key in header_txt) >= 3

        def __node_has_revision_marker(node: SrsNodeForm):
            for val in [getattr(node, "title", ""), getattr(node, "label", ""), getattr(node, "text", "")]:
                if __is_revision_label(val or ""):
                    return True
            return False

        async def __query_srs_reqs(type_code):
            resp = await srsreq_serv.list_srs_req(doc_id=doc_id, type_code=type_code, page_size=5000)
            reqs: List[SrsReq] = resp.data.rows or []
            def build_cells(table: Table):
                headers = getattr(table, "headers", None) or []
                rows = getattr(table, "rows", None) or []
                if not headers or not rows:
                    return None
                cells = [
                    [TableCell(value=getattr(header, "name", "") or "", row_span=1, col_span=1) for header in headers],
                    *[
                        [TableCell(value=str(row.get(getattr(header, "code", ""), "") or ""), row_span=1, col_span=1) for header in headers]
                        for row in rows
                    ],
                ]
                def header_index(match):
                    for idx, header in enumerate(headers):
                        name = self.__normalize_header(getattr(header, "name", "") or "")
                        if match(name):
                            return idx
                    return -1
                code_idx = header_index(lambda name: "需求编号" in name or name in ["srscode", "code"])
                module_idx = header_index(lambda name: "模块" in name)
                function_idx = header_index(lambda name: "功能" in name and "子功能" not in name)
                sub_function_idx = header_index(lambda name: "子功能" in name)
                if code_idx < 0 or module_idx < 0:
                    return cells
                def group_of(row):
                    code = self.__normalize_srs_code(str(row.get(getattr(headers[code_idx], "code", ""), "") or ""))
                    matched = re.match(r"^(SRS-[A-Z]+\d+)-\d+$", code)
                    return matched.group(1) if matched else code
                effective = []
                for idx, row in enumerate(rows):
                    prev = effective[idx - 1] if idx > 0 else {}
                    group = group_of(row)
                    same_group = bool(group and group == prev.get("group"))
                    raw_module = self.__keep_req_field(row.get(getattr(headers[module_idx], "code", ""))) if module_idx >= 0 else ""
                    raw_function = self.__keep_req_field(row.get(getattr(headers[function_idx], "code", ""))) if function_idx >= 0 else ""
                    raw_sub_function = self.__keep_req_field(row.get(getattr(headers[sub_function_idx], "code", ""))) if sub_function_idx >= 0 else ""
                    item = {
                        "group": group,
                        "module": raw_module or (prev.get("module") if same_group else ""),
                        "function": raw_function or (prev.get("function") if same_group and not raw_module else ""),
                        "sub_function": raw_sub_function or (prev.get("sub_function") if same_group and not raw_module and not raw_function else ""),
                    }
                    effective.append(item)
                def value_at(row_idx, col_idx):
                    if col_idx == module_idx:
                        return effective[row_idx].get("module") or ""
                    if col_idx == function_idx:
                        return effective[row_idx].get("function") or ""
                    if col_idx == sub_function_idx:
                        return effective[row_idx].get("sub_function") or ""
                    return self.__keep_req_field(rows[row_idx].get(getattr(headers[col_idx], "code", "")))
                def merge_column(col_idx, parent_indexes):
                    if col_idx < 0:
                        return
                    start = 0
                    while start < len(rows):
                        start_value = value_at(start, col_idx)
                        if not start_value:
                            start += 1
                            continue
                        end = start + 1
                        while end < len(rows):
                            if not effective[start].get("group") or effective[start].get("group") != effective[end].get("group"):
                                break
                            if value_at(end, col_idx) != start_value:
                                break
                            same_parent = all(parent_idx < 0 or value_at(end, parent_idx) == value_at(start, parent_idx) for parent_idx in parent_indexes)
                            if not same_parent:
                                break
                            end += 1
                        span = end - start
                        if span > 1:
                            cells[start + 1][col_idx].value = start_value
                            cells[start + 1][col_idx].row_span = span
                            for row_idx in range(start + 1, end):
                                cells[row_idx + 1][col_idx].value = ""
                                cells[row_idx + 1][col_idx].row_span = 0
                        start = end
                merge_column(module_idx, [])
                merge_column(function_idx, [module_idx])
                merge_column(sub_function_idx, [module_idx, function_idx])
                return cells
            if type_code == "2":
                headers = [TabHeader(code="srs_code", name="需求编号"), 
                        TabHeader(code="module", name="模块"), 
                        TabHeader(code="location", name="对应的章节")]
                rows = []
                for req in reqs:
                    row = dict()
                    row["srs_code"] = req.code
                    row["module"] = self.__keep_req_field(req.module)
                    row["location"] = self.__keep_req_field(req.location)
                    rows.append(row)
                table = Table(headers=headers, rows=rows)
                table.cells = build_cells(table)
                return table

            headers = [TabHeader(code="srs_code", name="需求编号"), 
                       TabHeader(code="module", name="模块"), 
                       TabHeader(code="function", name="功能"), 
                       TabHeader(code="sub_function", name="子功能")]
            rows = []
            for req in reqs:
                row = dict()
                row["srs_code"] = req.code
                row["module"] = self.__keep_req_field(req.module)
                row["function"] = self.__keep_req_field(req.function)
                row["sub_function"] = self.__keep_req_field(req.sub_function)
                rows.append(row)
            table = Table(headers=headers, rows=rows)
            table.cells = build_cells(table)
            return table
        
        async def __query_srs_reqs_x():
            srs_types: List[SrsType] = db.session.execute(select(SrsType).where(SrsType.doc_id == doc_id).order_by(SrsType.id)).scalars().all()
            change_req_rows: List[SrsReq] = db.session.execute(
                select(SrsReq).where(
                    SrsReq.doc_id == doc_id,
                    SrsReq.type_code.isnot(None),
                    SrsReq.type_code.notin_(["1", "2"]),
                ).order_by(SrsReq.id)
            ).scalars().all()
            type_names = {row.type_code: row.type_name for row in srs_types if row.type_code}
            type_codes = []
            for row in srs_types:
                if row.type_code and row.type_code not in type_codes:
                    type_codes.append(row.type_code)
            for row in change_req_rows:
                if row.type_code and row.type_code not in type_codes:
                    type_codes.append(row.type_code)
            results = []
            for type_code in type_codes:
                table = await __query_srs_reqs(type_code)
                if table and getattr(table, "rows", None):
                    results.append(SrsNodeForm(label=type_names.get(type_code) or "变更需求", table=table))
            return results
        
        async def __query_srs_change_req_exports_only():
            change_type_codes: List[str] = []
            type_names: Dict[str, str] = {}
            srs_types: List[SrsType] = db.session.execute(
                select(SrsType).where(SrsType.doc_id == doc_id).order_by(SrsType.id)
            ).scalars().all()
            for row in srs_types:
                code = str(row.type_code or "")
                if code and code not in ["1", "2"] and code not in change_type_codes:
                    change_type_codes.append(code)
                    type_names[code] = row.type_name or "变更需求"
            resp = await srsreq_serv.list_srs_req(doc_id=doc_id, page_size=5000)
            all_reqs: List[SrsReq] = resp.data.rows or []
            for req in all_reqs:
                code = str(getattr(req, "type_code", "") or "")
                if code and code not in ["1", "2"] and code not in change_type_codes:
                    change_type_codes.append(code)
            results = []
            for type_code in change_type_codes:
                table = await __query_srs_reqs(type_code)
                rows = getattr(table, "rows", None) or []
                if table and rows:
                    results.append(SrsNodeForm(label=type_names.get(type_code) or "变更需求", table=table))
            return results
        
        def __fix_chapter(p_title: str, nodes: List[SrsNodeForm]):
            chapter =re.search(r'(\d(\.\d)*)', p_title or "")
            chapter = chapter.group() if chapter else None
            chapter = f"{chapter}." if chapter else ""
            for idx, node in enumerate(nodes or []):
                if node.with_chapter == 1 and chapter and node.title:
                    node.title = f"{chapter}{idx+1} {node.title}"
                    __fix_chapter(node.title, node.children)
        
        async def __query_srs_reqds(p_title: str):
            reqds: List[SrsReqdObj] = (await srsreqd_serv.list_srs_reqd(doc_id=doc_id, page_size=2000)).data.rows
            parents = dict()
            for reqd in reqds:
                headers = [TabHeader(code="attr"), TabHeader(code="value")]
                rows = [
                    dict(attr="需求编号", value=reqd.code),
                    dict(attr="需求名称", value=reqd.name),
                    dict(attr="需求概述", value=reqd.overview),
                    dict(attr="主参加者", value=reqd.participant),
                    dict(attr="前置条件", value=reqd.pre_condition),
                    dict(attr="触发器", value=reqd.trigger),
                    dict(attr="工作流", value=reqd.work_flow),
                    dict(attr="后置条件", value=reqd.post_condition),
                    dict(attr="异常情况", value=reqd.exception),
                    dict(attr="约束", value=reqd.constraint),
                ]
                table = Table(headers=headers, rows=rows, show_header=0)
                p_node = find_parent(SrsNodeForm, [reqd.module, reqd.function], parents)
                with_chapter = 1 if reqd.sub_function else 0
                title = reqd.name if reqd.sub_function else None
                p_node.children.append(SrsNodeForm(with_chapter=with_chapter, title=title, table=table))
            p_nodes = [node for key, node in parents.items() if node.level == 0]
            __fix_chapter(p_title, p_nodes)
            return p_nodes

        image_caption_no = {"value": 0}

        use_snapshot_content = snapshot is not None

        def __child_label_text(child):
            return str(getattr(child, "label", "") or getattr(child, "title", "") or getattr(getattr(child, "table", None), "name", "") or "")

        def __is_snapshot_product_req_child(child):
            label = __child_label_text(child)
            return ("产品需求列表" in label or "产品需求" in label) and "其他需求" not in label

        def __is_snapshot_other_req_child(child):
            label = __child_label_text(child)
            return "其他需求列表" in label or label.strip() == "其他需求" or label.startswith("其他需求")

        def __is_snapshot_change_req_child(child):
            if "变更" in __child_label_text(child):
                return True
            return "变更" in str(getattr(getattr(child, "table", None), "name", "") or "")

        def __find_snapshot_req_child(node, matcher):
            for child in node.children or []:
                if matcher(child) and child.table and child.table.headers:
                    return child
            return None

        def __find_snapshot_change_req_children(node):
            results = []
            for child in node.children or []:
                if __is_snapshot_change_req_child(child) and child.table and child.table.headers:
                    results.append(child)
            return results

        def __write_snapshot_req_tables_after_text(node, docx, font_def, written_child_ids):
            if not use_snapshot_content:
                return
            for child in node.children or []:
                if id(child) in written_child_ids:
                    continue
                if not (
                    __is_snapshot_product_req_child(child) or
                    __is_snapshot_other_req_child(child) or
                    __is_other_req_export_table(getattr(child, "table", None)) or
                    __is_snapshot_change_req_child(child)
                ):
                    continue
                __write_snapshot_req_child(child, docx, font_def, written_child_ids)

        def __write_snapshot_req_child(child, docx, font_def, written_child_ids, skip_label: bool = False):
            label = getattr(child, "label", None) or ""
            table_name = str(getattr(getattr(child, "table", None), "name", "") or "")
            norm_label = __norm_title(label) or __norm_title(table_name)
            is_other_child = __is_snapshot_other_req_child(child) or __is_other_req_export_table(getattr(child, "table", None))
            if is_other_child and not skip_label:
                __write_other_req_title_if_needed(docx, font_def, label or table_name)
            elif label and not skip_label and norm_label not in exported_req_labels:
                exported_req_labels.add(norm_label)
                docx_util.save_txt2docx(label, docx, font_def)
            elif label and not skip_label:
                exported_req_labels.add(norm_label)
            if child.table and child.table.headers:
                rows = getattr(child.table, "rows", None) or []
                codes = []
                for row in rows[:5]:
                    if isinstance(row, dict):
                        codes.append(str(row.get("srs_code") or row.get("code") or ""))
                dedupe_key = (norm_label, tuple(codes), len(rows))
                if dedupe_key in exported_req_tables:
                    written_child_ids.add(id(child))
                    return
                exported_req_tables.add(dedupe_key)
                __save_tab2docx(child.table, docx, font_def, show_name=False)
            written_child_ids.add(id(child))

        def __change_req_table_dedupe_key(child):
            label = getattr(child, "label", None) or ""
            table_name = str(getattr(getattr(child, "table", None), "name", "") or "")
            norm_label = __norm_title(label) or __norm_title(table_name)
            rows = getattr(getattr(child, "table", None), "rows", None) or []
            codes = []
            for row in rows[:5]:
                if isinstance(row, dict):
                    codes.append(str(row.get("srs_code") or row.get("code") or ""))
            return (norm_label, tuple(codes), len(rows))

        def __is_change_req_table_exported(child):
            return __change_req_table_dedupe_key(child) in exported_req_tables

        def __write_change_req_table_node(table, docx, table_name: str = ""):
            if not table or not getattr(table, "headers", None):
                return False
            name = str(table_name or getattr(table, "name", "") or "变更需求").strip()
            norm_name = __norm_title(name)
            rows = getattr(table, "rows", None) or []
            codes = []
            for row in rows[:5]:
                if isinstance(row, dict):
                    codes.append(str(row.get("srs_code") or row.get("code") or ""))
            dedupe_key = (norm_name, tuple(codes), len(rows))
            if dedupe_key in exported_req_tables:
                return False
            exported_req_tables.add(dedupe_key)
            __save_tab2docx(table, docx, show_name=False)
            return True

        def __has_change_req_tables_exported():
            return any("变更" in str((key or [""])[0] or "") for key in exported_req_tables)

        async def __export_change_req_from_db(docx, font_def, title_already_written: bool = False):
            """变更需求表统一从数据库读取，在需求列表节点指定位置导出。"""
            if change_req_export_done["value"]:
                return False
            extras = await __query_srs_change_req_exports_only()
            if not extras:
                return False
            wrote_any = False
            for extra in extras:
                table = extra.table
                rows = getattr(table, "rows", None) or []
                if not table or not getattr(table, "headers", None) or not rows:
                    continue
                label = extra.label or "变更需求"
                norm_label = __norm_title(label)
                codes = []
                for row in rows[:5]:
                    if isinstance(row, dict):
                        codes.append(str(row.get("srs_code") or row.get("code") or ""))
                dedupe_key = (norm_label, tuple(codes), len(rows))
                if dedupe_key in exported_req_tables:
                    continue
                __write_change_req_title_if_needed(label, docx, font_def)
                exported_req_tables.add(dedupe_key)
                __save_tab2docx(table, docx, font_def, show_name=False)
                wrote_any = True
            if wrote_any:
                change_req_export_done["value"] = True
            return wrote_any

        async def __save_tab_and_export_change_if_other(table, docx, font_def, show_name: bool = False, prepared: bool = False):
            __save_tab2docx(table, docx, font_def, show_name=show_name, prepared=prepared)
            if __is_other_req_export_table(table) and not change_req_export_done["value"]:
                await __export_change_req_from_db(docx, font_def, "变更需求" in exported_req_labels)

        def __is_change_req_line(line: str):
            return "变更" in __norm_title(line)

        def __write_change_req_title_if_needed(title: str, docx, font_def):
            label = str(title or "变更需求").strip()
            norm_label = __norm_title(label)
            if not norm_label or norm_label in exported_req_labels:
                return False
            docx_util.save_txt2docx(label, docx, font_def)
            exported_req_labels.add(norm_label)
            return True

        def __is_req_list_export_slot(node_text: str, node_label: str, imported_table_children):
            text_blob = f"{node_text or ''}\n{node_label or ''}"
            has_product = "产品需求" in text_blob or "产品功能" in text_blob
            has_other = "其他需求" in text_blob
            has_pair = has_product and has_other
            has_imported_pair = len(imported_table_children or []) >= 2
            return has_pair or has_imported_pair

        db_doc_roots = []
        req_table_sources = []

        async def __writenodes(nodes: List[SrsNodeForm], docx: Document, level: int = 0):
            font_def = 10.5
            font_size = font_def
            if level == 0 :
                font_size = 16.0
            elif level == 1:
                font_size = 14.0
            font_name = "宋体"
            for node in nodes or []:
                if __is_imported_catalog_root(node):
                    continue
                raw_node_title = getattr(node, "title", "") or ""
                if __is_imported_catalog_title(raw_node_title):
                    continue
                node_title_for_export = __strip_imported_catalog_suffix(raw_node_title)
                written_child_ids = set()
                is_catalog_root = level == 0 and __norm_title(node_title_for_export) == "目录"
                node_image_caption = ""
                node_image_written = False
                if node_title_for_export:
                    if __is_imported_placeholder_title(node_title_for_export):
                        # 过滤系统中间占位标题，导出时不显示“导入表格X/导入图片X”等字样
                        pass
                    elif __is_image_caption_line(node_title_for_export) and node.img_url:
                        node_image_caption = node_title_for_export
                    elif is_catalog_root:
                        __write_catalog_page(docx, "")
                    elif level == 0 and __is_cover_section_title(node_title_for_export):
                        if __is_spec_title(node_title_for_export):
                            # 需求规格说明标题向上预留10行
                            __add_blank_lines(docx, 10)
                        __write_center_section_title(docx, node_title_for_export)
                        # 标题与其下方表格之间保留空白
                        __add_blank_lines(docx, 9 if __is_spec_title(node_title_for_export) else 2)
                    elif __is_imported_catalog_title(node_title_for_export) and not (node.text or node.table or node.img_url):
                        pass
                    else:
                        docx_util.save_title2docx(node_title_for_export, docx, level+1, font_size)
                if is_catalog_root:
                    # 目录页由TOC域生成，不再输出旧的目录文本和子节点
                    continue

                is_auto_req_node = str(getattr(node, "label", "") or "") in ["__auto_req_group", "__auto_req_detail"]
                if node.srs_code and not is_auto_req_node:
                    # 若正文文本已包含同一需求编号，避免重复导出“需求编号”行
                    text_norm = (node.text or "").replace("：", ":")
                    code_norm = (node.srs_code or "").strip()
                    has_code_in_text = bool(code_norm and code_norm in text_norm and ("需求编号" in text_norm or "SRS" in text_norm))
                    if not has_code_in_text:
                        docx_util.save_txt2docx("需求编号：" + node.srs_code, docx, font_def)
                if node.label:
                    if is_auto_req_node:
                        pass
                    elif __is_image_caption_line(node.label) and node.img_url:
                        node_image_caption = node_image_caption or node.label
                    else:
                        docx_util.save_txt2docx(node.label, docx, font_def)
                if (
                    node.table and
                    node.table.headers and
                    "变更" in str(getattr(node.table, "name", "") or "")
                ):
                    # 变更需求表仅从数据库导出，跳过快照中的独立变更表节点
                    continue
                node_text_for_export = __strip_imported_catalog_lines(node.text)
                imported_table_children = [
                    child for child in (node.children or [])
                    if __is_imported_table_title(child.title) and child.table and child.table.headers
                ]
                if node_text_for_export or imported_table_children:
                    imported_image_children = [
                        child for child in (node.children or [])
                        if (child.img_url and re.match(r"^导入图片\d*$", (child.title or "").strip()))
                    ]
                    lines = (node_text_for_export or "").splitlines()
                    has_caption = any(__is_table_caption_line(line) for line in lines)
                    has_image_caption = any(__is_image_caption_line(line) for line in lines)
                    has_change_req_marker = "变更需求" in node_text_for_export or "变更需求" in node_title_for_export
                    has_req_list_pair = __is_req_list_export_slot(
                        node_text_for_export,
                        getattr(node, "label", "") or "",
                        imported_table_children,
                    )
                    if has_req_list_pair:
                        if imported_table_children:
                            table_idx = 0
                            for raw_line in lines:
                                line = (raw_line or "").strip()
                                if not line:
                                    continue
                                if __is_change_req_line(line):
                                    if __norm_title(line) in exported_req_labels:
                                        continue
                                    exported_req_labels.add(__norm_title(line))
                                    docx_util.save_txt2docx(line, docx, font_def)
                                    continue
                                docx_util.save_txt2docx(line, docx, font_def)
                                if "其他需求列表" in line or ("其他需求" in line and "产品需求" not in line):
                                    __mark_other_req_title_from_text(line)
                                if ("产品需求列表" in line or "产品需求" in line or "其他需求列表" in line or "其他需求" in line) and table_idx < len(imported_table_children):
                                    tab_node = imported_table_children[table_idx]
                                    table_idx += 1
                                    if "其他需求列表" in line or ("其他需求" in line and "产品需求" not in line):
                                        await __export_db_req_list_table("other", docx, font_def, is_other=True)
                                    else:
                                        await __export_db_req_list_table("main", docx, font_def, is_other=False)
                                    written_child_ids.add(id(tab_node))
                            for tab_node in imported_table_children[table_idx:]:
                                table_name = str(getattr(tab_node.table, "name", "") or getattr(tab_node, "label", "") or getattr(tab_node, "title", "") or "")
                                if "变更" in table_name:
                                    continue
                                if __is_other_req_export_table(tab_node.table):
                                    await __export_db_req_list_table("other", docx, font_def, is_other=True)
                                elif __is_main_req_export_table(tab_node.table):
                                    __write_table_caption_once(tab_node.table, docx, font_def, node_text_for_export)
                                    await __export_db_req_list_table("main", docx, font_def, is_other=False)
                                else:
                                    __write_table_caption_once(tab_node.table, docx, font_def, node_text_for_export)
                                    __save_tab2docx(tab_node.table, docx, font_def, show_name=False)
                                written_child_ids.add(id(tab_node))
                            if not change_req_export_done["value"]:
                                await __export_change_req_from_db(docx, font_def)
                        else:
                            for raw_line in lines:
                                line = (raw_line or "").strip()
                                if not line:
                                    continue
                                if __is_change_req_line(line):
                                    if __norm_title(line) in exported_req_labels:
                                        continue
                                    exported_req_labels.add(__norm_title(line))
                                    docx_util.save_txt2docx(line, docx, font_def)
                                    continue
                                docx_util.save_txt2docx(line, docx, font_def)
                                if "其他需求列表" in line or ("其他需求" in line and "产品需求" not in line):
                                    __mark_other_req_title_from_text(line)
                                if "产品需求列表" in line or "产品需求" in line:
                                    await __export_db_req_list_table("main", docx, font_def, is_other=False)
                                elif "其他需求列表" in line or "其他需求" in line:
                                    await __export_db_req_list_table("other", docx, font_def, is_other=True)
                            if not change_req_export_done["value"]:
                                await __export_change_req_from_db(docx, font_def)
                        for child in (node.children or []):
                            if __is_snapshot_change_req_child(child):
                                written_child_ids.add(id(child))
                    elif has_change_req_marker:
                        if node_text_for_export:
                            for raw_line in lines:
                                line = (raw_line or "").strip()
                                if not line:
                                    continue
                                if __is_change_req_line(line):
                                    if __norm_title(line) in exported_req_labels:
                                        continue
                                    exported_req_labels.add(__norm_title(line))
                                docx_util.save_txt2docx(line, docx, font_def)
                        for child in (node.children or []):
                            if __is_snapshot_change_req_child(child):
                                written_child_ids.add(id(child))
                    elif (imported_table_children and has_caption) or (imported_image_children and has_image_caption) or (has_image_caption and node.img_url):
                        table_idx = 0
                        image_idx = 0
                        for raw_line in lines:
                            line = (raw_line or "").strip()
                            if not line:
                                continue
                            if __is_table_caption_line(line) and table_idx < len(imported_table_children):
                                docx_util.save_txt2docx(line, docx, font_def)
                                tab_node = imported_table_children[table_idx]
                                table_idx += 1
                                if __is_other_req_export_table(tab_node.table):
                                    await __export_db_req_list_table("other", docx, font_def, is_other=True)
                                elif __is_main_req_export_table(tab_node.table):
                                    await __export_db_req_list_table("main", docx, font_def, is_other=False)
                                else:
                                    __save_tab2docx(tab_node.table, docx, font_def, show_name=False)
                                written_child_ids.add(id(tab_node))
                            elif __is_image_caption_line(line):
                                img_url = None
                                if image_idx < len(imported_image_children):
                                    img_node = imported_image_children[image_idx]
                                    image_idx += 1
                                    img_url = img_node.img_url
                                    written_child_ids.add(id(img_node))
                                elif node.img_url:
                                    img_url = node.img_url
                                if img_url:
                                    docx_util.save_img2docx(img_url, docx, mw=500, mh=500)
                                    __save_image_caption_txt(docx, line, font_def)
                                    node_image_written = True
                                else:
                                    docx_util.save_txt2docx(line, docx, font_def)
                            else:
                                docx_util.save_txt2docx(line, docx, font_def)
                    else:
                        if imported_table_children:
                            table_idx = 0
                            for raw_line in lines:
                                line = (raw_line or "").strip()
                                if not line:
                                    continue
                                docx_util.save_txt2docx(line, docx, font_def)
                                if "其他需求列表" in line or ("其他需求" in line and "产品需求" not in line):
                                    __mark_other_req_title_from_text(line)
                                if __is_change_req_line(line):
                                    exported_req_labels.add(__norm_title(line))
                                elif "变更" in line:
                                    exported_req_labels.add(__norm_title(line))
                                if table_idx < len(imported_table_children) and (
                                    "产品需求列表" in line or "产品需求" in line or
                                    "其他需求列表" in line or "其他需求" in line or
                                    "变更需求" in line or "变更" in line
                                ):
                                    tab_node = imported_table_children[table_idx]
                                    table_idx += 1
                                    if "变更需求" in line or (__is_change_req_line(line) and "产品需求" not in line and "其他需求" not in line):
                                        __save_tab2docx(tab_node.table, docx, font_def, show_name=False)
                                    elif __is_other_req_export_table(tab_node.table) or "其他需求" in line:
                                        await __export_db_req_list_table("other", docx, font_def, is_other=True)
                                    else:
                                        await __export_db_req_list_table("main", docx, font_def, is_other=False)
                                    written_child_ids.add(id(tab_node))
                            for tab_node in imported_table_children[table_idx:]:
                                if id(tab_node) in written_child_ids:
                                    continue
                                if __is_other_req_export_table(tab_node.table):
                                    await __export_db_req_list_table("other", docx, font_def, is_other=True)
                                elif __is_main_req_export_table(tab_node.table):
                                    __write_table_caption_once(tab_node.table, docx, font_def, node_text_for_export)
                                    await __export_db_req_list_table("main", docx, font_def, is_other=False)
                                else:
                                    __write_table_caption_once(tab_node.table, docx, font_def, node_text_for_export)
                                    __save_tab2docx(tab_node.table, docx, font_def, show_name=False)
                                written_child_ids.add(id(tab_node))
                        else:
                            docx_util.save_txt2docx(node_text_for_export, docx, font_def)

                    if not change_req_export_done["value"]:
                        text_blob = f"{node_text_for_export or ''}\n{getattr(node, 'label', '') or ''}"
                        if "其他需求" in text_blob or any(
                            __is_other_req_export_table(getattr(child, "table", None))
                            for child in imported_table_children
                        ):
                            await __export_change_req_from_db(docx, font_def)

                if node.img_url and not node_image_written:
                    docx_util.save_img2docx(node.img_url, docx, mw=500, mh=500)
                    if not node_image_caption:
                        default_caption_name = __strip_heading_no(node_title_for_export)
                        if default_caption_name and not __is_imported_placeholder_title(default_caption_name):
                            image_caption_no["value"] += 1
                            node_image_caption = f"图{image_caption_no['value']} {default_caption_name}"
                    if node_image_caption:
                        __save_image_caption_txt(docx, node_image_caption, font_def)

                if node.ref_type == RefTypes.srs_reqs.value:
                    has_table_children = any(
                        child.table and child.table.headers
                        for child in (node.children or [])
                    )
                    if not has_table_children:
                        table = await __query_srs_reqs("1")
                        node1 = SrsNodeForm(label="产品需求列表:", table=table)

                        table = await __query_srs_reqs("2")
                        node2 = SrsNodeForm(label="其他需求列表:", table=table)

                        await __writenodes([node1, node2], docx, level + 1)
                        if not change_req_export_done["value"]:
                            await __export_change_req_from_db(docx, font_def, False)
                elif node.ref_type == RefTypes.srs_reqds.value:
                    reqds = await __query_srs_reqds(node.title)
                    await __writenodes(reqds, docx, level + 1)   
                else:
                    has_table_children = any(
                        child.table and child.table.headers
                        for child in (node.children or [])
                    )
                    if node.table and node.table.headers and not has_table_children:
                        if __is_main_req_export_table(node.table) or __is_other_req_export_table(node.table):
                            kind = "other" if __is_other_req_export_table(node.table) else "main"
                            await __export_db_req_list_table(kind, docx, font_def, is_other=(kind == "other"))
                        else:
                            __save_tab2docx(node.table, docx, font_def)

                if node.children:
                    next_children = [child for child in node.children if id(child) not in written_child_ids]
                    await __writenodes(next_children, docx, level + 1)

        resp = await self.get_srs_doc(doc_id, with_tree=True)
        srs_doc: SrsDocObj = resp.data
        db_doc_roots[:] = deepcopy(srs_doc.content or []) if srs_doc else []
        req_table_sources[:] = [deepcopy(db_doc_roots)]
        if srs_doc:
            if snapshot is not None:
                snapshot_content = getattr(snapshot, "content", None)
                if snapshot_content:
                    req_table_sources.append(deepcopy(snapshot_content))
                for attr in ["product_id", "version", "folder_name", "file_no", "change_log", "content"]:
                    value = getattr(snapshot, attr, None)
                    if value is not None:
                        setattr(srs_doc, attr, value)
            product = db.session.execute(select(Product).where(Product.id == srs_doc.product_id)).scalars().first()
            product_version = getattr(product, "full_version", "") or getattr(product, "name", "") or ""
            prod_imgs = self.__query_imgs(srs_doc.product_id, srs_doc.version, product_version)
            self.__hydrate_tree_product_doc_images(srs_doc.content or [], prod_imgs)
            docx = Document()
            docx_util.enable_update_fields_on_open(docx)

            header_para = docx.sections[0].header.add_paragraph()
            header_para.alignment = dox_enum.text.WD_ALIGN_PARAGRAPH.RIGHT
            docx_util.fonted_txt(header_para, srs_doc.file_no)

            roots = srs_doc.content or []
            spec_root = next((n for n in roots if "需求规格说明" in __norm_title(n.title)), None)
            rev_root = next((n for n in roots if "文件修订记录" in __norm_title(n.title)), None)
            catalog_root = next((n for n in roots if "目录" in __norm_title(n.title)), None)
            import_root = next(
                (
                    n for n in roots
                    if __norm_title(getattr(n, "title", "") or "") == "导入正文"
                    and ("需求规格说明" in str(getattr(n, "text", "") or "") or "目录" in str(getattr(n, "text", "") or ""))
                ),
                None,
            )
            used_ids = {id(node) for node in [spec_root, rev_root, catalog_root, import_root] if node is not None}
            remaining_roots = [n for n in roots if id(n) not in used_ids]

            # 参考详细设计导出：兼容历史导入数据把“封面/修订记录/正文”都挂在根节点下的情况。
            spec_section_nodes = [spec_root] if spec_root else [SrsNodeForm(title="需求规格说明", children=[])]
            rev_section_nodes = [rev_root] if rev_root else []
            body_from_spec = []
            if spec_root:
                cover_node = SrsNodeForm(title="需求规格说明", children=[])
                cover_table_picked = False
                rev_nodes_from_spec = []
                for child in (spec_root.children or []):
                    if __node_has_revision_marker(child) or __is_revision_table(getattr(child, "table", None)):
                        rev_nodes_from_spec.append(child)
                        continue
                    if (not cover_table_picked) and getattr(child, "table", None) and not __is_revision_table(child.table):
                        cover_node.children.append(child)
                        cover_table_picked = True
                        continue
                    body_from_spec.append(child)

                spec_section_nodes = [cover_node]
                if (not rev_section_nodes) and rev_nodes_from_spec:
                    rev_section_nodes = [SrsNodeForm(title="文件修订记录", children=rev_nodes_from_spec)]

            if (not spec_root) and import_root:
                cover_node = SrsNodeForm(title="需求规格说明", children=[])
                rev_nodes_from_import = []
                cover_table_picked = False
                for child in (import_root.children or []):
                    if __node_has_revision_marker(child) or __is_revision_table(getattr(child, "table", None)):
                        rev_nodes_from_import.append(child)
                        continue
                    if (not cover_table_picked) and getattr(child, "table", None):
                        cover_node.children.append(child)
                        cover_table_picked = True
                spec_section_nodes = [cover_node]
                if (not rev_section_nodes) and rev_nodes_from_import:
                    rev_section_nodes = [SrsNodeForm(title="文件修订记录", children=rev_nodes_from_import)]

            if not rev_section_nodes:
                rev_section_nodes = [SrsNodeForm(title="文件修订记录", children=[])]

            remaining_roots = body_from_spec + remaining_roots
            imported_catalog_text = __extract_imported_catalog_text(*(roots or []))

            export_sections = [
                ("spec", spec_section_nodes),
                ("rev", rev_section_nodes),
                ("catalog", [catalog_root] if catalog_root else [SrsNodeForm(title="目录", children=[])]),
                ("body", remaining_roots),
            ]
            first_section = True
            for section_name, section_nodes in export_sections:
                if not section_nodes:
                    continue
                if not first_section:
                    docx.add_page_break()
                if section_name == "catalog":
                    __write_catalog_page(docx, imported_catalog_text)
                else:
                    await __writenodes(section_nodes, docx, level=0)
                if section_name == "rev":
                    # 第二页文件修订记录表格后保留5行空白
                    __add_blank_lines(docx, 5)
                first_section = False

            docx.save(output)
            output.seek(0)

    async def add_doc_file(self, doc_id: int, file):
        size, path = await save_file("srs_node_img", doc_id, file)
        return Resp.resp_ok(data=path)  

    async def list_doc_trace(self, id: int):
        def __build_trace_rule_from_srs_code(srs_code: str):
            code = str(srs_code or "").strip().upper()
            matched = re.match(r"^SRS-([A-Z]+)(\d+)-(\d+)$", code)
            if not matched:
                return None
            prefix = matched.group(1)
            major = matched.group(2)
            minor = matched.group(3)
            if len(major) < 2:
                return None
            if_code = major[-2:]
            minor_int = str(int(minor)) if minor.isdigit() else minor
            minor3 = minor.zfill(3)
            return {
                "if_code": if_code,  # IF00 / IF06
                "unit_group": minor3,  # 005 / 003
                "sis_prefix": f"SDS-IF{if_code}-{prefix}{major}{minor_int}-",  # SDS-IF00-RCN3005-
                "unit_prefix": f"TU{if_code}-{minor3}-",  # TU00-005-
            }

        def __query_rcms(srs_ids: list[int]):
            sql = select(ReqRcm, Rcm).join(Rcm, ReqRcm.rcm_id == Rcm.id).where(ReqRcm.req_id.in_(srs_ids))
            rows: list[ReqRcm, Rcm] = db.session.execute(sql).all()
            req_rcms = dict()
            for row_req, row_rcm in rows:
                rcms = req_rcms.get(row_req.req_id) or []
                rcms.append(row_rcm)
                req_rcms[row_req.req_id] = rcms
            return req_rcms

        def __query_reqd_text_rcm_codes(req_ids: list[int]):
            if not req_ids:
                return {}
            sql = select(SrsReqd).where(SrsReqd.req_id.in_(req_ids))
            rows: List[SrsReqd] = db.session.execute(sql).scalars().all()
            result: dict[int, list[str]] = {}
            for row in rows:
                merged_text = "\n".join([
                    str(getattr(row, "name", "") or ""),
                    str(getattr(row, "overview", "") or ""),
                    str(getattr(row, "participant", "") or ""),
                    str(getattr(row, "pre_condition", "") or ""),
                    str(getattr(row, "trigger", "") or ""),
                    str(getattr(row, "work_flow", "") or ""),
                    str(getattr(row, "post_condition", "") or ""),
                    str(getattr(row, "exception", "") or ""),
                    str(getattr(row, "constraint", "") or ""),
                ])
                dedup = self.__extract_rcm_codes_from_text(merged_text)
                if dedup:
                    result[int(row.req_id)] = dedup
            return result

        def __query_node_rcm_codes(doc_id: int, srs_codes: list[str]):
            if not doc_id or not srs_codes:
                return {}
            srs_set = {str(code or "").strip().upper() for code in (srs_codes or []) if str(code or "").strip()}
            sql = select(
                SrsNode.n_id,
                SrsNode.p_id,
                SrsNode.srs_code,
                SrsNode.rcm_codes,
                SrsNode.title,
                SrsNode.label,
                SrsNode.text,
                SrsNode.table,
            ).where(SrsNode.doc_id == doc_id)
            rows = db.session.execute(sql).all()
            node_map = {
                int(n_id): {
                    "p_id": int(p_id or 0),
                    "srs_code": str(srs_code or "").strip().upper(),
                    "rcm_codes": str(rcm_codes or ""),
                    "title": str(title or ""),
                    "label": str(label or ""),
                    "text": str(text or ""),
                    "table": table,
                }
                for n_id, p_id, srs_code, rcm_codes, title, label, text, table in rows
            }

            def resolve_srs_code(start_nid: int) -> str:
                cur = node_map.get(start_nid)
                safety = 0
                while cur and safety < 200:
                    code = str(cur.get("srs_code") or "").strip().upper()
                    if code:
                        return code
                    pid = int(cur.get("p_id") or 0)
                    if pid <= 0:
                        break
                    cur = node_map.get(pid)
                    safety += 1
                return ""

            result: dict[str, list[str]] = {}
            for n_id, _p_id, own_code, raw_codes, title, label, text, table in rows:
                merged_text = "\n".join([
                    str(title or ""),
                    str(label or ""),
                    str(text or ""),
                    self.__table_search_text(table) if table is not None else "",
                ])
                parts = [str(item or "").strip() for item in re.split(r"[,，;；\s]+", str(raw_codes or "")) if str(item or "").strip()]
                if not parts:
                    parts = self.__extract_rcm_codes_from_text(merged_text)
                if not parts:
                    continue
                # 兜底：节点正文中显式出现了 SRS 编号时，直接按“编号->RCM”关联
                mentioned_srs_codes = []
                for hit in re.findall(r"SRS[\s\-]*[A-Z]+[\s\-]*\d{2,4}\s*-\s*\d{3}", merged_text, flags=re.IGNORECASE):
                    code_norm = self.__normalize_srs_code(re.sub(r"\s+", "", hit).replace("--", "-"))
                    if code_norm and code_norm not in mentioned_srs_codes:
                        mentioned_srs_codes.append(code_norm)
                for mcode in mentioned_srs_codes:
                    if srs_set and mcode not in srs_set:
                        continue
                    existed = result.get(mcode) or []
                    result[mcode] = list(dict.fromkeys(existed + parts))
                srs_code = str(own_code or "").strip().upper() or resolve_srs_code(int(n_id))
                if not srs_code:
                    continue
                if srs_set and srs_code not in srs_set:
                    continue
                existed = result.get(srs_code) or []
                merged = list(dict.fromkeys(existed + parts))
                result[srs_code] = merged
            return result
        
        def __query_tests(product_id: int) -> dict:
            def __normalize_stage(stage: str) -> str:
                txt = str(stage or "").strip()
                if "单元" in txt:
                    return "单元测试"
                if "集成" in txt:
                    return "集成测试"
                if "系统" in txt:
                    return "系统测试"
                if "用户" in txt:
                    return "用户测试"
                return txt

            def __to_srs_code(raw_code: str) -> list[str]:
                code = str(raw_code or "").strip().upper()
                if not code:
                    return []
                candidates: list[str] = []
                # 兼容历史接口编码：SDS-IF{xx}-RUS{yy}-...
                matched = re.match(r"SDS-IF(\d+)-RUS(\d+)-\d+", code)
                if matched:
                    candidates.append(f"SRS-RUS{matched.group(1)}-{matched.group(2)}")
                # 兼容接口编码：SDS-IF00-RCN300-005 / SDS-IF00-XXX...
                matched2 = re.match(r"SDS-IF\d+-([A-Z0-9]+-\d+)$", code)
                if matched2:
                    candidates.append(f"SRS-{matched2.group(1)}")
                    # 兼容压缩段：SDS-IF00-RCN3005-001 -> SRS-RCN300-005
                    seg = matched2.group(1)
                    seg_parts = seg.split("-")
                    if len(seg_parts) == 2:
                        left, right = seg_parts
                        m3 = re.match(r"^([A-Z]+)(\d{3,})(\d)$", left)
                        if m3:
                            req_group = str(m3.group(3) or "").zfill(3)
                            candidates.append(f"SRS-{m3.group(1)}{m3.group(2)}-{req_group}")
                # 若本身已是 SRS 编号也接受
                if code.startswith("SRS-"):
                    candidates.append(code)
                # 去重并保持顺序
                seen = set()
                uniq = []
                for item in candidates:
                    if item and item not in seen:
                        seen.add(item)
                        uniq.append(item)
                return uniq

            sql = select(TestCase, TestSet).join(TestSet, TestCase.set_id == TestSet.id).where(TestSet.product_id == product_id).order_by(TestCase.code)
            rows: List[Tuple[TestCase, TestSet]] = db.session.execute(sql).all()
            all_tests = dict()
            stage_code_index = dict()
            # 单元测试用例编号 -> 接口编号(SDS-IF...) 映射
            unit_case_to_sis = dict()
            for row_test, row_set in rows:
                norm_stage = __normalize_stage(row_set.stage)
                all_tests.setdefault(row_test.srs_code, {}).setdefault(norm_stage, []).append(row_test)
                code = str(row_test.code or "").strip().upper()
                if code:
                    stage_code_index.setdefault(norm_stage, []).append(code)
                code = str(row_test.code or "").strip().upper()
                srs_code_raw = str(row_test.srs_code or "").strip().upper()
                if code and srs_code_raw.startswith("SDS-IF"):
                    unit_case_to_sis.setdefault(code, set()).add(srs_code_raw)
            
            test_codes = dict()
            for srs_code, test_data in all_tests.items():
                for stage, items in test_data.items():
                    tests = list({item.code: item.code for item in items}.keys())
                    tests = [tests[0], tests[len(tests)-1]] if len(tests) > 1 else tests[:1]
                    test_codes.setdefault(srs_code, {}).setdefault(stage, []).extend(tests)

            req_pairs = dict()
            sds_uset = dict()
            for row_test, row_set in rows:
                src_code = str(row_test.srs_code or "").strip().upper()
                if not src_code:
                    continue
                # 接口编号列只取“单元测试”阶段中的接口测试编码（SDS-IF...）
                norm_stage = __normalize_stage(getattr(row_set, "stage", "") or "")
                if norm_stage != "单元测试":
                    continue
                if not src_code.startswith("SDS-IF"):
                    continue
                for srs_code in __to_srs_code(src_code):
                    uset = sds_uset.setdefault(srs_code, set())
                    if src_code not in uset:
                        uset.add(src_code)
                        req_pairs.setdefault(srs_code, []).append((src_code, row_test.code))

            # 补充链路：通过“单元测试用例编号”反查接口编号，填充到对应 SRS 需求
            for srs_code, test_data in all_tests.items():
                srs_code_norm = str(srs_code or "").strip().upper()
                if not srs_code_norm.startswith("SRS-"):
                    continue
                unit_items = test_data.get("单元测试") or []
                uset = sds_uset.setdefault(srs_code_norm, set())
                for item in unit_items:
                    unit_code = str(getattr(item, "code", "") or "").strip().upper()
                    if not unit_code:
                        continue
                    sis_codes = sorted(unit_case_to_sis.get(unit_code) or [])
                    for sis_code in sis_codes:
                        if sis_code in uset:
                            continue
                        uset.add(sis_code)
                        req_pairs.setdefault(srs_code_norm, []).append((sis_code, unit_code))

            for srs_code, codes in req_pairs.items():
                codes.sort(key=lambda x: x[0])
            for stage, codes in stage_code_index.items():
                uniq = list(dict.fromkeys(codes))
                stage_code_index[stage] = uniq
            all_stage_codes = []
            for _, codes in stage_code_index.items():
                all_stage_codes.extend(codes)
            stage_code_index["__all__"] = list(dict.fromkeys(all_stage_codes))
            return test_codes, req_pairs, stage_code_index
        
        def __resort_rows(rows: List[Tuple[SdsTrace, SdsDoc, SrsReq]], srsdoc_id: int):
            # 按 SRS 导入后的原始顺序（req_id 升序）返回，确保“从安装包打开SRS第一个开始”
            sortable_rows = sorted(rows, key=lambda row: (getattr(row[2], "id", 0) or 0))
            results = []
            exist_codes = set()
            for row in sortable_rows:
                code = (row[2].code or "").strip()
                if code and code not in exist_codes:
                    exist_codes.add(code)
                    results.append(row)
            return results
     
        sql = select(SdsDoc, SrsDoc).join(SrsDoc, SdsDoc.srsdoc_id == SrsDoc.id).where(SrsDoc.id == id).order_by(desc(SdsDoc.id)).limit(1)
        row_sdsdoc, row_srsdoc = db.session.execute(sql).first() or (None, None)
        if not row_sdsdoc:
            return Resp.resp_ok(data=[])
        
        sql = select(SdsTrace, SdsDoc, SrsReq).where(SdsTrace.doc_id==SdsDoc.id, SdsTrace.req_id==SrsReq.id, SdsDoc.srsdoc_id==SrsReq.doc_id)
        sql = sql.where(SdsDoc.id == row_sdsdoc.id)
        rows: List[Tuple[SdsTrace, SdsDoc, SrsReq]] = db.session.execute(sql).all()
        rows = __resort_rows(rows, row_sdsdoc.srsdoc_id)
        # 以当前SDS文档节点为准，建立“设计编号 -> 章节标题”映射，避免导出章节名错配
        sds_code_chapter_map = {}
        node_rows = db.session.execute(
            select(SdsNode.sds_code, SdsNode.title).where(SdsNode.doc_id == row_sdsdoc.id)
        ).all()
        for code_raw, title_raw in node_rows:
            code = str(code_raw or "").strip().upper()
            if not code:
                continue
            title = str(title_raw or "").strip()
            # 章节标题通常是“3 DataProcessing”，导出时去掉前缀序号
            title = re.sub(r"^\d+\s*[\.、\-]?\s*", "", title).strip()
            if title and code not in sds_code_chapter_map:
                sds_code_chapter_map[code] = title
        req_ids = [row.id for _, _, row in rows]
        req_rcms = __query_rcms(req_ids)
        reqd_text_rcm_codes = __query_reqd_text_rcm_codes(req_ids)
        node_rcm_codes = __query_node_rcm_codes(row_srsdoc.id, [str(row.code or "").strip().upper() for _, _, row in rows])
        req_tests, req_pairs, stage_code_index = __query_tests(row_srsdoc.product_id)
        row_product = db.session.execute(select(Product).where(Product.id == row_srsdoc.product_id)).scalars().first()
        type_rows: List[SrsType] = db.session.execute(select(SrsType).where(SrsType.doc_id == row_srsdoc.id)).scalars().all()
        type_names = {str(row.type_code or ""): row.type_name for row in type_rows}
        default_type_names = {"1": "标准需求", "2": "文档需求"}
        product_code = str(getattr(row_product, "product_code", "") or "").strip()
        fixed_note_text = self.__build_trace_fixed_note_text(product_code)
        results = []
        for row_trace, _, row in rows:
            rcms: List[Rcm] = req_rcms.get(row.id) or []
            relation_rcm_codes = [rcm.code for rcm in rcms]
            text_rcm_codes = reqd_text_rcm_codes.get(int(row.id)) or []
            node_rcm_fallback = node_rcm_codes.get(str(row.code or "").strip().upper()) or []
            rcm_codes = list(dict.fromkeys(relation_rcm_codes + text_rcm_codes + node_rcm_fallback))
            if not rcm_codes:
                rcm_codes = []
            rcm_flag = True if rcm_codes else False
            sds_code_norm = str(row_trace.sds_code or "").strip().upper()
            chapter = sds_code_chapter_map.get(sds_code_norm) or ""
            if not chapter:
                chapter = row_trace.chapter or row.sub_function or row.function or row.module or ""
                chapter = NAME_DICT.get(chapter) or chapter
                chapter = chapter if row.type_code == "2" else "NeoViewer"

            test_data = req_tests.get(row.code) or {}

            tests_unit = test_data.get("单元测试") or []
            tests_integ = test_data.get("集成测试") or []
            tests_sys = test_data.get("系统测试") or []
            tests_user = test_data.get("用户测试") or []

            srs_pairs = req_pairs.get(row.code) or []
            sis_codes = [sis[0] for sis in srs_pairs]
            test_codes = [sis[1] for sis in srs_pairs]

            # 关键规则：SRS-RCN300-005 -> IF00 + RCN3005（由 SRS 编号反推并过滤）
            trace_rule = __build_trace_rule_from_srs_code(row.code)
            if trace_rule:
                # 精准匹配：若不满足规则，直接清空，不保留旧值
                pair_filtered = [
                    (str(s or ""), str(t or ""))
                    for s, t in srs_pairs
                    if str(s or "").startswith(trace_rule["sis_prefix"])
                ]
                sis_codes = [item[0] for item in pair_filtered]
                test_codes = [item[1] for item in pair_filtered]

                # 单元测试记录按 TU{IF}-{group}- 严格过滤
                unit_filtered = [
                    str(code or "")
                    for code in (tests_unit or [])
                    if str(code or "").startswith(trace_rule["unit_prefix"])
                ]
                tests_unit = test_codes if test_codes else unit_filtered
                # 系统测试严格匹配：先过滤当前值，不命中再从“系统测试阶段/全量用例”反查
                sys_prefix = f"TS{trace_rule['if_code']}-{trace_rule['unit_group']}-"
                sys_filtered = [
                    str(code or "")
                    for code in (tests_sys or [])
                    if str(code or "").startswith(sys_prefix)
                ]
                if not sys_filtered:
                    sys_stage = stage_code_index.get("系统测试") or []
                    sys_filtered = [code for code in sys_stage if str(code or "").startswith(sys_prefix)]
                if not sys_filtered:
                    sys_all = stage_code_index.get("__all__") or []
                    sys_filtered = [code for code in sys_all if str(code or "").startswith(sys_prefix)]
                tests_sys = [sys_filtered[0], sys_filtered[-1]] if len(sys_filtered) > 1 else sys_filtered

            row_srs_code = str(row.code or "").strip().upper()
            note = fixed_note_text if row_srs_code == self.TRACE_FIXED_NOTE_CODE else None
            result = dict(
                srs_code=row.code,
                type_code=row.type_code,
                type_name=type_names.get(str(row.type_code or "")) or default_type_names.get(str(row.type_code or "")) or row.type_code,
                rcm_flag=rcm_flag,

                sds_code=row_trace.sds_code,

                sis_codes=sis_codes,
                test_codes=test_codes,

                chapter=chapter,

                tests_unit=tests_unit,
                tests_integ=tests_integ,
                tests_sys=tests_sys,
                tests_user=tests_user,

                rcm_codes=rcm_codes,

                note=note
            )
            results.append(result)
        return Resp.resp_ok(data=results)

    export_columns = [
        "srs_code",
        "rcm_flag",

        "sds_code",

        "sis_codes",

        "tests_unit",
        "tests_integ",
        "tests_sys",
        "tests_user",

        "rcm_codes",

        "note"
    ]

    arr_columns = set(["tests_integ", "tests_sys", "tests_user"])

    async def export_doc_trace(self, output, id: int):
        def __slash(v):
            txt = str(v or "").strip()
            return txt if txt else "/"

        def __is_change_trace_row(obj: dict):
            type_code = str((obj or {}).get("type_code") or "").strip()
            return bool(type_code) and type_code not in ["1", "2"]

        def __safe_sheet_title(value: str):
            txt = re.sub(r"[\[\]\*:/\\?]", "", str(value or "").strip()) or "变更追溯"
            return txt[:31]

        def __format_sds_code(value: str):
            codes = [item.strip() for item in re.split(r"[,，\s]+", str(value or "").strip()) if item.strip()]
            return "\n".join(codes) if len(codes) > 1 else str(value or "").strip()

        def __write_trace_rows(ws, rows: list[dict]):
            all_subs = 0
            for ridx, obj in enumerate(rows or [], 4):
                srs_code = __slash(obj.get("srs_code"))
                rcm_flag = ts("yes") if obj.get("rcm_flag") else ts("no")
                sds_code_raw = str(obj.get("sds_code") or "").strip()
                chapter_raw = str(obj.get("chapter") or "").strip()
                hide_chapter_codes = {
                    "SDS-RCN300-001",
                    "SDS-RCN300-002",
                    "SDS-RCN300-003",
                    "SDS-RCN300-008",
                    "SDS-RCN300-009",
                    "SDS-RCN300-010",
                }
                if sds_code_raw in hide_chapter_codes or not chapter_raw:
                    sds_code = __format_sds_code(sds_code_raw)
                else:
                    sds_code = f"{__format_sds_code(sds_code_raw)}（{chapter_raw}）"
                sis_codes = obj.get("sis_codes") or []

                test_codes = obj.get("test_codes") or []
                tests_unit = " ~ ".join(obj.get("tests_unit") or [])

                tests_integ = " ~ ".join(obj.get("tests_integ") or [])
                tests_sys = " ~ ".join(obj.get("tests_sys") or [])
                tests_user = " ~ ".join(obj.get("tests_user") or [])
                rcm_codes = "\n".join(obj.get("rcm_codes") or [])
                note_raw = str(obj.get("note") or "").strip()
                note = "\n".join([item.strip() for item in note_raw.split("、") if item and item.strip()]) if note_raw else ""

                if len(sis_codes) <= 1:
                    sis_code = sis_codes[0] if sis_codes else ""
                    tests_unit = test_codes[0] if test_codes else tests_unit
                    ws.append([
                        srs_code,
                        __slash(rcm_flag),
                        __slash(sds_code),
                        __slash(sis_code),
                        __slash(tests_unit),
                        __slash(tests_integ),
                        __slash(tests_sys),
                        __slash(tests_user),
                        __slash(rcm_codes),
                        __slash(note),
                    ])
                else:
                    temp_subs = len(sis_codes) - 1
                    for idx, sis_code in enumerate(sis_codes):
                        test_code = test_codes[idx] if idx < len(test_codes) else ""
                        ws.append([
                            srs_code,
                            __slash(rcm_flag),
                            __slash(sds_code),
                            __slash(sis_code),
                            __slash(test_code),
                            __slash(tests_integ),
                            __slash(tests_sys),
                            __slash(tests_user),
                            __slash(rcm_codes),
                            __slash(note),
                        ])
                    r_idx0 = ridx + all_subs
                    r_idx1 = ridx + all_subs + temp_subs
                    all_subs += temp_subs
                    ws.merge_cells(f"A{r_idx0}:A{r_idx1}")
                    ws.merge_cells(f"B{r_idx0}:B{r_idx1}")
                    ws.merge_cells(f"C{r_idx0}:C{r_idx1}")

                    ws.merge_cells(f"F{r_idx0}:F{r_idx1}")
                    ws.merge_cells(f"G{r_idx0}:G{r_idx1}")
                    ws.merge_cells(f"H{r_idx0}:H{r_idx1}")
                    ws.merge_cells(f"I{r_idx0}:I{r_idx1}")
                    ws.merge_cells(f"J{r_idx0}:J{r_idx1}")

            align = Alignment(vertical='top')
            for row in ws.iter_rows():
                for cell in row:
                    cell.alignment = align
            # SDS/RCM/备注列按行展示
            for row_idx in range(4, ws.max_row + 1):
                for col in ["C", "I", "J"]:
                    cell = ws[f"{col}{row_idx}"]
                    cell.alignment = Alignment(vertical='top', wrap_text=True)

        resp = await self.list_doc_trace(id)

        temp_path = os.path.join(os.path.dirname(__file__), "temp_srs_doc_trace.xlsx")
        wb = load_workbook(temp_path)
        ws = wb[wb.sheetnames[0]]
        rows = resp.data or []
        normal_rows = [obj for obj in rows if not __is_change_trace_row(obj)]
        change_rows = [obj for obj in rows if __is_change_trace_row(obj)]
        ws_change = wb.copy_worksheet(ws) if change_rows else None
        __write_trace_rows(ws, normal_rows)

        if change_rows and ws_change is not None:
            version = ""
            doc_resp = await self.get_srs_doc(id)
            doc = doc_resp.data if doc_resp else None
            if doc:
                version = str(getattr(doc, "product_version", "") or "").strip()
            ws_change.title = __safe_sheet_title(f"{version or '产品'}变更追溯")
            __write_trace_rows(ws_change, change_rows)
        wb.save(output)
        output.seek(0)

    async def export_doc_trace_word(self, output, id: int):
        if Document is None or Pt is None or dox_enum is None:
            return
        from .serv_utils import docx_util

        def __slash(v):
            txt = str(v or "").strip()
            return txt if txt else "/"

        def __is_change_trace_row(obj: dict):
            type_code = str((obj or {}).get("type_code") or "").strip()
            return bool(type_code) and type_code not in ["1", "2"]

        def __format_sds_code(value: str):
            codes = [item.strip() for item in re.split(r"[,，\s]+", str(value or "").strip()) if item.strip()]
            return "\n".join(codes) if len(codes) > 1 else str(value or "").strip()

        def __format_list(values):
            if isinstance(values, list):
                return "\n".join([str(item or "").strip() for item in values if str(item or "").strip()])
            return str(values or "").strip()

        def __format_range(values):
            if isinstance(values, list):
                return " ~ ".join([str(item or "").strip() for item in values if str(item or "").strip()])
            return str(values or "").strip()

        def __build_rows(rows: list[dict]):
            result = []
            hide_chapter_codes = {
                "SDS-RCN300-001",
                "SDS-RCN300-002",
                "SDS-RCN300-003",
                "SDS-RCN300-008",
                "SDS-RCN300-009",
                "SDS-RCN300-010",
            }
            for obj in rows or []:
                sds_code_raw = str(obj.get("sds_code") or "").strip()
                chapter_raw = str(obj.get("chapter") or "").strip()
                if sds_code_raw in hide_chapter_codes or not chapter_raw:
                    sds_code = __format_sds_code(sds_code_raw)
                else:
                    sds_code = f"{__format_sds_code(sds_code_raw)}（{chapter_raw}）"
                sis_codes = obj.get("sis_codes") or []
                test_codes = obj.get("test_codes") or []
                tests_unit = test_codes if test_codes else (obj.get("tests_unit") or [])
                note_raw = str(obj.get("note") or "").strip()
                note = "\n".join([item.strip() for item in note_raw.split("、") if item and item.strip()]) if note_raw else ""
                line_count = max(1, len(sis_codes))
                for idx in range(line_count):
                    is_first_line = idx == 0
                    tests_unit_text = tests_unit[idx] if idx < len(tests_unit) else ""
                    if not test_codes and is_first_line:
                        tests_unit_text = __format_range(tests_unit)
                    result.append({
                        "srs_code": __slash(obj.get("srs_code")) if is_first_line else "",
                        "rcm_flag": (ts("yes") if obj.get("rcm_flag") else ts("no")) if is_first_line else "",
                        "sds_code": __slash(sds_code) if is_first_line else "",
                        "sis_codes": __slash(sis_codes[idx] if idx < len(sis_codes) else ""),
                        "tests_unit": __slash(tests_unit_text),
                        "tests_integ": __slash(" ~ ".join(obj.get("tests_integ") or [])) if is_first_line else "",
                        "tests_sys": __slash(" ~ ".join(obj.get("tests_sys") or [])) if is_first_line else "",
                        "tests_user": __slash(" ~ ".join(obj.get("tests_user") or [])) if is_first_line else "",
                        "rcm_codes": __slash(__format_list(obj.get("rcm_codes") or [])) if is_first_line else "",
                        "note": __slash(note) if is_first_line else "",
                        "_row_span": line_count if idx == 0 else 0,
                    })
            return result

        trace_word_col_widths = [0.78, 0.35, 0.70, 0.70, 1.55, 1.35, 0.98, 0.98, 1.10, 1.15, 0.98]

        def __set_paragraph_text(paragraph, text: str, font_size: float = 10.5, bold: bool = False):
            for run in list(paragraph.runs):
                paragraph._element.remove(run._element)
            docx_util.fonted_txt(paragraph, text, font_size=font_size, bold=bold)

        def __force_paragraph_left_align(paragraph):
            paragraph.alignment = dox_enum.text.WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.right_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            if OxmlElement is None or qn is None:
                return
            p_pr = paragraph._p.get_or_add_pPr()
            for jc in list(p_pr.findall(qn("w:jc"))):
                p_pr.remove(jc)
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), "left")
            p_pr.append(jc)
            for ind in list(p_pr.findall(qn("w:ind"))):
                p_pr.remove(ind)
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), "0")
            ind.set(qn("w:leftChars"), "0")
            ind.set(qn("w:right"), "0")
            ind.set(qn("w:rightChars"), "0")
            ind.set(qn("w:firstLine"), "0")
            ind.set(qn("w:firstLineChars"), "0")
            ind.set(qn("w:hanging"), "0")
            ind.set(qn("w:hangingChars"), "0")
            p_pr.append(ind)

        def __force_cell_left_align(cell):
            try:
                cell.vertical_alignment = dox_enum.table.WD_CELL_VERTICAL_ALIGNMENT.TOP
            except Exception:
                pass
            if OxmlElement is not None and qn is not None:
                tc_pr = cell._tc.get_or_add_tcPr()
                for v_align in list(tc_pr.findall(qn("w:vAlign"))):
                    tc_pr.remove(v_align)
                v_align = OxmlElement("w:vAlign")
                v_align.set(qn("w:val"), "top")
                tc_pr.append(v_align)
                for tc_mar in list(tc_pr.findall(qn("w:tcMar"))):
                    tc_pr.remove(tc_mar)
                tc_mar = OxmlElement("w:tcMar")
                for side in ["top", "left", "bottom", "right"]:
                    item = OxmlElement(f"w:{side}")
                    item.set(qn("w:w"), "0")
                    item.set(qn("w:type"), "dxa")
                    tc_mar.append(item)
                tc_pr.append(tc_mar)
                for no_wrap_node in list(tc_pr.findall(qn("w:noWrap"))):
                    tc_pr.remove(no_wrap_node)
            for paragraph in cell.paragraphs:
                __force_paragraph_left_align(paragraph)

        def __set_cell_text(cell, text: str):
            # 清空模板数据行的原段落，避免继承居中、缩进、前置空格等格式。
            cell._tc.clear_content()
            paragraph = cell.add_paragraph()
            __force_paragraph_left_align(paragraph)
            docx_util.fonted_txt(paragraph, str(text or ""), font_size=10.5)
            __force_cell_left_align(cell)

        def __force_cell_font_size(cell, font_size: float = 10.5):
            if Pt is None:
                return
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)

        def __force_trace_table_font_size(table, font_size: float = 10.5):
            if table is None:
                return
            for row in table.rows:
                for cell in row.cells:
                    __force_cell_font_size(cell, font_size)

        def __set_table_cell_width(cell, width_inch: float):
            if Inches is not None:
                try:
                    cell.width = Inches(width_inch)
                except Exception:
                    pass
            if OxmlElement is None or qn is None:
                return
            tc_pr = cell._tc.get_or_add_tcPr()
            for tc_w in list(tc_pr.findall(qn("w:tcW"))):
                tc_pr.remove(tc_w)
            tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:w"), str(int(width_inch * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            tc_pr.append(tc_w)

        def __fit_trace_table_width(table):
            if table is None:
                return
            try:
                table.autofit = False
                table.allow_autofit = False
            except Exception:
                pass
            if OxmlElement is not None and qn is not None:
                tbl_pr = table._tbl.tblPr
                layout = tbl_pr.find(qn("w:tblLayout"))
                if layout is None:
                    layout = OxmlElement("w:tblLayout")
                    tbl_pr.append(layout)
                layout.set(qn("w:type"), "fixed")
                tbl_grid = table._tbl.tblGrid
                if tbl_grid is not None:
                    for idx, grid_col in enumerate(list(tbl_grid.gridCol_lst)):
                        if idx < len(trace_word_col_widths):
                            grid_col.set(qn("w:w"), str(int(trace_word_col_widths[idx] * 1440)))
            for row in table.rows:
                for idx, cell in enumerate(row.cells):
                    if idx < len(trace_word_col_widths):
                        __set_table_cell_width(cell, trace_word_col_widths[idx])

        def __find_trace_table(docx: Document):
            for table in docx.tables:
                if len(table.rows) >= 2 and len(table.columns) >= 10:
                    first_row = " ".join(cell.text for cell in table.rows[0].cells)
                    if "软件需求规格" in first_row and "软件详细设计" in first_row:
                        return table
            return docx.tables[-1] if docx.tables else None

        def __remove_grid_span(tc):
            if OxmlElement is None or qn is None:
                return 1
            tc_pr = tc.get_or_add_tcPr()
            spans = list(tc_pr.findall(qn("w:gridSpan")))
            span_val = 1
            for span in spans:
                try:
                    span_val = max(span_val, int(span.get(qn("w:val")) or "1"))
                except Exception:
                    span_val = max(span_val, 1)
                tc_pr.remove(span)
            return span_val

        def __split_trace_table_gridspans(table):
            if OxmlElement is None or qn is None:
                return
            for row in table.rows:
                for tc in list(row._tr.tc_lst):
                    span_val = __remove_grid_span(tc)
                    current = tc
                    for _idx in range(max(0, span_val - 1)):
                        cloned = deepcopy(tc)
                        __remove_grid_span(cloned)
                        current.addnext(cloned)
                        current = cloned

        def __fill_template_trace_table(table, rows: list[dict]):
            if table is None:
                return
            if len(table.rows) < 3:
                return
            __split_trace_table_gridspans(table)
            __fit_trace_table_width(table)
            if len(table.rows) >= 2 and len(table.rows[0].cells) >= 5 and len(table.rows[1].cells) >= 5:
                __set_cell_text(table.rows[0].cells[2], "软件详细设计")
                __set_cell_text(table.rows[0].cells[3], "")
                __set_cell_text(table.rows[0].cells[4], "接口编号")
                __set_cell_text(table.rows[1].cells[2], "《软件详细设计》")
                __set_cell_text(table.rows[1].cells[3], "")
                __set_cell_text(table.rows[1].cells[4], "接口编号")
                try:
                    table.rows[0].cells[2].merge(table.rows[0].cells[3])
                    table.rows[1].cells[2].merge(table.rows[1].cells[3])
                except Exception:
                    logger.exception("merge trace word SDS header failed")
            template_tr = deepcopy(table.rows[2]._tr)
            for row in list(table.rows[2:]):
                table._tbl.remove(row._tr)
            built_rows = __build_rows(rows)
            merge_ranges = []
            for item in built_rows:
                if item.get("_row_span"):
                    start = len(table.rows)
                    end = start + int(item.get("_row_span") or 1) - 1
                    if end > start:
                        merge_ranges.append((start, end))
                table._tbl.append(deepcopy(template_tr))
                cells = table.rows[-1].cells
                for idx, cell in enumerate(cells):
                    if idx < len(trace_word_col_widths):
                        __set_table_cell_width(cell, trace_word_col_widths[idx])
                values = [
                    item["srs_code"],
                    item["rcm_flag"],
                    item["sds_code"],
                    "",
                    item["sis_codes"],
                    item["tests_unit"],
                    item["tests_integ"],
                    item["tests_sys"],
                    item["tests_user"],
                    item["rcm_codes"],
                    item["note"],
                ]
                for idx, value in enumerate(values):
                    if idx < len(cells):
                        __set_cell_text(cells[idx], value)
                for cell in cells:
                    __force_cell_left_align(cell)
                try:
                    table.rows[-1].cells[2].merge(table.rows[-1].cells[3])
                    __force_cell_left_align(table.rows[-1].cells[2])
                except Exception:
                    logger.exception("merge trace word SDS row failed")
            for start, end in merge_ranges:
                # “接口编号”和“单元测试用例”按编号逐行展示，其它列按同一需求纵向合并。
                for col_idx in [0, 1, 2, 3, 6, 7, 8, 9, 10]:
                    try:
                        table.cell(start, col_idx).merge(table.cell(end, col_idx))
                        __force_cell_left_align(table.cell(start, col_idx))
                    except Exception:
                        logger.exception("merge trace word cell failed: row=%s-%s col=%s", start, end, col_idx)
            __force_trace_table_font_size(table, 10.5)

        def __append_template_trace_table(docx: Document, source_table, title: str, rows: list[dict]):
            paragraph = docx.add_paragraph()
            paragraph.style = docx.styles["Heading 1"]
            __set_paragraph_text(paragraph, title, font_size=14.0, bold=True)
            docx._body._element.append(deepcopy(source_table._tbl))
            new_table = docx.tables[-1]
            __fill_template_trace_table(new_table, rows)

        def __update_header_file_no(docx: Document, file_no: str):
            text = str(file_no or "").strip() or "/"
            for section in docx.sections:
                for header in [section.header, section.first_page_header, section.even_page_header]:
                    for paragraph in header.paragraphs:
                        raw = str(paragraph.text or "").strip()
                        if not raw:
                            continue
                        if raw == text or raw.startswith("TX-") or raw.startswith("QMS-"):
                            __set_paragraph_text(paragraph, text, font_size=10.5)

        resp = await self.list_doc_trace(id)
        row_srsdoc: SrsDoc = db.session.execute(select(SrsDoc).where(SrsDoc.id == id)).scalars().first()
        row_product: Product = None
        if row_srsdoc:
            row_product = db.session.execute(select(Product).where(Product.id == row_srsdoc.product_id)).scalars().first()
        product_name = str(getattr(row_product, "name", "") or "").strip()
        product_type_code = str(getattr(row_product, "type_code", "") or "").strip()
        product_full_version = str(getattr(row_product, "full_version", "") or "").strip()

        rows = resp.data or []
        normal_rows = [obj for obj in rows if not __is_change_trace_row(obj)]
        change_rows = [obj for obj in rows if __is_change_trace_row(obj)]

        temp_path = os.path.join(os.path.dirname(__file__), "temp_doc_trace_word.docx")
        docx = Document(temp_path) if os.path.exists(temp_path) else Document()
        __update_header_file_no(docx, getattr(row_srsdoc, "file_no", "") if row_srsdoc else "")
        for paragraph in docx.paragraphs:
            raw = str(paragraph.text or "").strip()
            if raw.startswith("产品名称："):
                __set_paragraph_text(paragraph, f"产品名称：{product_name}", font_size=10.5)
            elif raw.startswith("产品型号："):
                __set_paragraph_text(paragraph, f"产品型号：{product_type_code or '/'}", font_size=10.5)
            elif raw.startswith("产品版本："):
                __set_paragraph_text(paragraph, f"产品版本：{product_full_version or '/'}", font_size=10.5)
            elif raw.startswith("版本号："):
                __set_paragraph_text(paragraph, f"版本号：{product_full_version or '/'}", font_size=10.5)
        trace_table = __find_trace_table(docx)
        __fill_template_trace_table(trace_table, normal_rows)
        if change_rows:
            __append_template_trace_table(docx, trace_table, f"{product_full_version or '产品'}变更追溯", change_rows)
        docx.save(output)
        output.seek(0)
        
