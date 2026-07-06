#!/usr/bin/env python
# encoding: utf-8

# 自研软件网络安全研究报告服务层。
# 默认内容来自 src-res/nsr_default_content.json（章节树 + 表格 + 内置图，由原 Word 模板提取）。
# 自动获取：1.1 软件信息（产品名称/型号/发布版本/完整版本，安全级别保留模板默认）、全文产品名统一替换；
#          内置图按「正文 → 图 → 图题 → 正文」版式还原。其余章节模板化（保留模板原文）。

import copy
import json
import logging
import os
import re
from typing import List

from sqlalchemy import func, select, delete
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from ..model.product import Product
from ..model.nsr_doc import NsrDoc
from ..model.srs_doc import SrsDoc
from ..model.prod_dhf import ProdDhf
from ..obj import Page, Resp
from ..obj.tobj_role import Roles
from ..obj.vobj_user import UserObj
from ..obj.tobj_nsr_doc import NsrDocForm
from ..obj.vobj_nsr_doc import NsrDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from .serv_utils import new_version, sync_file_no_version, docx_util
from .serv_prod_haz import Server as ProdHazServer
from .serv_srs_doc import Server as SrsDocServer

logger = logging.getLogger(__name__)

# 文档大标题（封面 + 导出文件名）
DOC_TITLE = "自研软件网络安全研究报告"

# 模板占位产品名（含可能的空格变体），自动获取时全文统一替换为实际产品名
_TEMPLATE_PRODUCT_NAME_RE = re.compile(r"肿瘤\s*CT\s*图像随访与评估软件")
_TEMPLATE_PRODUCT_NAME = "肿瘤CT图像随访与评估软件"

DEFAULT_NSR_CONTENT = {"productName": "", "sections": []}

# 2.1 风险分布矩阵固定结构（严格按原 Word）：发生概率(经常5→非常少1) × 严重度(可忽略A→灾难性E)
_RISK_PROB_ROWS = [("经常", "5"), ("有时", "4"), ("偶然", "3"), ("很少", "2"), ("非常少", "1")]
_RISK_SEV_COLS = [("可忽略", "A"), ("轻度", "B"), ("严重", "C"), ("危重的", "D"), ("灾难性的", "E")]
# 风险等级配色矩阵（行=概率，列=严重度）：R=不可接受 O=进一步降低的研究 G=可忽略
_RISK_COLOR = [
    ["R", "R", "R", "R", "R"],
    ["R", "R", "R", "R", "R"],
    ["O", "O", "R", "R", "R"],
    ["G", "G", "O", "R", "R"],
    ["G", "G", "O", "O", "R"],
]
_RISK_HEX = {"R": "FF0000", "O": "FFC000", "G": "00B050"}
_RISK_LEGEND = [
    ("R", "红色", "不可接受：这类风险本质上不可接受。必须寻求风险降低措施，"),
    ("O", "橙色", "进一步降低的研究：这类风险必须降低到合理可行的最低限度才可视为可接受"),
    ("G", "绿色", "可忽略：这些风险本质上是可以接受的。即使它可以忽略不计，推想也需要尽可能地降低风险"),
]

_DEFAULT_CONTENT_FILE = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "src-res", "nsr_default_content.json"
)
try:
    with open(_DEFAULT_CONTENT_FILE, encoding="utf-8") as _f:
        _loaded = json.load(_f)
    if isinstance(_loaded, dict) and isinstance(_loaded.get("sections"), list) and _loaded["sections"]:
        DEFAULT_NSR_CONTENT = _loaded
except Exception:
    logger.exception("加载自研软件网络安全研究报告默认内容资源失败")

# 模板表名（标题 -> 各表表名），用于给旧文档补全表上方显示的表名（未人工编辑时）
_DEFAULT_TABLE_TITLES = {}


def _build_title_map(nodes):
    for n in nodes or []:
        if isinstance(n, dict):
            if n.get("table_titles") and n.get("title"):
                _DEFAULT_TABLE_TITLES[n["title"]] = n["table_titles"]
            _build_title_map(n.get("children"))


_build_title_map(DEFAULT_NSR_CONTENT.get("sections"))


class Server(object):
    # ---------------- 内容归一 ----------------
    def __fill_template_titles(self, nodes):
        for n in nodes or []:
            if isinstance(n, dict):
                if n.get("tables") and not n.get("table_titles"):
                    titles = _DEFAULT_TABLE_TITLES.get(n.get("title"))
                    if titles:
                        n["table_titles"] = list(titles)
                        text = n.get("text")
                        if isinstance(text, str) and text:
                            title_set = {re.sub(r"\s+", "", t) for t in titles if t}
                            n["text"] = "\n".join(
                                ln for ln in text.split("\n")
                                if re.sub(r"\s+", "", ln.strip()) not in title_set
                            )
                self.__fill_template_titles(n.get("children"))

    def __normalize_content(self, content):
        result = copy.deepcopy(DEFAULT_NSR_CONTENT)
        if isinstance(content, dict):
            result.update(content)
        result.setdefault("sections", copy.deepcopy(DEFAULT_NSR_CONTENT.get("sections", [])))
        result.setdefault("productName", "")
        self.__strip_blocks(result.get("sections"))
        self.__fill_template_titles(result.get("sections"))
        return result

    def __strip_blocks(self, nodes):
        # blocks 为读取时由 autofill 派生的展示结构，不持久化；保留原始 text/images/tables 以便重算
        for n in nodes or []:
            if isinstance(n, dict):
                n.pop("blocks", None)
                self.__strip_blocks(n.get("children"))

    def __strip_name(self, title):
        return re.sub(r"^[0-9．.、\s]+", "", str(title or "")).strip()

    # ---------------- 自动获取数据源 ----------------
    async def __cyber_haz_rows(self, product_id):
        # 附录A 网络安全风险分析列表：取产品HAZ管理中「分类=网络安全」的记录，
        # 复用 ProdHaz 服务（已合并 haz 基础信息 + prod_rcm 措施/证据展开）。
        resp = await ProdHazServer().list_prod_haz(None, export=True, prod_id=product_id, page_index=0, page_size=100000)
        objs = getattr(getattr(resp, "data", None), "rows", None) or []

        def s(v):
            return "" if v is None else str(v)

        def yn(v):
            if v in (None, ""):
                return ""
            return "Y" if s(v).strip() in ("1", "Y", "y", "是", "True", "true") else "N"

        rows = []
        for o in objs:
            cat = s(getattr(o, "category", ""))
            if "网络安全" not in cat:
                continue
            rows.append([
                s(o.code), s(o.source), s(o.event), s(o.situation), s(o.damage),
                s(o.init_rate), s(o.init_degree), s(o.init_level),
                s(o.deal), "\n".join(self.__split_rcm_codes(getattr(o, "rcms", ""))), s(o.evidence),
                s(o.cur_rate), s(o.cur_degree), s(o.cur_level),
                yn(getattr(o, "benefit_flag", None)), cat,
            ])
        return rows

    _RCM_CODE_RE = re.compile(r"\bRCM[\s\-_]*[A-Z0-9]+(?:[\-_][A-Z0-9]+)*\b", re.I)

    @classmethod
    def __split_rcm_codes(cls, value):
        return [re.sub(r"\s+", "", m).upper().replace("_", "-")
                for m in dict.fromkeys(cls._RCM_CODE_RE.findall(str(value or "")))]

    @staticmethod
    def __norm_code(value):
        return re.sub(r"\s+", "", str(value or "")).upper().replace("_", "-")

    @staticmethod
    def __test_ranges(codes):
        # 同前缀（去末尾数字）的连续用例合并为「首至尾」，多组按编号顺序换行罗列
        uniq = list(dict.fromkeys([str(c).strip() for c in (codes or []) if str(c).strip()]))

        def sort_key(c):
            m = re.match(r"^(.*?)(\d+)$", c)
            return (m.group(1), int(m.group(2))) if m else (c, -1)

        uniq.sort(key=sort_key)
        groups = {}
        order = []
        for c in uniq:
            m = re.match(r"^(.*?)(\d+)$", c)
            if not m:
                order.append(("raw", c))
                continue
            prefix = m.group(1)
            if prefix not in groups:
                groups[prefix] = []
                order.append(("grp", prefix))
            groups[prefix].append(c)
        out = []
        for typ, val in order:
            if typ == "raw":
                out.append(val)
                continue
            g = groups[val]
            out.append(g[0] if len(g) == 1 else f"{g[0]}至{g[-1]}")
        return "\n".join(out)

    async def __cyber_trace_rows(self, product_id):
        # 附录C 网络安全可追溯性分析表：结合附录A 的网络安全 HAZ → RCM，
        # 再用这些 RCM 过滤「可追溯性分析记录」(list_doc_trace)，命中的需求按 RCM 逐行展开，
        # SDS / 单元 / 集成 / 系统 / 用户测试用例直接取自该追溯记录（与可追溯性分析页面一致）。
        # 1) 网络安全 HAZ → RCM 集合（建立 RCM → 关联网络安全 HAZ 编号 的映射）
        haz_resp = await ProdHazServer().list_prod_haz(None, export=True, prod_id=product_id, page_index=0, page_size=100000)
        haz_objs = getattr(getattr(haz_resp, "data", None), "rows", None) or []
        rcm_to_hazs = {}
        for o in haz_objs:
            if "网络安全" not in str(getattr(o, "category", "") or ""):
                continue
            haz_code = str(getattr(o, "code", "") or "").strip()
            for rcm in self.__split_rcm_codes(getattr(o, "rcms", "")):
                hazs = rcm_to_hazs.setdefault(rcm, [])
                if haz_code and haz_code not in hazs:
                    hazs.append(haz_code)
        if not rcm_to_hazs:
            return []

        # 2) 可追溯性分析记录：取产品最新 SRS 文档的逐需求追溯（含 SDS 与各级测试用例）
        srs_doc = db.session.execute(
            select(SrsDoc).where(SrsDoc.product_id == product_id).order_by(SrsDoc.id.desc())
        ).scalars().first()
        if not srs_doc:
            return []
        trace_resp = await SrsDocServer().list_doc_trace(srs_doc.id)
        trace_rows = getattr(trace_resp, "data", None) or []

        # 3) 逐行展开：仅保留 RCM 命中网络安全的需求，严格取追溯记录实际值，无则「/」占位
        rows = []
        for tr in trace_rows:
            row_rcms = [self.__norm_code(c) for c in (tr.get("rcm_codes") or [])]
            cyber_rcms = [c for c in dict.fromkeys(row_rcms) if c in rcm_to_hazs]
            if not cyber_rcms:
                continue
            srs = str(tr.get("srs_code") or "").strip()
            sds = str(tr.get("sds_code") or "").strip() or "/"

            def col(existing):
                vals = [c for c in (existing or []) if str(c or "").strip()]
                return self.__test_ranges(vals) or "/"

            tu = col(tr.get("test_codes") or tr.get("tests_unit"))
            ti = col(tr.get("tests_integ"))
            tsv = col(tr.get("tests_sys"))
            ty = col(tr.get("tests_user"))
            # 同一 SRS 聚合为一行：危害编号、RCMID 各自把该需求下的多个值去重后按编号排序纵向堆叠显示
            def num(code):
                m = re.search(r"(\d+)", str(code or ""))
                return int(m.group(1)) if m else 0

            haz_all = []
            for rcm in cyber_rcms:
                for h in rcm_to_hazs.get(rcm, []):
                    if h not in haz_all:
                        haz_all.append(h)
            haz_all.sort(key=num)
            rows.append([
                srs, "是", "\n".join(haz_all), sds,
                tu, ti, tsv, ty, "\n".join(sorted(cyber_rcms, key=num)), "/",
            ])
        rows.sort(key=lambda r: r[0])
        return rows

    @staticmethod
    def __risk_matrix(objs, rate_attr, degree_attr):
        # 风险分布矩阵：行=发生概率(经常5→非常少1)，列=严重度(可忽略A→灾难性E)，格子=落入该组合的网络安全HAZ数量
        rate_to_row = {"5": 0, "4": 1, "3": 2, "2": 3, "1": 4}
        sev_to_col = {"A": 0, "B": 1, "C": 2, "D": 3, "E": 4}
        counts = [[0] * 5 for _ in range(5)]
        for o in objs:
            r = str(getattr(o, rate_attr, "") or "").strip()
            dg = str(getattr(o, degree_attr, "") or "").strip().upper()
            if r in rate_to_row and dg in sev_to_col:
                counts[rate_to_row[r]][sev_to_col[dg]] += 1
        row_totals = [sum(row) for row in counts]
        col_totals = [sum(counts[ri][ci] for ri in range(5)) for ci in range(5)]
        return {"counts": counts, "row_totals": row_totals,
                "col_totals": col_totals, "total": sum(row_totals)}

    async def __cyber_risk_matrix(self, product_id):
        # 初始风险矩阵(init_rate/init_degree) 与 采取措施后风险矩阵(cur_rate/cur_degree)
        resp = await ProdHazServer().list_prod_haz(None, export=True, prod_id=product_id, page_index=0, page_size=100000)
        objs = [o for o in (getattr(getattr(resp, "data", None), "rows", None) or [])
                if "网络安全" in str(getattr(o, "category", "") or "")]
        return (self.__risk_matrix(objs, "init_rate", "init_degree"),
                self.__risk_matrix(objs, "cur_rate", "cur_degree"))

    async def __collect_autofill(self, product_id):
        product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
        if not product:
            return None
        risk_init, risk_cur = await self.__cyber_risk_matrix(product_id)
        return {
            "product_name": (product.name or "").strip(),
            "type_code": (product.type_code or "").strip(),
            "release_version": (product.release_version or "").strip(),
            "full_version": (product.full_version or "").strip(),
            "cyber_haz_rows": await self.__cyber_haz_rows(product_id),
            "cyber_trace_rows": await self.__cyber_trace_rows(product_id),
            "risk_init": risk_init,
            "risk_cur": risk_cur,
        }

    _TABLE_PH_RE = re.compile(r"^\{\{TABLE:(\d+)\}\}$")
    _RISK_PH_RE = re.compile(r"^\{\{RISK:(\d+)\}\}$")

    @classmethod
    def __content_blocks(cls, text, image_urls, tables, table_titles, no_merge=False, risk_matrices=None):
        # 按原 Word 顺序还原章节版式：正文 → 表 → 正文 → 图 → 正文 …
        #   1) 正文含占位行「{{IMG}}」「{{TABLE:n}}」时，图/表按占位位置精确插入；
        #   2) 否则按「图N …」图题行定位（兼容旧数据）；
        #   3) 都没有时：正文 → 图 → 表 顺序兜底。
        urls = [u for u in (image_urls or []) if u]
        tables = tables or []
        titles = table_titles or []
        mats = risk_matrices or []
        lines = str(text or "").split("\n")
        blocks = []
        buf = []

        def flush():
            t = "\n".join(buf).strip()
            if t:
                blocks.append({"text": t})
            buf.clear()

        def table_block(idx):
            if 0 <= idx < len(tables):
                b = {"type": "table", "table_index": idx, "table": tables[idx],
                     "title": titles[idx] if idx < len(titles) else ""}
                if no_merge:
                    b["no_merge"] = True
                return b
            return None

        has_ph = any(ln.strip() == "{{IMG}}" or cls._TABLE_PH_RE.match(ln.strip())
                     or cls._RISK_PH_RE.match(ln.strip()) for ln in lines)
        if has_ph:
            ui = 0
            for ln in lines:
                s = ln.strip()
                if s == "{{IMG}}":
                    flush()
                    if ui < len(urls):
                        blocks.append({"type": "image", "url": urls[ui]})
                        ui += 1
                elif cls._RISK_PH_RE.match(s):
                    flush()
                    mi = int(cls._RISK_PH_RE.match(s).group(1))
                    if 0 <= mi < len(mats) and mats[mi]:
                        blocks.append({"type": "risk_matrix", "matrix": mats[mi]})
                elif cls._TABLE_PH_RE.match(s):
                    flush()
                    tb = table_block(int(cls._TABLE_PH_RE.match(s).group(1)))
                    if tb:
                        blocks.append(tb)
                elif re.match(r"^图\s*\d", s):
                    flush()
                    blocks.append({"type": "caption", "text": s})
                else:
                    buf.append(ln)
            flush()
            for u in urls[ui:]:
                blocks.append({"type": "image", "url": u})
            return blocks

        # 兜底（旧档无占位符）：按"图N"题注行定位，将每张图插到其题注前，
        # 保持"正文→图→图题→正文→图→图题→正文"的原始排版，支持多图。
        has_cap = any(re.match(r"^图\s*\d", ln.strip()) for ln in lines)
        if has_cap:
            ui = 0
            for ln in lines:
                s = ln.strip()
                if re.match(r"^图\s*\d", s):
                    flush()
                    if ui < len(urls):
                        blocks.append({"type": "image", "url": urls[ui]})
                        ui += 1
                    blocks.append({"type": "caption", "text": s})
                else:
                    buf.append(ln)
            flush()
            for u in urls[ui:]:
                blocks.append({"type": "image", "url": u})
        else:
            if str(text or "").strip():
                blocks.append({"text": str(text)})
            blocks.extend([{"type": "image", "url": u} for u in urls])
        for i in range(len(tables)):
            blocks.append(table_block(i))
        return blocks

    def __sw_info_text(self, node_text, auto):
        # 1.1 软件信息：逐行替换「产品名称/产品型号/软件发布版本/软件完整版本」，
        # 安全级别等其余行保留模板默认值。
        mapping = [
            ("产品名称", auto.get("product_name", "")),
            ("产品型号", auto.get("type_code", "")),
            ("软件发布版本", auto.get("release_version", "")),
            ("软件完整版本", auto.get("full_version", "")),
        ]
        out = []
        for ln in str(node_text or "").split("\n"):
            key = ln.replace("\t", "").strip()
            replaced = False
            for label, value in mapping:
                if key.startswith(label) and value:
                    out.append(f"{label}：{value}")
                    replaced = True
                    break
            if not replaced:
                out.append(ln)
        return "\n".join(out)

    def __apply_autofill(self, content, auto):
        if not auto:
            return content

        def walk(nodes):
            for node in nodes or []:
                rt = node.get("ref_type")
                title = str(node.get("title") or "")
                # 兼容旧文档（标记前已保存，持久化缺 ref_type）：按标题兜底识别自动获取章节
                if not rt:
                    if self.__strip_name(title) == "软件信息":
                        rt = "sw_info"
                    elif title.startswith("附录A"):
                        rt = "cyber_haz"
                    elif title.startswith("附录C") or "可追溯性分析表" in title:
                        rt = "cyber_trace"
                    elif self.__strip_name(title) == "风险管理":
                        rt = "risk_matrix"
                    if rt:
                        node["ref_type"] = rt
                if rt == "sw_info":
                    node["text"] = self.__sw_info_text(node.get("text"), auto)
                elif rt == "cyber_haz":
                    # 附录A：保留模板表头，数据行替换为产品HAZ管理中分类=网络安全的记录
                    tbls = node.get("tables") or []
                    if tbls and tbls[0]:
                        header = tbls[0][0]
                        node["tables"][0] = [header] + (auto.get("cyber_haz_rows") or [])
                elif rt == "cyber_trace":
                    # 附录C：保留模板表头（前3行：列名 + 两行来源说明），数据行按网络安全追溯动态生成
                    tbls = node.get("tables") or []
                    if tbls and tbls[0]:
                        header = tbls[0][:3]
                        node["tables"][0] = header + (auto.get("cyber_trace_rows") or [])
                elif rt == "risk_matrix":
                    # 2.1 风险管理：原 Word 两张「风险分布矩阵」图按原结构/配色还原为专用矩阵（{{RISK:0}}初始 / {{RISK:1}}措施后）
                    base = node.get("tables") or []
                    node["tables"] = [base[0]] if base else []  # 仅保留术语解释表
                    node["images"] = []
                    node["risk_matrices"] = [auto.get("risk_init"), auto.get("risk_cur")]
                    # 规范占位：剔除旧的 {{IMG}}/{{TABLE:n}}/{{RISK:n}}，按引导句重建术语表与两张矩阵占位（兼容无占位的旧文档）
                    out, has_t, has_r0, has_r1 = [], False, False, False
                    for ln in str(node.get("text") or "").split("\n"):
                        s = ln.strip()
                        if s in ("{{IMG}}", "{{TABLE:0}}", "{{TABLE:1}}", "{{TABLE:2}}", "{{RISK:0}}", "{{RISK:1}}"):
                            continue
                        out.append(ln)
                        if not has_t and "术语解释" in s:
                            out.append("{{TABLE:0}}"); has_t = True
                        elif not has_r0 and "初始风险分布矩阵" in s:
                            out.append("{{RISK:0}}"); has_r0 = True
                        elif not has_r1 and "采取风险措施后" in s:
                            out.append("{{RISK:1}}"); has_r1 = True
                    node["text"] = "\n".join(out)
                # 正文/表/图按原 Word 顺序还原为有序 blocks（cover/revision 走导出专用逻辑，不重排）
                if rt not in ("cover", "revision"):
                    txt = str(node.get("text") or "")
                    has_ph = "{{IMG}}" in txt or re.search(r"\{\{(TABLE|RISK):\d+\}\}", txt)
                    if has_ph or node.get("images") or node.get("tables"):
                        # 附录A(逐行HAZ)、附录C(逐行追溯)为独立行数据，禁止重复项纵向合并
                        no_merge = (rt in ("cyber_haz", "risk_matrix") or title.startswith("附录C")
                                    or "可追溯性分析表" in title or self.__strip_name(title) == "网络安全能力")
                        node["blocks"] = self.__content_blocks(
                            node.get("text"), node.get("images"), node.get("tables"), node.get("table_titles"),
                            no_merge=no_merge, risk_matrices=node.get("risk_matrices"),
                        )
                walk(node.get("children"))

        walk(content.get("sections"))
        product_name = (auto.get("product_name") or "").strip()
        if product_name and re.sub(r"\s", "", product_name) != _TEMPLATE_PRODUCT_NAME:
            self.__sync_product_name(content, product_name)
        return content

    def __sync_product_name(self, content, product_name):
        # 全文统一替换：把模板占位产品名（含空格变体）替换为实际产品名
        def rep(s):
            return _TEMPLATE_PRODUCT_NAME_RE.sub(product_name, s) if isinstance(s, str) else s

        def walk(nodes):
            for n in nodes or []:
                for k in ("text", "body"):
                    if isinstance(n.get(k), str):
                        n[k] = rep(n[k])
                for tbl in n.get("tables") or []:
                    for row in tbl:
                        for i, c in enumerate(row):
                            if isinstance(c, str):
                                row[i] = rep(c)
                for b in n.get("blocks") or []:
                    for k in ("text", "title"):
                        if isinstance(b.get(k), str):
                            b[k] = rep(b[k])
                    for row in b.get("table") or []:
                        for i, c in enumerate(row):
                            if isinstance(c, str):
                                row[i] = rep(c)
                walk(n.get("children"))

        walk(content.get("sections"))

    # ---------------- 转换 ----------------
    # 从产品 DHF 按文档名匹配文件编号（文件编号未手填时自动获取）
    def __dhf_file_no(self, prod_id):
        row = db.session.execute(
            select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name.like("%网络安全研究报告%"))
        ).scalars().first()
        return (row.code or "").strip() if row and row.code else ""

    async def __to_obj(self, row: NsrDoc, product: Product = None, with_autofill=True):
        obj = NsrDocObj(**row.dict())
        content = self.__normalize_content(obj.content)
        if with_autofill:
            auto = await self.__collect_autofill(row.product_id)
            content = self.__apply_autofill(content, auto)
        if product:
            content["productName"] = product.name or ""
        obj.content = content
        if product:
            obj.product_name = product.name
            obj.product_version = product.full_version
            obj.product_full_version = product.full_version
            obj.product_type_code = product.type_code
            if not (obj.file_no or "").strip():
                dhf_no = self.__dhf_file_no(product.id)
                if dhf_no:
                    obj.file_no = dhf_no
        return obj

    # ---------------- CRUD ----------------
    async def add_nsr_doc(self, form: NsrDocForm):
        try:
            sql = select(func.count(NsrDoc.id)).where(
                NsrDoc.product_id == form.product_id, NsrDoc.version == form.version
            )
            if db.session.execute(sql).scalar() > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            row = NsrDoc(**form.dict(exclude_none=True))
            row.id = None
            row.content = self.__normalize_content(row.content)
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=NsrDocForm(id=row.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def duplicate_nsr_doc(self, id: int, product_id: int = None):
        try:
            fromdoc: NsrDoc = db.session.execute(select(NsrDoc).where(NsrDoc.id == id)).scalars().first()
            if not fromdoc:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            target_pid = product_id or fromdoc.product_id
            all_versions = db.session.execute(select(NsrDoc.version).where(NsrDoc.product_id == target_pid)).scalars().all()
            existing = {v for v in all_versions if v}
            if target_pid == fromdoc.product_id:
                version = new_version(fromdoc.version)
            else:
                def _seq(v):
                    m = re.search(r"(\d+)(?!.*\d)", v or "")
                    return int(m.group(1)) if m else -1
                valid = [v for v in all_versions if v]
                version = new_version(max(valid, key=_seq)) if valid else fromdoc.version
            while version in existing:
                version = new_version(version)
            newdoc = NsrDoc(
                product_id=target_pid, version=version, file_no=sync_file_no_version((fromdoc.file_no or "").strip() or self.__dhf_file_no(target_pid), version) or None,
                change_log=fromdoc.change_log, content=copy.deepcopy(self.__normalize_content(fromdoc.content)),
            )
            db.session.add(newdoc)
            db.session.commit()
            return Resp.resp_ok(data=NsrDocForm(id=newdoc.id))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def update_nsr_doc(self, form: NsrDocForm):
        try:
            row: NsrDoc = db.session.execute(select(NsrDoc).where(NsrDoc.id == form.id)).scalars().first()
            if not row:
                return Resp.resp_err(msg=ts("msg_obj_null"))
            for key, value in form.dict(exclude_none=True).items():
                if key == "id":
                    continue
                if key == "content":
                    value = self.__normalize_content(value)
                setattr(row, key, value)
            db.session.commit()
            return Resp.resp_ok()
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts(msg_err_db))

    async def delete_nsr_doc(self, id: int):
        db.session.execute(delete(NsrDoc).where(NsrDoc.id == id))
        db.session.commit()
        return Resp.resp_ok()

    async def get_nsr_doc(self, id: int):
        sql = select(NsrDoc, Product).join(Product, NsrDoc.product_id == Product.id).where(NsrDoc.id == id)
        row = db.session.execute(sql).first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        doc, product = row
        return Resp.resp_ok(data=await self.__to_obj(doc, product))

    async def list_nsr_doc(self, op_user: UserObj = None, product_id: int = 0, version: str = None, page_index: int = 0, page_size: int = 10):
        wheres = []
        if product_id:
            wheres.append(NsrDoc.product_id == product_id)
        if version:
            wheres.append(NsrDoc.version.like(f"%{version}%"))
        if op_user and op_user.id != 1 and op_user.role_code == Roles.product_manager.value.code:
            wheres.append(Product.create_user_id == op_user.id)
        total = db.session.execute(
            select(func.count(NsrDoc.id)).join(Product, NsrDoc.product_id == Product.id).where(*wheres)
        ).scalar() or 0
        sql = (
            select(NsrDoc, Product).join(Product, NsrDoc.product_id == Product.id)
            .where(*wheres).order_by(NsrDoc.id.desc()).offset(page_index * page_size).limit(page_size)
        )
        rows: List[NsrDocObj] = [await self.__to_obj(doc, product, with_autofill=False) for doc, product in db.session.execute(sql).all()]
        return Resp.resp_ok(data=Page(total=total, rows=rows, page_index=page_index, page_size=page_size))

    async def nsr_autofill(self, product_id: int):
        # 新增页 / 切换产品预览：返回应用了自动获取的默认内容
        content = self.__normalize_content(None)
        auto = await self.__collect_autofill(product_id)
        content = self.__apply_autofill(content, auto)
        product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
        if product:
            content["productName"] = product.name or ""
        return Resp.resp_ok(data=content)

    # ---------------- 导出 ----------------
    async def export_nsr_doc(self, output, id: int):
        resp = await self.get_nsr_doc(id)
        obj: NsrDocObj = resp.data
        if obj is None:
            Document().save(output)
            output.seek(0)
            return
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        document.settings.element.append(update_fields)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, obj.file_no or "")
        docx_util.add_page_number_footer(section, obj.file_no or "")

        def normalized(value):
            return re.sub(r"\s+", "", str(value or ""))

        def is_cover(sec):
            return sec.get("ref_type") == "cover" or normalized(sec.get("title")) == normalized(DOC_TITLE)

        def is_revision(sec):
            return sec.get("ref_type") == "revision" or normalized(sec.get("title")) == "文件修订记录"

        def set_cell_text(cell, text, bold=False):
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.line_spacing = 1.5
            docx_util.fonted_txt(paragraph, str(text or ""), font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        def add_table_title(title):
            title = str(title or "").strip()
            if not title:
                return
            para = document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            docx_util.fonted_txt(para, title, font_size=10.5, bold=True)

        def add_plain_table(rows, title="", no_merge=False):
            rows = rows or []
            col_count = max([len(row or []) for row in rows] or [0])
            if col_count <= 0:
                return
            add_table_title(title)
            table = document.add_table(rows=0, cols=col_count)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for row in rows:
                cells = table.add_row().cells
                for idx in range(col_count):
                    set_cell_text(cells[idx], row[idx] if idx < len(row or []) else "", bold=(len(table.rows) == 1))
            if no_merge:
                document.add_paragraph()
                return
            # 重复项合并：每列纵向连续相同（非空）单元格合并并居中（表头行不合并）
            def cval(ri, ci):
                row = rows[ri] if ri < len(rows) else None
                return str(row[ci]) if row and ci < len(row or []) else ""

            for ci in range(col_count):
                ri = 1
                while ri < len(rows):
                    if not cval(ri, ci).strip():
                        ri += 1
                        continue
                    rj = ri + 1
                    while rj < len(rows) and cval(rj, ci) == cval(ri, ci):
                        rj += 1
                    if rj - ri > 1:
                        val = cval(ri, ci)
                        merged = table.cell(ri, ci).merge(table.cell(rj - 1, ci))
                        merged.text = ""
                        para = merged.paragraphs[0]
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        para.paragraph_format.line_spacing = 1.5
                        docx_util.fonted_txt(para, val, font_size=10.5)
                        merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    ri = rj
            document.add_paragraph()

        def shade_cell(cell, hexcolor):
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), hexcolor)
            tcPr.append(shd)

        def add_risk_matrix(data, title=""):
            # 2.1 风险分布矩阵：发生概率(经常5→非常少1) × 严重度(可忽略A→灾难性E)，带固定风险等级配色 + 行列总计 + 图例
            data = data or {}
            counts = data.get("counts") or [[0] * 5 for _ in range(5)]
            row_totals = data.get("row_totals") or [0] * 5
            col_totals = data.get("col_totals") or [0] * 5
            total = data.get("total") or 0
            add_table_title(title)
            table = document.add_table(rows=9 + len(_RISK_LEGEND), cols=9)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            def put(cell, text, bold=False, fill=None):
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 1.0
                docx_util.fonted_txt(p, "" if text is None else str(text), font_size=10.5, bold=bold)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if fill:
                    shade_cell(cell, fill)

            def merge_put(r1, c1, r2, c2, text, bold=False):
                put(table.cell(r1, c1).merge(table.cell(r2, c2)), text, bold=bold)

            merge_put(0, 0, 2, 2, "风险值", bold=True)
            merge_put(0, 3, 0, 7, "严重度", bold=True)
            merge_put(0, 8, 2, 8, "总计", bold=True)
            for ci, (sname, letter) in enumerate(_RISK_SEV_COLS):
                put(table.cell(1, 3 + ci), sname, bold=True)
                put(table.cell(2, 3 + ci), letter, bold=True)
            merge_put(3, 0, 7, 0, "发生概率", bold=True)
            for ri, (pname, pnum) in enumerate(_RISK_PROB_ROWS):
                r = 3 + ri
                put(table.cell(r, 1), pname, bold=True)
                put(table.cell(r, 2), pnum, bold=True)
                for ci in range(5):
                    put(table.cell(r, 3 + ci), counts[ri][ci], fill=_RISK_HEX.get(_RISK_COLOR[ri][ci]))
                put(table.cell(r, 8), row_totals[ri], bold=True)
            merge_put(8, 0, 8, 2, "总计", bold=True)
            for ci in range(5):
                put(table.cell(8, 3 + ci), col_totals[ci], bold=True)
            put(table.cell(8, 8), total, bold=True)
            # 图例并入同一张表（第0列色块标签，其余列合并为说明）
            for i, (key, label, desc) in enumerate(_RISK_LEGEND):
                r = 9 + i
                put(table.cell(r, 0), label, bold=True, fill=_RISK_HEX[key])
                dcell = table.cell(r, 1).merge(table.cell(r, 8))
                dcell.text = ""
                p1 = dcell.paragraphs[0]
                p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
                docx_util.fonted_txt(p1, desc, font_size=10.5)
                dcell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            document.add_paragraph()

        def add_image(url):
            raw = str(url or "").strip()
            if not raw:
                return
            try:
                # 流程图等内置图：限制最大边，避免单图过大占满版面
                docx_util.save_img2docx(raw, document, mw=340, mh=420)
            except Exception:
                logger.exception("导出网络安全研究报告图片失败")

        def add_section(sec, level=1):
            title = sec.get("title", "")
            if title and not is_cover(sec):
                docx_util.save_title2docx(title, document, level=max(1, min(level, 9)))
            blocks = sec.get("blocks")
            if isinstance(blocks, list) and blocks:
                for block in blocks:
                    if not isinstance(block, dict):
                        continue
                    if block.get("type") == "image":
                        add_image(block.get("url"))
                    elif block.get("type") == "caption":
                        cap = document.add_paragraph()
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        docx_util.fonted_txt(cap, str(block.get("text") or ""), font_size=10.5)
                    elif block.get("type") == "table":
                        add_plain_table(block.get("table") or [], block.get("title") or "", block.get("no_merge"))
                    elif block.get("type") == "risk_matrix":
                        add_risk_matrix(block.get("matrix") or {}, block.get("title") or "")
                    elif block.get("text"):
                        docx_util.save_txt2docx(str(block.get("text") or ""), document)
            else:
                if sec.get("text"):
                    docx_util.save_txt2docx(str(sec.get("text") or ""), document)
                for url in sec.get("images", []) or []:
                    add_image(url)
                titles = sec.get("table_titles") or []
                for idx, rows in enumerate(sec.get("tables", []) or []):
                    add_plain_table(rows, titles[idx] if idx < len(titles) else "")
            for child in sec.get("children", []) or []:
                add_section(child, level + 1)

        content = obj.content if isinstance(obj.content, dict) else self.__normalize_content(obj.content)
        sections = content.get("sections", [])
        cover = next((s for s in sections if is_cover(s)), None)
        revision = next((s for s in sections if is_revision(s)), None)
        body = [s for s in sections if not is_cover(s) and not is_revision(s)]

        # 封面：上方留白，使大标题落在封面页垂直中部
        for _ in range(10):
            document.add_paragraph("")
        title_para = document.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.line_spacing = 1.5
        docx_util.fonted_txt(title_para, DOC_TITLE, font_size=22.0, bold=False)
        for _ in range(6):
            document.add_paragraph("")
        for rows in (cover or {}).get("tables", []) or []:
            add_plain_table(rows, no_merge=True)
        document.add_page_break()
        # 文件修订记录
        if revision:
            docx_util.save_title2docx("文件修订记录", document, level=1)
            for rows in revision.get("tables", []) or []:
                add_plain_table(rows)
            document.add_page_break()
        # 目录
        docx_util.insert_toc_field(document, "1-4")
        document.add_page_break()
        # 正文
        for sec in body:
            add_section(sec, 1)

        document.save(output)
        output.seek(0)
