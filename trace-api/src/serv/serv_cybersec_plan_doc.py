#!/usr/bin/env python
# encoding: utf-8

# 网络安全风险管理计划服务层，参考网络安全风险管理报告（serv_cybersec_doc）+ 风险管理计划（serv_rmp_doc）。
# 整份文档以 content(JSON)「章节树」存储，字段用 text（与报告一致，非 body）。
# 与网络安全风险管理报告（serv_cybersec_doc）独立并行——本模块是计划阶段文档。
# 自动获取：产品信息（名称/型号/完整版本/适用范围）、物理拓扑图、体系结构图、网络安全流程图、封面/修订记录人员日期、阶段活动时间、网络安全风险管理小组。

import copy
import io
import base64
import json
import logging
import os
import re
from datetime import datetime
from typing import Any

from sqlalchemy import delete, func, select
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from ..model.product import Product
from ..model.cybersec_plan_doc import CybersecPlanDoc
from ..model.company_info import CompanyInfo
from ..model.project_member import ProjectMember
from ..model.project_timeline import ProjectTimelineRow, ProjectTimelineCell
from ..model.prod_runtime_env import ProdRuntimeEnv
from ..model.doc_file import DocFile
from ..model.prod_dhf import ProdDhf
from ..obj import Page, Resp
from ..obj.tobj_cybersec_plan_doc import CybersecPlanDocForm
from ..obj.vobj_cybersec_plan_doc import CybersecPlanDocObj
from ..utils.i18n import ts
from ..utils.sql_ctx import db
from . import msg_err_db
from . import serv_review_util
from .serv_utils import new_version, sync_file_no_version, docx_util
from .serv_prod_runtime_env import DEFAULT_RUNTIME_ENV
from .serv_doc_file import pick_doc_image_file_row

logger = logging.getLogger(__name__)

DOC_TITLE = "网络安全风险管理计划"
DATE_KEYWORDS = ["网络安全风险管理计划", "网络安全风险管理", "风险管理计划"]

TEAM_ROLE_KEYWORDS = {
    "网络安全负责人": ["网络安全负责人", "安全负责人"],
    "软件开发负责人": ["软件开发负责人", "研发负责人", "开发负责人", "产品开发负责人"],
    "QA负责人": ["QA负责人", "QA", "质量负责人"],
    "验证和确认负责人": ["验证和确认负责人", "验证负责人", "确认负责人", "测试负责人"],
}

DEFAULT_CONTENT = {"sections": []}
_DEFAULT_CONTENT_FILE = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "src-res", "cybersec_plan_default_content.json"
)
try:
    with open(_DEFAULT_CONTENT_FILE, encoding="utf-8") as _f:
        _loaded = json.load(_f)
    if isinstance(_loaded, dict) and isinstance(_loaded.get("sections"), list) and _loaded["sections"]:
        DEFAULT_CONTENT = _loaded
except Exception:
    logger.exception("加载网络安全风险管理计划默认内容资源失败")


class Server(object):

    # ---------------- 归一化（与报告一致，用 text 字段） ----------------
    def __normalize_node(self, node):
        if not isinstance(node, dict):
            return {"title": str(node or ""), "text": "", "text_before": "", "text_after": "", "tables": [], "children": []}
        result = dict(node)
        result["title"] = str(result.get("title") or "")
        result["text"] = str(result.get("text") or "")
        result["text_before"] = str(result.get("text_before") or "")
        result["text_after"] = str(result.get("text_after") or "")
        if "body" in result and not result.get("text"):
            result["text"] = str(result.get("body") or "")
        tables = result.get("tables")
        norm_tables = []
        if isinstance(tables, list):
            for table in tables:
                if isinstance(table, list):
                    norm_tables.append([[str(c) if c is not None else "" for c in (row or [])] for row in table if isinstance(row, list)])
        result["tables"] = norm_tables
        blocks = result.get("blocks")
        if isinstance(blocks, list):
            norm_blocks = []
            for blk in blocks:
                if not isinstance(blk, dict):
                    continue
                btype = blk.get("type")
                if btype == "text":
                    norm_blocks.append({"type": "text", "text": str(blk.get("text") or "")})
                elif btype == "table":
                    tbl = blk.get("table") or blk.get("rows") or []
                    norm_tbl = [[str(c) if c is not None else "" for c in (row or [])] for row in tbl if isinstance(row, list)]
                    norm_blocks.append({"type": "table", "table": norm_tbl})
            result["blocks"] = norm_blocks
        else:
            result.pop("blocks", None)
        children = result.get("children")
        result["children"] = [self.__normalize_node(c) for c in children] if isinstance(children, list) else []
        return result

    def __normalize_content(self, content):
        if not isinstance(content, dict) or not isinstance(content.get("sections"), list):
            return copy.deepcopy(DEFAULT_CONTENT)
        return {"sections": [self.__normalize_node(s) for s in content["sections"]], "productName": str(content.get("productName") or "")}

    # ---------------- 工具方法 ----------------
    def __normalize_title_text(self, value):
        return re.sub(r"\s+", "", str(value or ""))

    def __is_cover_section(self, section):
        title = self.__normalize_title_text((section or {}).get("title", ""))
        return (section or {}).get("ref_type") == "cover" or title == DOC_TITLE

    def __is_revision_section(self, section):
        title = self.__normalize_title_text((section or {}).get("title", ""))
        return (section or {}).get("ref_type") == "revision" or title == "文件修订记录"

    def __member_name(self, prod_id, keywords):
        for kw in keywords:
            row = db.session.execute(
                select(ProjectMember).where(ProjectMember.prod_id == prod_id, ProjectMember.role.like(f"%{kw}%"))
                .order_by(ProjectMember.sort_order.asc(), ProjectMember.id.asc())
            ).scalars().first()
            if row and (row.name or "").strip():
                return row.name.strip()
        return ""

    @staticmethod
    def __to_int(v):
        digits = re.sub(r"[^\d]", "", str(v or ""))
        return int(digits) if digits else None

    def __date_rows(self, prod_id):
        tl_rows = db.session.execute(
            select(ProjectTimelineRow).where(ProjectTimelineRow.prod_id == prod_id)
        ).scalars().all()
        return [r for r in tl_rows if (r.row_type or "date") == "date" and self.__to_int(r.year) and self.__to_int(r.month)]

    def __release_date(self, prod_id):
        date_rows = self.__date_rows(prod_id)
        if not date_rows:
            return ""
        cell_map = {}
        for c in db.session.execute(
            select(ProjectTimelineCell).where(ProjectTimelineCell.row_id.in_([r.id for r in date_rows]))
        ).scalars().all():
            cell_map.setdefault(c.row_id, []).append(c.output_result or "")

        def date_key(r):
            return self.__to_int(r.year) * 10000 + self.__to_int(r.month) * 100 + (self.__to_int(r.day) or 0)

        file_rows = [r for r in date_rows if any(k in str(v or "") for k in DATE_KEYWORDS for v in cell_map.get(r.id, []))]
        if not file_rows:
            return ""
        fr = min(file_rows, key=date_key)
        return f"{self.__to_int(fr.year)}年{self.__to_int(fr.month)}月{self.__to_int(fr.day)}日"

    def __stage_timeline_dates(self, product_id, keywords):
        if not product_id or not keywords:
            return (None, None)
        rows = db.session.execute(
            select(ProjectTimelineRow).where(ProjectTimelineRow.prod_id == product_id, ProjectTimelineRow.row_type == "date")
        ).scalars().all()
        if not rows:
            return (None, None)
        row_map = {r.id: r for r in rows}
        cells = db.session.execute(
            select(ProjectTimelineCell).where(ProjectTimelineCell.row_id.in_(list(row_map.keys())))
        ).scalars().all()
        matched_row_ids = set()
        for c in cells:
            txt = str(c.output_result or "")
            if any(k in txt for k in keywords):
                matched_row_ids.add(c.row_id)

        def parse_date(r):
            y, m = self.__to_int(r.year), self.__to_int(r.month)
            if y is None or m is None:
                return None
            try:
                from datetime import date as d
                return d(y, m, self.__to_int(r.day) or 1)
            except Exception:
                return None

        dates = [dd for dd in (parse_date(row_map[i]) for i in matched_row_ids) if dd]
        if not dates:
            return (None, None)
        fmt = lambda d: f"{d.year}年{d.month}月{d.day}日"
        return (fmt(min(dates)), fmt(max(dates)))

    # ---------------- 自动获取：封面/修订记录 ----------------
    def __autofill_front_matter(self, content, product_id, version):
        if not isinstance(content, dict):
            return content
        reviser = self.__member_name(product_id, ("网络安全负责人", "产品经理", "项目经理")) if product_id else ""
        auditor = self.__member_name(product_id, ("QA负责人", "质量负责人", "QA")) if product_id else ""
        approver = self.__member_name(product_id, ("产品负责人", "研发负责人", "管理者代表")) if product_id else ""
        rev_date = serv_review_util.review_date(product_id, DATE_KEYWORDS) if product_id else ""
        rev_date = rev_date or ""
        ver = str(version or "")
        for section in (content.get("sections") or []):
            if self.__is_cover_section(section):
                for table in (section.get("tables") or []):
                    for row in table:
                        if not isinstance(row, list) or not row:
                            continue
                        label = str(row[0]).strip()

                        def set_name(val):
                            if val and len(row) >= 2 and not str(row[1] or "").strip():
                                row[1] = val

                        def set_date(val):
                            if val and len(row) >= 4 and not str(row[3] or "").strip():
                                row[3] = val
                        if len(row) >= 4 and str(row[2]).strip() == "文件版本" and ver and not str(row[3] or "").strip():
                            row[3] = ver
                        if label == "编制部门":
                            set_name("研发部")
                        elif label in ("编制人", "审核人", "批准人"):
                            set_date(rev_date)
                        elif label == "生效日期":
                            set_name(rev_date)
            elif self.__is_revision_section(section):
                for table in (section.get("tables") or []):
                    if not isinstance(table, list) or not table:
                        continue
                    if len(table) >= 2 and isinstance(table[1], list) and len(table[1]) >= 5:
                        row = table[1]

                        def set_if(i, val):
                            if val and not str(row[i] if i < len(row) else "").strip():
                                row[i] = val
                        set_if(0, rev_date)
                        set_if(1, ver)
                        if not str(row[2] or "").strip():
                            row[2] = "首次发布"
                        set_if(3, reviser)
                        set_if(4, approver)
        serv_review_util.fill_cover_signers(content, serv_review_util.cover_signers(product_id, "cybersec")) if product_id else None
        return content

    # ---------------- 自动获取：阶段活动时间 ----------------
    def __autofill_stage_activity(self, content, product_id):
        if not isinstance(content, dict) or not product_id:
            return content

        def find_section(secs):
            for s in secs or []:
                if isinstance(s, dict):
                    title = re.sub(r"\s+", "", str(s.get("title") or ""))
                    if s.get("ref_type") == "stage_activity" or "阶段活动" in title:
                        return s
                    found = find_section(s.get("children"))
                    if found:
                        return found
            return None

        sec = find_section(content.get("sections"))
        if not sec:
            return content
        tables = sec.get("tables") or []
        if not tables or not isinstance(tables[0], list) or len(tables[0]) < 2:
            return content
        table = tables[0]
        header = table[0] if isinstance(table[0], list) else []

        def col_of(name):
            for i, c in enumerate(header):
                if name in re.sub(r"\s+", "", str(c or "")):
                    return i
            return -1

        c_start, c_end = col_of("开始"), col_of("结束")
        if c_start < 0 or c_end < 0:
            return content

        def phase_keywords(row_text):
            t = re.sub(r"\s+", "", row_text)
            if "更新维护" in t or "动态" in t:
                return ["网络安全维护计划", "网络安全维护"]
            if "验证" in t or "部署" in t:
                return ["网络安全风险管理报告", "网络安全扫描"]
            if "控制措施" in t:
                return ["网络安全风险管理报告", "风险管理报告"]
            if "风险评估" in t:
                return ["网络安全风险管理报告", "风险评估"]
            if "威胁识别" in t or "威胁建模" in t:
                return ["网络安全扫描", "SBOM", "网络安全风险管理计划"]
            if "准备" in t:
                return ["网络安全风险管理计划", "风险管理计划"]
            return None

        good_date = re.compile(r"^\d{4}年\d{1,2}月\d{1,2}日$")

        def norm_cell(val, new_val):
            if new_val:
                return new_val
            return "" if not good_date.match(str(val or "").strip()) else val

        for row in table[1:]:
            if not isinstance(row, list) or len(row) <= max(c_start, c_end):
                continue
            row_text = " ".join(str(c or "") for c in row)
            keywords = phase_keywords(row_text)
            start_str, end_str = self.__stage_timeline_dates(product_id, keywords) if keywords else (None, None)
            row[c_start] = norm_cell(row[c_start], start_str)
            row[c_end] = norm_cell(row[c_end], end_str)
        return content

    # ---------------- 自动获取：网络安全风险管理小组 ----------------
    def __autofill_team(self, content, product_id):
        if not isinstance(content, dict) or not product_id:
            return content

        def find_section(secs):
            for s in secs or []:
                if isinstance(s, dict):
                    if s.get("ref_type") == "team" or "网络安全风险管理小组" in str(s.get("title") or ""):
                        return s
                    found = find_section(s.get("children"))
                    if found:
                        return found
            return None

        sec = find_section(content.get("sections"))
        if not sec:
            return content
        for table in (sec.get("tables") or []):
            if not table or not isinstance(table[0], list):
                continue
            header = [str(c).strip() for c in table[0][:3]]
            # 兼容两种表头：项目角色/姓名/职责 或 项目角色/部门/职责
            if header[0] != "项目角色":
                continue
            # 部门映射：根据角色关键字映射到部门
            dept_map = {
                "网络安全负责人": "产品部",
                "软件开发负责人": "研发部",
                "QA负责人": "质量部",
                "验证和确认负责人": "研发部",
            }
            is_dept_col = header[1] == "部门"
            name_col = 1  # 第2列为姓名/部门列
            for row in table[1:]:
                if len(row) < 3:
                    continue
                role = str(row[0] or "").strip()
                if not role:
                    continue
                keywords = TEAM_ROLE_KEYWORDS.get(role)
                if not keywords:
                    continue
                name = self.__member_name(product_id, keywords)
                if is_dept_col:
                    # 部门列：只填部门名称
                    dept = dept_map.get(role, "")
                    if dept:
                        row[name_col] = dept
                else:
                    # 姓名列：直接填姓名
                    if name:
                        row[name_col] = name
        return content

    # ---------------- 自动获取：时间计划（阶段时间从时间逻辑线获取） ----------------
    def __autofill_time_plan(self, content, product_id):
        if not isinstance(content, dict) or not product_id:
            return content

        def find_section(secs):
            for s in secs or []:
                if isinstance(s, dict):
                    if s.get("ref_type") == "time_plan" or "时间计划" in str(s.get("title") or ""):
                        return s
                    found = find_section(s.get("children"))
                    if found:
                        return found
            return None

        sec = find_section(content.get("sections"))
        if not sec:
            return content
        tables = sec.get("tables") or []
        if not tables or not isinstance(tables[0], list) or len(tables[0]) < 2:
            return content
        table = tables[0]
        header = table[0] if isinstance(table[0], list) else []

        def col_of(name):
            for i, c in enumerate(header):
                if name in re.sub(r"\s+", "", str(c or "")):
                    return i
            return -1

        good_date = re.compile(r"^\d{4}年\d{1,2}月\d{1,2}日$")

        def norm_cell(val, new_val):
            if new_val:
                return new_val
            return "" if not good_date.match(str(val or "").strip()) else val

        def phase_keywords(row_text):
            t = re.sub(r"\s+", "", row_text)
            if "动态维护" in t or "持续" in t:
                return None
            if "网络安全扫描" in t:
                return ["网络安全扫描"]
            if "SBOM安全性" in t:
                return ["SBOM安全性风险评估"]
            if "SBOM信息" in t:
                return ["SBOM工具清单", "SBOM"]
            if "控制措施" in t or "制定控制" in t:
                return ["网络安全风险管理报告", "风险管理报告"]
            if "风险评估" in t:
                return ["网络安全风险管理报告", "风险评估"]
            if "威胁识别" in t or "威胁建模" in t:
                return ["网络安全扫描", "SBOM", "网络安全风险管理计划"]
            if "准备" in t:
                return ["网络安全风险管理计划", "风险管理计划"]
            return None

        c_plan = col_of("计划完成时间")
        if c_plan >= 0:
            for row in table[1:]:
                if not isinstance(row, list) or len(row) <= c_plan:
                    continue
                row_text = " ".join(str(c or "") for c in row)
                keywords = phase_keywords(row_text)
                if not keywords:
                    continue
                start_str, end_str = self.__stage_timeline_dates(product_id, keywords)
                plan_date = end_str or start_str or ""
                if plan_date:
                    row[c_plan] = norm_cell(row[c_plan], plan_date)
        return content

    # ---------------- 自动获取：运行环境（参考 serv_ftr_doc.__fill_runtime_env） ----------------
    def __autofill_runtime_env(self, content, product_id):
        if not isinstance(content, dict) or not product_id:
            return content
        env = dict(DEFAULT_RUNTIME_ENV)
        row = db.session.execute(select(ProdRuntimeEnv).where(ProdRuntimeEnv.prod_id == product_id)).scalars().first()
        if row:
            for key in DEFAULT_RUNTIME_ENV.keys():
                val = getattr(row, key, None)
                if val is not None and str(val).strip():
                    env[key] = val

        def overwrite_col1(table, label_map):
            for r in table:
                if not isinstance(r, list) or len(r) < 2:
                    continue
                key = str(r[0]).strip()
                if key in label_map and str(label_map[key] or "").strip():
                    r[1] = label_map[key]

        def fill_node(node):
            title = str(node.get("title") or "")
            tables = node.get("tables") or []
            if "运行环境" in title and len(tables) >= 4:
                # 表1 服务器硬件（2列：配置/要求）
                t1 = tables[0]
                for r in t1:
                    if not isinstance(r, list) or len(r) < 2:
                        continue
                    key = str(r[0]).strip()
                    if key == "CPU": r[1] = env.get("srv_cpu", "")
                    elif key == "内存": r[1] = env.get("srv_memory", "")
                    elif key == "GPU": r[1] = env.get("srv_gpu", "")
                    elif key == "硬盘": r[1] = env.get("srv_disk", "")
                    elif key == "网卡": r[1] = env.get("srv_nic", "")
                # 表2 服务器软件（2列：操作系统/CUDA）
                t2 = tables[1]
                if len(t2) >= 3 and isinstance(t2[2], list) and len(t2[2]) >= 2:
                    t2[2][0] = env.get("srv_os", "")
                    t2[2][1] = env.get("srv_cuda", "")
                # 表3 用户端（2列：配置/要求）
                t3 = tables[2]
                for r in t3:
                    if not isinstance(r, list) or len(r) < 2:
                        continue
                    key = str(r[0]).strip()
                    if key == "CPU": r[1] = env.get("cli_cpu", "")
                    elif key == "内存": r[1] = env.get("cli_memory", "")
                    elif key == "显示器分辨率": r[1] = env.get("cli_resolution", "")
                    elif key == "操作系统": r[1] = env.get("cli_os", "")
                    elif key == "浏览器": r[1] = env.get("cli_browser", "")
                # 表4 网络（3列：配置/局域网/广域网）
                t4 = tables[3]
                for r in t4:
                    if not isinstance(r, list) or len(r) < 3:
                        continue
                    if str(r[0]).strip() == "带宽":
                        r[1] = env.get("net_lan", "")
                        r[2] = env.get("net_wan", "")
            for c in (node.get("children") or []):
                fill_node(c)

        for s in (content.get("sections") or []):
            if "运行环境" in str(s.get("title") or ""):
                fill_node(s)
            else:
                for c in (s.get("children") or []):
                    if "运行环境" in str(c.get("title") or ""):
                        fill_node(c)
        return content

    # ---------------- 自动获取：产品信息/适用范围/产品描述 ----------------
    def __autofill_product_info(self, content, product_id):
        if not isinstance(content, dict) or not product_id:
            return content
        product = db.session.execute(select(Product).where(Product.id == product_id)).scalars().first()
        if not product:
            return content
        pname = (product.name or "").strip()
        content["productName"] = pname

        def build_desc():
            return f"本文件适用的产品为的软件产品：\n产品名称：{pname}\n完整版本：{product.full_version or ''}"

        def fill_sections(secs):
            for s in secs or []:
                if not isinstance(s, dict):
                    continue
                key = re.sub(r"^[0-9０-９.．\s、]+", "", str(s.get("title") or "")).strip()
                if key == "产品定义":
                    s["text"] = build_desc()
                fill_sections(s.get("children"))

        fill_sections(content.get("sections"))
        return content

    # ---------------- 统一自动获取入口 ----------------
    def __autofill(self, content, product_id, version=""):
        content = self.__autofill_front_matter(content, product_id, version)
        content = self.__autofill_time_plan(content, product_id)
        content = self.__autofill_team(content, product_id)
        content = self.__autofill_product_info(content, product_id)
        content = self.__autofill_runtime_env(content, product_id)
        return content

    def __dhf_file_no(self, prod_id):
        """从产品DHF中查找名称包含「网络安全风险管理计划」的记录，返回其编号"""
        if not prod_id:
            return ""
        row = db.session.execute(
            select(ProdDhf).where(ProdDhf.prod_id == prod_id, ProdDhf.name.like("%网络安全风险管理计划%"))
        ).scalars().first()
        return (row.code or "").strip() if row and row.code else ""

    # ---------------- 文档 CRUD ----------------
    def add_cybersec_plan_doc(self, form: CybersecPlanDocForm):
        try:
            # 前置查重：同产品下相同版本已存在则拒绝
            cnt = db.session.execute(
                select(func.count(CybersecPlanDoc.id)).where(
                    CybersecPlanDoc.product_id == form.product_id,
                    CybersecPlanDoc.version == form.version,
                )
            ).scalar()
            if cnt > 0:
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            content = form.content or copy.deepcopy(DEFAULT_CONTENT)
            content = self.__normalize_content(content)
            content = self.__autofill(content, form.product_id, form.version)
            row = CybersecPlanDoc(
                product_id=form.product_id, version=form.version, file_no=form.file_no or "",
                change_log=form.change_log or "", content=content,
            )
            db.session.add(row)
            db.session.commit()
            return Resp.resp_ok(data=CybersecPlanDocForm(id=row.id, product_id=row.product_id, version=row.version))
        except Exception:
            logger.exception("")
            db.session.rollback()
        return Resp.resp_err(msg=ts("msg_err_db"))

    def duplicate_cybersec_plan_doc(self, id: int, target_product_id: int = None):
        row = db.session.execute(select(CybersecPlanDoc).where(CybersecPlanDoc.id == id)).scalars().first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        target_pid = target_product_id or row.product_id
        new_ver = new_version(row.version)
        base_file_no = (row.file_no or "").strip() or self.__dhf_file_no(target_pid)
        new_row = CybersecPlanDoc(
            product_id=target_pid, version=new_ver,
            file_no=sync_file_no_version(base_file_no, new_ver) or base_file_no,
            change_log=row.change_log or "", content=copy.deepcopy(row.content),
        )
        db.session.add(new_row)
        try:
            db.session.flush()
        except Exception as e:
            db.session.rollback()
            if "Duplicate" in str(e) or "UNIQUE" in str(e):
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            logger.exception("")
            return Resp.resp_err()
        db.session.commit()
        return Resp.resp_ok(data=CybersecPlanDocForm(id=new_row.id, product_id=new_row.product_id, version=new_row.version))

    def update_cybersec_plan_doc(self, form: CybersecPlanDocForm):
        row = db.session.execute(select(CybersecPlanDoc).where(CybersecPlanDoc.id == form.id)).scalars().first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        if form.product_id is not None:
            row.product_id = form.product_id
        if form.version is not None:
            row.version = form.version
        if form.file_no is not None:
            row.file_no = form.file_no
        if form.change_log is not None:
            row.change_log = form.change_log
        if form.content is not None:
            row.content = self.__normalize_content(form.content)
        try:
            db.session.flush()
        except Exception as e:
            db.session.rollback()
            if "Duplicate" in str(e) or "UNIQUE" in str(e):
                return Resp.resp_err(msg=ts("msg_obj_exist"))
            logger.exception("")
            return Resp.resp_err()
        db.session.commit()
        return Resp.resp_ok()

    def delete_cybersec_plan_doc(self, id: int):
        row = db.session.execute(select(CybersecPlanDoc).where(CybersecPlanDoc.id == id)).scalars().first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        db.session.delete(row)
        db.session.commit()
        return Resp.resp_ok()

    def get_cybersec_plan_doc(self, id: int):
        row = db.session.execute(select(CybersecPlanDoc).where(CybersecPlanDoc.id == id)).scalars().first()
        if not row:
            return Resp.resp_err(msg=ts("msg_obj_null"))
        product = db.session.execute(select(Product).where(Product.id == row.product_id)).scalars().first()
        content = self.__normalize_content(row.content)
        content = self.__autofill(content, row.product_id, row.version)
        file_no = (row.file_no or "").strip() or self.__dhf_file_no(row.product_id)
        obj = CybersecPlanDocObj(
            id=row.id, product_id=row.product_id, version=row.version, file_no=file_no,
            change_log=row.change_log, content=content,
            product_name=(product.name if product else ""),
            product_version=(product.release_version if product else ""),
            product_full_version=(product.full_version if product else ""),
            product_type_code=(product.type_code if product else ""),
            create_time=row.create_time,
        )
        return Resp.resp_ok(data=obj)

    def list_cybersec_plan_doc(self, op_user=None, product_id: int = 0, version: str = None,
                                page_index: int = 0, page_size: int = 10):
        q = select(CybersecPlanDoc)
        if product_id:
            q = q.where(CybersecPlanDoc.product_id == product_id)
        if version:
            q = q.where(CybersecPlanDoc.version.like(f"%{version}%"))
        total = db.session.execute(select(func.count()).select_from(q.subquery())).scalar()
        rows = db.session.execute(q.order_by(CybersecPlanDoc.id.desc()).offset(page_index * page_size).limit(page_size)).scalars().all()
        prod_ids = list(set(r.product_id for r in rows))
        prod_map = {}
        if prod_ids:
            for p in db.session.execute(select(Product).where(Product.id.in_(prod_ids))).scalars().all():
                prod_map[p.id] = p
        objs = []
        for r in rows:
            p = prod_map.get(r.product_id)
            file_no = (r.file_no or "").strip() or self.__dhf_file_no(r.product_id)
            objs.append(CybersecPlanDocObj(
                id=r.id, product_id=r.product_id, version=r.version, file_no=file_no, change_log=r.change_log,
                product_name=(p.name if p else ""), product_version=(p.release_version if p else ""),
                product_full_version=(p.full_version if p else ""), product_type_code=(p.type_code if p else ""),
                create_time=r.create_time,
            ))
        return Resp.resp_ok(data=Page(total=total, rows=objs, page_index=page_index, page_size=page_size))

    def cybersec_plan_autofill(self, product_id: int, version: str = ""):
        content = self.__normalize_content(None)
        content = self.__autofill(content, product_id, version)
        return Resp.resp_ok(data=content)

    # ---------------- 导出 Word ----------------
    def export_cybersec_plan_doc(self, output, id: int):
        row = db.session.execute(select(CybersecPlanDoc).where(CybersecPlanDoc.id == id)).scalars().first()
        if not row:
            Document().save(output)
            output.seek(0)
            return
        content = self.__normalize_content(row.content)
        content = self.__autofill(content, row.product_id, row.version)
        sections = content.get("sections") or []
        product = db.session.execute(select(Product).where(Product.id == row.product_id)).scalars().first()

        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        document.settings.element.append(update_fields)
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        docx_util.fonted_txt(header_para, row.file_no or "")
        docx_util.add_page_number_footer(section, row.file_no or "")

        product_version = (product.full_version if product else "") or ""

        def normalized_title(value):
            return re.sub(r"\s+", "", str(value or ""))

        def is_cover_section(sec):
            return sec.get("ref_type") == "cover" or normalized_title(sec.get("title")) == DOC_TITLE

        def is_revision_section(sec):
            return sec.get("ref_type") == "revision" or normalized_title(sec.get("title")) == "文件修订记录"

        def write_center_section_title(title, font_size=16.0, bold=True):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.5
            docx_util.fonted_txt(p, title, font_size=font_size, bold=bold)

        def add_blank_lines(count):
            for _ in range(max(0, int(count or 0))):
                document.add_paragraph("")

        def insert_toc_field():
            p = document.add_paragraph()
            run_begin = p.add_run()
            fld_begin = OxmlElement("w:fldChar")
            fld_begin.set(qn("w:fldCharType"), "begin")
            fld_begin.set(qn("w:dirty"), "true")
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = ' TOC \\o "1-3" \\h \\z \\u '
            fld_separate = OxmlElement("w:fldChar")
            fld_separate.set(qn("w:fldCharType"), "separate")
            run_end = p.add_run()
            fld_end = OxmlElement("w:fldChar")
            fld_end.set(qn("w:fldCharType"), "end")
            run_begin._r.append(fld_begin)
            run_begin._r.append(instr)
            run_begin._r.append(fld_separate)
            p.add_run("目录将在打开文档后自动更新")
            run_end._r.append(fld_end)

        def set_cell_text(cell, text, bold=False):
            s = str(text or "")
            if s.startswith("data:image"):
                try:
                    b64 = s.split(",", 1)[1] if "," in s else ""
                    cell.text = ""
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.add_run().add_picture(io.BytesIO(base64.b64decode(b64)), height=Pt(33))
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    return
                except Exception:
                    pass
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.line_spacing = 1.5
            docx_util.fonted_txt(paragraph, s, font_size=10.5, bold=bold)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        def add_plain_table(rows):
            rows = rows or []
            col_count = max([len(r or []) for r in rows] or [0])
            if col_count <= 0:
                return
            table = document.add_table(rows=0, cols=col_count)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for ri, rrow in enumerate(rows):
                cells = table.add_row().cells
                for idx in range(col_count):
                    set_cell_text(cells[idx], rrow[idx] if idx < len(rrow or []) else "", bold=(ri == 0))
            document.add_paragraph()

        def set_cell_bg(cell, color_hex):
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), color_hex)
            tc_pr.append(shd)

        def is_matrix_table(rows):
            for r in rows or []:
                v = str((r or [""])[0] or "").strip()
                if v in ("红色", "黄色", "绿色"):
                    return True
            return False

        def add_matrix_table(rows):
            """风险矩阵表：表头/数据行正常；红色/黄色/绿色行第1列标签+第2列合并剩余列；分数单元格按值着色"""
            rows = rows or []
            col_count = max([len(r or []) for r in rows] or [0])
            if col_count <= 0:
                return
            label_colors = {"红色": "FF0000", "黄色": "FFFF00", "绿色": "00B050"}
            table = document.add_table(rows=0, cols=col_count)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for ri, rrow in enumerate(rows):
                label_cell = str((rrow or [""])[0] or "").strip()
                is_label_row = label_cell in ("红色", "黄色", "绿色")
                cells = table.add_row().cells
                if is_label_row:
                    # 第1列标签（着色）
                    set_cell_text(cells[0], rrow[0] if 0 < len(rrow or []) else "", bold=True)
                    set_cell_bg(cells[0], label_colors.get(label_cell, "FFFFFF"))
                    # 第2列合并剩余列
                    if col_count > 1:
                        merged = cells[1]
                        for ci in range(2, col_count):
                            merged = merged.merge(cells[ci])
                        set_cell_text(merged, rrow[1] if 1 < len(rrow or []) else "")
                else:
                    first_cell = (rows[0] or [""])[0] if (rows and len(rows[0] or []) > 0) else ""
                    is_header = ri == 0 or (ri == 1 and len(str(first_cell or "").strip()) == 0 and len(rows[0] or []) == 1)
                    for idx in range(col_count):
                        set_cell_text(cells[idx], rrow[idx] if idx < len(rrow or []) else "", bold=is_header)
                        # 分数单元格着色（第3列起，数字1~25）
                        if idx >= 2:
                            num_str = str(rrow[idx] if idx < len(rrow or []) else "").strip()
                            try:
                                num = int(num_str)
                                if num > 0:
                                    if num <= 5:
                                        set_cell_bg(cells[idx], "00B050")
                                    elif num <= 12:
                                        set_cell_bg(cells[idx], "FFFF00")
                                    else:
                                        set_cell_bg(cells[idx], "FF0000")
                            except (ValueError, TypeError):
                                pass
            document.add_paragraph()

        def add_section_image_by_category(category):
            """从图表文件管理获取图片并插入文档"""
            try:
                img_row = pick_doc_image_file_row(row.product_id, category, row.version or "", product_version)
                if img_row and img_row.file_url:
                    docx_util.save_img2docx(img_row.file_url, document, mw=520, mh=520)
                    return True
            except Exception:
                logger.exception("cybersec_plan_export_image_failed")
            return False

        # 图题关键词 -> 图表文件管理分类
        caption_cats = [("物理拓扑图", "img_topo"), ("体系结构图", "img_struct")]

        def render_text_with_anchors(text):
            """逐行渲染正文；遇到「图N xxx」图题行且能匹配到分类时，在该行后插入对应图片"""
            buf = []

            def flush():
                if buf:
                    docx_util.save_txt2docx("\n".join(buf), document)
                    buf.clear()

            for ln in str(text or "").split("\n"):
                buf.append(ln)
                s = ln.strip()
                if re.match(r"^图\s*\d", s):
                    cat = next((c for kw, c in caption_cats if kw in s), None)
                    if cat:
                        flush()
                        add_section_image_by_category(cat)
            flush()

        def add_section(sec: dict, level: int = 1):
            title = sec.get("title", "")
            if title and not is_cover_section(sec) and not is_revision_section(sec):
                docx_util.save_title2docx(title, document, level=max(1, min(level, 9)))
            # 有序内容块（text/table 交错）：按块顺序输出，忽略平铺 text/tables
            blocks = sec.get("blocks")
            if isinstance(blocks, list) and blocks:
                for block in blocks:
                    if not isinstance(block, dict):
                        continue
                    if block.get("type") == "table":
                        tbl_rows = block.get("table") or []
                        add_matrix_table(tbl_rows) if is_matrix_table(tbl_rows) else add_plain_table(tbl_rows)
                    elif block.get("type") == "text":
                        txt = str(block.get("text") or "")
                        if txt:
                            render_text_with_anchors(txt)
                for child in sec.get("children", []) or []:
                    add_section(child, level + 1)
                return
            # 普通正文（按锚点嵌图）
            text = str(sec.get("text") or "")
            if text:
                render_text_with_anchors(text)
            # 表格
            for table_rows in sec.get("tables", []) or []:
                add_plain_table(table_rows)
            for child in sec.get("children", []) or []:
                add_section(child, level + 1)

        cover_section = next((s for s in sections if is_cover_section(s)), None)
        revision_section = next((s for s in sections if is_revision_section(s)), None)
        body_sections = [s for s in sections if not is_cover_section(s) and not is_revision_section(s)]

        # 封面
        add_blank_lines(12)
        write_center_section_title(DOC_TITLE, font_size=22.0, bold=True)
        add_blank_lines(6)
        for table_rows in (cover_section or {}).get("tables", []) or []:
            add_plain_table(table_rows)

        # 修订记录
        document.add_page_break()
        write_center_section_title("文件修订记录", font_size=14.0, bold=True)
        add_blank_lines(2)
        for table_rows in (revision_section or {}).get("tables", []) or []:
            add_plain_table(table_rows)

        # 目录
        document.add_page_break()
        write_center_section_title("目录", font_size=16.0, bold=True)
        insert_toc_field()

        # 正文
        document.add_page_break()
        for sec in body_sections:
            add_section(sec)
        docx_util.fill_toc_cache(document)
        document.save(output)
        output.seek(0)