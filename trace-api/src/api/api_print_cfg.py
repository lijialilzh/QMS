#!/usr/bin/env python
# encoding: utf-8

# 打印服务配置接口层：支持多台打印机配置（列表CRUD），配置IPP/TCP9100打印机，支持测试连接和一键打印。

import io
import asyncio
from typing import Any
from fastapi import APIRouter

from sqlalchemy import select, func

from ..obj import Resp, Page
from ..obj.tobj_role import Perms
from ..model.print_service_cfg import PrintServiceCfg
from ..utils.sql_ctx import db
from ..utils.i18n import ts
from ..utils import ipp_client
from . import CtxUser, try_log

router = APIRouter()


def _build_test_page(printer_name: str, printer_host: str) -> bytes:
    """生成 PostScript 测试页内容，用于验证打印机端到端打印链路。
    使用纯 ASCII PostScript，确保兼容所有 PostScript 打印机。"""
    from datetime import datetime
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    # PostScript show 命令需要括号内的字符串，括号需转义
    safe_name = printer_name.replace("(", "\\(").replace(")", "\\)")
    safe_host = printer_host.replace("(", "\\(").replace(")", "\\)")
    ps = f"""%!PS-Adobe-3.0
%%Title: QMS Print Test Page
%%Creator: QMS Print Service
%%Pages: 1
%%EndComments
%%Page: 1 1

/Helvetica-Bold findfont 24 scalefont setfont
72 720 moveto (QMS Print Service Test Page) show

/Helvetica findfont 12 scalefont setfont
72 690 moveto (------------------------------------------------) show

72 660 moveto (Printer Name: {safe_name}) show
72 640 moveto (Printer Host: {safe_host}) show
72 620 moveto (Test Time: {now}) show

72 580 moveto (If you see this test page, the printer connection is OK.) show
72 560 moveto (You can now use the One-Click Print function.) show

1 setlinewidth
72 540 moveto 540 540 lineto stroke

showpage
%%EOF
"""
    return ps.encode("ascii", errors="replace")


def _to_obj(row: PrintServiceCfg) -> dict:
    return {
        "id": row.id,
        "printer_host": row.printer_host or "",
        "printer_port": row.printer_port or (9100 if row.protocol == "tcp9100" else 631),
        "printer_name": row.printer_name or "",
        "printer_uri": row.printer_uri or "",
        "protocol": row.protocol or "tcp9100",
        "is_default": row.is_default,
        "remark": row.remark or "",
        "create_time": str(row.create_time or ""),
    }


@router.get("/list_print_cfg", summary="查询打印服务配置列表")
@try_log(perm=Perms.print_cfg_view)
async def list_print_cfg(page_index: int = 0, page_size: int = 100):
    total = db.session.execute(select(func.count(PrintServiceCfg.id))).scalar() or 0
    rows = db.session.execute(
        select(PrintServiceCfg).order_by(PrintServiceCfg.is_default.desc(), PrintServiceCfg.id.asc())
        .offset(page_index * page_size).limit(page_size)
    ).scalars().all()
    return Resp.resp_ok(data=Page(total=total, rows=[_to_obj(r) for r in rows], page_index=page_index, page_size=page_size))


@router.get("/get_default_print_cfg", summary="获取默认打印服务配置")
@try_log(perm=Perms.print_cfg_view)
async def get_default_print_cfg():
    row = db.session.execute(select(PrintServiceCfg).where(PrintServiceCfg.is_default == 1)).scalars().first()
    return Resp.resp_ok(data=_to_obj(row) if row else None)


@router.post("/add_print_cfg", summary="添加打印服务配置")
@try_log(perm=Perms.print_cfg_edit)
async def add_print_cfg(form: dict):
    host = (form.get("printer_host") or "").strip()
    port = int(form.get("printer_port") or (9100 if form.get("protocol") == "tcp9100" else 631))
    name = (form.get("printer_name") or "").strip()
    uri = (form.get("printer_uri") or "").strip()
    protocol = (form.get("protocol") or "tcp9100").strip()
    remark = (form.get("remark") or "").strip()
    is_default = int(form.get("is_default") or 0)
    if not host:
        return Resp.resp_err(msg="请填写打印机IP/主机名")
    if not uri:
        uri = (f"tcp9100://{host}:{port}") if protocol == "tcp9100" else (f"ipp://{host}:{port}/ipp/print/{name}" if name else f"ipp://{host}:{port}/ipp/print")
    try:
        # 如果设为默认，先清除其他默认
        if is_default:
            for r in db.session.execute(select(PrintServiceCfg).where(PrintServiceCfg.is_default == 1)).scalars().all():
                r.is_default = 0
        row = PrintServiceCfg(printer_host=host, printer_port=port, printer_name=name,
                              printer_uri=uri, protocol=protocol, is_default=is_default, remark=remark)
        db.session.add(row)
        # 如果是第一条记录，自动设为默认
        if db.session.execute(select(func.count(PrintServiceCfg.id))).scalar() == 1:
            row.is_default = 1
        db.session.commit()
        return Resp.resp_ok(data={"id": row.id})
    except Exception:
        db.session.rollback()
        return Resp.resp_err(msg=ts("msg_err_db"))


@router.post("/update_print_cfg", summary="更新打印服务配置")
@try_log(perm=Perms.print_cfg_edit)
async def update_print_cfg(form: dict):
    cfg_id = int(form.get("id") or 0)
    row = db.session.execute(select(PrintServiceCfg).where(PrintServiceCfg.id == cfg_id)).scalars().first()
    if not row:
        return Resp.resp_err(msg="配置不存在")
    host = (form.get("printer_host") or "").strip()
    port = int(form.get("printer_port") or (9100 if form.get("protocol") == "tcp9100" else 631))
    name = (form.get("printer_name") or "").strip()
    uri = (form.get("printer_uri") or "").strip()
    protocol = (form.get("protocol") or "tcp9100").strip()
    remark = (form.get("remark") or "").strip()
    is_default = int(form.get("is_default") or 0)
    if not host:
        return Resp.resp_err(msg="请填写打印机IP/主机名")
    if not uri:
        uri = (f"tcp9100://{host}:{port}") if protocol == "tcp9100" else (f"ipp://{host}:{port}/ipp/print/{name}" if name else f"ipp://{host}:{port}/ipp/print")
    try:
        if is_default:
            for r in db.session.execute(select(PrintServiceCfg).where(PrintServiceCfg.is_default == 1, PrintServiceCfg.id != cfg_id)).scalars().all():
                r.is_default = 0
        row.printer_host = host
        row.printer_port = port
        row.printer_name = name
        row.printer_uri = uri
        row.protocol = protocol
        row.remark = remark
        row.is_default = is_default
        db.session.commit()
        return Resp.resp_ok()
    except Exception:
        db.session.rollback()
        return Resp.resp_err(msg=ts("msg_err_db"))


@router.delete("/delete_print_cfg", summary="删除打印服务配置")
@try_log(perm=Perms.print_cfg_edit)
async def delete_print_cfg(id: int):
    row = db.session.execute(select(PrintServiceCfg).where(PrintServiceCfg.id == id)).scalars().first()
    if not row:
        return Resp.resp_err(msg="配置不存在")
    try:
        was_default = row.is_default
        db.session.delete(row)
        db.session.commit()
        # 如果删的是默认打印机，把第一条设为默认
        if was_default:
            first = db.session.execute(select(PrintServiceCfg).order_by(PrintServiceCfg.id.asc()).limit(1)).scalars().first()
            if first:
                first.is_default = 1
                db.session.commit()
        return Resp.resp_ok()
    except Exception:
        db.session.rollback()
        return Resp.resp_err(msg=ts("msg_err_db"))


@router.post("/set_default_print_cfg", summary="设置默认打印机")
@try_log(perm=Perms.print_cfg_edit)
async def set_default_print_cfg(id: int):
    row = db.session.execute(select(PrintServiceCfg).where(PrintServiceCfg.id == id)).scalars().first()
    if not row:
        return Resp.resp_err(msg="配置不存在")
    try:
        for r in db.session.execute(select(PrintServiceCfg).where(PrintServiceCfg.is_default == 1)).scalars().all():
            r.is_default = 0
        row.is_default = 1
        db.session.commit()
        return Resp.resp_ok()
    except Exception:
        db.session.rollback()
        return Resp.resp_err(msg=ts("msg_err_db"))


@router.get("/test_print_conn", summary="测试打印机连接并打印测试页")
@try_log(perm=Perms.print_cfg_view)
async def test_print_conn(id: int = 0):
    row = db.session.execute(select(PrintServiceCfg).where(PrintServiceCfg.id == id)).scalars().first()
    if not row:
        return Resp.resp_err(msg="打印配置不存在")
    # 第一步：测试 TCP 连接
    if row.protocol == "tcp9100":
        ok, msg = ipp_client.test_printer_9100(row.printer_host, row.printer_port or 9100)
    else:
        ok, msg = ipp_client.test_printer(row.printer_uri, row.printer_host, row.printer_port or 631)
    if not ok:
        return Resp.resp_ok(data={"ok": False, "msg": msg})
    # 第二步：发送测试页到打印机
    # 生成简单 PostScript 测试页（兼容 Generic PostScript Printer）
    test_page = _build_test_page(row.printer_name or "", row.printer_host or "")
    if row.protocol == "tcp9100":
        ok2, msg2 = ipp_client.print_document_9100(row.printer_host, row.printer_port or 9100, "test_page", test_page)
    else:
        ok2, msg2 = ipp_client.print_document(row.printer_uri, row.printer_host, row.printer_port or 631, "test_page", test_page, document_format="application/postscript")
    if ok2:
        return Resp.resp_ok(data={"ok": True, "msg": "连接成功，测试页已发送到打印机，请查看打印机输出"})
    return Resp.resp_ok(data={"ok": False, "msg": f"连接成功但发送测试页失败：{msg2}"})


@router.get("/ipp_print_doc", summary="打印单个文档：生成docx后直接发到默认打印机")
@try_log(perm=Perms.product_view)
async def ipp_print_doc(module_key: str, doc_id: int, with_sign: bool = True):
    from .api_doc_integrate import _SERVERS, _DOC_MODULES, _build_doc_name
    srv = _SERVERS.get(module_key)
    if not srv:
        return Resp.resp_err(msg=f"不支持的文档模块：{module_key}")
    # 追溯分析：方法名 export_doc_trace
    if module_key == "srs_doc_trace":
        method = getattr(srv, "export_doc_trace", None)
    else:
        method = getattr(srv, f"export_{module_key}", None)
    if not method:
        return Resp.resp_err(msg=f"模块 {module_key} 无导出方法")
    cfg = db.session.execute(select(PrintServiceCfg).where(PrintServiceCfg.is_default == 1)).scalars().first()
    if not cfg:
        return Resp.resp_err(msg="未配置默认打印机，请先在基础配置中设置")
    # 设置签名模式（contextvar，仅影响本次打印请求）
    from ..serv.serv_review_util import set_export_sign_mode
    set_export_sign_mode(with_sign)
    try:
        out = io.BytesIO()
        result = method(out, doc_id)
        if asyncio.iscoroutine(result):
            await result
        doc_bytes = out.getvalue()
    except Exception as e:
        return Resp.resp_err(msg=f"生成文档失败：{str(e)[:80]}")
    job_name = _build_doc_name(module_key, doc_id).replace(".docx", "")
    # docx 直接发打印机无法识别，转成 PDF（保留格式）或 PostScript（兜底）再发送
    try:
        print_bytes, doc_format = _docx_to_printable(doc_bytes, job_name)
    except Exception as e:
        return Resp.resp_err(msg=f"文档转换失败：{str(e)[:80]}")
    if cfg.protocol == "tcp9100":
        ok, msg = ipp_client.print_document_9100(cfg.printer_host, cfg.printer_port or 9100, job_name, print_bytes)
    else:
        ok, msg = ipp_client.print_document(cfg.printer_uri, cfg.printer_host, cfg.printer_port or 631, job_name, print_bytes, document_format=doc_format)
    if ok:
        return Resp.resp_ok(data={"msg": msg})
    return Resp.resp_err(msg=msg)


def _docx_to_printable(docx_bytes: bytes, title: str = "") -> tuple:
    """把 docx 转成打印机能识别的格式。
    优先级1：调宿主机 docx2pdf 微服务转 PDF（保留完整格式，开发环境用）
    优先级2：用容器内 soffice 转 PDF（生产环境镜像预装）
    优先级3：用 python-docx 转 PostScript（兜底，仅文字）
    返回 (bytes, document_format)。"""
    import subprocess
    import tempfile
    import os

    # 方案1：调宿主机微服务转 PostScript（docx→PDF→PS via Ghostscript，保留格式+中文）
    try:
        import urllib.request
        boundary = "----qms_boundary"
        body = b"--" + boundary.encode() + b"\r\n"
        body += b'Content-Disposition: form-data; name="file"; filename="input.docx"\r\n'
        body += b"Content-Type: application/octet-stream\r\n\r\n"
        body += docx_bytes + b"\r\n"
        body += b"--" + boundary.encode() + b"--\r\n"
        req = urllib.request.Request(
            "http://host.docker.internal:8765/convert_ps",
            data=body,
            headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=120) as resp:
            ct = resp.headers.get("content-type", "")
            if resp.status == 200 and "postscript" in ct:
                return (resp.read(), "application/postscript")
            if resp.status == 200 and "pdf" in ct:
                return (resp.read(), "application/pdf")
    except Exception:
        pass

    # 方案2：用容器内 soffice 转 PDF（生产环境镜像预装）
    soffice_paths = ["/usr/bin/soffice", "/usr/bin/libreoffice", "/usr/local/bin/soffice"]
    soffice_bin = None
    for p in soffice_paths:
        if os.path.exists(p):
            soffice_bin = p
            break
    if not soffice_bin:
        try:
            subprocess.run(["which", "soffice"], capture_output=True, timeout=3)
            soffice_bin = "soffice"
        except Exception:
            pass

    if soffice_bin:
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                docx_path = os.path.join(tmpdir, "input.docx")
                with open(docx_path, "wb") as f:
                    f.write(docx_bytes)
                subprocess.run(
                    [soffice_bin, "--headless", "--convert-to", "pdf", "--outdir", tmpdir, docx_path],
                    capture_output=True, timeout=60,
                )
                pdf_path = os.path.join(tmpdir, "input.pdf")
                if os.path.exists(pdf_path):
                    with open(pdf_path, "rb") as f:
                        return (f.read(), "application/pdf")
        except Exception:
            pass

    # 方案3：用 python-docx 转 PostScript（兜底，仅文字）
    ps_bytes = _docx_to_postscript(docx_bytes, title)
    return (ps_bytes, "application/postscript")


def _docx_to_postscript(docx_bytes: bytes, title: str = "") -> bytes:
    """用 python-docx 读取 docx 内容，生成 PostScript 页面（兜底方案，仅文字）。"""
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(docx_bytes))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
        lines.append("")

    if not lines:
        lines = ["(empty document)"]

    PAGE_W, PAGE_H = 595, 842
    MARGIN = 50
    LINE_H = 16
    FONT_SIZE = 10
    usable_h = PAGE_H - 2 * MARGIN
    lines_per_page = int(usable_h / LINE_H)

    def esc(s):
        return s.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    ps_lines = [
        "%!PS-Adobe-3.0",
        f"%%Title: {esc(title)}",
        "%%Creator: QMS Print Service",
        f"%%Pages: {(len(lines) + lines_per_page - 1) // lines_per_page}",
        "%%EndComments",
    ]

    page_num = 0
    for start in range(0, len(lines), lines_per_page):
        page_num += 1
        page_lines = lines[start:start + lines_per_page]
        ps_lines.append(f"%%Page: {page_num} {page_num}")
        ps_lines.append("/Helvetica findfont %d scalefont setfont" % FONT_SIZE)
        y = PAGE_H - MARGIN
        for line in page_lines:
            if len(line) > 80:
                line = line[:77] + "..."
            ps_lines.append(f"{MARGIN} {y} moveto ({esc(line)}) show")
            y -= LINE_H
        ps_lines.append("showpage")

    ps_lines.append("%%EOF")
    return "\n".join(ps_lines).encode("ascii", errors="replace")
