#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将安装维护手册模板 docx 转为 imm_default_content.json。"""

import base64
import json
import re
import sys
import zipfile
from io import BytesIO
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "vendor"))
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

IMM_DOCX = Path("/tmp/imm_conv/TX-TF-RCN3V2000-VV-005-A0 安装维护手册.docx")
MD5_ATTACH = Path(
    "/Users/lijiali/Desktop/E/体系-工作文件/1·体系-文件/30·FDA-recist/1.0.1.0/体系文件/测试文件/"
    "TX-TF-RUSV1010-VV-005-A0 MD5值附件.docx"
)
MD5_REVIEW = Path("/Users/lijiali/Desktop/cursor/QMS/测试文件/MD5值-评审记录.docx")
OUT = Path(__file__).resolve().parents[1] / "src-res" / "imm_default_content.json"

CHAPTER1 = {"概述", "安装向导", "运行前系统配置", "安装后测试流程", "卸载向导"}


def rel_map(doc):
    rels = {}
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            rels[rel.rId] = rel.target_part.blob
    return rels


def para_images(p, rels):
    imgs = []
    for blip in p._element.iter(qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if rid and rid in rels:
            imgs.append(f"data:image/png;base64,{base64.b64encode(rels[rid]).decode()}")
    for im in p._element.iter("{urn:schemas-microsoft-com:vml}imagedata"):
        rid = im.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
        if rid and rid in rels:
            imgs.append(f"data:image/png;base64,{base64.b64encode(rels[rid]).decode()}")
    return imgs


def para_text(p):
    return p.text.replace("\u200b", "").strip()


def table_to_grid(tbl):
    return [[para_text(c.paragraphs[0]) if c.paragraphs else "" for c in row.cells] for row in tbl.rows]


def new_node(title="", body="", tables=None, images=None, children=None, ref_type=None):
    n = {"title": title, "body": body, "tables": tables or [], "children": children or [], "images": images or []}
    if ref_type:
        n["ref_type"] = ref_type
    return n


def heading_level(style, text):
    t = text.strip()
    if not t or t == "目录":
        return 0
    if t in CHAPTER1:
        return 1
    s = style or ""
    if s in ("样式1",):
        return 1
    if s == "样式2":
        return 2
    if s in ("样式3",):
        return 3
    if re.match(r"^\d+\.\d+\.\d+", t):
        return 3
    if re.match(r"^\d+\.\d+", t):
        return 2
    if re.match(r"^[1-5]\.\d", t):
        return 2
    return 0


def walk_body(doc, rels):
    """按文档顺序遍历段落与表格。"""
    for child in doc.element.body:
        if child.tag == qn("w:p"):
            yield ("p", Paragraph(child, doc))
        elif child.tag == qn("w:tbl"):
            yield ("t", Table(child, doc))


def parse_main(docx_path):
    doc = Document(str(docx_path))
    rels = rel_map(doc)
    cover_table = None
    revision_table = None
    in_body = False
    roots = []
    stack = []  # (level, node)

    def current_parent(level):
        while stack and stack[-1][0] >= level:
            stack.pop()
        if not stack:
            return None
        return stack[-1][1]

    def push_node(level, node):
        parent = current_parent(level)
        if parent is None:
            roots.append(node)
        else:
            parent["children"].append(node)
        stack.append((level, node))

    def append_body(text):
        if not stack:
            return
        node = stack[-1][1]
        node["body"] = (node["body"] + "\n" + text).strip() if node["body"] else text

    def append_images(imgs):
        if not stack or not imgs:
            return
        stack[-1][1]["images"].extend(imgs)

    def append_table(grid):
        if not stack:
            return
        stack[-1][1]["tables"].append(grid)

    for kind, obj in walk_body(doc, rels):
        if kind == "t":
            grid = table_to_grid(obj)
            if not any(any(c.strip() for c in row) for row in grid):
                continue
            first_cell = (grid[0][0] if grid and grid[0] else "").strip()
            if not in_body and first_cell in ("编制部门", "编写部门"):
                cover_table = grid
                continue
            if not in_body and first_cell == "修改日期":
                revision_table = grid
                continue
            append_table(grid)
            continue

        text = para_text(obj)
        style = obj.style.name if obj.style else ""
        imgs = para_images(obj, rels)

        if text == "文件修订记录":
            continue
        if text == "目录":
            in_body = True
            stack.clear()
            continue
        if not in_body:
            if text == "安装维护手册":
                continue
            continue

        lvl = heading_level(style, text)
        if lvl > 0 and text:
            node = new_node(title=text)
            push_node(lvl, node)
            if imgs:
                append_images(imgs)
            continue

        if imgs:
            append_images(imgs)
        if text:
            append_body(text)

    cover = new_node(
        title="安装维护手册",
        ref_type="cover",
        tables=[cover_table] if cover_table else [[
            ["编制部门", "产品开发部", "文件版本", "A0"],
            ["编制人", "", "日    期", ""],
            ["审核人", "", "日    期", ""],
            ["批准人", "", "日    期", ""],
            ["生效日期", "", "", ""],
        ]],
    )
    revision = new_node(
        title="文件修订记录",
        ref_type="revision",
        tables=[revision_table] if revision_table else [[
            ["修改日期", "版本号", "修订说明", "修订人", "批准人"],
            ["", "", "首次发布", "", ""],
        ]],
    )
    return cover, revision, roots


def parse_md5_attach(path):
    doc = Document(str(path))
    rels = rel_map(doc)
    node = new_node(title="安装维护手册附件：MD5值", ref_type="md5_attachment")
    for kind, obj in walk_body(doc, rels):
        if kind == "t":
            grid = table_to_grid(obj)
            if any(any(c.strip() for c in row) for row in grid):
                node["tables"].append(grid)
            continue
        text = para_text(obj)
        imgs = para_images(obj, rels)
        lvl = heading_level(obj.style.name if obj.style else "", text)
        if lvl == 1 and text:
            child = new_node(title=text)
            node["children"].append(child)
            if imgs:
                child["images"].extend(imgs)
            continue
        if node["children"] and lvl == 0:
            target = node["children"][-1]
        else:
            target = node
        if imgs:
            target["images"].extend(imgs)
        if text:
            target["body"] = (target["body"] + "\n" + text).strip() if target["body"] else text
    return node


def parse_md5_review(path):
    doc = Document(str(path))
    rels = rel_map(doc)
    node = new_node(title="附件一", ref_type="md5_review")
    for kind, obj in walk_body(doc, rels):
        if kind == "t":
            grid = table_to_grid(obj)
            if any(any(c.strip() for c in row) for row in grid):
                node["tables"].append(grid)
            continue
        text = para_text(obj)
        if text:
            node["body"] = (node["body"] + "\n" + text).strip() if node["body"] else text
    return node


def main():
    cover, revision, roots = parse_main(IMM_DOCX)
    md5_att = parse_md5_attach(MD5_ATTACH)
    md5_rev = parse_md5_review(MD5_REVIEW)
    content = {
        "md5_value": "af6524a73e3912f39e43fec45493392a",
        "package_name": "InferCare_RECIST-2.0.0.0.zip",
        "sections": [cover, revision] + roots + [md5_att, md5_rev],
    }
    OUT.parent.mkdir(parents=True, exist_ok=True)
    with open(OUT, "w", encoding="utf-8") as f:
        json.dump(content, f, ensure_ascii=False, indent=1)
    size_mb = OUT.stat().st_size / 1024 / 1024
    print(f"Wrote {OUT} ({size_mb:.2f} MB)")
    print(f"sections: {len(content['sections'])}, body chapters: {len(roots)}")


if __name__ == "__main__":
    main()
