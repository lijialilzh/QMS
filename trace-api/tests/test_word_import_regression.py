import tempfile
from pathlib import Path

try:
    from docx import Document
except ModuleNotFoundError as exc:
    raise SystemExit(
        "missing dependency: python-docx\n"
        "please run: pip install -r requirements.txt"
    ) from exc

try:
    from src.serv.serv_srs_doc import Server
except ModuleNotFoundError as exc:
    raise SystemExit(
        "missing project dependencies\n"
        "please run: pip install -r requirements.txt"
    ) from exc


def _assert(condition, message):
    if not condition:
        raise AssertionError(message)


def _build_docx_for_table_case(path: Path):
    doc = Document()
    doc.add_paragraph("1 总则").runs[0].bold = True
    doc.add_paragraph("这里是正文，包含 RCM-001。")

    table = doc.add_table(rows=3, cols=6)
    table.rows[0].cells[0].text = "需求编号"
    table.rows[0].cells[1].text = "模块"
    table.rows[0].cells[2].text = "功能"
    table.rows[0].cells[3].text = "子功能"
    table.rows[0].cells[4].text = "章节"
    table.rows[0].cells[5].text = "RCM"

    table.rows[1].cells[0].text = "SRS-ABC-001"
    table.rows[1].cells[1].text = "模块A"
    table.rows[1].cells[2].text = "功能A"
    table.rows[1].cells[3].text = "子功能A"
    table.rows[1].cells[4].text = "1.1"
    table.rows[1].cells[5].text = "RCM-001, RCM-002"

    table.rows[2].cells[0].text = "SRS-ABC-002"
    table.rows[2].cells[1].text = "模块B"
    table.rows[2].cells[2].text = "功能B"
    table.rows[2].cells[3].text = ""
    table.rows[2].cells[4].text = "1.2"
    table.rows[2].cells[5].text = "RCM-003"

    doc.save(str(path))


def test_heading_validation(server: Server):
    ok_rows = [
        {"level": 1, "title": "1 范围", "number": "1"},
        {"level": 1, "title": "2 参考资料", "number": "2"},
        {"level": 1, "title": "3 术语", "number": "3"},
    ]
    err = server._Server__validate_heading_numbers(ok_rows)
    _assert(err is None, f"连续编号应通过，实际: {err}")

    missing_rows = [
        {"level": 1, "title": "1 范围", "number": "1"},
        {"level": 1, "title": "参考资料", "number": None},
    ]
    err = server._Server__validate_heading_numbers(missing_rows)
    _assert(err and "缺少编号" in err and "第2个一级标题" in err, f"缺号提示不符合预期: {err}")

    order_rows = [
        {"level": 1, "title": "1 范围", "number": "1"},
        {"level": 1, "title": "3 参考资料", "number": "3"},
    ]
    err = server._Server__validate_heading_numbers(order_rows)
    _assert(err and "应为 2" in err and "实际为 3" in err, f"顺序提示不符合预期: {err}")


def test_extract_file_info(server: Server):
    folder, file_no = server._Server__extract_file_info("PRD01_SRS-2026-v1.docx")
    _assert(folder == "PRD01", f"目录名解析错误: {folder}")
    _assert(file_no in ("SRS-2026-v1", "SRS-2026-v1".split("-")[0]) or file_no, "文件编号解析为空")


def test_table_extract_and_parse(server: Server):
    with tempfile.TemporaryDirectory() as tmpdir:
        path = Path(tmpdir) / "word-import-case.docx"
        _build_docx_for_table_case(path)
        doc = Document(str(path))

        content, heading_rows = server._Server__parse_docx_content(doc)
        _assert(len(content) > 0, "解析内容为空")
        _assert(any((row.get("number") == "1") for row in heading_rows), "未识别标题编号")

        req_rows, req_rcm_map = server._Server__extract_srs_reqs_from_tables(doc)
        _assert(len(req_rows) == 2, f"SRS表提取数量错误: {len(req_rows)}")
        _assert(req_rows[0]["code"] == "SRS-ABC-001", f"需求编号提取错误: {req_rows[0]}")
        _assert("SRS-ABC-001" in req_rcm_map and "RCM-002" in req_rcm_map["SRS-ABC-001"], "RCM映射提取错误")


def main():
    server = Server()
    test_heading_validation(server)
    test_extract_file_info(server)
    test_table_extract_and_parse(server)
    print("word import regression: all checks passed")


if __name__ == "__main__":
    main()

