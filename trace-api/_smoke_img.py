import os, io, zipfile
from docx import Document
from docx.shared import Inches
# same path resolution as serv (src/serv/.. /.. /src-res/assets)
img_path = os.path.join('src','serv','..','..','src-res','assets','qc_pass.png')
print('exists', os.path.exists(img_path), 'size', os.path.getsize(img_path) if os.path.exists(img_path) else 0)
doc = Document()
doc.add_paragraph().add_run().add_picture(img_path, width=Inches(1.2))
out = io.BytesIO(); doc.save(out); out.seek(0)
z = zipfile.ZipFile(out)
print('media', [n for n in z.namelist() if n.startswith('word/media/')])
