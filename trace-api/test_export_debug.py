import asyncio
import io
import sys
import traceback

sys.path.insert(0, "/work/src")
sys.path.insert(0, "/work")


async def main():
    try:
        from sqlalchemy import create_engine
        from src import env
        from src.utils import sql_ctx
        from src.serv.serv_srs_doc import Server

        engine = create_engine(env.DB_URL)
        sql_ctx.init(engine)
        with sql_ctx.db():
            serv = Server()
            buf = io.BytesIO()
            await serv.export_srs_doc(buf, 173)
            print("=== direct export ===")
            _print_req_table(buf)

            resp = await serv.get_srs_doc(173, with_tree=True)
            from copy import deepcopy
            from src.obj.tobj_srs_doc import SrsDocForm

            snapshot = SrsDocForm(id=173, content=deepcopy(resp.data.content or []))

            def corrupt_table_modules(nodes):
                for node in nodes or []:
                    table = getattr(node, "table", None)
                    if table and getattr(table, "headers", None):
                        headers = table.headers or []
                        module_code = next(
                            (getattr(h, "code", "") for h in headers if "模块" in str(getattr(h, "name", "") or "")),
                            "",
                        )
                        if module_code:
                            for row in getattr(table, "rows", None) or []:
                                if isinstance(row, dict):
                                    row[module_code] = ""
                    corrupt_table_modules(getattr(node, "children", None) or [])

            corrupt_table_modules(snapshot.content)
            buf2 = io.BytesIO()
            await serv.export_srs_doc(buf2, 173, snapshot=snapshot)
            print("=== snapshot export (empty module rows) ===")
            _print_req_table(buf2)
    except Exception:
        traceback.print_exc()


def _print_req_table(buf):
    buf.seek(0)
    from docx import Document

    doc = Document(buf)
    print("export size", buf.tell(), "tables count", len(doc.tables))
    for ti, table in enumerate(doc.tables):
        if not table.rows:
            continue
        hdr = [c.text.strip() for c in table.rows[0].cells]
        if "模块" not in "".join(hdr):
            continue
        print("Table", ti, "cols", len(table.columns), "header", hdr)
        for ri in range(1, min(8, len(table.rows))):
            cells = [c.text.strip().replace("\n", " ") for c in table.rows[ri].cells]
            print(" row", ri, cells)
        for ri in range(1, len(table.rows)):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            if any("搜索与筛选" in c for c in cells):
                for r2 in range(ri, min(ri + 5, len(table.rows))):
                    print(" row303", r2, [c.text.strip() for c in table.rows[r2].cells])
                break
        break


asyncio.run(main())
